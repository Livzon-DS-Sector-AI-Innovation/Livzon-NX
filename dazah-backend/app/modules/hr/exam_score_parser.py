"""解析笔试成绩文件（.docx / .xlsx），提取 姓名+成绩 对."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Any

import openpyxl  # type: ignore[import-untyped]
from docx import Document
from fastapi import UploadFile

from app.core.config import get_settings
from app.core.upload_security import read_upload_secure


@dataclass
class ExamScore:
    name: str
    score: str  # 保留原始字符串，可能是数字也可能是等级


# ── 公共入口 ──────────────────────────────────────────────


async def parse_exam_scores(file: UploadFile) -> list[ExamScore]:
    """根据文件扩展名分发到对应解析器."""
    filename, content = await read_upload_secure(
        file,
        max_bytes=get_settings().MAX_UPLOAD_SIZE_MB * 1024 * 1024,
        allowed_extensions={".docx", ".xlsx"},
        what="笔试成绩文件",
    )
    filename = filename.lower()

    if filename.endswith(".xlsx"):
        return _parse_xlsx(content)
    if filename.endswith(".docx"):
        return _parse_docx(content)

    raise ValueError(f"不支持的文件格式：{filename}，仅支持 .docx 和 .xlsx")


# ── .xlsx 解析 ────────────────────────────────────────────

_NAME_HINTS = ["姓名", "名字", "人员", "学员", "员工", "参加人员"]
_SCORE_HINTS = ["成绩", "分数", "得分", "笔试成绩", "考核成绩", "考试成绩", "评分"]


def _find_column(header: tuple[Any, ...], hints: list[str]) -> int | None:
    """在表头行中按关键词匹配列索引."""
    for idx, cell in enumerate(header):
        if cell is None:
            continue
        text = str(cell).strip()
        for hint in hints:
            if hint in text:
                return idx
    return None


def _parse_xlsx(content: bytes) -> list[ExamScore]:
    """从 Excel 文件中提取姓名+成绩对.

    策略：
    1. 扫描所有工作表
    2. 表头关键词匹配"姓名"列和"成绩"列
    3. 找不到表头时，尝试相邻两列（左列中文=姓名，右列数字=成绩）
    """
    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    results: list[ExamScore] = []

    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        if len(rows) < 2:
            continue

        header = rows[0]
        name_col = _find_column(header, _NAME_HINTS)
        score_col = _find_column(header, _SCORE_HINTS)

        if name_col is not None and score_col is not None:
            # 精确匹配到表头列
            for row in rows[1:]:
                name = _safe_str(row, name_col)
                score = _safe_str(row, score_col)
                if name and score:
                    results.append(ExamScore(name=name, score=score))
        else:
            # fallback：扫描相邻两列
            found = False
            for c in range(len(header) - 1):
                # 检查数据行：左列像姓名（中文），右列像成绩（数字）
                candidates: list[ExamScore] = []
                for row in rows[1:]:
                    left = _safe_str(row, c)
                    right = _safe_str(row, c + 1)
                    if left and right and _is_chinese_name(left) and _is_score(right):
                        candidates.append(ExamScore(name=left, score=right))
                if len(candidates) >= 2:
                    results.extend(candidates)
                    found = True
                    break
            if not found and name_col is not None:
                # 有姓名列但没成绩列，尝试下一列作为成绩
                for row in rows[1:]:
                    name = _safe_str(row, name_col)
                    score = (
                        _safe_str(row, name_col + 1)
                        if name_col + 1 < len(row)
                        else None
                    )
                    if name and score and _is_score(str(score)):
                        results.append(ExamScore(name=name, score=str(score)))

    wb.close()
    return results


# ── .docx 解析 ────────────────────────────────────────────

# 匹配 "姓名+数字" 模式：中文姓名后跟可选空格/分隔符再跟数字成绩
# 例如：周亚学96、秦亚瑞 93、冯筱筱94
_DOCX_PATTERN = re.compile(
    r"([\u4e00-\u9fa5]{2,4})"  # 2-4 个中文字符（姓名）
    r"[\s:：]*"  # 可选分隔符
    r"(\d{1,3}(?:\.\d+)?)"  # 数字成绩（整数或小数）
)


def _parse_docx(content: bytes) -> list[ExamScore]:
    """从 Word 文档中提取姓名+成绩对.

    策略：
    1. 提取全部段落和表格文本
    2. 用正则匹配"姓名+数字"模式
    3. 也尝试解析表格中的姓名/成绩列
    """
    doc = Document(io.BytesIO(content))
    results: list[ExamScore] = []
    seen: set[str] = set()

    # 1. 从段落文本中提取
    full_text_parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text_parts.append(text)

    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            full_text_parts.append(" ".join(cells))

    full_text = "\n".join(full_text_parts)

    # 先尝试整体文本的正则匹配
    for match in _DOCX_PATTERN.finditer(full_text):
        name = match.group(1)
        score = match.group(2)
        key = f"{name}_{score}"
        if key not in seen:
            seen.add(key)
            results.append(ExamScore(name=name, score=score))

    # 2. 尝试从表格中结构化提取
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        header_tuple = tuple(header_cells)
        name_col = _find_column(header_tuple, _NAME_HINTS)
        score_col = _find_column(header_tuple, _SCORE_HINTS)

        if name_col is not None and score_col is not None:
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if name_col < len(cells) and score_col < len(cells):
                    name = cells[name_col]
                    score = cells[score_col]
                    key = f"{name}_{score}"
                    if name and score and key not in seen:
                        seen.add(key)
                        results.append(ExamScore(name=name, score=score))

    return results


# ── 工具函数 ──────────────────────────────────────────────


def _safe_str(row: tuple[Any, ...], idx: int) -> str | None:
    """安全获取行中某列的字符串值."""
    if idx >= len(row):
        return None
    val = row[idx]
    if val is None:
        return None
    text = str(val).strip()
    return text if text else None


def _is_chinese_name(text: str) -> bool:
    """判断是否像中文姓名（2-4 个汉字）."""
    if not (2 <= len(text) <= 4):
        return False
    return all("\u4e00" <= ch <= "\u9fa5" for ch in text)


def _is_score(text: str) -> bool:
    """判断是否像成绩（数字，0-100）."""
    try:
        val = float(text)
        return 0 <= val <= 100
    except ValueError:
        return False


def format_score_summary(scores: list[ExamScore]) -> str:
    """将成绩列表格式化为"姓名成绩、姓名成绩"格式."""
    parts = []
    for s in scores:
        # 去掉小数点后的 0（如 96.0 → 96）
        try:
            val = float(s.score)
            score_str = str(int(val)) if val == int(val) else str(val)
        except ValueError:
            score_str = s.score
        parts.append(f"{s.name}{score_str}")
    return "、".join(parts)


def extract_personal_score(
    name: str,
    score_summary: str | None,
    assessment_method: str | None,
) -> str | None:
    """从全场成绩汇总中提取个人考核结果（员工培训清单"考核结果"列）.

    - 口试/实操（或 score_summary 为 '/'）→ "合格"
    - 笔试 → 按姓名从"张三96、李四95"格式中提取个人分数；提取不到返回 None
    """
    summary = (score_summary or "").strip()
    method = (assessment_method or "").strip()
    if method in ("口试", "实操") or summary in ("/", "合格"):
        return "合格"
    if not summary:
        return None
    pattern = re.compile(rf"{re.escape(name)}[\s:：]*(\d+(?:\.\d+)?)")
    m = pattern.search(summary)
    if m:
        return m.group(1)
    return None
