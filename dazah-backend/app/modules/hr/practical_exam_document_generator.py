"""Generate practical exam result document (APP13-SMP-HR-002-14) from docx template.

一人一份：每个参训人员生成一份实操考核结果表（可多页），导出时打包成 zip。
另支持"导入实操试题"：解析 APP13 格式的 docx，提取实操考核情况描述与培训日期。
模板保真填充，仅改填写区。
"""

import copy
import logging
import re
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document

from app.modules.hr.date_format import fmt_date_str
from app.modules.hr.schemas import PracticalExamExportRequest, PracticalExamPersonItem
from app.modules.hr.template_filler import fill_whole_cell, header_fmt_source

logger = logging.getLogger(__name__)

TEMPLATE_NAME = "APP13-SMP-HR-002-14实操培训考核结果表.docx"


def _find_template() -> Path:
    candidates = [
        Path("员工培训教育管理规程") / TEMPLATE_NAME,
        Path(__file__).resolve().parent.parent.parent
        / "员工培训教育管理规程"
        / TEMPLATE_NAME,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {TEMPLATE_NAME}")


def _fill_table(
    table: Any, person: Any, data: PracticalExamExportRequest, fmt_src: Any
) -> None:
    # R1: 部门 | 姓名
    fill_whole_cell(table.rows[1].cells[1], person.department or "", fmt_src)
    fill_whole_cell(table.rows[1].cells[5], person.name or "", fmt_src)
    # R2: 培训内容 | 培训日期
    fill_whole_cell(table.rows[2].cells[1], data.training_content or "", fmt_src)
    fill_whole_cell(table.rows[2].cells[5], fmt_date_str(data.training_date), fmt_src)
    # R3: 实操考核情况描述（大块空白）
    fill_whole_cell(table.rows[3].cells[0], person.description or "", fmt_src)
    # R4: 评估人/日期（落款标签后空白格）
    fill_whole_cell(table.rows[4].cells[5], data.assessor or "", fmt_src)


def _fill_person_only(
    table: Any, person: Any, data: PracticalExamExportRequest, fmt_src: Any
) -> None:
    "以导入的实操试题文件为基底时，只填 部门/姓名/培训日期，其"
    "余（培训内容、试题结构、方框）原样保留."
    fill_whole_cell(table.rows[1].cells[1], person.department or "", fmt_src)
    fill_whole_cell(table.rows[1].cells[5], person.name or "", fmt_src)
    fill_whole_cell(table.rows[2].cells[5], fmt_date_str(data.training_date), fmt_src)


def _imported_template_path() -> Path:
    """最近一次导入的实操试题 docx 的保存路径（导出时作为基底，保留试题结构/方框）。"""
    return Path(__file__).resolve().parent / "_practical_imported_template.docx"


def _safe_filename(s: str | None, maxlen: int = 40) -> str:
    """去掉文件名非法字符，截断到 maxlen."""
    s = re.sub(r'[\\/:*?"<>|\r\n\t]+', "", s or "").strip()
    return s[:maxlen] or "未命名"


def _make_person_doc(person: Any, data: PracticalExamExportRequest) -> BytesIO:
    """生成单个人员的 APP13 文档（独立 docx，可多页）.

    若存在最近导入的实操试题文件，则以该文件为基底、只填部门/姓名/培训日期，
    完整保留试题的段落结构/方框；否则用空白模板完整填充。
    """
    imported = _imported_template_path()
    if imported.exists():
        doc = Document(str(imported))
        if doc.tables:
            fmt_src = header_fmt_source(doc.tables[0])
            _fill_person_only(doc.tables[0], person, data, fmt_src)
    else:
        doc = Document(str(_find_template()))
        if doc.tables:
            fmt_src = header_fmt_source(doc.tables[0])
            _fill_table(doc.tables[0], person, data, fmt_src)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_practical_exam_zip(data: PracticalExamExportRequest) -> BytesIO:
    """为每个参训人员生成一份实操考核结果表并打包成 zip."""
    persons = data.persons or []
    if not persons:
        persons = [PracticalExamPersonItem(name="")]
    zip_buf = BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, person in enumerate(persons, start=1):
            doc_buf = _make_person_doc(person, data)
            name = person.name or f"人员{i}"
            fname = (
                f"{_safe_filename(data.training_date)}"
                f"-{_safe_filename(data.training_content)}"
                f"-{_safe_filename(name)}实操.docx"
            )
            zf.writestr(fname, doc_buf.getvalue())
    zip_buf.seek(0)
    return zip_buf


def parse_practical_exam_questions(content: bytes) -> dict[str, Any]:
    """解析导入的 APP13 格式 docx，提取实操考核情况描述与培训日期.

    模板结构：R1 部门/姓名（导入时留空），R2 培训内容/培训日期，R3
    实操考核情况描述（整行大格）。
    """
    doc = Document(BytesIO(content))
    if not doc.tables:
        raise ValueError("导入文件中未找到表格")
    t = doc.tables[0]
    # R3 实操考核情况描述（整行合并大格，取首段文本）
    description = ""
    if len(t.rows) > 3:
        description = t.rows[3].cells[0].text.strip()
    # R2 培训日期
    training_date = ""
    if len(t.rows) > 2:
        training_date = t.rows[2].cells[5].text.strip()
    return {"description": description, "training_date": training_date}


def generate_practical_exam_result(data: PracticalExamExportRequest) -> BytesIO:
    """按 APP13 模板保真生成实操培训考核结果表 Word（每人一页，单文档）."""
    doc = Document(str(_find_template()))
    if not doc.tables:
        raise ValueError("模板文件中未找到表格")

    fmt_src = header_fmt_source(doc.tables[0])
    persons = data.persons or []
    base_tbl_xml = copy.deepcopy(doc.tables[0]._tbl)

    if not persons:
        _fill_table(doc.tables[0], PracticalExamPersonItem(name=""), data, fmt_src)
    for idx, person in enumerate(persons):
        if idx == 0:
            _fill_table(doc.tables[0], person, data, fmt_src)
        else:
            doc.add_page_break()  # type: ignore[no-untyped-call]
            doc.element.body.append(copy.deepcopy(base_tbl_xml))
            _fill_table(doc.tables[-1], person, data, fmt_src)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
