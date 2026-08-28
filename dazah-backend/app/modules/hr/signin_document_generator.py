"""Generate training sign-in sheet documents from APP3 docx template.

模板保真填充：仅填充填写区，保留模板 run 格式（字体/字号/加粗/sym 复选框）。

模板结构（实测）：
- 第 1 页：children[0:5]，28 行表格（r5 列表头，r6-r26 数据行 21 行，r27 备注）
- 第 2 页：children[5:10]，27 行表格（r5 列表头，r6-r25 数据行 20 行，r26 备注）
- 填写区：r0 日期/部门，r1 方式，r2 人数，r3 标签行，r4 空行（时间/题目/授课人）
"""

import logging
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from app.modules.hr.date_format import fmt_date_str
from app.modules.hr.schemas import TrainingSignInSheetInput
from app.modules.hr.template_filler import (
    fill_after_label,
    fill_after_phrase,
    fill_whole_cell,
    header_fmt_source,
    set_sym_group,
)

logger = logging.getLogger(__name__)

PER_PAGE = 42  # 第 1 页 42 人（21 行×2 栏）
PER_PAGE_2 = 40  # 第 2 页 40 人（20 行×2 栏）
HALF_PAGE_1 = 21
HALF_PAGE_2 = 20


def _find_template() -> Path:
    """Locate the docx template, trying several path candidates."""
    candidates = [
        Path("员工培训教育管理规程/APP3-SMP-HR-002-14培训签到表.docx"),
        Path("员工培训教育管理规程/APP3培训签到表.docx"),
        Path(__file__).resolve().parent.parent.parent.parent
        / "员工培训教育管理规程"
        / "APP3-SMP-HR-002-14培训签到表.docx",
        Path(__file__).resolve().parent.parent.parent.parent
        / "员工培训教育管理规程"
        / "APP3培训签到表.docx",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("模板文件未找到：APP3-SMP-HR-002-14培训签到表.docx")


def _fill_header(
    table: Any,
    data: TrainingSignInSheetInput,
    page: int,
    is_first_page: bool,
) -> None:
    """填充表头信息（r0-r4）。

    is_first_page=True: 填写 r0-r4
    is_first_page=False: 清空 r0-r4（续页不重复第一页内容）
    """
    fmt_src = header_fmt_source(table)

    if is_first_page:
        # r0: 培训日期 | 受训部门
        fill_after_label(table.rows[0].cells[0], f" {fmt_date_str(data.training_date)}")
        fill_after_label(table.rows[0].cells[5], f" {data.department or ''}")

        # r1: 培训方式（勾选 sym 组）
        method = data.training_method or ""
        other_txt = ""
        if method.startswith("其他"):
            other_txt = method.replace("其他", "", 1).lstrip("：: ")
        set_sym_group(
            table.rows[1].cells[0],
            method,
            ["面授", "实操", "函授", "远程教育", "其他"],
            write_trailing={"其他": other_txt},
        )

        # r2: 应受训人数（总人数）；实际受训人数合计留空待签名
        total = len(data.employee_names)
        fill_after_phrase(table.rows[2].cells[0], "应受训人数", str(total))

        # r3 是标签行（培训时间/培训题目或内容概要/授课人），r4 是填写行
        time_text = ""
        if data.training_time_start and data.training_time_end:
            time_text = f"{data.training_time_start} ~ {data.training_time_end}"
        fill_whole_cell(table.rows[4].cells[0], time_text, fmt_src)

        display_topic = data.topic
        if data.training_subject:
            display_topic = f"{data.training_subject} — {data.topic}"
        fill_whole_cell(table.rows[4].cells[2], display_topic, fmt_src)
        fill_whole_cell(table.rows[4].cells[8], data.instructor or "", fmt_src)
    else:
        # 续页：清空填写区（保留标签和格式）
        fill_after_label(table.rows[0].cells[0], "")
        fill_after_label(table.rows[0].cells[5], "")
        set_sym_group(
            table.rows[1].cells[0],
            None,
            ["面授", "实操", "函授", "远程教育", "其他"],
        )
        fill_after_phrase(
            table.rows[2].cells[0],
            "应受训人数",
            str(len(data.employee_names)),
        )
        fill_whole_cell(table.rows[4].cells[0], "", fmt_src)
        fill_whole_cell(table.rows[4].cells[2], "", fmt_src)
        fill_whole_cell(table.rows[4].cells[8], "", fmt_src)


def _fill_data_rows(
    table: Any,
    data: TrainingSignInSheetInput,
    page: int,
    half_page: int,
) -> None:
    """填充数据行（保留序号，空行只清空姓名/部门/签到）。"""
    fmt_src = header_fmt_source(table)
    n_rows = len(table.rows)

    # 找到列表头行（包含"序号"的行）
    header_row = None
    for i, row in enumerate(table.rows):
        if "序号" in row.cells[0].text:
            header_row = i
            break

    if header_row is None:
        return

    data_start_row = header_row + 1
    remarks_row = n_rows - 1

    # 本页人员
    start = page * PER_PAGE
    page_names = data.employee_names[start : start + PER_PAGE]

    # 填充数据行
    for i, name in enumerate(page_names):
        if i < half_page:  # 左栏
            row_idx = data_start_row + i
            if row_idx < remarks_row:
                row = table.rows[row_idx]
                # 序号保留模板预置值，只填姓名/部门
                fill_whole_cell(row.cells[1], name, fmt_src)
                fill_whole_cell(row.cells[3], data.department or "", fmt_src)
                fill_whole_cell(row.cells[4], "", fmt_src)  # 签到
        else:  # 右栏
            row_idx = data_start_row + (i - half_page)
            if row_idx < remarks_row:
                row = table.rows[row_idx]
                # 序号保留模板预置值，只填姓名/部门
                fill_whole_cell(row.cells[6], name, fmt_src)
                fill_whole_cell(row.cells[7], data.department or "", fmt_src)
                fill_whole_cell(row.cells[9], "", fmt_src)  # 签到

    # 空行：保留序号，只清空姓名/部门/签到
    left_filled = min(len(page_names), half_page)
    for i in range(left_filled, half_page):
        row_idx = data_start_row + i
        if row_idx < remarks_row:
            row = table.rows[row_idx]
            fill_whole_cell(row.cells[1], "", fmt_src)
            fill_whole_cell(row.cells[3], "", fmt_src)
            fill_whole_cell(row.cells[4], "", fmt_src)

    right_filled = max(0, len(page_names) - half_page)
    for i in range(right_filled, half_page):
        row_idx = data_start_row + i
        if row_idx < remarks_row:
            row = table.rows[row_idx]
            fill_whole_cell(row.cells[6], "", fmt_src)
            fill_whole_cell(row.cells[7], "", fmt_src)
            fill_whole_cell(row.cells[9], "", fmt_src)

    # 备注行（仅第一页填写）
    if data.remarks and page == 0:
        fill_after_phrase(table.rows[remarks_row].cells[0], "备注", data.remarks)


def generate_training_sign_in_sheet(data: TrainingSignInSheetInput) -> BytesIO:
    """Generate the complete sign-in sheet as ONE Word document.

    签到表始终是一份文档：
    - 第 1 页：42 人（21 行×2 栏），使用模板第一页块（children[0:5]）
    - 第 2 页：40 人（20 行×2 栏），使用模板第二页块（children[5:10]，空白续页）
    - 第 3 页及以后：深拷贝第二页块，动态调整数据行数量
    """
    template_path = _find_template()
    doc = Document(str(template_path))

    total = len(data.employee_names)
    pages = max(1, (total + PER_PAGE - 1) // PER_PAGE)

    body = doc.element.body
    children = [el for el in body if el.tag != qn("w:sectPr")]

    # 第 1 页块：children[0:5]（28 行表格）
    page1_block = children[0:5]
    # 第 2 页块：children[5:10]（27 行表格，空白续页）
    page2_block = children[5:10]

    # 移除模板自带的所有内容（保留 sectPr）
    for el in children:
        body.remove(el)

    # 第 1 页：使用 page1_block
    for el in page1_block:
        body.append(deepcopy(el))
    _set_para_text(body[-5], f"P1/{pages}")
    _fill_header(doc.tables[0], data, 0, is_first_page=True)
    _fill_data_rows(doc.tables[0], data, 0, HALF_PAGE_1)

    # 第 2 页及以后：深拷贝 page2_block
    for page in range(1, pages):
        for el in page2_block:
            body.append(deepcopy(el))
        _set_para_text(body[-5], f"P{page + 1}/{pages}")
        _fill_header(doc.tables[-1], data, page, is_first_page=False)
        _fill_data_rows(doc.tables[-1], data, page, HALF_PAGE_2)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _set_para_text(para_el: Any, text: str) -> None:
    """将段落文本整体替换（保留首个 run 的格式）。"""
    runs = para_el.findall(qn("w:r"))
    if not runs:
        return
    first = runs[0]
    for r in runs[1:]:
        para_el.remove(r)
    t_el = first.find(qn("w:t"))
    if t_el is None:
        return
    t_el.text = text
    t_el.set(qn("xml:space"), "preserve")
