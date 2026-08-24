"""岗位培训清单 API"""

import logging
import re
from io import BytesIO
from typing import Any
from urllib.parse import quote
from uuid import UUID

from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, File, Query, UploadFile
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.core.database import get_db
from app.core.deps import CurrentUser
from app.core.exceptions import AppException, NotFoundException
from app.core.response import paginated_response, success_response
from app.core.upload_security import read_upload_secure
from app.modules.hr.position_training_service import PositionTrainingListService
from app.modules.hr.schemas import (
    PositionTrainingListCreate,
    PositionTrainingListItemBatchUpdate,
    PositionTrainingListResponse,
    PositionTrainingListUpdate,
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/position-training-lists", tags=["人事-岗位培训清单"])


def _require_user(current_user: CurrentUser) -> Any:
    if current_user is None:
        raise AppException(status_code=401, message="请先登录")


@router.get("", summary="岗位培训清单列表")
async def list_position_training_lists(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    department: str | None = Query(None, description="部门筛选"),
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    from app.modules.hr.api import _assert_dept_in_scope

    alias_set = await _assert_dept_in_scope(db, current_user, department)
    service = PositionTrainingListService(db)
    lists, total = await service.list_lists(
        page=page, page_size=page_size, department=department, dept_alias_set=alias_set
    )
    return paginated_response(
        data=[
            PositionTrainingListResponse.model_validate(item).model_dump(mode="json")
            for item in lists
        ],
        page=page,
        page_size=page_size,
        total=total,
    )


@router.get("/departments", summary="已导入的部门列表")
async def list_imported_departments(
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    """返回岗位培训清单中已存在的部门名称列表（去重），用于前端部门按钮合并。"""
    from app.modules.hr.api import _resolve_visible_scope

    alias_set = await _resolve_visible_scope(db, current_user)
    service = PositionTrainingListService(db)
    depts = await service.repo.list_distinct_departments(dept_alias_set=alias_set)
    return success_response(data=depts)


@router.get("/departments/{department}/positions", summary="获取指定部门的岗位列表")
async def list_positions_by_department(
    department: str,
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    """返回指定部门的所有岗位名称（去重），用于新员工培训岗位选择。"""
    from app.modules.hr.api import _assert_dept_in_scope

    await _assert_dept_in_scope(db, current_user, department)
    service = PositionTrainingListService(db)
    positions = await service.list_positions_by_department(department)
    return success_response(data=positions)


@router.post("", summary="创建岗位培训清单")
async def create_position_training_list(
    data: PositionTrainingListCreate,
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    _require_user(current_user)
    service = PositionTrainingListService(db)
    payload = data.model_dump(exclude_unset=True)
    items_data = payload.pop("items", None)
    list_obj = await service.create(payload, items_data)
    return success_response(
        data=PositionTrainingListResponse.model_validate(list_obj).model_dump(
            mode="json"
        ),
        message="创建成功",
    )


@router.get("/{list_id}", summary="岗位培训清单详情")
async def get_position_training_list(
    list_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    _require_user(current_user)
    service = PositionTrainingListService(db)
    list_obj = await service.get_by_id(list_id)
    if not list_obj:
        raise NotFoundException(resource="岗位培训清单", resource_id=str(list_id))
    return success_response(
        data=PositionTrainingListResponse.model_validate(list_obj).model_dump(
            mode="json"
        )
    )


@router.put("/{list_id}", summary="更新岗位培训清单")
async def update_position_training_list(
    list_id: UUID,
    data: PositionTrainingListUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    _require_user(current_user)
    service = PositionTrainingListService(db)
    list_obj = await service.update(list_id, data.model_dump(exclude_unset=True))
    if not list_obj:
        raise NotFoundException(resource="岗位培训清单", resource_id=str(list_id))
    return success_response(
        data=PositionTrainingListResponse.model_validate(list_obj).model_dump(
            mode="json"
        ),
        message="更新成功",
    )


@router.put("/{list_id}/items/batch", summary="批量更新清单明细")
async def batch_update_items(
    list_id: UUID,
    data: PositionTrainingListItemBatchUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    _require_user(current_user)
    service = PositionTrainingListService(db)
    list_obj = await service.get_by_id(list_id)
    if not list_obj:
        raise NotFoundException(resource="岗位培训清单", resource_id=str(list_id))

    items_data = [item.model_dump(exclude_unset=True) for item in data.items]
    await service.batch_update_items(list_id, items_data)

    # 重新获取完整数据
    updated = await service.get_by_id(list_id)
    return success_response(
        data=PositionTrainingListResponse.model_validate(updated).model_dump(
            mode="json"
        ),
        message="更新成功",
    )


@router.get("/{list_id}/export", summary="导出岗位培训清单")
async def export_position_training_list(
    list_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    """导出岗位培训清单 Word 文档（APP9模板）"""
    _require_user(current_user)
    from app.modules.hr.position_training_document_generator import (
        generate_position_training_list,
    )

    service = PositionTrainingListService(db)
    list_obj = await service.get_by_id(list_id)
    if not list_obj:
        raise NotFoundException(resource="岗位培训清单", resource_id=str(list_id))

    buffer: BytesIO = generate_position_training_list(list_obj)
    dept = list_obj.department or "未知部门"
    pos = list_obj.position or "未知岗位"
    raw_name = f"{dept}-{pos}-岗位培训清单.docx"
    filename = quote(raw_name)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": (
                f"attachment; filename=\"{filename}\"; filename*=UTF-8''{filename}"
            )
        },
    )


@router.delete("/{list_id}", summary="删除岗位培训清单")
async def delete_position_training_list(
    list_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    _require_user(current_user)
    service = PositionTrainingListService(db)
    deleted = await service.delete(list_id)
    if not deleted:
        raise NotFoundException(resource="岗位培训清单", resource_id=str(list_id))
    return success_response(message="删除成功")


@router.delete("/by-dept/clear", summary="清除部门岗位培训清单")
async def clear_department_lists(
    department: str = Query(..., description="部门名称"),
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    """一键清除指定部门下的所有岗位培训清单及明细（软删除）。"""
    _require_user(current_user)
    service = PositionTrainingListService(db)
    lists, _ = await service.repo.list_lists(
        page=1, page_size=1000, department=department
    )
    if not lists:
        return success_response(message=f"部门「{department}」下暂无岗位培训清单")
    count = 0
    for lst in lists:
        await service.delete(lst.id)
        count += len([i for i in lst.items if not i.is_deleted])
    logger.info(
        "部门岗位培训清单已清除",
        extra={"department": department, "lists": len(lists), "items": count},
    )
    return success_response(
        data={
            "department": department,
            "deleted_lists": len(lists),
            "deleted_items": count,
        },
        message=f"已清除部门「{department}」下的 {len(lists)} 个清单、{count} 条明细",
    )


# ─── Word 导入 ───


def _parse_import_docx(content: bytes) -> dict[str, Any]:
    """解析 APP9 岗位培训清单 Word 文档。

    提取部门、岗位以及部门级/岗位级明细。

    Returns:
        {
            "department": "QA部",
            "position": "QA经理",
            "items": [
                {"level": "部门级", "textbook_name": "...", "textbook_code": "...",
                "assessment_method": "...", "remarks": "..."},
                ...
            ]
        }
    """
    doc = DocxDocument(BytesIO(content))

    department = ""
    position = ""
    items: list[dict[str, Any]] = []

    # 1. 从段落提取"部门"和"岗位"
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 段落格式: "部门：XXX          岗位：XXX"
        # 按 "岗位" 分割，分别提取部门和岗位
        if "部门" in text and "岗位" in text:
            # 找到 "岗位" 的位置做分割
            pos_idx = text.index("岗位")
            left_part = text[:pos_idx]
            right_part = text[pos_idx:]

            # 从左边提取部门
            dept_match = re.search(r"部门[：:]\s*(.*)", left_part)
            if dept_match:
                department = dept_match.group(1).strip()

            # 从右边提取岗位
            pos_match = re.search(r"岗位[：:]\s*(.*)", right_part)
            if pos_match:
                position = pos_match.group(1).strip()
        else:
            # 可能部门/岗位分行写
            dept_match = re.search(r"部门[：:]\s*(.+)", text)
            if dept_match:
                department = dept_match.group(1).strip()
            pos_match = re.search(r"岗位[：:]\s*(.+)", text)
            if pos_match:
                position = pos_match.group(1).strip()

    # 2. 从表格提取明细
    current_level = ""
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 5:
                continue
            first_cell_text = cells[0].text.strip()

            # 判断是否为级别标题行（"部门级"或"岗位级"）
            if first_cell_text in ("部门级", "岗位级"):
                current_level = first_cell_text
                continue

            # 跳过表头行
            if first_cell_text == "序号":
                continue

            # 跳过省略号行
            if first_cell_text == "……":
                continue

            # 跳过空行（所有单元格都为空）
            if not any(c.text.strip() for c in cells):
                continue

            # 数据行
            try:
                int(first_cell_text)  # 序号为数字
            except (ValueError, TypeError):
                continue

            textbook_name = cells[1].text.strip()
            textbook_code = cells[2].text.strip()
            assessment_method = cells[3].text.strip()
            remarks = cells[4].text.strip()

            items.append(
                {
                    "level": current_level or "岗位级",
                    "textbook_name": textbook_name,
                    "textbook_code": textbook_code,
                    "assessment_method": assessment_method,
                    "remarks": remarks,
                }
            )

    return {
        "department": department,
        "position": position,
        "items": items,
    }


@router.post("/import", summary="导入岗位培训清单（Word）")
async def import_position_training_list(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    """上传 Word 文件（.docx 或 .doc），解析并导入岗位培训清单。

    自动识别文件中的部门信息，如有同名清单则追加明细。
    """
    _require_user(current_user)

    try:
        safe_name, content = await read_upload_secure(
            file,
            max_bytes=get_settings().MAX_UPLOAD_SIZE_MB * 1024 * 1024,
            allowed_extensions={".docx", ".doc"},
            what="岗位培训清单",
        )
    except AppException:
        raise
    except Exception:
        logger.exception("读取上传文件失败")
        raise AppException(status_code=400, message="文件读取失败")

    # .doc 格式暂不支持，提示用户转换
    ext = safe_name.rsplit(".", 1)[-1].lower()
    if ext == "doc":
        raise AppException(
            status_code=400,
            message="暂不支持 .doc 格式，请将文件另存为 .docx 格式后上传",
        )

    # 解析文档
    try:
        parsed = _parse_import_docx(content)
    except Exception:
        logger.exception("解析 Word 文档失败")
        raise AppException(
            status_code=400, message="文档解析失败，请确认文件格式与 APP9 模板一致"
        )

    department = parsed.get("department", "")
    position = parsed.get("position", "")
    items = parsed.get("items", [])

    # 部门名归一化：将别名映射到系统标准名称（201 变体/别名等，与培训模块统一规则）
    from app.modules.hr.training_dept_resolver import resolve_training_department

    normalized = await resolve_training_department(db, department)
    if normalized != department:
        logger.info("部门名归一化: %s → %s", department, normalized)
        department = normalized

    if not department or not items:
        raise AppException(
            status_code=400, message="未能从文档中识别出部门或培训明细，请检查文件格式"
        )

    service = PositionTrainingListService(db)

    # 查找该部门下是否有同名清单（匹配部门+岗位）
    existing_lists, _ = await service.repo.list_lists(
        page=1, page_size=100, department=department
    )
    matched_list = None
    for lst in existing_lists:
        if lst.position == position:
            matched_list = lst
            break

    if matched_list:
        # 已有清单 → 追加明细
        existing_items_data = [
            {
                "level": item.level,
                "textbook_name": item.textbook_name,
                "textbook_code": item.textbook_code,
                "assessment_method": item.assessment_method,
                "remarks": item.remarks,
            }
            for item in matched_list.items
            if not item.is_deleted
        ]
        # 去重：按 (级别, 教材名称) 组合去重
        existing_keys = {
            (i["level"], i["textbook_name"])
            for i in existing_items_data
            if i["textbook_name"]
        }
        new_items = [
            i for i in items if (i["level"], i["textbook_name"]) not in existing_keys
        ]
        if new_items:
            merged = existing_items_data + new_items
            await service.batch_update_items(matched_list.id, merged)
            logger.info(
                "位置培训清单已更新",
                extra={"list_id": str(matched_list.id), "new_items": len(new_items)},
            )
            return success_response(
                data={
                    "matched": True,
                    "list_id": str(matched_list.id),
                    "department": department,
                    "position": position,
                    "imported": len(new_items),
                    "skipped": len(items) - len(new_items),
                },
                message=(
                    f"已追加 {len(new_items)} 条明细到已有清单"
                    f"（{department} - {position}），"
                    f"跳过 {len(items) - len(new_items)} 条重复"
                ),
            )
        else:
            return success_response(
                data={
                    "matched": True,
                    "list_id": str(matched_list.id),
                    "department": department,
                    "position": position,
                    "imported": 0,
                    "skipped": len(items),
                },
                message=f"所有明细已存在，无新增（{department} - {position}）",
            )

    # 新建清单
    new_list = await service.create(
        {"department": department, "position": position},
        items,
    )
    logger.info(
        "位置培训清单已创建",
        extra={
            "list_id": str(new_list.id),
            "department": department,
            "position": position,
            "items": len(items),
        },
    )
    return success_response(
        data={
            "matched": False,
            "list_id": str(new_list.id),
            "department": department,
            "position": position,
            "imported": len(items),
            "skipped": 0,
        },
        message=f"已创建新清单（{department} - {position}），共 {len(items)} 条明细",
    )
