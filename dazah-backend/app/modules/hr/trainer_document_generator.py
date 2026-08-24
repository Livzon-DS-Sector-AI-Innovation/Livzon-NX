"""培训师清单 Word 文档生成器（APP8-SMP-HR-002-14 模板）."""

import logging
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt

from app.modules.hr.date_format import fmt_date_obj

logger = logging.getLogger(__name__)

# 表头字段映射
HEADER_MAP = ["姓名", "部门", "岗位", "批准时间", "备注"]


def _find_template() -> Path:
    """查找 APP8 模板文件。"""
    candidates = [
        Path("员工培训教育管理规程/APP8-SMP-HR-002-14培训师清单.docx"),
        Path("员工培训教育管理规程/APP8培训师清单.docx"),
        Path("../员工培训教育管理规程/APP8-SMP-HR-002-14培训师清单.docx"),
        Path("../员工培训教育管理规程/APP8培训师清单.docx"),
        Path(__file__).resolve().parent.parent.parent.parent
        / "员工培训教育管理规程"
        / "APP8-SMP-HR-002-14培训师清单.docx",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("模板文件未找到: APP8-SMP-HR-002-14培训师清单.docx")


def generate_trainer_list(trainers: Any) -> BytesIO:
    """基于 APP8 模板生成培训师清单 docx。

    模板结构：
    - 标题：培训师清单（宋体 18pt 加粗居中）
    - 批准行：批准人/日期（居中）
    - 5 列表格：姓名 | 部门 | 岗位 | 批准时间 | 备注
    - 数据行字体：宋体 10.5pt，单元格垂直居中
    - 表格边框：实线
    """
    template_path = _find_template()
    doc = Document(str(template_path))

    if not doc.tables:
        raise ValueError("模板文件中未找到表格")

    table = doc.tables[0]
    template_rows = len(table.rows)

    # 填充已有行（跳过表头）
    for idx, trainer in enumerate(trainers):
        if idx < template_rows - 1:
            row = table.rows[idx + 1]
        else:
            # 动态新增行，复制模板行样式
            new_row = table.add_row()  # type: ignore[no-untyped-call]
            row = new_row

        vals = [
            trainer.name or "",
            trainer.department or "",
            trainer.position or "",
            fmt_date_obj(trainer.approval_date),
            trainer.remarks or "",
        ]

        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            # 清除旧文本，写入新内容
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(val)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(10.5)
            # 垂直居中
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 多余的空行清空（只保留数据行）
    extra_start = len(trainers) + 1
    while len(table.rows) > extra_start:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
