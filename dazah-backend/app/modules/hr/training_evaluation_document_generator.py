"""培训评估表 Word 文档生成器（APP4-SMP-HR-002-14 模板保真填充）.

以 dict 为主要入参（前端培训资料评估表页签导出），保留旧 TrainingEvaluation
ORM 对象入口（training-evaluations 导出复用同一填充）。

保真原则：仅填充填写区，保留模板 run 格式（字体/字号/加粗/sym 复选框）。
"""

from __future__ import annotations

import logging
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document

from app.modules.hr.date_format import fmt_date_obj, fmt_date_str
from app.modules.hr.template_filler import (
    fill_after_label,
    fill_after_phrase,
    fill_whole_cell,
    header_fmt_source,
    set_sym_group,
)

logger = logging.getLogger(__name__)

METHOD_OPTIONS = ["面授", "实操", "函授", "远程教育", "其他"]
ASSESS_OPTIONS = ["笔试", "口试", "实操", "写总结"]
EVAL_RESULT_OPTIONS = ["经考核，基本达到培训效果。", "经考核，不能达到预期培训效果。"]


def _find_template() -> Path:
    """Locate the docx template, trying several path candidates."""
    candidates = [
        Path("员工培训教育管理规程/APP4-SMP-HR-002-14培训评估表.docx"),
        Path("../员工培训教育管理规程/APP4-SMP-HR-002-14培训评估表.docx"),
        Path(__file__).resolve().parent.parent.parent.parent
        / "员工培训教育管理规程"
        / "APP4-SMP-HR-002-14培训评估表.docx",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("模板文件未找到: APP4-SMP-HR-002-14培训评估表.docx")


def _d(v: Any) -> Any:
    if v is None:
        return ""
    if isinstance(v, date):
        return fmt_date_obj(v)
    return fmt_date_str(v)


def _int(v: Any) -> Any:
    return "" if v is None else str(v)


def generate_training_evaluation_doc(values: dict[str, Any]) -> BytesIO:
    """按 APP4 模板保真填充生成培训评估表 Word 文档.

    Args:
        values: dict，键与前端 TrainingEvaluationInput 一致，支持 APP4 扩展字段：
            subject/training_date/duration_hours/training_method/other_method/
            instructor/target_dept_person/expected_count/actual_count/absent_count/
            textbook/absent_handling/need_retraining/retraining_info/assessment_method/
            excellent_count/good_count/pass_count/fail_count/absent_exam_count/
            fail_handling/makeup_count/makeup_pass_count/makeup_fail_count/
            makeup_fail_handling/evaluation_result/evaluation_comment/evaluator/
            evaluate_date/has_* 与 *_qty/other_attachment。
    """
    v = values or {}
    template_path = _find_template()
    doc = Document(str(template_path))
    if not doc.tables:
        raise ValueError("模板文件中未找到表格")
    table = doc.tables[0]
    fmt_src = header_fmt_source(table)

    def cell(ri: Any, ci: Any) -> Any:
        return table.rows[ri].cells[ci]

    # R0 培训内容
    fill_after_label(cell(0, 0), f" {v.get('subject') or ''}")

    # R1 培训日期 | 课时
    fill_whole_cell(cell(1, 1), _d(v.get("training_date")), fmt_src)
    fill_whole_cell(cell(1, 4), _int(v.get("duration_hours")), fmt_src)

    # R2 培训方式(sym) | 授课人
    method = v.get("training_method") or ""
    other_txt = v.get("other_method") or ""
    if method.startswith("其他") and not other_txt:
        other_txt = method.replace("其他", "", 1).lstrip("：: ")
    set_sym_group(
        cell(2, 1), method, METHOD_OPTIONS, write_trailing={"其他": other_txt}
    )
    fill_whole_cell(cell(2, 4), v.get("instructor") or "", fmt_src)

    # R3-R4 培训对象：部门/班组/人员 + 应到/实到/缺席
    fill_after_label(cell(3, 1), f" {v.get('target_dept_person') or ''}")
    fill_after_phrase(cell(4, 1), "应到", _int(v.get("expected_count")))
    fill_after_phrase(cell(4, 1), "实到", _int(v.get("actual_count")))
    fill_after_phrase(cell(4, 1), "缺席", _int(v.get("absent_count")))

    # R5 培训教材
    fill_whole_cell(cell(5, 1), v.get("textbook") or "", fmt_src)

    # R6 缺席人员处理方式 + 是否再培训 No/Yes + 再培训说明
    fill_after_phrase(cell(6, 0), "缺席人员处理方式", v.get("absent_handling") or "")
    need_re = v.get("need_retraining")
    retrain_sel = "Yes" if need_re is True else ("No" if need_re is False else None)
    set_sym_group(cell(6, 0), retrain_sel, ["No", "Yes"])
    fill_after_phrase(cell(6, 0), "再培训", v.get("retraining_info") or "")

    # R7 考核方式(sym)
    set_sym_group(cell(7, 2), v.get("assessment_method") or "", ASSESS_OPTIONS)

    # R8 考核标准及结果：优/良好/合格/不合格/缺考
    fill_after_phrase(cell(8, 2), "优", _int(v.get("excellent_count")))
    fill_after_phrase(cell(8, 2), "良好", _int(v.get("good_count")))
    fill_after_phrase(cell(8, 2), "合格", _int(v.get("pass_count")))
    fill_after_phrase(cell(8, 2), "不合格", _int(v.get("fail_count")))
    fill_after_phrase(cell(8, 2), "缺考", _int(v.get("absent_exam_count")))

    # R9 缺考及不合格人员处理方式
    fill_whole_cell(cell(9, 2), v.get("fail_handling") or "", fmt_src)

    # R10 补考结果（空值显示"—"，与纸质模板填写规范一致）
    fill_after_phrase(cell(10, 2), "补考", _int(v.get("makeup_count")) or "—")
    fill_after_phrase(cell(10, 2), "合格", _int(v.get("makeup_pass_count")) or "—")
    fill_after_phrase(cell(10, 2), "不合格", _int(v.get("makeup_fail_count")) or "—")

    # R11 缺考及补考不合格人员处理方式
    fill_after_phrase(cell(11, 0), "处理方式", v.get("makeup_fail_handling") or "")

    # R12 培训效果评估及其他：结论(sym) + 补充文本 + 培训评估人/日期
    set_sym_group(cell(12, 0), v.get("evaluation_result") or "", EVAL_RESULT_OPTIONS)
    comment = v.get("evaluation_comment") or ""
    if comment:
        fill_after_phrase(cell(12, 0), "培训效果评估及其他", comment)
    fill_after_phrase(cell(12, 0), "培训评估人", v.get("evaluator") or "")
    fill_after_phrase(cell(12, 0), "日期", _d(v.get("evaluate_date")))

    # R13-R18 附件：有/无(sym) + 数量
    attach_fields = [
        "has_notification",
        "has_signin_sheet",
        "has_textbook",
        "has_exam_paper",
        "has_score_summary",
    ]
    for idx, fld in enumerate(attach_fields):
        row_idx = 14 + idx
        has = v.get(fld)
        sel = "有" if has is True else ("无" if has is False else None)
        set_sym_group(cell(row_idx, 5), sel, ["有", "无"])
        qty = v.get(f"{fld}_qty")
        if qty:
            fill_whole_cell(cell(row_idx, 6), str(qty), fmt_src)

    # R19 其他附件说明
    fill_after_label(cell(19, 0), f" {v.get('other_attachment') or ''}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_training_evaluation_doc_from_orm(evaluation: Any) -> BytesIO:
    """兼容旧 training-evaluations 导出：把 ORM 对象映射为 dict 后填充."""
    v = {
        "subject": getattr(evaluation, "training_content", None),
        "training_date": getattr(evaluation, "training_date", None),
        "duration_hours": getattr(evaluation, "duration_hours", None),
        "training_method": getattr(evaluation, "training_method", None),
        "other_method": getattr(evaluation, "other_method", None),
        "instructor": getattr(evaluation, "instructor", None),
        "target_dept_person": getattr(evaluation, "target_dept_person", None),
        "expected_count": getattr(evaluation, "expected_count", None),
        "actual_count": getattr(evaluation, "actual_count", None),
        "absent_count": getattr(evaluation, "absent_count", None),
        "textbook": getattr(evaluation, "textbook", None),
        "absent_handling": getattr(evaluation, "absent_handling", None),
        "need_retraining": getattr(evaluation, "need_retraining", None),
        "retraining_info": getattr(evaluation, "retraining_info", None),
        "assessment_method": getattr(evaluation, "assessment_method", None),
        "excellent_count": getattr(evaluation, "excellent_count", None),
        "good_count": getattr(evaluation, "good_count", None),
        "pass_count": getattr(evaluation, "pass_count", None),
        "fail_count": getattr(evaluation, "fail_count", None),
        "absent_exam_count": getattr(evaluation, "absent_exam_count", None),
        "fail_handling": getattr(evaluation, "fail_handling", None),
        "makeup_count": getattr(evaluation, "makeup_count", None),
        "makeup_pass_count": getattr(evaluation, "makeup_pass_count", None),
        "makeup_fail_count": getattr(evaluation, "makeup_fail_count", None),
        "makeup_fail_handling": getattr(evaluation, "makeup_fail_handling", None),
        "evaluation_result": getattr(evaluation, "evaluation_result", None),
        "evaluation_comment": getattr(evaluation, "evaluation_comment", None),
        "evaluator": getattr(evaluation, "evaluator", None),
        "evaluate_date": getattr(evaluation, "evaluate_date", None),
        "has_notification": getattr(evaluation, "has_notification", None),
        "has_signin_sheet": getattr(evaluation, "has_signin_sheet", None),
        "has_textbook": getattr(evaluation, "has_textbook", None),
        "has_exam_paper": getattr(evaluation, "has_exam_paper", None),
        "has_score_summary": getattr(evaluation, "has_score_summary", None),
        "other_attachment": getattr(evaluation, "other_attachment", None),
    }
    return generate_training_evaluation_doc(v)
