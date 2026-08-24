"""Resume folder watcher — scans
configured folder for new PDF/DOCX and
submits processing
jobs."""

import hashlib
import json
import logging
from pathlib import Path
from typing import Any

from app.shared.config_reader import get_module_setting

logger = logging.getLogger(__name__)

PROCESSED_HASHES: set[str] = set()

# 支持的简历文件类型
SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".doc")

# 去重记录文件名（与 mail_fetcher 共享）
HASH_RECORD_FILE = ".processed_hashes.json"

# 服务器端简历存储目录默认值（相对后端项目根）
_PROJECT_ROOT = Path(__file__).resolve().parents[3]
DEFAULT_RESUME_WATCH_DIR = str(_PROJECT_ROOT / "data" / "hr" / "resumes")

# 历史遗留桌面路径：检测到则自动迁移为新默认值
_LEGACY_DESKTOP_WATCH_DIR = "C:\\Users\\Administrator\\Desktop\\简历"


async def _resolve_watch_dir() -> Path:
    """解析简历监控目录；若仍为旧桌面默认值则自动迁移为服务器目录。"""
    watch_dir = await get_module_setting(
        "hr", "HR_RESUME_WATCH_DIR", DEFAULT_RESUME_WATCH_DIR
    )
    if not watch_dir or watch_dir.strip() == _LEGACY_DESKTOP_WATCH_DIR:
        from app.core.database import async_session_factory
        from app.shared.config_reader import set_module_setting

        try:
            async with async_session_factory() as session:
                await set_module_setting(
                    session, "hr", "HR_RESUME_WATCH_DIR", DEFAULT_RESUME_WATCH_DIR
                )
                await session.commit()
            logger.info(
                "简历目录已迁移为服务器默认值",
                extra={"module": "hr", "watch_dir": DEFAULT_RESUME_WATCH_DIR},
            )
        except Exception:
            logger.exception("failed to migrate resume watch dir setting")
        watch_dir = DEFAULT_RESUME_WATCH_DIR
    return Path(watch_dir)


def _load_all_hashes(save_dir: Path) -> set[str]:
    """从 JSON 文件、根目录文件、processed 子文件夹合并加载所有哈希。

    与 mail_fetcher 共享同一份去重记录，防止重复处理。
    """
    hashes: set[str] = set()

    # 1. 从 JSON 文件加载（两者共享）
    hash_file = save_dir / HASH_RECORD_FILE
    if hash_file.exists():
        try:
            with open(hash_file, encoding="utf-8") as hash_stream:
                data = json.load(hash_stream)
                if isinstance(data, list):
                    hashes.update(data)
                elif isinstance(data, dict):
                    hashes.update(data.get("hashes", []))
        except Exception:
            logger.exception("failed to load hash record file")

    # 2. 从根目录文件名中提取哈希（仅兼容旧格式；不将根目录文件哈希计入已处理，
    # 否则处理失败的文件永远无法重试）
    if save_dir.exists():
        for file_path in save_dir.iterdir():
            if not file_path.name.lower().endswith(SUPPORTED_EXTENSIONS):
                continue
            # 兼容旧格式
            name = file_path.stem
            if name.startswith("简历_"):
                parts = name.split("_", 2)
                if len(parts) >= 2:
                    hashes.add(parts[1])

    # 3. 从 processed 子文件夹加载
    processed_dir = save_dir / "processed"
    if processed_dir.exists():
        for file_path in processed_dir.iterdir():
            if not file_path.name.lower().endswith(SUPPORTED_EXTENSIONS):
                continue
            try:
                file_hash = hashlib.sha256(file_path.read_bytes()).hexdigest()[:12]
                hashes.add(file_hash)
            except Exception:
                pass
            name = file_path.stem
            if name.startswith("简历_"):
                parts = name.split("_", 2)
                if len(parts) >= 2:
                    hashes.add(parts[1])

    return hashes


def _save_hashes(save_dir: Path, hashes: set[str] | None = None) -> None:
    """保存已处理哈希到 JSON 文件（与 mail_fetcher 共享）"""
    source = hashes if hashes is not None else PROCESSED_HASHES
    try:
        hash_file = save_dir / HASH_RECORD_FILE
        with open(hash_file, "w", encoding="utf-8") as f:
            json.dump(sorted(source), f, ensure_ascii=False)
    except Exception:
        logger.exception("failed to save hash record file")


async def scan_watched_folder() -> dict[str, Any]:
    """Scan the configured folder for new PDF/DOCX files and submit processing jobs.

    新流程：新文件提交后台任务后立即删除由处理函数完成（成功写飞书后删除本地文件）；
    哈希已记录的历史文件直接删除。处理失败的文件保留且不持久化哈希，下轮扫描重试。
    """
    save_dir = await _resolve_watch_dir()
    save_dir.mkdir(parents=True, exist_ok=True)

    processed_dir = save_dir / "processed"
    processed_dir.mkdir(parents=True, exist_ok=True)

    global PROCESSED_HASHES
    # 每次都重新加载合并哈希，确保包含 mail_fetcher 下载的的文件
    baseline_hashes = _load_all_hashes(save_dir)
    PROCESSED_HASHES = set(baseline_hashes)
    logger.info(
        "restored %d processed hashes (merged from json + files + processed/)",
        len(PROCESSED_HASHES),
    )

    # 只扫描根目录，不扫描 processed 子文件夹
    resume_files: list[Path] = []
    for ext in SUPPORTED_EXTENSIONS:
        resume_files.extend(save_dir.glob(f"*{ext}"))

    new_files: list[str] = []

    for file_path in resume_files:
        try:
            file_hash = hashlib.sha256(file_path.read_bytes()).hexdigest()[:12]
        except Exception:
            logger.exception(
                "failed to hash resume file", extra={"file": file_path.name}
            )
            continue
        if file_hash in PROCESSED_HASHES:
            # 历史已处理文件：直接删除（新流程不再移入 processed 子目录）
            try:
                file_path.unlink()
                logger.info(
                    "removed already-processed resume", extra={"file": file_path.name}
                )
            except Exception:
                logger.exception(
                    "failed to remove processed resume", extra={"file": file_path.name}
                )
            continue

        # 内存去重：防止 30 秒扫描窗口内重复提交（不持久化，失败可重试）
        PROCESSED_HASHES.add(file_hash)
        new_files.append(str(file_path))

        try:
            from app.core.jobs import submit_job

            await submit_job(
                process_single_resume,
                file_path=str(file_path),
                file_hash=file_hash,
            )
            logger.info("resume queued for processing", extra={"file": file_path.name})
        except Exception:
            # 提交失败：移除内存哈希，下轮扫描重试
            PROCESSED_HASHES.discard(file_hash)
            logger.exception(
                "failed to submit resume job", extra={"file": file_path.name}
            )

    # 持久化基线哈希（不含本轮新提交任务，保证失败重试语义）
    _save_hashes(save_dir, baseline_hashes)
    return {"status": "ok", "new_files": len(new_files), "files": new_files}


def _normalize_education(raw: str) -> str:
    """将 AI 提取的学历归一化为飞书单选选项（大专/本科/硕士/博士/其他）。"""
    edu = str(raw or "").strip()
    if not edu:
        return ""
    for option in ("大专", "本科", "硕士", "博士"):
        if option in edu:
            return option
    if any(k in edu for k in ("高中", "中专", "中职", "初中", "技校")):
        return "其他"
    return "其他"


def _normalize_fit_level(raw: str) -> str:
    """将 AI 返回的符合程度归一化为飞书单选选项（高/中/低/非常满足）。"""
    level = str(raw or "").strip()
    mapping = {
        "非常满足": "非常满足",
        "强烈推荐": "非常满足",
        "高": "高",
        "推荐": "高",
        "中": "中",
        "待定": "中",
        "低": "低",
        "不推荐": "低",
    }
    return mapping.get(level, "低")


async def _match_job_position(
    target_position: str, job_names: dict[str, str]
) -> str | None:
    """将简历意向职位与招聘职位表现有职位匹配，返回职位名称（文本字段）。

    匹配规则：先精确匹配，再 LLM 同义/近似匹配；禁止编造，
    匹配不到返回 None（应聘职位留空，由人工指定）。
    """
    target = (target_position or "").strip()
    if not target or not job_names:
        return None

    for _record_id, title in job_names.items():
        if title.strip() == target:
            return title.strip()

    from app.core.llm import (
        LLMOutputError,
        LLMProviderError,
        LLMRateLimitError,
        llm_client,
    )

    options = "\n".join(f"- {title}" for title in set(job_names.values()))
    system_prompt = (
        "你是职位匹配助手。判断候选人意向职位与职位列表中哪一项最接近"
        "（同义词、近似岗位视为匹配）。若无合适项返回空字符串。只输出JSON。"
    )
    user_prompt = (
        f"候选人意向职位：{target}\n\n现有职位列表：\n{options}\n\n"
        '请返回 {"job_title": "匹配的职位名称，无匹配则为空字符串"}，'
        "job_title 必须从列表中选取，严禁编造。"
    )
    try:
        result = await llm_client.chat_json(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            expected_keys=["job_title"],
            temperature=0.0,
        )
        matched = str(result.get("job_title") or "").strip()
        valid_titles = {t.strip() for t in job_names.values()}
        if matched in valid_titles:
            logger.info(
                "resume job position matched",
                extra={"target": target, "matched": matched},
            )
            return matched
        logger.info(
            "resume job position no match, leave empty", extra={"target": target}
        )
        return None
    except (LLMOutputError, LLMProviderError, LLMRateLimitError) as e:
        logger.warning(
            "job position matching LLM failed",
            extra={"target": target, "error": str(e)},
        )
        return None


def _extract_resume_text(resume_path: Path) -> str:
    """提取 PDF/DOCX 简历文本内容。"""
    file_ext = resume_path.suffix.lower()
    if file_ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(resume_path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if file_ext in (".docx", ".doc"):
        from docx import Document

        doc = Document(str(resume_path))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    raise ValueError(f"Unsupported file type: {file_ext}")


async def process_single_resume(file_path: str, file_hash: str) -> None:
    """Process a single resume PDF/DOCX: parse -> AI analyze -> write to Feishu.

    流程：文本提取 → AI 抽取结构化字段（含意向职位/邮箱）→ 应聘职位匹配 →
    创建候选人记录 → 上传简历附件 → 全部成功后删除本地文件并持久化哈希。
    附件上传失败降级不阻断；其余步骤失败抛异常（文件保留，下轮扫描重试）。
    """
    resume_path = Path(file_path)
    if not resume_path.exists():
        logger.warning("resume file missing, skip", extra={"file": file_path})
        return

    resume_text = _extract_resume_text(resume_path)
    if not resume_text.strip():
        raise ValueError(f"Resume content is empty: {resume_path.name}")

    from app.core.llm import (
        LLMOutputError,
        LLMProviderError,
        LLMRateLimitError,
        llm_client,
    )
    from app.modules.hr.recruitment_repository import RecruitmentBitableRepo

    repo = RecruitmentBitableRepo()

    try:
        extract_system = (
            "你是一个简历解析助手。请从简历文本中抽取结构化字段并评估。只输出JSON。"
        )
        extract_user = (
            f"简历文本：\n{resume_text}\n\n请提取："
            "name(姓名),phone(手机号),email(邮箱),"
            "target_position(意向/应聘职位，无则为空字符串),"
            "education_level(学历),education_school(毕业院校),"
            "work_years(工作经验年数),skills(技能列表),"
            "last_company(最近公司),current_title(当前职位),"
            "match_rate(综合评分0-100),resume_score(简历质量0-100),"
            "fit_level(符合程度:高/中/低),reason(评分理由)"
        )
        result = await llm_client.chat_json(
            messages=[
                {"role": "system", "content": extract_system},
                {"role": "user", "content": extract_user},
            ],
            expected_keys=[
                "name",
                "phone",
                "email",
                "target_position",
                "education_level",
                "education_school",
                "work_years",
                "skills",
                "last_company",
                "current_title",
                "match_rate",
                "resume_score",
                "fit_level",
                "reason",
            ],
            temperature=0.1,
        )
    except (LLMOutputError, LLMProviderError, LLMRateLimitError) as e:
        logger.error(
            "LLM analysis failed", extra={"file": resume_path.name, "error": str(e)}
        )
        result = {
            "name": resume_path.stem,
            "match_rate": 0,
            "resume_score": 0,
            "fit_level": "低",
            "reason": f"AI分析失败: {e}",
        }

    # 应聘职位匹配（匹配不到留空，由人工指定）
    matched_job_title = None
    target_position = result.get("target_position", "")
    if target_position:
        try:
            job_names = await repo.get_job_names()
            matched_job_title = await _match_job_position(target_position, job_names)
        except Exception:
            logger.exception(
                "failed to match job position", extra={"file": resume_path.name}
            )

    phone = str(result.get("phone") or "").strip()
    email = str(result.get("email") or "").strip()
    skills = result.get("skills") or []
    if isinstance(skills, str):
        skills_text = skills.strip()
    else:
        skills_text = "、".join(str(s).strip() for s in skills if str(s).strip())
    # 使用英文键（与 F_CANDIDATE_W 映射一致，由 create_candidate 转中文字段名）
    fields: dict[str, Any] = {
        "name": result.get("name", resume_path.stem),
        "contact": f"{phone} / {email}".strip(" /"),
        "email": email,
        "education": _normalize_education(result.get("education_level", "")),
        "work_years": result.get("work_years"),
        "skills": skills_text,
        "match_rate": result.get("match_rate"),
        "resume_score": result.get("resume_score"),
        "fit_level": _normalize_fit_level(result.get("fit_level", "低")),
        "interview_status": "待安排",
        "remark": result.get("reason", ""),
    }
    if matched_job_title:
        fields["job_id"] = matched_job_title
    fields = {k: v for k, v in fields.items() if v not in (None, "", [], 0)}
    record = await repo.create_candidate(fields)
    record_id = record.get("id", "")
    logger.info(
        "candidate created from resume",
        extra={
            "file": resume_path.name,
            "candidate_name": fields.get("name"),
            "record_id": record_id,
        },
    )

    # 简历附件上传（失败降级不阻断主流程）
    try:
        client = await repo._get_client()
        if client and record_id:
            file_token = await client.upload_attachment(
                resume_path.name, resume_path.read_bytes()
            )
            if file_token:
                await repo.update_candidate(
                    record_id, {"resume_attachment": [{"file_token": file_token}]}
                )
                logger.info(
                    "resume attachment uploaded", extra={"file": resume_path.name}
                )
    except Exception:
        logger.exception(
            "failed to upload resume attachment", extra={"file": resume_path.name}
        )

    # 全部成功：删除本地文件并持久化哈希（幂等去重）
    try:
        resume_path.unlink()
        logger.info(
            "local resume removed after upload", extra={"file": resume_path.name}
        )
    except Exception:
        logger.exception(
            "failed to remove local resume", extra={"file": resume_path.name}
        )
    save_dir = await _resolve_watch_dir()
    hashes = _load_all_hashes(save_dir)
    hashes.add(file_hash)
    _save_hashes(save_dir, hashes)

    logger.info("resume processed successfully", extra={"file": resume_path.name})
