"""年度培训计划 Word 文档生成器（APP1/APP2模板）.

支持任意数量的数据行：
- 数据少于模板行：清理多余行和"……"行
- 数据等于模板行：正好填满，"……"行替换为数据
- 数据多于模板行：在最后一个表格的备注行前插入新行，并插入新的备注+审批栏
"""

import copy
import logging
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import _Row

logger = logging.getLogger(__name__)


def _find_template(plan_level: str) -> Path:
    """根据计划级别定位模板文件."""
    if plan_level == "部门级":
        template_name = "APP1年度部门培训计划表.docx"
    else:
        template_name = "APP2年度公司培训计划表.docx"

    candidates = [
        Path(f"员工培训教育管理规程/{template_name}"),
        Path(f"../员工培训教育管理规程/{template_name}"),
        Path(__file__).resolve().parent.parent.parent.parent
        / "员工培训教育管理规程"
        / template_name,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {template_name}")


def _apply_font(run: Any, size: Any = None) -> Any:
    """应用字体：中文宋体，英文/数字 Times New Roman。

    size 为字号（Pt），仅在显式传入时设置；为 None 时不改动 run 原有字号，
    以保留调用方已设定的字号（如标题小二 18pt、部门行继承样式）。数据格等
    需要五号（10.5pt）的路径应显式传入 Pt(10.5)。
    """
    run.font.name = "Times New Roman"
    if size is not None:
        run.font.size = size
    run_properties = run._element.find(qn("w:rPr"))
    if run_properties is None:
        run_properties = run._element.get_or_add_rPr()
    run_fonts = run_properties.find(qn("w:rFonts"))
    if run_fonts is None:
        run_fonts = run_properties.makeelement(qn("w:rFonts"), {})
        run_properties.insert(0, run_fonts)
    run_fonts.set(qn("w:eastAsia"), "宋体")
    run_fonts.set(qn("w:ascii"), "Times New Roman")
    run_fonts.set(qn("w:hAnsi"), "Times New Roman")


def _set_cell_text(cell: Any, text: str) -> Any:
    """设置单元格文本，保持宋体/Times New Roman字体，字号五号（10.5pt）."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            run = p.runs[0]
            run.text = text
            _apply_font(run, Pt(10.5))
        else:
            run = p.add_run(text)
            _apply_font(run, Pt(10.5))
    else:
        p = cell.add_paragraph()
        run = p.add_run(text)
        _apply_font(run, Pt(10.5))


def _remark_content_cell(row: Any) -> Any:
    """定位备注行的"内容区"单元格。

    模板备注行结构：cell[0]+cell[1] 水平合并为"备注"标签区，
    cell[2..] 合并为内容区。返回标签区之后第一个独立单元格，
    即内容区，避免把备注内容写进标签区覆盖"备注"二字。
    若结构异常（无独立内容格）则回退到 cells[1]。
    """
    cells = row.cells
    label_tc = cells[0]._tc
    for c in cells[1:]:
        if c._tc is not label_tc:
            return c
    return cells[1] if len(cells) > 1 else cells[0]


def _make_type_sym_run(char_code: str) -> Any:
    """构造培训类型勾选框 run（Wingdings 2 符号，14pt），与模板原结构一致."""
    from docx.oxml import OxmlElement

    r = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    r.append(run_properties)
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:hint"), "default")
    run_fonts.set(qn("w:ascii"), "Wingdings 2")
    run_fonts.set(qn("w:hAnsi"), "Wingdings 2")
    run_fonts.set(qn("w:cs"), "Wingdings 2")
    run_properties.append(run_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "28")
    run_properties.append(sz)
    complex_size = OxmlElement("w:szCs")
    complex_size.set(qn("w:val"), "28")
    run_properties.append(complex_size)
    sym = OxmlElement("w:sym")
    sym.set(qn("w:font"), "Wingdings 2")
    sym.set(qn("w:char"), char_code)
    r.append(sym)
    return r


def _make_type_text_run(text: str, big: bool = False) -> Any:
    """构造培训类型文字/空格 run，big=True 时字号 14pt（内训/外训间空格）."""
    from docx.oxml import OxmlElement

    r = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    r.append(run_properties)
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), "Times New Roman")
    run_fonts.set(qn("w:hAnsi"), "Times New Roman")
    run_fonts.set(qn("w:eastAsia"), "宋体")
    run_properties.append(run_fonts)
    if big:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "28")
        run_properties.append(sz)
        complex_size = OxmlElement("w:szCs")
        complex_size.set(qn("w:val"), "28")
        run_properties.append(complex_size)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _fill_training_type_cell(cell: Any, training_type: str) -> Any:
    """填充培训类型单元格，保留模板的 Wingdings 勾选框结构与字号.

    模板数据行的培训类型格为 [框][内训][空格][框][外训]，框是 14pt 的
    Wingdings 2 符号（0052=勾选，00A3=未勾选）。优先只切换框的 char，
    完整保留字号/字体/换行布局；若格子无该结构（如"……"行）则按模板重建，
    保证所有数据行格式一致、行高对齐。
    """
    val = training_type or ""
    want_inner = "内训" in val
    want_outer = "外训" in val
    inner_char = "0052" if want_inner else "00A3"
    outer_char = "0052" if want_outer else "00A3"

    # 统一按模板规格重建该格：[框][内训][空格][框][外训]。
    # 不区分模板原有行与克隆插入行——克隆行 _clone_row_template 会清空文字，
    # 若仅切换框 char 会导致克隆行只剩框、缺"内训/外训"文字。统一重建保证全表一致。
    for p in cell.paragraphs:
        for r in list(p.runs):
            r._r.getparent().remove(r._r)
    target = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    for el in (
        _make_type_sym_run(inner_char),
        _make_type_text_run("内训"),
        _make_type_text_run("   ", big=True),
        _make_type_sym_run(outer_char),
        _make_type_text_run("外训"),
    ):
        target._p.append(el)


def _fill_year_in_paragraphs(doc: Any, year: int) -> Any:
    """在标题"年度...培训计划表"的年份占位处填入年份，保留下划线格式.

    模板标题形如"______年度部门培训计划表"，年份占位由带下划线
    (underline=single) 的空格 run 表示。直接就地把年份写入最长的下划线占位
    run，完整保留其下划线/字号/加粗等格式，使年份仍显示在下划线上；
    若无下划线占位 run 则在"年度"前插入一个带下划线的年份 run 兜底。
    """
    from docx.oxml import OxmlElement

    year_str = str(year)

    def _underline_val(r: Any) -> Any:
        run_properties = r._element.find(qn("w:rPr"))
        if run_properties is None:
            return None
        u = run_properties.find(qn("w:u"))
        return u.get(qn("w:val")) if u is not None else None

    for para in doc.paragraphs:
        text = para.text
        if "年度" not in text or "培训计划表" not in text or year_str in text:
            continue
        runs = list(para.runs)

        # 1) 优先就地填下划线占位 run（仅改文本，rPr 原样保留 → 下划线不丢）
        ul_runs = [r for r in runs if _underline_val(r) == "single"]
        if ul_runs:
            target = max(ul_runs, key=lambda r: len(r.text))
            target.text = year_str
            for r in ul_runs:
                if r is not target:
                    r.text = ""  # 清空其余占位空格，避免年份两侧多余下划线
            continue

        # 2) 兜底：在"年度"所在 run 前插入带下划线/加粗/标题字号的年份 run
        for r in runs:
            if "年度" not in r.text:
                continue
            idx = r.text.find("年度")
            prefix, suffix = r.text[:idx], r.text[idx:]
            r.text = suffix

            year_run = OxmlElement("w:r")
            year_properties = OxmlElement("w:rPr")
            year_run.append(year_properties)
            year_properties.append(OxmlElement("w:b"))
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            year_properties.append(u)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "36")
            year_properties.append(sz)
            complex_size = OxmlElement("w:szCs")
            complex_size.set(qn("w:val"), "36")
            year_properties.append(complex_size)
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = year_str
            year_run.append(t)

            if prefix:
                pre_run = OxmlElement("w:r")
                prefix_properties = OxmlElement("w:rPr")
                pre_run.append(prefix_properties)
                prefix_properties.append(OxmlElement("w:b"))
                psz = OxmlElement("w:sz")
                psz.set(qn("w:val"), "36")
                prefix_properties.append(psz)
                pt = OxmlElement("w:t")
                pt.set(qn("xml:space"), "preserve")
                pt.text = prefix
                pre_run.append(pt)
                r._r.addprevious(pre_run)
            r._r.addprevious(year_run)
            break


def _get_table_info(table: Any) -> Any:
    """解析表格结构，返回表头行索引、数据行索引列表、备注行索引、审批行索引列表."""
    header_row = 0
    for ri, row in enumerate(table.rows):
        row_text = " ".join(cell.text.strip() for cell in row.cells)
        if "序号" in row_text and "培训类型" in row_text:
            header_row = ri
            break

    data_rows = []
    remark_row = None
    approve_rows = []

    for ri in range(header_row + 1, len(table.rows)):
        row_text = " ".join(cell.text.strip() for cell in table.rows[ri].cells)
        if "备注" in row_text:
            remark_row = ri
        elif "制表人" in row_text or "签名" in row_text:
            approve_rows.append(ri)
        elif "……" in row_text:
            data_rows.append(ri)
        else:
            data_rows.append(ri)

    return header_row, data_rows, remark_row, approve_rows


def _get_column_map(table: Any, header_row: int, plan_level: Any) -> dict[int, str]:
    """解析表头行，建立列索引到字段名的映射.

    处理合并单元格：同一字段占多列时，只取第一列。
    """
    row = table.rows[header_row]
    cells = row.cells
    col_map: dict[int, str] = {}  # 列索引 -> 字段名
    seen_fields: set[str] = set()

    for ci in range(len(cells)):
        text = cells[ci].text.strip()
        if not text:
            continue
        field = None
        if "序号" in text:
            field = "seq"
        elif "培训类型" in text or ("类型" in text and field is None):
            field = "training_type"
        elif "培训时间" in text or "月度" in text:
            field = "training_month"
        elif "培训内容" in text or "教材" in text or "内容" in text:
            field = "content_textbook"
        elif "培训对象" in text or "对象" in text:
            field = "target_audience"
        elif "授课" in text:
            field = "instructor"
        elif "考核" in text:
            field = "assessment_method"

        if field and field not in seen_fields:
            col_map[ci] = field
            seen_fields.add(field)

    return col_map


def _fill_row(row: Any, item: Any, seq: int, col_map: dict[int, str]) -> Any:
    """根据列映射填充一行数据."""
    cells = row.cells

    for ci, field in col_map.items():
        if ci >= len(cells):
            continue
        cell = cells[ci]

        if field == "seq":
            _set_cell_text(cell, str(seq))
        elif field == "training_type":
            _fill_training_type_cell(cell, item.training_type or "")
        elif field == "training_month":
            _set_cell_text(cell, item.training_month or "")
        elif field == "content_textbook":
            _set_cell_text(cell, item.content_textbook or "")
        elif field == "target_audience":
            _set_cell_text(cell, item.target_audience_new or "")
        elif field == "instructor":
            _set_cell_text(cell, item.instructor or "")
        elif field == "assessment_method":
            _set_cell_text(cell, item.assessment_method or "")


def _clear_row(row: Any) -> Any:
    """清空一行的所有单元格内容."""
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.text = ""


def _clone_row_template(table: Any, template_row_idx: Any) -> Any:
    """深拷贝模板行，返回新的 _tr 元素（已清空内容）."""
    template_tr = table.rows[template_row_idx]._tr
    new_tr = copy.deepcopy(template_tr)
    for tc in new_tr.findall(qn("w:tc")):
        for p in tc.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                for t in r.findall(qn("w:t")):
                    t.text = ""
    return new_tr


def generate_annual_plan_doc(plan: Any, items: Any) -> BytesIO:
    """填充年度培训计划 Word 模板.

    支持任意数量的数据：
    - 数据少于模板行：清理多余行
    - 数据多于模板行：插入新行+审批栏

    Args:
        plan: AnnualTrainingPlan ORM 对象（含 year, department, plan_level）
        items: AnnualTrainingPlanItem 列表

    Returns:
        BytesIO buffer containing the generated docx
    """
    template_path = _find_template(plan.plan_level or "公司级")
    doc = Document(str(template_path))
    plan_level = plan.plan_level or "公司级"

    # 1. 填充年份
    _fill_year_in_paragraphs(doc, plan.year)

    # 2. 填充部门名称 + 版本号 + 行内布局（部门行可能跨多页，逐页处理）
    # 模板部门行形如"部门：______ 版本：__"，中间预留大段空白占位；若直接填值，
    # "部门名+空白+版本"会超出一行宽，使版本号折到下一行（"跑下去"）。
    # 故填值后清空中间空白，改用右对齐制表位让"版本：xx"靠右并与部门名同行。
    from docx.oxml import OxmlElement

    version = getattr(plan, "version", None)

    def _norm(t: str) -> str:
        return (
            (t or "")
            .replace(" ", "")
            .replace("　", "")
            .replace("\n", "")
            .replace("\r", "")
        )

    for para in doc.paragraphs:
        if "培训计划表" in _norm(para.text):
            continue
        runs = list(para.runs)
        if not runs:
            continue

        # 定位部门冒号 run 与版本标签 run（"部门"/"："、"版本"/"：" 可能分处不同 run）
        dept_colon_idx = None
        ver_idx = None
        seen_dept = False
        for i, r in enumerate(runs):
            n = _norm(r.text)
            if "部门" in n:
                seen_dept = True
            if (
                seen_dept
                and dept_colon_idx is None
                and "：" in r.text
                and "版本" not in n
            ):
                dept_colon_idx = i
            if ver_idx is None and "：" in r.text and "版本" in n:
                ver_idx = i
        if ver_idx is None:
            for i, r in enumerate(runs):
                if "版本" in _norm(r.text):
                    ver_idx = i
                    break

        # 填部门名称（仅部门级且本段含部门冒号）
        dept_name_idx = None
        if plan_level == "部门级" and dept_colon_idx is not None:
            end = ver_idx if ver_idx is not None else len(runs)
            colon_run = runs[dept_colon_idx]
            after = colon_run.text.split("：", 1)[1] if "：" in colon_run.text else ""
            if after.strip():
                colon_run.text = (
                    colon_run.text.split("：", 1)[0] + "：" + plan.department
                )
                _apply_font(colon_run)
                dept_name_idx = dept_colon_idx
            else:
                for j in range(dept_colon_idx + 1, end):
                    if not runs[j].text.strip():
                        runs[j].text = plan.department
                        _apply_font(runs[j])
                        dept_name_idx = j
                        break

        # 填版本号
        if version and ver_idx is not None:
            vr = runs[ver_idx]
            if "：" in vr.text:
                vr.text = vr.text.split("：", 1)[0] + "：" + version
            else:
                vr.text = vr.text + "：" + version
            _apply_font(vr)

        # 行内布局：右对齐制表位 + 清空中间空白 + 制表符，使版本号靠右且不折行
        if ver_idx is not None:
            sec = doc.sections[0]
            page_width = sec.page_width
            left_margin = sec.left_margin
            right_margin = sec.right_margin
            if page_width is None or left_margin is None or right_margin is None:
                twips = 0
            else:
                twips = int(
                    (int(page_width) - int(left_margin) - int(right_margin)) / 635
                )
            paragraph_properties = para._p.get_or_add_pPr()
            tabs = paragraph_properties.find(qn("w:tabs"))
            if tabs is None:
                tabs = OxmlElement("w:tabs")
                paragraph_properties.append(tabs)
            for t in list(tabs.findall(qn("w:tab"))):
                if t.get(qn("w:val")) == "right":
                    tabs.remove(t)
            te = OxmlElement("w:tab")
            te.set(qn("w:val"), "right")
            te.set(qn("w:pos"), str(twips))
            tabs.append(te)

            anchor_idx = dept_name_idx if dept_name_idx is not None else dept_colon_idx
            clear_start = (anchor_idx + 1) if anchor_idx is not None else 0
            for j in range(clear_start, ver_idx):
                runs[j].text = ""
            tab_run = OxmlElement("w:r")
            tab_run.append(OxmlElement("w:tab"))
            if anchor_idx is not None:
                runs[anchor_idx]._r.addnext(tab_run)
            else:
                runs[ver_idx]._r.addprevious(tab_run)

    # 3. 收集有效明细
    valid_items = [i for i in items if not i.is_deleted]

    # 4. 逐表填充数据
    item_idx = 0
    total_tables = len(doc.tables)

    for table_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue

        header_row, data_rows, remark_row, approve_rows = _get_table_info(table)
        if not data_rows:
            continue

        # 解析列映射
        col_map = _get_column_map(table, header_row, plan_level)

        # 计算本表需要填充多少条数据
        # 策略：均匀分配到各表，最后一个表兜底
        remaining = len(valid_items) - item_idx
        if table_idx < total_tables - 1:
            # 非最后一个表：填满数据行（含"……"行）
            fill_count = min(remaining, len(data_rows))
        else:
            # 最后一个表：填满后剩余的插入新行
            fill_count = remaining

        # 填充模板中的数据行
        rows_to_fill = min(len(data_rows), fill_count)
        for i in range(rows_to_fill):
            if item_idx >= len(valid_items):
                break
            row = table.rows[data_rows[i]]
            _fill_row(row, valid_items[item_idx], item_idx + 1, col_map)
            item_idx += 1

        # 清理多余的模板数据行（含"……"行）：先收集_tr再删除，避免索引偏移
        if rows_to_fill < len(data_rows):
            trs_to_remove = [
                table.rows[data_rows[i]]._tr
                for i in range(rows_to_fill, len(data_rows))
            ]
            for tr in trs_to_remove:
                tr.getparent().remove(tr)

        # 最后一个表格：插入超出部分的新行（在备注行前插入）
        if table_idx == total_tables - 1 and item_idx < len(valid_items):
            if remark_row is not None:
                remark_tr = table.rows[remark_row]._tr
                template_row_idx = data_rows[0]

                while item_idx < len(valid_items):
                    new_tr = _clone_row_template(table, template_row_idx)
                    remark_tr.addprevious(new_tr)
                    new_row = _Row(new_tr, table)
                    _fill_row(new_row, valid_items[item_idx], item_idx + 1, col_map)
                    item_idx += 1
                # 原始备注行和审批行已在正确位置，无需复制

    # 4.5 填充备注行内容（写入"备注"标签区之后的内容区，保留"备注"二字）
    plan_remarks = getattr(plan, "remarks", None)
    if plan_remarks:
        for table in doc.tables:
            for row in table.rows:
                cells_text = [c.text.strip() for c in row.cells]
                if cells_text and cells_text[0] == "备注":
                    _set_cell_text(_remark_content_cell(row), plan_remarks)
                    break

    # 5. 清理空表格页：数据行全部被删除的表格（连同标题段落）整页移除
    _remove_empty_table_pages(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _remove_empty_table_pages(doc: Any) -> Any:
    """移除没有数据行的表格及其关联的标题段落（整页清理）.

    模板结构：标题段落(含"年度公司/部门培训计划表") + permEnd + Table + 分隔段落
    当某表格的所有数据行被清理后，该页只剩表头+审批栏，应整页移除。
    但至少保留一个表格（防止全部删空）。
    """
    body = doc.element.body
    children = list(body)

    # 找出所有表格元素及其在 body 中的索引
    table_indices = [i for i, c in enumerate(children) if c.tag == qn("w:tbl")]

    # 统计每个表格的数据行数（序号列是数字的行）
    tables_to_remove = []
    for ti, tbl_idx in enumerate(table_indices):
        tbl = children[tbl_idx]
        rows = tbl.findall(qn("w:tr"))
        data_count = 0
        for tr in rows:
            cells = tr.findall(qn("w:tc"))
            if not cells:
                continue
            texts = cells[0].findall(".//" + qn("w:t"))
            cell_text = "".join(t.text or "" for t in texts).strip()
            if cell_text.isdigit():
                data_count += 1
        if data_count == 0:
            tables_to_remove.append((ti, tbl_idx))

    # 至少保留最后一个表格（兜底）
    if len(tables_to_remove) >= len(table_indices):
        tables_to_remove = tables_to_remove[:-1]

    if not tables_to_remove:
        return

    # 为每个要移除的表格计算其页面块范围 [start, end]
    page_blocks = []
    for ti, tbl_idx in tables_to_remove:
        # 向前查找关联的标题段落（直到上一个表格或body开头）
        start = tbl_idx
        for j in range(tbl_idx - 1, -1, -1):
            if children[j].tag == qn("w:tbl"):
                break
            start = j

        # 向后查找分隔段落（空段落，直到下一个标题或表格）
        end = tbl_idx
        for j in range(tbl_idx + 1, len(children)):
            child = children[j]
            tag = child.tag
            if tag in (qn("w:tbl"), qn("w:sectPr")):
                break
            if tag == qn("w:p"):
                text = "".join(child.itertext())
                if "年度" in text and "培训计划表" in text:
                    break
                end = j
            else:
                break

        page_blocks.append((start, end))

    # 合并重叠的页面块
    page_blocks.sort()
    merged = [page_blocks[0]]
    for start, end in page_blocks[1:]:
        if start <= merged[-1][1] + 1:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))

    # 收集所有要删除的元素，然后一次性删除
    elements_to_remove = set()
    for start, end in merged:
        for j in range(start, end + 1):
            elements_to_remove.add(id(children[j]))

    for child in list(body):
        if id(child) in elements_to_remove:
            body.remove(child)
