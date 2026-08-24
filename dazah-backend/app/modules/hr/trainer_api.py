"""培训师管理 API"""

import logging
from datetime import date
from datetime import datetime as dt
from io import BytesIO
from typing import Any
from urllib.parse import quote
from uuid import UUID

from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, File, Query, UploadFile
from fastapi.responses import StreamingResponse
from openpyxl import load_workbook  # type: ignore[import-untyped]
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.core.database import get_db
from app.core.deps import CurrentUser
from app.core.exceptions import AppException, NotFoundException
from app.core.response import paginated_response, success_response
from app.core.upload_security import read_upload_secure
from app.modules.hr.schemas import (
    TrainerCreate,
    TrainerResponse,
    TrainerUpdate,
)
from app.modules.hr.trainer_service import TrainerService

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/trainers", tags=["人事-培训师管理"])


def _require_user(current_user: CurrentUser) -> Any:
    if current_user is None:
        raise AppException(status_code=401, message="请先登录")


def _parse_date_value(raw: Any) -> date | None:
    """解析多种日期格式（2024-03-21 / 2024.03.21 / 2024/3/21 / 2024年3月21日）。"""
    if raw is None:
        return None
    if isinstance(raw, dt):
        return raw.date()
    if isinstance(raw, date):
        return raw
    text = str(raw).strip()
    if not text:
        return None
    for fmt in ("%Y-%m-%d", "%Y.%m.%d", "%Y/%m/%d", "%Y年%m月%d日"):
        try:
            return dt.strptime(text, fmt).date()
        except ValueError:
            continue
    logger.warning("日期解析失败: %s", text)
    return None


def _parse_docx_trainers(content: bytes) -> list[dict[str, Any]]:
    """解析 APP8 培训师清单 Word 文档中的表格。

    自动定位表头含「姓名」「部门」的表格，按表头映射列。
    """
    doc = DocxDocument(BytesIO(content))
    if not doc.tables:
        raise AppException(
            status_code=400,
            message="Word 文档中未找到表格，请使用 APP8-SMP-HR-002-14 培训师清单模板",
        )

    table = None
    col_map: dict[str, int] = {}
    for t in doc.tables:
        header = [cell.text.strip() for cell in t.rows[0].cells]
        mapping = _match_trainer_header(header)
        if mapping:
            table = t
            col_map = mapping
            break

    if table is None:
        raise AppException(
            status_code=400,
            message=(
                "文档表格表头与培训师清单不符（需含 姓名/部门/岗位/批准时"
                "间/备注），请检查文"
                "件"
            ),
        )

    rows_data: list[dict[str, Any]] = []
    for row in table.rows[1:]:
        vals = [cell.text.strip() for cell in row.cells]
        name = _cell(vals, col_map, "name")
        if not name:
            continue
        rows_data.append(
            {
                "name": name,
                "department": _cell(vals, col_map, "department") or None,
                "position": _cell(vals, col_map, "position") or None,
                "approval_date": _parse_date_value(
                    _cell(vals, col_map, "approval_date")
                ),
                "approver": _cell(vals, col_map, "approver") or None,
                "remarks": _cell(vals, col_map, "remarks") or None,
            }
        )
    return rows_data


def _parse_excel_trainers(content: bytes) -> list[dict[str, Any]]:
    """解析 Excel 培训师清单，按首行表头映射列。"""
    wb = load_workbook(BytesIO(content), data_only=True)
    ws = wb.active

    header_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), None)
    header = [str(v).strip() if v is not None else "" for v in (header_row or [])]
    col_map = _match_trainer_header(header)
    if not col_map:
        raise AppException(
            status_code=400,
            message="Excel 表头与培训师清单不符（需含 姓名/部门 列），请检查文件",
        )

    rows_data: list[dict[str, Any]] = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        vals = [str(v).strip() if v is not None else "" for v in row]
        name = _cell(vals, col_map, "name")
        if not name:
            continue
        rows_data.append(
            {
                "name": name,
                "department": _cell(vals, col_map, "department") or None,
                "position": _cell(vals, col_map, "position") or None,
                "approval_date": _parse_date_value(
                    row[col_map["approval_date"]]
                    if "approval_date" in col_map
                    and len(row) > col_map["approval_date"]
                    else None
                ),
                "approver": _cell(vals, col_map, "approver") or None,
                "remarks": _cell(vals, col_map, "remarks") or None,
            }
        )
    return rows_data


def _match_trainer_header(header: list[str]) -> dict[str, int]:
    """表头必须含「姓名」「部门」才认为是培训师清单表格，返回字段→列下标映射。"""
    idx = {h: i for i, h in enumerate(header) if h}
    mapping: dict[str, int] = {}
    for field, names in {
        "name": ["姓名", "培训师姓名"],
        "department": ["部门"],
        "position": ["岗位", "职位"],
        "approval_date": ["批准时间", "批准日期"],
        "approver": ["批准人"],
        "remarks": ["备注"],
    }.items():
        for n in names:
            if n in idx:
                mapping[field] = idx[n]
                break
    if "name" not in mapping or "department" not in mapping:
        return {}
    return mapping


def _cell(vals: list[str], col_map: dict[str, int], field: str) -> str:
    i = col_map.get(field)
    if i is None or i >= len(vals):
        return ""
    return vals[i]


@router.get("", summary="培训师列表")
async def list_trainers(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    keyword: str | None = Query(None, description="姓名搜索"),
    department: str | None = Query(None, description="部门筛选"),
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    from app.modules.hr.api import _assert_dept_in_scope

    alias_set = await _assert_dept_in_scope(db, current_user, department)
    service = TrainerService(db)
    trainers, total = await service.list_trainers(
        page=page,
        page_size=page_size,
        keyword=keyword,
        department=department,
        dept_alias_set=alias_set,
    )
    return paginated_response(
        data=[
            TrainerResponse.model_validate(t).model_dump(mode="json") for t in trainers
        ],
        page=page,
        page_size=page_size,
        total=total,
    )


@router.post("", summary="创建培训师")
async def create_trainer(
    data: TrainerCreate,
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    _require_user(current_user)
    service = TrainerService(db)
    trainer = await service.create(data.model_dump(exclude_unset=True))
    return success_response(
        data=TrainerResponse.model_validate(trainer).model_dump(mode="json"),
        message="创建成功",
    )


@router.post("/import", summary="导入培训师清单")
async def import_trainers(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    """导入培训师清单，支持 Word 文档（APP8 模板）和 Excel。

    Word 文档解析表格：姓名 | 部门 | 岗位 | 批准时间 | 备注（首行表头）
    Excel 表头：姓名 | 部门 | 岗位 | 批准时间 | 批准人 | 备注
    重复判断：姓名 + 部门相同则更新，否则新增。
    """
    _require_user(current_user)
    safe_name, content = await read_upload_secure(
        file,
        max_bytes=get_settings().MAX_UPLOAD_SIZE_MB * 1024 * 1024,
        allowed_extensions={".docx", ".xlsx", ".xls"},
        what="培训师清单",
    )
    ext = safe_name.rsplit(".", 1)[-1].lower()

    try:
        if ext == "docx":
            rows_data = _parse_docx_trainers(content)
        elif ext in ("xlsx", "xls"):
            rows_data = _parse_excel_trainers(content)
        else:
            raise AppException(
                status_code=400,
                message="仅支持 Word 文档（.docx）或 Excel 文件（.xlsx / .xls）",
            )
    except AppException:
        raise
    except Exception:
        logger.exception("解析导入文件失败")
        raise AppException(status_code=400, message="文件格式错误，无法读取")

    if not rows_data:
        raise AppException(status_code=400, message="文件中无有效数据")

    service = TrainerService(db)
    result = await service.import_trainers(rows_data)
    logger.info(
        "培训师导入完成",
        extra={
            "import_created": result["created"],
            "import_updated": result["updated"],
            "import_skipped": result["skipped"],
            "import_total": result["total"],
        },
    )
    return success_response(
        data=result,
        message=(
            f"导入完成：新增 {result['created']} 条，更新 "
            f"{result['updated']} 条，跳过 {result['skipped']} 条"
        ),
    )


@router.get("/export", summary="导出培训师清单")
async def export_trainers(
    department: str | None = Query(None, description="部门筛选后导出"),
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    """导出培训师清单 Word 文档（APP8-SMP-HR-002-14）"""
    _require_user(current_user)
    from app.modules.hr.trainer_document_generator import generate_trainer_list

    service = TrainerService(db)
    trainers, _ = await service.list_trainers(
        page=1, page_size=500, department=department
    )

    buffer: BytesIO = generate_trainer_list(trainers)
    # 按部门导出时文件名带部门名，如：101二车间培训师清单.docx
    base_name = (
        f"{department}培训师清单" if department else "APP8-SMP-HR-002-14培训师清单"
    )
    filename = quote(f"{base_name}.docx")

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )


@router.get("/{trainer_id}", summary="培训师详情")
async def get_trainer(
    trainer_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    _require_user(current_user)
    service = TrainerService(db)
    trainer = await service.get_by_id(trainer_id)
    if not trainer:
        raise NotFoundException(resource="培训师", resource_id=str(trainer_id))
    return success_response(
        data=TrainerResponse.model_validate(trainer).model_dump(mode="json")
    )


@router.put("/{trainer_id}", summary="更新培训师")
async def update_trainer(
    trainer_id: UUID,
    data: TrainerUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    _require_user(current_user)
    service = TrainerService(db)
    trainer = await service.update(trainer_id, data.model_dump(exclude_unset=True))
    if not trainer:
        raise NotFoundException(resource="培训师", resource_id=str(trainer_id))
    return success_response(
        data=TrainerResponse.model_validate(trainer).model_dump(mode="json"),
        message="更新成功",
    )


@router.delete("/{trainer_id}", summary="删除培训师")
async def delete_trainer(
    trainer_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: CurrentUser = None,
) -> Any:
    _require_user(current_user)
    service = TrainerService(db)
    deleted = await service.delete(trainer_id)
    if not deleted:
        raise NotFoundException(resource="培训师", resource_id=str(trainer_id))
    return success_response(message="删除成功")
