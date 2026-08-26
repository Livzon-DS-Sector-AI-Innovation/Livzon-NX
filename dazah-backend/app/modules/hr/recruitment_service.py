"""HR recruitment and onboarding service layer."""

import asyncio
import logging
from typing import Any, cast

from app.core.llm import LLMOutputError, LLMProviderError, LLMRateLimitError, llm_client
from app.modules.hr.recruitment_repository import TBL_CANDIDATE, RecruitmentBitableRepo

logger = logging.getLogger(__name__)

# ─── Field Maps (Pydantic schema field → Feishu Chinese field name) ───

JOB_FIELD_MAP: dict[str, str] = {
    "title": "职位名称",
    "description": "岗位描述",
    "requirement": "任职要求",
    "salary_range": "薪资范围",
    "location": "工作地点",
    "req_skills": "要求技能",
    "status": "招聘状态",
    "publish_date": "发布时间",
}

CANDIDATE_FIELD_MAP: dict[str, str] = {
    "name": "姓名",
    "contact": "联系方式",
    "email": "邮箱",
    "job_id": "应聘职位",
    "education": "学历",
    "work_years": "工作经验(年)",
    "skills": "技能标签",
    "match_rate": "技能匹配度",
    "resume_score": "简历评分",
    "fit_level": "招聘符合程度",
    "interview_status": "面试状态",
    "interview_time": "面试时间",
    "interviewer": "面试官",
    "remark": "备注",
    "source_channel": "来源渠道",
}


def _translate(data: dict[str, Any], field_map: dict[str, str]) -> dict[str, Any]:
    """Translate dict keys from Pydantic field names to Feishu Chinese field names."""
    return {field_map.get(k, k): v for k, v in data.items()}


def _untranslate(data: dict[str, Any], field_map: dict[str, str]) -> dict[str, Any]:
    "Translate dict keys from Feishu Chinese field names back to Pydantic field names."
    reverse = {v: k for k, v in field_map.items()}
    return {reverse.get(k, k): v for k, v in data.items()}


# ─── RecruitmentService ───


class RecruitmentService:
    def __init__(self, app_token: str | None = None) -> None:
        self.repo = RecruitmentBitableRepo(app_token=app_token)

    # -- Jobs --

    async def list_jobs(
        self, keyword: str | None = None, page: int = 1, page_size: int = 20
    ) -> tuple[list[dict[str, Any]], int]:
        logger.info("listing jobs", extra={"keyword": keyword, "page": page})
        result = await self.repo.list_jobs(
            keyword=keyword, page=page, page_size=page_size
        )
        items = result.get("items", []) if isinstance(result, dict) else []
        total = result.get("total", 0) if isinstance(result, dict) else 0
        if isinstance(result, tuple):
            items = result[0] if result else []
            total = result[1] if len(result) > 1 else len(items)
        elif isinstance(result, list):
            items = result
            total = len(items)
        # Translate Feishu field names back to schema names
        jobs = [_untranslate(item, JOB_FIELD_MAP) for item in items]
        # Attach candidate count per job
        all_cands, _ = await self.repo.list_candidates()
        count_map: dict[str, int] = {}
        for c in all_cands:
            jid = c.get("job_id", "")
            if jid:
                count_map[jid] = count_map.get(jid, 0) + 1
        for j in jobs:
            j["candidate_count"] = count_map.get(j["id"], 0)
        return jobs, total

    async def create_job(self, data: dict[str, Any]) -> dict[str, Any]:
        logger.info("creating job", extra={"title": data.get("title")})
        result = await self.repo.create_job(data)
        return _untranslate(result, JOB_FIELD_MAP)

    async def get_job(self, record_id: str) -> dict[str, Any]:
        logger.info("getting job", extra={"record_id": record_id})
        result = await self.repo.get_job(record_id)
        return _untranslate(result, JOB_FIELD_MAP)

    async def update_job(self, record_id: str, data: dict[str, Any]) -> dict[str, Any]:
        logger.info("updating job", extra={"record_id": record_id})
        result = await self.repo.update_job(record_id, data)
        return _untranslate(result, JOB_FIELD_MAP)

    # -- Candidates --

    async def list_candidates(
        self,
        keyword: str | None = None,
        fit_level: str | None = None,
        interview_status: str | None = None,
        job_id: str | None = None,
        page: int = 1,
        page_size: int = 20,
    ) -> tuple[list[dict[str, Any]], int]:
        logger.info(
            "listing candidates",
            extra={
                "keyword": keyword,
                "fit_level": fit_level,
                "interview_status": interview_status,
                "job_id": job_id,
            },
        )
        result = await self.repo.list_candidates(
            keyword=keyword,
            fit_level=fit_level,
            interview_status=interview_status,
            job_id=job_id,
            page=page,
            page_size=page_size,
        )
        if isinstance(result, dict):
            items, total = result.get("items", []), result.get("total", 0)
        elif isinstance(result, tuple):
            items = result[0] if result else []
            total = result[1] if len(result) > 1 else len(items)
        elif isinstance(result, list):
            items, total = result, len(result)
        else:
            items, total = [], 0
        items = [_untranslate(item, CANDIDATE_FIELD_MAP) for item in items]
        # resolve job_id (record id) to job_position (display name)
        job_names = await self.repo.get_job_names()
        for item in items:
            jid = item.get("job_id")
            if jid and jid in job_names:
                item["job_position"] = job_names[jid]
        return items, total

    async def create_candidate(self, data: dict[str, Any]) -> dict[str, Any]:
        logger.info("creating candidate", extra={"candidate_name": data.get("name")})
        result = await self.repo.create_candidate(data)
        return _untranslate(result, CANDIDATE_FIELD_MAP)

    async def get_candidate(self, record_id: str) -> dict[str, Any]:
        logger.info("getting candidate", extra={"record_id": record_id})
        result = await self.repo.get_candidate(record_id)
        item = _untranslate(result, CANDIDATE_FIELD_MAP)
        jid = item.get("job_id")
        if jid:
            job_names = await self.repo.get_job_names()
            if jid in job_names:
                item["job_position"] = job_names[jid]
        return item

    async def update_candidate(
        self, record_id: str, data: dict[str, Any]
    ) -> dict[str, Any]:
        logger.info("updating candidate", extra={"record_id": record_id})
        result = await self.repo.update_candidate(record_id, data)
        untrans = _untranslate(result, CANDIDATE_FIELD_MAP)

        # 检查面试状态变更，自动发送邮件
        new_status = data.get("interview_status")
        if new_status in ("通过", "不符合"):
            email = untrans.get("email")
            name = untrans.get("name", "")
            if email:
                try:
                    await self._send_status_email(email, name, new_status, untrans)
                except Exception as e:
                    logger.exception(
                        "failed to send status email",
                        extra={
                            "candidate_id": record_id,
                            "status": new_status,
                            "error": str(e),
                        },
                    )

        return untrans

    async def _send_status_email(
        self, email: str, name: str, status: str, candidate_data: dict[str, Any]
    ) -> None:
        """根据面试状态发送录用或拒绝邮件。

        注意：Offer 邮件只发给 HR 自己（配置的邮箱），不发给候选人。
        """
        from app.modules.hr.mail_sender import send_email_with_template
        from app.shared.config_reader import get_module_setting

        if status == "通过":
            subject_tpl = await get_module_setting(
                "hr", "HR_MAIL_OFFER_SUBJECT", "录用通知 - {name}"
            )
            body_tpl = await get_module_setting("hr", "HR_MAIL_OFFER_BODY", "")
            template_path = await get_module_setting(
                "hr", "HR_MAIL_OFFER_TEMPLATE_PATH", ""
            )

            # 替换变量
            variables = {
                "name": name,
                "position": candidate_data.get("job_id", ""),
                "department": candidate_data.get("department", ""),
                "interview_time": candidate_data.get("interview_time", ""),
                "onboard_date": candidate_data.get("onboard_date", "入职日"),
            }
            subject = subject_tpl.format(**variables)
            body = (
                body_tpl.format(**variables)
                if body_tpl
                else self._get_default_body(status, variables)
            )

            # 只发给 HR 自己（配置的邮箱）
            hr_email = await get_module_setting("hr", "HR_MAIL_SMTP_USER", "")
            if hr_email:
                await send_email_with_template(
                    hr_email,  # 发给 HR 自己
                    subject,
                    body,
                    attachment_path=template_path if template_path else None,
                )
        else:  # 不符合
            subject_tpl = await get_module_setting(
                "hr", "HR_MAIL_REJECT_SUBJECT", "面试结果通知 - {name}"
            )
            body_tpl = await get_module_setting("hr", "HR_MAIL_REJECT_BODY", "")

            # 替换变量
            variables = {
                "name": name,
                "position": candidate_data.get("job_id", ""),
                "department": candidate_data.get("department", ""),
                "interview_time": candidate_data.get("interview_time", ""),
                "onboard_date": candidate_data.get("onboard_date", "入职日"),
            }
            subject = subject_tpl.format(**variables)
            body = (
                body_tpl.format(**variables)
                if body_tpl
                else self._get_default_body(status, variables)
            )

            # 拒绝邮件只发给 HR 自己
            hr_email = await get_module_setting("hr", "HR_MAIL_SMTP_USER", "")
            if hr_email:
                await send_email_with_template(hr_email, subject, body)

    def _get_default_body(self, status: str, variables: dict[str, Any]) -> str:
        """获取默认邮件正文"""
        if status == "通过":
            return f"""<!DOCTYPE html>
<html>
<body style="font-family: 'Microsoft YaHei', sans-serif;
max-width: 600px; margin: 0 auto;">
<h2 style="color: #1890ff;">录用通知书</h2>
<p>{variables["name"]}：</p>
<p>很高兴通知您，您已通过我司面试，被正式录用为 <strong>{
                variables[("position")]
            }</strong>。</p>
<p>入职部门：<strong>{variables["department"]}</strong></p>
<p>请于 <strong>{variables["onboard_date"]}</strong> 携带以下材料报到：</p>
<ul>
  <li>身份证原件及复印件</li>
  <li>学历证明</li>
  <li>离职证明（如有）</li>
  <li>银行卡</li>
</ul>
<p>如有疑问请联系 HR。</p>
<hr/>
<p style="color: #999; font-size: 12px;">此邮件由系统自动发送，请勿回复。</p>
</body>
</html>"""
        else:
            return f"""<!DOCTYPE html>
<html>
<body style="font-family: 'Microsoft YaHei', sans-serif;
max-width: 600px; margin: 0 auto;">
<h2 style="color: #ff4d4f;">面试结果通知</h2>
<p>{variables["name"]}：</p>
<p>感谢您参加我司 <strong>{variables["position"]}</strong> 职位的面试。</p>
<p>经过综合评估，很遗憾地通知您，您未能通过本次面试。</p>
<p>我们已将您的简历纳入人才库，未来如有合适职位将优先联系您。</p>
<p>祝您求职顺利！</p>
<hr/>
<p style="color: #999; font-size: 12px;">此邮件由系统自动发送，请勿回复。</p>
</body>
</html>"""

    async def delete_candidate(self, record_id: str) -> None:
        logger.info("soft-deleting candidate", extra={"record_id": record_id})
        await self.repo.soft_delete_candidate(record_id)

    async def analyze_resume(
        self, resume_text: str, job_requirements: str = ""
    ) -> dict[str, Any]:
        """Single resume AI analysis with LLM, with 3 retries.

        Uses a multi-step scoring methodology: hard filters → weighted scorecard → match
        reference.
        """
        has_job = bool(job_requirements and job_requirements.strip())
        system_prompt = (
            "你是资深招聘专家，负责筛选候选人简历。\n\n"
            "# 判断规则（最高优先级，必须遵守违者重罚）\n"
            "1. **所见即所得，禁止编造**：简历写了什么就是什么，简历没写的一律标「未提"
            "及」，严禁脑补任何信息\n"
            "2. 不因性别、年龄、婚育、籍贯、照片等因素影响评分\n"
            "3. 不因排版格式、信息完整度扣分\n"
            "4. 期望薪资仅做备注不参与评分\n"
            "5. **工作断档不扣分**：职场空窗期（离职休息、照顾家人、生病等）是正常人生"
            "阶段，可以标注但不扣分\n\n"
            "你只输出 JSON，不要输出其他任何内容。"
        )
        if has_job:
            user_prompt = (
                f"# 岗位信息卡\n"
                f"{job_requirements}\n\n"
                f"# 简历文本\n"
                f"{resume_text}\n\n"
                f"# 第一步：提炼筛选标准\n"
                f"从「任职要求 + 要求技能」中提炼：\n"
                f"- 硬性门槛：学历、年限、必备技能、证书等\n"
                f"- 加分项：按重要性排序\n\n"
                "# 第二步：按评分卡打分（每项 1–5 分，"
                "3=基本符合，5=完美匹配/超出预期）\n"
                f"- 技能匹配度（对照要求技能）——权重 30%\n"
                "- 经验相关性（对照岗位描述，逐段分析工作/项目/"
                "在校经历是否相关或相似）——权重 25%\n"
                f"- 项目深度与量化成果——权重 20%\n"
                f"- 职业稳定性与成长路径——权重 15%\n"
                f"- 学历与专业背景——权重 10%\n"
                f"评审铁律：不因信息缺失扣分、不因排版格式扣分、不因薪资期望扣分\n\n"
                f"match_rate 计算公式：sum((各维度得分/5) × 权重 × 100)\n"
                f"例：全部3分 → 60分；全部4分 → 80分\n\n"
                f"# 第三步：匹配参考（不扣分，只备注）\n"
                "- 期望薪资 vs 薪资范围：简历有提及则注明，"
                "未提及标「未提及」。注意：薪资期望不参与评分\n"
                "- 当前城市 vs 工作地点：同上，异地的注明"
                "「需确认 relocation 意愿」\n\n"
                f"# 返回 JSON（必须严格遵守）\n"
                f"{{\n"
                f'  "name": "候选人姓名",\n'
                f'  "phone": "手机号",\n'
                f'  "email": "邮箱",\n'
                f'  "education_level": "学历（大专/本科/硕士/博士/其他）",\n'
                f'  "education_school": "毕业院校",\n'
                f'  "work_years": 工作年数(整数),\n'
                f'  "skills": "技能列表（逗号分隔，只取前5个核心技能）",\n'
                f'  "last_company": "最近就职公司",\n'
                f'  "current_title": "当前职位",\n'
                f'  "match_rate": 加权综合得分(0-100整数，按上述公式计算),\n'
                '  "resume_score": 简历内容质量评分(0-100整数，'
                "基于项目深度、经验含金量、专业能力，不看信息完整度)\n"
                '  "fit_level": "推荐度（≥90=强烈推荐, 70-89=推荐, '
                '50-69=待定, <50=不推荐）",\n'
                f'  "reason": "结构化分析文本，按以下格式输出：\\n'
                "    1. 硬性门槛判断：列举每条门槛是否满足（引用原文），"
                "不满足时标注风险\\n"
                "    2. 各维度得分明细：技能匹配度X分 / 经验相关性X分 / "
                "项目深度X分 / 职业稳定性X分 / 学历专业X分\\n"
                f"    3. 亮点（引用原文）\\n"
                f"    4. 风险点与疑点（引用原文）\\n"
                f'    5. 薪资 / 地点匹配备注"\n'
                f"}}"
            )
        else:
            user_prompt = (
                f"# 简历文本\n"
                f"{resume_text}\n\n"
                f"# 评估要求\n"
                f"无岗位信息，请基于经验、技能、学历等维度进行综合质量评估。\n\n"
                f"# 返回 JSON（必须严格遵守）\n"
                f"{{\n"
                f'  "name": "候选人姓名",\n'
                f'  "phone": "手机号",\n'
                f'  "email": "邮箱",\n'
                f'  "education_level": "学历（大专/本科/硕士/博士/其他）",\n'
                f'  "education_school": "毕业院校",\n'
                f'  "work_years": 工作年数(整数),\n'
                f'  "skills": "技能列表（逗号分隔，只取前5个核心技能）",\n'
                f'  "last_company": "最近就职公司",\n'
                f'  "current_title": "当前职位",\n'
                f'  "match_rate": 综合评分(0-100整数),\n'
                '  "resume_score": 简历内容质量评分(0-100整数，'
                "基于项目深度、经验含金量、专业能力，不看完整度)\n"
                f'  "fit_level": "评级（高/中/低）",\n'
                f'  "reason": "结构化分析：\\n'
                f"    1. 硬性条件总结（学历/年限等）\\n"
                "    2. 各维度评分：技能X分 / 经验X分 / 项目深度X分 / "
                "稳定性X分 / 学历X分\\n"
                f"    3. 亮点\\n"
                f'    4. 风险点"\n'
                f"}}"
            )
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]
        expected_keys = [
            "name",
            "phone",
            "email",
            "education_level",
            "education_school",
            "work_years",
            "skills",
            "last_company",
            "current_title",
            "match_rate",
            "resume_score",
            "fit_level",
            "reason",
        ]

        last_error: Exception | None = None
        for attempt in range(3):
            try:
                result = await llm_client.chat_json(
                    messages=messages,
                    expected_keys=expected_keys,
                    temperature=0.1,
                )
                logger.info("resume analysis success", extra={"attempt": attempt + 1})
                return result
            except LLMRateLimitError as e:
                logger.warning(
                    "LLM rate limit on attempt %d, waiting longer", attempt + 1
                )
                last_error = e
                await asyncio.sleep(10 * (attempt + 1))  # 10s/20s/30s
                continue
            except LLMProviderError as e:
                logger.warning(
                    "LLM provider error on attempt %d", attempt + 1, exc_info=True
                )
                last_error = e
            except LLMOutputError as e:
                logger.warning(
                    "LLM output error on attempt %d", attempt + 1, exc_info=True
                )
                last_error = e

            if attempt < 2:
                await asyncio.sleep(2**attempt)

        logger.error("resume analysis failed after 3 attempts")
        raise last_error  # type: ignore[misc]

    async def _try_vision_analysis(
        self, file_bytes: bytes, filename: str, job_requirements: str
    ) -> dict[str, Any] | None:
        """Try vision-based analysis (PDF → images → vision model).
        Returns None if vision model doesn't support or fails, so caller can fall
        back."""

        # Only PDF supported for vision; DOCX falls back immediately
        if not filename.lower().endswith(".pdf"):
            return None

        if len(file_bytes) < 100:
            return None

        import base64

        try:
            # PyMuPDF's compatibility import has no mypy metadata in CI.
            import fitz  # type: ignore[import-not-found]  # PyMuPDF
        except ImportError:
            logger.warning("PDF vision analysis unavailable: PyMuPDF is not installed")
            return None

        try:
            # Convert PDF pages to images (max 4 pages to avoid token overflow)
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            page_count = min(len(doc), 4)
            images = []
            for i in range(page_count):
                page = doc[i]
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                images.append(base64.b64encode(img_bytes).decode())
            doc.close()

            if not images:
                return None

            logger.info(
                "vision: converted PDF to images",
                extra={"pages": page_count, "image_count": len(images)},
            )

            # Build multimodal message
            content_parts = []
            for img_b64 in images:
                content_parts.append(
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{img_b64}"},
                    }
                )
            content_parts.append(
                {
                    "type": "text",
                    "text": (
                        f"# 岗位信息卡\n{job_requirements}\n\n"
                        f"# 分析以上简历图片\n"
                        f"你是资深招聘专家，请三步分析：\n\n"
                        f"1. 提炼硬性门槛 + 加分项\n"
                        "2. 按评分卡打分（技能匹配30%/经验相关25%/"
                        "项目深度20%/职业稳定性15%/学历专业10%，每项1-5分）\n"
                        "   经验相关：逐段分析工作经历/项目经历/"
                        "在校项目是否与岗位相关或相似\n"
                        "   评审铁律：不因信息缺失扣分、不因排版扣分、"
                        "不因薪资期望扣分\n"
                        f"3. 薪资/地点仅做备注参考（不参与评分）\n\n"
                        f"返回 JSON：\n"
                        f'{{"name":"姓名","phone":"手机","email":"邮箱",'
                        f'"education_level":"学历","education_school":"学校",'
                        f'"work_years":年数(整数),"skills":"技能列表",'
                        f'"last_company":"最近公司","current_title":"当前职位",'
                        f'"match_rate":加权分(0-100),"resume_score":内容质量(0-100,不看完整度),'
                        f'"fit_level":"≥90强烈推荐/70-89推荐/50-69待定/<50不推荐",'
                        f'"reason":"结构化分析报告"}}'
                    ),
                }
            )

            vision_messages = [{"role": "user", "content": content_parts}]

            # Try vision model with thinking mode enabled
            raw_text = await llm_client.chat(
                vision_messages,
                config_type="vision",
                temperature=0.1,
                max_tokens=4096,
                response_format="json_object",
            )

            # Parse JSON from response
            import json as _json

            cleaned = raw_text.strip()
            if cleaned.startswith("```"):
                lines = cleaned.split("\n")
                for i, line in enumerate(lines):
                    if not line.strip().startswith("```"):
                        cleaned = "\n".join(lines[i:])
                        break
                if cleaned.endswith("```"):
                    cleaned = cleaned[:-3].strip()
            result = cast(dict[str, Any], _json.loads(cleaned))
            logger.info(
                "vision analysis success",
                extra={"match_rate": result.get("match_rate")},
            )
            return result

        except Exception as e:
            logger.warning(
                "vision analysis failed, will fall back to text",
                extra={"error": str(e)[:100]},
            )
            return None

    async def batch_analyze(
        self, candidate_ids: list[str] | None = None, job_id: str | None = None
    ) -> dict[str, Any]:
        """Run AI analysis on candidates.  If candidate_ids is empty/None, all
        unanalyzed
        candidates are processed automatically.  Returns progress summary."""

        def _build_job_req_text(job: dict[str, Any]) -> str:
            parts = []
            if job.get("title"):
                parts.append(f"- 职位名称：{job['title']}")
            if job.get("description"):
                parts.append(f"- 岗位描述：{job['description']}")
            if job.get("requirement"):
                parts.append(f"- 任职要求：{job['requirement']}")
            if job.get("req_skills"):
                sks = job["req_skills"]
                parts.append(
                    f"- 要求技能：{', '.join(sks) if isinstance(sks, list) else sks}"
                )
            if job.get("salary_range"):
                parts.append(f"- 薪资范围：{job['salary_range']}")
            if job.get("location"):
                parts.append(f"- 工作地点：{job['location']}")
            return "\n".join(parts)

        def _normalize_fit(raw: str) -> str:
            "Map LLM's recommendation level"
            " to Feishu's 招聘符合程度 (非常满足/高/中/"
            "低)."
            raw = raw.strip()
            if not raw:
                return "低"
            mapping = {
                "强烈推荐": "非常满足",
                "推荐": "高",
                "待定": "中",
                "不推荐": "低",
                "非常满足": "非常满足",
                "高": "高",
                "中": "中",
                "低": "低",
            }
            return mapping.get(raw, "低")

        async def _save_analysis(candidate_id: str, analysis: dict[str, Any]) -> None:
            """Write AI analysis result (scoring fields only) to bitable."""
            match_rate = analysis.get("match_rate")
            # 根据 match_rate 强制计算 fit_level，不依赖 LLM 返回
            if match_rate is not None:
                if match_rate >= 90:
                    fit_level = "非常满足"
                elif match_rate >= 70:
                    fit_level = "高"
                elif match_rate >= 50:
                    fit_level = "中"
                else:
                    fit_level = "低"
            else:
                fit_level = _normalize_fit(analysis.get("fit_level", ""))

            mapped = {
                "match_rate": match_rate,
                "resume_score": analysis.get("resume_score"),
                "fit_level": fit_level,
                "remark": analysis.get("reason", ""),
            }
            mapped = {
                k: v
                for k, v in mapped.items()
                if v not in (None, "", []) and v != " / "
            }
            await self.repo.update_candidate(candidate_id, mapped)

        def _already_analyzed(c: dict[str, Any]) -> bool:
            """A candidate is considered analyzed when match_rate is populated."""
            return c.get("match_rate") is not None and c.get("match_rate") != ""

        # ── Resolve candidate list ──────────────────────────────────
        if not candidate_ids:
            all_cands, _ = await self.repo.list_candidates()
            # Filter to only unanalyzed candidates
            target = [c for c in all_cands if not _already_analyzed(c)]
            candidate_ids = [c["id"] for c in target]
            logger.info(
                "batch: auto-selected unanalyzed candidates",
                extra={"total": len(all_cands), "unanalyzed": len(candidate_ids)},
            )
        else:
            logger.info(
                "batch: explicit candidate_ids", extra={"count": len(candidate_ids)}
            )

        if not candidate_ids:
            return {
                "total": 0,
                "success": 0,
                "skipped": 0,
                "failed": 0,
                "message": "没有需要分析的候选人",
            }

        # ── Import dependencies ─────────────────────────────────────
        from io import BytesIO

        from docx import Document
        from pypdf import PdfReader

        # Pre-fetch global job if provided
        global_job_req = ""
        if job_id:
            try:
                job = await self.repo.get_job(job_id)
                global_job_req = _build_job_req_text(_untranslate(job, JOB_FIELD_MAP))
            except Exception:
                logger.warning("failed to fetch global job", extra={"job_id": job_id})

        success = 0
        skipped = 0
        failed = 0
        total = len(candidate_ids)

        for idx, cid in enumerate(candidate_ids):
            try:
                # Get raw feishu record directly to access 简历附件 field
                bitable = await self.repo._get_client()
                if bitable is None:
                    raise RuntimeError("飞书招聘数据源未配置")
                resp = await bitable.client.request(
                    "GET",
                    bitable._path(TBL_CANDIDATE, f"/records/{cid}"),
                )
                raw = resp.get("record", {})
                fields_data = raw.get("fields", {})
                c = _untranslate(raw, CANDIDATE_FIELD_MAP)

                # ── 1. Resume: prefer vision analysis, fallback to text extraction ──
                resume_text = ""
                file_bytes = None
                filename_raw = ""

                attachments = fields_data.get("简历附件")
                if (
                    attachments
                    and isinstance(attachments, list)
                    and len(attachments) > 0
                ):
                    file_token = attachments[0].get("file_token")
                    filename_raw = attachments[0].get("name", "")
                    if file_token:
                        try:
                            file_bytes = await bitable.client.download_file(file_token)
                            logger.info(
                                "batch: downloaded file",
                                extra={"cid": cid, "bytes": len(file_bytes)},
                            )
                        except Exception:
                            logger.warning(
                                "batch: failed to download file", extra={"cid": cid}
                            )

                # Try vision analysis first (PDFs only with image-based model)
                if file_bytes:
                    try:
                        # Resolve job requirements for this candidate
                        vision_job_req = global_job_req
                        if not vision_job_req and c.get("job_id"):
                            try:
                                job = await self.repo.get_job(c["job_id"])
                                vision_job_req = _build_job_req_text(
                                    _untranslate(job, JOB_FIELD_MAP)
                                )
                            except Exception:
                                pass
                        vision_result = await self._try_vision_analysis(
                            file_bytes, filename_raw, vision_job_req or ""
                        )
                        if vision_result:
                            analysis = vision_result
                            await _save_analysis(cid, analysis)
                            success += 1
                            logger.info(
                                "batch: vision analysis saved",
                                extra={
                                    "cid": cid,
                                    "match_rate": analysis.get("match_rate"),
                                    "progress": f"{idx + 1}/{total}",
                                },
                            )
                            continue  # Vision success, skip text path
                    except Exception:
                        logger.warning(
                            "batch: vision path exception, falling back to text",
                            extra={"cid": cid},
                        )

                # ── Text extraction (fallback) ─────────────────────
                if file_bytes and not resume_text.strip():
                    try:
                        filename_lower = filename_raw.lower()
                        if filename_lower.endswith(".docx"):
                            from io import BytesIO as DocBytesIO

                            doc = Document(DocBytesIO(file_bytes))
                            parts = []
                            for p in doc.paragraphs:
                                if p.text.strip():
                                    parts.append(p.text.strip())
                            for table in doc.tables:
                                for row in table.rows:
                                    cells = [
                                        cell.text.strip()
                                        for cell in row.cells
                                        if cell.text.strip()
                                    ]
                                    if cells:
                                        parts.append(" | ".join(cells))
                            resume_text = "\n".join(parts)
                            logger.info("batch: parsed DOCX", extra={"cid": cid})
                        else:
                            reader = PdfReader(BytesIO(file_bytes))
                            text_parts = []
                            for page in reader.pages:
                                t = page.extract_text()
                                if t:
                                    text_parts.append(t)
                                try:
                                    extract_tables = getattr(
                                        page, "extract_tables", None
                                    )
                                    if callable(extract_tables):
                                        tables = extract_tables()
                                        for tbl in tables or []:
                                            for row in tbl:
                                                cells = [c or "" for c in row if c]
                                                if cells:
                                                    text_parts.append(" | ".join(cells))
                                except Exception as e:
                                    logger.warning("PDF表格提取失败: %s", e)
                            resume_text = "\n".join(text_parts)
                            logger.info("batch: parsed PDF", extra={"cid": cid})
                    except Exception:
                        logger.warning(
                            "batch: failed to parse resume file, using fields",
                            extra={"cid": cid, "file": filename_raw},
                        )

                if not resume_text.strip():
                    # Fallback: build from text fields
                    parts = []
                    if c.get("name"):
                        parts.append(f"姓名：{c['name']}")
                    if c.get("contact"):
                        parts.append(f"联系方式：{c['contact']}")
                    if c.get("education"):
                        parts.append(f"学历：{c['education']}")
                    if c.get("work_years"):
                        parts.append(f"工作经验：{c['work_years']}年")
                    if c.get("skills"):
                        parts.append(f"技能：{c['skills']}")
                    resume_text = "\n".join(parts)

                if not resume_text.strip():
                    logger.warning(
                        "batch: no resume content", extra={"candidate_id": cid}
                    )
                    skipped += 1
                    continue

                # ── 2. Job requirements ──────────────────────────────
                job_req_text = global_job_req
                if not job_req_text and c.get("job_id"):
                    try:
                        job = await self.repo.get_job(c["job_id"])
                        job_req_text = _build_job_req_text(
                            _untranslate(job, JOB_FIELD_MAP)
                        )
                    except Exception:
                        logger.warning(
                            "batch: failed to fetch linked job",
                            extra={"job_id": c["job_id"]},
                        )

                # ── 3. AI analysis ───────────────────────────────────
                analysis = await self.analyze_resume(resume_text, job_req_text)
                await _save_analysis(cid, analysis)
                success += 1
                logger.info(
                    "batch: analyzed",
                    extra={"candidate_id": cid, "idx": f"{idx + 1}/{total}"},
                )

            except Exception:
                logger.exception("batch: failed", extra={"candidate_id": cid})
                failed += 1

        summary = {
            "total": total,
            "success": success,
            "skipped": skipped,
            "failed": failed,
        }
        logger.info("batch analysis finished", extra=summary)
        return summary

    async def analyze_and_create(
        self, resume_text: str, source_channel: str, file_hash: str
    ) -> dict[str, Any]:
        """AI分析简历文本并创建候选人记录。"""
        from app.core.llm import (
            LLMOutputError,
            LLMProviderError,
            LLMRateLimitError,
            llm_client,
        )

        # AI analysis with retry
        result = None
        for attempt in range(3):
            try:
                result = await llm_client.chat_json(
                    messages=[
                        {
                            "role": "system",
                            ("content"): (
                                "你是一个简历解析助手。请从简历文本中抽取结构化字段并评估。只"
                                "输出JSON。"
                            ),
                        },
                        {
                            "role": "user",
                            "content": (
                                f"简历文本：\n{resume_text}\n\n请提取："
                                "name,phone,email,education_level(学历),"
                                "education_school(毕业院校),work_years(工作经验年数),"
                                "skills(技能列表),last_company(最近公司),"
                                "current_title(当前职位),match_rate(综合评分0-100),"
                                "resume_score(内容质量0-100,不看完整度),"
                                "fit_level(符合程度:高/中/低),reason(评分理由)"
                            ),
                        },
                    ],
                    expected_keys=[
                        "name",
                        "phone",
                        "email",
                        "education_level",
                        "education_school",
                        "work_years",
                        "skills",
                        "last_company",
                        "current_title",
                        "match_rate",
                        "resume_score",
                        "fit_level",
                        "reason",
                    ],
                    temperature=0.1,
                )
                break
            except (LLMOutputError, LLMProviderError) as e:
                if attempt == 2:
                    raise
                logger.warning("LLM retry %d: %s", attempt + 1, e)
                await asyncio.sleep(2**attempt)
            except LLMRateLimitError as e:
                logger.warning("LLM rate limit, waiting: %s", e)
                await asyncio.sleep(5)

        # Build candidate fields
        fields: dict[str, Any]
        if result is None:
            fields = {"interview_status": "待安排", "source_channel": source_channel}
        else:
            skills = result.get("skills", [])
            if isinstance(skills, str):
                skills = [s.strip() for s in skills.split(",")]

            fields = {
                "name": result.get("name", ""),
                "contact": f"{result.get('phone', '')} / {result.get('email', '')}",
                "email": result.get("email", ""),
                "education": result.get("education_level", ""),
                "work_years": result.get("work_years"),
                "skills": skills[:5] if skills else [],
                "match_rate": result.get("match_rate"),
                "resume_score": result.get("resume_score"),
                "fit_level": result.get("fit_level", "低"),
                "interview_status": "待安排",
                "source_channel": source_channel,
                "remark": result.get("reason", ""),
            }
        # Filter out None/empty
        fields = {
            k: v for k, v in fields.items() if v not in (None, "", [], 0) and v != " / "
        }

        # Create candidate via repo
        record = await self.repo.create_candidate(fields)
        logger.info(
            "candidate created from upload",
            extra={
                "name": result.get("name") if result else "",
                "source": source_channel,
            },
        )

        return {"candidate": record, "ai_result": result}


# ─── OnboardingService ───


class OnboardingService:
    def __init__(self, app_token: str | None = None) -> None:
        self.repo = RecruitmentBitableRepo(app_token=app_token)

    async def create_from_interview(self, candidate_id: str) -> dict[str, Any]:
        """Create onboarding record from an interview-passed candidate."""
        from datetime import date as _date

        logger.info(
            "creating onboarding from interview", extra={"candidate_id": candidate_id}
        )

        # 用 repo.list_candidates 关联查出 job_position
        raw_list, _ = await self.repo.list_candidates(page=1, page_size=500)
        candidate = None
        for raw in raw_list:
            if raw.get("id") == candidate_id:
                candidate = _untranslate(raw, CANDIDATE_FIELD_MAP)
                break
        if not candidate:
            raw = await self.repo.get_candidate(candidate_id)
            candidate = _untranslate(raw, CANDIDATE_FIELD_MAP)

        # 关联查出职位名称
        jid = candidate.get("job_id")
        if jid and not candidate.get("job_position"):
            job_names = await self.repo.get_job_names()
            if jid in job_names:
                candidate["job_position"] = job_names[jid]

        onboarding_data = {
            "name": candidate.get("name", ""),
            "onboard_date": _date.today().strftime("%Y-%m-%d"),
            "status": "进行中",
            "level": candidate.get("job_position", "") or candidate.get("job_id", ""),
            "department": candidate.get("department", ""),
            "health_status": "未进行",
            "resignation_cert": "未提供",
            "id_card": "未提供",
            "education_cert": "未提供",
        }
        onboarding = await self.repo.create_onboarding(onboarding_data)

        logger.info(
            "onboarding created from interview",
            extra={"onboarding_id": onboarding.get("id", "")},
        )
        return onboarding

    async def list_onboarding(
        self,
        keyword: str | None = None,
        page: int = 1,
        page_size: int = 20,
        dept_alias_set: set[str] | None = None,
    ) -> tuple[list[dict[str, Any]], int]:
        logger.info(
            "listing onboarding records", extra={"keyword": keyword, "page": page}
        )
        result = await self.repo.list_onboarding(
            keyword=keyword,
            page=page,
            page_size=page_size,
            dept_alias_set=dept_alias_set,
        )
        if isinstance(result, dict):
            items, total = result.get("items", []), result.get("total", 0)
        else:
            items = result[0] if isinstance(result, (list, tuple)) else []
            total = (
                result[1]
                if isinstance(result, (list, tuple)) and len(result) > 1
                else len(items)
            )
        return items, total

    async def get_onboarding(self, record_id: str) -> dict[str, Any]:
        """Get onboarding detail."""
        logger.info("getting onboarding detail", extra={"record_id": record_id})
        return await self.repo.get_onboarding(record_id)

    async def update_onboarding(
        self, record_id: str, data: dict[str, Any]
    ) -> dict[str, Any]:
        logger.info("updating onboarding", extra={"record_id": record_id})
        result = await self.repo.update_onboarding(record_id, data)
        return result

    async def get_dashboard(
        self, dept_alias_set: set[str] | None = None
    ) -> dict[str, Any]:
        logger.info("getting recruitment dashboard stats")
        return await self.repo.get_dashboard_stats(dept_alias_set=dept_alias_set)
