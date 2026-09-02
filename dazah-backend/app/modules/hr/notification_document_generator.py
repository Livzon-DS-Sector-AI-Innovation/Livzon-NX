"""培训通知 Word 文档生成器（桌面版 培训通知.docx 模板保真填充）.

仅填充填写区（空白 run 段落 / 段内尾随空白 run），保留模板字体、字号、行距、
对齐。模板段落结构（实测）：
  P2 培训内容 | P4 培训对象(段首空白 run) | P6 培训时间 | P8 培训地点
  P14 六、培训考核(段尾空白 run) | P18 落款单位 | P19 落款日期(年/月/日)
"""

from __future__ import annotations

import logging
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from app.modules.hr.date_format import HR_EXPORT_DATE_FORMAT
from app.modules.hr.schemas import TrainingNotificationInput
from app.modules.hr.template_filler import (
    _apply_std_font,
    set_paragraph_value,
)

logger = logging.getLogger(__name__)


def _find_template() -> Path:
    """Locate the docx template, trying several path candidates."""
    candidates = [
        Path("员工培训教育管理规程/培训通知.docx"),
        Path("../员工培训教育管理规程/培训通知.docx"),
        Path(__file__).resolve().parent.parent.parent.parent
        / "员工培训教育管理规程"
        / "培训通知.docx",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("模板文件未找到: 培训通知.docx")


def _set_right_align(paragraph: Any) -> None:
    """落款段右对齐（部门在日期上方，右侧对齐成块）。"""
    ppr = paragraph._p.get_or_add_pPr()
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = ppr.makeelement(qn("w:jc"), {})
        ppr.append(jc)
    jc.set(qn("w:val"), "right")


def _set_fill_indent(paragraph: Any) -> None:
    """填写段左缩进2字符（leftChars=266），与培训内容/对象/时间填写区对齐。

    模板中"培训地点"填写段无缩进，填值会顶到左边距，需右移对齐其他填写区。
    """
    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = ppr.makeelement(qn("w:ind"), {})
        ppr.append(ind)
    ind.set(qn("w:left"), "559")
    ind.set(qn("w:leftChars"), "266")
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:firstLineChars"), "0")


def _locate(paras: Any, label: str, fallback_idx: int) -> Any:
    """按段落文本定位标签段及其填写段（标签下一段），模板段落增删时不越界。"""
    for i, p in enumerate(paras):
        if p.text.strip().startswith(label):
            return i, i + 1
    return fallback_idx, fallback_idx + 1


def _insert_training_level(paras: Any, level: str | None) -> None:
    """五、培训要求 第1条："1.以上课程属于培训课程的重要内容"
    → 在"属于"后
    插入培训级别，成为"属于公司级培训…"/"属于部门级培训…"。

    模板 run 结构（实测）：run0="1.以上课程属于" run1="培训课程的重要内容，要求"…
    级别为空时不改动（保持模板原文）。
    """
    lv = (level or "").strip()
    if lv not in ("公司级", "部门级"):
        return
    for p in paras:
        if "以上课程属于" not in p.text:
            continue
        runs = p.runs
        for i, r in enumerate(runs):
            if (
                r.text.startswith("培训")
                and i > 0
                and runs[i - 1].text.endswith("属于")
            ):
                r.text = lv + r.text
                _apply_std_font(r, SZ_XIAOSAN)
                return


# 培训通知模板正文为小三（15pt=sz30），填写值需与正文同号
SZ_XIAOSAN = "30"


def generate_training_notification(data: TrainingNotificationInput) -> BytesIO:
    """按桌面版 培训通知 模板保真填充生成 Word 文档.

    格式要求：填写内容不加粗、黄色底纹区统一小三字(15pt)、
    地点填值与其他填写区对齐右移、落款单位位于日期上方、
    五、培训要求第1条按培训级别插入"公司级/部门级"。
    """
    template_path = _find_template()
    doc = Document(str(template_path))
    paras = doc.paragraphs

    # 一、培训内容：填写段所有 run 统一小三、不加粗
    _, idx = _locate(paras, "一、培训内容", 1)
    set_paragraph_value(paras[idx], data.subject or "", SZ_XIAOSAN)
    for r in paras[idx].runs:
        _apply_std_font(r, SZ_XIAOSAN)

    # 二、培训对象：段首空白 run 填入，保留"（具体人员名单详见培训签到表）"，
    # 整段 run 统一小三、不加粗
    _, idx = _locate(paras, "二、培训对象", 3)
    people_str = "、".join(data.trainee_names) if data.trainee_names else ""
    if paras[idx].runs:
        paras[idx].runs[0].text = people_str
        for r in paras[idx].runs:
            _apply_std_font(r, SZ_XIAOSAN)

    # 三、培训时间：日期 + 括号包时间（如 2026.08.26 （16:00 ~ 18:00））
    _, idx = _locate(paras, "三、培训时间", 5)
    time_parts = [x for x in (data.training_time_start, data.training_time_end) if x]
    time_str = " ~ ".join(time_parts)
    if data.training_date:
        date_str = data.training_date.strftime(HR_EXPORT_DATE_FORMAT)
        time_str = f"{date_str} （{time_str}）" if time_str else date_str
    set_paragraph_value(paras[idx], time_str, SZ_XIAOSAN)
    for r in paras[idx].runs:
        _apply_std_font(r, SZ_XIAOSAN)

    # 四、培训地点：填值右移（补2字符左缩进，与其余填写区对齐）
    _, idx = _locate(paras, "培训地点", 7)
    set_paragraph_value(paras[idx], data.location or "", SZ_XIAOSAN)
    for r in paras[idx].runs:
        _apply_std_font(r, SZ_XIAOSAN)
    _set_fill_indent(paras[idx])

    # 五、培训要求 第1条：在"属于"与"培训"之间插入培训级别
    # （公司级开展→"属于公司级培训…"，部门级开展→"属于部门级培训…"）
    _insert_training_level(paras, data.training_level)

    # 六、培训考核：考核方式写入标签段内最后一个空白 run（段尾填写区），
    # 不能用 runs[0]/runs[-1] 盲取（模板标签段首 run 为"六、培训考核："标签）
    idx, _ = _locate(paras, "六、培训考核", 14)
    if data.assessment_method:
        blanks = [r for r in paras[idx].runs if r.text.strip() == ""]
        if blanks:
            blanks[-1].text = data.assessment_method
            _apply_std_font(blanks[-1], SZ_XIAOSAN)
        else:
            new_run = paras[idx].add_run(data.assessment_method)
            _apply_std_font(new_run, SZ_XIAOSAN)

    # 落款单位 + 落款日期：单位段在前、日期段在后（单位在日期上方），均右对齐
    issuer = data.issuer_department or data.department or ""
    issue_date = data.issue_date or data.training_date
    unit_idx = 18
    date_idx = 19
    for i, p in enumerate(paras):
        if p.text.strip().startswith("特此通知"):
            # 特此通知后：空行、单位段、日期段
            unit_idx = i + 2
            date_idx = i + 3
            break
    set_paragraph_value(paras[unit_idx], issuer, SZ_XIAOSAN)
    for r in paras[unit_idx].runs:
        _apply_std_font(r, SZ_XIAOSAN)
    # 清除模板单位段的首行缩进，改右对齐，与日期块右侧成组
    ppr = paras[unit_idx]._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is not None:
        ppr.remove(ind)
    _set_right_align(paras[unit_idx])

    # 落款日期：直接写入 YYYY.MM.DD 格式
    set_paragraph_value(
        paras[date_idx], issue_date.strftime(HR_EXPORT_DATE_FORMAT), SZ_XIAOSAN
    )
    for r in paras[date_idx].runs:
        _apply_std_font(r, SZ_XIAOSAN)
    _set_right_align(paras[date_idx])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
