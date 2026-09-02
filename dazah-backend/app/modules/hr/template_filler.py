"""模板保真填充工具。

培训资料导出（APP3/APP4/APP10/APP13/培训通知）要求生成文档与 SMP-HR-002-14
模板格式完全一致：字体（Times New Roman/宋体）、字号（10.5pt）、行距、对齐、
表格边框逐格相同，仅"填写区"写入内容。

因此本模块只提供 run 级填充原语：修改已有 run 的文本、或克隆原 run 的 rPr
追加新 run、或改写 w:sym 字符码。**禁止**使用 ``cell.text=`` /
``paragraph.text=``（会清空模板 run 格式并以默认格式重建）。

填写区形态（实测模板 XML）：
- 标签后空白 run：``培训日期``(加粗) + ``：`` + ``   `` ← 值写入空白 run
- 整格可编辑：数据行姓名/部门等，首 run 写值
- 复选框：``w:sym`` Wingdings 2，00A3=□ / 0052=☑，sym run 后紧跟选项文本 run
- 文本复选框（APP10）：``合格□不合格□`` 中的字面 ``□`` 替换为 ``☑``
"""

from __future__ import annotations

import logging
from copy import deepcopy
from typing import Any

from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# Wingdings 2 复选框字符码（模板实测：0052=☑ 选中样例，00A3=□ 未选）
SYM_CHECKED = "0052"
SYM_UNCHECKED = "00A3"


def _clone_rpr_to(source_run: Any, target_run: Any) -> None:
    """把 source_run 的 rPr 深拷贝到 target_run（保留字体/字号/加粗）。"""
    if source_run is None:
        return
    rpr = source_run._r.find(qn("w:rPr"))
    if rpr is not None:
        old = target_run._r.find(qn("w:rPr"))
        if old is not None:
            target_run._r.remove(old)
        target_run._r.insert(0, deepcopy(rpr))


def _append_clone_run(paragraph: Any, source_run: Any, text: str) -> Any:
    """在段落末尾追加新 run，克隆 source_run 的格式，并统一为标准填值字体。"""
    new_run = paragraph.add_run(text)
    _clone_rpr_to(source_run, new_run)
    _apply_std_font(new_run)
    return new_run


def _apply_std_font(run: Any, size_half_points: str = "21") -> None:
    """统一填值字体：中文宋体、数字/英文 Times New Roman、默认5号(10.5pt=sz21)、不加粗。

    填写的值（姓名/部门/日期/题目/授课人/说明等）统一用模板正文字体，
    避免克隆到标题行（如 APP10 标题 sz=36）导致字号异常。
    模板填写区 run 常继承标签的加粗格式（如培训通知填写段 w:b），
    填写内容为正文，必须显式去除加粗。

    size_half_points 为半磅值：5号=21、小三=30。培训通知等正文为小三的
    模板传 "30"，使填写值与模板正文字号一致。
    """
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:cs"), "Times New Roman")
    for tag in ("w:sz", "w:szCs"):
        el = rpr.find(qn(tag))
        if el is None:
            el = rpr.makeelement(qn(tag), {})
            rpr.append(el)
        el.set(qn("w:val"), size_half_points)
    # 去除加粗：填写内容为正文样式，不继承模板标签/标题的粗体
    for tag in ("w:b", "w:bCs"):
        el = rpr.find(qn(tag))
        if el is None:
            el = rpr.makeelement(qn(tag), {})
            rpr.append(el)
        el.set(qn("w:val"), "0")


def _is_blank_run(run: Any) -> bool:
    """是否为可填写的空白 run（文本为空/空白，且本身不是复选框 w:sym）。"""
    return run.text.strip() == "" and run._r.find(qn("w:sym")) is None


def _clear_paragraph_except(paragraph: Any, keep_run: Any = None) -> None:
    keep_element = keep_run._r if keep_run is not None else None
    for r in paragraph.runs:
        if r._r is not keep_element:
            r.text = ""


def fill_after_label(cell: Any, value: Any) -> None:
    """标签+填写区同格（如 ``培训日期：____``）：保留标签 run，值写入标签后
    第一个空白/空 run，其余空白 run 清空；无空白 run 时克隆末 run 格式追加。"""
    if value is None:
        value = ""
    value = str(value)
    para = cell.paragraphs[0]
    runs = para.runs
    start = 0
    for i, r in enumerate(runs):
        if r.text.rstrip().endswith(("：", ":")):
            start = i + 1
            break
    targets = [r for r in runs[start:] if _is_blank_run(r)]
    if targets:
        targets[0].text = value
        _apply_std_font(targets[0])
        for r in targets[1:]:
            r.text = ""
    elif runs:
        _append_clone_run(para, runs[-1], value)
    else:
        para.add_run(value)
    for p in cell.paragraphs[1:]:
        _clear_paragraph_except(p)


def fill_after_phrase(
    cell: Any, phrase: str, value: Any, center: bool = False
) -> None:
    """在匹配 ``phrase`` 的 run 之后的空白 run 写入值（一格内多处填写/跨段落的场景）。

    如 APP3 ``实际受训人数合计：___人``、APP4 ``应到：__人 实到：__人``。
    按段落逐个查找短语所在 run，写入其后的空白 run；无空白 run 时克隆该 run
    格式追加。**未找到短语时不做任何修改**（避免破坏单元格内其他勾选/文本）。

    center=True 时把值写入短语后连续空白 run 的中间一个，使填写值在空白区
    居中显示（如 APP3 应受训人数），而不是紧贴短语左侧。
    """
    if value is None:
        value = ""
    value = str(value)
    for para in cell.paragraphs:
        runs = para.runs
        for i, r in enumerate(runs):
            if phrase in r.text:
                # 收集该 run 之后第一组连续空白 run（先跳过"："等非空白，
                # 遇到"人"等非空白后缀即止）
                blanks: list[int] = []
                started = False
                for j in range(i + 1, len(runs)):
                    if _is_blank_run(runs[j]):
                        started = True
                        blanks.append(j)
                    elif started:
                        break
                if blanks:
                    target = blanks[len(blanks) // 2] if center else blanks[0]
                    runs[target].text = value
                    _apply_std_font(runs[target])
                    return
                # 无空白 run：克隆当前 run 格式追加
                _append_clone_run(para, r, value)
                return
    # 未找到短语：不修改


def fill_whole_cell(cell: Any, value: Any, fmt_source: Any = None) -> None:
    """整格可编辑（数据行姓名/部门/描述等）：首 run 写值、其余清空；
    无 run 时克隆 fmt_source（默认表格表头 run）格式新建。"""
    if value is None:
        value = ""
    value = str(value)
    para = cell.paragraphs[0]
    runs = para.runs
    if runs and runs[0]._r.find(qn("w:sym")) is None:
        runs[0].text = value
        _apply_std_font(runs[0])
        for r in runs[1:]:
            r.text = ""
    else:
        # 首 run 为复选框或无 run：克隆格式追加，避免破坏 w:sym
        _append_clone_run(para, runs[-1] if runs else fmt_source, value)
    for p in cell.paragraphs[1:]:
        _clear_paragraph_except(p)


def rewrite_cell_runs(cell: Any, text: str) -> None:
    """整格文本重写但保留首个 run 的格式（如 ``应受训人数：N 人 …`` 合成串）。"""
    para = cell.paragraphs[0]
    runs = para.runs
    if runs:
        runs[0].text = text
        _apply_std_font(runs[0])
        for r in runs[1:]:
            r.text = ""
    else:
        para.add_run(text)
    for p in cell.paragraphs[1:]:
        _clear_paragraph_except(p)


def append_value(cell: Any, value: Any) -> None:
    """在格末追加值 run，克隆本格末位的 run 格式（用于标签+值同格、值在标签后的场景）。

    如 APP3 ``培训时间`` / ``培训题目或内容概要`` / ``授课人`` 单元格。
    """
    if value is None:
        value = ""
    value = str(value)
    para = cell.paragraphs[0]
    runs = para.runs
    if runs:
        _append_clone_run(para, runs[-1], value)
    else:
        para.add_run(value)


def set_sym_group(
    cell: Any,
    selected: str | None,
    options: list[str],
    *,
    write_trailing: dict[str, str] | None = None,
) -> None:
    """按选项勾选 w:sym 复选框组。

    sym run 后紧跟的文本 run 以某选项名开头即归属该选项；选中写
    SYM_CHECKED，其余 SYM_UNCHECKED。
    write_trailing: 选中某选项时把附加文本写入该 sym 组后的空白 run
    （如 APP3 ``其他：`` 后填写说明）。
    """
    selected = selected or ""
    for para in cell.paragraphs:
        runs = para.runs
        for i, r in enumerate(runs):
            sym = r._r.find(qn("w:sym"))
            if sym is None:
                continue
            nxt = runs[i + 1].text.strip() if i + 1 < len(runs) else ""
            matched = next((o for o in options if o and nxt.startswith(o)), None)
            if matched is None:
                continue
            is_on = selected.startswith(matched)
            sym.set(qn("w:char"), SYM_CHECKED if is_on else SYM_UNCHECKED)
            if is_on and write_trailing and matched in write_trailing:
                extra = write_trailing[matched]
                for after in runs[i + 2 :]:
                    if after.text.strip() == "":
                        after.text = extra
                        _apply_std_font(after)
                        break


def replace_text_in_cell(cell: Any, old: str, new: str) -> None:
    """run 级文本替换（保留 rPr），用于 APP10 ``合格□不合格□`` 勾选。"""
    for para in cell.paragraphs:
        for r in para.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)


def set_paragraph_value(
    paragraph: Any, value: Any, size_half_points: str = "21"
) -> None:
    """独立空格段填写（培训通知的内容/时间/地点/落款单位段）。"""
    if value is None:
        value = ""
    runs = paragraph.runs
    if runs:
        runs[0].text = str(value)
        _apply_std_font(runs[0], size_half_points)
        for r in runs[1:]:
            r.text = ""
    else:
        new_run = paragraph.add_run(str(value))
        _apply_std_font(new_run, size_half_points)


def fill_date_paragraph(paragraph: Any, year: Any, month: Any, day: Any) -> None:
    """``    年     月     日`` 段：把 年/月/日 前的空白 run 写入数字。"""
    runs = paragraph.runs
    filled = False
    for i, r in enumerate(runs):
        t = r.text.strip()
        if t in ("年", "月", "日") and i > 0:
            prev = runs[i - 1]
            if prev.text.strip() == "":
                prev.text = str({"年": year, "月": month, "日": day}[t])
                _apply_std_font(prev)
                filled = True
    if not filled:
        set_paragraph_value(paragraph, f"{year}年{month}月{day}日")


def header_fmt_source(table: Any) -> Any:
    """取表格表头行首个非空 run 作为格式来源（数据行空格新建 run 时克隆）。"""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    if r.text.strip():
                        return r
    return None


def clone_row_after(table: Any, row_idx: int) -> None:
    """在指定行后深拷贝一行（保留行格式），APP10/APP13 动态增行用。"""
    tr = table.rows[row_idx]._tr
    new_tr = deepcopy(tr)
    tr.addnext(new_tr)


def delete_row(table: Any, row_idx: int) -> None:
    """整行删除（含其格式），用于移除模板多余预置行。"""
    tr = table.rows[row_idx]._tr
    tr.getparent().remove(tr)
