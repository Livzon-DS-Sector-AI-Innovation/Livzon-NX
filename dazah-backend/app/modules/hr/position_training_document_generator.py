"""岗位培训清单 Word 文档生成器（APP9模板）.

严格按照 APP9-SMP-HR-002-14岗位培训清单.docx 模板格式导出：
- 保留模板的黄色可编辑区域标记（w:permStart/permEnd），填充内容写在标记内
- 级别标题行: 灰色背景 shade=D7D7D7, 11pt, 加粗, 居中
- 表头/数据行: 10.5pt, 居中, 宋体 + Times New Roman
- 有数据的级别区块自动删除"……"行；无数据时保留模板空行+"……"
- 数据行数动态扩展（克隆模板数据行）
- 页眉页码: 模板中硬编码的总页数"1"替换为 NUMPAGES 域，并强制打开时更新（updateFields）

格式保真: 仅对模板 XML 做外科手术式最小修改（表格行/页眉域/settings），
其余 part（样式、字体、主题、页脚、图片）原样保留，确保导出与模板格式一致。
不使用 Word COM 回存（Word 会物化 first/even 页眉页脚等结构，导致格式漂移）。
"""

import logging
import zipfile
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _find_template() -> Path:
    candidates = [
        Path("员工培训教育管理规程/APP9岗位培训清单.docx"),
        Path("../员工培训教育管理规程/APP9岗位培训清单.docx"),
        Path(__file__).resolve().parent.parent.parent.parent
        / "员工培训教育管理规程"
        / "APP9岗位培训清单.docx",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("模板文件未找到: APP9岗位培训清单.docx")


def _make_run(text: str, size_pt: float = 10.5, bold: bool = False) -> Any:
    """Create a run element with 宋体 + Times New Roman fonts."""
    r = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), "Times New Roman")
    run_fonts.set(qn("w:hAnsi"), "Times New Roman")
    run_fonts.set(qn("w:eastAsia"), "宋体")
    run_fonts.set(qn("w:cs"), "Times New Roman")
    run_properties.append(run_fonts)
    if bold:
        b = OxmlElement("w:b")
        b.set(qn("w:val"), "0")
        run_properties.append(b)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    run_properties.append(sz)
    complex_size = OxmlElement("w:szCs")
    complex_size.set(qn("w:val"), str(int(size_pt * 2)))
    run_properties.append(complex_size)
    r.append(run_properties)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _fill_cell(
    cell: Any,
    text: str,
    size_pt: float = 10.5,
    bold: bool = False,
    shade: str | None = None,
) -> Any:
    """Fill a table cell, preserving yellow editable markers (permStart/permEnd)."""
    p = cell.paragraphs[0]
    p_el = p._p
    # Remove existing runs but keep permStart/permEnd markers
    for child in list(p_el):
        tag = child.tag.split("}")[-1]
        if tag == "r":
            p_el.remove(child)
    # Insert new run after permStart if present, else at end after pPr
    run = _make_run(text, size_pt, bold)
    perm_start = p_el.find(qn("w:permStart"))
    paragraph_properties = p_el.find(qn("w:pPr"))
    if perm_start is not None:
        perm_start.addnext(run)
    elif paragraph_properties is not None:
        paragraph_properties.addnext(run)
    else:
        p_el.insert(0, run)

    # Background shade
    cell_properties = cell._tc.get_or_add_tcPr()
    shd_el = cell_properties.find(qn("w:shd"))
    if shade is None:
        if shd_el is not None:
            cell_properties.remove(shd_el)
    else:
        if shd_el is None:
            shd_el = OxmlElement("w:shd")
            cell_properties.append(shd_el)
        shd_el.set(qn("w:val"), "clear")
        shd_el.set(qn("w:fill"), shade)


def _fill_dept_pos_paragraph(doc: Any, department: str, position: str) -> Any:
    """Fill 部门/岗位 values inside the yellow editable regions (permStart..permEnd)."""
    for para in doc.paragraphs:
        runs_text = "".join(r.text for r in para.runs)
        if "部门" in runs_text and "岗位" in runs_text:
            p_el = para._p
            # Collect perm regions: between permStart and permEnd
            regions: list[list[Any]] = []  # list of lists of elements
            current: list[Any] | None = None
            for child in list(p_el):
                tag = child.tag.split("}")[-1]
                if tag == "permStart":
                    current = []
                elif tag == "permEnd":
                    if current is not None:
                        regions.append(current)
                    current = None
                elif current is not None:
                    current.append(child)

            values = [department or "", position or ""]
            for idx, region in enumerate(regions[:2]):
                # Remove old runs in region
                for el in region:
                    if el.tag == qn("w:r"):
                        p_el.remove(el)
                # Insert value run right after the permStart that opens this region
                perm_starts = [c for c in p_el if c.tag == qn("w:permStart")]
                if idx < len(perm_starts):
                    perm_starts[idx].addnext(_make_run(values[idx], 14))
            return


def _fill_signature_paragraphs(doc: Any, list_obj: Any) -> Any:
    """Fill 制定人/审核人/批准人 rows.

    模板结构: Run0=标签("制定人/日期：") Run1=空格(可填写区)。
    保留标签 run，仅替换空格 run 为实际值。
    """
    mapping = [
        ("制定人", list_obj.creator or ""),
        ("审核人", list_obj.reviewer or ""),
        ("批准人", list_obj.approver or ""),
    ]
    for para in doc.paragraphs:
        runs_text = "".join(r.text for r in para.runs)
        for key, value in mapping:
            if runs_text.startswith(key):
                runs = para.runs
                if runs:
                    # 模板最后一个 run 是空白填写区，替换为实际值
                    runs[-1].text = f"{value}            "
                break


def _enable_update_fields(doc: Any) -> Any:
    """Force Word/WPS to refresh PAGE fields on open."""
    settings_el = doc.settings.element
    uf = settings_el.find(qn("w:updateFields"))
    if uf is None:
        uf = OxmlElement("w:updateFields")
        settings_el.append(uf)
    uf.set(qn("w:val"), "true")


def _make_field_runs(instr: str, rpr: Any) -> Any:
    """Build begin/instrText/separate/cached/end runs for a Word field."""
    runs = []
    for tag in ("begin", "separate", "end"):
        r = OxmlElement("w:r")
        if rpr is not None:
            r.append(deepcopy(rpr))
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), tag)
        r.append(fc)
        runs.append(r)
    # instrText run
    ri = OxmlElement("w:r")
    if rpr is not None:
        ri.append(deepcopy(rpr))
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    ri.append(it)
    # cached result run (plain, refreshed on field update)
    rc = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    rc.append(t)
    return [runs[0], ri, runs[1], rc, runs[2]]


def _inject_numpages_field(doc: Any) -> Any:
    """Replace the hard-coded total-page "1" in the header with a NUMPAGES field.

    模板页眉结构: P [perm#14]{PAGE}[/][perm#15]"1"。
    总页数区域的 run 是 permStart..permEnd 之间的纯文本 "1"，
    将其替换为 NUMPAGES 域，打开文档时随 updateFields 一起刷新。
    """
    for section in doc.sections:
        for para in section.header.paragraphs:
            p_el = para._p
            children = list(p_el)
            for i, child in enumerate(children):
                if child.tag != qn("w:permStart"):
                    continue
                nxt = children[i + 1] if i + 1 < len(children) else None
                after = children[i + 2] if i + 2 < len(children) else None
                if nxt is None or nxt.tag != qn("w:r"):
                    continue
                if after is None or after.tag != qn("w:permEnd"):
                    continue
                t = nxt.find(qn("w:t"))
                if t is None or (t.text or "").strip() != "1":
                    continue
                rpr = nxt.find(qn("w:rPr"))
                for fr in _make_field_runs(" NUMPAGES  \\* MERGEFORMAT ", rpr):
                    nxt.addprevious(fr)
                p_el.remove(nxt)
                logger.debug("header total-page text replaced with NUMPAGES field")


# 允许重新序列化的 part（本次导出的预期修改点），其余 part 从模板原样拷贝
_PATCHED_PARTS = {"word/document.xml", "word/header1.xml", "word/settings.xml"}


def _save_minimal_patch(doc: Any, template_path: Path) -> BytesIO:
    """Zip-level minimal patch: 只替换表格/页眉/settings 三个 part。

    python-docx 的 save 会重新序列化全部 XML part（引号风格、Content_Types
    默认条目、core 属性等漂移），与模板产生非必要差异。这里先在内存生成
    修改后的 part，再拼回模板的 zip 包，保证样式/字体/页脚/图片/Content_Types
    等与用户模板逐字节一致。
    """
    generated = BytesIO()
    doc.save(generated)
    generated.seek(0)

    out = BytesIO()
    with (
        zipfile.ZipFile(template_path) as tpl,
        zipfile.ZipFile(generated) as gen,
        zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as dst,
    ):
        for item in tpl.infolist():
            if item.is_dir():
                continue
            if item.filename in _PATCHED_PARTS:
                dst.writestr(item, gen.read(item.filename))
            else:
                dst.writestr(item, tpl.read(item.filename))
        # 生成包中新增的 part（当前预期没有，防御性保留）
        tpl_names = {i.filename for i in tpl.infolist()}
        for item in gen.infolist():
            if not item.is_dir() and item.filename not in tpl_names:
                dst.writestr(item, gen.read(item.filename))
    out.seek(0)
    return out


def generate_position_training_list(list_obj: Any) -> BytesIO:
    """Fill the position training list template with data."""
    template_path = _find_template()
    doc = Document(str(template_path))

    # ── 段落填充（保留黄色可编辑标记）──
    _fill_dept_pos_paragraph(doc, list_obj.department or "", list_obj.position or "")
    _fill_signature_paragraphs(doc, list_obj)

    # ── 页眉页码：总页数改为 NUMPAGES 域 + 打开文档时更新域 ──
    _inject_numpages_field(doc)
    _enable_update_fields(doc)

    # ── 表格填充 ──
    dept_items = sorted(
        [i for i in list_obj.items if i.level == "部门级" and not i.is_deleted],
        key=lambda x: x.sort_order or 0,
    )
    pos_items = sorted(
        [i for i in list_obj.items if i.level == "岗位级" and not i.is_deleted],
        key=lambda x: x.sort_order or 0,
    )

    _build_table(doc.tables[0], dept_items, pos_items)

    return _save_minimal_patch(doc, template_path)


def _build_table(table: Any, dept_items: list[Any], pos_items: list[Any]) -> Any:
    """Rebuild table: data rows = exactly len(items); remove "……" & blank rows when data
    exists.

    Template layout (16 rows):
      Row0 部门级标题(灰) Row1 表头 Row2-6 数据(5行) Row7 ……
      Row8 岗位级标题(灰) Row9 表头 Row10-14 数据(5行) Row15 ……
    """
    tbl = table._tbl

    # ── 1. 删除模板的数据行和"……"行（bottom-up，避免索引偏移）──
    # 岗位级: Row10-14 数据 + Row15 …… ; 部门级: Row2-6 数据 + Row7 ……
    removal = list(range(10, 16)) + list(range(2, 8))
    for idx in sorted(removal, reverse=True):
        tbl.remove(table.rows[idx]._tr)

    # 模板的可编辑区 permEnd 标记是 <w:tbl> 的直接子元素（行与行之间），
    # 删行后它们会残留，必须一并清除，新行会自带配对的标记
    for el in list(tbl):
        tag = el.tag.split("}")[-1]
        if tag in ("permStart", "permEnd"):
            tbl.remove(el)
    # 剩余: [0]部门级标题 [1]表头 [2]岗位级标题 [3]表头

    # 用表头行作为克隆模板（数据行从表头行克隆后重新填充）
    dept_header_tr = table.rows[1]._tr
    pos_header_tr = table.rows[3]._tr

    # ── 2. 部门级数据行（插在部门表头后、岗位标题前）──
    prev = dept_header_tr
    for i, item in enumerate(dept_items):
        clone = _clone_data_row(dept_header_tr)
        prev.addnext(clone)
        prev = clone
        _fill_row_element(clone, i + 1, item)

    # ── 3. 岗位级数据行（插在岗位表头后）──
    prev = pos_header_tr
    for i, item in enumerate(pos_items):
        clone = _clone_data_row(pos_header_tr)
        prev.addnext(clone)
        prev = clone
        _fill_row_element(clone, i + 1, item)

    # ── 4. 无数据的级别保留模板原样（5空行+……）──
    # 由于已删除全部数据行，无数据时补回 5 个空行 + …… 行
    _ensure_placeholder(table, "部门级", dept_items)
    _ensure_placeholder(table, "岗位级", pos_items)

    # ── 5. 重设灰色标题行格式 ──
    for ri, row in enumerate(table.rows):
        first = row.cells[0].text.strip()
        if first in ("部门级", "岗位级"):
            for ci in range(5):
                _fill_cell(row.cells[ci], first, 11, bold=True, shade="D7D7D7")


# 克隆行 perm 标记的全局唯一 id 起始值（模板已占用 0-13）
_next_perm_id = [100]


def _add_row_perm_markers(clone: Any, with_cols: bool = True) -> Any:
    """按模板结构给行添加黄色可编辑区标记: 首格 permStart + 末格 permEnd。"""
    perm_id = str(_next_perm_id[0])
    _next_perm_id[0] += 1
    tcs = clone.findall(qn("w:tc"))
    if not tcs:
        return
    # 首格: permStart
    p0 = tcs[0].find(qn("w:p"))
    if p0 is not None:
        perm = OxmlElement("w:permStart")
        perm.set(qn("w:id"), perm_id)
        perm.set(qn("w:edGrp"), "everyone")
        if with_cols:
            perm.set(qn("w:colFirst"), "1")
            perm.set(qn("w:colLast"), "5")
        paragraph_properties = p0.find(qn("w:pPr"))
        if paragraph_properties is not None:
            paragraph_properties.addnext(perm)
        else:
            p0.insert(0, perm)
    # 末格: permEnd
    last_paragraph = tcs[-1].find(qn("w:p"))
    if last_paragraph is not None:
        perm_end = OxmlElement("w:permEnd")
        perm_end.set(qn("w:id"), perm_id)
        last_paragraph.append(perm_end)


def _clone_data_row(header_tr: Any) -> Any:
    """Clone a header row and strip its text to make an empty data row."""
    clone = deepcopy(header_tr)
    for tc in clone.findall(qn("w:tc")):
        p = tc.find(qn("w:p"))
        if p is None:
            continue
        for child in list(p):
            if child.tag == qn("w:r"):
                p.remove(child)
    _add_row_perm_markers(clone, with_cols=True)
    return clone


def _ensure_placeholder(table: Any, level_name: str, items: list[Any]) -> Any:
    """如果该级别没有数据，补回 5 个空白行 + …… 行（保持模板原样）。"""
    if items:
        return
    # 找到该级别的表头行
    header_idx = None
    title_idx = None
    for ri, row in enumerate(table.rows):
        first = row.cells[0].text.strip()
        if first == level_name:
            title_idx = ri
        elif first == "序号" and title_idx is not None and header_idx is None:
            header_idx = ri
            break
    if header_idx is None:
        return
    header_tr = table.rows[header_idx]._tr
    prev = header_tr
    for _ in range(5):
        clone = _clone_data_row(header_tr)
        prev.addnext(clone)
        prev = clone
    # …… 行（模板中 …… 行的 permStart 不带 colFirst/colLast）
    ell = deepcopy(header_tr)
    for tc in ell.findall(qn("w:tc")):
        p = tc.find(qn("w:p"))
        if p is None:
            continue
        for child in list(p):
            if child.tag == qn("w:r"):
                p.remove(child)
    # 先移除 _clone 未加的标记，再按模板格式添加
    _add_row_perm_markers(ell, with_cols=False)
    prev.addnext(ell)
    # 填充 …… 文字
    for tc in ell.findall(qn("w:tc")):
        p = tc.find(qn("w:p"))
        if p is None:
            continue
        run = _make_run("……")
        perm = p.find(qn("w:permStart"))
        if perm is not None:
            perm.addnext(run)
        else:
            p.append(run)


def _fill_row_element(tr: Any, seq: int, item: Any) -> Any:
    """Fill a cloned <w:tr> element's cells."""
    cells = tr.findall(qn("w:tc"))
    values = [
        str(seq),
        item.textbook_name or "",
        item.textbook_code or "",
        item.assessment_method or "",
        item.remarks or "",
    ]
    for ci, tc in enumerate(cells[:5]):
        p = tc.find(qn("w:p"))
        if p is None:
            continue
        # Remove runs, keep perm markers
        for child in list(p):
            if child.tag == qn("w:r"):
                p.remove(child)
        run = _make_run(values[ci])
        perm_start = p.find(qn("w:permStart"))
        paragraph_properties = p.find(qn("w:pPr"))
        if perm_start is not None:
            perm_start.addnext(run)
        elif paragraph_properties is not None:
            paragraph_properties.addnext(run)
        else:
            p.append(run)
