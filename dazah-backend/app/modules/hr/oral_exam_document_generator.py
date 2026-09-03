"""Generate oral exam result document (APP10-SMP-HR-002-14) from docx template.

模板保真填充 + 动态行克隆：问题行/人员行按数据增删，合并单元格版式保持不变。
仅填充填写区，保留模板 run 格式（字体/字号/加粗），复选框用 run 级文本替换。
模板自带"限制编辑"保护（settings.xml 随保存继承）；填写区例外标记在填充后
统一重建——克隆行原样复制模板标记会产生重复 w:id，Word 视为无效导致新行锁定。
"""

import copy
import logging
from collections.abc import Iterator
from io import BytesIO
from itertools import count
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.modules.hr.date_format import fmt_date_str
from app.modules.hr.schemas import OralExamExportRequest
from app.modules.hr.template_filler import (
    fill_whole_cell,
    header_fmt_source,
    replace_text_in_cell,
)

logger = logging.getLogger(__name__)

TEMPLATE_NAME = "APP10-SMP-HR-002-14口试培训考核结果表.docx"

_WML = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _find_template() -> Path:
    candidates = [
        Path("员工培训教育管理规程") / TEMPLATE_NAME,
        Path(__file__).resolve().parent.parent.parent
        / "员工培训教育管理规程"
        / TEMPLATE_NAME,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {TEMPLATE_NAME}")


def _row_texts(row: Any) -> list[str]:
    return [c.text.strip() for c in row.cells]


def _clone_row_after(table: Any, row_idx: int) -> None:
    tr = table.rows[row_idx]._tr
    new_tr = copy.deepcopy(tr)
    tr.addnext(new_tr)


def _find_row(table: Any, pred: Any, start: int = 0) -> int:
    for i in range(start, len(table.rows)):
        if pred(table.rows[i]):
            return i
    return -1


def _strip_permission_markers(doc: Any) -> None:
    """清空模板预置的可编辑例外标记（permStart/permEnd）。

    动态克隆行会原样复制模板标记，重复 w:id 会让 Word 判定区域无效，
    因此填充后统一清掉、按最终版式重建。
    """
    for tag in ("permStart", "permEnd"):
        for el in list(doc.element.body.iter(_WML + tag)):
            el.getparent().remove(el)


def _mark_cell_editable(cell: Any, next_id: Iterator[int]) -> None:
    """整格包上"每个人可编辑"例外区；受保护文档中其余内容保持锁定。"""
    paragraphs = cell.paragraphs
    if not paragraphs:
        return
    marker_id = str(next(next_id))
    start = OxmlElement("w:permStart")
    start.set(qn("w:id"), marker_id)
    start.set(qn("w:edGrp"), "everyone")
    first_p = paragraphs[0]._p
    ppr = first_p.find(qn("w:pPr"))
    if ppr is not None:
        ppr.addnext(start)
    else:
        first_p.insert(0, start)
    end = OxmlElement("w:permEnd")
    end.set(qn("w:id"), marker_id)
    paragraphs[-1]._p.append(end)


def _row_unique_cells(row: Any) -> list[Any]:
    """行内实体单元格（合并单元格去重）。"""
    cells: list[Any] = []
    for cell in row.cells:
        if any(cell._tc is tc for tc in cells):
            continue
        cells.append(cell)
    return cells


def generate_oral_exam_result(data: OralExamExportRequest) -> BytesIO:
    """按 APP10 模板保真生成口试培训考核结果表 Word."""
    doc = Document(str(_find_template()))
    if not doc.tables:
        raise ValueError("模板文件中未找到表格")
    table = doc.tables[0]
    fmt_src = header_fmt_source(table)

    # R1: 培训内容 | 培训日期
    fill_whole_cell(table.rows[1].cells[1], data.training_content or "", fmt_src)
    fill_whole_cell(table.rows[1].cells[5], fmt_date_str(data.training_date), fmt_src)

    # ── 问题区：首行题号"1"，末尾"……"行 ──
    q_anchor = _find_row(table, lambda r: _row_texts(r)[0] == "1")
    if q_anchor >= 0:
        ell1 = _find_row(table, lambda r: _row_texts(r)[0] == "……", q_anchor)
        have = ell1 - q_anchor
        need = len(data.questions)
        while have < need:
            idx = _find_row(table, lambda r: _row_texts(r)[0] == "……", q_anchor)
            _clone_row_after(table, idx - 1)
            have += 1
        while have > need:
            idx = _find_row(table, lambda r: _row_texts(r)[0] == "……", q_anchor)
            table.rows[idx - 1]._tr.getparent().remove(table.rows[idx - 1]._tr)
            have -= 1
        ell1 = _find_row(table, lambda r: _row_texts(r)[0] == "……", q_anchor)
        capacity = ell1 - q_anchor
        for i in range(capacity):
            row = table.rows[q_anchor + i]
            if i < need:
                q = data.questions[i]
                fill_whole_cell(row.cells[0], q.no or str(i + 1), fmt_src)
                fill_whole_cell(row.cells[1], q.question or "", fmt_src)
                fill_whole_cell(row.cells[4], q.answer or "", fmt_src)
            else:
                fill_whole_cell(row.cells[0], "", fmt_src)
                fill_whole_cell(row.cells[1], "", fmt_src)
                fill_whole_cell(row.cells[4], "", fmt_src)

    # ── 人员区：表头"序号"，数据行含"合格□"，末尾"……"行 ──
    p_header = _find_row(table, lambda r: _row_texts(r)[0] == "序号")
    if p_header >= 0:
        p_first = p_header + 1
        ell2 = _find_row(table, lambda r: _row_texts(r)[0] == "……", p_first)
        have_p = ell2 - p_first
        need_p = len(data.persons)
        while have_p < need_p:
            idx = _find_row(table, lambda r: _row_texts(r)[0] == "……", p_first)
            _clone_row_after(table, idx - 1)
            have_p += 1
        while have_p > need_p:
            idx = _find_row(table, lambda r: _row_texts(r)[0] == "……", p_first)
            table.rows[idx - 1]._tr.getparent().remove(table.rows[idx - 1]._tr)
            have_p -= 1
        ell2 = _find_row(table, lambda r: _row_texts(r)[0] == "……", p_first)
        capacity_p = ell2 - p_first
        for i in range(capacity_p):
            row = table.rows[p_first + i]
            if i < need_p:
                p = data.persons[i]
                fill_whole_cell(row.cells[0], str(i + 1), fmt_src)
                fill_whole_cell(row.cells[1], p.name, fmt_src)
                fill_whole_cell(row.cells[2], p.department or "", fmt_src)
                fill_whole_cell(row.cells[3], p.question_nos or "", fmt_src)
                # 复选框：合格□不合格□ → 勾选对应项
                if p.result == "合格":
                    replace_text_in_cell(row.cells[4], "合格□", "合格☑")
                elif p.result == "不合格":
                    replace_text_in_cell(row.cells[4], "不合格□", "不合格☑")
                fill_whole_cell(row.cells[6], p.remark or "—", fmt_src)
            else:
                fill_whole_cell(row.cells[0], "", fmt_src)
                fill_whole_cell(row.cells[1], "", fmt_src)
                fill_whole_cell(row.cells[2], "", fmt_src)
                fill_whole_cell(row.cells[3], "", fmt_src)
                fill_whole_cell(row.cells[6], "—", fmt_src)

    # ── 落款：评估人/日期（标签在后栏，值填入其后空白格） ──
    sig = _find_row(table, lambda r: "评估人" in "".join(_row_texts(r)))
    if sig >= 0:
        fill_whole_cell(table.rows[sig].cells[5], data.assessor or "", fmt_src)

    # ── 可编辑例外区：填充后统一重建（克隆行复制模板标记会产生重复 ID） ──
    # 培训内容/培训日期、全部问题行（含克隆行与备用空行）、全部人员行、评估人值；
    # 其余内容（表头/标签/固定文字）保持模板锁定状态不动。
    _strip_permission_markers(doc)
    targets = [table.rows[1].cells[1], table.rows[1].cells[5]]
    if q_anchor >= 0:
        q_end = _find_row(table, lambda r: _row_texts(r)[0] == "……", q_anchor)
        for idx in range(q_anchor, q_end if q_end >= 0 else len(table.rows)):
            targets.extend(_row_unique_cells(table.rows[idx]))
    p_header = _find_row(table, lambda r: _row_texts(r)[0] == "序号")
    if p_header >= 0:
        p_first = p_header + 1
        p_end = _find_row(table, lambda r: _row_texts(r)[0] == "……", p_first)
        for idx in range(p_first, p_end if p_end >= 0 else len(table.rows)):
            targets.extend(_row_unique_cells(table.rows[idx]))
    if sig >= 0:
        targets.append(table.rows[sig].cells[5])
    next_id = count(1)
    marked: list[Any] = []
    for cell in targets:
        if any(cell._tc is tc for tc in marked):
            continue
        marked.append(cell._tc)
        _mark_cell_editable(cell, next_id)

    # ── 清理模板预置的“……”占位行（导出文档不保留） ──
    for i in range(len(table.rows) - 1, -1, -1):
        if _row_texts(table.rows[i])[0] == "……":
            tr = table.rows[i]._tr
            tr.getparent().remove(tr)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
