"""HR business workflows live here."""

import asyncio
import copy
import logging
import re
from datetime import UTC, date, datetime, timedelta
from io import BytesIO
from typing import Any, cast
from uuid import UUID

from sqlalchemy import desc, func, select, update
from sqlalchemy.ext.asyncio import AsyncSession

from app.core import storage
from app.core.exceptions import AppException, DuplicateException, NotFoundException
from app.core.llm import (
    LLMOutputError,
    LLMProviderError,
    LLMRateLimitError,
    llm_client,
)
from app.core.upload_security import safe_upload_filename
from app.modules.hr.attachment_parser import (
    ANNEX_RE,
    SectionDraft,
    build_outline,
    build_preview,
    extract_annex_refs,
    normalize_annex_no,
    parse_sections,
    strip_punct,
)
from app.modules.hr.feishu import FeishuBitableSync
from app.modules.hr.feishu.departure_datasource import DepartureBitableDataSource
from app.modules.hr.feishu.employee_datasource import (
    EmployeeBitableDataSource,
)
from app.modules.hr.feishu.onboarding_datasource import OnboardingBitableDataSource
from app.modules.hr.feishu_settings_service import get_hr_feishu_app_credentials
from app.modules.hr.legacy_models import (
    DepartureRecord as LegacyDepartureRecord,
)
from app.modules.hr.legacy_models import (
    OnboardingRecord as LegacyOnboardingRecord,
)
from app.modules.hr.models import (
    AnnualTrainingPlan,
    AnnualTrainingPlanItem,
    Employee,
    HrDepartment,
    HrDocumentTemplate,
    HrFeishuEntitySetting,
    HrFeishuMember,
    OffboardingRecord,
    PlanAttachment,
    PlanAttachmentSection,
    PositionTransferRecord,
    Team,
    Trainer,
    TrainingLedger,
    TrainingLedgerPage,
    TrainingPersonnelConfig,
    TrainingSession,
)
from app.modules.hr.repository import (
    AnnualTrainingPlanItemRepository,
    AnnualTrainingPlanRepository,
    ContractManagementRepository,
    DepartmentRepository,
    DepartureRecordRepository,
    EmployeeRepository,
    EmployeeTrainingListRepository,
    OffboardingRecordRepository,
    OnboardingRecordRepository,
    PlanAttachmentRepository,
    PlanAttachmentSectionRepository,
    PositionTransferRecordRepository,
    TeamRepository,
    TrainingLedgerPageRepository,
    TrainingLedgerRepository,
    TrainingPersonnelConfigRepository,
)
from app.modules.hr.schemas import (
    AnnualTrainingPlanCreate,
    AnnualTrainingPlanItemBatchUpdate,
    AnnualTrainingPlanItemCreate,
    AnnualTrainingPlanUpdate,
    DepartmentCreate,
    DepartmentUpdate,
    EmployeeCreate,
    EmployeeUpdate,
    NewHireOut,
    OffboardingRecordCreate,
    OffboardingRecordUpdate,
    PositionTransferRecordCreate,
    PositionTransferRecordUpdate,
    SyncStatusResponse,
    TeamCreate,
    TeamUpdate,
    TrainingLedgerCreate,
    TrainingLedgerUpdate,
    TrainingPersonnelConfigCreate,
)

logger = logging.getLogger(__name__)


def _dept_mapping_to_dict(mapping: Any) -> dict[str, Any]:
    """映射配置 ORM → dict（供 API 返回）"""
    return {
        "id": str(mapping.id),
        "source_name": mapping.source_name,
        "target_name": mapping.target_name,
        "match_level": mapping.match_level,
        "mapping_type": mapping.mapping_type,
        "priority": mapping.priority,
        "enabled": mapping.enabled,
        "remark": mapping.remark,
        "created_at": mapping.created_at.isoformat() if mapping.created_at else None,
        "updated_at": mapping.updated_at.isoformat() if mapping.updated_at else None,
    }


# ─── Feishu field mapping helpers ───


def _extract_text(value: Any) -> str:
    """Extract text from Feishu array format or plain string."""
    if isinstance(value, list) and len(value) > 0 and isinstance(value[0], dict):
        return str(value[0].get("text") or "")
    if isinstance(value, dict):
        if "text" in value:
            return str(value["text"] or "")
        if "value" in value and isinstance(value["value"], list):
            inner = value["value"]
            if len(inner) > 0 and isinstance(inner[0], dict):
                return str(inner[0].get("text") or "")
    if value is None:
        return ""
    return str(value)


def _extract_number(value: Any) -> int | None:
    """Extract number from Feishu format."""
    if isinstance(value, (int, float)):
        return int(value)
    if isinstance(value, dict):
        if (
            "value" in value
            and isinstance(value["value"], list)
            and len(value["value"]) > 0
        ):
            return int(value["value"][0])
    return None


def _ms_to_date(value: Any) -> date | None:
    """Convert Feishu millisecond timestamp to Python date.

    Feishu stores date fields as the timestamp of the calendar date in the
    user's local timezone (UTC+8 for China). Converting with the server's
    local timezone reproduces the calendar date shown in Feishu. Do NOT use
    UTC here — it would shift the date back by one day.
    """
    if isinstance(value, (int, float)) and value > 0:
        return datetime.fromtimestamp(value / 1000).date()
    return None


def _parse_date(value: Any) -> date | None:
    """Parse Feishu date value (ms timestamp or ISO string) to Python date."""
    if value is None:
        return None
    if isinstance(value, (int, float)):
        if value > 0:
            return _ms_to_date(value)
        return None
    if isinstance(value, str):
        try:
            return datetime.strptime(value[:10], "%Y-%m-%d").date()
        except (ValueError, IndexError):
            return None
    return None


def _parse_contract_date(value: Any) -> date | None:
    """Parse contract end date from Feishu text array or string formats."""
    text = _extract_text(value)
    if not text:
        return None
    for fmt in ("%Y/%m/%d", "%Y-%m-%d", "%Y%m%d"):
        try:
            return datetime.strptime(text.strip()[:10], fmt).date()
        except (ValueError, IndexError):
            continue
    return None


def _parse_feishu_record(record: dict[str, Any]) -> dict[str, Any]:
    """Convert a raw Feishu record into Employee constructor kwargs."""
    fields = record.get("fields", {})
    rid = record.get("record_id", "")
    updated_time = record.get("updated_time", "")

    def gt(key: str) -> Any:
        return fields.get(key)

    qualifications = gt("技能证书") or gt("职称／职业资格")
    data = {
        "feishu_record_id": rid,
        "seq_number": _extract_number(gt("序号")),
        "employee_number": _extract_text(gt("工号")),
        "name": _extract_text(gt("姓名")),
        "domain_account": _extract_text(gt("域账户")),
        "department": _extract_text(gt("一级部门") or gt("部门")),
        "sub_department": _extract_text(gt("二级部门")),
        "position": _extract_text(gt("职务|岗位") or gt("职位")),
        "level": _extract_text(gt("职级")),
        "qualifications": (
            qualifications
            if isinstance(qualifications, list)
            else ([qualifications] if isinstance(qualifications, str) else None)
        ),
        "qualification_type": _extract_text(gt("职称")),
        "certificate_number": _extract_text(gt("证书编号")),
        "certificate_review_date": _parse_contract_date(gt("技能证书复审时间")),
        "gender": _extract_text(gt("性别")),
        "ethnic_group": _extract_text(gt("民族")),
        "native_place": _extract_text(gt("籍贯")),
        "political_status": _extract_text(gt("政治面貌")),
        "marital_status": _extract_text(gt("婚姻状况")),
        "household_type": _extract_text(gt("户口类别")),
        "employment_type": _extract_text(gt("人员就业方式")),
        "probation_status": _extract_text(gt("转正状态")),
        "planned_probation_date": _parse_date(gt("拟转正日期")),
        "probation_effective_date": _parse_date(gt("转正生效日期")),
        "work_start_date": _parse_date(gt("参加工作时间")),
        "factory_entry_date": _parse_date(gt("进本公司时间")),
        "hire_date": _parse_date(gt("入职日期")) or _parse_date(gt("进厂时间")),
        "graduation_date": _parse_date(gt("毕业时间")),
        "education": _extract_text(gt("学历")),
        "degree": _extract_text(gt("学位")),
        "school": _extract_text(gt("毕业院校")),
        "major": _extract_text(gt("专业")),
        "id_card": _extract_text(gt("身份证号")),
        "id_card_expiry": _extract_text(gt("身份证有效期截止日期")),
        "current_address": _extract_text(gt("现居住地址")),
        "contract_start_date": _ms_to_date(gt("首次签订合同日期")),
        "contract_end_date": _ms_to_date(gt("首次签订合同截止日期")),
        "contract_start_2": _ms_to_date(gt("第二次续签合同日期")),
        "contract_end_2": _parse_contract_date(gt("合同截止日期（2）")),
        "contract_start_3": _ms_to_date(gt("第三次续签合同日期")),
        "contract_end_3": _parse_contract_date(gt("合同截止日期（3）")),
        "contract_start_4": _ms_to_date(gt("第四次续签合同日期")),
        "contract_end_4": _parse_contract_date(gt("合同截止日期4")),
        "contract_start_5": _ms_to_date(gt("第五次续签合同日期")),
        "contract_end_5": _extract_text(gt("合同截止日期5")),
        "contract_start_6": _extract_text(gt("第六次续签合同日期")),
        "contract_end_6": _extract_text(gt("合同截止日期6")),
        "phone": _extract_text(gt("联系电话") or gt("手机")),
        "email": _extract_text(gt("电子邮箱")),
        "emergency_contact_name": _extract_text(gt("紧急联系人姓名")),
        "emergency_contact_phone": _extract_text(gt("紧急联系人电话")),
        "emergency_contact_relation": _extract_text(gt("与本人关系")),
        "archive_number": _extract_text(gt("档案编号")),
        "health_status": _extract_text(gt("健康情况")),
        "status_category": _extract_text(gt("人员类别")),
        "remarks": (
            gt("备注")
            if isinstance(gt("备注"), list)
            else ([gt("备注")] if isinstance(gt("备注"), str) and gt("备注") else None)
        ),
        "work_experience_1": _extract_text(gt("工作经验一")),
        "work_experience_2": _extract_text(gt("工作经验二")),
        "work_experience_3": _extract_text(gt("工作经验三")),
        "work_experience_4": _extract_text(gt("工作经验四")),
        "offboarding_type": _extract_text(gt("离职类型")),
        "offboarding_reason": _extract_text(gt("离职原因")),
        "status": _extract_text(gt("在职状态")) or "在职",
    }
    # Parse updated_time for sync tracking
    if updated_time:
        try:
            # Feishu returns ISO format string like "2024-01-15T08:30:00.000000Z"
            dt = datetime.fromisoformat(updated_time.replace("Z", "+00:00"))
            data["feishu_synced_at"] = dt.date()
        except Exception:
            data["feishu_synced_at"] = date.today()
    else:
        data["feishu_synced_at"] = date.today()

    # 以飞书为主：飞书为空的字段也返回空字符串，用于覆盖本地旧值。
    # 不再过滤空值，保证"飞书没有的，本地也删除"。
    cleaned = {k: v for k, v in data.items()}
    return cleaned


# ─── Services ───


async def _resolve_feishu_sync_session(session: Any) -> FeishuBitableSync:
    """从 DB 解析飞书应用凭据，构造 FeishuBitableSync 实例。

    供 EmployeeService 和 DepartmentService 复用，避免 _ensure_feishu_creds 重复。
    """
    app_id, app_secret = await get_hr_feishu_app_credentials(session)
    return FeishuBitableSync(app_id=app_id or None, app_secret=app_secret or None)


class EmployeeService:
    def __init__(self, session: AsyncSession) -> None:
        self.repo = EmployeeRepository(session)
        self.session = session
        self._feishu: FeishuBitableSync | None = None

    async def _ensure_feishu_creds(self) -> FeishuBitableSync:
        """Lazy-load FeishuBitableSync cache (delegates to module helper)."""
        legacy_feishu = getattr(self, "feishu", None)
        if legacy_feishu is not None:
            return cast(FeishuBitableSync, legacy_feishu)
        if self._feishu is None:
            self._feishu = await _resolve_feishu_sync_session(self.session)
        return self._feishu

    async def _get_bitable(self) -> EmployeeBitableDataSource:
        """Read employee entity settings from DB and return configured datasource."""
        legacy_bitable = getattr(self, "bitable", None)
        if legacy_bitable is not None:
            return cast(EmployeeBitableDataSource, legacy_bitable)
        query = select(HrFeishuEntitySetting).where(
            HrFeishuEntitySetting.entity_code == "employee",
            HrFeishuEntitySetting.is_enabled.is_(True),
        )
        result = await self.session.execute(query)
        setting = result.scalar_one_or_none()
        if not setting or not setting.app_token or not setting.base_table_id:
            raise AppException(message="员工飞书多维表格未配置或未启用")
        return EmployeeBitableDataSource(
            app_token=setting.app_token,
            table_id=setting.base_table_id,
        )

    async def get_employee(self, employee_id: UUID) -> Employee:
        employee = await self.repo.get_by_id(employee_id)
        if not employee:
            raise NotFoundException("员工", str(employee_id))
        return employee

    async def get_employee_by_number(self, employee_number: str) -> Employee:
        employee = await self.repo.get_by_employee_number(employee_number)
        if not employee:
            raise NotFoundException("员工", employee_number)
        return employee

    async def get_employee_stats(
        self, dept_alias_set: set[str] | None = None
    ) -> dict[str, Any]:
        return await self.repo.get_stats(dept_alias_set=dept_alias_set)

    async def create_employee(
        self, data: EmployeeCreate
    ) -> Employee | tuple[Employee, str]:
        # 空字符串工号转为 None，避免 unique 约束冲突
        employee_number = getattr(data, "employee_number", None)
        if employee_number is not None and employee_number.strip() == "":
            data.employee_number = None

        if data.employee_number:
            existing = await self.repo.get_by_employee_number(data.employee_number)
            if existing:
                raise DuplicateException("工号", data.employee_number)

        employee = Employee(**data.model_dump())
        employee.status = "在职"

        # 根据手机号获取飞书 open_id（非阻塞，失败仅记录日志）
        if data.phone:
            try:
                from app.modules.hr.feishu.im import FeishuIM

                im = FeishuIM()
                # 飞书接口要求手机号带 +86 区号
                mobile = (
                    data.phone if data.phone.startswith("+") else f"+86{data.phone}"
                )
                mapping = await im.batch_get_open_ids_by_mobile([mobile])
                open_id = mapping.get(mobile) or mapping.get(data.phone)
                if open_id:
                    employee.feishu_open_id = open_id
                    logger.info(
                        "Fetched feishu_open_id for employee %s: %s",
                        data.employee_number,
                        open_id,
                    )
            except Exception as e:
                logger.warning(
                    "Failed to fetch feishu open_id for phone %s: %s",
                    data.phone,
                    e,
                )

        result = await self.repo.create(employee)

        # Sync to Feishu
        sync_status = "success"
        try:
            bitable = await self._get_bitable()
            rid = await bitable.create(self._to_bitable_fields(result))
            if rid:
                result.feishu_record_id = rid
                result.feishu_synced_at = date.today()
                await self.repo.update(result)
        except Exception as e:
            logger.warning("Feishu sync failed for employee created: %s", e)
            sync_status = f"failed: {str(e)}"

        if not hasattr(self, "session"):
            return result
        return result, sync_status

    async def approve_employee(self, employee_number: str) -> Employee:
        employee = await self.repo.get_by_employee_number(employee_number)
        if not employee:
            raise NotFoundException("员工", employee_number)
        if employee.status != "待审批":
            raise DuplicateException("审批", "该员工已审批完成")

        employee.status = "在职"
        result = await self.repo.update(employee)

        try:
            await self._sync_single_to_feishu(result)
        except Exception as e:
            logger.warning("Feishu sync failed for employee approved: %s", e)

        return result

    async def update_employee(
        self, employee_id: UUID, data: EmployeeUpdate
    ) -> Employee | tuple[Employee, str]:
        employee = await self.get_employee(employee_id)
        employee_number = getattr(data, "employee_number", None)
        if employee_number is not None and employee_number.strip() == "":
            data.employee_number = None
        update_data = data.model_dump(exclude_unset=True)

        if "employee_number" in update_data:
            existing = await self.repo.get_by_employee_number(
                update_data["employee_number"]
            )
            if existing and existing.id != employee_id:
                raise DuplicateException("工号", update_data["employee_number"])

        # 日期字段自动转换：前端传字符串 → date 对象
        date_fields = {
            "work_start_date",
            "factory_entry_date",
            "livo_entry_date",
            "hire_date",
            "graduation_date",
            "contract_start_date",
            "contract_end_date",
            "contract_start_2",
            "contract_end_2",
            "contract_start_3",
            "contract_end_3",
            "contract_start_4",
            "contract_end_4",
            "planned_probation_date",
            "probation_effective_date",
            "last_working_day",
        }
        for field in date_fields:
            if field in update_data and isinstance(update_data[field], str):
                try:
                    update_data[field] = datetime.strptime(
                        update_data[field], "%Y-%m-%d"
                    ).date()
                except ValueError:
                    update_data[field] = None  # 非法日期置空

        for field, value in update_data.items():
            setattr(employee, field, value)

        result = await self.repo.update(employee)

        sync_status = "success"
        try:
            await self._sync_single_to_feishu(result)
        except Exception as e:
            logger.warning("Feishu sync failed for employee updated: %s", e)
            sync_status = f"failed: {str(e)}"

        if not hasattr(self, "session"):
            return result
        return result, sync_status

    async def delete_employee(self, employee_id: UUID) -> str:
        employee = await self.get_employee(employee_id)
        employee_number = employee.employee_number
        await self.repo.soft_delete(employee)

        sync_status = "success"
        try:
            if employee.feishu_record_id:
                bitable = await self._get_bitable()
                await bitable.delete(employee.feishu_record_id)
            else:
                feishu = await self._ensure_feishu_creds()
                if employee_number:
                    await feishu.sync_employee_deleted(employee_number)
        except Exception as e:
            logger.warning("Feishu sync failed for employee deleted: %s", e)
            sync_status = f"failed: {str(e)}"

        return sync_status

    async def list_employees(
        self,
        *,
        department: str | None = None,
        sub_department: str | None = None,
        status: str | None = None,
        keyword: str | None = None,
        gender: str | None = None,
        level: str | None = None,
        position: str | None = None,
        page: int = 1,
        page_size: int = 20,
        sort_by: str = "created_at",
        sort_order: str = "desc",
        dept_alias_set: set[str] | None = None,
    ) -> tuple[list[Employee], int]:
        return await self.repo.list_employees(
            department=department,
            sub_department=sub_department,
            status=status,
            keyword=keyword,
            gender=gender,
            level=level,
            position=position,
            page=page,
            page_size=page_size,
            sort_by=sort_by,
            sort_order=sort_order,
            dept_alias_set=dept_alias_set,
        )

    async def notify_training(self, payload: Any) -> dict[str, Any]:
        """Send a training notice while retaining the legacy employee API."""
        from app.modules.hr.feishu.im import FeishuIM

        employees: list[Employee] = []
        for employee_number in payload.employee_numbers:
            employee = await self.repo.get_by_employee_number(employee_number)
            if employee:
                employees.append(employee)

        time_range = ""
        if payload.training_time_start and payload.training_time_end:
            time_range = f"{payload.training_time_start} ~ {payload.training_time_end}"
        content_lines = [
            "【培训通知】",
            f"主题：{payload.subject}",
            f"时间：{payload.training_date} {time_range}",
            f"地点：{payload.location or '待定'}",
            f"培训师：{payload.trainer or '待定'}",
        ]
        if payload.content:
            content_lines.append(f"内容：{payload.content}")
        content_lines.append("请准时参加，自带笔记本笔，不得无故缺席。")
        content = "\n".join(content_lines)

        sent = 0
        failed = 0
        details: list[dict[str, Any]] = []
        im = FeishuIM()
        for employee in employees:
            if not employee.feishu_open_id:
                failed += 1
                details.append(
                    {
                        "employee_number": employee.employee_number,
                        "name": employee.name,
                        "status": "failed",
                        "reason": "数据库中缺少 feishu_open_id，请先同步",
                    }
                )
                continue
            try:
                await im.send_text_message(employee.feishu_open_id, content)
                sent += 1
                details.append(
                    {
                        "employee_number": employee.employee_number,
                        "name": employee.name,
                        "status": "sent",
                    }
                )
            except Exception as exc:
                failed += 1
                details.append(
                    {
                        "employee_number": employee.employee_number,
                        "name": employee.name,
                        "status": "failed",
                        "reason": str(exc),
                    }
                )

        found_numbers = {employee.employee_number for employee in employees}
        for employee_number in payload.employee_numbers:
            if employee_number not in found_numbers:
                failed += 1
                details.append(
                    {
                        "employee_number": employee_number,
                        "status": "failed",
                        "reason": "未找到员工",
                    }
                )
        return {"sent": sent, "failed": failed, "details": details}

    # ─── Bi-directional sync ───

    async def sync_from_feishu(self) -> dict[str, Any]:
        """Pull all records from Feishu Bitable and upsert into local PG.

        同步策略：全量覆盖。飞书有的新增/更新，飞书没有的本地删除。
        Returns:
            {"created": N, "updated": N, "deleted": N, "failed": N, "total": N}
        """
        if not hasattr(self, "session"):
            legacy_bitable = getattr(self, "bitable", None)
            if legacy_bitable is None:
                raise AppException(message="员工飞书数据源未配置")
            raw_records = await legacy_bitable.client.search_records(
                legacy_bitable.table_id,
                page_size=500,
            )
            legacy_stats = {
                "created": 0,
                "updated": 0,
                "failed": 0,
                "total": len(raw_records),
            }
            for record in raw_records:
                try:
                    parsed = _parse_feishu_record(record)
                    employee_number = parsed.get("employee_number")
                    if not employee_number:
                        legacy_stats["failed"] += 1
                        continue
                    await self.repo.upsert_by_employee_number(parsed)
                    existing = await self.repo.get_by_employee_number(employee_number)
                    recently_created = (
                        existing
                        and existing.created_at
                        and (
                            datetime.utcnow() - existing.created_at.replace(tzinfo=None)
                        ).total_seconds()
                        < 60
                    )
                    legacy_stats["created" if recently_created else "updated"] += 1
                except Exception as exc:
                    logger.error(
                        "Failed to sync Feishu record %s: %s",
                        record.get("record_id"),
                        exc,
                    )
                    legacy_stats["failed"] += 1
            return legacy_stats

        bitable = await self._get_bitable()
        raw_records = await bitable.client.search_records(
            bitable.table_id,
            page_size=500,
        )
        stats = {
            "created": 0,
            "updated": 0,
            "deleted": 0,
            "failed": 0,
            "skipped": 0,
            "total": len(raw_records),
        }

        # 收集飞书中的所有 record_id
        feishu_record_ids: set[str] = set()

        for rec in raw_records:
            try:
                parsed = _parse_feishu_record(rec)
                emp_no = parsed.get("employee_number")
                rid = rec.get("record_id", "")
                if not emp_no:
                    stats["skipped"] += 1
                    logger.warning("飞书同步跳过：record_id=%s 工号为空", rid)
                    continue

                if rid:
                    feishu_record_ids.add(rid)

                await self.repo.upsert_by_employee_number(parsed)
                existing = await self.repo.get_by_employee_number(emp_no)
                if (
                    existing
                    and existing.created_at
                    and (
                        datetime.utcnow() - existing.created_at.replace(tzinfo=None)
                    ).total_seconds()
                    < 60
                ):
                    stats["created"] += 1
                else:
                    stats["updated"] += 1
            except Exception as e:
                logger.error(
                    "飞书同步异常：record_id=%s, error=%s, traceback:",
                    rec.get("record_id"),
                    e,
                    exc_info=True,
                )
                stats["failed"] += 1
                await self.session.rollback()

        # 删除飞书中已不存在的记录（按 feishu_record_id 判断）
        deleted = await self.repo.delete_not_in_feishu(feishu_record_ids)
        stats["deleted"] = deleted

        return stats

    async def sync_to_feishu(self, employee_id: UUID) -> str:
        employee = await self.get_employee(employee_id)
        return await self._sync_single_to_feishu(employee)

    async def list_contract_expiring(
        self,
        start_date: date,
        end_date: date,
        department: str | None = None,
        page: int = 1,
        page_size: int = 20,
        dept_alias_set: set[str] | None = None,
    ) -> tuple[list[dict[str, Any]], int]:
        """筛选合同到期人员。"""
        return await self.repo.list_contract_expiring(
            start_date=start_date,
            end_date=end_date,
            department=department,
            page=page,
            page_size=page_size,
            dept_alias_set=dept_alias_set,
        )

    async def get_sync_status(self) -> SyncStatusResponse:
        local_total = await self.repo.count_total()
        synced_count = await self.repo.count_synced()
        unsynced_count = local_total - synced_count

        # Use local synced count as feishu_total proxy to avoid expensive
        # real-time Feishu API calls (fetching 500 records takes ~5s).
        # Data consistency is ensured by the sync process itself.
        feishu_total = synced_count

        return SyncStatusResponse(
            local_total=local_total,
            feishu_total=feishu_total,
            synced_count=synced_count,
            unsynced_count=unsynced_count,
            conflict_count=0,  # TODO: implement conflict detection
            last_sync_at=None,
        )

    # ─── Internal helpers ───

    async def _sync_single_to_feishu(self, employee: Employee) -> str:
        """Sync one employee to Feishu, creating or updating as needed."""
        bitable = await self._get_bitable()
        fields = self._to_bitable_fields(employee)
        if employee.feishu_record_id:
            await bitable.update(employee.feishu_record_id, fields)
            return employee.feishu_record_id
        if not hasattr(self, "session"):
            record_id = await bitable.create(fields)
            employee.feishu_record_id = record_id
            employee.feishu_synced_at = date.today()
            await self.repo.update(employee)
            return record_id
        else:
            # 如果没有 feishu_record_id，先根据工号查找现有记录
            if not employee.employee_number:
                return ""
            existing = await bitable.find_by_employee_number(employee.employee_number)
            if existing:
                # 找到现有记录，更新它
                await bitable.update(existing.record_id, fields)
                # 回写 feishu_record_id 到本地数据库
                employee.feishu_record_id = existing.record_id
                employee.feishu_synced_at = date.today()
                await self.repo.update(employee)
                return existing.record_id
            else:
                # 没有找到现有记录，创建新记录
                rid = await bitable.create(fields)
                employee.feishu_record_id = rid
                employee.feishu_synced_at = date.today()
                await self.repo.update(employee)
                return rid

    def _to_bitable_fields(self, employee: Employee) -> dict[str, Any]:
        """Convert Employee ORM object to Feishu Bitable field dict.

        Filters out empty values to avoid Feishu validation errors
        (especially for phone fields which reject empty strings).
        """
        from app.modules.hr.feishu.bitable import _to_ms_timestamp

        if not hasattr(employee, "archive_number"):
            raw_legacy = {
                "姓名": employee.name,
                "工号": employee.employee_number,
                "部门": employee.department,
                "职位": employee.position,
                "手机": employee.phone,
                "邮箱地址": employee.email,
                "性别": employee.gender,
                "籍贯": employee.native_place,
                "政治面貌": employee.political_status,
                "婚姻状况": employee.marital_status,
                "学历": employee.education,
                "分类": employee.classification,
                "专业": employee.major,
                "身份证号": employee.id_card,
                "银行卡号": employee.bank_account,
                "培训档案编号": employee.training_id,
                "域账号": employee.domain_account,
                "班组": employee.team,
                "职类": employee.job_category,
                "级别": employee.level,
                "职称类型": employee.qualification_type,
                "户籍类型": employee.household_type,
                "统计类别": employee.status_category,
                "身份证到期日": employee.id_card_expiry,
                "合同期限": employee.contract_type,
                "身份证地址|家庭地址": employee.id_card_address,
                "现住址": employee.current_address,
                "紧急联系人电话": employee.emergency_contact_phone,
                "紧急联系人|关系": employee.emergency_contact_relation,
                "异动（含曾经工作部门、岗位)": employee.transfer_history,
            }
            fields: dict[str, Any] = {
                key: value
                for key, value in raw_legacy.items()
                if value not in (None, "", [])
            }
            if employee.qualifications:
                fields["职称／职业资格"] = employee.qualifications
            if employee.remarks:
                fields["备注"] = employee.remarks
            legacy_dates = {
                "参加工作时间": employee.work_start_date,
                "进厂时间": employee.factory_entry_date,
                "入丽珠时间": employee.livo_entry_date,
                "毕业时间": employee.graduation_date,
                "第一次合同起点时间": employee.contract_start_date,
                "第一次合同终止时间": employee.contract_end_date,
                "第二次合同起点时间": employee.contract_start_2,
                "第二次合同终止时间": employee.contract_end_2,
                "第三次合同起点时间": employee.contract_start_3,
                "第三次合同终止时间": employee.contract_end_3,
                "第四次合同起点时间": employee.contract_start_4,
                "第四次合同终止时间": employee.contract_end_4,
            }
            fields.update(
                {
                    key: _to_ms_timestamp(value)
                    for key, value in legacy_dates.items()
                    if value
                }
            )
            for key, value in (
                ("年", employee.birth_year),
                ("月", employee.birth_month),
                ("日", employee.birth_day),
            ):
                if value:
                    fields[key] = value
            return fields

        def date_to_str(d: Any) -> Any:
            """Convert date to string for text fields."""
            if d is None:
                return None
            if hasattr(d, "isoformat"):
                return d.isoformat()
            return str(d) if d else None

        # Build raw fields, keeping None/empty filtering for later
        # Field names must match Feishu Bitable column names exactly
        raw: dict[str, Any] = {
            "档案编号": employee.archive_number,
            "工号": int(employee.employee_number)
            if employee.employee_number and employee.employee_number.isdigit()
            else employee.employee_number,
            "域账户": employee.domain_account,
            "一级部门": employee.department,
            "二级部门": employee.sub_department,
            "职务|岗位": employee.position,
            "姓名": employee.name,
            "职级": employee.level,
            "性别": employee.gender,
            "民族": employee.ethnic_group,
            "籍贯": employee.native_place,
            "身份证号": employee.id_card,
            "身份证有效期截止日期": date_to_str(employee.id_card_expiry),  # 文本类型
            "婚姻状况": employee.marital_status,
            "政治面貌": employee.political_status,
            "现居住地址": employee.current_address,
            "户口类别": employee.household_type,
            "学历": employee.education,
            "学位": employee.degree,
            "毕业院校": employee.school,
            "专业": employee.major,
            "毕业时间": _to_ms_timestamp(employee.graduation_date)
            if employee.graduation_date
            else None,
            "证书编号": employee.certificate_number,
            "职称": employee.qualification_type,
            "技能证书复审时间": date_to_str(
                employee.certificate_review_date
            ),  # 文本类型
            "参加工作时间": _to_ms_timestamp(employee.work_start_date)
            if employee.work_start_date
            else None,
            "入职日期": _to_ms_timestamp(employee.hire_date)
            if employee.hire_date
            else None,
            "进本公司时间": _to_ms_timestamp(employee.factory_entry_date)
            if employee.factory_entry_date
            else None,
            "工龄1": employee.work_years,
            "在职状态": employee.status,
            "拟转正日期": _to_ms_timestamp(employee.planned_probation_date)
            if employee.planned_probation_date
            else None,
            "转正状态": employee.probation_status,
            "转正生效日期": _to_ms_timestamp(employee.probation_effective_date)
            if employee.probation_effective_date
            else None,
            "首次签订合同日期": _to_ms_timestamp(employee.contract_start_date)
            if employee.contract_start_date
            else None,
            "首次签订合同截止日期": _to_ms_timestamp(employee.contract_end_date)
            if employee.contract_end_date
            else None,
            "第二次续签合同日期": _to_ms_timestamp(employee.contract_start_2)
            if employee.contract_start_2
            else None,
            "合同截止日期（2）": date_to_str(employee.contract_end_2),  # 文本类型
            "第三次续签合同日期": _to_ms_timestamp(employee.contract_start_3)
            if employee.contract_start_3
            else None,
            "合同截止日期（3）": date_to_str(employee.contract_end_3),  # 文本类型
            "第四次续签合同日期": _to_ms_timestamp(employee.contract_start_4)
            if employee.contract_start_4
            else None,
            "合同截止日期4": date_to_str(employee.contract_end_4),  # 文本类型
            "第五次续签合同日期": _to_ms_timestamp(employee.contract_start_5)
            if employee.contract_start_5
            else None,
            "合同截止日期5": date_to_str(employee.contract_end_5),  # 文本类型
            "第六次续签合同日期": date_to_str(employee.contract_start_6),  # 文本类型
            "合同截止日期6": date_to_str(employee.contract_end_6),  # 文本类型
            "人员就业方式": employee.employment_type,
            "联系电话": employee.phone,
            "电子邮箱": employee.email,
            "紧急联系人姓名": employee.emergency_contact_name,
            "与本人关系": employee.emergency_contact_relation,
            "紧急联系人电话": employee.emergency_contact_phone,
            "健康情况": employee.health_status,
            "最后工作日": date_to_str(employee.last_working_day),  # 文本类型
            "工作经验一": employee.work_experience_1,
            "工作经验二": employee.work_experience_2,
            "工作经验三": employee.work_experience_3,
            "工作经验四": employee.work_experience_4,
            "人员类别": employee.status_category,
            "离职类型": employee.offboarding_type,
            "离职原因": employee.offboarding_reason,
        }
        # Filter out empty strings / None / empty lists
        fields = {k: v for k, v in raw.items() if v not in (None, "", [])}

        if employee.qualifications:
            fields["技能证书"] = employee.qualifications
        if employee.remarks:
            fields["备注"] = employee.remarks

        # Birth date - Feishu uses a single "出生年月" field (formula, read-only)
        # No need to set this field as it's computed by Bitable

        return fields


async def run_sync_from_feishu_background() -> dict[str, Any]:
    """后台任务：创建独立 DB session 执行飞书同步。

    供 jobs.submit_job 调用，不与 HTTP 请求共享 session。
    """
    from app.core.database import async_session_factory

    async with async_session_factory() as session:
        service = EmployeeService(session)
        result = await service.sync_from_feishu()
        await session.commit()
        return result


class DepartmentService:
    def __init__(self, session: AsyncSession) -> None:
        self.repo = DepartmentRepository(session)
        self.session = session
        self._feishu: FeishuBitableSync | None = None

    async def _ensure_feishu_creds(self) -> FeishuBitableSync:
        """Lazy-load FeishuBitableSync cache (delegates to module helper)."""
        if self._feishu is None:
            self._feishu = await _resolve_feishu_sync_session(self.session)
        return self._feishu

    async def get_department(self, department_id: UUID) -> HrDepartment:
        department = await self.repo.get_by_id(department_id)
        if not department:
            raise NotFoundException("部门", str(department_id))
        return department

    async def create_department(self, data: DepartmentCreate) -> HrDepartment:
        existing = await self.repo.get_by_code(data.code)
        if existing:
            raise DuplicateException("部门编码", data.code)

        department = HrDepartment(**data.model_dump())
        result = await self.repo.create(department)

        try:
            feishu = await self._ensure_feishu_creds()
            await feishu.sync_department_created(result.__dict__)
        except Exception as e:
            logger.warning("Feishu sync failed for department created: %s", e)

        return result

    async def update_department(
        self, department_id: UUID, data: DepartmentUpdate
    ) -> HrDepartment:
        department = await self.get_department(department_id)
        update_data = data.model_dump(exclude_unset=True)

        if "code" in update_data:
            existing = await self.repo.get_by_code(update_data["code"])
            if existing and existing.id != department_id:
                raise DuplicateException("部门编码", update_data["code"])

        # 显式处理 parent_id：确保清空上级部门时能生效（null → 设为 None）
        if "parent_id" in update_data:
            department.parent_id = update_data["parent_id"]
            update_data.pop("parent_id")

        for field, value in update_data.items():
            setattr(department, field, value)

        result = await self.repo.update(department)

        try:
            feishu = await self._ensure_feishu_creds()
            await feishu.sync_department_updated(result.__dict__)
        except Exception as e:
            logger.warning("Feishu sync failed for department updated: %s", e)

        return result

    async def delete_department(self, department_id: UUID) -> None:
        department = await self.get_department(department_id)
        code = department.code
        await self.repo.soft_delete(department)

        try:
            feishu = await self._ensure_feishu_creds()
            await feishu.sync_department_deleted(code)
        except Exception as e:
            logger.warning("Feishu sync failed for department deleted: %s", e)

    async def list_departments(
        self,
        *,
        keyword: str | None = None,
        parent_id: UUID | None = None,
        leader_name: str | None = None,
        page: int = 1,
        page_size: int = 20,
    ) -> tuple[list[HrDepartment], int]:
        return await self.repo.list_departments(
            keyword=keyword,
            parent_id=parent_id,
            leader_name=leader_name,
            page=page,
            page_size=page_size,
        )

    @staticmethod
    def _build_dept_tree(all_depts: list[HrDepartment]) -> list[HrDepartment]:
        """Build a sorted tree from a flat department list.

        Returns the root nodes with nested ``_children_list`` attributes,
        sorted by (workshop, priority, sort_order, name).
        """
        if not all_depts:
            return []

        dept_map = {str(d.id): d for d in all_depts}
        roots: list[HrDepartment] = []
        for dept in all_depts:
            if dept.parent_id and str(dept.parent_id) in dept_map:
                parent = dept_map[str(dept.parent_id)]
                children = getattr(parent, "_children_list", None)
                if not isinstance(children, list):
                    children = []
                    setattr(parent, "_children_list", children)
                children.append(dept)
            else:
                roots.append(dept)

        dept_priority = {"总经办": 1, "财务部": 2, "安全部": 3, "质量管理部": 4}

        def sort_key(d: HrDepartment) -> tuple[Any, ...]:
            priority = dept_priority.get(d.name, 99)
            is_workshop = 1 if "车间" in d.name else 0
            return (is_workshop, priority, d.sort_order, d.name)

        roots.sort(key=sort_key)
        for dept in all_depts:
            children = getattr(dept, "_children_list", None)
            if isinstance(children, list):
                children.sort(key=sort_key)

        return roots

    async def get_department_tree(self) -> list[dict[str, Any]]:
        """获取部门树形结构。"""
        departments = await self.repo.list_all_departments()
        roots = self._build_dept_tree(departments)
        if not roots:
            return []

        def dept_to_dict(dept: HrDepartment) -> dict[str, Any]:
            children = getattr(dept, "_children_list", [])
            hc = dept.headcount if dept.headcount is not None else None
            cc = dept.current_count if dept.current_count is not None else None
            vac = (hc - cc) if (hc is not None and cc is not None) else None
            return {
                "id": str(dept.id),
                "name": dept.name,
                "code": dept.code,
                "description": dept.description,
                "leader_name": dept.leader_name,
                "parent_id": str(dept.parent_id) if dept.parent_id else None,
                "feishu_open_department_id": dept.feishu_open_department_id,
                "sort_order": dept.sort_order,
                "headcount": hc,
                "current_count": cc,
                "vacancy": vac,
                "responsibilities": dept.responsibilities,
                "category": dept.category,
                "created_at": dept.created_at.isoformat() if dept.created_at else None,
                "updated_at": dept.updated_at.isoformat() if dept.updated_at else None,
                "children": [dept_to_dict(c) for c in children],
            }

        return [dept_to_dict(r) for r in roots]

    async def get_org_tree(self) -> list[dict[str, Any]]:
        """获取组织架构树（含人员）。"""
        from app.modules.hr.feishu.contact import FeishuContact

        departments = await self.repo.list_all_departments()
        roots = self._build_dept_tree(departments)
        if not roots:
            return []

        contact = FeishuContact()

        # 并发获取所有叶子部门的用户列表，单次超时 5 秒
        leaf_depts = [
            d
            for d in departments
            if not getattr(d, "_children_list", []) and d.feishu_open_department_id
        ]
        user_cache: dict[str, list[dict[str, Any]]] = {}

        async def _fetch_users(dept_id: str) -> None:
            try:
                users = await asyncio.wait_for(
                    contact.get_department_users(dept_id), timeout=5.0
                )
                user_cache[dept_id] = users
            except Exception as e:
                logger.warning("Failed to get users for dept %s: %s", dept_id, e)
                user_cache[dept_id] = []

        await asyncio.gather(
            *[
                _fetch_users(str(d.feishu_open_department_id))
                for d in leaf_depts
                if d.feishu_open_department_id
            ]
        )

        async def dept_to_org_node(dept: HrDepartment) -> dict[str, Any]:
            children_depts = getattr(dept, "_children_list", [])
            nodes = []
            for child in children_depts:
                nodes.append(await dept_to_org_node(child))
            if not children_depts and dept.feishu_open_department_id:
                users = user_cache.get(dept.feishu_open_department_id, [])
                for user in users:
                    # 过滤公用账号（无工号 = 非真实人员）
                    if not user.get("employee_no"):
                        continue
                    name = user.get("name", "")
                    nodes.append(
                        {
                            "id": f"emp_{user.get('open_id', '')}",
                            "name": name,
                            "type": "employee",
                            "leader_name": user.get("job_title", ""),
                            "current_count": None,
                            "headcount": None,
                            "vacancy": None,
                            "category": None,
                            "sort_order": None,
                            "children": [],
                        }
                    )
            hc = dept.headcount if dept.headcount is not None else None
            cc = dept.current_count if dept.current_count is not None else None
            vac = (hc - cc) if (hc is not None and cc is not None) else None
            return {
                "id": str(dept.id),
                "name": dept.name,
                "type": "department",
                "leader_name": dept.leader_name,
                "current_count": cc,
                "headcount": hc,
                "vacancy": vac,
                "category": dept.category,
                "sort_order": dept.sort_order,
                "children": nodes,
            }

        return [await dept_to_org_node(r) for r in roots]

    async def sync_departments_from_feishu(self) -> dict[str, Any]:
        """从飞书通讯录同步部门数据（Redis 缓存 + 增量更新）。

        仅同步本公司（本地部门树顶层节点的飞书父部门）整棵子树，
        不会把飞书租户下其他公司/集团的部门拉进来。
        """
        import json

        from app.core.redis import cache_get, cache_set
        from app.modules.hr.constants import match_department_category

        cache_key = "hr:feishu:departments"
        cache_ttl = 86400  # 24 小时缓存，部门结构变更频率低

        # 1. 定位本公司根部门：本地顶层部门（parent_id 为空）自身。
        # 从它开始 BFS 即可覆盖本公司整棵子树。
        # 注意：不能取它的飞书父部门——父部门通常是集团节点（如"丽珠医药集团股份
        # 有限公司"），会把整个集团（其他子公司）全部拉进来，导致同名部门
        # （如"总经办""财务部"）触发 code 唯一约束冲突，最终整个事务被回滚，
        # 数据库停留在旧快照，页面人数与飞书不一致。
        root_department_id: str | None = None
        all_depts_local = await self.repo.list_all_departments()
        top_depts = [d for d in all_depts_local if d.parent_id is None]
        probe = next((d for d in top_depts if d.feishu_open_department_id), None)
        if probe is not None and probe.feishu_open_department_id:
            root_department_id = probe.feishu_open_department_id
        if not root_department_id:
            logger.warning("无法定位本公司根部门，回退为全租户遍历（可能拉到其他公司）")

        # 2. 尝试从 Redis 缓存获取飞书部门数据
        cached_raw = await cache_get(cache_key)
        if cached_raw:
            departments_data = json.loads(cached_raw)
            logger.info(
                "Using cached Feishu departments (%d items)", len(departments_data)
            )
        else:
            # 缓存 miss，从飞书 BFS 获取本公司子树
            from app.platform.integrations.feishu.contact import get_all_departments

            departments_data = await get_all_departments(
                root_department_id=root_department_id or "0"
            )
            # 写入缓存
            await cache_set(
                cache_key,
                json.dumps(departments_data, ensure_ascii=False),
                ex=cache_ttl,
            )
            logger.info(
                "Fetched and cached Feishu departments (%d items)",
                len(departments_data),
            )

        stats = {
            "created": 0,
            "updated": 0,
            "skipped": 0,
            "failed": 0,
            "total": len(departments_data),
        }

        # 2. 加载本地所有部门，建立映射
        all_departments = await self.repo.list_all_departments()
        dept_map = {
            d.feishu_open_department_id: d
            for d in all_departments
            if d.feishu_open_department_id
        }
        feishu_to_local_id: dict[str, UUID] = {}
        for d in all_departments:
            if d.feishu_open_department_id:
                feishu_to_local_id[d.feishu_open_department_id] = d.id

        # 3. 增量更新：阶段1 批量计算变更 → 阶段2 并发查负责人 → 阶段3 批量写库，
        #    避免逐部门串行请求飞书导致同步过慢
        pending_updates: list[tuple[HrDepartment, str]] = []
        pending_creates: list[tuple[dict[str, Any], str]] = []
        # 按名称兜底匹配：飞书部门 ID 变化（如部门重建）时按"名称+父级"匹配并回写 ID，
        # 避免把已存在的部门误当新增 INSERT，触发 code 唯一约束并污染整个事务。
        # 注意：飞书允许同名父子部门（如"仓储部"下还有"仓储部"），同名时必须同时
        # 校验父级，否则会把不同的同名部门错误合并（导致父级缺失/自引用）。
        name_map: dict[str, list[HrDepartment]] = {}
        for d in all_departments:
            name_map.setdefault(d.name, []).append(d)
        used_codes = {d.code for d in all_departments}
        for dept_data in departments_data:
            open_dept_id = dept_data.get("department_id")
            name = dept_data.get("name", "")
            if not open_dept_id or not name:
                stats["failed"] += 1
                continue

            member_count = dept_data.get("member_count")
            sort_order = dept_data.get("order", 0)
            parent_feishu_id = dept_data.get("parent_department_id", "")

            existing = dept_map.get(open_dept_id)
            re_id = False
            if existing is None and name in name_map:
                parent_local_id = feishu_to_local_id.get(parent_feishu_id)
                # 只合并"确认是同一部门"的本地行：本地行无飞书 ID、飞书 ID 相同、
                # 或父级一致（同名父子部门场景下父级不同 → 不合并，走创建分支）
                candidate = next(
                    (
                        d
                        for d in name_map[name]
                        if not d.feishu_open_department_id
                        or d.feishu_open_department_id == open_dept_id
                        or (parent_local_id and d.parent_id == parent_local_id)
                    ),
                    None,
                )
                if candidate is None:
                    # 同名但父级不匹配（同名父子部门）→ 视为新部门，走创建分支
                    pass
                elif (
                    candidate.feishu_open_department_id
                    and candidate.feishu_open_department_id != open_dept_id
                ):
                    # 飞书 ID 变化（部门重建）但名称+父级一致：合并映射到已有本地行，
                    # 不覆盖其飞书 ID、不新建重复行
                    dept_map[open_dept_id] = candidate
                    feishu_to_local_id[open_dept_id] = candidate.id
                    stats["skipped"] += 1
                    continue
                else:
                    existing = candidate
                    existing.feishu_open_department_id = open_dept_id
                    dept_map[open_dept_id] = existing
                    feishu_to_local_id[open_dept_id] = existing.id
                    re_id = True
            if existing:
                # 增量更新：检测字段是否有变化
                changed = re_id
                if existing.name != name:
                    existing.name = name
                    changed = True
                if member_count is not None and existing.current_count != member_count:
                    existing.current_count = member_count
                    changed = True
                if sort_order is not None and existing.sort_order != sort_order:
                    existing.sort_order = sort_order
                    changed = True

                # 更新 parent_id（通过飞书 ID 映射；同名部门合并后避免自挂父级）
                if parent_feishu_id:
                    parent_local_id = feishu_to_local_id.get(parent_feishu_id)
                    if (
                        parent_local_id
                        and parent_local_id != existing.id
                        and existing.parent_id != parent_local_id
                    ):
                        existing.parent_id = parent_local_id
                        changed = True
                elif existing.parent_id is not None:
                    # 飞书无上级，本地有上级 → 清空
                    existing.parent_id = None
                    changed = True

                if changed:
                    category = match_department_category(name)
                    if category is not None and existing.category != category:
                        existing.category = category
                    pending_updates.append(
                        (existing, dept_data.get("leader_user_id", ""))
                    )
                else:
                    stats["skipped"] += 1
            else:
                pending_creates.append((dept_data, dept_data.get("leader_user_id", "")))

        # 阶段2：并发查询负责人姓名
        leader_ids = {lid for _, lid in pending_updates if lid} | {
            lid for _, lid in pending_creates if lid
        }
        leader_names: dict[str, str | None] = {}
        if leader_ids:
            from app.modules.hr.feishu.contact import FeishuContact

            contact = FeishuContact()
            sem = asyncio.Semaphore(10)

            async def _fetch_leader_name(uid: str) -> tuple[str, str | None]:
                async with sem:
                    try:
                        return uid, await contact.get_user_name(uid)
                    except Exception as e:
                        logger.warning("Failed to get user name for %s: %s", uid, e)
                        return uid, None

            for uid, lname in await asyncio.gather(
                *[_fetch_leader_name(u) for u in leader_ids]
            ):
                leader_names[uid] = lname

        # 阶段3：写库（单条失败用 savepoint 回滚，不影响其他写入；
        # 不能用 session.rollback()——它会回滚整个事务，丢失之前所有已 flush 的变更）
        for existing, leader_user_id in pending_updates:
            try:
                leader_name = leader_names.get(leader_user_id)
                if leader_name and existing.leader_name != leader_name:
                    existing.leader_name = leader_name
                async with self.session.begin_nested():
                    await self.repo.update(existing)
                stats["updated"] += 1
            except Exception as e:
                logger.error("Failed to sync department %s: %s", existing.name, e)
                stats["failed"] += 1

        for dept_data, leader_user_id in pending_creates:
            try:
                # 新增部门
                from uuid import uuid4

                open_dept_id = dept_data.get("department_id")
                if not isinstance(open_dept_id, str) or not open_dept_id:
                    continue
                name = dept_data.get("name", "")
                # code 需唯一（飞书允许同名父子部门），同名时用飞书 ID 短码做后缀
                code = name
                if code in used_codes:
                    code = f"{name}-{open_dept_id[-8:]}"
                used_codes.add(code)
                new_dept = HrDepartment(
                    id=uuid4(),
                    name=name,
                    code=code,
                    leader_name=leader_names.get(leader_user_id),
                    feishu_open_department_id=open_dept_id,
                    current_count=dept_data.get("member_count"),
                    category=match_department_category(name),
                    sort_order=dept_data.get("order", 0) or 0,
                )
                # 设置 parent_id
                parent_feishu_id = dept_data.get("parent_department_id", "")
                if parent_feishu_id:
                    parent_local_id = feishu_to_local_id.get(parent_feishu_id)
                    if parent_local_id:
                        new_dept.parent_id = parent_local_id

                async with self.session.begin_nested():
                    await self.repo.create(new_dept)
                    feishu_to_local_id[open_dept_id] = new_dept.id
                stats["created"] += 1
            except Exception as e:
                logger.error(
                    "Failed to sync department %s: %s", dept_data.get("name"), e
                )
                stats["failed"] += 1

        # 4. 第二轮：补全新增部门的 parent_id（新增部门的父部门可能也是本轮新增的）
        for dept_data in departments_data:
            try:
                open_dept_id = dept_data.get("department_id")
                parent_feishu_id = dept_data.get("parent_department_id", "")
                if not open_dept_id or not parent_feishu_id:
                    continue
                local_id = feishu_to_local_id.get(open_dept_id)
                parent_local_id = feishu_to_local_id.get(parent_feishu_id)
                if not local_id or not parent_local_id:
                    continue
                existing = dept_map.get(open_dept_id)
                if existing and existing.parent_id != parent_local_id:
                    existing.parent_id = parent_local_id
                    await self.repo.update(existing)
            except Exception as e:
                logger.error(
                    "Failed to set parent_id for %s: %s", dept_data.get("name"), e
                )

        # 5. 物理清理：彻底删除不在本公司子树内的部门（如历史误同步进来的其他公司/集团
        # 部门），
        #    以及所有历史软删除残留，确保不留下其他公司的部门数据
        feishu_ids = {d["department_id"] for d in departments_data}
        stale_ids = [
            d.id
            for d in all_depts_local
            if d.feishu_open_department_id
            and d.feishu_open_department_id not in feishu_ids
            and not d.is_deleted
        ]
        try:
            if stale_ids:
                # 先解除被删部门的子部门父级引用，避免外键约束失败
                await self.session.execute(
                    update(HrDepartment)
                    .where(HrDepartment.parent_id.in_(stale_ids))
                    .values(parent_id=None)
                )
                await self.session.execute(
                    update(HrDepartment)
                    .where(HrDepartment.id.in_(stale_ids))
                    .values(is_deleted=True)
                )
                await self.session.flush()
                logger.info(
                    "Soft-deleted %d stale departments not in company subtree",
                    len(stale_ids),
                )
            # 历史软删除残留改为仅记录日志，不再物理删除
            result = await self.session.execute(
                select(func.count())
                .select_from(HrDepartment)
                .where(HrDepartment.is_deleted.is_(True))
            )
            leftover_count = result.scalar() or 0
            if leftover_count:
                logger.info(
                    "Found %d soft-deleted leftovers (skipped physical deletion)",
                    leftover_count,
                )
        except Exception as e:
            await self.session.rollback()
            logger.error("Failed to hard-clean stale departments: %s", e)

        return stats


class TeamService:
    def __init__(self, session: AsyncSession) -> None:
        self.repo = TeamRepository(session)
        self.department_repo = DepartmentRepository(session)

    async def get_team(self, team_id: UUID) -> Team:
        team = await self.repo.get_by_id(team_id)
        if not team:
            raise NotFoundException("班组", str(team_id))
        return team

    async def create_team(self, data: TeamCreate) -> Team:
        department = await self.department_repo.get_by_id(data.department_id)
        if not department:
            raise NotFoundException("部门", str(data.department_id))

        team = Team(**data.model_dump())
        result = await self.repo.create(team)
        return result

    async def update_team(self, team_id: UUID, data: TeamUpdate) -> Team:
        team = await self.get_team(team_id)
        update_data = data.model_dump(exclude_unset=True)

        if "department_id" in update_data:
            department = await self.department_repo.get_by_id(
                update_data["department_id"]
            )
            if not department:
                raise NotFoundException("部门", str(update_data["department_id"]))

        for field, value in update_data.items():
            setattr(team, field, value)

        result = await self.repo.update(team)
        return result

    async def delete_team(self, team_id: UUID) -> None:
        team = await self.get_team(team_id)
        await self.repo.soft_delete(team)

    async def list_teams(
        self,
        *,
        department_id: UUID | None = None,
        keyword: str | None = None,
        page: int = 1,
        page_size: int = 20,
    ) -> tuple[list[Team], int]:
        return await self.repo.list_teams(
            department_id=department_id,
            keyword=keyword,
            page=page,
            page_size=page_size,
        )


class OffboardingRecordService:
    def __init__(self, session: AsyncSession) -> None:
        self.repo = OffboardingRecordRepository(session)
        self.employee_repo = EmployeeRepository(session)
        self.session = session

    async def get_record(self, record_id: UUID) -> OffboardingRecord:
        record = await self.repo.get_by_id(record_id)
        if not record:
            raise NotFoundException("离职记录", str(record_id))
        return record

    async def create_record(self, data: OffboardingRecordCreate) -> OffboardingRecord:
        if not hasattr(self, "session"):
            legacy_employee_id = data.employee_id
            if legacy_employee_id is None:
                raise NotFoundException("员工", "未知")
            employee = await self.employee_repo.get_by_id(legacy_employee_id)
            if not employee:
                raise NotFoundException("员工", str(data.employee_id))
            record = OffboardingRecord(**data.model_dump())
            result = await self.repo.create(record)
            employee.status = "离职"
            await self.employee_repo.update(employee)
            try:
                legacy_feishu = getattr(self, "feishu", None)
                if legacy_feishu is not None:
                    await legacy_feishu.sync_offboarding_created(result.__dict__)
            except Exception as exc:
                logger.warning("Feishu sync failed for offboarding created: %s", exc)
            return result

        employee = None
        employee_id = getattr(data, "employee_id", None)
        if employee_id is not None:
            employee = await self.employee_repo.get_by_id(employee_id)
        employee_number = getattr(data, "employee_number", None)
        if not employee and employee_number:
            employee = await self.employee_repo.get_by_employee_number(employee_number)
        if not employee:
            raise NotFoundException(
                "员工",
                str(
                    getattr(data, "employee_id", None)
                    or getattr(data, "employee_number", None)
                ),
            )

        record = OffboardingRecord(**data.model_dump())
        if not record.employee_id:
            record.employee_id = employee.id
        record = await self.repo.create(record)

        # 自动将员工状态更新为离职
        employee.status = "离职"
        await self.employee_repo.update(employee)

        # 同步到飞书（失败不影响本地创建）
        try:
            logger.info(
                "开始同步离职记录到飞书: employee_number=%s, name=%s",
                record.employee_number,
                record.name,
                extra={"hr_module": "hr"},
            )
            await self._sync_to_feishu(record, employee, is_create=True)
            logger.info(
                "离职记录同步到飞书成功: employee_number=%s, feishu_record_id=%s",
                record.employee_number,
                record.feishu_record_id,
                extra={"hr_module": "hr"},
            )
        except Exception as e:
            logger.error(
                "飞书同步失败，但离职记录已创建: employee_number=%s, error=%s",
                record.employee_number,
                str(e),
                extra={"employee_id": str(employee.id), "hr_module": "hr"},
                exc_info=True,
            )

        # 自动发送离职材料
        try:
            if employee.feishu_open_id:
                from app.platform.integrations.feishu import notification as feishu_n

                content = """请按模版规范填写以下离职材料：

• 附表6：员工离职申请单
• 附表7：员工离职面谈表
• 附表8：员工离职清退表

（请到人力资源部领取纸质表单）

来自：离职资料领取登记表"""

                await feishu_n.send_user_card(
                    open_id=employee.feishu_open_id,
                    title="离职单填写模版请查收！",
                    content=content,
                )
                record.materials_sent = True
                record.materials_sent_at = datetime.now()
                await self.repo.update(record)
                logger.info(
                    "已向员工 %s 发送离职材料",
                    employee.name,
                    extra={"employee_id": str(employee.id), "hr_module": "hr"},
                )
        except Exception:
            logger.warning(
                "发送离职材料失败",
                extra={"employee_id": str(employee.id), "hr_module": "hr"},
                exc_info=True,
            )

        return record

    async def update_record(
        self, record_id: UUID, data: OffboardingRecordUpdate
    ) -> OffboardingRecord:
        record = await self.get_record(record_id)
        update_data = data.model_dump(exclude_unset=True)
        for field, value in update_data.items():
            setattr(record, field, value)
        result = await self.repo.update(record)

        if not hasattr(self, "session"):
            try:
                legacy_feishu = getattr(self, "feishu", None)
                if legacy_feishu is not None:
                    await legacy_feishu.sync_offboarding_updated(result.__dict__)
            except Exception as exc:
                logger.warning("Feishu sync failed for offboarding updated: %s", exc)
            return result

        # 同步到飞书
        employee = None
        if record.employee_id:
            try:
                employee = await self.employee_repo.get_by_id(record.employee_id)
            except Exception:
                logger.exception("update_record 查找员工失败，将跳过飞书同步")
        await self._sync_to_feishu(result, employee, is_create=False)
        return result

    async def delete_record(self, record_id: UUID) -> None:
        record = await self.get_record(record_id)
        if not hasattr(self, "session"):
            await self.repo.soft_delete(record)
            return
        await self._delete_from_feishu(record)
        await self.repo.soft_delete(record)

    async def generate_termination_certificate(
        self, record_id: UUID
    ) -> tuple[BytesIO, str, OffboardingRecord]:
        """生成解除劳动合同通知单，返回 (文件缓冲区, 文件名, 离职记录)"""
        from app.modules.hr.offboarding_document_generator import (
            generate_termination_notice,
        )

        record = await self.get_record(record_id)

        employee = None
        if record.employee_id:
            employee = await self.employee_repo.get_by_id(record.employee_id)

        if not employee:
            raise NotFoundException(
                "员工", str(record.employee_id) if record.employee_id else "未知"
            )

        employee_data = {
            "name": employee.name or "",
            "gender": employee.gender or "",
            "id_card": employee.id_card or "",
            "hire_date": employee.hire_date.strftime("%Y年%m月%d日")
            if employee.hire_date
            else "",
            "current_address": employee.current_address or "",
        }

        doc_buffer = generate_termination_notice(employee_data)

        # 上传到 MinIO
        from app.core.storage import is_enabled, upload_object

        filename = f"{employee.name}_解除劳动合同通知单.docx"
        if is_enabled():
            upload_object(
                "hr",
                f"offboarding/{filename}",
                doc_buffer.read(),
                doc_buffer.getbuffer().nbytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            doc_buffer.seek(0)

        # Update completion status
        record.completed_date = date.today()
        record.handover_status = "已完成"
        await self.repo.update(record)

        return doc_buffer, filename, record

    # ─── Feishu 读写方法 ───

    async def _get_offboarding_bitable(self) -> Any:
        from app.modules.hr.feishu.bitable import BitableClient
        from app.modules.hr.feishu_settings_service import _get_entity_prefill

        result = await self.session.execute(
            select(HrFeishuEntitySetting).where(
                HrFeishuEntitySetting.entity_code == "offboarding_record",
                HrFeishuEntitySetting.is_deleted.is_(False),
            )
        )
        entity = result.scalar_one_or_none()

        # Fallback 到默认配置
        prefill = _get_entity_prefill("offboarding_record")
        app_token = (
            entity.app_token
            if entity and entity.app_token
            else prefill.get("app_token")
        )
        table_id = (
            entity.base_table_id
            if entity and entity.base_table_id
            else prefill.get("table_id")
        )

        if not app_token or not table_id:
            logger.warning(
                (
                    "离职记录飞书实体 app_token 或 base_table_id 未配置，无"
                    "法同步到飞书多维表格"
                ),
                extra={"hr_module": "hr"},
            )
            return None
        app_id, app_secret = await get_hr_feishu_app_credentials(self.session)
        return (
            BitableClient(
                app_token=app_token,
                app_id=app_id or None,
                app_secret=app_secret or None,
            ),
            table_id,
        )

    async def _sync_to_feishu(
        self, record: OffboardingRecord, employee: Any = None, *, is_create: bool
    ) -> None:
        """同步离职记录到飞书多维表格，推送全部字段"""
        from app.modules.hr.feishu.bitable import _to_ms_timestamp as _dt

        pair = await self._get_offboarding_bitable()
        if not pair:
            logger.warning(
                "飞书多维表格未配置，跳过离职记录同步: employee_number=%s",
                record.employee_number,
                extra={"hr_module": "hr"},
            )
            return
        client, table_id = pair
        logger.info(
            "已获取飞书离职表配置: employee_number=%s, table_id=%s",
            record.employee_number,
            table_id,
            extra={"hr_module": "hr"},
        )

        # 字段配置：(飞书字段名, 本地值, type=文本/数字/日期)
        # 文本直接写字符串，日期写 ms 时间戳，数字写 int
        def _s(v: Any) -> Any:
            return v if v else ""

        def _d(v: Any) -> Any:
            return _dt(v) if v else ""

        def _n(v: Any) -> Any:
            """转数字，飞书 number 字段需要 int"""
            if v is None or v == "":
                return None
            if isinstance(v, int):
                return v
            if isinstance(v, str) and v.isdigit():
                return int(v)
            return None

        # 组合出生年月
        birth_str = ""
        if record.birth_year and record.birth_month:
            birth_str = f"{record.birth_year}-{record.birth_month:02d}"
            if record.birth_day:
                birth_str = (
                    f"{record.birth_year}-{record.birth_month:02d}-"
                    f"{record.birth_day:02d}"
                )

        # 技能证书 JSON 数组转字符串
        qualifications_str = ""
        if record.qualifications:
            if isinstance(record.qualifications, list):
                qualifications_str = "、".join(record.qualifications)
            else:
                qualifications_str = str(record.qualifications)

        # 在职状态（从关联的 employee 读取）
        status_str = ""
        if employee and hasattr(employee, "status"):
            status_str = _s(employee.status)

        # 拟转正日期（从关联的 employee 读取，OffboardingRecord 没有此字段）
        planned_probation_date_str = ""
        if (
            employee
            and hasattr(employee, "planned_probation_date")
            and employee.planned_probation_date
        ):
            planned_probation_date_str = _d(employee.planned_probation_date)

        fields_config = [
            # 基本信息 - Number 类型
            ("序号", _n(record.seq_number), "number"),
            ("工号", _n(record.employee_number), "number"),
            ("联系电话", _n(record.phone), "number"),
            ("紧急联系人电话", _n(record.emergency_contact_phone), "number"),
            # 基本信息 - Text 类型
            ("姓名", _s(record.name), "text"),
            ("域账户", _s(record.domain_account), "text"),
            ("性别", _s(record.gender), "text"),
            ("民族", _s(record.ethnic_group), "text"),
            ("籍贯", _s(record.native_place), "text"),
            ("政治面貌", _s(record.political_status), "text"),
            ("婚姻状况", _s(record.marital_status), "text"),
            ("健康状况", _s(record.health_status), "text"),
            ("户口类别", _s(record.household_type), "text"),
            ("人员类别", _s(record.status_category), "text"),
            ("身份证号码", _s(record.id_card), "text"),
            ("身份证有效期截止日期", _s(record.id_card_expiry), "text"),
            ("现家庭住址", _s(record.current_address), "text"),
            ("电子邮箱", _s(record.email), "text"),
            ("紧急联系人", _s(record.emergency_contact_name), "text"),
            ("与本人关系", _s(record.emergency_contact_relation), "text"),
            ("年龄", _s(record.age) if record.age is not None else "", "text"),
            ("出生年月", birth_str, "text"),
            ("在职状态", status_str, "text"),
            # 组织信息
            ("一级部门", _s(record.department), "text"),
            ("二级部门", _s(record.sub_department), "text"),
            ("职位/岗位", _s(record.position), "text"),
            ("职级", _s(record.level), "text"),
            ("人员就业方式", _s(record.employment_type), "text"),
            ("转正状态", _s(record.probation_status), "text"),
            (
                "转正生效日期",
                record.probation_effective_date.isoformat()
                if isinstance(record.probation_effective_date, date)
                else (
                    _s(record.probation_effective_date)
                    if record.probation_effective_date
                    else ""
                ),
                "text",
            ),
            ("拟转正日期", planned_probation_date_str, "date"),
            # 工作经历
            ("工作经验一", _s(record.work_experience_1), "text"),
            ("工作经验二", _s(record.work_experience_2), "text"),
            ("工作经验三", _s(record.work_experience_3), "text"),
            ("工作经验四", _s(record.work_experience_4), "text"),
            # 日期字段 - Datetime 类型
            ("入职日期", _d(record.hire_date), "date"),
            ("参加工作时间", _d(record.work_start_date), "date"),
            ("进入本公司时间", _d(record.factory_entry_date), "date"),
            ("工龄", _s(record.work_years), "text"),
            # 离职专属
            ("最后工作日", _d(record.offboarding_date), "date"),
            ("离职类型", _s(record.offboarding_type), "text"),
            ("离职原因", _s(record.reason), "text"),
            ("备注", _s(record.notes), "text"),
            # 档案与证书
            ("档案号", _s(record.archive_number), "text"),
            ("技能证书", qualifications_str, "text"),
            ("证书编号", _s(record.certificate_number), "text"),
            ("技能证书复审时间", _d(record.certificate_review_date), "date"),
            # 教育信息
            ("学历", _s(record.education), "text"),
            ("学位", _s(record.degree), "text"),
            ("毕业院校", _s(record.school), "text"),
            ("专业", _s(record.major), "text"),
            ("毕业时间", _d(record.graduation_date), "date"),
            ("职称", _s(record.qualification_type), "text"),
            # 合同日期 - Datetime 类型（首次和第二次）
            ("首次签订合同日期", _d(record.contract_start_date), "date"),
            ("首次签订合同截止日期", _d(record.contract_end_date), "date"),
            ("第二次续签合同日期", _d(record.contract_start_2), "date"),
            # 合同日期 - Text 类型（第三次及以后）
            (
                "合同截止日期",
                _s(record.contract_end_date)
                if isinstance(record.contract_end_date, str)
                else (
                    record.contract_end_date.isoformat()
                    if record.contract_end_date
                    else ""
                ),
                "text",
            ),
            ("合同截止日期2", _s(record.contract_end_2), "text"),
            ("第三次续签合同日期", _s(record.contract_start_3), "text"),
            ("合同截止日期3", _s(record.contract_end_3), "text"),
            ("第四次续签合同日期", _s(record.contract_start_4), "text"),
            ("合同截止日期4", _s(record.contract_end_4), "text"),
            ("第五次续签合同日期", _s(record.contract_start_5), "text"),
            ("合同截止日期5", _s(record.contract_end_5), "text"),
            ("第六次续签合同日期", _s(record.contract_start_6), "text"),
        ]

        fields = {}
        for name, value, ftype in fields_config:
            if value is None:
                continue
            if ftype == "number" and value == 0:
                continue  # 数字 0 不写入（可能是空值）
            if ftype in ("text", "date") and value == "":
                continue
            fields[name] = value

        logger.info(
            "准备写入飞书离职表: employee_number=%s, fields_count=%d",
            record.employee_number,
            len(fields),
            extra={"hr_module": "hr"},
        )

        if record.feishu_record_id:
            logger.info(
                "更新飞书离职记录: employee_number=%s, feishu_record_id=%s",
                record.employee_number,
                record.feishu_record_id,
                extra={"hr_module": "hr"},
            )
            await client.update_record(table_id, record.feishu_record_id, fields)
        else:
            logger.info(
                "创建飞书离职记录: employee_number=%s",
                record.employee_number,
                extra={"hr_module": "hr"},
            )
            rec = await client.create_record(table_id, fields)
            record.feishu_record_id = rec.get("record_id")
            record.feishu_synced_at = date.today()
            await self.repo.update(record)
            logger.info(
                "飞书离职记录创建成功: employee_number=%s, feishu_record_id=%s",
                record.employee_number,
                record.feishu_record_id,
                extra={"hr_module": "hr"},
            )

    async def _delete_from_feishu(self, record: OffboardingRecord) -> None:
        if not record.feishu_record_id:
            return
        pair = await self._get_offboarding_bitable()
        if not pair:
            return
        client, table_id = pair
        try:
            await client.delete_record(table_id, record.feishu_record_id)
        except Exception as e:
            logger.error(
                "Feishu delete failed for offboarding record %s: %s",
                record.id,
                str(e),
                extra={"hr_module": "hr"},
                exc_info=True,
            )
            raise

    async def list_records(
        self,
        *,
        employee_id: UUID | None = None,
        keyword: str | None = None,
        page: int = 1,
        page_size: int = 20,
        dept_alias_set: set[str] | None = None,
    ) -> tuple[list[OffboardingRecord], int]:
        return await self.repo.list_records(
            employee_id=employee_id,
            keyword=keyword,
            page=page,
            page_size=page_size,
            dept_alias_set=dept_alias_set,
        )

    async def sync_from_feishu(self) -> dict[str, Any]:
        """Pull all records from Feishu Bitable and upsert into local PG.
        Records in local DB but not in Feishu will be soft-deleted.

        Returns:
            {"created": N, "updated": N, "deleted": N, "failed": N, "total": N}
        """
        query = select(HrFeishuEntitySetting).where(
            HrFeishuEntitySetting.entity_code == "offboarding_record",
            HrFeishuEntitySetting.is_enabled.is_(True),
        )
        result = await self.session.execute(query)
        setting = result.scalar_one_or_none()
        if not setting or not setting.app_token or not setting.base_table_id:
            raise AppException(message="离职记录飞书多维表格未配置或未启用")

        from app.modules.hr.feishu.bitable import BitableClient

        app_id, app_secret = await get_hr_feishu_app_credentials(self.session)
        client = BitableClient(
            app_token=setting.app_token,
            app_id=app_id or None,
            app_secret=app_secret or None,
        )
        raw_records = await client.search_records(
            setting.base_table_id,
            page_size=500,
        )
        stats = {
            "created": 0,
            "updated": 0,
            "deleted": 0,
            "failed": 0,
            "total": len(raw_records),
        }
        feishu_record_ids: set[str] = set()

        for rec in raw_records:
            try:
                fields = rec.get("fields", {})
                rid = rec.get("record_id", "")
                if not rid:
                    stats["failed"] += 1
                    continue

                feishu_record_ids.add(rid)

                def gt(key: str) -> Any:
                    return fields.get(key)

                # Try to resolve employee_id from employee_number
                emp_no = _extract_text(gt("工号"))
                employee_id = None
                if emp_no:
                    emp = await self.employee_repo.get_by_employee_number(emp_no)
                    if emp:
                        employee_id = emp.id

                data = {
                    "feishu_record_id": rid,
                    "employee_id": employee_id,
                    "seq_number": _extract_number(gt("序号")),
                    "employee_number": emp_no,
                    "name": _extract_text(gt("姓名")),
                    "domain_account": _extract_text(gt("域账号")),
                    "gender": _extract_text(gt("性别")),
                    "ethnic_group": _extract_text(gt("民族")),
                    "native_place": _extract_text(gt("籍贯")),
                    "political_status": _extract_text(gt("政治面貌")),
                    "marital_status": _extract_text(gt("婚姻状况")),
                    "health_status": _extract_text(gt("健康状况")),
                    "household_type": _extract_text(gt("户籍类型")),
                    "status_category": _extract_text(gt("统计类别")),
                    "id_card": _extract_text(gt("身份证号")),
                    "id_card_expiry": _extract_text(gt("身份证到期日")),
                    "current_address": _extract_text(gt("现住址")),
                    "phone": _extract_text(gt("手机")),
                    "email": _extract_text(gt("邮箱地址")),
                    "emergency_contact_name": _extract_text(gt("紧急联系人")),
                    "emergency_contact_phone": _extract_text(gt("紧急联系人电话")),
                    "emergency_contact_relation": _extract_text(gt("紧急联系人关系")),
                    "department": _extract_text(gt("一级部门")),
                    "sub_department": _extract_text(gt("二级部门")),
                    "position": _extract_text(gt("职务|岗位")),
                    "level": _extract_text(gt("职级")),
                    "employment_type": _extract_text(gt("人员就业方式")),
                    "probation_status": _extract_text(gt("转正状态")),
                    "probation_effective_date": _parse_date(gt("转正生效日期")),
                    "hire_date": _parse_date(gt("进厂时间")),
                    "work_start_date": _parse_date(gt("参加工作时间")),
                    "factory_entry_date": _parse_date(gt("进厂时间")),
                    "work_years": _extract_text(gt("工龄"))
                    or _extract_text(gt("工作年限")),
                    "offboarding_date": _parse_date(gt("最后工作日")),
                    "offboarding_type": _extract_text(gt("离职类型")) or "辞职",
                    "reason": _extract_text(gt("离职原因")),
                    "handover_status": _extract_text(gt("交接状态")) or "待交接",
                    "education": _extract_text(gt("学历")),
                    "degree": _extract_text(gt("学位")),
                    "major": _extract_text(gt("专业")),
                    "school": _extract_text(gt("毕业学校")),
                    "graduation_date": _parse_date(gt("毕业时间")),
                    "qualification_type": _extract_text(gt("职称类型")),
                    "qualifications": gt("职称／职业资格")
                    if isinstance(gt("职称／职业资格"), list)
                    else None,
                    "certificate_number": _extract_text(gt("证书编号")),
                    "certificate_review_date": _parse_date(gt("技能证书复审时间")),
                    "contract_start_date": _parse_date(gt("第一次合同起点时间")),
                    "contract_end_date": _parse_date(gt("第一次合同终止时间")),
                    "contract_end_2": _extract_text(gt("合同截止日期2")),
                    "contract_end_3": _extract_text(gt("合同截止日期3")),
                    "contract_end_4": _extract_text(gt("合同截止日期4")),
                    "contract_end_5": _extract_text(gt("合同截止日期5")),
                    "contract_start_2": _parse_date(gt("第二次合同起点时间")),
                    "contract_start_3": _extract_text(gt("第三次续签合同日期")),
                    "contract_start_4": _extract_text(gt("第四次续签合同日期")),
                    "contract_start_5": _extract_text(gt("第五次续签合同日期")),
                    "contract_start_6": _extract_text(gt("第六次续签合同日期")),
                    "work_experience_1": _extract_text(gt("工作经验一")),
                    "work_experience_2": _extract_text(gt("工作经验二")),
                    "work_experience_3": _extract_text(gt("工作经验三")),
                    "work_experience_4": _extract_text(gt("工作经验四")),
                    "archive_number": _extract_text(gt("档案编号")),
                    "notes": _extract_text(gt("备注")),
                    "feishu_synced_at": date.today(),
                }

                # Upsert by feishu_record_id（以飞书为主：空值也覆盖本地旧值）
                existing = await self.repo.get_by_feishu_record_id(rid)
                if existing:
                    for key, value in data.items():
                        if key != "id":
                            setattr(existing, key, value)
                    stats["updated"] += 1
                else:
                    record = OffboardingRecord(**data)
                    self.session.add(record)
                    stats["created"] += 1
            except Exception:
                logger.exception(
                    "Failed to sync offboarding record %s", rec.get("record_id")
                )
                stats["failed"] += 1

        # 软删除飞书中不存在的本地记录（以飞书为主）
        local_records = await self.repo.list_all()
        for local in local_records:
            if (
                local.feishu_record_id
                and local.feishu_record_id not in feishu_record_ids
            ):
                local.is_deleted = True
                stats["deleted"] += 1

        await self.session.commit()
        return stats


class PositionTransferRecordService:
    """岗位调动管理服务"""

    def __init__(self, session: AsyncSession) -> None:
        self.repo = PositionTransferRecordRepository(session)
        self.employee_repo = EmployeeRepository(session)
        self.session = session

    async def get_record(self, record_id: UUID) -> PositionTransferRecord:
        record = await self.repo.get_by_id(record_id)
        if not record:
            raise NotFoundException("岗位调动记录", str(record_id))
        return record

    async def list_records(
        self,
        *,
        employee_id: UUID | None = None,
        approval_status: str | None = None,
        keyword: str | None = None,
        page: int = 1,
        page_size: int = 20,
        dept_alias_set: set[str] | None = None,
    ) -> tuple[list[PositionTransferRecord], int]:
        return await self.repo.list_records(
            employee_id=employee_id,
            approval_status=approval_status,
            keyword=keyword,
            page=page,
            page_size=page_size,
            dept_alias_set=dept_alias_set,
        )

    async def create_record(
        self, data: PositionTransferRecordCreate
    ) -> PositionTransferRecord:
        """飞书直连模式：先写飞书，再创建本地缓存。"""
        record = PositionTransferRecord(**data.model_dump())
        # 先写飞书多维表格
        await self._sync_to_feishu(record, is_create=True)
        # 再创建本地缓存
        result = await self.repo.create(record)
        return result

    async def update_record(
        self, record_id: UUID, data: PositionTransferRecordUpdate
    ) -> PositionTransferRecord:
        """飞书直连模式：先写飞书，再更新本地缓存。"""
        record = await self.get_record(record_id)
        for field, value in data.model_dump(exclude_unset=True).items():
            setattr(record, field, value)
        # 先写飞书多维表格
        await self._sync_to_feishu(record, is_create=False)
        # 再更新本地缓存
        result = await self.repo.update(record)
        return result

    async def delete_record(self, record_id: UUID) -> None:
        """飞书直连模式：先删飞书，再本地软删除。"""
        record = await self.get_record(record_id)
        await self._delete_from_feishu(record)
        await self.repo.soft_delete(record)

    # ─── Feishu 直连同步 ───

    async def _get_position_bitable(self) -> Any:
        from app.modules.hr.feishu.bitable import BitableClient

        result = await self.session.execute(
            select(HrFeishuEntitySetting).where(
                HrFeishuEntitySetting.entity_code == "position_transfer",
                HrFeishuEntitySetting.is_deleted.is_(False),
            )
        )
        entity = result.scalar_one_or_none()
        if not entity or not entity.app_token or not entity.base_table_id:
            return None
        return BitableClient(app_token=entity.app_token), entity.base_table_id

    async def _build_feishu_fields(
        self, record: PositionTransferRecord
    ) -> dict[str, Any]:
        """构建飞书多维表格字段映射（飞书为主源）。

        飞书字段类型：
        - type=1 文本：传字符串
        - type=5 日期：传毫秒时间戳（整数），空值不传
        - type=11 人员：传 [{"id": "ou_xxx"}]，找不到不传
        - type=13 电话：传字符串，空值不传
        """
        from app.platform.integrations.feishu.bitable import _to_ms_timestamp

        fields: dict[str, Any] = {
            "申请人": record.employee_name or "",
            "原部门": record.department_before or "",
            "原职位": record.original_position or "",
            "申请部门": record.apply_department or "",
            "申请职位": record.apply_position or "",
            "申请人确认说明": record.applicant_confirmation_text or "",
        }
        # 电话号码类型（type=13）：空值不传
        if record.contact_phone:
            fields["联系电话"] = record.contact_phone
        # 日期字段（type=5）：有值才传毫秒时间戳
        if record.effective_date:
            fields["生效日期"] = _to_ms_timestamp(record.effective_date)
            fields["申请日期"] = _to_ms_timestamp(record.effective_date)
        if record.applicant_confirmation_date:
            fields["申请人确认日期"] = _to_ms_timestamp(
                record.applicant_confirmation_date
            )
        # 人员字段（type=11）：传 [{"id": open_id}]
        if record.applicant_signature:
            open_id = await self._get_open_id_by_name(record.applicant_signature)
            if open_id:
                fields["申请人签名"] = [{"id": open_id}]

        # 审批进度同步到多维表格
        if record.approval_flow:
            fields.update(await self._build_approval_fields(record.approval_flow))

        return fields

    # 审批节点 → 飞书多维表格字段映射
    # (意见字段, 签名字段, 签名是否人员类型type=11, 日期字段)
    _APPROVAL_FIELD_MAP = {
        "origin_direct_leader": (
            "原部门直属领导意见",
            "原部门直属领导签名",
            False,
            "原部门直属领导日期",
        ),
        "origin_manager": (
            "原部门/经理总监意见",
            "原部门/经理总监签名",
            False,
            "原部门/经理总监日期",
        ),
        "origin_director": (
            "原部门/经理总监意见",
            "原部门/经理总监签名",
            False,
            "原部门/经理总监日期",
        ),
        "origin_vp": (
            "原部门分管领导意见",
            "原部门分管领导签名",
            False,
            "原部门分管领导日期",
        ),
        "target_direct_leader": (
            "接收部门直属领导意见",
            "接收部门直属领导签名",
            False,
            "接收部门直属领导日期",
        ),
        "target_manager": (
            "接收部门经理/总监意见",
            "接收部门/经理总监签名",
            True,
            "接收部门经理/总监日期",
        ),
        "target_director": (
            "接收部门经理/总监意见",
            "接收部门/经理总监签名",
            True,
            "接收部门经理/总监日期",
        ),
        "target_vp": (
            "接收部门分管领导意见",
            "接收部门分管领导签名",
            True,
            "接收部门分管领导日期",
        ),
        "hr": ("人力资源部意见", "人力资源部签名", True, "人力资源部日期"),
        "executive_vp": (
            "常务副总经理意见",
            "常务副总经理签名",
            True,
            "常务副总经理日期",
        ),
        "general_manager": ("总经理意见", "总经理签名", True, "总经理日期"),
    }

    async def _build_approval_fields(self, flow: dict[str, Any]) -> dict[str, Any]:
        """从 approval_flow 构建审批进度字段，同步到多维表格。"""
        from datetime import datetime as _dt

        from app.platform.integrations.feishu.bitable import _to_ms_timestamp

        fields: dict[str, Any] = {}
        # 防止同字段被多个审批节点覆盖（origin_manager/origin_dir
        # ector 等映射同一飞书列）
        written: set[str] = set()
        for step in flow.get("steps", []):
            if step["status"] not in ("approved", "rejected"):
                continue
            node = step.get("node", "")
            mapping = self._APPROVAL_FIELD_MAP.get(node)
            if not mapping:
                continue
            opinion_field, sig_field, sig_is_person, date_field = mapping

            # 意见
            opinion = step.get("opinion", "")
            if opinion and opinion_field not in written:
                fields[opinion_field] = opinion
                written.add(opinion_field)

            # 签名
            signer = step.get("signer", "")
            signer_open_id = step.get("signer_open_id", "")
            if sig_field not in written:
                if sig_is_person and signer_open_id:
                    fields[sig_field] = [{"id": signer_open_id}]
                elif signer:
                    fields[sig_field] = signer
                if signer or signer_open_id:
                    written.add(sig_field)

            # 日期
            step_date = step.get("date", "")
            if step_date and date_field not in written:
                try:
                    d = _dt.strptime(step_date, "%Y.%m.%d").date()
                    fields[date_field] = _to_ms_timestamp(d)
                    written.add(date_field)
                except (ValueError, TypeError):
                    pass

        return fields

    async def _sync_to_feishu(
        self, record: PositionTransferRecord, *, is_create: bool
    ) -> None:
        """写飞书多维表格（主源）。"""
        try:
            pair = await self._get_position_bitable()
            if not pair:
                logger.warning(
                    "Feishu bitable not configured for position_transfer, skipping sync"
                )
                return
            client, table_id = pair
            fields = await self._build_feishu_fields(record)
            logger.info("飞书同步字段数: %d", len(fields))

            if record.feishu_record_id:
                await client.update_record(table_id, record.feishu_record_id, fields)
            else:
                rec = await client.create_record(table_id, fields)
                record.feishu_record_id = rec.get("record_id")
                record.feishu_synced_at = date.today()
                logger.info(
                    "飞书创建成功: feishu_record_id=%s", record.feishu_record_id
                )
        except Exception:
            logger.exception(
                "Feishu sync FAILED for position transfer: name=%s",
                record.employee_name,
            )

    async def _delete_from_feishu(self, record: PositionTransferRecord) -> None:
        """从飞书多维表格删除。"""
        if not record.feishu_record_id:
            return
        try:
            pair = await self._get_position_bitable()
            if not pair:
                return
            client, table_id = pair
            await client.delete_record(table_id, record.feishu_record_id)
        except Exception:
            logger.warning("Feishu delete failed for position transfer", exc_info=True)

    async def sync_from_feishu(self) -> dict[str, Any]:
        """从飞书多维表格全量同步到本地缓存（飞书为主源）。

        飞书有的记录 → upsert 到本地；
        飞书没有但本地有的记录 → 本地软删除。
        """
        pair = await self._get_position_bitable()
        if not pair:
            raise AppException(status_code=503, message="岗位调动飞书设置未配置")
        client, table_id = pair

        raw_records = await client.search_records(table_id, page_size=500)
        feishu_record_ids = {rec.get("record_id", "") for rec in raw_records}
        stats = {
            "created": 0,
            "updated": 0,
            "deleted": 0,
            "failed": 0,
            "total": len(raw_records),
        }

        # 1. 飞书有 → upsert 本地
        for rec in raw_records:
            try:
                fields = rec.get("fields", {})
                rid = rec.get("record_id", "")
                data = {
                    "feishu_record_id": rid,
                    "employee_name": _extract_text(fields.get("申请人", "")),
                    "department_before": _extract_text(fields.get("原部门", "")),
                    "original_position": _extract_text(fields.get("原职位", "")),
                    "effective_date": _parse_date(fields.get("生效日期")),
                    "apply_department": _extract_text(fields.get("申请部门", "")),
                    "apply_position": _extract_text(fields.get("申请职位", "")),
                    "contact_phone": _extract_text(fields.get("联系电话", "")),
                    "applicant_confirmation_text": _extract_text(
                        fields.get("申请人确认说明", "")
                    ),
                    "applicant_signature": _extract_text(fields.get("申请人签名", "")),
                    "applicant_confirmation_date": _parse_date(
                        fields.get("申请人确认日期")
                    ),
                    "approval_status": "已通过",  # 飞书同步的记录默认为已通过
                    "feishu_synced_at": date.today(),
                }
                # 以飞书为主：不过滤空值，空值也覆盖本地旧值
                existing = await self.repo.get_by_feishu_record_id(rid) if rid else None
                if existing:
                    for k, v in data.items():
                        if k != "id":
                            setattr(existing, k, v)
                    await self.repo.update(existing)
                    stats["updated"] += 1
                else:
                    record = PositionTransferRecord(**data)
                    await self.repo.create(record)
                    stats["created"] += 1
            except Exception:
                logger.exception(
                    "Failed to sync position transfer record %s", rec.get("record_id")
                )
                stats["failed"] += 1

        # 2. 飞书没有但本地有 → 软删除（以飞书为主）
        local_records = await self.repo.list_all_with_feishu_id()
        for local_rec in local_records:
            if local_rec.feishu_record_id not in feishu_record_ids:
                local_rec.is_deleted = True
                stats["deleted"] += 1
                logger.info(
                    (
                        "岗位调动记录已在飞书中删除，本地软删除: feishu_record_id=%s"
                        ", name=%s"
                    ),
                    local_rec.feishu_record_id,
                    local_rec.employee_name,
                )

        return stats

    # ─── Multi-step approval workflow ───
    #
    # 审批流程：
    # 主管以下：直属领导 →
    # 部门经理/总监（经理签了就不需要总监，没有经理的签总监，同一人自动跳过）
    #   主管以上：部门经理 → 部门总监 → 主管领导
    #   原部门和接收部门各走一遍，最后 HR → 常务副总 → 总经理
    #
    APPROVAL_NODES: list[dict[str, Any]] = [
        {"node": "applicant", "label": "申请人确认", "level": "applicant"},
        # 原部门审批
        {
            "node": "origin_direct_leader",
            "label": "原部门直属领导",
            "level": "direct_leader",
            "skip_for_supervisor": True,
        },
        {"node": "origin_manager", "label": "原部门经理", "level": "manager"},
        {
            "node": "origin_director",
            "label": "原部门总监",
            "level": "director",
            "skip_if_manager_present": True,
        },
        {
            "node": "origin_vp",
            "label": "原部门主管领导",
            "level": "vp",
            "skip_for_junior": True,
        },
        # 接收部门审批
        {
            "node": "target_direct_leader",
            "label": "接收部门直属领导",
            "level": "direct_leader",
            "skip_for_supervisor": True,
        },
        {"node": "target_manager", "label": "接收部门经理", "level": "manager"},
        {
            "node": "target_director",
            "label": "接收部门总监",
            "level": "director",
            "skip_if_manager_present": True,
        },
        {
            "node": "target_vp",
            "label": "接收部门主管领导",
            "level": "vp",
            "skip_for_junior": True,
        },
        # 固定角色（仅主管级以上）
        {"node": "hr", "label": "人力资源部", "level": "hr", "skip_for_junior": True},
        {
            "node": "executive_vp",
            "label": "常务副总经理",
            "level": "executive_vp",
            "skip_for_junior": True,
        },
        {
            "node": "general_manager",
            "label": "总经理",
            "level": "general_manager",
            "skip_for_junior": True,
        },
    ]

    async def submit_for_approval(
        self,
        record_id: UUID,
        request: Any,
    ) -> PositionTransferRecord:
        """提交审批：将草稿转为审批流程，初始化审批节点。"""
        # Redis 锁防重复提交
        from app.core.redis import cache_get, cache_set

        lock_key = f"hr:position_transfer:submit_lock:{record_id}"
        if await cache_get(lock_key):
            raise AppException(message="正在提交中，请勿重复操作")
        await cache_set(lock_key, "1", ex=30)  # 30秒锁

        try:
            return await self._do_submit_for_approval(record_id, request)
        finally:
            from app.core.redis import redis_client

            await redis_client.delete(lock_key)

    async def _do_submit_for_approval(
        self,
        record_id: UUID,
        request: Any,
    ) -> PositionTransferRecord:
        """提交审批内部实现。"""
        # 用 FOR UPDATE 锁行，防止并发重复提交
        query_result = await self.session.execute(
            select(PositionTransferRecord)
            .where(
                PositionTransferRecord.id == record_id,
                PositionTransferRecord.is_deleted.is_(False),
            )
            .with_for_update()
        )
        record = query_result.scalar_one_or_none()
        if not record:
            raise NotFoundException("岗位调动记录", str(record_id))
        if record.approval_status != "草稿":
            raise AppException(message="只有草稿状态的记录才能提交审批")

        is_supervisor = request.is_supervisor_level
        custom_approvers = getattr(request, "custom_approvers", None) or {}
        steps: list[dict[str, Any]] = []

        # 先收集所有节点的审批人（用于 skip_if_manager_present 判断）
        resolved: list[tuple[dict[str, Any], str | None, str | None]] = []
        for node_def in self.APPROVAL_NODES:
            # 跳过规则：
            # - 主管以下(skip_for_junior)：跳过主管领导
            # - 主管以上(skip_for_supervisor)：跳过直属领导
            if node_def.get("skip_for_junior") and not is_supervisor:
                continue
            if node_def.get("skip_for_supervisor") and is_supervisor:
                continue

            node_key = node_def["node"]
            # 优先使用手动指定的审批人，否则自动解析
            if node_key in custom_approvers:
                approver_name = custom_approvers[node_key]
                approver_open_id = (
                    await self._get_open_id_by_name(approver_name)
                    if approver_name
                    else None
                )
            else:
                approver_name, approver_open_id = await self._resolve_approver(
                    record, node_key
                )

            resolved.append((node_def, approver_name, approver_open_id))

        # 第二轮：根据 skip_if_manager_present 和同一人去重，生成最终 steps
        # 收集每个部门前缀的 manager 是否有值
        manager_signed: dict[str, bool] = {}  # {"origin": True, "target": False}
        for node_def, approver_name, _ in resolved:
            level = node_def.get("level", "")
            if level == "manager":
                prefix = (
                    "origin" if node_def["node"].startswith("origin_") else "target"
                )
                manager_signed[prefix] = bool(approver_name)

        for node_def, approver_name, approver_open_id in resolved:
            # 经理有值时跳过总监（经理签了就不需要总监）
            if node_def.get("skip_if_manager_present"):
                prefix = (
                    "origin" if node_def["node"].startswith("origin_") else "target"
                )
                if manager_signed.get(prefix):
                    continue

            # 没有审批人的节点直接标记为跳过（不卡流程）
            if node_def["node"] != "applicant" and not approver_name:
                steps.append(
                    {
                        "node": node_def["node"],
                        "label": node_def["label"],
                        "status": "skipped",
                        "signer": "",
                        "signer_open_id": "",
                        "date": date.today().strftime("%Y.%m.%d"),
                        "opinion": "未配置审批人，自动跳过",
                    }
                )
                continue

            steps.append(
                {
                    "node": node_def["node"],
                    "label": node_def["label"],
                    "status": "pending",
                    "signer": approver_name or "",
                    "signer_open_id": approver_open_id or "",
                    "date": None,
                    "opinion": None,
                }
            )

        # 申请人确认节点自动通过，current_step 指向第一个 pending 节点
        if steps and steps[0]["node"] == "applicant":
            steps[0]["status"] = "approved"
            steps[0]["date"] = date.today().strftime("%Y.%m.%d")
            steps[0]["opinion"] = "提交申请"

        # 找到第一个 pending 的节点作为 current_step
        current_step = 0
        for i, s in enumerate(steps):
            if s["status"] == "pending":
                current_step = i
                break
        else:
            # 所有节点都已完成/跳过
            current_step = len(steps)

        record.approval_flow = {
            "current_step": current_step,
            "applicant_name": record.employee_name,
            "applicant_date": date.today().strftime("%Y.%m.%d"),
            "is_supervisor_level": is_supervisor,
            "steps": steps,
        }
        record.approval_status = "待审批"
        result = await self.repo.update(record)
        # 同步审批状态到飞书
        await self._sync_to_feishu(result, is_create=False)

        # 通知第一个审批节点
        await self._notify_next_approver(result)
        return result

    async def approve_current_node(
        self,
        record_id: UUID,
        request: Any,
    ) -> PositionTransferRecord:
        """审批通过当前节点，推进到下一节点。"""
        from app.core.redis import cache_get, cache_set, redis_client

        lock_key = f"hr:position_transfer:approve_lock:{record_id}"
        if await cache_get(lock_key):
            raise AppException(message="正在处理中，请勿重复操作")
        await cache_set(lock_key, "1", ex=30)
        try:
            return await self._do_approve_current_node(record_id, request)
        finally:
            await redis_client.delete(lock_key)

    async def _do_approve_current_node(
        self,
        record_id: UUID,
        request: Any,
    ) -> PositionTransferRecord:
        """审批通过当前节点，推进到下一节点。"""
        query_result = await self.session.execute(
            select(PositionTransferRecord)
            .where(
                PositionTransferRecord.id == record_id,
                PositionTransferRecord.is_deleted.is_(False),
            )
            .with_for_update()
        )
        record = query_result.scalar_one_or_none()
        if not record:
            raise NotFoundException("岗位调动记录", str(record_id))
        if not record.approval_flow:
            raise AppException(message="该记录未提交审批流程")

        # 深拷贝后再修改，避免 in-place 修改导致 SQLAlchemy JSONB 变更检测失败
        flow = copy.deepcopy(record.approval_flow)
        current = flow.get("current_step", 0)
        steps = flow.get("steps", [])

        if current >= len(steps):
            raise AppException(message="所有节点已审批完毕")

        # 标记当前节点为已通过
        step = steps[current]
        step["status"] = "approved"
        step["date"] = date.today().strftime("%Y.%m.%d")
        step["opinion"] = request.opinion or "同意"

        # 推进到下一个 pending 节点（跳过 skipped）
        new_current = current + 1
        while new_current < len(steps) and steps[new_current]["status"] == "skipped":
            new_current += 1
        flow["current_step"] = new_current

        if new_current >= len(steps):
            # 所有节点完成
            record.approval_status = "已通过"
            record.approval_date = date.today()
        else:
            record.approval_status = "待审批"

        record.approval_flow = flow
        result = await self.repo.update(record)
        # 同步审批状态到飞书
        await self._sync_to_feishu(result, is_create=False)

        # 通知下一节点审批人
        if record.approval_status == "待审批":
            await self._notify_next_approver(result)

        return result

    async def reject_current_node(
        self,
        record_id: UUID,
        request: Any,
    ) -> PositionTransferRecord:
        """审批拒绝当前节点。"""
        query_result = await self.session.execute(
            select(PositionTransferRecord)
            .where(
                PositionTransferRecord.id == record_id,
                PositionTransferRecord.is_deleted.is_(False),
            )
            .with_for_update()
        )
        record = query_result.scalar_one_or_none()
        if not record:
            raise NotFoundException("岗位调动记录", str(record_id))
        if not record.approval_flow:
            raise AppException(message="该记录未提交审批流程")

        # 深拷贝后再修改
        flow = copy.deepcopy(record.approval_flow)
        current = flow.get("current_step", 0)
        steps = flow.get("steps", [])

        if current >= len(steps):
            raise AppException(message="所有节点已审批完毕")

        step = steps[current]
        step["status"] = "rejected"
        step["date"] = date.today().strftime("%Y.%m.%d")
        step["opinion"] = request.opinion or "不同意"

        record.approval_status = "已拒绝"
        record.approval_flow = flow
        result = await self.repo.update(record)
        # 同步审批状态到飞书
        await self._sync_to_feishu(result, is_create=False)

        # 通知申请人审批被拒绝
        await self._notify_applicant_rejected(result)
        return result

    async def _resolve_approver(
        self, record: PositionTransferRecord, node: str
    ) -> tuple[str | None, str | None]:
        """根据节点类型从部门级审批人配置表解析审批人 (姓名, open_id)。"""
        if node == "applicant":
            return record.employee_name, None

        # 确定是原部门还是接收部门
        dept_name = None
        if node.startswith("origin_"):
            dept_name = record.department_before
        elif node.startswith("target_"):
            dept_name = record.apply_department

        # 从部门级审批人配置表查找
        if dept_name:
            config = await self._get_dept_approval_config(dept_name)
            if config:
                level = node.split("_", 1)[1] if "_" in node else node
                # origin_direct_leader → direct_leader, origin_manager → manager 等
                if level == "direct_leader":
                    return config.direct_leader_name, config.direct_leader_open_id
                elif level == "manager":
                    return config.manager_name, config.manager_open_id
                elif level == "director":
                    return config.director_name, config.director_open_id
                elif level == "vp":
                    return config.vp_name, config.vp_open_id

        # 固定角色：按对应部门负责人解析（人力资源部 / 常务副总经理 / 总经理）
        # 不再硬编码测试人员，保证审批人按真实组织架构指派人。
        role_dept = {
            "hr": "人力资源部",
            "executive_vp": "常务副总经理",
            "general_manager": "总经理",
        }
        if node in role_dept:
            dept = await self._get_department_by_name(role_dept[node])
            if dept and dept.leader_name:
                return dept.leader_name, await self._get_open_id_by_name(
                    dept.leader_name
                )
            logger.warning(
                "岗位调动审批节点 %s 未找到部门(%s)负责人，中断审批",
                node,
                role_dept[node],
            )
            return None, None

        return None, None

    async def _get_dept_approval_config(self, dept_name: str) -> Any:
        """根据部门名称查询部门级审批人配置。"""
        from app.modules.hr.models import HrDeptApprovalConfig

        result = await self.session.execute(
            select(HrDeptApprovalConfig)
            .where(
                HrDeptApprovalConfig.department_name == dept_name,
                HrDeptApprovalConfig.is_deleted.is_(False),
            )
            .limit(1)
        )
        return result.scalar_one_or_none()

    async def _get_open_id_by_name(self, name: str) -> str | None:
        """根据姓名从飞书成员缓存表查找 open_id。"""
        if not name:
            return None
        result = await self.session.execute(
            select(HrFeishuMember.open_id)
            .where(
                HrFeishuMember.name == name,
                HrFeishuMember.is_deleted.is_(False),
                HrFeishuMember.status == "1",  # 在职
            )
            .limit(1)
        )
        return result.scalar_one_or_none()

    async def _get_member_by_job_title(
        self, job_title: str
    ) -> tuple[str | None, str | None]:
        """根据职位从飞书成员缓存表查找 (open_id, name)。"""
        result = await self.session.execute(
            select(HrFeishuMember.open_id, HrFeishuMember.name)
            .where(
                HrFeishuMember.job_title == job_title,
                HrFeishuMember.is_deleted.is_(False),
                HrFeishuMember.status == "1",  # 在职
            )
            .limit(1)
        )
        row = result.first()
        if row:
            return row[0], row[1]
        return None, None

    async def _get_department_by_name(self, name: str) -> Any:
        """根据部门名称查询部门记录（优先顶级部门）。"""
        result = await self.session.execute(
            select(HrDepartment)
            .where(
                HrDepartment.name == name,
                HrDepartment.is_deleted.is_(False),
            )
            .order_by(HrDepartment.parent_id.asc().nullsfirst())
            .limit(1)
        )
        return result.scalar_one_or_none()

    async def _notify_next_approver(self, record: PositionTransferRecord) -> None:
        """通知下一节点审批人 - 直接使用步骤中解析的 open_id。"""
        if not record.approval_flow:
            return
        flow = record.approval_flow
        current = flow.get("current_step", 0)
        steps = flow.get("steps", [])
        if current >= len(steps):
            logger.info("[通知] 所有节点已完成, record=%s", record.id)
            return

        next_step = steps[current]
        open_id = next_step.get("signer_open_id", "")
        approver_name = next_step.get("signer", "")
        node = next_step.get("node", "")

        logger.info(
            (
                "[通知] 发送审批卡片: record=%s, step=%d/%d, node"
                "=%s, signer=%s, open_id=%s"
            ),
            record.id,
            current,
            len(steps),
            node,
            approver_name,
            open_id[:20] if open_id else "NONE",
        )

        if not open_id:
            logger.warning(
                "岗位调动审批通知跳过: node=%s, signer=%s, 未找到open_id",
                node,
                approver_name,
            )
            return

        # Redis 去重：同一 record+step 只发一次卡片
        from app.core.redis import cache_get, cache_set

        notify_key = f"hr:position_transfer:notify:{record.id}:{current}"
        if await cache_get(notify_key):
            logger.warning(
                "[通知] 重复通知已跳过: record=%s, step=%d", record.id, current
            )
            return
        await cache_set(notify_key, "1", ex=3600)  # 1小时过期

        try:
            from app.platform.integrations.feishu.notification import (
                send_user_card_with_message_id,
            )

            card_elements: list[dict[str, Any]] = [
                {
                    "tag": "div",
                    "text": {
                        "tag": "plain_text",
                        "content": f"申请人：{record.employee_name}",
                    },
                },
                {
                    "tag": "div",
                    "text": {
                        "tag": "plain_text",
                        "content": (
                            f"原部门：{record.department_before} -> "
                            f"申请部门：{record.apply_department}"
                        ),
                    },
                },
                {
                    "tag": "div",
                    "text": {
                        "tag": "plain_text",
                        "content": (
                            f"原职位：{record.original_position} -> "
                            f"申请职位：{record.apply_position}"
                        ),
                    },
                },
                {
                    "tag": "div",
                    "text": {
                        "tag": "plain_text",
                        "content": f"当前节点：{next_step.get('label')}",
                    },
                },
                {"tag": "hr"},
                {
                    "tag": "action",
                    "actions": [
                        {
                            "tag": "button",
                            "text": {"tag": "plain_text", "content": "✅ 通过"},
                            "type": "primary",
                            "value": {
                                "module": "position_transfer_approval",
                                "action": "approve",
                                "record_id": str(record.id),
                                "node": node,
                                "signer": approver_name,
                            },
                        },
                        {
                            "tag": "button",
                            "text": {"tag": "plain_text", "content": "❌ 拒绝"},
                            "type": "danger",
                            "value": {
                                "module": "position_transfer_approval",
                                "action": "reject",
                                "record_id": str(record.id),
                                "node": node,
                                "signer": approver_name,
                            },
                        },
                    ],
                },
            ]

            message_id = await send_user_card_with_message_id(
                open_id=open_id,
                title="岗位调动审批 - 需您审批",
                content="",
                elements=card_elements,
            )
            if message_id:
                record.feishu_approval_message_id = message_id
            logger.info(
                "岗位调动审批通知已发送: record=%s, node=%s, signer=%s, open_id=%s",
                record.id,
                node,
                approver_name,
                (open_id[:20] + "..." if open_id else "NONE"),
            )
        except Exception:
            logger.warning(
                "飞书通知失败 - 审批人: %s, 节点: %s",
                approver_name,
                next_step.get("label"),
            )

    async def _notify_applicant_rejected(self, record: PositionTransferRecord) -> None:
        """通知申请人审批被拒绝。"""
        try:
            # 从飞书成员缓存表查找申请人 open_id
            applicant_open_id = await self._get_open_id_by_name(record.employee_name)
            if not applicant_open_id:
                logger.warning(
                    "飞书拒绝通知跳过: 未找到申请人 %s 的 open_id",
                    record.employee_name,
                )
                return

            from app.platform.integrations.feishu.notification import (
                send_user_card_with_message_id,
            )

            await send_user_card_with_message_id(
                open_id=applicant_open_id,
                title="岗位调动审批 - 已拒绝",
                content="",
                elements=[
                    {
                        "tag": "div",
                        "text": {
                            "tag": "plain_text",
                            "content": "您的岗位调动申请已被拒绝",
                        },
                    },
                    {
                        "tag": "div",
                        "text": {
                            "tag": "plain_text",
                            "content": (
                                f"原部门：{record.department_before} → "
                                f"申请部门：{record.apply_department}"
                            ),
                        },
                    },
                    {
                        "tag": "div",
                        "text": {
                            "tag": "plain_text",
                            "content": (
                                f"原职位：{record.original_position} → "
                                f"申请职位：{record.apply_position}"
                            ),
                        },
                    },
                ],
            )
        except Exception:
            logger.warning("飞书拒绝通知失败")

    async def list_approvals(
        self,
        *,
        current_user: Any = None,
        tab: str = "my_applications",
        page: int = 1,
        page_size: int = 20,
    ) -> tuple[list[PositionTransferRecord], int]:
        """获取审批列表（按 tab + 当前用户筛选）。

        - my_applications：我发起的申请（employee_name == 当前用户姓名）
        - pending_approval：待我审批（当前步骤 signer 为当前用户）
        - approved：我审批过的（任一步骤 signer 为当前用户）
        """
        user_name = (getattr(current_user, "name", None) or "").strip()
        user_open_id = (getattr(current_user, "feishu_open_id", None) or "").strip()

        stmt = select(PositionTransferRecord).where(
            PositionTransferRecord.is_deleted.is_(False)
        )

        if tab == "my_applications":
            stmt = stmt.where(PositionTransferRecord.employee_name == user_name)
        elif tab == "pending_approval":
            stmt = stmt.where(PositionTransferRecord.approval_status == "待审批")
        elif tab == "approved":
            stmt = stmt.where(
                PositionTransferRecord.approval_status.in_(["已通过", "已拒绝"])
            )

        # 按创建时间倒序
        stmt = stmt.order_by(desc(PositionTransferRecord.created_at))

        # 我的申请：SQL 层过滤 + 分页
        if tab == "my_applications":
            count_stmt = select(func.count()).select_from(stmt.subquery())
            total_result = await self.session.execute(count_stmt)
            total = total_result.scalar() or 0
            stmt = stmt.offset((page - 1) * page_size).limit(page_size)
            result = await self.session.execute(stmt)
            return list(result.scalars().all()), total

        # 待我审批 / 我审批过的：signer / signer_open_id 存于 approval_flow JSONB，
        # 岗位调动审批记录量小，拉取后按当前用户匹配，避免复杂 JSONB 下标 SQL。
        def _step_matches(step: dict[str, Any]) -> bool:
            if user_open_id and step.get("signer_open_id") == user_open_id:
                return True
            if user_name and step.get("signer") == user_name:
                return True
            return False

        result = await self.session.execute(stmt)
        records = list(result.scalars().all())

        if tab == "pending_approval":
            filtered = []
            for r in records:
                flow = r.approval_flow or {}
                steps = flow.get("steps", []) or []
                current = flow.get("current_step", 0)
                if 0 <= current < len(steps) and _step_matches(steps[current]):
                    filtered.append(r)
            records = filtered
        else:  # approved
            filtered = []
            for r in records:
                flow = r.approval_flow or {}
                steps = flow.get("steps", []) or []
                if any(_step_matches(s) for s in steps):
                    filtered.append(r)
            records = filtered

        total = len(records)
        start = (page - 1) * page_size
        return records[start : start + page_size], total

    async def export_approval_pdf(self, record_id: UUID) -> tuple[bytes, str]:
        """导出内调申请表，填充审批数据到 Word 模板。"""
        record = await self.get_record(record_id)
        docx_bytes = await self._fill_word_template(record)
        filename = (
            f"内调申请表_{record.employee_name}_{date.today().strftime('%Y%m%d')}.docx"
        )
        return docx_bytes, filename

    async def _fill_word_template(self, record: PositionTransferRecord) -> bytes:
        """从数据库读取模板，填充审批数据到 Word 模板，不修改任何格式/字体/大小。"""
        import io as _io

        from docx import Document

        # 从数据库读取模板（部署时预置）
        template_code = "position_transfer_application"
        result = await self.session.execute(
            select(HrDocumentTemplate)
            .where(
                HrDocumentTemplate.template_code == template_code,
                HrDocumentTemplate.is_deleted.is_(False),
            )
            .limit(1)
        )
        template = result.scalar_one_or_none()
        if not template:
            raise AppException(message=f"模板未配置: {template_code}，请联系管理员上传")

        flow: dict[str, Any] = record.approval_flow or {}
        steps: list[dict[str, Any]] = cast(
            list[dict[str, Any]], flow.get("steps") or []
        )

        def _get_step(node: str) -> dict[str, Any]:
            for s in steps:
                if s.get("node") == node:
                    return s
            return {}

        def _fmt(v: Any) -> Any:
            return str(v) if v else ""

        doc = Document(_io.BytesIO(template.template_data))
        table = doc.tables[0]

        # ── 全部单元格垂直居中 ──
        for row in table.rows:
            for cell in row.cells:
                _set_cell_v_center(cell)

        # ── Row 0: 申请人 / 原部门 / 原职位 / 生效日期 ──
        # 模板要求：宋体 10号 居中
        _set_cell_text_v(table.rows[0].cells[1], _fmt(record.employee_name))
        _set_cell_text_v(table.rows[0].cells[3], _fmt(record.department_before))
        _set_cell_text_v(table.rows[0].cells[5], _fmt(record.original_position))
        _set_cell_text_v(
            table.rows[0].cells[7],
            _fmt(str(record.effective_date) if record.effective_date else ""),
        )

        # ── Row 1: 申请部门 / 申请职位 / 申请日期 / 联系电话 ──
        _set_cell_text_v(table.rows[1].cells[1], _fmt(record.apply_department))
        _set_cell_text_v(table.rows[1].cells[3], _fmt(record.apply_position))
        _set_cell_text_v(table.rows[1].cells[5], _fmt(flow.get("applicant_date", "")))
        _set_cell_text_v(table.rows[1].cells[7], _fmt(record.contact_phone))

        # ── Row 2: 申请人确认 ──
        # 模板格式: P0 "申请人确认："  P1 "此申请经审批..."  P2 "签名：    日期："
        _replace_signer_date_cell(
            table.rows[2].cells[0],
            _fmt(record.applicant_signature),
            _fmt(
                str(record.applicant_confirmation_date)
                if record.applicant_confirmation_date
                else ""
            ),
        )

        # ── Row 3-11: 审批节点 ──
        _fill_approval_node(table.rows[3].cells[0], _get_step("origin_direct_leader"))
        _fill_approval_node_mgr_dir(
            table.rows[4].cells[0],
            _get_step("origin_manager"),
            _get_step("origin_director"),
        )
        _fill_approval_node(table.rows[5].cells[0], _get_step("origin_vp"))
        _fill_approval_node(table.rows[6].cells[0], _get_step("target_direct_leader"))
        _fill_mgr_dir_split(
            table.rows[7].cells[0],
            table.rows[7].cells[3],
            _get_step("target_manager"),
            _get_step("target_director"),
        )
        _fill_approval_node(table.rows[8].cells[0], _get_step("target_vp"))
        _fill_hr_row(table.rows[9].cells[0], table.rows[9].cells[2], _get_step("hr"))
        _fill_approval_node(table.rows[10].cells[0], _get_step("executive_vp"))
        _fill_approval_node(table.rows[11].cells[0], _get_step("general_manager"))

        buf = _io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()


# ── Word 模板填充辅助函数（不修改任何格式/字体/大小） ──


def _set_cell_text(cell: Any, text: str) -> Any:
    """设置单元格文字，保留原有格式。"""
    paragraphs = cell.paragraphs
    if paragraphs:
        p = paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.add_run(text)
    else:
        p = cell.add_paragraph()
        p.add_run(text)


def _set_cell_text_v(cell: Any, text: str) -> Any:
    """设置单元格文字：中文宋体、英文数字 Times New Roman、10号、水平+垂直全居中。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    _set_cell_v_center(cell)
    paragraphs = cell.paragraphs
    if paragraphs:
        p = paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].text = text
            for run in p.runs[1:]:
                run.text = ""
            run = p.runs[0]
        else:
            run = p.add_run(text)
    else:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)

    # 中文宋体 + 英文数字 Times New Roman + 10号
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    east_asia = r_pr.find(qn("w:rFonts"))
    if east_asia is None:
        from lxml import etree  # type: ignore[import-untyped]

        east_asia = etree.SubElement(r_pr, qn("w:rFonts"))
    east_asia.set(qn("w:eastAsia"), "宋体")


def _replace_signer_date_cell(cell: Any, signer: str, signdate: str) -> Any:
    """替换 '签名：XXX  日期：XXX' 中的占位空格，保留格式。"""
    # 设置垂直居中
    _set_cell_v_center(cell)
    paragraphs = cell.paragraphs
    for p in paragraphs:
        full = p.text
        if "签名：" in full and "日期：" in full:
            runs = p.runs
            if len(runs) >= 3:
                # Run0=签名：xxx, Run1=空格, Run2=日期：xxx
                runs[0].text = f"签名：{signer}" + " " * max(
                    0, len(runs[0].text) - len(f"签名：{signer}")
                )
                runs[-1].text = f"日期：{signdate}"
                for r in runs[1:-1]:
                    r.text = " " * len(r.text)
            elif len(runs) >= 2:
                # Run0=签名：xxx, Run1=日期：xxx
                runs[0].text = f"签名：{signer}"
                runs[1].text = f"日期：{signdate}"
            break


def _set_cell_v_center(cell: Any) -> Any:
    """设置单元格垂直居中。"""
    from docx.oxml.ns import qn

    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        tc_pr = cell._tc.get_or_add_tcPr()
    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is None:
        from lxml import etree

        v_align = etree.SubElement(tc_pr, qn("w:vAlign"))
    v_align.set(qn("w:val"), "center")


def _fill_approval_node(cell: Any, step: dict[str, Any]) -> Any:
    """填充审批节点：签名和日期替换到模板原有段落中，不修改格式。"""
    signer = str(step.get("signer", "")) if step.get("signer") else ""
    signdate = str(step.get("date", "")) if step.get("date") else ""
    _replace_signer_date_cell(cell, signer, signdate)


def _fill_approval_node_mgr_dir(
    cell: Any, mgr_step: dict[str, Any], dir_step: dict[str, Any]
) -> Any:
    """填充经理/总监合并节点（签名日期同一行）。"""
    signer = str(mgr_step.get("signer") or dir_step.get("signer") or "")
    signdate = str(mgr_step.get("date") or dir_step.get("date") or "")
    _replace_signer_date_cell(cell, signer, signdate)


def _fill_mgr_dir_split(
    left_cell: Any, right_cell: Any, mgr_step: dict[str, Any], dir_step: dict[str, Any]
) -> Any:
    """填充接收部门经理/总监（签名在左，日期在右）。"""
    signer = str(mgr_step.get("signer") or dir_step.get("signer") or "")
    signdate = str(mgr_step.get("date") or dir_step.get("date") or "")
    for p in left_cell.paragraphs:
        for run in p.runs:
            if "签名：" in run.text:
                run.text = f"签名：{signer}"
                break
    for p in right_cell.paragraphs:
        for run in p.runs:
            if "日期：" in run.text:
                run.text = f"日期：{signdate}"
                break


def _fill_hr_row(left_cell: Any, right_cell: Any, hr_step: dict[str, Any]) -> Any:
    """填充人力资源部行（签名在左单元格，日期在右单元格）。"""
    signer = str(hr_step.get("signer", "")) if hr_step.get("signer") else ""
    signdate = str(hr_step.get("date", "")) if hr_step.get("date") else ""
    # 左单元格：替换 "签名："
    for p in left_cell.paragraphs:
        for run in p.runs:
            if "签名：" in run.text:
                run.text = f"签名：{signer}"
                break
    # 右单元格：替换 "日期："
    for p in right_cell.paragraphs:
        for run in p.runs:
            if "日期：" in run.text:
                run.text = f"日期：{signdate}"
                break


class _LegacyFeishuRecordService:
    repo: Any
    bitable: Any
    record_label = "记录"

    async def get_record(self, record_id: UUID) -> Any:
        record = await self.repo.get_by_id(record_id)
        if not record:
            raise NotFoundException(self.record_label, str(record_id))
        return record

    async def _parse_feishu_record(self, record: dict[str, Any]) -> dict[str, Any]:
        raise NotImplementedError

    async def sync_from_feishu(self) -> dict[str, int]:
        raw_records = await self.bitable.client.search_records(
            self.bitable.table_id,
            page_size=500,
        )
        stats = {"created": 0, "updated": 0, "failed": 0, "total": len(raw_records)}
        for record in raw_records:
            try:
                data = await self._parse_feishu_record(record)
                data["feishu_synced_at"] = date.today()
                record_id = data.get("feishu_record_id")
                if not record_id:
                    stats["failed"] += 1
                    continue
                await self.repo.upsert_by_feishu_record_id(data)
                existing = await self.repo.get_by_feishu_record_id(record_id)
                recently_created = (
                    existing
                    and existing.created_at
                    and (
                        datetime.utcnow() - existing.created_at.replace(tzinfo=None)
                    ).total_seconds()
                    < 60
                )
                stats["created" if recently_created else "updated"] += 1
            except Exception:
                logger.exception(
                    "Failed to sync legacy HR record %s", record.get("record_id")
                )
                stats["failed"] += 1
        return stats

    async def get_sync_status(self) -> SyncStatusResponse:
        local_total = await self.repo.count_total()
        synced_count = await self.repo.count_synced()
        return SyncStatusResponse(
            local_total=local_total,
            feishu_total=synced_count,
            synced_count=synced_count,
            unsynced_count=local_total - synced_count,
            conflict_count=0,
            last_sync_at=None,
        )


class OnboardingRecordService(_LegacyFeishuRecordService):
    record_label = "入职记录"

    def __init__(self, session: AsyncSession) -> None:
        self.repo = OnboardingRecordRepository(session)
        self.bitable = OnboardingBitableDataSource()

    async def _parse_feishu_record(self, record: dict[str, Any]) -> dict[str, Any]:
        from app.modules.hr.feishu.onboarding_datasource import OnboardingRecord

        return OnboardingRecord.from_api(record).to_dict()

    async def list_records(
        self,
        *,
        department: str | None = None,
        position: str | None = None,
        is_employed: str | None = None,
        keyword: str | None = None,
        page: int = 1,
        page_size: int = 20,
        sort_by: str = "hire_date",
        sort_order: str = "desc",
        dept_alias_set: set[str] | None = None,
    ) -> tuple[list[LegacyOnboardingRecord], int]:
        return cast(
            tuple[list[LegacyOnboardingRecord], int],
            await self.repo.list_records(
                department=department,
                position=position,
                is_employed=is_employed,
                keyword=keyword,
                page=page,
                page_size=page_size,
                sort_by=sort_by,
                sort_order=sort_order,
                dept_alias_set=dept_alias_set,
            ),
        )


class DepartureRecordService(_LegacyFeishuRecordService):
    record_label = "离职台账记录"

    def __init__(self, session: AsyncSession) -> None:
        self.repo = DepartureRecordRepository(session)
        self.bitable = DepartureBitableDataSource()

    async def _parse_feishu_record(self, record: dict[str, Any]) -> dict[str, Any]:
        from app.modules.hr.feishu.departure_datasource import DepartureRecord

        return DepartureRecord.from_api(record).to_dict()

    async def list_records(
        self,
        *,
        department: str | None = None,
        offboarding_type: str | None = None,
        keyword: str | None = None,
        page: int = 1,
        page_size: int = 20,
        sort_by: str = "offboarding_date",
        sort_order: str = "desc",
        dept_alias_set: set[str] | None = None,
    ) -> tuple[list[LegacyDepartureRecord], int]:
        return cast(
            tuple[list[LegacyDepartureRecord], int],
            await self.repo.list_records(
                department=department,
                offboarding_type=offboarding_type,
                keyword=keyword,
                page=page,
                page_size=page_size,
                sort_by=sort_by,
                sort_order=sort_order,
                dept_alias_set=dept_alias_set,
            ),
        )

    async def create_record(self, data: Any) -> LegacyDepartureRecord:
        values = data.model_dump()
        allowed = {column.name for column in LegacyDepartureRecord.__table__.columns}
        record = LegacyDepartureRecord(
            **{key: value for key, value in values.items() if key in allowed}
        )
        return cast(LegacyDepartureRecord, await self.repo.create(record))

    async def update_record(self, record_id: UUID, data: Any) -> LegacyDepartureRecord:
        record = await self.get_record(record_id)
        allowed = {column.name for column in LegacyDepartureRecord.__table__.columns}
        for field, value in data.model_dump(exclude_unset=True).items():
            if field in allowed and field not in {"id", "created_at", "updated_at"}:
                setattr(record, field, value)
        return cast(LegacyDepartureRecord, await self.repo.update(record))

    async def delete_record(self, record_id: UUID) -> None:
        record = await self.get_record(record_id)
        await self.repo.soft_delete(record)


class TrainingLedgerService:
    def __init__(self, session: AsyncSession) -> None:
        self.session = session
        self.repo = TrainingLedgerRepository(session)

    async def get_record(self, record_id: UUID) -> TrainingLedger:
        record = await self.repo.get_by_id(record_id)
        if not record:
            raise NotFoundException("培训台账记录", str(record_id))
        return record

    async def _resolve_teaching_dept(
        self, instructor: str | None, fallback_dept: str | None
    ) -> str | None:
        """按培训师姓名查 Trainer 表确定授课部门，找不到则 fallback 到主办部门。"""
        if not instructor:
            return fallback_dept
        result = await self.session.execute(
            select(Trainer.department)
            .where(
                Trainer.name == instructor,
                Trainer.is_deleted.is_(False),
                Trainer.department.is_not(None),
                Trainer.department != "",
            )
            .limit(1)
        )
        dept = result.scalar_one_or_none()
        return dept or fallback_dept

    async def create_record(self, data: TrainingLedgerCreate) -> TrainingLedger:
        if not hasattr(self, "session"):
            record = TrainingLedger(**data.model_dump())
            return await self.repo.create(record)

        from app.modules.hr.training_dept_resolver import split_ledger_departments

        # ① 按培训师姓名查 Trainer 表确定 teaching_dept
        # （真实授课部门，所有副本一致，不再被篡改）
        teaching_dept = await self._resolve_teaching_dept(
            data.instructor, data.teaching_dept
        )
        data.teaching_dept = teaching_dept

        # ①.5 归属部门写端归一：裸名/历史变体存规范名
        # （裸名主记录归 MC，DR 副本由下面拆分补齐）
        if data.ledger_department:
            data.ledger_department = (
                await split_ledger_departments(self.session, data.ledger_department)
            )[0]

        # ② 多部门培训二级确认兜底：前端未传 pending 时按本记录涉及部门数判断
        if data.second_level_status != "pending":
            depts_all = [
                d.strip() for d in (data.involved_depts or "").split("、") if d.strip()
            ]
            if len(depts_all) >= 2:
                data.second_level_status = "pending"

        record = TrainingLedger(**data.model_dump())
        await self.repo.create(record)

        # ③ 按涉及部门拆分：为每个涉及部门创建一条完整副本
        # 核心规则：每条记录内容完全一致（授课部门/涉及部门/培训对象等
        # 都不变），仅 ledger_department 不同；涉及部门名先按
        # split_ledger_departments 归一（裸 201二车间 拆 MC+DR 两副本）
        if data.involved_depts and data.session_id:
            depts = [d.strip() for d in data.involved_depts.split("、") if d.strip()]
            seen = {data.ledger_department}
            for dept in depts:
                for canonical in await split_ledger_departments(self.session, dept):
                    if canonical in seen:
                        continue
                    seen.add(canonical)
                    copy_data = data.model_dump()
                    copy_data["ledger_department"] = canonical
                    copy_record = TrainingLedger(**copy_data)
                    await self.repo.create(copy_record)

        return record

    # 需要同步的核心字段（teaching_dept 除外，每条记录保持自己的部门）
    _SYNC_FIELDS = {
        "training_datetime",
        "training_date",
        "duration_hours",
        "training_content",
        "instructor",
        "trainer",
        "level_category",
        "involved_depts",
        "trainees",
        "training_type",
        "ledger_assessment_method",
        "plan_source",
        "drug_category",
        "score_summary",
        "assessment_result",
    }

    async def update_record(
        self, record_id: UUID, data: TrainingLedgerUpdate
    ) -> TrainingLedger:
        record = await self.get_record(record_id)
        update_data = data.model_dump(exclude_unset=True)
        for field, value in update_data.items():
            setattr(record, field, value)
        await self.repo.update(record)

        if not hasattr(self, "session"):
            return record

        # 通过 session_id 同步到其他部门的相同培训记录
        if record.session_id:
            await self.repo.sync_by_session_id(
                session_id=record.session_id,
                exclude_id=record_id,
                update_data={
                    k: v for k, v in update_data.items() if k in self._SYNC_FIELDS
                },
            )
        return record

    async def delete_record(self, record_id: UUID) -> None:
        record = await self.get_record(record_id)
        if not hasattr(self, "session"):
            await self.repo.soft_delete(record)
            return
        # 主办方（落款部门）删除 → 同培训其他部门副本标记
        # owner_deleted（变红提示）
        if record.session_id:
            session = await self.session.get(TrainingSession, record.session_id)
            if (
                session
                and session.department
                and record.ledger_department == session.department
            ):
                await self.repo.mark_owner_deleted(
                    session_id=record.session_id, exclude_id=record_id
                )
        await self.repo.soft_delete(record)

    async def delete_by_department(self, department: str) -> int:
        """清空部门全部台账（软删除），返回删除条数."""
        return await self.repo.delete_all_by_department(department)

    # ── 培训时间冲突检测 ──

    async def check_conflict(
        self,
        training_date: date,
        time_start: str,
        time_end: str,
        instructor: str | None,
        trainees: list[str],
        exclude_session_id: str | None = None,
    ) -> dict[str, Any]:
        """检测指定日期/时间段内授课人和参训人员是否与已有台账及培训会话冲突."""
        ledgers = await self.repo.list_by_date(training_date)

        # 同时查询 TrainingSession 表中同日期的会话（未归档的培训也参与冲突检测）
        session_result = await self.session.execute(
            select(TrainingSession).where(
                TrainingSession.training_date == training_date,
            )
        )
        sessions = list(session_result.scalars().all())
        # 排除当前正在编辑的会话自身
        if exclude_session_id:
            from uuid import UUID as _UUID

            exclude_uuid = (
                _UUID(exclude_session_id)
                if isinstance(exclude_session_id, str)
                else exclude_session_id
            )
            sessions = [s for s in sessions if s.id != exclude_uuid]

        new_start = datetime.strptime(time_start, "%H:%M").time()
        new_end = datetime.strptime(time_end, "%H:%M").time()

        # ── 收集原始冲突 ──
        raw_instructor: list[dict[str, Any]] = []
        raw_trainee: list[dict[str, Any]] = []
        occupied_slots: list[tuple[Any, ...]] = []

        # ── 1. 检查台账记录 ──
        for ledger in ledgers:
            existing_start, existing_end = self._parse_datetime_range(
                ledger.training_datetime
            )
            if not existing_start or not existing_end:
                continue

            # 判断时间是否重叠
            if new_start < existing_end and new_end > existing_start:
                # 检查授课人冲突
                if instructor and ledger.instructor and instructor == ledger.instructor:
                    raw_instructor.append(
                        {
                            "training_name": ledger.training_subject or "",
                            "time_range": (
                                f"{existing_start.strftime('%H:%M')}~"
                                f"{existing_end.strftime('%H:%M')}"
                            ),
                            "dept": ledger.teaching_dept or "",
                        }
                    )

                # 检查参训人员冲突（授课人也视为占用）
                if ledger.trainees:
                    existing_trainees = [
                        n.strip() for n in ledger.trainees.split("、") if n.strip()
                    ]
                    all_occupied = set(existing_trainees)
                    if ledger.instructor:
                        all_occupied.add(ledger.instructor)

                    for trainee in trainees:
                        if trainee in all_occupied:
                            raw_trainee.append(
                                {
                                    "name": trainee,
                                    "training_name": ledger.training_subject or "",
                                    "time_range": (
                                        f"{existing_start.strftime('%H:%M')}~"
                                        f"{existing_end.strftime('%H:%M')}"
                                    ),
                                    "dept": ledger.teaching_dept or "",
                                }
                            )

                occupied_slots.append((existing_start, existing_end))

        # ── 2. 检查 TrainingSession 记录（未归档的培训会话） ──
        for sess in sessions:
            if not sess.time_start or not sess.time_end:
                continue
            try:
                sess_start = datetime.strptime(sess.time_start, "%H:%M").time()
                sess_end = datetime.strptime(sess.time_end, "%H:%M").time()
            except (ValueError, TypeError):
                continue

            if new_start < sess_end and new_end > sess_start:
                sess_topic = sess.topic or ""
                sess_dept = sess.department or ""

                # 检查授课人冲突
                if instructor and sess.instructor and instructor == sess.instructor:
                    raw_instructor.append(
                        {
                            "training_name": sess_topic,
                            "time_range": (
                                f"{sess_start.strftime('%H:%M')}~"
                                f"{sess_end.strftime('%H:%M')}"
                            ),
                            "dept": sess_dept,
                        }
                    )

                # 检查参训人员冲突
                existing_names: list[str] = []
                if sess.employee_names and isinstance(sess.employee_names, list):
                    existing_names = [
                        str(n).strip() for n in sess.employee_names if str(n).strip()
                    ]
                all_occupied_sess = set(existing_names)
                if sess.instructor:
                    all_occupied_sess.add(sess.instructor)

                for trainee in trainees:
                    if trainee in all_occupied_sess:
                        raw_trainee.append(
                            {
                                "name": trainee,
                                "training_name": sess_topic,
                                "time_range": (
                                    f"{sess_start.strftime('%H:%M')}~"
                                    f"{sess_end.strftime('%H:%M')}"
                                ),
                                "dept": sess_dept,
                            }
                        )

                occupied_slots.append((sess_start, sess_end))

        # ── 3. 去重汇总 ──
        # 授课人冲突：按 (training_name, time_range) 分组，收集所有涉及部门
        _inst_map: dict[tuple[str, str], set[str]] = {}
        for item in raw_instructor:
            key = (item["training_name"], item["time_range"])
            dept_set = _inst_map.setdefault(key, set())
            if item["dept"]:
                dept_set.add(item["dept"])
        instructor_conflicts = [
            {
                "training_name": k[0],
                "time_range": k[1],
                "conflict_depts": sorted(v),
                "conflict_count": len(v),
            }
            for k, v in _inst_map.items()
        ]

        # 参训人员冲突：按 (training_name, time_range) 分组，收集所有冲突人员
        _trainee_map: dict[tuple[str, str], set[str]] = {}
        for item in raw_trainee:
            key = (item["training_name"], item["time_range"])
            _trainee_map.setdefault(key, set()).add(item["name"])
        trainee_conflicts = [
            {
                "training_name": k[0],
                "time_range": k[1],
                "names": sorted(v),
                "conflict_count": len(v),
            }
            for k, v in _trainee_map.items()
        ]

        # 推荐不冲突的时间段（occupied_slots 先去重）
        occupied_unique = list(dict.fromkeys(occupied_slots))
        suggested_times = self._suggest_free_slots(new_start, new_end, occupied_unique)

        return {
            "has_conflict": bool(instructor_conflicts or trainee_conflicts),
            "instructor_conflicts": instructor_conflicts,
            "trainee_conflicts": trainee_conflicts,
            "suggested_times": suggested_times,
        }

    @staticmethod
    def _parse_datetime_range(dt_str: str | None) -> tuple[Any, ...]:
        """解析 '2026.08.09 19:00~22:00' 格式，返回 (start_time, end_time)."""
        if not dt_str:
            return None, None
        match = re.search(r"(\d{2}:\d{2})\s*[~\-]\s*(\d{2}:\d{2})", dt_str)
        if not match:
            return None, None
        start = datetime.strptime(match.group(1), "%H:%M").time()
        end = datetime.strptime(match.group(2), "%H:%M").time()
        return start, end

    @staticmethod
    def _suggest_free_slots(
        requested_start: Any,
        requested_end: Any,
        occupied: list[tuple[Any, ...]],
    ) -> list[dict[str, Any]]:
        """根据已有占用时段，推荐不冲突的时间段."""
        if not occupied:
            return []

        occupied_sorted = sorted(set(occupied))

        # 工作时段
        work_slots = [
            (
                datetime.strptime("08:00", "%H:%M").time(),
                datetime.strptime("12:00", "%H:%M").time(),
            ),
            (
                datetime.strptime("14:00", "%H:%M").time(),
                datetime.strptime("18:00", "%H:%M").time(),
            ),
            (
                datetime.strptime("19:00", "%H:%M").time(),
                datetime.strptime("22:00", "%H:%M").time(),
            ),
        ]

        today = date.today()
        duration_minutes = (
            datetime.combine(today, requested_end)
            - datetime.combine(today, requested_start)
        ).seconds // 60

        suggestions = []
        for slot_start, slot_end in work_slots:
            current = slot_start
            for occ_start, occ_end in occupied_sorted:
                if occ_start <= current:
                    current = max(current, occ_end)
            if current < slot_end:
                free_minutes = (
                    datetime.combine(today, slot_end) - datetime.combine(today, current)
                ).seconds // 60
                if free_minutes >= duration_minutes:
                    end_time = (
                        datetime.combine(today, current)
                        + timedelta(minutes=duration_minutes)
                    ).time()
                    suggestions.append(
                        {
                            "start": current.strftime("%H:%M"),
                            "end": end_time.strftime("%H:%M"),
                        }
                    )

        return suggestions[:3]

    async def list_records(
        self,
        *,
        employee_number: str | None = None,
        date_from: date | None = None,
        date_to: date | None = None,
        session_id: UUID | None = None,
        page: int = 1,
        page_size: int = 20,
        sort_by: str = "training_date",
        sort_order: str = "asc",
        dept_alias_set: set[str] | None = None,
    ) -> tuple[list[TrainingLedger], int]:
        return await self.repo.list_records(
            employee_number=employee_number,
            date_from=date_from,
            date_to=date_to,
            session_id=session_id,
            page=page,
            page_size=page_size,
            sort_by=sort_by,
            sort_order=sort_order,
            dept_alias_set=dept_alias_set,
        )

    async def list_by_department(
        self, department: str, page: int = 1, page_size: int = 200
    ) -> tuple[list[TrainingLedger], int]:
        return await self.repo.list_by_department(
            department=department, page=page, page_size=page_size
        )

    async def list_training_departments(self) -> list[str]:
        "培训模块所有有数据的部门（台账/ESG/年度计划/岗位清单/培训师/培训会话 并集）."
        return await self.repo.list_all_training_departments()

    async def list_custom_training_departments(self) -> list[str]:
        """获取手动添加的自定义培训部门列表"""
        return await self.repo.list_custom_training_departments()

    async def add_custom_training_department(self, name: str) -> dict[str, Any]:
        """添加自定义部门（含去重校验）"""
        from app.modules.hr.training_dept_resolver import resolve_training_department

        # 1. 标准化输入名（复用配置表解析规则）
        normalized = (await resolve_training_department(self.session, name)) or name

        # 2. 检查是否已存在（数据驱动 + 自定义部门并集）
        existing = await self.repo.list_all_training_departments()
        if normalized in existing:
            raise AppException(status_code=409, message=f"部门「{normalized}」已存在")

        # 3. 创建
        dept = await self.repo.add_custom_training_department(normalized)
        return {"name": dept.name, "id": str(dept.id)}

    async def delete_custom_training_department(self, name: str) -> bool:
        """删除自定义部门（软删除）"""
        # 不能删除数据驱动部门（UNION 查询产生的部门，不含自定义部门）
        # 先获取所有部门，再减去自定义部门，得到数据驱动部门
        all_depts = await self.repo.list_all_training_departments()
        custom_depts = await self.repo.list_custom_training_departments()
        data_driven = set(all_depts) - set(custom_depts)

        if name in data_driven:
            raise AppException(
                status_code=400, message=f"部门「{name}」有培训数据，不可删除"
            )
        return await self.repo.delete_custom_training_department(name)

    # ─── 培训部门映射配置（HR 设置维护）───

    async def list_dept_mappings(self) -> list[dict[str, Any]]:
        """全部映射配置（解析层与前端共用）"""
        mappings = await self.repo.list_dept_mappings()
        return [_dept_mapping_to_dict(m) for m in mappings]

    async def create_dept_mapping(
        self, payload: Any, user_id: Any = None
    ) -> dict[str, Any]:
        "新增映射（source+match_level+mapping_type+target 四重查重，含软删除行兜底）"
        from sqlalchemy.exc import IntegrityError

        existing = await self.repo.list_dept_mappings()
        for m in existing:
            if (
                m.source_name == payload.source_name
                and m.match_level == payload.match_level
                and m.mapping_type == payload.mapping_type
                and m.target_name == payload.target_name
            ):
                raise AppException(
                    status_code=409,
                    message=(
                        f"映射已存在：{payload.source_name}"
                        f"（{payload.mapping_type}/{payload.match_level}/"
                        f"{payload.target_name or '∅'}）"
                    ),
                )
        try:
            mapping = await self.repo.create_dept_mapping(
                {
                    "source_name": payload.source_name,
                    "target_name": payload.target_name,
                    "match_level": payload.match_level,
                    "mapping_type": payload.mapping_type,
                    "priority": payload.priority,
                    "enabled": payload.enabled,
                    "remark": payload.remark,
                    "created_by": user_id,
                }
            )
        except IntegrityError:
            # 软删除行仍在表中（唯一约束含软删行），查重未命中时兜底转 409
            await self.session.rollback()
            raise AppException(
                status_code=409,
                message=f"映射已存在（含已删除记录）：{payload.source_name}（{payload.mapping_type}/{payload.match_level}），如需恢复请直接编辑同源记录",
            )
        # 解析层缓存失效，下次解析自动重载
        from app.modules.hr.training_dept_resolver import (
            invalidate_training_dept_mapping_cache,
        )

        invalidate_training_dept_mapping_cache()
        return _dept_mapping_to_dict(mapping)

    async def update_dept_mapping(
        self, mapping_id: Any, payload: Any, user_id: Any = None
    ) -> dict[str, Any]:
        """更新映射（仅非 None 字段生效）"""
        mapping = await self.repo.get_dept_mapping(mapping_id)
        if not mapping:
            raise AppException(status_code=404, message="映射不存在")
        for field in (
            "source_name",
            "target_name",
            "match_level",
            "mapping_type",
            "priority",
            "enabled",
            "remark",
        ):
            val = getattr(payload, field, None)
            if val is not None:
                setattr(mapping, field, val)
        mapping.updated_by = user_id
        await self.session.flush()
        # flush 后 updated_at（onupdate 服务端值）已过期，需刷新才能安全序列化
        await self.session.refresh(mapping)
        # 解析层缓存失效
        from app.modules.hr.training_dept_resolver import (
            invalidate_training_dept_mapping_cache,
        )

        invalidate_training_dept_mapping_cache()
        return _dept_mapping_to_dict(mapping)

    async def delete_dept_mapping(self, mapping_id: Any) -> bool:
        """软删除映射，返回是否找到"""
        mapping = await self.repo.get_dept_mapping(mapping_id)
        if not mapping:
            return False
        await self.repo.delete_dept_mapping(mapping)
        # 解析层缓存失效
        from app.modules.hr.training_dept_resolver import (
            invalidate_training_dept_mapping_cache,
        )

        invalidate_training_dept_mapping_cache()
        return True

    async def create_from_notification(
        self,
        *,
        employee_number: str,
        training_date: date,
        training_subject: str,
        training_method: str | None,
        trainer: str | None,
        source_id: str | None = None,
    ) -> TrainingLedger | None:
        """当培训通知包含特定员工时，自动创建培训台账记录。"""
        if source_id:
            existing = await self.repo.get_by_source("notification", source_id)
            if existing:
                return existing

        record = TrainingLedger(
            employee_number=employee_number,
            training_date=training_date,
            training_subject=training_subject,
            training_method=training_method,
            trainer=trainer,
            source_type="notification",
            source_id=source_id,
        )
        return await self.repo.create(record)


# ─── 员工培训清单（配置表方案）───


def _split_trainees(trainees: str | None) -> list[str]:
    """拆分培训对象名单文本（支持、，空格分隔），去空去重."""
    if not trainees:
        return []
    names: list[str] = []
    for part in re.split(r"[、，,；;\s]+", trainees):
        part = part.strip()
        if part and part not in names:
            names.append(part)
    return names


def _dedupe_training_records(records: list[Any]) -> list[Any]:
    """同一场培训按涉及部门拆分会产生多条副本，按培训场次去重（session_id 优先，
    无 session 时按 日期+时间+内容 归一）."""
    seen: set[Any] = set()
    result: list[Any] = []
    for r in records:
        if getattr(r, "session_id", None):
            key = f"s:{r.session_id}"
        else:
            key = f"t:{r.training_date}|{r.training_datetime}|{r.training_content}"
        if key in seen:
            continue
        seen.add(key)
        result.append(r)
    return result


class EmployeeTrainingListService:
    """员工培训清单：部门→人员配置维护 + 按人汇总培训台账记录."""

    def __init__(self, session: AsyncSession) -> None:
        self.session = session
        self.member_repo = EmployeeTrainingListRepository(session)
        self.ledger_repo = TrainingLedgerRepository(session)
        self.employee_repo = EmployeeRepository(session)

    async def list_employee_members(self, department: str) -> list[dict[str, Any]]:
        """配置表人员 ∪ 动态新员工（在职档案解析后属于该部门且不在配置表）."""
        configured = await self.member_repo.list_by_department(department)
        result: list[dict[str, Any]] = [
            {
                "id": m.id,
                "department": m.department,
                "name": m.name,
                "employee_number": m.employee_number,
                "source": m.source,
            }
            for m in configured
            # 过滤公用账号（非真实人员，与 HR 联系人展示规范一致）
            if "公用账号" not in (m.name or "")
        ]
        configured_names = {m.name for m in configured}
        # 动态新员工：在职员工档案解析后属于该部门且未在配置表
        auto = await self.employee_repo.list_by_department_for_auto(department)
        for name, emp_no in auto:
            if name not in configured_names:
                result.append(
                    {
                        "id": None,
                        "department": department,
                        "name": name,
                        "employee_number": emp_no,
                        "source": "auto",
                    }
                )
        return result

    async def import_feishu_members(
        self, department: str | None = None
    ) -> dict[str, Any]:
        """从飞书联系人缓存表一键导入人员配置（全部部门或单部门）."""
        from sqlalchemy import select as _select

        from app.modules.hr.training_dept_resolver import (
            _load_mappings,
            resolve_training_department,
        )

        # 候选来源覆写（配置表 candidate_source：目标培训部门 ← 飞书源部门）
        candidate_source = {
            m["source_name"]: m["target_name"]
            for m in await _load_mappings(self.session)
            if m["mapping_type"] == "candidate_source" and m["target_name"]
        }

        stmt = _select(HrFeishuMember).where(HrFeishuMember.is_deleted.is_(False))
        rows = (await self.session.execute(stmt)).scalars().all()

        per_department: dict[str, int] = {}
        total = 0

        async def _upsert(target: str, member: HrFeishuMember) -> None:
            nonlocal total
            await self.member_repo.upsert_member(
                department=target,
                name=member.name,
                employee_number=member.employee_no,
                source="feishu",
            )
            per_department[target] = per_department.get(target, 0) + 1
            total += 1

        for member in rows:
            feishu_dept = (member.department or "").strip()
            if not feishu_dept or "冻结" in feishu_dept:
                continue
            # 过滤公用账号（非真实人员，与 HR 联系人展示规范一致）
            if "公用账号" in (member.name or ""):
                continue
            # 部门映射：SUB_201_MAP → 特殊映射 → 别名归并 → 培训部门列表匹配 → 回退同名
            target = await resolve_training_department(self.session, feishu_dept)
            if not target:
                continue
            if department and target != department:
                # 候选来源覆写（如 IT←AI创新部）：指定部门导入时也检查目标部门
                for tgt, src in candidate_source.items():
                    if tgt == department and target == src:
                        await _upsert(tgt, member)
                continue
            await _upsert(target, member)
            # 候选来源覆写：目标部门的人同时归入覆写部门
            # （如 102二车间（DR）←201二车间（DR），用解析后的培训部门名比较）
            for tgt, src in candidate_source.items():
                if tgt != target and target == src:
                    await _upsert(tgt, member)

        logger.info(
            "employee training list feishu import",
            extra={"department": department or "*", "count": total},
        )
        return {"total": total, "per_department": per_department}

    async def add_member(
        self, department: str, name: str, employee_number: str | None
    ) -> dict[str, Any]:
        """手动添加人员（离职等不在飞书联系人的人员）."""
        member = await self.member_repo.upsert_member(
            department=department,
            name=name,
            employee_number=employee_number,
            source="manual",
        )
        logger.info(
            "employee training list member added",
            extra={"department": department, "member_name": name},
        )
        return {
            "id": member.id,
            "department": member.department,
            "name": member.name,
            "employee_number": member.employee_number,
            "source": member.source,
        }

    async def remove_member(self, member_id: UUID) -> None:
        member = await self.member_repo.get_by_id(member_id)
        if member is None:
            raise NotFoundException("员工培训清单人员", str(member_id))
        await self.member_repo.soft_delete(member)
        logger.info(
            "employee training list member removed",
            extra={"department": member.department, "member_name": member.name},
        )

    async def update_member_name(self, member_id: UUID, name: str) -> dict[str, Any]:
        """编辑人员姓名（改名，保留原来源与部门）."""
        member = await self.member_repo.get_by_id(member_id)
        if member is None:
            raise NotFoundException("员工培训清单人员", str(member_id))
        member.name = name.strip()
        await self.session.flush()
        logger.info(
            "employee training list member renamed",
            extra={"department": member.department, "member_name": name},
        )
        return {
            "id": member.id,
            "department": member.department,
            "name": member.name,
            "employee_number": member.employee_number,
            "source": member.source,
        }

    async def list_employee_training_summary(
        self,
        department: str,
        name: str | None = None,
        date_from: date | None = None,
        date_to: date | None = None,
    ) -> list[dict[str, Any]]:
        """该部门成员逐人按姓名从全量台账匹配记录，返回汇总列表."""
        members = await self.list_employee_members(department)
        if name:
            members = [m for m in members if name in m["name"]]
        ledgers = await self.ledger_repo.list_all_for_employee_list()
        # 姓名 → 匹配的台账记录（trainees 精确拆分匹配）
        records_by_name: dict[str, list[Any]] = {m["name"]: [] for m in members}
        for ledger in ledgers:
            names = _split_trainees(ledger.trainees)
            for n in names:
                if n in records_by_name:
                    if (
                        date_from
                        and ledger.training_date
                        and ledger.training_date < date_from
                    ):
                        continue
                    if (
                        date_to
                        and ledger.training_date
                        and ledger.training_date > date_to
                    ):
                        continue
                    records_by_name[n].append(ledger)
        # 同一培训多部门副本去重
        records_by_name = {
            k: _dedupe_training_records(v) for k, v in records_by_name.items()
        }

        result: list[dict[str, Any]] = []
        for m in members:
            recs = records_by_name[m["name"]]
            dates = [r.training_date for r in recs if r.training_date]
            result.append(
                {
                    "name": m["name"],
                    "employee_number": m.get("employee_number"),
                    "source": m["source"],
                    "record_count": len(recs),
                    "first_training_date": min(dates) if dates else None,
                    "last_training_date": max(dates) if dates else None,
                }
            )
        return result

    async def get_employee_training_records(
        self,
        name: str,
        date_from: date | None = None,
        date_to: date | None = None,
    ) -> list[dict[str, Any]]:
        """某员工全部匹配台账记录（含个人考核结果，导出用）."""
        from app.modules.hr.exam_score_parser import extract_personal_score

        ledgers = await self.ledger_repo.list_all_for_employee_list()
        result: list[dict[str, Any]] = []
        seen: set[Any] = set()
        for ledger in ledgers:
            if name not in _split_trainees(ledger.trainees):
                continue
            if date_from and ledger.training_date and ledger.training_date < date_from:
                continue
            if date_to and ledger.training_date and ledger.training_date > date_to:
                continue
            # 同一培训多部门副本去重（仅保留一条）
            if getattr(ledger, "session_id", None):
                key = f"s:{ledger.session_id}"
            else:
                key = (
                    f"t:{ledger.training_date}|{ledger.training_datetime}|"
                    f"{ledger.training_content}"
                )
            if key in seen:
                continue
            seen.add(key)
            result.append(
                {
                    "training_datetime": ledger.training_datetime,
                    "training_date": ledger.training_date,
                    "training_content": ledger.training_content,
                    "personal_score": extract_personal_score(
                        name, ledger.score_summary, ledger.ledger_assessment_method
                    ),
                    "remarks": ledger.remarks,
                }
            )
        return result


class TrainingLedgerPageService:
    def __init__(self, session: AsyncSession) -> None:
        self.repo = TrainingLedgerPageRepository(session)

    async def list_pages(self) -> list[TrainingLedgerPage]:
        return await self.repo.list_pages()

    async def list_pages_with_department(
        self,
    ) -> list[tuple[TrainingLedgerPage, str | None]]:
        return await self.repo.list_pages_with_department()

    async def create_page(self, data: Any) -> TrainingLedgerPage:
        existing = await self.repo.get_by_employee_number(data.employee_number)
        if existing:
            raise DuplicateException("培训台账页面", data.employee_number)
        page = TrainingLedgerPage(**data.model_dump())
        return await self.repo.create(page)


# ─── 年度培训计划 Word 文档解析辅助（APP1/APP2 格式）───

_TRAINING_CHECKED_SYMBOLS = ("☑", "√", "R", "V")
_TRAINING_UNCHECKED_SYMBOLS = ("□", "£", "×")

# Wingdings 2 字体符号映射到 Unicode 勾选/未勾选符号
_WINGDINGS2_MAP = {
    "0052": "☑",  # R → 勾选框
    "00A3": "□",  # £ → 空框
    "0050": "□",  # P → 空框（备选）
    "00FE": "☑",  # þ → 勾选框（备选）
}


def _get_cell_full_text(cell: Any) -> str:
    """提取单元格完整文本，包含 w:sym 符号字符。

    python-docx 的 cell.text 只读 w:t 元素，忽略 w:sym（Word 插入符号）。
    本函数按 XML 顺序遍历 w:sym 和 w:t，将符号映射为 Unicode 字符后拼接。
    """
    from docx.oxml.ns import qn

    parts = []
    for elem in cell._tc.iter():
        tag = elem.tag
        if tag == qn("w:t"):
            parts.append(elem.text or "")
        elif tag == qn("w:sym"):
            char_code = elem.get(qn("w:char"), "")
            font = elem.get(qn("w:font"), "")
            if font == "Wingdings 2" and char_code in _WINGDINGS2_MAP:
                parts.append(_WINGDINGS2_MAP[char_code])
            elif char_code:
                # 尝试将十六进制字符码转为 Unicode 字符
                try:
                    parts.append(chr(int(char_code, 16)))
                except (ValueError, OverflowError):
                    pass
    return "".join(parts)


def _parse_training_type(val: str) -> str | None:
    """从勾选文本中解析培训类型.

    判断"内训"/"外训"前相邻的符号：
    - ☑/√/R/V 视为勾选
    - □/£/× 视为未勾选
    - 无勾选符号且文本为"内训外训"（空白模板）时返回 None

    Returns:
        "内训" / "外训" / "内训+外训" / None（未填写）
    """
    compact = re.sub(r"\s+", "", val or "")
    if compact == "内训外训":
        return None

    has_symbol = any(
        sym in compact
        for sym in _TRAINING_CHECKED_SYMBOLS + _TRAINING_UNCHECKED_SYMBOLS
    )

    def checked(keyword: str) -> bool:
        idx = compact.find(keyword)
        if idx < 0:
            return False
        before = compact[max(0, idx - 1) : idx]
        return any(sym in before for sym in _TRAINING_CHECKED_SYMBOLS)

    if has_symbol:
        has_neixun = checked("内训")
        has_waixun = checked("外训")
    else:
        has_neixun = "内训" in compact
        has_waixun = "外训" in compact

    if has_neixun and has_waixun:
        return "内训+外训"
    if has_neixun:
        return "内训"
    if has_waixun:
        return "外训"
    return None


def _detect_plan_meta_from_doc(doc: Any) -> tuple[str | None, str | None, str | None]:
    """从年度培训计划文档中自动识别计划级别、部门和版本号。

    识别规则：
    - 计划级别：标题段落含"公司培训计划表"→公司级；含"部门培训计划表"→部门级
    - 部门：段落中"部门：XXX"后的内容，截止到"版本"或空白
    - 版本号：段落中"版本：XX"后的内容

    Returns:
        (plan_level, department, version)，识别不到对应项时返回 None
    """
    plan_level = None
    department = None
    version = None

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue

        # 计划级别：标题（可能带年份前缀，如"2026年度部门培训计划表"）
        if "培训计划表" in text:
            if "公司" in text:
                plan_level = "公司级"
            elif "部门" in text:
                plan_level = "部门级"
            continue

        # 部门 + 版本号：提取"部门：XXX"和"版本：XX"
        if "部门" in text and "：" in text and "培训计划表" not in text:
            seg = text.split("部门", 1)[1]
            if "：" not in seg:
                continue
            rest = seg.split("：", 1)[1]
            # 提取版本号
            for ver_kw in ("版本", "版 本"):
                if ver_kw in rest:
                    ver_part = rest.split(ver_kw, 1)[1]
                    # 版本号在"："后面
                    if "：" in ver_part:
                        ver_val = ver_part.split("：", 1)[1].strip()
                    else:
                        ver_val = ver_part.strip()
                    ver_val = re.sub(r"\s+", "", ver_val)
                    if ver_val:
                        version = ver_val
                    rest = rest.split(ver_kw, 1)[0]
            dept = rest.replace(" ", "").replace("　", "").strip("：: ")
            if dept:
                department = dept

    return plan_level, department, version


def _parse_plan_items_from_doc(doc: Any) -> tuple[list[dict[str, Any]], str | None]:
    """解析文档表格，返回 (明细列表, 计划级备注)。

    遍历所有表格，按表头列名映射字段，跳过空行/审批栏行，
    提取"备注"行作为计划级备注。
    """
    items_data: list[dict[str, Any]] = []
    plan_remarks: str | None = None

    for table in doc.tables:
        # 找到表头行（包含"序号"和"培训类型"的行）
        header_row_idx = None
        col_map: dict[int, str] = {}  # 列索引 -> 字段名

        for ri, row in enumerate(table.rows):
            cells_text = [_get_cell_full_text(cell).strip() for cell in row.cells]
            row_text = " ".join(cells_text)

            if "序号" in row_text and ("培训类型" in row_text or "类型" in row_text):
                header_row_idx = ri
                for ci, text in enumerate(cells_text):
                    if "序号" in text:
                        col_map[ci] = "seq"
                    elif "培训类型" in text or "类型" in text:
                        col_map[ci] = "training_type"
                    elif "培训时间" in text or "月度" in text:
                        col_map[ci] = "training_month"
                    elif "培训内容" in text or "教材" in text or "内容" in text:
                        col_map[ci] = "content_textbook"
                    elif "培训对象" in text or "对象" in text:
                        col_map[ci] = "target_audience_new"
                    elif "授课" in text:
                        col_map[ci] = "instructor"
                    elif "考核" in text:
                        col_map[ci] = "assessment_method"
                break

        if header_row_idx is None:
            continue

        # 从表头行下一行开始读取数据
        for ri in range(header_row_idx + 1, len(table.rows)):
            row = table.rows[ri]
            cells_text = [_get_cell_full_text(cell).strip() for cell in row.cells]

            # 跳过空行
            if all(t == "" for t in cells_text):
                continue

            row_text = " ".join(cells_text)

            # 备注行：提取备注内容（"备注"开头的行）
            # 注意："备注"标签常合并占多格，内容格因水平合并会重复，
            # 且多页表格中后续页的备注行可能只是破折号占位符，需过滤并避免
            # 覆盖首个有效备注。
            if cells_text and cells_text[0].strip().startswith("备注"):
                seen: set[str] = set()
                uniq: list[str] = []
                for t in cells_text:
                    s = t.strip()
                    if not s or s == "备注" or s == "：" or s in seen:
                        continue
                    seen.add(s)
                    body = re.sub(r"\s+", "", s)
                    # 纯破折号/下划线等占位符丢弃
                    if body and all(ch in "—－-_＿~～.·•―" for ch in body):
                        continue
                    uniq.append(s)
                remark_text = " ".join(uniq).strip()
                if remark_text and plan_remarks is None:
                    plan_remarks = remark_text
                continue

            # 跳过审批栏行（包含"制表人"、"签名"、"……"等）
            if any(
                kw in row_text
                for kw in [
                    "制表人",
                    "签名",
                    "……",
                    "审批",
                    "复核",
                    "审核",
                    "批准",
                    "意见",
                ]
            ):
                continue

            # 跳过只有"内训 外训"没有其他内容的空白模板行（兼容不同空格/全角空白）
            non_empty = [
                t for t in cells_text if t and re.sub(r"\s+", "", t) != "内训外训"
            ]
            if not non_empty:
                continue

            # 提取数据 — 同一字段映射多列时只取第一个非空值
            item: dict[str, Any] = {}
            for ci, field in col_map.items():
                if field == "seq":
                    continue
                if field in item and item[field]:
                    continue  # 已有值，跳过
                if ci < len(cells_text):
                    val = cells_text[ci]
                    if field == "training_type":
                        # 培训类型特殊处理：从勾选文本中提取内训/外训（支持
                        # ☑/□、R/£、√/× 等）
                        parsed = _parse_training_type(val)
                        if parsed:
                            item[field] = parsed
                    elif val and re.sub(r"\s+", "", val) != "内训外训":
                        item[field] = val

            # 至少有一个有效字段才算数据行
            if any(v for v in item.values()):
                item["sort_order"] = len(items_data)
                items_data.append(item)

    return items_data, plan_remarks


class AnnualTrainingPlanService:
    def __init__(self, session: AsyncSession) -> None:
        self.repo = AnnualTrainingPlanRepository(session)
        self.item_repo = AnnualTrainingPlanItemRepository(session)
        self.session = session

    async def get_plan(self, plan_id: UUID) -> AnnualTrainingPlan:
        plan = await self.repo.get_by_id(plan_id)
        if not plan:
            raise NotFoundException("年度培训计划", str(plan_id))
        return plan

    async def create_plan(self, data: AnnualTrainingPlanCreate) -> AnnualTrainingPlan:
        if not hasattr(self, "session"):
            existing = await self.repo.get_by_year_and_department(
                data.year, data.department
            )
            if existing:
                raise DuplicateException(
                    "年度培训计划", f"{data.year}年-{data.department}"
                )
            return await self.repo.create(AnnualTrainingPlan(**data.model_dump()))

        # 落库前归一：手输部门写法（201二车间/动力科）→ 培训规范名，与 Word 导入一致；
        # 公司级计划部门固定为「公司」（不在映射表，幂等）
        from app.modules.hr.training_dept_resolver import resolve_training_department

        department = data.department or "公司"
        normalized = (
            await resolve_training_department(self.session, department) or department
        )
        existing = await self.repo.get_by_year_and_department(data.year, normalized)
        if existing:
            raise DuplicateException("年度培训计划", f"{data.year}年-{normalized}")
        plan = AnnualTrainingPlan(**{**data.model_dump(), "department": normalized})
        return await self.repo.create(plan)

    async def update_plan(
        self, plan_id: UUID, data: AnnualTrainingPlanUpdate
    ) -> AnnualTrainingPlan:
        plan = await self.get_plan(plan_id)
        update_data = data.model_dump(exclude_unset=True)
        if not hasattr(self, "session"):
            for field, value in update_data.items():
                setattr(plan, field, value)
            return await self.repo.update(plan)

        from app.modules.hr.training_dept_resolver import resolve_training_department

        if update_data.get("department"):
            dept = await resolve_training_department(
                self.session, update_data["department"]
            )
            update_data["department"] = dept or update_data["department"]
        for field, value in update_data.items():
            setattr(plan, field, value)
        return await self.repo.update(plan)

    async def delete_plan(self, plan_id: UUID) -> None:
        plan = await self.get_plan(plan_id)
        await self.repo.soft_delete(plan)
        if not hasattr(self, "session"):
            return
        # 级联软删该计划名下的明细，避免删除头表后明细残留为孤儿数据
        await self.item_repo.delete_by_plan_id(plan_id)

    async def list_plans(
        self,
        *,
        year: int | None = None,
        department: str | None = None,
        page: int = 1,
        page_size: int = 20,
        dept_alias_set: set[str] | None = None,
    ) -> tuple[list[AnnualTrainingPlan], int]:
        return await self.repo.list_plans(
            year=year,
            department=department,
            page=page,
            page_size=page_size,
            dept_alias_set=dept_alias_set,
        )

    async def import_from_docx(
        self,
        content: bytes,
        *,
        year: int,
        plan_level: str | None,
        department: str | None,
        item_service: "AnnualTrainingPlanItemService",
    ) -> dict[str, Any]:
        """从 Word 文档字节流导入年度培训计划。

        自动识别计划级别/部门/版本号，解析表格明细与备注，
        创建或复用计划并全量替换明细。

        Returns:
            {"plan_id": str, "imported_count": int}
        """
        from docx import Document

        doc = Document(BytesIO(content))

        # 文档识别优先，识别不到再回退到调用方参数
        detected_level, detected_dept, detected_version = _detect_plan_meta_from_doc(
            doc
        )
        if detected_level:
            plan_level = detected_level
        if detected_dept:
            # 落库前归一：文档中的部门写法（如 201二车间（多拉）/201二车间）→ 培训规范名
            from app.modules.hr.training_dept_resolver import (
                resolve_training_department,
            )

            department = await resolve_training_department(self.session, detected_dept)
        plan_level = plan_level or "公司级"
        department = department or ("公司" if plan_level == "公司级" else None)
        if department is None:
            raise AppException(
                status_code=400, message="无法从文档中识别部门，请手动指定部门"
            )

        items_data, plan_remarks = _parse_plan_items_from_doc(doc)
        if not items_data:
            raise AppException(
                status_code=400,
                message=(
                    "文档中未找到有效的培训计划数据，请确保使用 APP1/APP2 格式的 Word"
                    " 文档"
                ),
            )

        # 查找或创建计划
        try:
            plan = await self.create_plan(
                AnnualTrainingPlanCreate(
                    year=year,
                    department=department,
                    plan_level=plan_level,
                    version=detected_version,
                    remarks=plan_remarks,
                )
            )
        except DuplicateException:
            # 同年同部门计划已存在，复用并更新版本号/备注
            existing = await self.repo.get_by_year_and_department(year, department)
            if not existing:
                raise
            plan = existing
            update_fields: dict[str, Any] = {}
            if detected_version:
                update_fields["version"] = detected_version
            if plan_remarks:
                update_fields["remarks"] = plan_remarks
            if update_fields:
                plan = await self.update_plan(
                    plan.id, AnnualTrainingPlanUpdate(**update_fields)
                )

        # 全量替换明细
        batch_data = AnnualTrainingPlanItemBatchUpdate(
            items=[AnnualTrainingPlanItemCreate(**item) for item in items_data]
        )
        await item_service.batch_update_items(plan.id, batch_data)

        return {"plan_id": str(plan.id), "imported_count": len(items_data)}


class AnnualTrainingPlanItemService:
    def __init__(self, session: AsyncSession) -> None:
        self.repo = AnnualTrainingPlanItemRepository(session)
        self.plan_repo = AnnualTrainingPlanRepository(session)

    async def list_items(self, plan_id: UUID) -> list[AnnualTrainingPlanItem]:
        return await self.repo.list_items(plan_id)

    async def batch_update_items(
        self, plan_id: UUID, data: AnnualTrainingPlanItemBatchUpdate
    ) -> list[AnnualTrainingPlanItem]:
        plan = await self.plan_repo.get_by_id(plan_id)
        if not plan:
            raise NotFoundException("年度培训计划", str(plan_id))

        # 删除旧明细
        await self.repo.delete_by_plan_id(plan_id)

        # 创建新明细
        results: list[AnnualTrainingPlanItem] = []
        for idx, item_data in enumerate(data.items):
            item = AnnualTrainingPlanItem(
                plan_id=plan_id,
                sort_order=idx,
                **item_data.model_dump(exclude={"sort_order"}),
            )
            created = await self.repo.create(item)
            results.append(created)
        return results


class TrainingPersonnelConfigService:
    def __init__(self, session: AsyncSession) -> None:
        self.repo = TrainingPersonnelConfigRepository(session)
        self.contract_repo = ContractManagementRepository(session)
        self.session = session

    async def list_configs(
        self,
        *,
        level: str | None = None,
        department: str | None = None,
    ) -> list[TrainingPersonnelConfig]:
        return await self.repo.list_configs(level=level, department=department)

    async def upsert_config(
        self, data: TrainingPersonnelConfigCreate
    ) -> TrainingPersonnelConfig:
        "按 (level, department, config_n"
        "ame) upsert：存在则更新 personnel/re"
        "marks，否则创建"
        existing = await self.repo.get_by_key(
            data.level, data.department, data.config_name
        )
        if existing:
            return await self.repo.update_fields(
                existing,
                personnel=data.personnel,
                remarks=data.remarks,
            )
        record = TrainingPersonnelConfig(
            level=data.level,
            department=data.department,
            config_name=data.config_name,
            personnel=data.personnel,
            remarks=data.remarks,
        )
        return await self.repo.create(record)

    async def delete_config(self, config_id: UUID) -> None:
        """软删除一条人员配置"""
        from app.core.exceptions import NotFoundException

        record = await self.repo.get_by_id(config_id)
        if not record:
            raise NotFoundException("培训人员配置", str(config_id))
        await self.repo.soft_delete(record)

    async def list_new_hires(
        self, days: int = 7, dept_alias_set: set[str] | None = None
    ) -> list[NewHireOut]:
        """最近 days 天内进厂的新员工，按 (姓名, 部门) 去重"""
        today = date.today()
        start = today - timedelta(days=days)
        emp_repo = EmployeeRepository(self.session)
        employees = await emp_repo.list_recent_entries(
            start_date=start, end_date=today, dept_alias_set=dept_alias_set
        )
        seen: set[tuple[str, str]] = set()
        result: list[NewHireOut] = []
        for emp in employees:
            key = (emp.name, emp.department or "")
            if key in seen:
                continue
            seen.add(key)
            result.append(
                NewHireOut(
                    employee_number=emp.employee_number or "",
                    name=emp.name,
                    department=emp.department,
                    factory_entry_date=emp.factory_entry_date,
                )
            )
        return result


# ─── 年度培训计划附件 ───


async def _llm_chat_json_with_retry(
    messages: Any, expected_keys: Any, temperature: Any = 0.1, max_retries: Any = 3
) -> Any:
    """LLM chat_json 带重试：最多 3 次，指数退避 (1s, 2s, 4s)。"""
    import asyncio as _asyncio

    for attempt in range(max_retries):
        try:
            return await llm_client.chat_json(
                messages, expected_keys=expected_keys, temperature=temperature
            )
        except LLMRateLimitError:
            if attempt == max_retries - 1:
                raise
            await _asyncio.sleep(2**attempt)
        except (LLMOutputError, LLMProviderError) as e:
            if attempt == max_retries - 1:
                raise
            logger.warning("LLM retry %d/%d: %s", attempt + 1, max_retries, e)
            await _asyncio.sleep(2**attempt)


class PlanAttachmentService:
    def __init__(self, session: AsyncSession) -> None:
        self.repo = PlanAttachmentRepository(session)
        self.section_repo = PlanAttachmentSectionRepository(session)

    async def list_by_plan(self, plan_id: UUID) -> list[PlanAttachment]:
        return await self.repo.list_by_plan(plan_id)

    async def upload(
        self,
        plan_id: UUID,
        file_name: str,
        data: bytes,
        annex_no: str | None = None,
    ) -> PlanAttachment:
        """保存附件并自动拆分附件条目.

        条目拆分顺序：规则解析（sheet名/标题段）→ AI 兜底 → 文件名编号（整文件一条）。
        条目落 plan_attachment_sections，供跨模块按 计划+附件号 索引。
        文件存储：MinIO 启用时走 app.core.storage（hr bucket），否则回退数据库。
        """
        import mimetypes
        from uuid import uuid4

        file_name = safe_upload_filename(file_name)
        storage_key: str | None = None
        file_data: bytes | None = data
        if storage.is_enabled():
            ct = mimetypes.guess_type(file_name)[0] or "application/octet-stream"
            storage_key = f"plan_attachments/{plan_id}/{uuid4().hex}_{file_name}"
            storage.upload_object("hr", storage_key, data, len(data), ct)
            file_data = None
        attachment = PlanAttachment(
            plan_id=plan_id,
            annex_no=annex_no or normalize_annex_no(file_name),
            file_name=file_name,
            file_data=file_data,
            storage_key=storage_key,
            file_size=len(data),
        )
        attachment = await self.repo.create(attachment)

        drafts = parse_sections(file_name, data)
        if not drafts:
            drafts = await self._ai_infer_sections(file_name, data)
        if not drafts and attachment.annex_no:
            drafts = [
                SectionDraft(
                    annex_no=attachment.annex_no,
                    title=file_name,
                    source_kind="whole_file",
                    source_ref=None,
                )
            ]
        if not drafts:
            # 单文件无编号：与计划行"附件X"引用匹配（确定性子串优先，AI 语义兜底）
            drafts = await self._match_to_plan_refs(plan_id, file_name, data)

        seen: set[str] = set()
        for d in drafts:
            if d.annex_no in seen:
                continue
            seen.add(d.annex_no)
            await self.section_repo.create(
                PlanAttachmentSection(
                    attachment_id=attachment.id,
                    plan_id=plan_id,
                    annex_no=d.annex_no,
                    title=d.title,
                    source_kind=d.source_kind,
                    source_ref=d.source_ref,
                )
            )
        return attachment

    async def _ai_infer_sections(
        self, file_name: str, data: bytes
    ) -> list[SectionDraft]:
        """规则解析失败时的 AI 兜底：基于结构大纲推断附件编号/位置.

        任何异常都降级为空列表（不阻塞上传）。
        """
        outline = build_outline(file_name, data)
        if outline.get("kind") == "unknown":
            return []
        try:
            if outline["kind"] == "xlsx":
                prompt = (
                    "以下是Excel的sheet名及前两行摘要。请判断哪些sheet对应培训附件，"
                    '并给出附件序号（阿拉伯数字）。只返回JSON：{"sections":[{"sheet":sheet名,"n":序号,"title":标题}]}。'
                    f"sheet信息：{outline['sheets']}"
                )
                result = await _llm_chat_json_with_retry(
                    [{"role": "user", "content": prompt}],
                    expected_keys=["sections"],
                    temperature=0.1,
                )
                names = {s["name"] for s in outline["sheets"]}
                drafts = []
                for s in result.get("sections", []):
                    if (
                        s.get("sheet") in names
                        and isinstance(s.get("n"), int)
                        and s["n"] > 0
                    ):
                        drafts.append(
                            SectionDraft(
                                annex_no=f"附件{s['n']}",
                                title=s.get("title") or s["sheet"],
                                source_kind="xlsx_sheet",
                                source_ref=s["sheet"],
                            )
                        )
                return drafts
            prompt = (
                "以下是Word文档中带索引的短行文本。请判断哪些行是附件标题（如附件一/附件1），"
                '返回JSON：{"sections":[{"index":行索引,"n":附件序号,"title":该行文本}]}。'
                f"行列表：{outline['lines']}"
            )
            result = await _llm_chat_json_with_retry(
                [{"role": "user", "content": prompt}],
                expected_keys=["sections"],
                temperature=0.1,
            )
            idx_map = {line["index"]: line["text"] for line in outline["lines"]}
            drafts = []
            for s in result.get("sections", []):
                if (
                    s.get("index") in idx_map
                    and isinstance(s.get("n"), int)
                    and s["n"] > 0
                ):
                    drafts.append(
                        SectionDraft(
                            annex_no=f"附件{s['n']}",
                            title=s.get("title") or idx_map[s["index"]],
                            source_kind="docx_section",
                            source_ref=str(s["index"]),
                        )
                    )
            return drafts
        except LLMOutputError as e:
            logger.error("AI 附件兜底解析输出格式错误: %s", e)
            return []
        except LLMRateLimitError as e:
            logger.warning("AI 附件兜底解析速率限制: %s", e)
            return []
        except LLMProviderError as e:
            logger.error("AI 附件兜底解析服务调用失败: %s", e)
            return []
        except Exception as e:
            logger.exception("AI 附件兜底解析未预期异常: %s", e)
            return []

    async def _match_to_plan_refs(
        self, plan_id: UUID, file_name: str, data: bytes
    ) -> list[SectionDraft]:
        """无编号单文件 → 匹配计划行的"附件X"引用.

        先确定性子串比对（文件名 vs 行内容去标点后互相包含），
        再 AI 语义匹配；只认领尚未被占用的附件号。
        """
        items = await AnnualTrainingPlanItemRepository(self.repo.session).list_items(
            plan_id
        )
        existing = {s.annex_no for s in await self.section_repo.list_by_plan(plan_id)}
        ref_rows: list[tuple[str, str]] = []
        for it in items:
            content = it.content_textbook or it.content_and_textbook or ""
            for ref in extract_annex_refs(content):
                if ref not in existing and ref not in [r for r, _ in ref_rows]:
                    ref_rows.append((ref, content))
        if not ref_rows:
            return []

        stem = strip_punct(file_name.rsplit(".", 1)[0])
        for ref, content in ref_rows:
            core = strip_punct(ANNEX_RE.sub("", content))
            if len(stem) >= 4 and len(core) >= 4 and (stem in core or core in stem):
                return [
                    SectionDraft(
                        annex_no=ref,
                        title=file_name,
                        source_kind="whole_file",
                        source_ref=None,
                    )
                ]

        # AI 语义匹配（单文件 → 单个附件号或空）
        try:
            outline = build_outline(file_name, data)
            if outline.get("kind") == "xlsx":
                heads = [s["name"] for s in outline.get("sheets", [])][:10]
            elif outline.get("kind") == "docx":
                heads = [line["text"] for line in outline.get("lines", [])][:10]
            else:
                heads = []
            ref_payload = [{"ref": r, "content": c} for r, c in ref_rows]
            prompt = (
                f"年度培训计划中有以下带附件编号的培训内容行：{ref_payload}。"
                f"现上传了一个未编号的附件文件，文件名：{file_name}，内容概要：{heads}。"
                '请判断该文件最可能对应哪个附件编号；若都不对应返回空字符串。只返回JSON：{"ref":"附件N或空"}。'
            )
            result = await _llm_chat_json_with_retry(
                [{"role": "user", "content": prompt}],
                expected_keys=["ref"],
                temperature=0.1,
            )
            ref = result.get("ref") or ""
            if ref in [r for r, _ in ref_rows]:
                return [
                    SectionDraft(
                        annex_no=ref,
                        title=file_name,
                        source_kind="whole_file",
                        source_ref=None,
                    )
                ]
        except LLMOutputError as e:
            logger.error("AI 附件匹配输出格式错误: %s", e)
        except LLMRateLimitError as e:
            logger.warning("AI 附件匹配速率限制: %s", e)
        except LLMProviderError as e:
            logger.error("AI 附件匹配服务调用失败: %s", e)
        except Exception as e:
            logger.exception("AI 附件匹配未预期异常: %s", e)
        return []

    async def get(self, attachment_id: UUID) -> PlanAttachment:
        attachment = await self.repo.get_by_id(attachment_id)
        if not attachment:
            raise NotFoundException("计划附件", str(attachment_id))
        return attachment

    def read_data(self, attachment: PlanAttachment) -> bytes:
        """读取附件二进制：优先 MinIO（storage_key），回退数据库字段."""
        if attachment.storage_key and storage.is_enabled():
            obj = storage.get_object("hr", attachment.storage_key)
            if obj is not None:
                return obj[0]
        return attachment.file_data or b""

    async def delete(self, attachment_id: UUID) -> None:
        attachment = await self.get(attachment_id)
        if attachment.storage_key and storage.is_enabled():
            try:
                storage.delete_object("hr", attachment.storage_key)
            except Exception:
                logger.exception("删除 MinIO 附件对象失败: %s", attachment.storage_key)
        await self.section_repo.soft_delete_by_attachment(attachment_id)
        await self.repo.soft_delete(attachment)

    async def mark_ledger_imported(self, ids: list[UUID]) -> int:
        """标记附件已导入培训台账（幂等），返回本次标记数量."""
        now = datetime.now(UTC)
        count = 0
        for aid in ids:
            attachment = await self.repo.get_by_id(aid)
            if attachment and attachment.ledger_imported_at is None:
                attachment.ledger_imported_at = now
                count += 1
        await self.repo.session.flush()
        return count

    # ── 附件条目 / 预览 ──

    async def list_sections(self, plan_id: UUID) -> list[PlanAttachmentSection]:
        return await self.section_repo.list_by_plan(plan_id)

    async def get_section(self, section_id: UUID) -> PlanAttachmentSection:
        section = await self.section_repo.get_by_id(section_id)
        if not section:
            raise NotFoundException("附件条目", str(section_id))
        return section

    async def preview_section(self, section_id: UUID) -> dict[str, Any]:
        section = await self.get_section(section_id)
        attachment = await self.get(section.attachment_id)
        return build_preview(
            attachment.file_name,
            self.read_data(attachment),
            section.source_kind,
            section.source_ref,
        )

    async def preview_attachment(self, attachment_id: UUID) -> dict[str, Any]:
        attachment = await self.get(attachment_id)
        return build_preview(
            attachment.file_name, self.read_data(attachment), "whole_file", None
        )
