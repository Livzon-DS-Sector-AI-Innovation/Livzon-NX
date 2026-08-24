"""培训附件 Word 文档生成器（模板保真填充）.

模板结构：加粗段落"附件：" + Table Grid 三列表格（序号/文件名称/文件编号），
模板预置 4 行数据行。仅填充数据区，保留表头格式与表格边框；
数据行多于预置行时按最后一行克隆扩展，少于时删除多余行。
"""

import logging
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.modules.hr.template_filler import (
    clone_row_after,
    delete_row,
    fill_whole_cell,
    header_fmt_source,
)

logger = logging.getLogger(__name__)

TEMPLATE_NAME = "培训附件.docx"

# 数据区正文字号：四号 = 14pt（业务口径 2026-08-13）
BODY_FONT_SIZE = Pt(14)


def _find_template() -> Path:
    candidates = [
        Path("员工培训教育管理规程") / TEMPLATE_NAME,
        Path(__file__).resolve().parent.parent.parent
        / "员工培训教育管理规程"
        / TEMPLATE_NAME,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {TEMPLATE_NAME}")


def generate_training_attachment(items: list[Any]) -> BytesIO:
    """按模板保真生成培训附件 Word（附件： + 序号/文件名称/文件编号表格）.

    items: TrainingAttachmentItem 列表（或含 name/code 的对象/dict）。
    模板结构：R0 表头 + R1..R4 预置数据行，最终数据行数 = len(items)。
    """
    doc = Document(str(_find_template()))
    if not doc.tables:
        raise ValueError("模板文件中未找到表格")

    table = doc.tables[0]
    fmt_src = header_fmt_source(table)

    need = len(items)
    # 行数调整：不足克隆最后一行，多余删除尾部预置行
    while len(table.rows) - 1 < need:
        clone_row_after(table, len(table.rows) - 1)
    while len(table.rows) - 1 > need:
        delete_row(table, len(table.rows) - 1)

    for i, item in enumerate(items, start=1):
        row = table.rows[i]
        fill_whole_cell(row.cells[0], str(i), fmt_src)
        fill_whole_cell(row.cells[1], getattr(item, "name", "") or "", fmt_src)
        fill_whole_cell(row.cells[2], getattr(item, "code", "") or "", fmt_src)
        # 数据区正文字体统一四号；序号/文件编号列内容完全居中（水平 + 垂直）
        for ci, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = BODY_FONT_SIZE
                if ci in (0, 2):
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ci in (0, 2):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    logger.info(
        "training attachment generated",
        extra={"item_count": len(items), "module_name": "hr"},
    )
    return buffer
