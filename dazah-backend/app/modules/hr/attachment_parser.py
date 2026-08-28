"""年度培训计划附件解析器：附件条目拆分（规则）+ 预览结构化.

规则解析优先（准确、零成本）：
- xlsx：sheet 名 / 首格含"附件X" → 每 sheet 一条条目；名为"全部附件"的汇总 sheet
  跳过；若仅有汇总 sheet 则按首列"附件X"标记行在 sheet 内切分。
- docx：以"附件X"开头的短标题段切分条目。
解析不出时由 service 层走 AI 兜底（本模块只提供结构大纲提取）。
"""

import re
from dataclasses import dataclass
from io import BytesIO
from typing import Any

from docx import Document
from docx.oxml.ns import qn as docx_qn
from openpyxl import load_workbook  # type: ignore[import-untyped]

# 附件编号匹配（阿拉伯/全角/中文数字）
ANNEX_RE = re.compile(r"附件\s*([0-9０-９一二三四五六七八九十]+)")

_CN_DIGIT = {
    "一": 1,
    "二": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
}

# 预览行数上限，防止响应过大
PREVIEW_ROW_LIMIT = 300
PREVIEW_BLOCK_LIMIT = 400


def _cn_to_int(s: str) -> int | None:
    """中文数字转整数（支持 1-99，如一/十/十五/二十/二十一）."""
    if not s:
        return None
    if s.isdigit():
        return int(s)
    if "十" not in s:
        return _CN_DIGIT.get(s)
    tens, _, ones = s.partition("十")
    t = _CN_DIGIT.get(tens, 1) if tens else 1
    o = _CN_DIGIT.get(ones, 0) if ones else 0
    return t * 10 + o


def _to_half_digits(s: str) -> str:
    return s.translate(str.maketrans("０１２３４５６７８９", "0123456789"))


def normalize_annex_no(text: str | None) -> str | None:
    """从文本中提取"附件X"引用并归一化为"附件{n}"（n 为阿拉伯数字）.

    支持 附件一 / 附件1 / 附件 12 等写法；提取不到返回 None。
    """
    m = ANNEX_RE.search(text or "")
    if not m:
        return None
    n = _cn_to_int(_to_half_digits(m.group(1)))
    return f"附件{n}" if n else None


def extract_annex_refs(text: str | None) -> list[str]:
    """提取文本中所有"附件X"引用（归一化、去重、保序）."""
    refs: list[str] = []
    for m in ANNEX_RE.finditer(text or ""):
        n = _cn_to_int(_to_half_digits(m.group(1)))
        key = f"附件{n}" if n else ""
        if key and key not in refs:
            refs.append(key)
    return refs


def strip_punct(s: str) -> str:
    """去除空白与常见标点，用于文件名与行内容的模糊比对."""
    return re.sub(
        r"[\s（）()【】\[\]、，,。．.·\-—_：:；;！!？?%％0-9０-９]+", "", s or ""
    )


@dataclass
class SectionDraft:
    """解析出的附件条目草稿（尚未落库）."""

    annex_no: str
    title: str | None
    source_kind: str  # xlsx_sheet / docx_section / whole_file
    source_ref: str | None  # sheet 名 / 起始段落下标 / None


def _file_ext(file_name: str) -> str:
    return (
        (file_name or "").rsplit(".", 1)[-1].lower() if "." in (file_name or "") else ""
    )


# ─── 规则解析 ───


def parse_sections(file_name: str, data: bytes) -> list[SectionDraft]:
    """按规则拆分附件条目；解析不到返回空列表（交由 AI 兜底）."""
    ext = _file_ext(file_name)
    if ext == "xlsx":
        return _parse_xlsx_sections(data)
    if ext == "docx":
        return _parse_docx_sections(data)
    return []


def _parse_xlsx_sections(data: bytes) -> list[SectionDraft]:
    wb = load_workbook(BytesIO(data), data_only=True, read_only=True)
    drafts: list[SectionDraft] = []
    try:
        for name in wb.sheetnames:
            if "全部" in name:
                continue  # 汇总 sheet 跳过
            annex = normalize_annex_no(name)
            if not annex:
                ws = wb[name]
                first_cell = None
                for row in ws.iter_rows(
                    min_row=1, max_row=1, max_col=1, values_only=True
                ):
                    first_cell = row[0] if row else None
                annex = normalize_annex_no(str(first_cell)) if first_cell else None
            if annex:
                drafts.append(
                    SectionDraft(
                        annex_no=annex,
                        title=name,
                        source_kind="xlsx_sheet",
                        source_ref=name,
                    )
                )
        if not drafts:
            # 仅有汇总 sheet：按首列"附件X"标记行在 sheet 内切分
            for name in wb.sheetnames:
                drafts.extend(_split_sheet_by_first_col(wb, name))
    finally:
        wb.close()
    return drafts


def _split_sheet_by_first_col(wb: Any, sheet_name: str) -> list[SectionDraft]:
    """汇总 sheet 内按首列"附件X"标记切分（只记录标记，预览时再切行）."""
    ws = wb[sheet_name]
    drafts: list[SectionDraft] = []
    seen: set[str] = set()
    for row in ws.iter_rows(min_col=1, max_col=1, values_only=True):
        cell = row[0] if row else None
        if not cell:
            continue
        annex = normalize_annex_no(str(cell))
        if (
            annex
            and annex not in seen
            and ANNEX_RE.fullmatch(str(cell).strip()) is not None
        ):
            # 仅当单元格就是"附件X"标记本身（非数据行误匹配）
            seen.add(annex)
            drafts.append(
                SectionDraft(
                    annex_no=annex,
                    title=str(cell).strip(),
                    source_kind="xlsx_sheet",
                    source_ref=sheet_name,
                )
            )
    return drafts


def _parse_docx_sections(data: bytes) -> list[SectionDraft]:
    doc = Document(BytesIO(data))
    drafts: list[SectionDraft] = []
    for idx, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t or len(t) > 30:
            continue
        if t.startswith("附件") and ANNEX_RE.match(t):
            annex = normalize_annex_no(t)
            if annex:
                drafts.append(
                    SectionDraft(
                        annex_no=annex,
                        title=t,
                        source_kind="docx_section",
                        source_ref=str(idx),
                    )
                )
    return drafts


# ─── AI 兜底用结构大纲 ───


def build_outline(file_name: str, data: bytes) -> dict[str, Any]:
    """提取供 AI 推断编号用的结构大纲（sheet 名 / 带索引的标题段）."""
    ext = _file_ext(file_name)
    if ext == "xlsx":
        wb = load_workbook(BytesIO(data), data_only=True, read_only=True)
        sheets = []
        try:
            for name in wb.sheetnames:
                ws = wb[name]
                first_cells = []
                for row in ws.iter_rows(min_row=1, max_row=2, values_only=True):
                    first_cells.append(
                        [str(c)[:30] for c in (row or [])[:4] if c is not None]
                    )
                sheets.append({"name": name, "head": first_cells})
        finally:
            wb.close()
        return {"kind": "xlsx", "sheets": sheets}
    if ext == "docx":
        doc = Document(BytesIO(data))
        lines = []
        for idx, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            if t and len(t) <= 40:
                lines.append({"index": idx, "text": t})
            if len(lines) >= 150:
                break
        return {"kind": "docx", "lines": lines}
    return {"kind": "unknown"}


# ─── 预览结构化 ───


def build_preview(
    file_name: str, data: bytes, source_kind: str, source_ref: str | None
) -> dict[str, Any]:
    """把附件内容转成前端可渲染的结构化预览.

    返回:
      {kind: 'table', title, header, rows} 或
      {kind: 'doc', title, blocks: [{type:'p',text}|{type:'table',rows}]} 或
      {kind: 'tables', title, tables: [{title, header, rows}]}
    """
    ext = _file_ext(file_name)
    if ext == "xlsx":
        return _preview_xlsx(data, source_kind, source_ref)
    if ext == "docx":
        return _preview_docx(data, source_kind, source_ref)
    return {
        "kind": "doc",
        "title": file_name,
        "blocks": [{"type": "p", "text": "该文件类型暂不支持预览，请下载查看。"}],
    }


def _sheet_to_table(ws: Any) -> dict[str, Any]:
    header: list[str] = []
    rows: list[list[str]] = []
    for ri, row in enumerate(ws.iter_rows(values_only=True)):
        cells = [("" if c is None else str(c)) for c in row]
        if not any(cells):
            continue
        if not header and ("序号" in cells or "名称" in "".join(cells)):
            header = cells
            continue
        if not header and ri <= 1:
            continue  # 跳过顶部标题行
        rows.append(cells)
        if len(rows) >= PREVIEW_ROW_LIMIT:
            break
    return {"header": header, "rows": rows}


def _preview_xlsx(
    data: bytes, source_kind: str, source_ref: str | None
) -> dict[str, Any]:
    wb = load_workbook(BytesIO(data), data_only=True, read_only=True)
    try:
        if source_kind == "xlsx_sheet" and source_ref in wb.sheetnames:
            ws = wb[source_ref]
            tbl = _sheet_to_table(ws)
            return {"kind": "table", "title": source_ref, **tbl}
        # whole_file：多 sheet 预览（限 5 个）
        tables = []
        for name in wb.sheetnames[:5]:
            tbl = _sheet_to_table(wb[name])
            tables.append({"title": name, **tbl})
        return {"kind": "tables", "title": "全部附件", "tables": tables}
    finally:
        wb.close()


def _doc_body_blocks(doc: Any) -> list[Any]:
    """按文档顺序提取 段落/表格 块."""
    blocks = []
    body = doc.element.body
    para_map = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}
    for child in body:
        if child.tag == docx_qn("w:p"):
            p = para_map.get(child)
            text = (p.text.strip() if p else "") or ""
            if text:
                blocks.append({"type": "p", "text": text})
        elif child.tag == docx_qn("w:tbl"):
            t = table_map.get(child)
            if t is not None:
                rows = [[c.text.strip() for c in row.cells] for row in t.rows][
                    :PREVIEW_ROW_LIMIT
                ]
                blocks.append({"type": "table", "rows": rows})
        if len(blocks) >= PREVIEW_BLOCK_LIMIT:
            break
    return blocks


def _preview_docx(
    data: bytes, source_kind: str, source_ref: str | None
) -> dict[str, Any]:
    doc = Document(BytesIO(data))
    if source_kind == "docx_section" and source_ref is not None:
        start = int(source_ref)
        # 收集从起始段落到下一个"附件X"标题段之间的块
        blocks: list[dict[str, Any]] = []
        started = False
        para_idx = -1
        body = doc.element.body
        para_map = {p._p: p for p in doc.paragraphs}
        table_map = {t._tbl: t for t in doc.tables}
        for child in body:
            if child.tag == docx_qn("w:p"):
                p = para_map.get(child)
                if p is not None:
                    para_idx += 1
                    if para_idx == start:
                        started = True
                        paragraph_text = p.text.strip()
                        blocks.append({"type": "p", "text": paragraph_text})
                        continue
                    if (
                        started
                        and p.text.strip().startswith("附件")
                        and ANNEX_RE.match(p.text.strip())
                        and len(p.text.strip()) <= 30
                    ):
                        break  # 下一附件段，停止
                if started and p is not None and p.text.strip():
                    blocks.append({"type": "p", "text": p.text.strip()})
            elif child.tag == docx_qn("w:tbl") and started:
                table = table_map.get(child)
                if table is not None:
                    rows = [[c.text.strip() for c in row.cells] for row in table.rows][
                        :PREVIEW_ROW_LIMIT
                    ]
                    blocks.append({"type": "table", "rows": rows})
            if len(blocks) >= PREVIEW_BLOCK_LIMIT:
                break
        title = blocks[0]["text"] if blocks and blocks[0]["type"] == "p" else "附件"
        return {"kind": "doc", "title": title, "blocks": blocks[:PREVIEW_BLOCK_LIMIT]}
    # whole_file
    return {"kind": "doc", "title": "附件全文", "blocks": _doc_body_blocks(doc)}
