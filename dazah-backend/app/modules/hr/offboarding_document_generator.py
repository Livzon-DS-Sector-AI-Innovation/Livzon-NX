"""离职文档生成器 - 使用 python-docx 模板渲染"""

import logging
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document

from app.modules.hr.date_format import HR_EXPORT_DATE_FORMAT

logger = logging.getLogger(__name__)

TEMPLATE_DIR = Path(__file__).parent.parent.parent.parent / "templates" / "offboarding"


def generate_termination_notice(employee: dict[str, Any]) -> BytesIO:
    """生成《解除劳动合同通知单》

    Args:
        employee: 员工信息字典，包含 name, id_card, hire_date, department, position 等

    Returns:
        BytesIO: 生成的 Word 文档缓冲区
    """
    template_path = TEMPLATE_DIR / "解除劳动合同单.docx"
    doc = Document(str(template_path))

    replacements = {
        "{姓名}": employee.get("name", ""),
        "{性别}": employee.get("gender", ""),
        "{身份证号}": employee.get("id_card", ""),
        "{入职日期}": employee.get("hire_date", ""),
        "{现家庭地址}": employee.get("current_address", ""),
    }

    today_str = date.today().strftime(HR_EXPORT_DATE_FORMAT)
    today_year = today_str[:4]
    today_month = today_str[5:7]
    today_day = today_str[8:10]

    # 逐 run 替换，保持格式不变
    def replace_in_runs(paragraph: Any) -> Any:
        for run in paragraph.runs:
            # 常规占位符替换
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, str(value))
            # 处理被拆分的日期 YYYY年MM月DD日
            if "YYYY" in run.text:
                run.text = run.text.replace("YYYY", today_year)
            if "MM" in run.text:
                run.text = run.text.replace("MM", today_month)
            if "DD" in run.text:
                run.text = run.text.replace("DD", today_day)

    for paragraph in doc.paragraphs:
        replace_in_runs(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_runs(paragraph)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
