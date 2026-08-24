"""Authorization export helpers based on Word templates."""

from __future__ import annotations

import copy
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.core.config import get_settings
from app.core.exceptions import AppException, NotFoundException
from app.modules.registration.schemas import (
    AuthorizationFdaRecord,
    AuthorizationLedgerMainRead,
)

logger = logging.getLogger(__name__)

FDA_SUFFIX = "-list-of-authorized-parties-to-incorporate-by-reference"
_SOFFICE_CANDIDATES: list[str] = []
if sys.platform == "win32":
    _SOFFICE_CANDIDATES = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
else:
    _SOFFICE_CANDIDATES = [
        "/usr/bin/soffice",
        "/usr/local/bin/soffice",
        "/opt/libreoffice/program/soffice",
    ]


@dataclass
class AuthorizationExportArtifact:
    file_path: Path
    download_name: str
    temp_dir: Path


@dataclass
class MarketExportLine:
    row_type: str
    sequence: str = ""
    authorization_file_name: str = ""
    quality_standard: str = ""
    company_country: str = ""
    customer_code: str = ""
    purpose: str = ""
    authorization_date: str = ""
    handler: str = ""
    remarks: str = ""


@dataclass
class MarketTemplateLayout:
    base_template: Any
    update_template: Any
    repeat_base_columns: set[int]


def _get_authorization_source_dir() -> Path:
    settings = get_settings()
    return Path(settings.REGISTRATION_AUTHORIZATION_SOURCE_DIR).expanduser()


def _get_upload_dir() -> Path:
    settings = get_settings()
    upload_dir_setting = getattr(settings, "UPLOAD_DIR", "uploads")
    base = Path(upload_dir_setting)
    upload_dir = base / "authorization_letters"
    upload_dir.mkdir(parents=True, exist_ok=True)
    return upload_dir


def _extract_product_name(file_path: Path, source_root: Path) -> str:
    stem = file_path.stem
    if stem.endswith(FDA_SUFFIX):
        return stem[: -len(FDA_SUFFIX)]

    if stem.startswith("授权书台帐-"):
        remainder = stem.removeprefix("授权书台帐-")
        return remainder.split("-", 1)[0]

    for parent in file_path.parents:
        if parent == source_root.parent:
            break
        if "授权客户明细" in parent.name:
            return parent.name.replace("授权客户明细", "").strip("- ")

    return file_path.parent.name if file_path.parent != source_root else stem


def _extract_market_name(file_path: Path) -> str | None:
    stem = file_path.stem
    if stem.endswith(FDA_SUFFIX):
        return "FDA"

    if stem.startswith("授权书台帐-"):
        remainder = stem.removeprefix("授权书台帐-")
        parts = remainder.split("-")
        if len(parts) > 1:
            return parts[-1]

    if "授权客户明细" in str(file_path.parent):
        parts = stem.split("-")
        if len(parts) > 2:
            return parts[-1]

    return None


def _should_skip_source_file(file_path: Path) -> bool:
    return file_path.name.startswith("~$")


def _normalize_match_text(value: str | None) -> str:
    return re.sub(r"\s+", "", (value or "").strip()).casefold()


def resolve_fda_template(product_name: str) -> Path:
    source_root = _get_authorization_source_dir()
    if not source_root.exists():
        raise NotFoundException("授权书模板目录", str(source_root))

    target_key = _normalize_match_text(product_name)
    for file_path in sorted(source_root.rglob("*")):
        if not file_path.is_file() or _should_skip_source_file(file_path):
            continue
        if not file_path.stem.endswith(FDA_SUFFIX):
            continue
        if (
            _normalize_match_text(_extract_product_name(file_path, source_root))
            == target_key
        ):
            return file_path

    raise AppException(message="当前产品未配置 FDA 授权模板，无法导出")


def resolve_market_template(product_name: str, market_name: str) -> Path:
    source_root = _get_authorization_source_dir()
    if not source_root.exists():
        raise NotFoundException("授权书模板目录", str(source_root))

    target_product = _normalize_match_text(product_name)
    target_market = _normalize_match_text(market_name)

    for file_path in sorted(source_root.rglob("*")):
        if not file_path.is_file() or _should_skip_source_file(file_path):
            continue
        if "授权书台帐-" not in file_path.stem:
            continue

        file_product = _normalize_match_text(
            _extract_product_name(file_path, source_root)
        )
        file_market = _normalize_match_text(_extract_market_name(file_path))
        if file_product == target_product and file_market == target_market:
            return file_path

    raise AppException(message="当前产品和市场未找到对应授权模板，无法导出")


def _normalize_line_value(value: str | None) -> str:
    normalized = (value or "").strip()
    return normalized if normalized else "-"


def _build_market_document_lines(
    record: AuthorizationLedgerMainRead,
) -> list[MarketExportLine]:
    updates = sorted(
        record.updates, key=lambda item: (item.sort_order, item.created_at)
    )
    first_update = updates[0] if updates else None

    company_country_parts: list[str] = [
        part
        for part in [record.company_name, record.country]
        if isinstance(part, str) and part.strip()
    ]
    company_country = "/".join(company_country_parts) if company_country_parts else "-"
    lines = [
        MarketExportLine(
            row_type="base",
            sequence=(record.source_sequence or "").strip() or "-",
            authorization_file_name=_normalize_line_value(
                record.authorization_file_name
            ),
            quality_standard=_normalize_line_value(record.quality_standard),
            company_country=company_country,
            customer_code=_normalize_line_value(record.customer_code),
            purpose=_normalize_line_value(record.purpose),
            authorization_date=_normalize_line_value(
                first_update.authorization_date if first_update else None
            ),
            handler=_normalize_line_value(
                first_update.handler if first_update else None
            ),
            remarks=_normalize_line_value(
                first_update.remarks if first_update else None
            ),
        )
    ]

    for update in updates[1:]:
        lines.append(
            MarketExportLine(
                row_type="update",
                authorization_date=_normalize_line_value(update.authorization_date),
                handler=_normalize_line_value(update.handler),
                remarks=_normalize_line_value(update.remarks),
            )
        )
    return lines


def _clear_text_in_paragraph(paragraph_elem: Any) -> None:
    for run_elem in paragraph_elem.findall(qn("w:r")):
        for text_elem in run_elem.findall(qn("w:t")):
            text_elem.text = ""


def _set_docx_cell_text(cell_elem: Any, text: str) -> None:
    lines = (text or "").split("\n")
    paragraph_elems = cell_elem.findall(qn("w:p"))

    if not paragraph_elems:
        paragraph_elem = cell_elem.makeelement(qn("w:p"), {})
        cell_elem.append(paragraph_elem)
        paragraph_elems = [paragraph_elem]

    while len(paragraph_elems) > len(lines):
        paragraph_elems[-1].getparent().remove(paragraph_elems[-1])
        paragraph_elems = cell_elem.findall(qn("w:p"))

    while len(paragraph_elems) < len(lines):
        new_paragraph = copy.deepcopy(paragraph_elems[0])
        _clear_text_in_paragraph(new_paragraph)
        cell_elem.append(new_paragraph)
        paragraph_elems = cell_elem.findall(qn("w:p"))

    for index, line in enumerate(lines):
        paragraph_elem = paragraph_elems[index]
        run_elems = paragraph_elem.findall(qn("w:r"))
        for run_elem in run_elems:
            for text_elem in run_elem.findall(qn("w:t")):
                text_elem.text = ""

        if not run_elems:
            run_elem = paragraph_elem.makeelement(qn("w:r"), {})
            paragraph_elem.append(run_elem)
            run_elems = [run_elem]

        text_elems = run_elems[0].findall(qn("w:t"))
        if not text_elems:
            text_elem = run_elems[0].makeelement(qn("w:t"), {})
            run_elems[0].append(text_elem)
            text_elems = [text_elem]

        text_elems[0].text = line


def _get_row_cell_elements(row_elem: Any) -> list[Any]:
    return list(row_elem.findall(qn("w:tc")))


def _get_cell_text(cell_elem: Any) -> str:
    return "".join(node.text or "" for node in cell_elem.iter(qn("w:t"))).strip()


def _set_row_values(row_elem: Any, values: list[str]) -> None:
    cell_elems = _get_row_cell_elements(row_elem)
    if len(cell_elems) < len(values):
        raise AppException(message="当前模板列数不足，无法导出")

    for cell_elem, value in zip(cell_elems, values, strict=False):
        _set_docx_cell_text(cell_elem, value)


def _get_vmerge_state(cell_elem: Any) -> str | None:
    tc_pr = cell_elem.find(qn("w:tcPr"))
    if tc_pr is None:
        return None

    vmerge = tc_pr.find(qn("w:vMerge"))
    if vmerge is None:
        return None

    return vmerge.get(qn("w:val")) or "continue"


def _set_vmerge_state(cell_elem: Any, state: str | None) -> None:
    tc_pr = cell_elem.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        cell_elem.insert(0, tc_pr)

    current_vmerge = tc_pr.find(qn("w:vMerge"))
    if current_vmerge is not None:
        tc_pr.remove(current_vmerge)

    if state is None:
        return

    vmerge = OxmlElement("w:vMerge")
    vmerge.set(qn("w:val"), state)
    tc_pr.append(vmerge)


def _is_market_update_template_row(row_elem: Any) -> bool:
    cell_elems = _get_row_cell_elements(row_elem)
    if len(cell_elems) < 9:
        return False
    return all(_get_vmerge_state(cell_elems[index]) == "continue" for index in range(6))


def _is_market_base_template_row(row_elem: Any) -> bool:
    cell_elems = _get_row_cell_elements(row_elem)
    if len(cell_elems) < 9:
        return False
    return all(_get_vmerge_state(cell_elems[index]) == "restart" for index in range(6))


def _is_blank_market_row(row_elem: Any) -> bool:
    return not any(
        _get_cell_text(cell_elem) for cell_elem in _get_row_cell_elements(row_elem)
    )


def _has_continue_merge(row_elem: Any) -> bool:
    return any(
        _get_vmerge_state(cell_elem) == "continue"
        for cell_elem in _get_row_cell_elements(row_elem)[:6]
    )


def _has_any_vertical_merge(row_elem: Any, *, column_count: int = 6) -> bool:
    return any(
        _get_vmerge_state(cell_elem) is not None
        for cell_elem in _get_row_cell_elements(row_elem)[:column_count]
    )


def _apply_vertical_merge(row_elem: Any, *, state: str, column_count: int = 6) -> None:
    for cell_elem in _get_row_cell_elements(row_elem)[:column_count]:
        _set_vmerge_state(cell_elem, state)


def _build_market_base_row_values(line: MarketExportLine) -> list[str]:
    return [
        line.sequence,
        line.authorization_file_name,
        line.quality_standard,
        line.company_country,
        line.customer_code,
        line.purpose,
        line.authorization_date,
        line.handler,
        line.remarks,
    ]


def _resolve_market_template_layout(row_elements: list[Any]) -> MarketTemplateLayout:
    if len(row_elements) < 2:
        raise AppException(message="市场授权模板格式不正确，无法导出")

    body_rows = row_elements[1:]
    non_blank_rows = [row for row in body_rows if not _is_blank_market_row(row)]
    if not non_blank_rows:
        raise AppException(message="市场授权模板格式不正确，无法导出")

    base_template = copy.deepcopy(non_blank_rows[0])
    update_candidate = next(
        (row for row in body_rows[1:] if _has_continue_merge(row)), None
    )
    if update_candidate is None:
        update_candidate = next(
            (row for row in body_rows[1:] if _is_blank_market_row(row)), None
        )

    if update_candidate is None:
        update_candidate = base_template
    else:
        update_candidate = copy.deepcopy(update_candidate)

    repeat_base_columns: set[int] = set()
    update_cells = _get_row_cell_elements(update_candidate)
    for index, cell_elem in enumerate(update_cells[:6]):
        if _get_vmerge_state(cell_elem) == "continue":
            continue
        if _get_cell_text(cell_elem):
            repeat_base_columns.add(index)

    return MarketTemplateLayout(
        base_template=base_template,
        update_template=update_candidate,
        repeat_base_columns=repeat_base_columns,
    )


def _find_soffice() -> Path | None:
    for candidate in _SOFFICE_CANDIDATES:
        path = Path(candidate)
        if path.exists():
            return path

    for cmd in ("soffice", "libreoffice"):
        found = shutil.which(cmd)
        if found:
            return Path(found)

    return None


def _convert_doc_with_soffice(source_path: Path, output_path: Path) -> bool:
    soffice = _find_soffice()
    if not soffice:
        return False

    generated_path = output_path.parent / f"{source_path.stem}.docx"
    generated_path.unlink(missing_ok=True)
    output_path.unlink(missing_ok=True)

    try:
        result = subprocess.run(
            [
                str(soffice),
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(output_path.parent),
                str(source_path),
            ],
            capture_output=True,
            text=True,
            timeout=120,
            # subprocess 例外：传递环境变量给 LibreOffice 子进程
            env={**os.environ, "SAL_USE_VCLPLUGIN": "gen"},
        )
    except (subprocess.TimeoutExpired, FileNotFoundError):
        return False

    if result.returncode != 0:
        logger.warning(
            "LibreOffice 转换 .doc 失败: %s",
            result.stderr.strip() or result.stdout.strip(),
        )
        return False

    if generated_path.exists():
        if generated_path != output_path:
            shutil.move(str(generated_path), str(output_path))
        return True

    return output_path.exists()


def _convert_doc_with_word(source_path: Path, output_path: Path) -> bool:
    # 仅在遗留 .doc 模板归一化时调用 Word，后续排版与写表全部走 docx/xml。
    quoted_source = str(source_path.resolve()).replace("'", "''")
    quoted_output = str(output_path.resolve()).replace("'", "''")
    script = (
        "$word = $null; $doc = $null; "
        "try { "
        "$word = New-Object -ComObject Word.Application; "
        "$word.Visible = $false; "
        "$word.DisplayAlerts = 0; "
        f"$doc = $word.Documents.Open('{quoted_source}', $false, $true); "
        f"$doc.SaveAs2('{quoted_output}', 16); "
        "$doc.Close([ref]0); "
        "$doc = $null; "
        "exit 0 "
        "} catch { "
        "Write-Error $_.Exception.Message; "
        "exit 1 "
        "} finally { "
        "if ($doc -ne $null) { try { $doc.Close([ref]0) } catch {} } "
        "if ($word -ne $null) { try { $word.Quit() } catch {} } "
        "}"
    )
    try:
        result = subprocess.run(
            ["powershell", "-NoProfile", "-Command", script],
            capture_output=True,
            text=True,
            timeout=120,
        )
    except (subprocess.TimeoutExpired, FileNotFoundError):
        return False

    if result.returncode != 0:
        logger.warning(
            "Word 转换 .doc 失败: %s", result.stderr.strip() or result.stdout.strip()
        )
        return False

    return output_path.exists()


def _copy_to_temp(
    template_path: Path, download_name: str
) -> AuthorizationExportArtifact:
    temp_dir = Path(
        tempfile.mkdtemp(prefix="registration-auth-export-", dir=str(_get_upload_dir()))
    )
    output_path = temp_dir / download_name
    shutil.copy2(template_path, output_path)
    return AuthorizationExportArtifact(
        file_path=output_path, download_name=download_name, temp_dir=temp_dir
    )


def _build_fda_download_name(product_name: str, template_path: Path) -> str:
    return f"FDA授权-{product_name}{template_path.suffix.lower()}"


def _build_market_download_name(
    product_name: str, market_name: str, template_path: Path
) -> str:
    del template_path
    return f"市场授权-{product_name}-{market_name}.docx"


def _prepare_market_export_artifact(
    template_path: Path, product_name: str, market_name: str
) -> AuthorizationExportArtifact:
    download_name = _build_market_download_name(
        product_name, market_name, template_path
    )
    temp_dir = Path(
        tempfile.mkdtemp(prefix="registration-auth-export-", dir=str(_get_upload_dir()))
    )
    output_path = temp_dir / download_name

    if template_path.suffix.lower() == ".docx":
        shutil.copy2(template_path, output_path)
        return AuthorizationExportArtifact(
            file_path=output_path, download_name=download_name, temp_dir=temp_dir
        )

    converted = _convert_doc_with_soffice(template_path, output_path)
    if not converted:
        converted = _convert_doc_with_word(template_path, output_path)

    if not converted:
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise AppException(message="当前 .doc 模板无法归一化为 .docx，无法导出")

    return AuthorizationExportArtifact(
        file_path=output_path, download_name=download_name, temp_dir=temp_dir
    )


def render_fda_export(
    *,
    product_name: str,
    records: list[AuthorizationFdaRecord],
) -> AuthorizationExportArtifact:
    template_path = resolve_fda_template(product_name)
    artifact = _copy_to_temp(
        template_path, _build_fda_download_name(product_name, template_path)
    )

    try:
        doc = Document(str(artifact.file_path))
        if not doc.tables:
            raise AppException(message="FDA 授权模板格式不正确，无法导出")

        table = doc.tables[0]
        row_elements = table._tbl.findall(qn("w:tr"))
        if len(row_elements) < 2:
            raise AppException(message="FDA 授权模板格式不正确，无法导出")

        template_cells = _get_row_cell_elements(row_elements[1])
        template_column_count = len(template_cells)
        if template_column_count not in {6, 7}:
            raise AppException(message="FDA 授权模板格式不正确，无法导出")

        template_row = copy.deepcopy(row_elements[1])
        for row_elem in list(row_elements[1:]):
            table._tbl.remove(row_elem)

        for index, record in enumerate(records, start=1):
            row_elem = copy.deepcopy(template_row)
            table._tbl.append(row_elem)
            values = [
                (record.company_name or "").strip() or "-",
                (record.address or "").strip() or "-",
                (record.reference_number or "").strip() or "-",
                (record.loa_date or "").strip() or "-",
                (record.submission_date or "").strip() or "-",
                (record.referenced_sections or "").strip() or "-",
            ]
            if template_column_count == 7:
                values = [str(index), *values]
            _set_row_values(row_elem, values)

        doc.save(str(artifact.file_path))
        return artifact
    except Exception:
        shutil.rmtree(artifact.temp_dir, ignore_errors=True)
        raise


def render_market_export(
    *,
    product_name: str,
    market_name: str,
    records: list[AuthorizationLedgerMainRead],
) -> AuthorizationExportArtifact:
    template_path = resolve_market_template(product_name, market_name)
    artifact = _prepare_market_export_artifact(template_path, product_name, market_name)

    try:
        doc = Document(str(artifact.file_path))
        if not doc.tables:
            raise AppException(message="市场授权模板格式不正确，无法导出")

        table = doc.tables[0]
        row_elements = table._tbl.findall(qn("w:tr"))
        if len(row_elements) < 2:
            raise AppException(message="市场授权模板格式不正确，无法导出")
        layout = _resolve_market_template_layout(row_elements)

        for row_elem in list(row_elements[1:]):
            table._tbl.remove(row_elem)

        for record in records:
            lines = _build_market_document_lines(record)
            base_line = lines[0]
            base_values = _build_market_base_row_values(base_line)
            update_lines = lines[1:]
            synthesize_grouped_merge = bool(update_lines) and not (
                _has_any_vertical_merge(layout.base_template)
                or _has_any_vertical_merge(layout.update_template)
            )

            base_row = copy.deepcopy(layout.base_template)
            table._tbl.append(base_row)
            _set_row_values(base_row, base_values)
            if synthesize_grouped_merge:
                _apply_vertical_merge(base_row, state="restart")

            for update_line in update_lines:
                update_row = copy.deepcopy(layout.update_template)
                table._tbl.append(update_row)
                cell_elems = _get_row_cell_elements(update_row)
                if len(cell_elems) < 9:
                    raise AppException(message="当前模板列数不足，无法导出")

                update_values = [""] * 9
                for index in layout.repeat_base_columns:
                    update_values[index] = base_values[index]
                update_values[6] = update_line.authorization_date
                update_values[7] = update_line.handler
                update_values[8] = update_line.remarks
                _set_row_values(update_row, update_values)
                if synthesize_grouped_merge:
                    _apply_vertical_merge(update_row, state="continue")

        doc.save(str(artifact.file_path))
        return artifact
    except Exception:
        shutil.rmtree(artifact.temp_dir, ignore_errors=True)
        raise
