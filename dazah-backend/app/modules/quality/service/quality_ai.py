"""Unified quality AI service."""

from __future__ import annotations

import base64
import io
import os
import uuid
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any

import pytesseract
from docx import Document
from fastapi import UploadFile
from openpyxl import load_workbook
from PIL import Image
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.core.database import async_session_factory
from app.core.llm import llm_client
from app.core.llm.config import get_config
from app.core.storage import delete_object, upload_object
from app.core.storage import is_enabled as minio_enabled
from app.modules.quality.models import (
    CAPA,
    ChangeControl,
    Deviation,
    DeviationAiSession,
    DeviationAiSessionAttachment,
    QualityAiAnalysisLog,
)
from app.modules.quality.schemas.deviation_ai_session import (
    DeviationAiSessionAttachmentOut,
    DeviationAiSessionOut,
    DeviationAiSessionResultPayload,
)
from app.modules.quality.schemas.quality_ai import (
    QualityAiApplicableField,
    QualityAiAnalysisLogOut,
)

QUALITY_AI_RESPONSE_KEYS = [
    "summary",
    "risk_level",
    "risks",
    "suggestions",
    "missing_info",
    "structured_fields",
]

DEVIATION_AI_SESSION_RESPONSE_KEYS = [
    "deviation_analysis",
    "capa_suggestion",
]


def _to_iso(value: Any) -> Any:
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if isinstance(value, uuid.UUID):
        return str(value)
    if isinstance(value, dict):
        return {k: _to_iso(v) for k, v in value.items()}
    if isinstance(value, list):
        return [_to_iso(v) for v in value]
    return value


def _build_applicable_fields(
    entity_type: str, analysis_type: str, output_payload: dict | None
) -> list[QualityAiApplicableField]:
    structured_fields = (output_payload or {}).get("structured_fields") or {}
    field_map: dict[str, tuple[str, str]] = {}

    if entity_type == "deviation" and analysis_type == "deviation_analysis":
        field_map = {
            "preliminary_cause_analysis": ("root_cause_analysis", "应用到根因分析"),
            "capa_suggestions": ("corrective_actions", "应用到纠正预防措施"),
        }
    elif entity_type == "deviation" and analysis_type == "capa_suggestion":
        field_map = {
            "capa_suggestions": ("corrective_actions", "应用到纠正预防措施"),
        }
    elif entity_type == "capa":
        field_map = {
            "root_cause_analysis": ("root_cause_analysis", "应用到根因分析"),
            "capa_content": ("capa_content", "应用到CAPA措施"),
            "effectiveness_review": ("evaluation_target", "应用到有效性评估"),
        }
    elif entity_type == "change":
        field_map = {
            "impact_assessment": ("impact_assessment", "应用到影响评估"),
        }

    fields: list[QualityAiApplicableField] = []
    for source_key, (field_key, label) in field_map.items():
        value = structured_fields.get(source_key)
        if value:
            fields.append(
                QualityAiApplicableField(
                    field_key=field_key,
                    label=label,
                    description=f"来自 AI 结构化字段 {source_key}",
                )
            )
    return fields


def _log_to_schema(log: QualityAiAnalysisLog) -> QualityAiAnalysisLogOut:
    return QualityAiAnalysisLogOut(
        id=log.id,
        entity_type=log.entity_type,
        entity_id=log.entity_id,
        analysis_type=log.analysis_type,
        input_snapshot=log.input_snapshot,
        output_payload=log.output_payload,
        model_name=log.model_name,
        status=log.status,
        error_message=log.error_message,
        is_applied=log.is_applied,
        created_at=log.created_at,
        created_by=log.created_by,
        applied_at=log.applied_at,
        applied_by=log.applied_by,
        applicable_fields=_build_applicable_fields(
            log.entity_type, log.analysis_type, log.output_payload
        ),
    )


def _normalize_result(raw: dict) -> dict[str, Any]:
    structured_fields = raw.get("structured_fields") or {}
    risks = raw.get("risks") or []
    suggestions = raw.get("suggestions") or []
    missing_info = raw.get("missing_info") or []

    return {
        "summary": raw.get("summary") or "",
        "risk_level": raw.get("risk_level") or "",
        "risks": risks if isinstance(risks, list) else [str(risks)],
        "suggestions": suggestions
        if isinstance(suggestions, list)
        else [str(suggestions)],
        "missing_info": missing_info
        if isinstance(missing_info, list)
        else [str(missing_info)],
        "structured_fields": structured_fields if isinstance(structured_fields, dict) else {},
        "disclaimer": "AI 结果仅供辅助判断，需人工复核后使用",
    }


def _result_payload_to_schema(
    entity_type: str,
    analysis_type: str,
    payload: dict[str, Any] | None,
) -> DeviationAiSessionResultPayload | None:
    if payload is None:
        return None
    normalized = _normalize_result(payload)
    return DeviationAiSessionResultPayload(
        summary=normalized["summary"],
        risk_level=normalized["risk_level"],
        risks=normalized["risks"],
        suggestions=normalized["suggestions"],
        missing_info=normalized["missing_info"],
        structured_fields=normalized["structured_fields"],
        disclaimer=normalized["disclaimer"],
        applicable_fields=_build_applicable_fields(
            entity_type,
            analysis_type,
            normalized,
        ),
    )


def _attachment_to_schema(
    attachment: DeviationAiSessionAttachment,
) -> DeviationAiSessionAttachmentOut:
    return DeviationAiSessionAttachmentOut(
        id=attachment.id,
        file_name=attachment.file_name,
        file_type=attachment.file_type,
        file_size=attachment.file_size,
        parse_status=attachment.parse_status,
        parse_error=attachment.parse_error,
        parsed_summary=attachment.parsed_summary,
    )


def _build_deviation_snapshot(deviation: Deviation) -> dict[str, Any]:
    return _to_iso(
        {
            "deviation_code": deviation.deviation_code,
            "title": deviation.title,
            "department": deviation.department,
            "discovery_date": deviation.discovery_date,
            "discovery_time": deviation.discovery_time,
            "discovery_location": deviation.discovery_location,
            "level": deviation.level,
            "description": deviation.description,
            "immediate_actions": deviation.immediate_actions,
            "affected_items": deviation.affected_items,
            "batch_number": deviation.batch_number,
            "root_cause_analysis": deviation.root_cause_analysis,
            "corrective_actions": deviation.corrective_actions,
        }
    )


def _deviation_conversation_prompt(
    deviation_snapshot: dict[str, Any],
    supplement_text: str,
    attachment_summary: str,
) -> str:
    supplement_block = supplement_text.strip() or "无"
    attachment_block = attachment_summary.strip() or "无"
    return f"""
你是原料药工厂质量管理助理，只能基于提供的记录、补充说明和附件摘要进行分析，禁止编造事实。

当前任务：针对同一条偏差同时输出两块内容：
1. 偏差分析
2. CAPA建议

偏差原始记录：
{deviation_snapshot}

用户补充信息：
{supplement_block}

附件解析摘要：
{attachment_block}

请严格输出 JSON，字段必须完整：
{{
  "deviation_analysis": {{
    "summary": "一句话概括当前偏差结论",
    "risk_level": "低/中/高 或留空",
    "risks": ["风险点1", "风险点2"],
    "suggestions": ["建议1", "建议2"],
    "missing_info": ["仍缺失的信息1"],
    "structured_fields": {{
      "preliminary_cause_analysis": "根因分析建议",
      "capa_suggestions": "建议写入偏差纠正预防措施字段的内容"
    }}
  }},
  "capa_suggestion": {{
    "summary": "一句话概括建议的CAPA方向",
    "risk_level": "低/中/高 或留空",
    "risks": ["CAPA相关风险1"],
    "suggestions": ["CAPA建议1", "CAPA建议2"],
    "missing_info": ["CAPA仍缺失的信息1"],
    "structured_fields": {{
      "capa_suggestions": "建议写入偏差纠正预防措施字段的CAPA内容"
    }}
  }}
}}
""".strip()


def _build_capa_snapshot(capa: CAPA) -> dict[str, Any]:
    return _to_iso(
        {
            "capa_code": capa.capa_code,
            "title": capa.title,
            "deviation_id": capa.deviation_id,
            "source": capa.source,
            "source_code": capa.source_code,
            "category": capa.category,
            "department": capa.department,
            "affected_product": capa.affected_product,
            "non_conformity_description": capa.non_conformity_description,
            "root_cause_analysis": capa.root_cause_analysis,
            "capa_content": capa.capa_content,
            "capa_items": capa.capa_items,
            "execution_status": capa.execution_status,
            "evaluation_result": capa.evaluation_result,
            "closure_date": capa.closure_date,
        }
    )


def _extract_deviation_code_from_capa_code(capa_code: str | None) -> str | None:
    if not capa_code or not capa_code.startswith("CAPA-"):
        return None
    deviation_code = capa_code.removeprefix("CAPA-").strip()
    return deviation_code or None


async def _resolve_related_deviation_for_capa(
    db: AsyncSession, capa: CAPA
) -> tuple[Deviation | None, str | None]:
    if capa.deviation_id:
        deviation = await db.get(Deviation, capa.deviation_id)
        if deviation and not deviation.is_deleted:
            return deviation, "deviation_id"

    if capa.source == "deviation" and capa.source_code:
        result = await db.execute(
            select(Deviation).where(
                Deviation.is_deleted == False,
                Deviation.deviation_code == capa.source_code,
            )
        )
        deviation = result.scalar_one_or_none()
        if deviation:
            return deviation, "source_code"

    derived_deviation_code = _extract_deviation_code_from_capa_code(capa.capa_code)
    if derived_deviation_code:
        result = await db.execute(
            select(Deviation).where(
                Deviation.is_deleted == False,
                Deviation.deviation_code == derived_deviation_code,
            )
        )
        deviation = result.scalar_one_or_none()
        if deviation:
            return deviation, "capa_code"

    return None, None


async def _build_capa_analysis_snapshot(
    db: AsyncSession, capa: CAPA
) -> dict[str, Any]:
    capa_snapshot = _build_capa_snapshot(capa)
    linked_deviation, match_rule = await _resolve_related_deviation_for_capa(db, capa)
    deviation_snapshot = (
        _build_deviation_snapshot(linked_deviation) if linked_deviation else None
    )
    return _to_iso(
        {
            "capa": capa_snapshot,
            "linked_deviation": deviation_snapshot,
            "analysis_context": {
                "has_linked_deviation": linked_deviation is not None,
                "match_rule": match_rule,
                "analysis_basis": (
                    "capa_and_deviation"
                    if linked_deviation is not None
                    else "capa_only"
                ),
            },
        }
    )


def _build_change_snapshot(change: ChangeControl) -> dict[str, Any]:
    return _to_iso(
        {
            "change_code": change.change_code,
            "serial_number": change.serial_number,
            "applicant_department": change.applicant_department,
            "change_object": change.change_object,
            "change_content": change.change_content,
            "impact_assessment": change.impact_assessment,
            "change_level": change.change_level,
            "application_date": change.application_date,
            "planned_approval_date": change.planned_approval_date,
            "execution_date": change.execution_date,
            "closure_date": change.closure_date,
        }
    )


def _deviation_prompt(snapshot: dict[str, Any], analysis_type: str) -> str:
    if analysis_type == "capa_suggestion":
        task_text = (
            "请只重点生成偏差对应的纠正预防措施建议，强调措施是否对因、是否可执行、是否可验证。"
        )
    else:
        task_text = (
            "请完成偏差摘要、可能根因、风险评估、临时控制措施建议和 CAPA 建议。"
        )
    return f"""
你是原料药工厂质量管理助理，只能基于提供数据分析，禁止编造事实或法规条文。
当前任务：{task_text}

偏差数据：
{snapshot}

请严格返回 JSON，字段必须完整：
{{
  "summary": "一句话概括事件",
  "risk_level": "低/中/高 或留空",
  "risks": ["风险点1", "风险点2"],
  "suggestions": ["建议1", "建议2"],
  "missing_info": ["缺失信息1"],
  "structured_fields": {{
    "preliminary_cause_analysis": "根因分析建议",
    "capa_suggestions": "纠正预防措施建议"
  }}
}}
""".strip()


def _capa_prompt(snapshot: dict[str, Any]) -> str:
    analysis_context = snapshot.get("analysis_context") or {}
    linked_deviation = snapshot.get("linked_deviation")
    if linked_deviation:
        context_text = f"""
本次分析已识别到关联偏差，必须结合 CAPA 与偏差完整详情一起判断措施是否真正对因。

关联规则：
{analysis_context.get("match_rule") or ""}

关联偏差完整详情：
{linked_deviation}
""".strip()
    else:
        context_text = (
            "当前未识别到关联偏差，本次只能基于 CAPA 自身内容分析；"
            "请在结论中明确指出缺少偏差背景。"
        )
    return f"""
你是原料药工厂质量管理助理，只能基于提供数据分析，禁止编造事实或法规条文。
请审查以下 CAPA 记录，重点判断措施是否对因、是否缺责任人/时限、是否可验证，并给出有效性评估建议。

分析上下文：
{context_text}

CAPA 数据：
{snapshot.get("capa") or {}}

请严格返回 JSON，字段必须完整：
{{
  "summary": "一句话概括 CAPA 当前状态",
  "risk_level": "低/中/高 或留空",
  "risks": ["风险点1", "风险点2"],
  "suggestions": ["建议1", "建议2"],
  "missing_info": ["缺失信息1"],
  "structured_fields": {{
    "root_cause_analysis": "更完整的根因分析建议",
    "capa_content": "更可执行的 CAPA 措施建议",
    "effectiveness_review": "有效性评估建议"
  }}
}}
""".strip()


def _change_prompt(snapshot: dict[str, Any]) -> str:
    return f"""
你是原料药工厂质量管理助理，只能基于提供数据分析，禁止编造事实或法规条文。
请对以下变更记录进行影响分析，输出受影响范围、文件/培训/验证建议和遗漏项提醒。

变更数据：
{snapshot}

请严格返回 JSON，字段必须完整：
{{
  "summary": "一句话概括变更事项",
  "risk_level": "低/中/高 或留空",
  "risks": ["风险点1", "风险点2"],
  "suggestions": ["建议1", "建议2"],
  "missing_info": ["缺失信息1"],
  "structured_fields": {{
    "impact_assessment": "影响评估建议",
    "validation_actions": ["验证活动1", "验证活动2"],
    "affected_documents": ["文件1", "文件2"]
  }}
}}
""".strip()


async def _create_log(
    db: AsyncSession,
    *,
    entity_type: str,
    entity_id: uuid.UUID,
    analysis_type: str,
    input_snapshot: dict[str, Any],
    model_name: str,
    status: str,
    output_payload: dict[str, Any] | None = None,
    error_message: str | None = None,
    user_id: str = "system",
) -> QualityAiAnalysisLog:
    log = QualityAiAnalysisLog(
        entity_type=entity_type,
        entity_id=entity_id,
        analysis_type=analysis_type,
        input_snapshot=input_snapshot,
        output_payload=output_payload,
        model_name=model_name,
        status=status,
        error_message=error_message,
        created_by=uuid.UUID(user_id) if user_id != "system" else None,
    )
    db.add(log)
    await db.flush()
    return log


async def analyze_deviation_record(
    db: AsyncSession,
    deviation_id: uuid.UUID,
    user_id: str,
    *,
    analysis_type: str = "deviation_analysis",
    transition_status: bool = False,
) -> QualityAiAnalysisLogOut:
    deviation = await db.get(Deviation, deviation_id)
    if not deviation or deviation.is_deleted:
        raise ValueError("偏差不存在")

    snapshot = _build_deviation_snapshot(deviation)
    prompt = _deviation_prompt(snapshot, analysis_type)
    config = await get_config("text")

    try:
        raw = await llm_client.chat_json(
            [{"role": "user", "content": prompt}],
            expected_keys=QUALITY_AI_RESPONSE_KEYS,
            temperature=0.2,
        )
        result = _normalize_result(raw)
        log = await _create_log(
            db,
            entity_type="deviation",
            entity_id=deviation.id,
            analysis_type=analysis_type,
            input_snapshot=snapshot,
            output_payload=result,
            model_name=config.model_name,
            status="completed",
            user_id=user_id,
        )
        if analysis_type == "deviation_analysis":
            deviation.ai_analysis = result
        if transition_status and deviation.status == "pending_ai_analysis":
            deviation.status = "pending_investigation"
            deviation.status_updated_at = datetime.now(timezone.utc)
            deviation.updated_by = uuid.UUID(user_id) if user_id != "system" else None
        await db.commit()
        await db.flush()
        return _log_to_schema(log)
    except Exception as exc:
        log = await _create_log(
            db,
            entity_type="deviation",
            entity_id=deviation.id,
            analysis_type=analysis_type,
            input_snapshot=snapshot,
            model_name=config.model_name,
            status="failed",
            error_message=str(exc),
            user_id=user_id,
        )
        await db.commit()
        await db.flush()
        raise RuntimeError(str(exc)) from exc


async def suggest_capa_for_deviation(
    db: AsyncSession, deviation_id: uuid.UUID, user_id: str
) -> QualityAiAnalysisLogOut:
    return await analyze_deviation_record(
        db,
        deviation_id,
        user_id,
        analysis_type="capa_suggestion",
        transition_status=False,
    )


async def analyze_capa_record(
    db: AsyncSession, capa_id: uuid.UUID, user_id: str
) -> QualityAiAnalysisLogOut:
    capa = await db.get(CAPA, capa_id)
    if not capa or capa.is_deleted:
        raise ValueError("CAPA 不存在")

    snapshot = await _build_capa_analysis_snapshot(db, capa)
    prompt = _capa_prompt(snapshot)
    config = await get_config("text")

    try:
        raw = await llm_client.chat_json(
            [{"role": "user", "content": prompt}],
            expected_keys=QUALITY_AI_RESPONSE_KEYS,
            temperature=0.2,
        )
        result = _normalize_result(raw)
        log = await _create_log(
            db,
            entity_type="capa",
            entity_id=capa.id,
            analysis_type="capa_review",
            input_snapshot=snapshot,
            output_payload=result,
            model_name=config.model_name,
            status="completed",
            user_id=user_id,
        )
        await db.commit()
        await db.flush()
        return _log_to_schema(log)
    except Exception as exc:
        log = await _create_log(
            db,
            entity_type="capa",
            entity_id=capa.id,
            analysis_type="capa_review",
            input_snapshot=snapshot,
            model_name=config.model_name,
            status="failed",
            error_message=str(exc),
            user_id=user_id,
        )
        await db.commit()
        await db.flush()
        raise RuntimeError(str(exc)) from exc


async def analyze_change_record(
    db: AsyncSession, change_id: uuid.UUID, user_id: str
) -> QualityAiAnalysisLogOut:
    change = await db.get(ChangeControl, change_id)
    if not change or change.is_deleted:
        raise ValueError("变更不存在")

    snapshot = _build_change_snapshot(change)
    prompt = _change_prompt(snapshot)
    config = await get_config("text")

    try:
        raw = await llm_client.chat_json(
            [{"role": "user", "content": prompt}],
            expected_keys=QUALITY_AI_RESPONSE_KEYS,
            temperature=0.2,
        )
        result = _normalize_result(raw)
        log = await _create_log(
            db,
            entity_type="change",
            entity_id=change.id,
            analysis_type="change_impact",
            input_snapshot=snapshot,
            output_payload=result,
            model_name=config.model_name,
            status="completed",
            user_id=user_id,
        )
        await db.commit()
        await db.flush()
        return _log_to_schema(log)
    except Exception as exc:
        log = await _create_log(
            db,
            entity_type="change",
            entity_id=change.id,
            analysis_type="change_impact",
            input_snapshot=snapshot,
            model_name=config.model_name,
            status="failed",
            error_message=str(exc),
            user_id=user_id,
        )
        await db.commit()
        await db.flush()
        raise RuntimeError(str(exc)) from exc


async def list_ai_logs(
    db: AsyncSession,
    *,
    entity_type: str | None = None,
    entity_id: uuid.UUID | None = None,
    page: int = 1,
    page_size: int = 20,
) -> dict[str, Any]:
    filters = [QualityAiAnalysisLog.is_deleted == False]
    if entity_type:
        filters.append(QualityAiAnalysisLog.entity_type == entity_type)
    if entity_id:
        filters.append(QualityAiAnalysisLog.entity_id == entity_id)

    query = select(QualityAiAnalysisLog).where(*filters)
    count_query = select(func.count()).select_from(QualityAiAnalysisLog).where(*filters)

    total = (await db.execute(count_query)).scalar() or 0
    result = await db.execute(
        query.order_by(QualityAiAnalysisLog.created_at.desc())
        .offset((page - 1) * page_size)
        .limit(page_size)
    )
    items = result.scalars().all()
    return {
        "items": [_log_to_schema(item).model_dump() for item in items],
        "total": total,
        "page": page,
        "page_size": page_size,
    }


async def get_ai_log_detail(
    db: AsyncSession, log_id: uuid.UUID
) -> QualityAiAnalysisLogOut:
    log = await db.get(QualityAiAnalysisLog, log_id)
    if not log or log.is_deleted:
        raise ValueError("AI 分析记录不存在")
    return _log_to_schema(log)


async def apply_ai_log(
    db: AsyncSession,
    log_id: uuid.UUID,
    field_keys: list[str],
    user_id: str,
) -> QualityAiAnalysisLogOut:
    log = await db.get(QualityAiAnalysisLog, log_id)
    if not log or log.is_deleted:
        raise ValueError("AI 分析记录不存在")
    if log.status != "completed" or not log.output_payload:
        raise ValueError("只有成功的 AI 分析记录才可以应用")

    structured_fields = (log.output_payload or {}).get("structured_fields") or {}
    allowed = {
        field.field_key: field
        for field in _build_applicable_fields(
            log.entity_type, log.analysis_type, log.output_payload
        )
    }
    if not field_keys:
        raise ValueError("请选择要应用的字段")
    invalid = [field for field in field_keys if field not in allowed]
    if invalid:
        raise ValueError(f"存在不允许应用的字段: {', '.join(invalid)}")

    if log.entity_type == "deviation":
        entity = await db.get(Deviation, log.entity_id)
        if not entity or entity.is_deleted:
            raise ValueError("偏差不存在")
        if "root_cause_analysis" in field_keys:
            entity.root_cause_analysis = structured_fields.get(
                "preliminary_cause_analysis"
            )
        if "corrective_actions" in field_keys:
            entity.corrective_actions = structured_fields.get("capa_suggestions")
    elif log.entity_type == "capa":
        entity = await db.get(CAPA, log.entity_id)
        if not entity or entity.is_deleted:
            raise ValueError("CAPA 不存在")
        if "root_cause_analysis" in field_keys:
            entity.root_cause_analysis = structured_fields.get("root_cause_analysis")
        if "capa_content" in field_keys:
            entity.capa_content = structured_fields.get("capa_content")
        if "evaluation_target" in field_keys:
            entity.evaluation_target = structured_fields.get("effectiveness_review")
    else:
        entity = await db.get(ChangeControl, log.entity_id)
        if not entity or entity.is_deleted:
            raise ValueError("变更不存在")
        if "impact_assessment" in field_keys:
            entity.impact_assessment = structured_fields.get("impact_assessment")

    entity.updated_by = uuid.UUID(user_id) if user_id != "system" else None
    log.is_applied = True
    log.applied_at = datetime.now(timezone.utc)
    log.applied_by = uuid.UUID(user_id) if user_id != "system" else None
    await db.commit()
    await db.flush()
    return _log_to_schema(log)


def _quality_ai_local_upload_dir() -> Path:
    upload_dir = Path(get_settings().UPLOAD_DIR) / "quality" / "deviation_ai"
    upload_dir.mkdir(parents=True, exist_ok=True)
    return upload_dir


def _truncate_text(value: str, limit: int = 2000) -> str:
    cleaned = value.strip()
    if len(cleaned) <= limit:
        return cleaned
    return f"{cleaned[:limit]}..."


def _build_attachment_summary_from_rows(
    attachments: list[DeviationAiSessionAttachment],
) -> str:
    parts: list[str] = []
    for attachment in attachments:
        if attachment.is_deleted:
            continue
        if attachment.parse_status != "completed":
            continue
        if attachment.parsed_summary:
            parts.append(f"{attachment.file_name}：{attachment.parsed_summary}")
        elif attachment.parsed_text:
            parts.append(f"{attachment.file_name}：{_truncate_text(attachment.parsed_text, 400)}")
    return "\n\n".join(parts)


async def _list_deviation_ai_session_attachments(
    db: AsyncSession,
    session_id: uuid.UUID,
) -> list[DeviationAiSessionAttachment]:
    result = await db.execute(
        select(DeviationAiSessionAttachment)
        .where(
            DeviationAiSessionAttachment.session_id == session_id,
            DeviationAiSessionAttachment.is_deleted == False,
        )
        .order_by(
            DeviationAiSessionAttachment.sort_order.asc(),
            DeviationAiSessionAttachment.created_at.asc(),
        )
    )
    return list(result.scalars().all())


async def _session_to_schema(
    db: AsyncSession,
    session: DeviationAiSession,
) -> DeviationAiSessionOut:
    attachments = await _list_deviation_ai_session_attachments(db, session.id)
    return DeviationAiSessionOut(
        id=session.id,
        deviation_id=session.deviation_id,
        supplement_text=session.supplement_text,
        status=session.status,
        error_message=session.error_message,
        attachments=[_attachment_to_schema(item) for item in attachments],
        deviation_analysis_payload=_result_payload_to_schema(
            "deviation",
            "deviation_analysis",
            session.deviation_analysis_payload,
        ),
        capa_suggestion_payload=_result_payload_to_schema(
            "deviation",
            "capa_suggestion",
            session.capa_suggestion_payload,
        ),
        last_generated_at=session.last_generated_at,
        created_at=session.created_at,
        updated_at=session.updated_at,
    )


async def _get_deviation_or_raise(
    db: AsyncSession,
    deviation_id: uuid.UUID,
) -> Deviation:
    deviation = await db.get(Deviation, deviation_id)
    if not deviation or deviation.is_deleted:
        raise ValueError("偏差不存在")
    return deviation


async def _get_or_create_deviation_ai_session_record(
    db: AsyncSession,
    deviation_id: uuid.UUID,
    *,
    user_id: str = "system",
) -> DeviationAiSession:
    await _get_deviation_or_raise(db, deviation_id)
    result = await db.execute(
        select(DeviationAiSession).where(
            DeviationAiSession.deviation_id == deviation_id,
            DeviationAiSession.is_deleted == False,
        )
    )
    session = result.scalar_one_or_none()
    if session:
        return session

    session = DeviationAiSession(
        deviation_id=deviation_id,
        supplement_text="",
        attachment_summary="",
        status="idle",
        created_by=uuid.UUID(user_id) if user_id != "system" else None,
        updated_by=uuid.UUID(user_id) if user_id != "system" else None,
    )
    db.add(session)
    await db.flush()
    return session


async def _get_deviation_ai_session_by_id(
    db: AsyncSession,
    session_id: uuid.UUID,
) -> DeviationAiSession | None:
    result = await db.execute(
        select(DeviationAiSession)
        .where(
            DeviationAiSession.id == session_id,
            DeviationAiSession.is_deleted == False,
        )
        .execution_options(populate_existing=True)
    )
    return result.scalar_one_or_none()


def _extract_docx_text(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_values:
                parts.append(" | ".join(row_values))
    return "\n".join(parts).strip()


def _extract_legacy_doc_text(content: bytes) -> str:
    for encoding in ("utf-8", "gbk", "gb18030", "latin1"):
        try:
            decoded = content.decode(encoding)
            cleaned = decoded.replace("\x00", "").strip()
            if cleaned:
                return cleaned
        except UnicodeDecodeError:
            continue
    return ""


def _extract_excel_text(content: bytes) -> str:
    workbook = load_workbook(io.BytesIO(content), data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"[{sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value not in (None, "")]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines).strip()


async def _extract_image_text(content: bytes, content_type: str) -> tuple[str, str | None]:
    try:
        image = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
        if text:
            return text, None
    except Exception as exc:  # noqa: BLE001
        ocr_error = str(exc)
    else:
        ocr_error = None

    try:
        encoded = base64.b64encode(content).decode("ascii")
        image_url = f"data:{content_type};base64,{encoded}"
        vision_result = await llm_client.chat_vision_json(
            "请提取图片中的可读文字，并返回 JSON：{\"text\": \"...\", \"summary\": \"...\"}",
            image_urls=[image_url],
            expected_keys=["text", "summary"],
        )
        text = str(vision_result.get("text") or "").strip()
        summary = str(vision_result.get("summary") or "").strip() or None
        return text, summary
    except Exception as exc:  # noqa: BLE001
        if ocr_error:
            raise RuntimeError(f"{ocr_error}; {exc}") from exc
        raise RuntimeError(str(exc)) from exc


async def _parse_attachment_content(
    file_name: str,
    content: bytes,
    content_type: str,
) -> tuple[str | None, str | None]:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".docx":
        parsed_text = _extract_docx_text(content)
        return parsed_text or None, _truncate_text(parsed_text, 600) or None
    if suffix == ".doc":
        parsed_text = _extract_legacy_doc_text(content)
        return parsed_text or None, _truncate_text(parsed_text, 600) or None
    if suffix in {".xls", ".xlsx"}:
        parsed_text = _extract_excel_text(content)
        return parsed_text or None, _truncate_text(parsed_text, 600) or None
    if suffix in {".png", ".jpg", ".jpeg"}:
        parsed_text, summary = await _extract_image_text(content, content_type)
        return parsed_text or None, summary or _truncate_text(parsed_text, 600) or None
    raise ValueError("仅支持 Word、Excel、图片附件")


async def _read_upload_file_with_limit(file: UploadFile) -> bytes:
    max_bytes = get_settings().MAX_UPLOAD_SIZE_MB * 1024 * 1024
    chunks: list[bytes] = []
    total_size = 0
    while True:
        chunk = await file.read(1024 * 1024)
        if not chunk:
            break
        total_size += len(chunk)
        if total_size > max_bytes:
            raise ValueError(f"附件不能超过 {get_settings().MAX_UPLOAD_SIZE_MB}MB")
        chunks.append(chunk)
    return b"".join(chunks)


def _store_deviation_ai_attachment(
    file_name: str,
    content: bytes,
    content_type: str,
) -> str:
    safe_name = f"{uuid.uuid4()}_{Path(file_name).name}"
    object_key = f"deviation-ai/{safe_name}"
    if minio_enabled():
        upload_object(
            module="quality",
            object_key=object_key,
            data=content,
            length=len(content),
            content_type=content_type,
        )
        return object_key

    file_path = os.path.normpath(str(_quality_ai_local_upload_dir() / safe_name))
    allowed_root = os.path.normpath(str(_quality_ai_local_upload_dir()))
    if not file_path.startswith(allowed_root):
        raise ValueError("非法文件路径")
    with open(file_path, "wb") as file_obj:
        file_obj.write(content)
    return file_path


def _delete_deviation_ai_attachment_file(storage_path: str) -> None:
    if minio_enabled():
        try:
            delete_object("quality", storage_path)
        except Exception:  # noqa: BLE001
            return
        return
    if os.path.exists(storage_path):
        os.remove(storage_path)


async def get_or_create_deviation_ai_session(
    db: AsyncSession,
    deviation_id: uuid.UUID,
) -> DeviationAiSessionOut:
    session = await _get_or_create_deviation_ai_session_record(db, deviation_id)
    session_id = session.id
    await db.commit()
    session = await _get_deviation_ai_session_by_id(db, session_id)
    if not session:
        raise ValueError("AI 会话不存在")
    return await _session_to_schema(db, session)


async def update_deviation_ai_session(
    db: AsyncSession,
    deviation_id: uuid.UUID,
    supplement_text: str,
    user_id: str,
) -> DeviationAiSessionOut:
    session = await _get_or_create_deviation_ai_session_record(
        db, deviation_id, user_id=user_id
    )
    session.supplement_text = supplement_text.strip()
    session.updated_by = uuid.UUID(user_id) if user_id != "system" else None
    session_id = session.id
    await db.commit()
    session = await _get_deviation_ai_session_by_id(db, session_id)
    if not session:
        raise ValueError("AI 会话不存在")
    return await _session_to_schema(db, session)


async def upload_deviation_ai_session_attachment(
    db: AsyncSession,
    deviation_id: uuid.UUID,
    file: UploadFile,
    user_id: str,
) -> DeviationAiSessionAttachmentOut:
    if not file.filename:
        raise ValueError("附件文件名不能为空")
    content = await _read_upload_file_with_limit(file)
    session = await _get_or_create_deviation_ai_session_record(
        db, deviation_id, user_id=user_id
    )
    existing_attachments = await _list_deviation_ai_session_attachments(db, session.id)
    stored_path = _store_deviation_ai_attachment(
        file.filename,
        content,
        file.content_type or "application/octet-stream",
    )
    try:
        parsed_text, parsed_summary = await _parse_attachment_content(
            file.filename,
            content,
            file.content_type or "application/octet-stream",
        )
        parse_status = "completed"
        parse_error = None
    except Exception as exc:  # noqa: BLE001
        parsed_text = None
        parsed_summary = None
        parse_status = "failed"
        parse_error = str(exc)

    attachment = DeviationAiSessionAttachment(
        session_id=session.id,
        file_name=file.filename,
        file_type=file.content_type or "application/octet-stream",
        file_size=len(content),
        storage_path=stored_path,
        parsed_text=parsed_text,
        parsed_summary=parsed_summary,
        parse_status=parse_status,
        parse_error=parse_error,
        sort_order=len(existing_attachments),
        created_by=uuid.UUID(user_id) if user_id != "system" else None,
        updated_by=uuid.UUID(user_id) if user_id != "system" else None,
    )
    db.add(attachment)
    await db.flush()
    attachment_id = attachment.id
    current_attachments = existing_attachments + [attachment]
    session.attachment_summary = _build_attachment_summary_from_rows(current_attachments)
    session.updated_by = uuid.UUID(user_id) if user_id != "system" else None
    await db.commit()
    attachment = await db.get(DeviationAiSessionAttachment, attachment_id)
    if not attachment:
        raise ValueError("附件不存在")
    return _attachment_to_schema(attachment)


async def delete_deviation_ai_session_attachment(
    db: AsyncSession,
    deviation_id: uuid.UUID,
    attachment_id: uuid.UUID,
    user_id: str,
) -> DeviationAiSessionOut:
    session = await _get_or_create_deviation_ai_session_record(
        db, deviation_id, user_id=user_id
    )
    attachment = await db.get(DeviationAiSessionAttachment, attachment_id)
    if (
        not attachment
        or attachment.is_deleted
        or attachment.session_id != session.id
    ):
        raise ValueError("附件不存在")

    _delete_deviation_ai_attachment_file(attachment.storage_path)
    attachment.is_deleted = True
    attachment.updated_by = uuid.UUID(user_id) if user_id != "system" else None
    remaining_attachments = [
        item
        for item in await _list_deviation_ai_session_attachments(db, session.id)
        if item.id != attachment.id
    ]
    session.attachment_summary = _build_attachment_summary_from_rows(remaining_attachments)
    session.updated_by = uuid.UUID(user_id) if user_id != "system" else None
    session_id = session.id
    await db.commit()
    session = await _get_deviation_ai_session_by_id(db, session_id)
    if not session:
        raise ValueError("AI 会话不存在")
    return await _session_to_schema(db, session)


async def regenerate_deviation_ai_session(
    db: AsyncSession,
    deviation_id: uuid.UUID,
    user_id: str,
) -> DeviationAiSessionOut:
    deviation = await _get_deviation_or_raise(db, deviation_id)
    session = await _get_or_create_deviation_ai_session_record(
        db, deviation_id, user_id=user_id
    )
    snapshot = _build_deviation_snapshot(deviation)
    prompt = _deviation_conversation_prompt(
        snapshot,
        session.supplement_text,
        session.attachment_summary,
    )
    config = await get_config("text")
    session.status = "processing"
    session.error_message = None
    session.updated_by = uuid.UUID(user_id) if user_id != "system" else None
    await db.flush()

    try:
        raw = await llm_client.chat_json(
            [{"role": "user", "content": prompt}],
            expected_keys=DEVIATION_AI_SESSION_RESPONSE_KEYS,
            temperature=0.2,
        )
        session.deviation_analysis_payload = _normalize_result(
            raw.get("deviation_analysis") or {}
        )
        session.capa_suggestion_payload = _normalize_result(
            raw.get("capa_suggestion") or {}
        )
        session.model_name = config.model_name
        session.status = "completed"
        session.error_message = None
        session.last_generated_at = datetime.now(timezone.utc)
        session_id = session.id
        await db.commit()
    except Exception as exc:  # noqa: BLE001
        session.model_name = config.model_name
        session.status = "failed"
        session.error_message = str(exc)
        await db.commit()
        raise RuntimeError(str(exc)) from exc

    session = await _get_deviation_ai_session_by_id(db, session_id)
    if not session:
        raise ValueError("AI 会话不存在")
    return await _session_to_schema(db, session)


async def apply_deviation_ai_session(
    db: AsyncSession,
    deviation_id: uuid.UUID,
    section: str,
    field_keys: list[str],
    user_id: str,
) -> DeviationAiSessionOut:
    session = await _get_or_create_deviation_ai_session_record(
        db, deviation_id, user_id=user_id
    )
    deviation = await _get_deviation_or_raise(db, deviation_id)
    if section not in {"deviation_analysis", "capa_suggestion"}:
        raise ValueError("section 不合法")
    if not field_keys:
        raise ValueError("请选择要应用的字段")

    payload = (
        session.deviation_analysis_payload
        if section == "deviation_analysis"
        else session.capa_suggestion_payload
    )
    if not payload:
        raise ValueError("当前会话暂无可应用结果")

    analysis_type = "deviation_analysis" if section == "deviation_analysis" else "capa_suggestion"
    allowed = {
        field.field_key: field
        for field in _build_applicable_fields("deviation", analysis_type, payload)
    }
    invalid = [field for field in field_keys if field not in allowed]
    if invalid:
        raise ValueError(f"存在不允许应用的字段: {', '.join(invalid)}")

    structured_fields = (payload or {}).get("structured_fields") or {}
    if "root_cause_analysis" in field_keys:
        deviation.root_cause_analysis = structured_fields.get("preliminary_cause_analysis")
    if "corrective_actions" in field_keys:
        deviation.corrective_actions = structured_fields.get("capa_suggestions")
    if section == "deviation_analysis":
        deviation.ai_analysis = payload
    deviation.updated_by = uuid.UUID(user_id) if user_id != "system" else None
    session.updated_by = uuid.UUID(user_id) if user_id != "system" else None
    session_id = session.id
    await db.commit()
    session = await _get_deviation_ai_session_by_id(db, session_id)
    if not session:
        raise ValueError("AI 会话不存在")
    return await _session_to_schema(db, session)


async def analyze_deviation_async(deviation_id: uuid.UUID, user_id: str) -> None:
    async with async_session_factory() as db:
        await analyze_deviation_record(
            db, deviation_id, user_id, transition_status=True
        )
