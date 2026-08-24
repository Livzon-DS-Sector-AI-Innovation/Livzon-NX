"""Quality module import/export service for CAPA and Deviation (Word docx format)."""

import logging
import re
from copy import deepcopy
from datetime import date, datetime, timedelta
from io import BytesIO
from pathlib import Path
from typing import Any

from docx.oxml.ns import qn
from sqlalchemy.ext.asyncio import AsyncSession

from app.modules.quality import repository as repo
from app.modules.quality.models import Deviation

logger = logging.getLogger(__name__)


# 偏差登记表导出模板（与桌面《2026年偏差登记表.docx》一致），
# 含标题/日期区/签名区、文档保护（readOnly）与 permStart/permEnd 黄色可编辑区域
TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent
    / "templates"
    / "deviations_register_template.docx"
)


def _parse_date(value: Any) -> datetime | None:
    """Parse date from various formats."""
    if not value:
        return None
    if isinstance(value, datetime):
        return value
    s = str(value).strip().replace(".", "-").replace("/", "-")
    for fmt in ("%Y-%m-%d", "%Y%m%d", "%Y-%m-%d %H:%M:%S"):
        try:
            return datetime.strptime(s, fmt)
        except (ValueError, TypeError):
            continue
    return None


def _parse_bool(value: Any) -> bool | None:
    """Parse boolean from various formats."""
    if value is None or value == "":
        return None
    s = str(value).strip().lower()
    if s in ("是", "yes", "true", "1", "y"):
        return True
    if s in ("否", "no", "false", "0", "n"):
        return False
    return None


def _parse_occurred_text(text: Any) -> tuple[bool | None, str | None]:
    """解析"偏差是否曾发生"列文本 → (是否曾发生, 曾发生编号)。

    对齐桌面模板勾选格式：选中项不带□、未选中项带□。
    未发生：'□是 编号：\\n否'；曾发生：'是 编号：PC-xxx\\n□否'。
    支持多个曾发生编号（换行分隔），如 '是 编号：PC-2502001\\nPC-2502003\\n□否'。
    """
    if not text:
        return None, None
    t = str(text).strip()
    # 曾发生：存在"是"且"是"前无□（即"□是"不存在）
    if "是" in t and "□是" not in t:
        code = None
        # 提取"编号："到"□否"/"否"之间的全部内容（支持换行的多个编号）
        m = re.search(r"编号[:：]\s*(.*?)(?:\n*□否|\n*否\s*$|$)", t, re.DOTALL)
        if m:
            code_text = m.group(1).strip()
            if code_text:
                # 提取所有 PC-数字 格式的编号
                codes = re.findall(r"PC-\d+", code_text)
                if codes:
                    code = "\n".join(codes)
                else:
                    code = code_text
        return True, code
    if "否" in t:
        return False, None
    return None, None


def _clean_text(value: Any) -> str:
    """Clean text value."""
    if not value:
        return ""
    return str(value).strip()


def _parse_date_value(value: Any) -> date | None:
    parsed = _parse_date(value)
    if parsed is None:
        return None
    return parsed.date()


def _is_empty_row(row_data: dict[str, str]) -> bool:
    return not any(_clean_text(value) for value in row_data.values())


CHANGE_HEADERS = [
    "序号",
    "变更控制号",
    "变更申请部门",
    "变更对象",
    "变更内容",
    "变更等级",
    "变更申请日期",
    "变更计划批准日期",
    "变更正式执行日期",
    "变更关闭日期",
]


def _change_row_to_data(
    row_data: dict[str, str], change_type: str = "technical"
) -> dict[str, Any]:
    return {
        "change_type": change_type,
        "serial_number": row_data.get("序号") or None,
        "change_code": row_data.get("变更控制号", ""),
        "applicant_department": row_data.get("变更申请部门") or None,
        "change_object": row_data.get("变更对象") or None,
        "change_content": row_data.get("变更内容") or None,
        "change_level": row_data.get("变更等级") or None,
        "application_date": _parse_date_value(row_data.get("变更申请日期", "")),
        "planned_approval_date": _parse_date_value(
            row_data.get("变更计划批准日期", "")
        ),
        "execution_date": _parse_date_value(row_data.get("变更正式执行日期", "")),
        "closure_date": _parse_date_value(row_data.get("变更关闭日期", "")),
    }


async def preview_change_import(
    db: AsyncSession,
    file_content: bytes,
    change_type: str = "technical",
) -> dict[str, Any]:
    """Preview change control import from Word docx."""
    import docx

    doc = docx.Document(BytesIO(file_content))
    if not doc.tables:
        return {
            "headers": [],
            "valid_rows": 0,
            "error_rows": [{"row_number": 0, "error_message": "文档中未找到表格"}],
            "total_rows": 0,
        }

    table = doc.tables[0]
    headers = [_clean_text(cell.text) for cell in table.rows[0].cells]

    valid_rows = 0
    error_rows = []

    for row_idx, row in enumerate(table.rows[1:], start=2):
        row_data = {
            _clean_text(headers[i]): _clean_text(row.cells[i].text)
            for i in range(len(headers))
        }
        if _is_empty_row(row_data):
            continue

        change_code = row_data.get("变更控制号", "")
        errors = []
        if not change_code:
            errors.append("变更控制号不能为空")
        else:
            exists = await repo.exists_by_change_code(db, change_code)
            if exists:
                errors.append(f"变更控制号已存在: {change_code}")

        if errors:
            error_rows.append(
                {
                    "row_number": row_idx,
                    "error_message": "; ".join(errors),
                    "row_data": row_data,
                }
            )
        else:
            valid_rows += 1

    return {
        "headers": headers,
        "valid_rows": valid_rows,
        "error_rows": error_rows,
        "total_rows": valid_rows + len(error_rows),
    }


async def confirm_change_import(
    db: AsyncSession,
    file_content: bytes,
    skip_duplicates: bool = True,
    update_existing: bool = False,
    change_type: str = "technical",
) -> dict[str, Any]:
    """Confirm change control import from Word docx."""
    import docx

    doc = docx.Document(BytesIO(file_content))
    if not doc.tables:
        return {
            "success_count": 0,
            "update_count": 0,
            "skip_count": 0,
            "error_count": 1,
            "error_details": [{"row": 0, "error": "文档中未找到表格"}],
        }

    table = doc.tables[0]
    headers = [_clean_text(cell.text) for cell in table.rows[0].cells]

    success_count = 0
    update_count = 0
    skip_count = 0
    error_count = 0
    error_details = []

    for row_idx, row in enumerate(table.rows[1:], start=2):
        row_data = {
            _clean_text(headers[i]): _clean_text(row.cells[i].text)
            for i in range(len(headers))
        }
        if _is_empty_row(row_data):
            continue

        try:
            change_code = row_data.get("变更控制号", "")
            if not change_code:
                error_count += 1
                error_details.append({"row": row_idx, "error": "变更控制号为空"})
                continue

            # 查找记录（包含软删除记录，避免唯一约束冲突）
            existing = await repo.get_change_by_code_include_deleted(db, change_code)
            data = _change_row_to_data(row_data, change_type)
            if existing:
                if existing.is_deleted:
                    # 软删除记录：恢复并更新（重新启用该记录，避免唯一约束冲突）
                    existing.is_deleted = False
                    existing.deleted_by = None
                    existing.deleted_at = None
                    existing.change_type = change_type
                    await repo.update_change(db, existing, data)
                    update_count += 1
                elif update_existing:
                    await repo.update_change(db, existing, data)
                    update_count += 1
                elif skip_duplicates:
                    skip_count += 1
                else:
                    error_count += 1
                    error_details.append(
                        {"row": row_idx, "error": f"变更控制号已存在: {change_code}"}
                    )
                continue

            await repo.create_change(db, data)
            success_count += 1
        except Exception as e:
            error_count += 1
            error_details.append({"row": row_idx, "error": str(e)})

    await db.commit()

    return {
        "success_count": success_count,
        "update_count": update_count,
        "skip_count": skip_count,
        "error_count": error_count,
        "error_details": error_details,
    }


async def export_changes(
    db: AsyncSession,
    template_content: bytes | None = None,
    change_code: str | None = None,
    applicant_department: str | None = None,
    change_object: str | None = None,
    change_level: str | None = None,
    application_date_from: str | None = None,
    application_date_to: str | None = None,
    planned_approval_date_from: str | None = None,
    planned_approval_date_to: str | None = None,
    execution_date_from: str | None = None,
    execution_date_to: str | None = None,
    closure_date_from: str | None = None,
    closure_date_to: str | None = None,
    content_keyword: str | None = None,
) -> bytes:
    """Export change control data to Word docx."""
    import docx

    changes, _ = await repo.get_changes(
        db,
        change_code=change_code,
        applicant_department=applicant_department,
        change_object=change_object,
        change_level=change_level,
        application_date_from=_parse_date_value(application_date_from),
        application_date_to=_parse_date_value(application_date_to),
        planned_approval_date_from=_parse_date_value(planned_approval_date_from),
        planned_approval_date_to=_parse_date_value(planned_approval_date_to),
        execution_date_from=_parse_date_value(execution_date_from),
        execution_date_to=_parse_date_value(execution_date_to),
        closure_date_from=_parse_date_value(closure_date_from),
        closure_date_to=_parse_date_value(closure_date_to),
        content_keyword=content_keyword,
        page=1,
        page_size=10000,
    )

    if template_content:
        doc = docx.Document(BytesIO(template_content))
        table = doc.tables[0]
        for i in range(len(table.rows) - 1, 0, -1):
            tr = table.rows[i]._tr
            table._tbl.remove(tr)
    else:
        doc = docx.Document()
        doc.add_heading("变更管理台账", level=1)
        table = doc.add_table(rows=1, cols=len(CHANGE_HEADERS))
        table.style = "Table Grid"
        for i, header in enumerate(CHANGE_HEADERS):
            table.rows[0].cells[i].text = header

    for idx, change in enumerate(changes, start=1):
        row = table.add_row()  # type: ignore[no-untyped-call]
        row.cells[0].text = change.serial_number or str(idx)
        row.cells[1].text = change.change_code or ""
        row.cells[2].text = change.applicant_department or ""
        row.cells[3].text = change.change_object or ""
        row.cells[4].text = change.change_content or ""
        row.cells[5].text = change.change_level or ""
        row.cells[6].text = (
            change.application_date.strftime("%Y.%m.%d")
            if change.application_date
            else ""
        )
        row.cells[7].text = (
            change.planned_approval_date.strftime("%Y.%m.%d")
            if change.planned_approval_date
            else ""
        )
        row.cells[8].text = (
            change.execution_date.strftime("%Y.%m.%d") if change.execution_date else ""
        )
        row.cells[9].text = (
            change.closure_date.strftime("%Y.%m.%d") if change.closure_date else ""
        )

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


# ============ CAPA Import/Export (Word docx) ============


async def preview_capa_import(
    db: AsyncSession,
    file_content: bytes,
) -> dict[str, Any]:
    """Preview CAPA import from Word docx."""
    import docx

    doc = docx.Document(BytesIO(file_content))
    if not doc.tables:
        return {
            "valid_rows": 0,
            "error_rows": [{"row_number": 0, "error_message": "文档中未找到表格"}],
            "total_rows": 0,
        }

    table = doc.tables[0]
    headers = [_clean_text(cell.text) for cell in table.rows[0].cells]

    valid_rows = 0
    error_rows = []

    for row_idx, row in enumerate(table.rows[1:], start=2):
        row_data = {
            _clean_text(headers[i]): _clean_text(row.cells[i].text)
            for i in range(len(headers))
        }
        errors = []

        capa_code = row_data.get("CAPA编号", "")
        if not capa_code:
            errors.append("CAPA编号不能为空")
        else:
            exists = await repo.exists_by_capa_code(db, capa_code)
            if exists:
                errors.append(f"CAPA编号已存在: {capa_code}")

        if errors:
            error_rows.append(
                {
                    "row_number": row_idx,
                    "error_message": "; ".join(errors),
                    "row_data": row_data,
                }
            )
        else:
            valid_rows += 1

    return {
        "valid_rows": valid_rows,
        "error_rows": error_rows,
        "total_rows": valid_rows + len(error_rows),
    }


async def confirm_capa_import(
    db: AsyncSession,
    file_content: bytes,
    skip_duplicates: bool = True,
    update_existing: bool = False,
) -> dict[str, Any]:
    """Confirm CAPA import from Word docx with deduplication."""
    import re

    import docx

    doc = docx.Document(BytesIO(file_content))
    if not doc.tables:
        return {
            "success_count": 0,
            "error_count": 1,
            "error_details": [{"row": 0, "error": "文档中未找到表格"}],
        }

    table = doc.tables[0]
    headers = [_clean_text(cell.text) for cell in table.rows[0].cells]

    success_count = 0
    skip_count = 0
    update_count = 0
    error_count = 0
    error_details = []

    for row_idx, row in enumerate(table.rows[1:], start=2):
        row_data = {
            _clean_text(headers[i]): _clean_text(row.cells[i].text)
            for i in range(len(headers))
        }
        try:
            capa_code = row_data.get("CAPA编号", "")
            if not capa_code:
                error_count += 1
                error_details.append({"row": row_idx, "error": "CAPA编号为空"})
                continue

            existing = await repo.get_capa_by_code(db, capa_code)
            if existing:
                if update_existing:
                    # Update existing record
                    qa_info = row_data.get("QA质量员/日期", "")
                    qa_confirmer = None
                    qa_confirm_date = None
                    if qa_info:
                        match = re.match(
                            r"([^\d]+)(\d{4}[\.\-/]\d{2}[\.\-/]\d{2})", qa_info
                        )
                        if match:
                            qa_confirmer = match.group(1).strip()
                            qa_confirm_date = _parse_date(match.group(2))
                        else:
                            qa_confirmer = qa_info

                    update_data = {
                        "title": row_data.get("CAPA简述", ""),
                        "source_code": row_data.get("来源编号", ""),
                        "department": row_data.get("事件部门", ""),
                        "affected_product": row_data.get("涉及产品", ""),
                        "evaluation_result": row_data.get("CAPA效果评估", ""),
                        "closure_date": _parse_date(row_data.get("关闭日期", "")),
                        "qa_confirmer": qa_confirmer,
                        "qa_confirm_date": qa_confirm_date,
                    }
                    await repo.update_capa(db, existing, update_data)
                    update_count += 1
                elif skip_duplicates:
                    skip_count += 1
                else:
                    error_count += 1
                    error_details.append(
                        {"row": row_idx, "error": f"CAPA编号已存在: {capa_code}"}
                    )
                continue

            # Parse QA质量员/日期 (format: "杨小芹2026.03.11")
            qa_info = row_data.get("QA质量员/日期", "")
            qa_confirmer = None
            qa_confirm_date = None
            if qa_info:
                match = re.match(r"([^\d]+)(\d{4}[\.\-/]\d{2}[\.\-/]\d{2})", qa_info)
                if match:
                    qa_confirmer = match.group(1).strip()
                    qa_confirm_date = _parse_date(match.group(2))
                else:
                    qa_confirmer = qa_info

            data = {
                "capa_code": capa_code,
                "title": row_data.get("CAPA简述", ""),
                "source_code": row_data.get("来源编号", ""),
                "department": row_data.get("事件部门", ""),
                "affected_product": row_data.get("涉及产品", ""),
                "evaluation_result": row_data.get("CAPA效果评估", ""),
                "closure_date": _parse_date(row_data.get("关闭日期", "")),
                "qa_confirmer": qa_confirmer,
                "qa_confirm_date": qa_confirm_date,
                "status": "draft",
            }

            await repo.create_capa(db, data)
            success_count += 1
        except Exception as e:
            error_count += 1
            error_details.append({"row": row_idx, "error": str(e)})

    await db.commit()

    return {
        "success_count": success_count,
        "update_count": update_count,
        "skip_count": skip_count,
        "error_count": error_count,
        "error_details": error_details,
    }


def export_capas_template() -> bytes:
    """生成 CAPA 导入模板（仅含表头）"""
    from io import BytesIO

    import docx
    from docx.shared import Pt

    doc = docx.Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)

    doc.add_heading("CAPA 导入模板", level=1)
    doc.add_paragraph("请按照以下表头格式填充数据，保留表头行。")
    doc.add_paragraph("注意事项：编号、日期、部门等字段请勿留空。")

    headers = [
        "CAPA编号",
        "启动日期",
        "事件部门",
        "涉及产品",
        "来源编号",
        "CAPA简述",
        "CAPA效果评估",
        "关闭日期",
        "QA质量员/日期",
    ]
    table = doc.add_table(rows=2, cols=len(headers), style="Table Grid")
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    # 示例行
    sample = [
        "C20260001",
        "2026-01-15",
        "质检部",
        "样品A",
        "SRC001",
        "改进流程",
        "有效",
        "2026-03-01",
        "张三/2026-01-15",
    ]
    for i, v in enumerate(sample):
        table.rows[1].cells[i].text = v

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def export_capas(
    db: AsyncSession,
    template_content: bytes | None = None,
    status: str | None = None,
    source: str | None = None,
    category: str | None = None,
    keyword: str | None = None,
    capa_code: str | None = None,
    affected_product: str | None = None,
    source_code: str | None = None,
    evaluation_result: str | None = None,
    closure_date_from: str | None = None,
    closure_date_to: str | None = None,
    department: str | None = None,
    qa_confirmer: str | None = None,
) -> bytes:
    """Export CAPA data to Word docx, preserving template format."""
    import docx

    # Get data
    closure_start = (
        datetime.fromisoformat(closure_date_from) if closure_date_from else None
    )
    closure_end = (
        datetime.fromisoformat(closure_date_to) + timedelta(days=1)
        if closure_date_to
        else None
    )
    capas, _ = await repo.get_capas(
        db,
        status,
        source,
        category,
        None,
        keyword,
        capa_code,
        affected_product,
        source_code,
        evaluation_result,
        closure_start,
        closure_end,
        department,
        qa_confirmer,
        1,
        10000,
    )

    if template_content:
        # Use template
        doc = docx.Document(BytesIO(template_content))
        table = doc.tables[0]

        # Clear existing data rows (keep header)
        for i in range(len(table.rows) - 1, 0, -1):
            tr = table.rows[i]._tr
            table._tbl.remove(tr)

        # Add data rows
        for capa in capas:
            row = table.add_row()  # type: ignore[no-untyped-call]
            row.cells[0].text = capa.capa_code
            row.cells[1].text = (
                capa.created_at.strftime("%Y.%m.%d") if capa.created_at else ""
            )
            row.cells[2].text = capa.department or ""
            row.cells[3].text = capa.affected_product or ""
            row.cells[4].text = capa.source_code or ""
            row.cells[5].text = capa.title or ""
            row.cells[6].text = capa.evaluation_result or ""
            row.cells[7].text = (
                capa.closure_date.strftime("%Y.%m.%d") if capa.closure_date else ""
            )
            qa_info = ""
            if capa.qa_confirmer:
                qa_info = capa.qa_confirmer
                if capa.qa_confirm_date:
                    qa_info += capa.qa_confirm_date.strftime("%Y.%m.%d")
            row.cells[8].text = qa_info
    else:
        # Create new document
        doc = docx.Document()
        doc.add_heading("CAPA登记汇总表", level=1)
        table = doc.add_table(rows=1, cols=9)
        table.style = "Table Grid"

        # Header
        headers = [
            "CAPA编号",
            "启动日期",
            "事件部门",
            "涉及产品",
            "来源编号",
            "CAPA简述",
            "CAPA效果评估",
            "关闭日期",
            "QA质量员/日期",
        ]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        # Data
        for capa in capas:
            row = table.add_row()
            row.cells[0].text = capa.capa_code
            row.cells[1].text = (
                capa.created_at.strftime("%Y.%m.%d") if capa.created_at else ""
            )
            row.cells[2].text = capa.department or ""
            row.cells[3].text = capa.affected_product or ""
            row.cells[4].text = capa.source_code or ""
            row.cells[5].text = capa.title or ""
            row.cells[6].text = capa.evaluation_result or ""
            row.cells[7].text = (
                capa.closure_date.strftime("%Y.%m.%d") if capa.closure_date else ""
            )
            qa_info = ""
            if capa.qa_confirmer:
                qa_info = capa.qa_confirmer
                if capa.qa_confirm_date:
                    qa_info += capa.qa_confirm_date.strftime("%Y.%m.%d")
            row.cells[8].text = qa_info

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


# ============ Deviation Import/Export (Word docx) ============


async def preview_deviation_import(
    db: AsyncSession,
    file_content: bytes,
) -> dict[str, Any]:
    """Preview Deviation import from Word docx."""
    import docx

    doc = docx.Document(BytesIO(file_content))
    if not doc.tables:
        return {
            "valid_rows": 0,
            "error_rows": [{"row_number": 0, "error_message": "文档中未找到表格"}],
            "total_rows": 0,
        }

    table = doc.tables[0]
    headers = [_clean_text(cell.text) for cell in table.rows[0].cells]

    valid_rows = 0
    error_rows = []

    for row_idx, row in enumerate(table.rows[1:], start=2):
        row_data = {
            _clean_text(headers[i]): _clean_text(row.cells[i].text)
            for i in range(len(headers))
        }
        errors = []

        deviation_code = row_data.get("偏差编号", "")
        if not deviation_code:
            errors.append("偏差编号不能为空")
        else:
            exists = await repo.exists_by_deviation_code(db, deviation_code)
            if exists:
                errors.append(f"偏差编号已存在: {deviation_code}")

        if errors:
            error_rows.append(
                {
                    "row_number": row_idx,
                    "error_message": "; ".join(errors),
                    "row_data": row_data,
                }
            )
        else:
            valid_rows += 1

    return {
        "valid_rows": valid_rows,
        "error_rows": error_rows,
        "total_rows": valid_rows + len(error_rows),
    }


async def confirm_deviation_import(
    db: AsyncSession,
    file_content: bytes,
    skip_duplicates: bool = True,
    update_existing: bool = False,
) -> dict[str, Any]:
    """Confirm Deviation import from Word docx with deduplication."""
    import docx

    doc = docx.Document(BytesIO(file_content))
    if not doc.tables:
        return {
            "success_count": 0,
            "error_count": 1,
            "error_details": [{"row": 0, "error": "文档中未找到表格"}],
        }

    table = doc.tables[0]
    headers = [_clean_text(cell.text) for cell in table.rows[0].cells]

    success_count = 0
    skip_count = 0
    update_count = 0
    error_count = 0
    error_details = []

    for row_idx, row in enumerate(table.rows[1:], start=2):
        row_data = {
            _clean_text(headers[i]): _clean_text(row.cells[i].text)
            for i in range(len(headers))
        }
        try:
            deviation_code = row_data.get("偏差编号", "")
            if not deviation_code:
                error_count += 1
                error_details.append({"row": row_idx, "error": "偏差编号为空"})
                continue

            existing = await repo.get_deviation_by_code_include_deleted(
                db, deviation_code
            )
            if existing:
                if existing.is_deleted:
                    # 软删除记录：恢复并更新（重新启用该记录，避免唯一约束冲突）
                    existing.is_deleted = False
                    existing.deleted_by = None
                    existing.deleted_at = None
                    product_batch = row_data.get("产品名称/批号", "")
                    parts = product_batch.split("\n") if product_batch else []
                    affected_items = parts[0].strip() if parts else None
                    batch_number = parts[1].strip() if len(parts) > 1 else None
                    has_occurred_before, previous_occurrence_code = (
                        _parse_occurred_text(row_data.get("偏差是否曾发生", ""))
                    )
                    level_text = row_data.get("偏差等级", "")
                    level = {
                        "次要偏差": "minor",
                        "中等偏差": "moderate",
                        "严重偏差": "major",
                    }.get(level_text, level_text.lower() if level_text else None)
                    update_data = {
                        "title": row_data.get("偏差简要描述", "")[:100]
                        or deviation_code,
                        "department": row_data.get("事件部门", ""),
                        "description": row_data.get("偏差简要描述", ""),
                        "batch_number": batch_number,
                        "affected_items": affected_items,
                        "has_occurred_before": has_occurred_before,
                        "previous_occurrence_code": previous_occurrence_code,
                        "root_cause_analysis": row_data.get("根本原因", ""),
                        "level": level,
                        "corrective_actions": row_data.get("纠正预防措施", ""),
                        "material_disposition": row_data.get("产品/物料处理结果", ""),
                        "investigation_completed_at": _parse_date(
                            row_data.get("调查完成时间", "")
                        ),
                    }
                    await repo.update_deviation(db, existing, update_data)
                    update_count += 1
                elif update_existing:
                    # Update existing record
                    product_batch = row_data.get("产品名称/批号", "")
                    parts = product_batch.split("\n") if product_batch else []
                    affected_items = parts[0].strip() if len(parts) > 0 else None
                    batch_number = parts[1].strip() if len(parts) > 1 else None

                    has_occurred_before, previous_occurrence_code = (
                        _parse_occurred_text(row_data.get("偏差是否曾发生", ""))
                    )

                    level_text = row_data.get("偏差等级", "")
                    level_map = {
                        "次要偏差": "minor",
                        "中等偏差": "moderate",
                        "严重偏差": "major",
                    }
                    level = level_map.get(
                        level_text, level_text.lower() if level_text else None
                    )

                    update_data = {
                        "title": row_data.get("偏差简要描述", "")[:100]
                        or deviation_code,
                        "department": row_data.get("事件部门", ""),
                        "description": row_data.get("偏差简要描述", ""),
                        "batch_number": batch_number,
                        "affected_items": affected_items,
                        "has_occurred_before": has_occurred_before,
                        "previous_occurrence_code": previous_occurrence_code,
                        "root_cause_analysis": row_data.get("根本原因", ""),
                        "level": level,
                        "corrective_actions": row_data.get("纠正预防措施", ""),
                        "material_disposition": row_data.get("产品/物料处理结果", ""),
                        "investigation_completed_at": _parse_date(
                            row_data.get("调查完成时间", "")
                        ),
                    }
                    await repo.update_deviation(db, existing, update_data)
                    update_count += 1
                elif skip_duplicates:
                    skip_count += 1
                else:
                    error_count += 1
                    error_details.append(
                        {"row": row_idx, "error": f"偏差编号已存在: {deviation_code}"}
                    )
                continue

            # Parse 产品名称/批号
            product_batch = row_data.get("产品名称/批号", "")
            parts = product_batch.split("\n") if product_batch else []
            affected_items = parts[0].strip() if len(parts) > 0 else None
            batch_number = parts[1].strip() if len(parts) > 1 else None

            # Parse 偏差是否曾发生 (format: "□是 编号：\n否" / "是 编号：PC-xxx\n□否")
            has_occurred_before, previous_occurrence_code = _parse_occurred_text(
                row_data.get("偏差是否曾发生", "")
            )

            # Parse 偏差等级
            level_text = row_data.get("偏差等级", "")
            level_map = {
                "次要偏差": "minor",
                "中等偏差": "moderate",
                "严重偏差": "major",
            }
            level = level_map.get(
                level_text, level_text.lower() if level_text else None
            )

            data = {
                "deviation_code": deviation_code,
                "title": row_data.get("偏差简要描述", "")[:100] or deviation_code,
                "department": row_data.get("事件部门", ""),
                "description": row_data.get("偏差简要描述", ""),
                "batch_number": batch_number,
                "affected_items": affected_items,
                "has_occurred_before": has_occurred_before,
                "previous_occurrence_code": previous_occurrence_code,
                "root_cause_analysis": row_data.get("根本原因", ""),
                "level": level,
                "corrective_actions": row_data.get("纠正预防措施", ""),
                "material_disposition": row_data.get("产品/物料处理结果", ""),
                "investigation_completed_at": _parse_date(
                    row_data.get("调查完成时间", "")
                ),
                "status": "draft",
            }

            await repo.create_deviation(db, data)
            success_count += 1
        except Exception as e:
            error_count += 1
            error_details.append({"row": row_idx, "error": str(e)})

    await db.commit()

    return {
        "success_count": success_count,
        "update_count": update_count,
        "skip_count": skip_count,
        "error_count": error_count,
        "error_details": error_details,
    }


def _strip_auto_numbering(elem: Any) -> None:
    """移除元素内各段落的 w:numPr（Word 自动编号）。

    模板数据行部分单元格依赖自动编号；写入手动文本后若不清理，
    Word 会叠加渲染自动编号，导致导出内容重复。
    """
    for p_elem in elem.iter(qn("w:p")):
        ppr_elem = p_elem.find(qn("w:pPr"))
        if ppr_elem is not None:
            num_pr = ppr_elem.find(qn("w:numPr"))
            if num_pr is not None:
                ppr_elem.remove(num_pr)


def _clear_text_in_paragraph(paragraph_elem: Any) -> None:
    """清空段落内所有 run 的文本。"""
    for run_elem in paragraph_elem.findall(qn("w:r")):
        for text_elem in run_elem.findall(qn("w:t")):
            text_elem.text = ""


def _set_cell_text(cell_elem: Any, text: str) -> None:
    """按行写入单元格文本，保留单元格格式与可编辑区域标记。

    直接在 XML 层修改已有段落/run 的文本，不使用 python-docx 的 cell.text
    赋值（其会清空 tc 内容，导致 permStart/permEnd 等保护元素丢失）。
    """
    lines = (text or "").split("\n")
    paragraph_elems = cell_elem.findall(qn("w:p"))
    if not paragraph_elems:
        paragraph_elem = cell_elem.makeelement(qn("w:p"), {})
        cell_elem.append(paragraph_elem)
        paragraph_elems = [paragraph_elem]

    # 段落数量与行数对齐（多余删除、不足复制首段补足）
    while len(paragraph_elems) > len(lines):
        paragraph_elems[-1].getparent().remove(paragraph_elems[-1])
        paragraph_elems = cell_elem.findall(qn("w:p"))
    while len(paragraph_elems) < len(lines):
        new_paragraph = deepcopy(paragraph_elems[0])
        _clear_text_in_paragraph(new_paragraph)
        cell_elem.append(new_paragraph)
        paragraph_elems = cell_elem.findall(qn("w:p"))

    for index, line in enumerate(lines):
        paragraph_elem = paragraph_elems[index]
        run_elems = paragraph_elem.findall(qn("w:r"))
        for run_elem in run_elems:
            for text_elem in run_elem.findall(qn("w:t")):
                text_elem.text = ""

        if not run_elems:
            run_elem = paragraph_elem.makeelement(qn("w:r"), {})
            paragraph_elem.append(run_elem)
            run_elems = [run_elem]

        text_elems = run_elems[0].findall(qn("w:t"))
        if not text_elems:
            text_elem = run_elems[0].makeelement(qn("w:t"), {})
            run_elems[0].append(text_elem)
            text_elems = [text_elem]
        text_elems[0].text = line


def _deviation_row_values(
    idx: int, d: Deviation, level_map: dict[str, Any]
) -> list[str]:
    """构建偏差数据行 11 列文本值（序号/偏差编号/产品名称批号/偏差简要描述/是否曾发生/
    根本原因/偏差等级/调查完成时间/纠正预防措施/产品物料处理结果/是否关闭）。"""
    # 产品名称/批号：产品为空显示"—"，批号非空时换行附加
    product_batch = ""
    if d.affected_items and d.affected_items != "—":
        product_batch = d.affected_items
    if d.batch_number and d.batch_number != "—":
        product_batch += f"\n{d.batch_number}"
    # 偏差是否曾发生：曾发生→"是 编号：[编号]\n□否"；未发生→"□是
    # 编号：\n否"；未知→"□是 编号：\n否"
    # （对齐模板：选中项不带□，未选中项带□）
    if d.has_occurred_before is True:
        code = (d.previous_occurrence_code or "").strip()
        occurred_text = f"是 编号：{code}\n□否"
    elif d.has_occurred_before is False:
        occurred_text = "□是 编号：\n否"
    else:
        occurred_text = "□是 编号：\n否"
    # 调查完成时间
    completed_text = (
        d.investigation_completed_at.strftime("%Y.%m.%d")
        if d.investigation_completed_at
        else ""
    )
    return [
        str(idx),
        d.deviation_code,
        product_batch or "—",
        d.description or d.title or "",
        occurred_text,
        d.root_cause_analysis or "",
        level_map.get(d.level or "", d.level or ""),
        completed_text,
        d.corrective_actions or "",
        d.material_disposition or "—",
        "是" if d.status == "closed" else "否",
    ]


def _fill_template_row(
    row_elem: Any, idx: int, d: Deviation, level_map: dict[str, Any]
) -> None:
    """将偏差数据写入模板数据行（XML 级文本替换，保留单元格格式与 perm 标记）。"""
    cell_elems = row_elem.findall(qn("w:tc"))
    for ci, value in enumerate(_deviation_row_values(idx, d, level_map)):
        if ci < len(cell_elems):
            _set_cell_text(cell_elems[ci], value)


async def export_deviations(
    db: AsyncSession,
    template_content: bytes | None = None,
    status: str | None = None,
    level: str | None = None,
    department: str | None = None,
    keyword: str | None = None,
) -> bytes:
    """导出偏差数据到 Word
    docx，保留模板的标题/日期区/签名区、文档保护与黄色可编辑区域。

    模板来源：优先使用显式传入的 template_content；为 None 时读取固定模板文件
    TEMPLATE_PATH。填充数据行时深拷贝模板数据行格式（而非 add_row），避免丢失
    permStart/permEnd、单元格格式与列宽设置。
    """
    import docx

    deviations, _ = await repo.get_deviations(
        db, status, level, department, keyword, None, 1, 10000
    )
    level_map = {"minor": "次要偏差", "moderate": "中等偏差", "major": "严重偏差"}

    if template_content:
        doc = docx.Document(BytesIO(template_content))
    elif TEMPLATE_PATH.exists():
        doc = docx.Document(str(TEMPLATE_PATH))
    else:
        # 模板文件缺失时兜底：从头创建简单表格（无标题/签名/保护）
        doc = docx.Document()
        doc.add_heading("偏差登记表", level=1)
        table = doc.add_table(rows=1, cols=11)
        table.style = "Table Grid"

        headers = [
            "序号",
            "偏差编号",
            "产品名称/批号",
            "偏差简要描述",
            "偏差是否曾发生",
            "根本原因",
            "偏差等级",
            "调查完成时间",
            "纠正预防措施",
            "产品/物料处理结果",
            "是否关闭",
        ]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for idx, d in enumerate(deviations, start=1):
            row = table.add_row()
            for ci, value in enumerate(_deviation_row_values(idx, d, level_map)):
                row.cells[ci].text = value

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()

    # 模板导出：保留 1 行表头，深拷贝首行数据行作为格式模板行，删除其余数据行，
    # 再按数据条数复制格式行生成新行（数据条数超过模板行数时继续追加）。
    table = doc.tables[0]
    table_element = table._tbl
    row_elements = table_element.findall(qn("w:tr"))
    if len(row_elements) >= 2:
        template_row = deepcopy(row_elements[1])
        for row_elem in row_elements[1:]:
            table_element.remove(row_elem)

        # 清理自动编号，避免与手写文本叠加渲染
        _strip_auto_numbering(table_element)

        # 首行直接使用格式模板行（保留其 permStart 可编辑标记），其余副本移除
        # permStart，避免同一 id 重复；可编辑区域仍由首行 permStart 与表格后的
        # permEnd 界定，覆盖整个表格数据区。
        for idx, d in enumerate(deviations, start=1):
            if idx == 1:
                row_elem = template_row
            else:
                row_elem = deepcopy(template_row)
                for perm_start in row_elem.iter(qn("w:permStart")):
                    perm_start.getparent().remove(perm_start)
            table_element.append(row_elem)
            _fill_template_row(row_elem, idx, d, level_map)
    else:
        # 模板没有可复制的数据行时退回简单追加
        for idx, d in enumerate(deviations, start=1):
            row = table.add_row()  # type: ignore[no-untyped-call]
            for ci, value in enumerate(_deviation_row_values(idx, d, level_map)):
                row.cells[ci].text = value

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
