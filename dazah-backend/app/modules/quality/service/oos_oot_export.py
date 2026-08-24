"""Export OOS/OOT ledger data to docx using templates."""

from __future__ import annotations

import io
import logging
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from sqlalchemy.ext.asyncio import AsyncSession

from app.modules.quality.service.quality_feishu_pages_oos_oot import (
    list_oos_ledger_records,
    list_oot_ledger_records,
)

logger = logging.getLogger(__name__)

# Template paths
_TEMPLATE_DIR = (
    Path(__file__).resolve().parent.parent.parent.parent.parent / "templates"
)
OOS_TEMPLATE = "APP2-SMP-QA-048-02 检验结果OOS调查列表.docx"
OOT_TEMPLATE = "APP3-SMP-QA-048-02 检验结果OOT调查列表.docx"

# Column mapping: template header -> data field
LEDGER_COLUMNS = [
    ("序号", "serial_number"),
    ("日期", "date"),
    ("物料名称", "material_name"),
    ("批号", "batch_number"),
    ("调查编号", "investigation_code"),
    ("问题描述", "problem_description"),
    ("产生原因", "root_cause"),
    ("纠正预防措施", "corrective_actions"),
    ("最终处理结果", "final_disposition"),
    ("登记人", "registrant"),
    ("备注", "remark"),
]


def _format_date(value: str | None) -> str:
    """Format ISO date string to YYYY/MM/DD or YYYY-MM-DD."""
    if not value:
        return ""
    try:
        dt = datetime.fromisoformat(value.replace("Z", "+00:00"))
        return dt.strftime("%Y/%m/%d")
    except (ValueError, AttributeError):
        return str(value)[:10] if value else ""


def _fill_template(
    template_name: str, records: list[dict[str, Any]], max_rows: int = 50
) -> io.BytesIO:
    """Fill a docx template with ledger data.

    Args:
        template_name: Template filename (OOS or OOT)
        records: Data records from Feishu
        max_rows: Maximum data rows (templates have pre-numbered rows)
    """
    template_path = _TEMPLATE_DIR / template_name
    if not template_path.exists():
        # Fallback to desktop
        template_path = Path.home() / "Desktop" / template_name

    doc = Document(str(template_path))
    table = doc.tables[0]

    header_row = table.rows[0]
    data_rows = list(table.rows[1:])

    # Map column index to field name
    col_map: dict[int, str] = {}
    for i, cell in enumerate(header_row.cells):
        hdr = cell.text.strip().replace("\n", "").replace(" ", "")
        for tpl_name, field_name in LEDGER_COLUMNS:
            if tpl_name.replace(" ", "") == hdr:
                col_map[i] = field_name
                break

    # Create a style reference from the first data row
    ref_row = data_rows[0]

    # Remove all existing data rows
    for row in data_rows:
        table._tbl.remove(row._tr)

    # Ensure we have enough rows
    total_needed = max(len(records), 1)
    needed = min(total_needed, max_rows)

    for idx in range(needed):
        # Add a new row copying the style of the first data row
        new_row = table.add_row()  # type: ignore[no-untyped-call]

        record = records[idx] if idx < len(records) else {}

        for col_idx, field_name in col_map.items():
            value = record.get(field_name) or ""

            # Format date
            if field_name == "date":
                value = _format_date(str(value))
            elif field_name == "serial_number":
                value = str(idx + 1)  # Auto-number starting from 1

            # Set cell value and preserve formatting
            cell = new_row.cells[col_idx]

            # Copy paragraph formatting from reference
            ref_cell = ref_row.cells[col_idx]
            ref_paragraph = ref_cell.paragraphs[0]

            # Clear existing paragraphs
            for p in cell.paragraphs:
                p.clear()

            # Set text
            cell.paragraphs[0].text = str(value) if value else ""

            # Copy alignment from reference
            cell.paragraphs[0].alignment = ref_paragraph.alignment

            # Copy font from reference
            if ref_paragraph.runs:
                ref_run = ref_paragraph.runs[0]
                run = (
                    cell.paragraphs[0].runs[0]
                    if cell.paragraphs[0].runs
                    else cell.paragraphs[0].add_run(str(value) if value else "")
                )
                if ref_run.font.name:
                    run.font.name = ref_run.font.name
                    # Set East Asian font
                    run_properties = run._r.get_or_add_rPr()
                    run_fonts = run_properties.find(qn("w:rFonts"))
                    if run_fonts is None:
                        run_fonts = run_properties.makeelement(qn("w:rFonts"), {})
                        run_properties.insert(0, run_fonts)
                    run_fonts.set(qn("w:eastAsia"), ref_run.font.name)
                if ref_run.font.size:
                    run.font.size = ref_run.font.size

            # Copy cell background and borders by copying tcPr
            reference_cell_properties = ref_cell._tc.find(qn("w:tcPr"))
            cell_properties = cell._tc.find(qn("w:tcPr"))
            if reference_cell_properties is not None and cell_properties is not None:
                # Copy borders
                ref_borders = reference_cell_properties.findall(qn("w:tcBorders"))
                cell_borders = cell_properties.findall(qn("w:tcBorders"))
                # Remove existing borders and copy from reference
                for b in cell_borders:
                    cell_properties.remove(b)
                for ref_b in ref_borders:
                    cell_properties.append(deepcopy(ref_b))

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


async def export_oos_ledger(db: AsyncSession) -> io.BytesIO:
    """Export OOS ledger data to docx."""
    try:
        result = await list_oos_ledger_records(db, page=1, page_size=10000)
        records = result.get("items", [])
    except Exception:
        records = []
    return _fill_template(OOS_TEMPLATE, records)


async def export_oot_ledger(db: AsyncSession) -> io.BytesIO:
    """Export OOT ledger data to docx."""
    try:
        result = await list_oot_ledger_records(db, page=1, page_size=10000)
        records = result.get("items", [])
    except Exception:
        records = []
    return _fill_template(OOT_TEMPLATE, records)
