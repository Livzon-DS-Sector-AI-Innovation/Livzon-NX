"""Document catalog docx→markdown 模板化转换器。

转换规则移植自公司文件库批量脚本 `docx_to_md_converter.py`（V4，2026-08-27，
2986 个 GMP 文件实测成功率 99.6%），针对公司文件模板：
  - 首页仅保留文件名称、文件编号、生效日期；
  - 删除页眉页脚、审批签名格、分发部门表；
  - 序号（1、1.1、2.3.4）转 Markdown 标题层级；
  - 正文表格与嵌套表格（如修订简历表）保留为 Markdown 表格语法；
  - 内嵌图片提取为独立文件并在原位置引用（存储由调用方适配，md 中先以
    `img_000.png` 等占位名引用）。
与桌面脚本的差异：不落盘、按字节输入输出、图片提取失败记录日志而非静默。
"""

from __future__ import annotations

import io
import logging
import os
import re
from dataclasses import dataclass
from typing import TYPE_CHECKING, Any

from docx import Document
from docx.oxml.ns import qn

from app.core.exceptions import AppException

if TYPE_CHECKING:
    from docx.table import Table, _Cell

logger = logging.getLogger(__name__)

# 首页正文开始的关键词（页眉区跳过，直到命中这些词认为进入正文）
_CONTENT_KEYWORDS = (
    "范围",
    "Scope",
    "目的",
    "Purpose",
    "职责",
    "Responsibilities",
    "定义",
    "Definitions",
    "程序",
    "Procedures",
    "内容",
    "Content",
    "依据",
    "相关文件",
    "Basis",
    "related files",
    "适用",
    "质量指标",
)

# 审批格 / 分发表等小表格的首页关键词（命中且无序号列则整表跳过）
_HEADER_KEYWORDS = (
    "生效日期",
    "Effective Date",
    "审核间隔",
    "Review Frequency",
    "部门",
    "签名",
    "起草",
    "审核",
    "批准",
    "Department",
    "Signature",
    "Prepared by",
    "Reviewed by",
    "Approved by",
    "分发",
    "Distribution",
)

_IMAGE_EXT_CONTENT_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".bmp": "image/bmp",
    ".tiff": "image/tiff",
    ".tif": "image/tiff",
}

_MD_IMAGE_RE = re.compile(r"!\[image\]\((img_\d+\.[A-Za-z0-9]+)\)")


@dataclass(slots=True)
class ExtractedImage:
    """从 docx 中提取的内嵌图片，name 为 md 中的占位引用名。"""

    name: str
    data: bytes
    content_type: str


def _is_section_number(text: str) -> bool:
    """判断文本是否为序号格式：1、1.、1.1、1.2.3、1、2.3.4 等。"""
    text = text.strip()
    if not text:
        return False
    return bool(re.match(r"^\d+(\.\d+)*[.、]?$", text))


def _heading_level(seq_num: str) -> int:
    """序号转标题级别：1→2 级、1.1→3 级、1.1.1→4 级，更深层级不转标题。"""
    num = seq_num.rstrip(".").rstrip("、")
    depth = len(num.split("."))
    if depth == 1:
        return 2
    if depth == 2:
        return 3
    if depth == 3:
        return 4
    return 0


def _find_seq_column(table: Table, max_rows_to_check: int = 30) -> int:
    """在表格前几列中查找包含序号的列，未找到返回 -1。"""
    if not table.rows:
        return -1

    cols = len(table.rows[0].cells)
    max_check_cols = min(cols, 4)

    best_col = -1
    best_count = 0
    for col_idx in range(max_check_cols):
        count = 0
        for r in range(min(max_rows_to_check, len(table.rows))):
            if col_idx < len(table.rows[r].cells):
                cell_text = table.rows[r].cells[col_idx].text.strip()
                if _is_section_number(cell_text):
                    count += 1
        if count >= 2 and count > best_count:
            best_count = count
            best_col = col_idx
    return best_col


def _find_content_start_row(table: Table, max_rows_to_check: int = 30) -> int:
    """查找正文内容开始的行号（第一个有序号的行）。"""
    seq_col = _find_seq_column(table, max_rows_to_check)
    if seq_col < 0:
        return 0

    for r in range(min(max_rows_to_check, len(table.rows))):
        if seq_col < len(table.rows[r].cells):
            cell_text = table.rows[r].cells[seq_col].text.strip()
            if _is_section_number(cell_text):
                return r
    return 0


def _is_small_header_table(table: Table) -> bool:
    """审批格、分发表等小表格：≤10 行、无序号列、首页命中审批/分发关键词。"""
    rows = len(table.rows)
    if rows > 10 or not table.rows:
        return False
    if _find_seq_column(table) >= 0:
        return False

    first_row_text = " ".join(cell.text.strip() for cell in table.rows[0].cells)
    return any(kw in first_row_text for kw in _HEADER_KEYWORDS)


def _find_main_content_table(doc: Any) -> Table | None:
    """按评分找主内容表格：行数 + 序号列 +500 + 内容字符数 - 小页眉表 500。"""
    best_table: Table | None = None
    best_score = 0

    for table in doc.tables:
        rows = len(table.rows)
        if rows < 3:
            continue

        score = rows
        if _find_seq_column(table) >= 0:
            score += 500

        content_chars = 0
        for r in range(min(rows, 50)):
            if len(table.rows[r].cells) >= 2:
                content_chars += len(table.rows[r].cells[1].text)
        if content_chars > 1000:
            score += 200
        elif content_chars > 500:
            score += 100

        if _is_small_header_table(table):
            score -= 500

        if score > best_score:
            best_score = score
            best_table = table
    return best_table


def _extract_images(doc: Any) -> dict[str, ExtractedImage]:
    """提取文档全部图片，返回 {relationship_id: ExtractedImage}。"""
    image_map: dict[str, ExtractedImage] = {}
    counter = 0
    for rel_id, rel in doc.part.rels.items():
        if "image" not in rel.reltype:
            continue
        try:
            data = rel.target_part.blob
            ext = os.path.splitext(rel.target_part.partname)[1].lower()
            if not ext:
                ext = ".png"
            name = f"img_{counter:03d}{ext}"
            image_map[rel_id] = ExtractedImage(
                name=name,
                data=data,
                content_type=_IMAGE_EXT_CONTENT_TYPES.get(
                    ext, "application/octet-stream"
                ),
            )
            counter += 1
        except Exception:  # noqa: BLE001 单张图片损坏不阻断整体转换
            logger.warning(
                "docx image extraction failed for one relationship",
                extra={"component": "quality", "rel_id": rel_id},
                exc_info=True,
            )
    return image_map


def _image_rels_in_cell(cell: _Cell) -> list[str]:
    """获取单元格中全部图片 relationship ID。"""
    rel_ids: list[str] = []
    for blip in cell._element.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid:
            rel_ids.append(rid)
    return rel_ids


def _table_to_2d_list(table: Table) -> list[list[str]]:
    result: list[list[str]] = []
    for row in table.rows:
        result.append(
            [cell.text.strip().replace("\n", " ") for cell in row.cells]
        )
    return result


def _dedupe_table_columns(rows_data: list[list[str]]) -> list[list[str]]:
    """整列内容完全相同的重复列只保留一列（如分发表 4 列相同）。"""
    if not rows_data or len(rows_data[0]) <= 1:
        return rows_data

    num_cols = len(rows_data[0])
    cols_to_keep: list[int] = []
    for col_idx in range(num_cols):
        is_duplicate = False
        for prev_col in cols_to_keep:
            all_same = True
            for row in rows_data:
                if col_idx < len(row) and prev_col < len(row):
                    if row[col_idx] != row[prev_col]:
                        all_same = False
                        break
                else:
                    all_same = False
                    break
            if all_same:
                is_duplicate = True
                break
        if not is_duplicate:
            cols_to_keep.append(col_idx)

    return [
        [row[i] if i < len(row) else "" for i in cols_to_keep]
        for row in rows_data
    ]


def _table_data_to_md(table_data: list[list[str]]) -> str:
    """二维表格数据转 Markdown 表格。"""
    if not table_data:
        return ""
    lines = [""]
    for ri, row in enumerate(table_data):
        if any(cell.strip() for cell in row):
            lines.append("| " + " | ".join(cell.strip() for cell in row) + " |")
            if ri == 0:
                lines.append("|" + "|".join(["---"] * len(row)) + "|")
    lines.append("")
    return "\n".join(lines)


def _nested_table_to_md(
    table: Table,
    image_map: dict[str, ExtractedImage],
) -> str:
    """嵌套表格转 Markdown：图片嵌入对应单元格，重复列去重。"""
    if not table.rows:
        return ""

    cell_data: list[list[dict[str, Any]]] = []
    for row in table.rows:
        row_data: list[dict[str, Any]] = []
        for cell in row.cells:
            row_data.append(
                {
                    "text": cell.text.strip().replace("\n", " "),
                    "img_rids": _image_rels_in_cell(cell),
                }
            )
        cell_data.append(row_data)

    if cell_data and len(cell_data[0]) > 1:
        num_cols = len(cell_data[0])
        cols_to_keep: list[int] = []
        for col_idx in range(num_cols):
            is_duplicate = False
            for prev_col in cols_to_keep:
                all_same = True
                for row in cell_data:
                    if col_idx < len(row) and prev_col < len(row):
                        if row[col_idx]["text"] != row[prev_col]["text"]:
                            all_same = False
                            break
                    else:
                        all_same = False
                        break
                if all_same:
                    is_duplicate = True
                    break
            if not is_duplicate:
                cols_to_keep.append(col_idx)
        cell_data = [
            [row[c] for c in cols_to_keep if c < len(row)] for row in cell_data
        ]

    lines = [""]
    for ri, row in enumerate(cell_data):
        cells_md: list[str] = []
        for cell_info in row:
            parts: list[str] = []
            if cell_info["text"]:
                parts.append(str(cell_info["text"]))
            for rid in cell_info["img_rids"]:
                image = image_map.get(str(rid))
                if image is not None:
                    parts.append(f"![image]({image.name})")
            cells_md.append(" ".join(parts) if parts else "")
        if any(c.strip() for c in cells_md):
            lines.append("| " + " | ".join(c.strip() for c in cells_md) + " |")
            if ri == 0:
                lines.append("|" + "|".join(["---"] * len(cells_md)) + "|")
    lines.append("")
    return "\n".join(lines)


def _tables_identical(a: list[list[str]], b: list[list[str]]) -> bool:
    if len(a) != len(b):
        return False
    for row_a, row_b in zip(a, b, strict=False):
        if row_a != row_b:
            return False
    return True


def _unique_nested_tables(cell_tables_list: list[list[Table]]) -> list[Table]:
    """多个单元格中的嵌套表格按内容去重。"""
    unique_tables: list[Table] = []
    seen_data: list[list[list[str]]] = []
    for tables in cell_tables_list:
        for nested in tables:
            nested_data = _table_to_2d_list(nested)
            if not any(_tables_identical(nested_data, seen) for seen in seen_data):
                seen_data.append(nested_data)
                unique_tables.append(nested)
    return unique_tables


def _is_data_table_header_row(row: Any, seq_col: int) -> bool:
    """数据表表头行（如修订简历）：序号列无序号且序号列后有 ≥2 个唯一内容列。"""
    cells = row.cells
    if seq_col < len(cells) and _is_section_number(cells[seq_col].text.strip()):
        return False

    seen_content: set[str] = set()
    for c in range(seq_col + 1, len(cells)):
        text = cells[c].text.strip()
        if text:
            seen_content.add(text)
    return len(seen_content) >= 2


def _find_column_groups(
    main_table: Table, seq_col: int, start_row: int = 0, check_rows: int = 20
) -> list[int]:
    """识别重复列组（如分发的多列相同），返回每组的代表列号。"""
    num_cols = len(main_table.rows[0].cells) if main_table.rows else 0
    content_cols = list(range(seq_col + 1, num_cols))
    if not content_cols:
        return []

    check_rows = min(check_rows, len(main_table.rows) - start_row)
    group_reps: list[int] = []
    for col_idx in content_cols:
        is_dup = False
        for rep_col in group_reps:
            all_same = True
            has_content = False
            for r in range(start_row, start_row + check_rows):
                if r >= len(main_table.rows):
                    break
                cells = main_table.rows[r].cells
                if col_idx < len(cells) and rep_col < len(cells):
                    t1 = cells[col_idx].text.strip()
                    t2 = cells[rep_col].text.strip()
                    if t1 or t2:
                        has_content = True
                        if t1 != t2:
                            all_same = False
                            break
            if has_content and all_same:
                is_dup = True
                break
        if not is_dup:
            group_reps.append(col_idx)
    return group_reps


def _cell_paragraphs(cell: _Cell) -> list[str]:
    return [p.text.strip() for p in cell.paragraphs if p.text.strip()]


def _merge_column_group_contents(row: Any, group_reps: list[int]) -> str:
    """合并列组内容：完全相同只留一份；段落数相近配对为"名称（编号）"；否则拼接唯一项。"""
    if not group_reps:
        return ""

    cells = row.cells
    group_paragraphs: list[list[str]] = []
    for col_idx in group_reps:
        if col_idx < len(cells):
            paragraphs = _cell_paragraphs(cells[col_idx])
            if paragraphs:
                group_paragraphs.append(paragraphs)

    if not group_paragraphs:
        return ""
    if len(group_paragraphs) == 1:
        return "\n".join(group_paragraphs[0])

    first_group = group_paragraphs[0]
    if all(g == first_group for g in group_paragraphs[1:]):
        return "\n".join(first_group)

    para_counts = [len(g) for g in group_paragraphs]
    min_count = min(para_counts)
    if len(set(para_counts)) <= 2 and max(para_counts) - min(para_counts) <= 1:
        paired: list[str] = []
        for i in range(min_count):
            parts = [g[i] for g in group_paragraphs if i < len(g)]
            if len(parts) == 2:
                paired.append(f"{parts[0]}（{parts[1]}）")
            else:
                paired.append(" ".join(parts))
        for i in range(min_count, max(para_counts)):
            for g in group_paragraphs:
                if i < len(g):
                    paired.append(g[i])
        return "\n".join(paired)

    all_paras: list[str] = []
    seen: set[str] = set()
    for g in group_paragraphs:
        for p in g:
            if p not in seen:
                seen.add(p)
                all_paras.append(p)
    return "\n".join(all_paras)


def _emit_nested_tables_and_images(
    md_lines: list[str],
    unique_nested_tables: list[Table],
    unique_images: list[str],
    image_map: dict[str, ExtractedImage],
) -> set[str]:
    """输出嵌套表格与独立图片（排除已被嵌套表使用的图片），返回已用图片 rid。"""
    used_rids: set[str] = set()
    for nested in unique_nested_tables:
        md_lines.append(_nested_table_to_md(nested, image_map))
    for nested in unique_nested_tables:
        for row in nested.rows:
            for cell in row.cells:
                used_rids.update(_image_rels_in_cell(cell))
    for rid in unique_images:
        image = image_map.get(rid)
        if image is not None and rid not in used_rids:
            md_lines.append(f"\n![image]({image.name})\n")
    return used_rids


def _process_main_table(
    main_table: Table, image_map: dict[str, ExtractedImage]
) -> list[str]:
    """主内容表格处理：序号转标题、数据表转 md 表格、嵌套表格与图片原位输出。"""
    md_lines: list[str] = []

    seq_col = _find_seq_column(main_table)
    if seq_col < 0:
        seq_col = 0

    start_row = _find_content_start_row(main_table)
    num_cols = len(main_table.rows[0].cells) if main_table.rows else 0

    col_group_reps = _find_column_groups(main_table, seq_col, start_row)
    if not col_group_reps:
        col_group_reps = [seq_col + 1] if seq_col + 1 < num_cols else [seq_col]

    i = start_row
    while i < len(main_table.rows):
        row = main_table.rows[i]
        cells = row.cells

        seq_num = cells[seq_col].text.strip() if seq_col < len(cells) else ""
        content_text = _merge_column_group_contents(row, col_group_reps)
        first_para = content_text.split("\n")[0] if content_text else ""
        other_paras = (
            content_text.split("\n")[1:] if "\n" in content_text else []
        )

        all_nested_tables_by_col: list[list[Table]] = []
        all_images_by_col: list[list[str]] = []
        for c in range(seq_col + 1, num_cols):
            if c < len(cells):
                cell = cells[c]
                if len(cell.tables) > 0:
                    all_nested_tables_by_col.append(list(cell.tables))
                cell_imgs = _image_rels_in_cell(cell)
                if cell_imgs:
                    all_images_by_col.append(cell_imgs)

        unique_nested_tables = _unique_nested_tables(all_nested_tables_by_col)
        seen_image_rids: set[str] = set()
        unique_images: list[str] = []
        for img_list in all_images_by_col:
            for rid in img_list:
                if rid not in seen_image_rids:
                    seen_image_rids.add(rid)
                    unique_images.append(rid)

        has_content = bool(
            seq_num or content_text or unique_nested_tables or unique_images
        )
        if not has_content:
            i += 1
            continue

        if _is_section_number(seq_num):
            heading_level = _heading_level(seq_num)
            prefix = "#" * heading_level
            heading_text = first_para if first_para else content_text
            md_lines.append(f"{prefix} {seq_num} {heading_text}")
            md_lines.append("")

            for para in other_paras:
                md_lines.append(f"- {para}")
            if other_paras:
                md_lines.append("")

            # 后续行构成数据表（如修订简历）时整体转 md 表格
            peek_rows = []
            for k in range(i + 1, min(i + 6, len(main_table.rows))):
                peek_row = main_table.rows[k]
                peek_seq = (
                    peek_row.cells[seq_col].text.strip()
                    if seq_col < len(peek_row.cells)
                    else ""
                )
                if _is_section_number(peek_seq):
                    break
                peek_rows.append(peek_row)

            if len(peek_rows) >= 2 and _is_data_table_header_row(
                peek_rows[0], seq_col
            ):
                data_rows: list[list[str]] = []
                j = i + 1
                while j < len(main_table.rows):
                    data_row = main_table.rows[j]
                    data_seq = (
                        data_row.cells[seq_col].text.strip()
                        if seq_col < len(data_row.cells)
                        else ""
                    )
                    if _is_section_number(data_seq):
                        break
                    data_rows.append(
                        [
                            data_row.cells[c].text.strip().replace("\n", " ")
                            for c in range(seq_col, min(num_cols, len(data_row.cells)))
                        ]
                    )
                    j += 1
                md_lines.append(_table_data_to_md(_dedupe_table_columns(data_rows)))
                i = j
                continue

            _emit_nested_tables_and_images(
                md_lines, unique_nested_tables, unique_images, image_map
            )
            i += 1
        else:
            _emit_nested_tables_and_images(
                md_lines, unique_nested_tables, unique_images, image_map
            )
            if content_text:
                if "\n" in content_text:
                    for para in content_text.split("\n"):
                        if para.strip():
                            md_lines.append(para.strip())
                            md_lines.append("")
                else:
                    md_lines.append(content_text)
                    md_lines.append("")
            i += 1

    return md_lines


def _extract_file_number(file_name: str) -> str:
    """从文件名提取文件编号，如 STP-QS-MC-001-06xxx.docx → STP-QS-MC-001-06。"""
    base = os.path.splitext(file_name)[0].strip()
    match = re.match(r"^([A-Z0-9]+-[A-Z0-9()（）]+-\d+-\d+)", base, re.IGNORECASE)
    if match:
        return match.group(1)
    match = re.match(r"^([A-Z0-9]+-[A-Z0-9()（）]+-\d+)", base, re.IGNORECASE)
    if match:
        return match.group(1)
    return base


def _extract_file_title(file_name: str) -> str:
    """从文件名提取标题（去掉编号部分）。"""
    base = os.path.splitext(file_name)[0].strip()
    for pattern in (
        r"^[A-Z0-9]+-[A-Z0-9()（）]+-\d+-\d+\s*(.*)",
        r"^[A-Z0-9]+-[A-Z0-9()（）]+-\d+\s*(.*)",
    ):
        match = re.match(pattern, base, re.IGNORECASE)
        if match and match.group(1).strip():
            return match.group(1).strip()
    return base


def _extract_effective_date(doc: Any) -> str:
    """从首页表格提取生效日期（"生效日期/Effective Date" 右侧单元格）。"""
    if not doc.tables:
        return ""
    for row in doc.tables[0].rows:
        cells = row.cells
        for i, cell in enumerate(cells):
            text = cell.text.strip()
            if "生效日期" in text or "Effective Date" in text:
                if i + 1 < len(cells):
                    date_text = cells[i + 1].text.strip()
                    if date_text and date_text != text:
                        return str(date_text)
    return ""


def _is_inline_small_header_table(tbl_rows: list[Any]) -> bool:
    """XML 级判断首页审批格/分发表：≤10 行、命中关键词且无序号行。"""
    if len(tbl_rows) > 10 or not tbl_rows:
        return False

    first_row_cells = tbl_rows[0].findall(qn("w:tc"))
    first_row_text = " ".join(
        "".join(t.text or "" for t in cell.findall(".//" + qn("w:t"))).strip()
        for cell in first_row_cells
    )
    if not any(kw in first_row_text for kw in _HEADER_KEYWORDS):
        return False

    for tr in tbl_rows[:20]:
        tcs = tr.findall(qn("w:tc"))
        if tcs:
            first_cell_text = "".join(
                t.text or "" for t in tcs[0].findall(".//" + qn("w:t"))
            ).strip()
            if _is_section_number(first_cell_text):
                return False
    return True


def _body_table_rows(element: Any) -> list[list[str]]:
    """XML 级提取表格全部行文本（主内容之后的普通表格用）。"""
    all_rows: list[list[str]] = []
    for tr in element.findall(qn("w:tr")):
        cells_text = [
            "".join(t.text or "" for t in tc.findall(".//" + qn("w:t")))
            .strip()
            .replace("\n", " ")
            for tc in tr.findall(qn("w:tc"))
        ]
        if cells_text:
            all_rows.append(cells_text)
    return all_rows


def convert_docx_content_to_md(
    content: bytes, file_name: str
) -> tuple[str, list[ExtractedImage]]:
    """docx 字节内容 → (标准 MD 文本, 内嵌图片列表)。

    md 中图片以 `![image](img_000.png)` 占位名引用，由调用方替换为实际
    存储 URL。文件损坏无法解析时抛出 AppException。
    """
    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001 伪 docx/损坏文件统一转为业务异常
        raise AppException(message="word 文档解析失败，文件可能已损坏") from exc

    image_map = _extract_images(doc)

    md_lines: list[str] = [
        f"# {_extract_file_title(file_name)}",
        "",
        f"**文件编号**: {_extract_file_number(file_name)}",
        "",
    ]
    effective_date = _extract_effective_date(doc)
    if effective_date:
        md_lines.append(f"**生效日期**: {effective_date}")
        md_lines.append("")
    md_lines.append("---")
    md_lines.append("")

    main_content_table = _find_main_content_table(doc)
    main_content_elem = main_content_table._element if main_content_table else None

    body = doc.element.body
    header_section_passed = False
    seen_header_table = False
    header_marker_count = 0

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            full_text = "".join(
                t.text or "" for t in child.findall(".//" + qn("w:t"))
            ).strip()
            if not full_text:
                continue

            is_header_marker = bool(
                re.match(r"^[一二三四五六七八九十][、.]\s*", full_text)
            )
            if is_header_marker:
                header_marker_count += 1
                continue
            if re.match(r"^P\d+/\d+$", full_text):
                continue

            if not header_section_passed:
                looks_like_content = any(
                    kw in full_text and len(full_text) < 80
                    for kw in _CONTENT_KEYWORDS
                )
                if seen_header_table and looks_like_content:
                    header_section_passed = True
                if header_marker_count >= 3 and not is_header_marker:
                    header_section_passed = True

            if header_section_passed:
                sec_match = re.match(r"^(\d+(?:\.\d+)*)([.、]?)\s*(.*)", full_text)
                if sec_match and sec_match.group(3):
                    sec_num = sec_match.group(1)
                    sec_suffix = sec_match.group(2)
                    sec_content = sec_match.group(3).strip()
                    has_marker = bool(sec_suffix) or "." in sec_num
                    level = _heading_level(sec_num) if has_marker else 0
                    if level > 0:
                        md_lines.append(f"{'#' * level} {sec_num} {sec_content}")
                        md_lines.append("")
                    else:
                        md_lines.append(full_text)
                        md_lines.append("")
                else:
                    md_lines.append(full_text)
                    md_lines.append("")

        elif tag == "tbl":
            tbl_rows = child.findall(qn("w:tr"))

            if _is_inline_small_header_table(tbl_rows):
                seen_header_table = True
                continue

            is_main_content = child is main_content_elem
            if seen_header_table and not is_main_content and not header_section_passed:
                header_section_passed = True

            if is_main_content and main_content_table:
                header_section_passed = True
                try:
                    has_section_nums = _find_seq_column(main_content_table) >= 0
                    if not has_section_nums and len(main_content_table.rows) >= 3:
                        # 纯数据表格（无序号列）
                        all_rows_data = _table_to_2d_list(main_content_table)
                        md_lines.append(
                            _table_data_to_md(_dedupe_table_columns(all_rows_data))
                        )
                    else:
                        table_md = _process_main_table(main_content_table, image_map)
                        md_lines.extend(table_md)
                except Exception:  # noqa: BLE001 主表处理失败不阻断其余内容
                    logger.warning(
                        "main content table convert failed",
                        extra={"component": "quality", "file": file_name},
                        exc_info=True,
                    )
                    md_lines.append("\n*[主内容表格处理失败]*\n")
                continue

            if header_section_passed and len(tbl_rows) > 0:
                md_lines.append(
                    _table_data_to_md(_dedupe_table_columns(_body_table_rows(child)))
                )

    images = sorted(image_map.values(), key=lambda img: img.name)
    used_names = set(_MD_IMAGE_RE.findall("\n".join(md_lines)))
    images = [img for img in images if img.name in used_names]
    return "\n".join(md_lines), images
