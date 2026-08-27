"""Document catalog export service: 标准模板留存 + 导出 docx 生成。

导出格式保真：以留存模板的数据行作为格式原型（行高/单元格对齐/字体/居中），
通过深拷贝行 XML 生成新行，仅替换文本内容，保证导出与模板格式完全一致。
"""

from __future__ import annotations

import copy
import io
import logging
import shutil
from pathlib import Path
from typing import Any

from docx.oxml import parse_xml
from docx.oxml.ns import nsmap, qn
from lxml.etree import _Element as Element  # type: ignore[import-untyped]

from app.modules.quality.models.document_catalog import DocumentEntry

logger = logging.getLogger(__name__)

HEADERS = ["序号", "文件名称", "文件编码", "生效日期"]

# 导出字体要求：中文（eastAsia）宋体，数字/英文（ascii/hAnsi/cs）Times New Roman
REQUIRED_FONT_ATTRS = (
    ("w:ascii", "Times New Roman"),
    ("w:hAnsi", "Times New Roman"),
    ("w:eastAsia", "宋体"),
    ("w:cs", "Times New Roman"),
)


def _replace_cell_text(tc: Element, new_text: str) -> None:
    """替换 w:tc 单元格文本，保留段落、行内文字的全部格式（字体/大小/对齐/居中）。

    策略：找到该单元格下第一个 w:r → w:t，将文本设为新值；
    清空同段落内其他 run 的文本；若单元格为空则创建最小 w:p → w:r → w:t 结构。
    """
    w_ns = nsmap["w"]
    p_els = tc.findall(qn("w:p"))
    if not p_els:
        p_el = parse_xml(f'<w:p xmlns:w="{w_ns}"><w:r><w:t></w:t></w:r></w:p>')
        tc.append(p_el)
        r_el = p_el.find(qn("w:r"))
        t_el = r_el.find(qn("w:t"))
        t_el.text = new_text
        t_el.set(qn("xml:space"), "preserve")
        _ensure_run_fonts(r_el, p_el)
        return

    # 仅保留第一个段落（其格式），移除多余段落，避免旧文本残留（如序号列双段落）
    first_p, rest_ps = p_els[0], p_els[1:]
    for extra_p in rest_ps:
        tc.remove(extra_p)

    # 清空该段落内所有 run 的文本，保留格式元素，并记录承载文本的 run
    target_run = None
    for r_el in first_p.findall(qn("w:r")):
        t_el = r_el.find(qn("w:t"))
        if t_el is not None:
            if target_run is None:
                t_el.text = new_text
                t_el.set(qn("xml:space"), "preserve")
                target_run = r_el
            else:
                t_el.text = ""
    if target_run is None:
        target_run = first_p.find(qn("w:r"))
        if target_run is None:
            target_run = parse_xml(f'<w:r xmlns:w="{w_ns}"><w:t></w:t></w:r>')
            first_p.append(target_run)
        t_el = target_run.find(qn("w:t"))
        if t_el is None:
            t_el = parse_xml(f'<w:t xmlns:w="{w_ns}"></w:t>')
            target_run.append(t_el)
        t_el.text = new_text
        t_el.set(qn("xml:space"), "preserve")
    _ensure_run_fonts(target_run, first_p)


def _ensure_run_fonts(r_el: Any, p_el: Any) -> None:
    """保证承载文本的 run 满足导出字体要求（宋体 + Times New Roman）。

    run 无 rPr 时优先继承段落 rPr（保留模板字号等格式），
    再强制 rFonts 四属性符合导出字体要求，避免依赖文档默认样式。
    """
    rpr_el = r_el.find(qn("w:rPr"))
    if rpr_el is None:
        ppr_el = p_el.find(qn("w:pPr"))
        src_rpr = ppr_el.find(qn("w:rPr")) if ppr_el is not None else None
        rpr_el = (
            copy.deepcopy(src_rpr)
            if src_rpr is not None
            else parse_xml(f'<w:rPr xmlns:w="{nsmap["w"]}"/>')
        )
        r_el.insert(0, rpr_el)
    fonts_el = rpr_el.find(qn("w:rFonts"))
    if fonts_el is None:
        fonts_el = parse_xml(f'<w:rFonts xmlns:w="{nsmap["w"]}"/>')
        rpr_el.insert(0, fonts_el)
    for attr, value in REQUIRED_FONT_ATTRS:
        fonts_el.set(qn(attr), value)


def _sanitize_copied_row(tr_el: Any) -> None:
    """清理深拷贝模板行中会导致渲染异常的元素。

    - 移除各段落的 w:numPr：留存模板「序号」列使用 Word 自动编号（单元格文本为空），
      拷贝行写入手动序号后，Word 会在手动序号旁再渲染一个自动编号，
      导致导出文档序号列出现两个数字；
    - 移除书签/编辑保护范围元素：模板首数据行的这类元素带固定 id，
      逐行拷贝后 id 大量重复，易触发 Word 文档修复提示。
    """
    for p_el in tr_el.iter(qn("w:p")):
        ppr_el = p_el.find(qn("w:pPr"))
        if ppr_el is not None:
            num_pr = ppr_el.find(qn("w:numPr"))
            if num_pr is not None:
                ppr_el.remove(num_pr)
    for tag in ("w:bookmarkStart", "w:bookmarkEnd", "w:permStart", "w:permEnd"):
        for el in list(tr_el.iter(qn(tag))):
            el.getparent().remove(el)


def get_templates_dir() -> Path:
    """标准格式模板留存目录（不进 git，业务数据）。"""
    path = Path(__file__).resolve().parent.parent / "templates" / "document_catalog"
    path.mkdir(parents=True, exist_ok=True)
    return path


def seed_template_from_source(source_path: str, department_name: str = "") -> Path:
    """将标准文件格式 docx 复制到留存目录（初始化脚本用），返回留存路径。"""
    source = Path(source_path)
    if not source.exists():
        raise FileNotFoundError(f"标准模板源文件不存在：{source_path}")
    stem = department_name or source.stem
    destination = get_templates_dir() / f"{stem}.docx"
    shutil.copyfile(source, destination)
    logger.info(
        "document catalog template seeded",
        extra={"component": "quality", "source": str(source), "dest": str(destination)},
    )
    return destination


def find_template(department_name: str = "") -> Path | None:
    """查找留存模板：优先按部门名匹配，其次取留存目录第一份；无留存返回 None。"""
    templates_dir = get_templates_dir()
    files = sorted(templates_dir.glob("*.docx"))
    if not files:
        return None
    if department_name:
        for f in files:
            if department_name in f.stem:
                return f
    return files[0]


def _find_target_table(document: Any) -> Any:
    """定位目标表格：优先表头含「文件名称」的表格，否则取第一个表格。"""
    for table in document.tables:
        if len(table.rows) > 0:
            header_text = "".join(cell.text for cell in table.rows[0].cells)
            if "文件名称" in header_text:
                return table
    return document.tables[0] if document.tables else None


def export_document_catalog_docx(
    entries: list[DocumentEntry], department_name: str
) -> bytes:
    """按留存标准模板生成导出 docx；无模板时程序化生成标准格式。

    与留存模板格式一致：标题/部门段落 + [序号, 文件名称, 文件编码, 生效日期] 表格，
    日期格式 YYYY.MM.DD。
    """
    import docx

    template = find_template(department_name)
    if template is not None:
        document = docx.Document(str(template))
        table = _find_target_table(document)
    else:
        document = docx.Document()
        document.add_paragraph("文件目录清单")
        document.add_paragraph(f"部门：{department_name}")
        document.add_paragraph("确认人/日期：")
        table = document.add_table(rows=1, cols=len(HEADERS))
        table.style = "Table Grid"
        for i, header in enumerate(HEADERS):
            table.rows[0].cells[i].text = header

    # 更新「部门：XXX」段落（如有）
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("部门") and ("：" in text or ":" in text):
            paragraph.text = f"部门：{department_name}"
            break

    # 清空模板数据行（保留表头行0），但先保留 row 1 作为格式原型
    if table is not None and len(table.rows) > 1:
        proto_row_el = copy.deepcopy(table.rows[1]._tr)
        for i in range(len(table.rows) - 1, 0, -1):
            tr = table.rows[i]._tr
            table._tbl.remove(tr)

        for idx, entry in enumerate(entries, start=1):
            row_el = copy.deepcopy(proto_row_el)
            _sanitize_copied_row(row_el)
            # 修改单元格文本内容（保留全部格式），并替换 _t element 中的 te:t 文本元素
            cells_el = row_el.findall(qn("w:tc"))
            # 导出序号从 1 开始连续编号，不沿用台账中可能跳号的 seq_no
            texts = [
                str(idx),
                entry.name or "",
                entry.code or "",
                (
                    entry.effective_date.strftime("%Y.%m.%d")
                    if entry.effective_date
                    else (entry.effective_date_text or "")
                ),
            ]
            for ci, cell_text in enumerate(texts):
                if ci < len(cells_el):
                    _replace_cell_text(cells_el[ci], cell_text)
            table._tbl.append(row_el)
    elif table is not None:
        # 无数据行可参考 → fallback add_row
        for i in range(len(table.rows) - 1, 0, -1):
            tr = table.rows[i]._tr
            table._tbl.remove(tr)
        for idx, entry in enumerate(entries, start=1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = entry.name or ""
            row.cells[2].text = entry.code or ""
            row.cells[3].text = (
                entry.effective_date.strftime("%Y.%m.%d")
                if entry.effective_date
                else (entry.effective_date_text or "")
            )

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()
