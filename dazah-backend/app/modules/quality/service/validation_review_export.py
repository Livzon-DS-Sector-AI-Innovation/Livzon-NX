"""验证方案/报告 AI 审核结果导出为 docx 审核报告。"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

CATEGORY_LABELS: dict[str, str] = {
    "reference_missing": "引用文件缺失",
    "version_mismatch": "引用版本不一致",
    "plan_report_mismatch": "方案报告不一致",
    "content_consistency": "内容一致性",
    "format_issue": "格式/编号问题",
    "numeric_check": "数值核对",
    "basis_content_mismatch": "依据内容不一致",
}
SEVERITY_LABELS: dict[str, str] = {
    "high": "高",
    "medium": "中",
    "low": "低",
}


def _set_run_font(run: Any, name: str = "宋体", size: float = 10.5) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def _add_paragraph(
    doc: Any,
    text: str,
    *,
    bold: bool = False,
    size: float = 10.5,
    align: int | None = None,
) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    _set_run_font(run, size=size)
    if align is not None:
        paragraph.alignment = align
    return None


def _add_heading(doc: Any, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        _set_run_font(run, name="黑体", size=14 if level == 1 else 12)
    return None


def _add_table(doc: Any, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            _set_run_font(run, size=9)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            for run in cells[index].paragraphs[0].runs:
                _set_run_font(run, size=9)
    return None


def _status_label(status: str) -> str:
    return {
        "draft": "草稿",
        "processing": "审核中",
        "completed": "已完成",
        "failed": "失败",
    }.get(status, status)


def _format_time(value: Any) -> str:
    """兼容 datetime 对象与 ISO 字符串。"""
    if not value:
        return "—"
    if isinstance(value, str):
        return value[:16].replace("T", " ")
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M")
    return str(value)[:16]


def build_review_docx(out_data: dict[str, Any]) -> bytes:
    """根据审核结果组装 docx 报告字节流（横向 A4，适配宽表格）。"""
    doc = Document()
    section = doc.sections[0]
    page_width, page_height = section.page_width, section.page_height
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = page_height
    section.page_height = page_width

    _add_paragraph(
        doc, "验证方案与报告 AI 审核报告", bold=True, size=16,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    meta = [
        ("审核记录", str(out_data.get("id") or "")[:8]),
        ("标题", out_data.get("title") or ""),
        ("来源", "页面上传"),
        ("状态", _status_label(out_data.get("status") or "")),
        ("生成时间", _format_time(out_data.get("last_generated_at"))),
    ]
    for label, value in meta:
        _add_paragraph(doc, f"{label}：{value}")

    summary = out_data.get("summary")
    if summary:
        _add_heading(doc, "审核结论", level=1)
        _add_paragraph(doc, summary)

    stats = out_data.get("stats") or {}
    if isinstance(stats, dict) and stats:
        _add_heading(doc, "问题统计", level=1)
        _add_table(
            doc,
            ["问题总数", "高", "中", "低", "引用核对数", "命中数", "方案报告核对"],
            [
                [
                    str(stats.get("total_findings", 0)),
                    str(stats.get("high", 0)),
                    str(stats.get("medium", 0)),
                    str(stats.get("low", 0)),
                    str(stats.get("references_checked", 0)),
                    str(stats.get("references_matched", 0)),
                    "已核对" if stats.get("plan_report_checked") else "未核对",
                ]
            ],
        )

    findings = out_data.get("findings") or []
    if isinstance(findings, list) and findings:
        _add_heading(doc, "发现问题", level=1)
        rows: list[list[str]] = []
        for index, finding in enumerate(findings, start=1):
            if not isinstance(finding, dict):
                continue
            rows.append(
                [
                    str(index),
                    CATEGORY_LABELS.get(
                        str(finding.get("category")), str(finding.get("category"))
                    ),
                    SEVERITY_LABELS.get(
                        str(finding.get("severity")), str(finding.get("severity"))
                    ),
                    str(finding.get("location") or ""),
                    str(finding.get("quote") or ""),
                    "已核" if finding.get("quote_verified") else "未核",
                    str(finding.get("validation_quote") or "")
                    if finding.get("category") == "basis_content_mismatch"
                    else "",
                    str(finding.get("basis_quote") or "")
                    if finding.get("category") == "basis_content_mismatch"
                    else "",
                    str(finding.get("detail") or ""),
                ]
            )
        _add_table(
            doc,
            [
                "序号",
                "分类",
                "严重度",
                "位置",
                "原文",
                "引文校验",
                "验证文档原文",
                "依据规程原文",
                "说明",
            ],
            rows,
        )

    basis_used = out_data.get("basis_used") or []
    if isinstance(basis_used, list) and basis_used:
        _add_heading(doc, "引用文件核对明细", level=1)
        rows = []
        for index, item in enumerate(basis_used, start=1):
            if not isinstance(item, dict):
                continue
            status_text = {
                "version_mismatch": "版本不一致",
                "missing": "目录缺失",
                "none": "一致",
            }.get(str(item.get("issue")), str(item.get("issue")))
            rows.append(
                [
                    str(index),
                    str(item.get("code") or ""),
                    str(item.get("entry_name") or ""),
                    str(item.get("entry_code") or ""),
                    status_text,
                ]
            )
        _add_table(doc, ["序号", "引用编号", "目录文件", "目录编号", "核对结果"], rows)

    _add_paragraph(doc, "")
    _add_paragraph(
        doc,
        "说明：本报告由 AI 辅助生成，仅作为审核辅助依据，"
        "最终审批仍按线下纸质流程执行。",
        size=9,
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
