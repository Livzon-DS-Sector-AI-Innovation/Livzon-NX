"""Generate change ledger exports from the local Word template."""

from __future__ import annotations

import copy
import logging
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from app.core.exceptions import AppException

logger = logging.getLogger(__name__)


_TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent / "templates" / "变更台账-模板.docx"
)


def _safe_str(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    return str(value)


def _format_date(value: Any) -> str:
    if value in (None, ""):
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y.%m.%d")
    if isinstance(value, date):
        return value.strftime("%Y.%m.%d")

    text = _safe_str(value).strip()
    if not text:
        return ""

    normalized = text.replace("/", "-").replace(".", "-").replace("T", " ")
    if normalized.endswith("Z"):
        normalized = normalized[:-1] + "+00:00"

    for parser in (datetime.fromisoformat,):
        try:
            return parser(normalized).strftime("%Y.%m.%d")
        except ValueError:
            continue

    return text.replace("-", ".").replace("/", ".")


# 导出字体要求：中文（eastAsia）宋体，数字/英文（ascii/hAnsi/cs）Times New Roman
_REQUIRED_FONT_ATTRS = (
    ("w:ascii", "Times New Roman"),
    ("w:hAnsi", "Times New Roman"),
    ("w:eastAsia", "宋体"),
    ("w:cs", "Times New Roman"),
)


def _ensure_run_fonts(run_elem: Any, paragraph_elem: Any) -> None:
    """保证承载文本的 run 满足导出字体要求（宋体 + Times New Roman）。

    run 无 rPr 时优先继承段落 rPr（保留模板字号等格式），
    再强制 rFonts 四属性符合导出字体要求，避免依赖文档默认样式。
    """
    rpr_elem = run_elem.find(qn("w:rPr"))
    if rpr_elem is None:
        ppr_elem = paragraph_elem.find(qn("w:pPr"))
        src_rpr = ppr_elem.find(qn("w:rPr")) if ppr_elem is not None else None
        rpr_elem = (
            copy.deepcopy(src_rpr)
            if src_rpr is not None
            else run_elem.makeelement(qn("w:rPr"), {})
        )
        run_elem.insert(0, rpr_elem)
    fonts_elem = rpr_elem.find(qn("w:rFonts"))
    if fonts_elem is None:
        fonts_elem = rpr_elem.makeelement(qn("w:rFonts"), {})
        rpr_elem.insert(0, fonts_elem)
    for attr, value in _REQUIRED_FONT_ATTRS:
        fonts_elem.set(qn(attr), value)


def _clear_text_in_paragraph(paragraph_elem: Any) -> None:
    for run_elem in paragraph_elem.findall(qn("w:r")):
        for text_elem in run_elem.findall(qn("w:t")):
            text_elem.text = ""


def _set_cell_text(cell_elem: Any, text: str) -> None:
    lines = (text or "").split("\n")
    paragraph_elems = cell_elem.findall(qn("w:p"))

    if not paragraph_elems:
        paragraph_elem = cell_elem.makeelement(qn("w:p"), {})
        cell_elem.append(paragraph_elem)
        paragraph_elems = [paragraph_elem]

    while len(paragraph_elems) > len(lines):
        paragraph_elems[-1].getparent().remove(paragraph_elems[-1])
        paragraph_elems = cell_elem.findall(qn("w:p"))

    while len(paragraph_elems) < len(lines):
        new_paragraph = copy.deepcopy(paragraph_elems[0])
        _clear_text_in_paragraph(new_paragraph)
        cell_elem.append(new_paragraph)
        paragraph_elems = cell_elem.findall(qn("w:p"))

    for index, line in enumerate(lines):
        paragraph_elem = paragraph_elems[index]
        run_elems = paragraph_elem.findall(qn("w:r"))
        for run_elem in run_elems:
            for text_elem in run_elem.findall(qn("w:t")):
                text_elem.text = ""

        if not run_elems:
            run_elem = paragraph_elem.makeelement(qn("w:r"), {})
            paragraph_elem.append(run_elem)
            run_elems = [run_elem]

        text_elems = run_elems[0].findall(qn("w:t"))
        if not text_elems:
            text_elem = run_elems[0].makeelement(qn("w:t"), {})
            run_elems[0].append(text_elem)
            text_elems = [text_elem]

        text_elems[0].text = line
        _ensure_run_fonts(run_elems[0], paragraph_elem)


def _strip_auto_numbering(elem: Any) -> None:
    """移除元素内各段落的 w:numPr（Word 自动编号）。

    模板数据行单元格文本为空、序号等列依赖自动编号；
    拷贝行写入手动文本后，Word 会叠加渲染自动编号，导致导出内容重复。
    """
    for p_elem in elem.iter(qn("w:p")):
        ppr_elem = p_elem.find(qn("w:pPr"))
        if ppr_elem is not None:
            num_pr = ppr_elem.find(qn("w:numPr"))
            if num_pr is not None:
                ppr_elem.remove(num_pr)


def _resolve_serial_number(item: dict[str, Any], index: int) -> str:
    serial_number = _safe_str(item.get("serial_number")).strip()
    return serial_number or str(index)


def _resolve_closure_value(item: dict[str, Any]) -> str:
    closure_date = _format_date(item.get("closure_date"))
    if closure_date:
        return closure_date
    return _safe_str(item.get("status")).strip()


def _find_max_perm_id(doc: Any) -> int:
    """Find the maximum permStart/permEnd id in the whole document.

    Searches the entire document body so new ids never collide with ids
    already used in the title area (ids 0/1) or any other region.
    """
    max_id = -1
    body = doc.element.body
    for perm_elem in body.iter(qn("w:permStart")):
        id_str = perm_elem.get(qn("w:id"))
        if id_str:
            max_id = max(max_id, int(id_str))
    for perm_elem in body.iter(qn("w:permEnd")):
        id_str = perm_elem.get(qn("w:id"))
        if id_str:
            max_id = max(max_id, int(id_str))
    return max_id


def _remove_perm_markers(row_elem: Any) -> None:
    """Remove all permStart/permEnd markers from a row element.

    Handles both direct children of w:tr and descendants inside cells.
    """
    # Remove from direct children of row (permStart/permEnd between <w:tr> and <w:tc>)
    for perm_elem in list(row_elem):
        tag = perm_elem.tag.split("}")[-1] if "}" in perm_elem.tag else perm_elem.tag
        if tag in ("permStart", "permEnd"):
            row_elem.remove(perm_elem)

    # Remove from descendants (inside cells/paragraphs)
    for perm_elem in list(row_elem.iter(qn("w:permStart"))):
        perm_elem.getparent().remove(perm_elem)
    for perm_elem in list(row_elem.iter(qn("w:permEnd"))):
        perm_elem.getparent().remove(perm_elem)


def _insert_perm_markers(row_elem: Any, perm_id: int) -> None:
    """Insert permStart at first cell and permEnd at last cell of a row.

    Inserts at row level (between cells) to match the original template structure.
    """
    cells = row_elem.findall(qn("w:tc"))
    if not cells:
        return

    # Insert permStart after the first cell
    perm_start = row_elem.makeelement(
        qn("w:permStart"),
        {
            qn("w:id"): str(perm_id),
            qn("w:edGrp"): "everyone",
        },
    )
    first_cell = cells[0]
    # Insert after first cell
    next_sibling = first_cell.getnext()
    if next_sibling is not None:
        next_sibling.addprevious(perm_start)
    else:
        row_elem.append(perm_start)

    # Insert permEnd after the last cell
    perm_end = row_elem.makeelement(
        qn("w:permEnd"),
        {
            qn("w:id"): str(perm_id),
        },
    )
    row_elem.append(perm_end)


def _set_enforcement(doc: Any) -> None:
    """Ensure document protection enforcement is set to 1."""
    settings_elem = doc.element.body.getparent().find(".//" + qn("w:settings"))
    if settings_elem is None:
        return

    protection = settings_elem.find(qn("w:documentProtection"))
    if protection is not None:
        protection.set(qn("w:enforcement"), "1")


def _remove_all_perm_markers(doc: Any) -> None:
    """Remove every permStart/permEnd marker in the document body.

    The template carries editable-range markers around the title area
    (ids 0/1) and at the table level (ids 2-8). Only the data rows we
    add below should be editable, so all existing markers are cleared
    and only fresh per-row markers are inserted.
    """
    body = doc.element.body
    for perm_elem in list(body.iter(qn("w:permStart"))):
        perm_elem.getparent().remove(perm_elem)
    for perm_elem in list(body.iter(qn("w:permEnd"))):
        perm_elem.getparent().remove(perm_elem)


def _set_ledger_subtitle(doc: Any, subtitle: str) -> None:
    """将模板副标题段（"技术变更"）替换为目标类型文案，保留原字体格式。"""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "技术变更":
            for run in paragraph.runs:
                if "技术变更" in run.text:
                    run.text = subtitle
                else:
                    run.text = ""
            return


def generate_change_ledger_export_docx(
    items: list[dict[str, Any]], ledger_title: str = "技术变更"
) -> bytes:
    """Render change ledger records into the local Word template."""
    doc = Document(str(_TEMPLATE_PATH))
    _set_ledger_subtitle(doc, ledger_title)
    if not doc.tables:
        raise AppException(message="变更台账模板中未找到表格")

    table = doc.tables[0]
    table_element = table._tbl
    row_elements = table_element.findall(qn("w:tr"))
    if len(row_elements) < 2:
        raise AppException(message="变更台账模板缺少可复用的数据行")

    # 模板表头/数据行带自动编号，先整体清理，避免导出序号等内容重复渲染
    _strip_auto_numbering(table_element)

    # 先记录模板中已有的最大 perm id（标题区 0/1 + 数据行 2-8），新 id 从其后递增
    max_perm_id = _find_max_perm_id(doc)

    # 找到第一个数据行作为模板行
    template_row = copy.deepcopy(row_elements[1])

    # 删除所有数据行（保留表头行）
    for row_elem in list(row_elements[1:]):
        table_element.remove(row_elem)

    # 清除模板中所有 permStart/permEnd（标题区 + 表级残留），只保留表头行
    _remove_all_perm_markers(doc)

    for index, item in enumerate(items, start=1):
        row_elem = copy.deepcopy(template_row)

        # 移除拷贝行中的 permStart/permEnd 标记
        _remove_perm_markers(row_elem)

        table_element.append(row_elem)
        cell_elems = row_elem.findall(qn("w:tc"))

        values = [
            _resolve_serial_number(item, index),
            _safe_str(item.get("change_code")).strip(),
            _safe_str(item.get("applicant_department")).strip(),
            _safe_str(item.get("change_object")).strip(),
            _safe_str(item.get("change_content")).strip(),
            _safe_str(item.get("change_level")).strip(),
            _format_date(item.get("application_date")),
            _format_date(item.get("planned_approval_date")),
            _format_date(item.get("execution_date")),
            _resolve_closure_value(item),
        ]

        for cell_index, value in enumerate(values):
            if cell_index >= len(cell_elems):
                break
            _set_cell_text(cell_elems[cell_index], value)

        # 为拷贝行插入新的 permStart/permEnd 标记（唯一 id）
        max_perm_id += 1
        _insert_perm_markers(row_elem, max_perm_id)

    # 确保文档保护启用
    _set_enforcement(doc)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
