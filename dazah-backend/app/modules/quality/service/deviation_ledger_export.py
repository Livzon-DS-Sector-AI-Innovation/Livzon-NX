"""Generate deviation ledger exports from a local Word template."""

from __future__ import annotations

import copy
import logging
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from app.core.exceptions import AppException

logger = logging.getLogger(__name__)


_TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent / "templates" / "偏差登记表-模板.docx"
)


def _safe_str(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    return str(value)


def _format_date(value: Any) -> str:
    if value in (None, ""):
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y.%m.%d")
    if isinstance(value, date):
        return value.strftime("%Y.%m.%d")

    text = _safe_str(value).strip()
    if not text:
        return ""

    normalized = text.replace("/", "-").replace(".", "-").replace("T", " ")
    if normalized.endswith("Z"):
        normalized = normalized[:-1] + "+00:00"

    for parser in (datetime.fromisoformat,):
        try:
            return parser(normalized).strftime("%Y.%m.%d")
        except ValueError:
            continue

    return text.replace("-", ".").replace("/", ".")


def _normalize_yes_no(value: Any) -> str:
    """通用是/否转换（用于是否关闭等列）。"""
    if isinstance(value, bool):
        return "是" if value else "否"

    text = _safe_str(value).strip().lower()
    if text in {"1", "true", "yes", "y", "是", "已关闭", "closed"}:
        return "是"
    if text in {"0", "false", "no", "n", "否", "未关闭", "draft", "open"}:
        return "否"
    return ""


def _format_occurred_column(
    has_occurred_before: Any, previous_occurrence_code: Any
) -> str:
    """生成"偏差是否曾发生"列的勾选格式文本。

    对齐桌面文档勾选格式：选中项不带方框，未选中项带方框。
    未发生：'[ ]是 编号：\\n否'；曾发生：'是 编号：PC-xxx\\n[ ]否'。
    使用 [ ] 而非 □，确保在所有字体/WPS 环境下都能正常显示。
    """
    if has_occurred_before is True:
        code = (
            _safe_str(previous_occurrence_code).strip()
            if previous_occurrence_code
            else ""
        )
        if code:
            return f"是 编号：{code}\n[ ]否"
        return "是 编号：\n[ ]否"
    elif has_occurred_before is False:
        return "[ ]是 编号：\n否"
    else:
        # 未知/空值
        return "[ ]是 编号：\n[ ]否"


def _normalize_level(value: Any) -> str:
    text = _safe_str(value).strip()
    if not text:
        return ""

    mapping = {
        "major": "重大",
        "moderate": "次要",
        "minor": "微小",
        "严重偏差": "重大",
        "中等偏差": "次要",
        "次要偏差": "微小",
    }
    return mapping.get(text.lower(), mapping.get(text, text))


def _build_product_batch(item: dict[str, Any]) -> str:
    product_batch = _safe_str(item.get("product_batch")).strip()
    if product_batch:
        return product_batch

    affected_items = _safe_str(item.get("affected_items")).strip()
    batch_number = _safe_str(item.get("batch_number")).strip()
    if affected_items and batch_number:
        return f"{affected_items}\n{batch_number}"
    return affected_items or batch_number


# 导出字体要求：中文（eastAsia）宋体，数字/英文（ascii/hAnsi/cs）Times New Roman
_REQUIRED_FONT_ATTRS = (
    ("w:ascii", "Times New Roman"),
    ("w:hAnsi", "Times New Roman"),
    ("w:eastAsia", "宋体"),
    ("w:cs", "Times New Roman"),
)


def _ensure_run_fonts(run_elem: Any, paragraph_elem: Any) -> None:
    """保证承载文本的 run 满足导出字体要求（宋体 + Times New Roman）。

    run 无 rPr 时优先继承段落 rPr（保留模板字号等格式），
    再强制 rFonts 四属性符合导出字体要求，避免依赖文档默认样式。
    """
    rpr_elem = run_elem.find(qn("w:rPr"))
    if rpr_elem is None:
        ppr_elem = paragraph_elem.find(qn("w:pPr"))
        src_rpr = ppr_elem.find(qn("w:rPr")) if ppr_elem is not None else None
        rpr_elem = (
            copy.deepcopy(src_rpr)
            if src_rpr is not None
            else run_elem.makeelement(qn("w:rPr"), {})
        )
        run_elem.insert(0, rpr_elem)
    fonts_elem = rpr_elem.find(qn("w:rFonts"))
    if fonts_elem is None:
        fonts_elem = rpr_elem.makeelement(qn("w:rFonts"), {})
        rpr_elem.insert(0, fonts_elem)
    for attr, value in _REQUIRED_FONT_ATTRS:
        fonts_elem.set(qn(attr), value)


def _clear_text_in_paragraph(paragraph_elem: Any) -> None:
    for run_elem in paragraph_elem.findall(qn("w:r")):
        for text_elem in run_elem.findall(qn("w:t")):
            text_elem.text = ""


def _set_cell_text(cell_elem: Any, text: str) -> None:
    lines = (text or "").split("\n")
    paragraph_elems = cell_elem.findall(qn("w:p"))

    if not paragraph_elems:
        paragraph_elem = cell_elem.makeelement(qn("w:p"), {})
        cell_elem.append(paragraph_elem)
        paragraph_elems = [paragraph_elem]

    while len(paragraph_elems) > len(lines):
        paragraph_elems[-1].getparent().remove(paragraph_elems[-1])
        paragraph_elems = cell_elem.findall(qn("w:p"))

    while len(paragraph_elems) < len(lines):
        new_paragraph = copy.deepcopy(paragraph_elems[0])
        _clear_text_in_paragraph(new_paragraph)
        cell_elem.append(new_paragraph)
        paragraph_elems = cell_elem.findall(qn("w:p"))

    for index, line in enumerate(lines):
        paragraph_elem = paragraph_elems[index]
        run_elems = paragraph_elem.findall(qn("w:r"))
        for run_elem in run_elems:
            for text_elem in run_elem.findall(qn("w:t")):
                text_elem.text = ""

        if not run_elems:
            run_elem = paragraph_elem.makeelement(qn("w:r"), {})
            paragraph_elem.append(run_elem)
            run_elems = [run_elem]

        text_elems = run_elems[0].findall(qn("w:t"))
        if not text_elems:
            text_elem = run_elems[0].makeelement(qn("w:t"), {})
            run_elems[0].append(text_elem)
            text_elems = [text_elem]

        text_elems[0].text = line
        _ensure_run_fonts(run_elems[0], paragraph_elem)


def _strip_auto_numbering(elem: Any) -> None:
    """移除元素内各段落的 w:numPr（Word 自动编号）。

    模板数据行单元格文本为空、部分列依赖自动编号；
    拷贝行写入手动文本后，Word 会叠加渲染自动编号，导致导出内容重复。
    """
    for p_elem in elem.iter(qn("w:p")):
        ppr_elem = p_elem.find(qn("w:pPr"))
        if ppr_elem is not None:
            num_pr = ppr_elem.find(qn("w:numPr"))
            if num_pr is not None:
                ppr_elem.remove(num_pr)


def _find_max_perm_id(doc: Any) -> int:
    """查找文档中已有的最大 permStart/permEnd id（标题区 id=0/1 等）。"""
    max_id = -1
    body = doc.element.body
    for perm_elem in body.iter(qn("w:permStart")):
        id_str = perm_elem.get(qn("w:id"))
        if id_str:
            max_id = max(max_id, int(id_str))
    for perm_elem in body.iter(qn("w:permEnd")):
        id_str = perm_elem.get(qn("w:id"))
        if id_str:
            max_id = max(max_id, int(id_str))
    return max_id


def _remove_all_perm_markers(doc: Any) -> None:
    """移除文档正文中所有 permStart/permEnd 标记。

    模板标题区/表头区可能带有原始标记，导出时仅对数据行重新添加，
    确保标题与表头锁定、数据行可编辑。
    """
    body = doc.element.body
    for perm_elem in list(body.iter(qn("w:permStart"))):
        perm_elem.getparent().remove(perm_elem)
    for perm_elem in list(body.iter(qn("w:permEnd"))):
        perm_elem.getparent().remove(perm_elem)


def _remove_perm_markers(row_elem: Any) -> None:
    """移除行内所有 permStart/permEnd（直接子级 + 单元格后代）。"""
    for perm_elem in list(row_elem):
        tag = perm_elem.tag.split("}")[-1] if "}" in perm_elem.tag else perm_elem.tag
        if tag in ("permStart", "permEnd"):
            row_elem.remove(perm_elem)
    for perm_elem in list(row_elem.iter(qn("w:permStart"))):
        perm_elem.getparent().remove(perm_elem)
    for perm_elem in list(row_elem.iter(qn("w:permEnd"))):
        perm_elem.getparent().remove(perm_elem)


def _insert_perm_markers(row_elem: Any, perm_id: int) -> None:
    """在行内插入一对 permStart（cell#0 之后）与 permEnd（行尾），标记可编辑范围。"""
    cells = row_elem.findall(qn("w:tc"))
    if not cells:
        return

    perm_start = row_elem.makeelement(
        qn("w:permStart"),
        {
            qn("w:id"): str(perm_id),
            qn("w:edGrp"): "everyone",
        },
    )
    first_cell = cells[0]
    next_sibling = first_cell.getnext()
    if next_sibling is not None:
        next_sibling.addprevious(perm_start)
    else:
        row_elem.append(perm_start)

    perm_end = row_elem.makeelement(
        qn("w:permEnd"),
        {
            qn("w:id"): str(perm_id),
        },
    )
    row_elem.append(perm_end)


def _set_enforcement(doc: Any) -> None:
    """确保文档保护强制启用（enforcement=1）。"""
    settings_elem = doc.element.body.getparent().find(".//" + qn("w:settings"))
    if settings_elem is None:
        return
    protection = settings_elem.find(qn("w:documentProtection"))
    if protection is not None:
        protection.set(qn("w:enforcement"), "1")


def generate_deviation_ledger_export_docx(items: list[dict[str, Any]]) -> bytes:
    """Render deviation ledger records into the local Word template."""
    doc = Document(str(_TEMPLATE_PATH))
    if not doc.tables:
        raise AppException(message="偏差台账模板中未找到表格")

    table = doc.tables[0]
    table_element = table._tbl
    row_elements = table_element.findall(qn("w:tr"))
    if len(row_elements) < 2:
        raise AppException(message="偏差台账模板缺少可复用的数据行")

    # 模板数据行带自动编号，先整体清理，避免导出内容重复渲染
    _strip_auto_numbering(table_element)

    # 先记录模板中已有的最大 perm id（标题区 0/1 + 数据行 2-8），新 id 从其后递增
    max_perm_id = _find_max_perm_id(doc)

    # 找到第一个数据行作为模板行
    template_row = copy.deepcopy(row_elements[1])

    # 删除所有数据行（保留表头行）
    for row_elem in list(row_elements[1:]):
        table_element.remove(row_elem)

    # 清除模板中所有 permStart/permEnd（标题区 + 表级残留），只保留表头行
    _remove_all_perm_markers(doc)

    for index, item in enumerate(items, start=1):
        row_elem = copy.deepcopy(template_row)

        # 移除拷贝行中的 permStart/permEnd 标记
        _remove_perm_markers(row_elem)

        table_element.append(row_elem)
        cell_elems = row_elem.findall(qn("w:tc"))

        values = [
            str(index),
            _safe_str(item.get("deviation_code")).strip(),
            _build_product_batch(item),
            _safe_str(item.get("description") or item.get("title")).strip(),
            _normalize_yes_no(item.get("has_occurred_before")),
            _safe_str(item.get("root_cause_analysis")).strip(),
            _normalize_level(item.get("level")),
            _format_date(item.get("investigation_completed_at")),
            _safe_str(item.get("corrective_actions")).strip(),
            _safe_str(item.get("material_disposition")).strip(),
            _normalize_yes_no(
                item.get("is_closed")
                if item.get("is_closed") is not None
                else item.get("status")
            ),
        ]

        for cell_index, value in enumerate(values):
            if cell_index >= len(cell_elems):
                break
            _set_cell_text(cell_elems[cell_index], value)

        # 为拷贝行插入新的 permStart/permEnd 标记（唯一 id）
        max_perm_id += 1
        _insert_perm_markers(row_elem, max_perm_id)

    # 确保文档保护启用
    _set_enforcement(doc)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
