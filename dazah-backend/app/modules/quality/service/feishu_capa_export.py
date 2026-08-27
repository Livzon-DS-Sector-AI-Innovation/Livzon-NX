"""Feishu CAPA ledger export — fill data into local template .docx."""

from __future__ import annotations

import copy
import logging
import os
from io import BytesIO
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell

logger = logging.getLogger(__name__)


_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "_capa_export_template.docx")


def _date_dot(value: str | None) -> str:
    """Convert YYYY-MM-DD -> YYYY.MM.DD"""
    if not value:
        return ""
    return value.replace("-", ".")


def _safe_str(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    return str(value)


def _set_cell_text(cell: _Cell, text: str) -> None:
    """Replace text in first paragraph of a cell, preserving formatting."""
    p = cell.paragraphs[0]
    # preserve formatting from the first run if exists
    existing_runs = p.runs
    if existing_runs:
        # keep style, just replace text
        for run in existing_runs:
            run.text = ""
        existing_runs[0].text = text
    else:
        p.add_run(text)


def _set_cell_multiline(cell: _Cell, text: str) -> None:
    """Handle multiline text. Preserve styling of existing first paragraph,
    add extra paragraphs for additional lines."""
    lines = text.split("\n")
    paras = cell.paragraphs

    # first line -> reuse first paragraph, preserve style
    if paras:
        p0 = paras[0]
        for run in p0.runs:
            run.text = ""
        if p0.runs:
            p0.runs[0].text = lines[0]
        else:
            p0.add_run(lines[0])
    else:
        cell.add_paragraph(lines[0])

    # remove extra paragraphs beyond needed
    needed = len(lines)
    while len(cell.paragraphs) > needed:
        p_elem = cell.paragraphs[-1]._element
        p_elem.getparent().remove(p_elem)

    # add missing paragraphs, copying style from first
    template_elem = cell.paragraphs[0]._element if cell.paragraphs else None
    for i in range(len(cell.paragraphs), needed):
        if template_elem is not None:
            new_elem = copy.deepcopy(template_elem)
            # add after last paragraph
            cell.paragraphs[-1]._element.addnext(new_elem)
        else:
            cell.add_paragraph(lines[i])

    # set text
    for i, line in enumerate(lines):
        p = cell.paragraphs[i]
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = line
        else:
            p.add_run(line)


def generate_capa_export_docx(items: list[dict[str, Any]]) -> bytes:
    """将飞书CAPA台账记录填入本地 Word 模板，返回 .docx 字节流。

    基于模板文件 ``_capa_export_template.docx`` 的首张表格，清除表头以外的
    数据行后，复制表头行样式逐条写入 ``items`` 数据。日期字段会被转换为
    ``YYYY.MM.DD`` 格式。

    表格列映射关系（从 0 开始）：

    | 列 | 内容 | 数据来源字段 | 备注 |
    |----|------|-------------|------|
    | 0 | CAPA编号 | ``CAPA编号`` | |
    | 1 | 启动日期 | ``启动日期`` | 转为 ``YYYY.MM.DD`` |
    | 2 | 事件部门 | ``事件部门`` | |
    | 3 | 涉及产品 | ``涉及产品`` | |
    | 4 | （空） | — | 预留列 |
    | 5 | CAPA简述 | ``CAPA简述`` | 支持多行，按换行符拆分到多个段落 |
    | 6 | CAPA效果评估 | ``CAPA效果评估`` | |
    | 7 | 关闭日期 | ``关闭日期`` | 转为 ``YYYY.MM.DD`` |
    | 8 | QA质量员/确认日期 | ``QA质量员`` +
    ``QA质量员确认日期`` | 拼接为 ``姓名+日期`` |

    Args:
        items: 已解析的 CAPA 台账 item 列表，键为中文字段名。

    Returns:
        生成的 .docx 文件字节内容。
    """
    doc = Document(_TEMPLATE_PATH)
    table = doc.tables[0]

    # ── 1. Remove all existing data rows (keep only header row 0) ──
    tbl_element = table._tbl
    row_elements = tbl_element.findall(qn("w:tr"))
    header_elem = row_elements[0] if row_elements else None
    # delete everything after header
    for elem in list(row_elements[1:]):
        tbl_element.remove(elem)

    # ── 2. Copy header row style as template for data rows ──
    def _clone_row(src_elem: Any) -> Any:
        new_elem = copy.deepcopy(src_elem)
        # clear cell texts in cloned row
        for tc in new_elem.findall(qn("w:tc")):
            for p_elem in tc.findall(qn("w:p")):
                for r_elem in p_elem.findall(qn("w:r")):
                    for t_elem in r_elem.findall(qn("w:t")):
                        t_elem.text = ""
        return new_elem

    header_row_elem = copy.deepcopy(header_elem) if header_elem is not None else None

    # ── 3. Add data rows ──
    for item in items:
        # clone header row for styling
        row_elem = (
            copy.deepcopy(header_row_elem)
            if header_row_elem is not None
            else _clone_row(header_elem)
        )
        tbl_element.append(row_elem)

        # now write data into cells
        cells = row_elem.findall(qn("w:tc"))

        capa_code = _safe_str(item.get("CAPA编号"))
        start_date = _date_dot(_safe_str(item.get("启动日期")))
        department = _safe_str(item.get("事件部门"))
        product = _safe_str(item.get("涉及产品"))
        summary = _safe_str(item.get("CAPA简述"))
        evaluation = _safe_str(item.get("CAPA效果评估"))
        close_date = _date_dot(_safe_str(item.get("关闭日期")))
        qa_name = _safe_str(item.get("QA质量员"))
        qa_date = _date_dot(_safe_str(item.get("QA质量员确认日期")))
        qa_combined = f"{qa_name}{qa_date}" if qa_name or qa_date else ""

        col_values = [
            capa_code,
            start_date,
            department,
            product,
            "",
            summary,
            evaluation,
            close_date,
            qa_combined,
        ]

        for ci, val in enumerate(col_values):
            if ci >= len(cells):
                continue
            tc = cells[ci]
            p_elems = tc.findall(qn("w:p"))
            if not p_elems:
                continue

            if ci == 5 and "\n" in val:  # CAPA简述 - multiline
                lines = val.split("\n")
                # first line in first paragraph
                r_elems_0 = p_elems[0].findall(qn("w:r"))
                for r in r_elems_0:
                    for t in r.findall(qn("w:t")):
                        t.text = ""
                if r_elems_0:
                    for t in r_elems_0[0].findall(qn("w:t")):
                        t.text = lines[0]
                else:
                    new_r = p_elems[0].makeelement(qn("w:r"), {})
                    new_t = new_r.makeelement(qn("w:t"), {})
                    new_t.text = lines[0]
                    new_r.append(new_t)
                    p_elems[0].append(new_r)

                # remove extra paragraphs beyond needed
                while len(p_elems) > len(lines):
                    p_elems[-1].getparent().remove(p_elems[-1])

                # add missing paragraphs, clone style from first
                for i in range(len(p_elems), len(lines)):
                    new_p = copy.deepcopy(p_elems[0])
                    for r in new_p.findall(qn("w:r")):
                        for t in r.findall(qn("w:t")):
                            t.text = ""
                    if new_p.findall(qn("w:r")):
                        for t in new_p.findall(qn("w:r"))[0].findall(qn("w:t")):
                            t.text = lines[i]
                    tc.append(new_p)
                for i, line in enumerate(lines):
                    if i < len(p_elems):
                        r_list = p_elems[i].findall(qn("w:r"))
                        if r_list:
                            for t in r_list[0].findall(qn("w:t")):
                                t.text = line
                        else:
                            new_r = p_elems[i].makeelement(qn("w:r"), {})
                            new_t = new_r.makeelement(qn("w:t"), {})
                            new_t.text = line
                            new_r.append(new_t)
                            p_elems[i].append(new_r)
            else:
                # Single line cell
                r_elems = p_elems[0].findall(qn("w:r"))
                for r in r_elems:
                    for t in r.findall(qn("w:t")):
                        t.text = ""
                if r_elems:
                    for t in r_elems[0].findall(qn("w:t")):
                        t.text = val
                else:
                    new_r = p_elems[0].makeelement(qn("w:r"), {})
                    new_t = new_r.makeelement(qn("w:t"), {})
                    new_t.text = val
                    new_r.append(new_t)
                    p_elems[0].append(new_r)

    # ── 4. Save ──
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
