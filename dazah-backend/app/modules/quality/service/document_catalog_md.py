"""Document catalog attachment markdown conversion service.

两段式转换（规则同步自 `公司文件md处理脚本/transform.py`，2026-08-05）：
  ① word（.doc/.docx）→ markdown：按文档顺序提取段落与表格，**保留表格结构**（顶层表格 +
     单元格内嵌套表格递归提取为 markdown 表格语法 `|…|`，不做拍平）；
  ② 对第①步生成的 md 应用转换规则：
     1. 删除「审核及颁发 / Review and
     issue」整段（含批准签名行：起草/审核/批准/颁发等）；
     2. 删除「分发 / Distribution」段的分发清单子表（含「分发-1
     … 部门」条目），保留该段正文；
     3. 其余所有表格（正文表、日期表、记录表单模板）转成带原有序号的独立成行文本大纲，
        保留序号顺序与层级。

图片/PDF 附件不进入本流程（原样存储，直接预览）。
"""

from __future__ import annotations

import io
import logging
import os
import re
import shutil
import subprocess
import tempfile
from typing import Any

from app.core.exceptions import AppException

logger = logging.getLogger(__name__)

# ============ 转换规则（移植自 transform.py） ============

NUM_RE = re.compile(r"^(\d+)(?:\.(\d+))*(?:\.)?$")
HEADING_RE = re.compile(r"^\*\*.+\*\*\s*$")

# Tokens that mark an approval-signature row (content-based, avoids prose like "经批准")
APPROVAL_TOKENS = [
    "起草",
    "颁发",
    "签名",
    "prepared by",
    "reviewed by",
    "approved by",
    "issued by",
    "signature",
]
REVIEW_DIGIT_RE = re.compile(r"审核\d")


def strip_bold(s: str) -> str:
    return s.replace("**", "")


def non_empty(cells: list[str]) -> list[str]:
    return [c.strip() for c in cells if c.strip()]


def parse_row(s: str) -> list[str]:
    s = s.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [p.strip() for p in s.split("|")]


def is_sep_row(cells: list[str]) -> bool:
    if not cells:
        return False
    return all(re.match(r"^:?-+:?$", c) for c in cells)


def is_dash_row(cells: list[str]) -> bool:
    if not any(c.strip() for c in cells):
        return False
    return all(re.match(r"^[\s\-—:]*$", c) for c in cells)


def is_section_marker_to_delete(line: str) -> bool:
    """Detect the two boilerplate section markers no matter how they are formatted
    (bold heading, numbered outline item, with/without full-width slash)."""
    s = line.strip()
    low = s.lower()
    if "审核及颁发" in s or "review and issue" in low:
        return True
    if "分发" in s and "distribution" in low:
        return True
    return False


CJK_RE = re.compile(r"[\u4e00-\u9fff]")


def is_boilerplate_text_line(line: str) -> bool:
    """Catch plain-text fragments of the approval/distribution blocks left behind by
    badly-converted files. Body text (which contains CJK) is never matched."""
    s = line.strip()
    if not s:
        return False
    low = s.lower()
    if "审核及颁发" in s or "review and issue" in low or "and issue" in low:
        return True
    if "分发-" in s or "distribution-" in low:
        return True
    if any(
        t in low
        for t in (
            "prepared by",
            "issued by",
            "approved by",
            "reviewed by",
            "preparedby",
            "issuedby",
            "approvedby",
            "reviewedby",
        )
    ):
        return True
    if re.fullmatch(r"颁发[\s:：]*", s):
        return True
    # standalone English boilerplate fragments only when no Chinese present
    if not CJK_RE.search(s):
        if (
            re.fullmatch(r"[\s,./()\-]*department[\s,./()\-]*", low)
            or re.fullmatch(r"[\s,./()\-]*signature[\s,./()\-]*", low)
            or re.fullmatch(r"[\s,./()\-]*date[\s,./()\-]*", low)
        ):
            return True
    return False


def is_distribution_list(rows_cells: list[list[str]]) -> bool:
    for cells in rows_cells:
        for c in cells:
            cl = c.lower()
            if "分发-" in cl or "distribution-" in cl:
                return True
    return False


def is_approval_row(cells: list[str]) -> bool:
    low = " ".join(cells).lower()
    for tok in APPROVAL_TOKENS:
        if tok in low:
            return True
    if REVIEW_DIGIT_RE.search(low):
        return True
    return False


def is_placeholder(ne: list[str]) -> bool:
    """A signature/blank placeholder row: contains slashes but no CJK and no letters."""
    text = " ".join(ne)
    if "/" not in text:
        return False
    if re.search(r"[\u4e00-\u9fffA-Za-z]", text):
        return False
    return True


def is_num_row(cells: list[str]) -> bool:
    if not cells:
        return False
    return bool(NUM_RE.match(strip_bold(cells[0]).strip().rstrip(".")))


def render_outline(rows_cells: list[list[str]]) -> list[str]:
    out: list[str] = []
    prev_depth: int | None = None
    for cells in rows_cells:
        ne = non_empty(cells)
        if not ne:
            continue
        c0 = ne[0]
        m = NUM_RE.match(strip_bold(c0).strip().rstrip("."))
        if not m:
            out.append(" ".join(ne))
            prev_depth = None
            continue
        num = strip_bold(c0).strip()
        parts = num.rstrip(".").split(".")
        depth = len(parts)
        indent = "  " * (depth - 1)
        rest = " ".join(ne[1:])
        line = f"{indent}{num} {rest}".rstrip()
        # Blank line before EVERY numbered item (any depth) so each number sits on its
        # own line.
        if prev_depth is not None:
            out.append("")
        out.append(line)
        prev_depth = depth
    return out


def render_flat(rows_cells: list[list[str]]) -> list[str]:
    out: list[str] = []
    for cells in rows_cells:
        ne = non_empty(cells)
        if not ne:
            continue
        if is_placeholder(ne):
            continue
        if len(ne) == 2:
            out.append(f"{ne[0]}：{ne[1]}")
        elif len(ne) >= 4 and len(ne) % 2 == 0:
            pairs = [f"{ne[i]}：{ne[i + 1]}" for i in range(0, len(ne), 2)]
            out.append("；".join(pairs))
        else:
            out.append(" ".join(ne))
    return out


def convert_table(raw_rows: list[str]) -> list[str]:
    """raw_rows: list of '|...|' strings. Returns list of text lines."""
    data: list[list[str]] = []
    for r in raw_rows:
        cells = parse_row(r)
        if is_sep_row(cells):
            continue
        if is_dash_row(cells):
            continue
        data.append(cells)
    if not data:
        return []
    # Whole-table distribution list -> drop entirely
    if is_distribution_list(data):
        return []
    # Approval block -> strip approval rows (handles embedded-in-date-table case)
    if any(is_approval_row(c) for c in data):
        data = [c for c in data if not is_approval_row(c)]
    if not data:
        return []
    # split into contiguous segments by is_num_row
    segments: list[tuple[str, list[list[str]]]] = []
    cur_type: str | None = None
    cur: list[list[str]] = []
    for cells in data:
        t = "num" if is_num_row(cells) else "flat"
        if t != cur_type:
            if cur and cur_type is not None:
                segments.append((cur_type, cur))
            cur_type = t
            cur = []
        cur.append(cells)
    if cur and cur_type is not None:
        segments.append((cur_type, cur))
    out: list[str] = []
    for typ, seg in segments:
        if typ == "num":
            out.extend(render_outline(seg))
        else:
            out.extend(render_flat(seg))
    return out


def transform_text(lines: list[str]) -> list[str]:
    out: list[str] = []
    n = len(lines)
    idx = 0

    def collect_table(start: int) -> int:
        k = start
        while k < n:
            s = lines[k].strip()
            if s.startswith("|") and s.endswith("|"):
                k += 1
            else:
                break
        return k

    while idx < n:
        line = lines[idx]
        stripped = line.strip()

        if is_section_marker_to_delete(stripped):
            idx += 1
            continue
        if HEADING_RE.match(stripped):
            out.append(line)
            idx += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            k = collect_table(idx)
            tbl = lines[idx:k]
            out.extend(convert_table(tbl))
            idx = k
            continue

        if is_boilerplate_text_line(line):
            idx += 1
            continue

        out.append(line)
        idx += 1

    text = "\n".join(out)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.split("\n")


# ============ word → markdown（保留表格，含嵌套表格） ============


def _dedup_merged_cells(cells: list[Any]) -> list[Any]:
    """合并单元格在 python-docx 的 row.cells 中会重复出现，做相邻去重。"""
    deduped: list[Any] = []
    prev_tc = None
    for cell in cells:
        tc = cell._tc
        if tc is not prev_tc:
            deduped.append(cell)
            prev_tc = tc
    return deduped


def _cell_text(cell: Any) -> str:
    parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    return "<br>".join(parts) if parts else ""


def _table_to_markdown(table: Any) -> list[str]:
    lines: list[str] = []
    for ri, row in enumerate(table.rows):
        cells = _dedup_merged_cells(row.cells)
        rendered = " | ".join(_cell_text(c) for c in cells)
        lines.append(f"| {rendered} |")
        if ri == 0:
            lines.append("|" + "|".join([" --- "] * len(cells)) + "|")
    return lines


def _collect_nested_tables(table: Any, depth: int = 0) -> list[Any]:
    """递归收集表格内所有嵌套表格（含更深层级）。"""
    if depth > 5:
        return []
    nested: list[Any] = []
    seen: set[int] = set()
    for row in table.rows:
        for cell in _dedup_merged_cells(row.cells):
            for t in cell.tables:
                if id(t._tbl) in seen:
                    continue
                seen.add(id(t._tbl))
                nested.append(t)
                nested.extend(_collect_nested_tables(t, depth + 1))
    return nested


def word_to_markdown(content: bytes) -> str:
    "docx → markdown 文本：按文档顺序输出段落与表格（保留表格，嵌套表格紧随父表输出）。"
    import docx
    from docx.oxml.ns import qn

    document = docx.Document(io.BytesIO(content))
    out: list[str] = []
    body = document.element.body
    para_map = {p._p: p for p in document.paragraphs}
    table_map = {t._tbl: t for t in document.tables}
    for child in body:
        if child.tag == qn("w:p"):
            p = para_map.get(child)
            if p is not None:
                text = p.text.strip()
                if text:
                    out.append(text)
        elif child.tag == qn("w:tbl"):
            t = table_map.get(child)
            if t is not None:
                if out and out[-1] != "":
                    out.append("")
                out.extend(_table_to_markdown(t))
                nested = _collect_nested_tables(t)
                for nt in nested:
                    out.append("")
                    out.extend(_table_to_markdown(nt))
    return "\n".join(out)


# ============ .doc → .docx 转换层 ============

DOC_CONVERTER_BIN = os.environ.get("DOC_CONVERTER_BIN", "")


def _find_soffice() -> str | None:
    if DOC_CONVERTER_BIN and os.path.exists(DOC_CONVERTER_BIN):
        return DOC_CONVERTER_BIN
    for name in ("soffice", "libreoffice"):
        path = shutil.which(name)
        if path:
            return path
    return None


def _convert_doc_via_soffice(content: bytes, file_name: str) -> bytes:
    soffice = _find_soffice()
    if not soffice:
        return b""
    with tempfile.TemporaryDirectory() as tmp_dir:
        src_path = os.path.join(tmp_dir, file_name)
        out_dir = os.path.join(tmp_dir, "out")
        os.makedirs(out_dir, exist_ok=True)
        with open(src_path, "wb") as f:
            f.write(content)
        try:
            result = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    out_dir,
                    src_path,
                ],
                capture_output=True,
                timeout=120,
                check=False,
            )
        except (subprocess.SubprocessError, OSError) as exc:
            logger.warning(
                "soffice convert failed",
                extra={"business_module": "quality", "error": str(exc)},
            )
            return b""
        dst_path = os.path.join(out_dir, os.path.splitext(file_name)[0] + ".docx")
        if result.returncode != 0 or not os.path.exists(dst_path):
            logger.warning(
                "soffice convert non-zero",
                extra={"module": "quality", "returncode": result.returncode},
            )
            return b""
        with open(dst_path, "rb") as f:
            return f.read()


def convert_doc_to_docx(content: bytes, file_name: str) -> bytes:
    """.doc → .docx：优先 LibreOffice headless，其次 Windows Word COM（可选）。"""
    converted = _convert_doc_via_soffice(content, file_name)
    if converted:
        return converted
    # Windows Word COM（需 pywin32，可选）
    try:
        import win32com.client  # type: ignore[import-untyped]

        with tempfile.TemporaryDirectory() as tmp_dir:
            src_path = os.path.join(tmp_dir, file_name)
            dst_path = os.path.join(tmp_dir, os.path.splitext(file_name)[0] + ".docx")
            with open(src_path, "wb") as f:
                f.write(content)
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(src_path)
                doc.SaveAs2(dst_path, FileFormat=16)  # 16 = wdFormatXMLDocument (.docx)
                doc.Close(False)
            finally:
                word.Quit()
            with open(dst_path, "rb") as f:
                return f.read()
    except ImportError:
        pass
    except Exception as exc:  # noqa: BLE001
        logger.warning(
            "Word COM convert failed",
            extra={"business_module": "quality", "error": str(exc)},
        )
    raise AppException(message="当前环境不支持 .doc 文件，请转换为 .docx 后上传")


# ============ 统一入口 ============


def convert_word_attachment(file_name: str, content: bytes) -> str:
    """word（.doc/.docx）→ 标准 MD 文本（两段式：保表格 → transform 规则）。"""
    ext = os.path.splitext(file_name)[1].lower()
    if ext == ".doc":
        content = convert_doc_to_docx(content, file_name)
    md_text = word_to_markdown(content)
    lines = transform_text(md_text.split("\n"))
    return "\n".join(lines)
