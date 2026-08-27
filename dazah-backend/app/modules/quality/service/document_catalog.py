"""Document catalog service: 汇总 Excel/Word 解析与按部门导入。"""

from __future__ import annotations

import io
import logging
import os
import re
from datetime import date, datetime
from typing import Any

import openpyxl  # type: ignore[import-untyped]
from sqlalchemy import select, update
from sqlalchemy.ext.asyncio import AsyncSession

from app.modules.quality.models.document_catalog import (
    DocumentDepartment,
    DocumentEntry,
)

logger = logging.getLogger(__name__)

EFFECTIVE_DATE_PATTERN = re.compile(
    r"^(\d{4})\s*[.\-/年]\s*(\d{1,2})\s*[.\-/月]\s*(\d{1,2})"
)


def _clean_cell(value: str | None) -> str:
    return (value or "").strip()


def parse_effective_date(value: Any) -> tuple[date | None, str | None]:
    """解析生效日期，返回 (日期, 原始文本)。无法解析时日期为 None，保留原始文本。"""
    if value in (None, ""):
        return None, None
    if isinstance(value, datetime):
        return value.date(), None
    if isinstance(value, date):
        return value, None
    text = str(value).strip()
    match = EFFECTIVE_DATE_PATTERN.match(text)
    if match:
        try:
            return (
                date(int(match.group(1)), int(match.group(2)), int(match.group(3))),
                None,
            )
        except ValueError:
            return None, text
    return None, text


def parse_document_catalog_workbook(content: bytes) -> dict[str, list[dict[str, Any]]]:
    """解析《各部门文件目录汇总》工作簿。

    每个 Sheet 对应一个部门，统一表头：序号 / 文件名称 / 文件编码 / 生效日期。
    """
    workbook = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
    result: dict[str, list[dict[str, Any]]] = {}
    for sheet_name in workbook.sheetnames:
        worksheet = workbook[sheet_name]
        entries: list[dict[str, Any]] = []
        for row in worksheet.iter_rows(min_row=2, values_only=True):
            if not any(v not in (None, "") for v in row):
                continue
            seq_raw = row[0] if len(row) > 0 else None
            name_raw = row[1] if len(row) > 1 else None
            code_raw = row[2] if len(row) > 2 else None
            date_raw = row[3] if len(row) > 3 else None

            name = str(name_raw).strip() if name_raw not in (None, "") else ""
            if not name:
                continue

            seq_no: int | None = None
            if isinstance(seq_raw, (int, float)):
                seq_no = int(seq_raw)
            else:
                seq_text = str(seq_raw).strip() if seq_raw is not None else ""
                seq_no = int(seq_text) if seq_text.isdigit() else None

            code = str(code_raw).strip() if code_raw not in (None, "") else None
            effective_date, effective_date_text = parse_effective_date(date_raw)

            entries.append(
                {
                    "seq_no": seq_no,
                    "name": name,
                    "code": code,
                    "effective_date": effective_date,
                    "effective_date_text": effective_date_text,
                }
            )
        if entries:
            result[sheet_name.strip()] = entries
    return result


def _extract_department_from_docx(document: Any, fallback_name: str = "") -> str:
    """从 docx 段落「部门：XXX」解析部门名称；缺失时回退文件名。"""
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        for sep in ("：", ":"):
            if text.startswith("部门") and sep in text:
                name = text.split(sep, 1)[1].strip()
                if name:
                    return str(name)
    if fallback_name:
        base = os.path.splitext(os.path.basename(fallback_name))[0]
        base = re.sub(r"(文件目录|清单)$", "", base).strip("（()）()-_  ")
        return base
    return ""


def parse_document_catalog_docx(
    content: bytes, fallback_name: str = ""
) -> tuple[str, list[dict[str, Any]]]:
    """解析《各部门文件目录》docx，返回 (部门名称, 条目列表)。

    文档结构（已调研确认）：段落「部门：XXX」 + 一个四列表格
    [序号, 文件名称, 文件编码, 生效日期]。
    """
    import docx

    document = docx.Document(io.BytesIO(content))
    department_name = _extract_department_from_docx(document, fallback_name)
    entries: list[dict[str, Any]] = []
    if not document.tables:
        return department_name, entries

    table = document.tables[0]
    for row in table.rows[1:]:
        cells = [_clean_cell(c.text) for c in row.cells]
        if len(cells) < 4:
            continue
        seq_raw, name_raw, code_raw, date_raw = cells[0], cells[1], cells[2], cells[3]
        if not name_raw:
            continue

        seq_no: int | None = None
        if seq_raw.isdigit():
            seq_no = int(seq_raw)

        code = code_raw or None
        effective_date, effective_date_text = parse_effective_date(date_raw)
        entries.append(
            {
                "seq_no": seq_no,
                "name": name_raw,
                "code": code,
                "effective_date": effective_date,
                "effective_date_text": effective_date_text,
            }
        )
    return department_name, entries


async def import_document_catalog(
    db: AsyncSession,
    content: bytes,
    source_file: str,
    filename: str = "",
) -> dict[str, Any]:
    """按部门全量替换导入：支持 xlsx（sheet=部门）与 docx（单部门）两种格式。"""
    lower_name = (filename or source_file).lower()
    if lower_name.endswith((".xlsx", ".xls")):
        sheet_map = parse_document_catalog_workbook(content)
        dept_entries: list[tuple[str, list[dict[str, Any]]]] = list(sheet_map.items())
    else:
        dept_name, entries = parse_document_catalog_docx(
            content, fallback_name=filename
        )
        dept_entries = [(dept_name, entries)] if dept_name and entries else []

    # 序号自动填充：源文件序号缺失时按部门内顺序补 1..N
    for _name, entry_list in dept_entries:
        for index, entry in enumerate(entry_list, start=1):
            if entry.get("seq_no") is None:
                entry["seq_no"] = index

    sheet_results: list[dict[str, Any]] = []
    total_entries = 0

    for index, (dept_name, entries) in enumerate(dept_entries, start=1):
        if not dept_name:
            continue
        result = await db.execute(
            select(DocumentDepartment).where(DocumentDepartment.name == dept_name)
        )
        department = result.scalar_one_or_none()
        if department is None:
            department = DocumentDepartment(name=dept_name, sort_order=index)
            db.add(department)
            await db.flush()
        elif department.is_deleted:
            # 复活已软删除的同名部门，避免触发唯一约束
            department.is_deleted = False
            department.sort_order = index
            await db.flush()

        # 收集旧条目附件（按 code），全量替换时保留已绑定附件，避免重复导入丢失
        old_result = await db.execute(
            select(DocumentEntry).where(
                DocumentEntry.department_id == department.id,
                DocumentEntry.is_deleted.is_(False),
            )
        )
        old_attachments_by_code = {
            entry.code: entry.attachments
            for entry in old_result.scalars().all()
            if entry.code and entry.attachments
        }

        await db.execute(
            update(DocumentEntry)
            .where(
                DocumentEntry.department_id == department.id,
                DocumentEntry.is_deleted.is_(False),
            )
            .values(is_deleted=True)
        )

        for entry in entries:
            code = entry.get("code")
            db.add(
                DocumentEntry(
                    department_id=department.id,
                    source_file=source_file,
                    attachments=(
                        old_attachments_by_code.get(code) or []
                        if isinstance(code, str)
                        else []
                    ),
                    **entry,
                )
            )
        await db.flush()

        sheet_results.append(
            {
                "sheet_name": dept_name,
                "department_id": department.id,
                "imported_count": len(entries),
            }
        )
        total_entries += len(entries)
        logger.info(
            "document catalog sheet imported",
            extra={"component": "quality", "sheet": dept_name, "count": len(entries)},
        )

    return {
        "source_file": source_file,
        "department_count": len(sheet_results),
        "entry_count": total_entries,
        "sheets": sheet_results,
    }
