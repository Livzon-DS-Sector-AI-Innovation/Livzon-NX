"""培训文档文本提取：统一解析 docx/doc/wps/pdf/txt/md 为纯文本。

供 AI 笔试出题使用（上传补充文件后自动解析全文，无需用户粘贴文本）。
"""

import logging
from io import BytesIO

logger = logging.getLogger(__name__)

# 单文件提取文本上限（防止 prompt 超长）
MAX_TEXT_LEN = 60000


class DocumentParseError(Exception):
    """文档解析失败（格式不支持或内容无法读取）。"""


def _extract_txt(content: bytes) -> str:
    """纯文本文件：优先 UTF-8，回退 GBK（Windows 中文 txt 常见编码）。"""
    for enc in ("utf-8", "gbk"):
        try:
            return content.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    raise DocumentParseError("文本文件编码无法识别，请另存为 UTF-8 编码后重试")


def _extract_docx(content: bytes) -> str:
    """docx：逐段 + 表格逐行提取文本。"""
    from docx import Document

    doc = Document(BytesIO(content))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    if not parts:
        raise DocumentParseError("文档内容为空或无法读取")
    return "\n".join(parts)


def _extract_pdf(content: bytes) -> str:
    """pdf：逐页提取文本（pypdf layout 模式保留版面结构）。"""
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text(extraction_mode="layout") or ""
        if text.strip():
            parts.append(text.strip())
    if not parts:
        raise DocumentParseError("PDF 文本提取为空（可能是扫描件图片，请粘贴文本）")
    return "\n\n".join(parts)


def extract_document_text(filename: str, content: bytes) -> str:
    """根据文件扩展名分派解析，返回纯文本。

    Raises:
        DocumentParseError: 解析失败时抛出，附明确处理建议。
    """
    name = (filename or "").lower()
    try:
        if name.endswith((".txt", ".md")):
            text = _extract_txt(content)
        elif name.endswith(".pdf"):
            text = _extract_pdf(content)
        elif name.endswith((".docx", ".doc", ".wps")):
            # .doc/.wps 先按 OOXML 尝试（WPS 保存的常兼容）；旧格式会抛错并给出建议
            try:
                text = _extract_docx(content)
            except DocumentParseError:
                raise
            except Exception as e:
                if name.endswith((".doc", ".wps")):
                    raise DocumentParseError(
                        "该 .doc/.wps 格式解析失败，请用 WPS/Word 另存为 .docx"
                        " 或 .pdf 后重新上传，或粘贴文本"
                    ) from e
                raise DocumentParseError("文档解析失败，请确认文件未损坏") from e
        else:
            raise DocumentParseError(
                "不支持的文件格式，请上传 .docx/.doc/.wps/.pdf/.txt/.md"
            )
    except DocumentParseError:
        raise
    except Exception as e:
        logger.warning("document text extraction failed: %s", e)
        raise DocumentParseError("文档解析失败，请确认文件未损坏或改用粘贴文本") from e

    text = text.strip()
    if not text:
        raise DocumentParseError("文档内容为空")
    return text[:MAX_TEXT_LEN]
