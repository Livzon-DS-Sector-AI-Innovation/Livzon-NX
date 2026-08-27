"""AI 出题：根据文件内容生成题目，并导出试卷 Word 文档."""

import re
import zipfile
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.platform.ai.schemas import (
    ExamExportRequest,
    WrittenExamExportRequest,
)

_PROMPT_TEMPLATE = """你是一位专业的培训考核出题专家。请根
据以下文件内容，生成一份新员工入职培训考核试卷。

要求：
1. 生成 5 道选择题，每题 10 分，共 50 分。每道题有 A、B、C、D 四个选项。
2. 生成 5 道判断题，每题 10 分，共 50 分。
3. 题目必须严格基于文件内容，不能编造文件中没有的信息。
4. 题目难度适中，适合新员工入职培训考核。
5. 每道题都要给出正确答案。

请严格按照以下 JSON 格式返回，不要包含任何其他文字：

{
  "choice_questions": [
    {
      "number": 1,
      "question": "题目内容",
      "options": [
        {"label": "A", "text": "选项内容"},
        {"label": "B", "text": "选项内容"},
        {"label": "C", "text": "选项内容"},
        {"label": "D", "text": "选项内容"}
      ],
      "answer": "A"
    }
  ],
  "true_false_questions": [
    {
      "number": 1,
      "question": "题目内容",
      "answer": "√"
    }
  ]
}

文件内容如下：

{content}
"""


def build_generate_prompt(file_content: str) -> str:
    """构建让 AI 根据文件内容出题的 prompt."""
    return _PROMPT_TEMPLATE.replace("{content}", file_content)


_ORAL_PROMPT_TEMPLATE = """你是一位具有 20
年原料药生产与质量检验经验的资深 GMP 专家，精通 ICH
Q7、药品生产质量管理规范（GMP）及相关法规。现在请根据以下培训文件内容，为口试培训考核生成问答题（口试为考官口头提问、受训人员口头作答）。

要求：
1. 生成填空题（考核问题 + 参考答案），不要生成选择题或判断题。
2. 题目形式：从文件原文中选取关键句子，**用 [______] 挖掉关键参数
、名词、术语或短语**，受训者需填写被挖掉的原文内容；参考答案只写被挖掉的那部分
原文。
3. {count_rule}
4. 出题范围必须限定在文件的"程序/操作步骤/正文条款"等主体内容，聚焦操作要
求、关键参数、质量控制点、管理要求等知识点；禁止从"目的、适用范围、职责"等概述
性章节出题。
5. 题目必须与 GMP 相关，贴合原料药生产与质量检验的实际操作。
6. 题目和参考答案必须严格引用文件原文（直接摘录文件中的原句或关键条款），可从
文件原文组合摘录相关原句；不得自行发挥、概括或补充文件之外的信息。
7. 题目难度适中，覆盖文件程序部分的关键知识点。
8. 所有题目必须互不相同（不得重复）。

请严格按照以下 JSON 格式返回，不要包含任何其他文字：

{
  "questions": [
    {
      "question": "考核问题",
      "answer": "参考答案（必须为文件原文摘录）"
    }
  ]
}

培训文件内容如下：

{content}
"""


def build_oral_generate_prompt(
    files: list[dict[str, Any]], question_count: int | None = None
) -> str:
    """构建口试（问答）出题 prompt：每份文件标注名称+编码，内容分段输出。

    question_count：用户指定的总题数；为空时按默认逻辑（每份文件
    2~3 题，超 5 份文件≤15 题）。
    """
    sections: list[str] = []
    for i, file in enumerate(files, start=1):
        name = file.get("name") or ""
        code = file.get("code") or ""
        content = file.get("content") or ""
        header = f"文件{i}：{name}" + (f"（编号 {code}）" if code else "")
        sections.append(f"{header}\n{content}")
    if question_count:
        count_rule = (
            f"共生成 {question_count} 道题"
            "（在各文件间尽量均匀分布，每份文件至少 1 题）。"
        )
    else:
        count_rule = (
            "每份文件生成 2~3 道问答题；如果培训文件超过 5 份，题目总数不得超过 15 道。"
        )
    return _ORAL_PROMPT_TEMPLATE.replace("{count_rule}", count_rule).replace(
        "{content}", "\n\n---\n\n".join(sections)
    )


# ─── AI 笔试（选择+填空）提示词 ───

_WRITTEN_PROMPT_TEMPLATE = """你是一位具有 20
年原料药生产与质量检验经验的资深 GMP 专家，精通 ICH
Q7、药品生产质量管理规范（GMP）及相关法规。现在请根据以下培训文件内容，生成一份培训考核笔试试卷。

{rules}

出题要求（严格遵守）：
1. 所有题目必须严格基于所提供的文件内容出题，禁止编造文件中不存在的信息。
2. 出题范围必须限定在文件的"程序/操作步骤/正文条款"等主体内容，聚焦操作要
求、关键参数、质量控制点、管理要求等知识点；禁止从"目的、适用范围、职责"等概述
性章节出题。
3. 题目必须与 GMP 相关，贴合原料药生产与质量检验的实际操作。
4. 选择题的干扰项（错误选项）必须合理，与正确答案属于同一知识领域，不得出现明
显荒谬的选项。
5. 填空题的答案必须能从文件原文中直接找到对应表述。
6. 每道题都要给出正确答案，答案必须引用或基于文件原文。
7. 题目难度适中，覆盖文件程序部分的关键知识点。
8. 题目表述直接以题目内容开始，禁止添加"根据文件内容，""根据上述材料，"等任何前缀。

请严格按照以下 JSON 格式返回，不要包含任何其他文字：

{
  "choice_questions": [
    {
      "number": 1,
      "question": "题目内容",
      "options": [
        {"label": "A", "text": "选项内容"},
        {"label": "B", "text": "选项内容"},
        {"label": "C", "text": "选项内容"},
        {"label": "D", "text": "选项内容"}
      ],
      "answer": "A"
    }
  ],
  "true_false_questions": [
    {
      "number": 1,
      "question": "题目内容",
      "answer": "√"
    }
  ],
  "fill_blank_questions": [
    {
      "number": 1,
      "question": "题目内容（填空处用______表示）",
      "answer": "填空答案"
    }
  ]
}

培训文件内容如下：

{content}
"""


def _build_rules(
    single_choice_count: int,
    multiple_choice_count: int,
    fill_blank_count: int,
    true_false_count: int = 0,
) -> str:
    "根据用户选择动态生成题型规则描述（单选/多选/判断/填空分别"
    "指定数量；分值与出题无关，不写入 prompt）。"
    rules: list[str] = []
    if single_choice_count > 0:
        rules.append(
            f"必须生成恰好 {single_choice_count} 道单选题"
            "（每题只有一个正确答案，不得多也不得少）。"
            "每道题有 A、B、C、D 四个选项，answer 字段只写唯一正确选项字母（如 A）。"
        )
    if multiple_choice_count > 0:
        rules.append(
            f"必须生成恰好 {multiple_choice_count} 道多选题"
            "（每题有两个或以上正确答案，不得多也不得少）。"
            "每道题有 A、B、C、D 四个选项，answer 字段写全部正确选项字母（如 ABD）。"
        )
    if true_false_count > 0:
        rules.append(
            f"必须生成恰好 {true_false_count} 道判断题（不得多也不得少）。"
            "每题陈述一个事实，answer 字段写 √（正确）或 ×（错误）。"
        )
    if fill_blank_count > 0:
        rules.append(
            f"必须生成恰好 {fill_blank_count} 道填空题（不得多也不得少）。"
            "填空处用______表示，答案写在 answer 字段中。"
        )
    return "\n".join(rules)


def _calc_question_scores(total: int) -> list[int]:
    """试卷满分固定 100 分，按总题数均分：前余数道题各 base+1 分，其余各 base 分。"""
    if total <= 0:
        return []
    base = 100 // total
    rem = 100 - base * total
    return [base + 1] * rem + [base] * (total - rem)


def _section_score_text(scores: list[int]) -> str:
    """大题分值描述：共 X 分（分值统一时追加"每题 Y 分"）。"""
    if not scores:
        return "共 0 分"
    total = sum(scores)
    if len(set(scores)) == 1:
        return f"共 {total} 分，每题 {scores[0]} 分"
    return f"共 {total} 分"


# 出题素材总字数上限（1M token 上下文模型约可容纳 60 万字，此处 12 万字留足余量）
_TOTAL_CONTENT_MAX = 120000


def build_written_generate_prompt(
    files: list[dict[str, Any]],
    uploaded_content: str,
    manual_content: str,
    single_choice_count: int,
    multiple_choice_count: int,
    fill_blank_count: int,
    true_false_count: int = 0,
    total_content_max: int = _TOTAL_CONTENT_MAX,
) -> str:
    """构建笔试出题 prompt：文件内容 + 上传附件 + 手动粘贴 →
    拼接为出题素材（总量封顶）。

    Args:
        total_content_max: 素材总字符上限，默认 12 万；由调用方从
            HR_EXAM_TOTAL_CONTENT_MAX 配置读取后透传（配置变更即时生效）。
    """
    sections: list[str] = []
    for i, f in enumerate(files, start=1):
        name = f.get("name") or ""
        code = f.get("code") or ""
        content = f.get("content") or ""
        header = f"文件{i}：{name}" + (f"（编号 {code}）" if code else "")
        sections.append(f"{header}\n{content}")
    if uploaded_content:
        sections.append(f"上传附件内容：\n{uploaded_content}")
    if manual_content:
        sections.append(f"补充培训内容：\n{manual_content}")

    # 素材总量封顶：超出时按比例截断每个分段，避免 LLM 上下文爆炸/响应超时
    total_len = sum(len(s) for s in sections)
    if total_len > total_content_max:
        ratio = total_content_max / total_len
        sections = [s[: max(500, int(len(s) * ratio))] for s in sections]

    content = "\n\n---\n\n".join(sections)
    rules = _build_rules(
        single_choice_count, multiple_choice_count, fill_blank_count, true_false_count
    )
    return _WRITTEN_PROMPT_TEMPLATE.replace("{content}", content).replace(
        "{rules}", rules
    )


def _safe_filename(s: str, maxlen: int = 60) -> str:
    """去掉文件名非法字符，截断到 maxlen."""
    s = re.sub(r'[\\/:*?"<>|\r\n\t]+', "", s or "").strip()
    return s[:maxlen] or "未命名"


def _build_question_paper(data: WrittenExamExportRequest) -> BytesIO:
    """生成试卷卷（不含答案）：选择题保留空白括号，判断题保留括号，填空题保留填空线。"""
    doc = Document()

    # 分值分配：满分 100 分按总题数均分
    total_questions = (
        len(data.choice_questions)
        + len(data.true_false_questions)
        + len(data.fill_blank_questions)
    )
    all_scores = _calc_question_scores(total_questions)
    choice_scores = all_scores[: len(data.choice_questions)]
    tf_scores = all_scores[
        len(data.choice_questions) : len(data.choice_questions)
        + len(data.true_false_questions)
    ]
    fill_scores = all_scores[
        len(data.choice_questions) + len(data.true_false_questions) :
    ]

    # ── 试卷标题（黑体 四号 居中 加粗）──
    _add_para(
        doc,
        data.title,
        western="Times New Roman",
        east_asian="黑体",
        size=SIZE_4,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # ── 表头：姓名、部门、分数 ──
    _add_para(
        doc,
        "姓 名：                部 门：                    分 数：        ",
        east_asian="宋体",
        size=SIZE_S4,
    )

    # 落款（出卷人/出卷时间/考核时间）按业务要求不打印在试卷上

    _add_para(doc, "", size=SIZE_S4, space_after=Pt(0))

    section_num = 1
    section_labels = ["一", "二", "三", "四", "五"]

    # ── 选择题 ──
    if data.choice_questions:
        label = section_labels[section_num - 1]
        _add_para(
            doc,
            f"{label}、选择题：（{_section_score_text(choice_scores)}）",
            east_asian="宋体",
            size=SIZE_X3,
            bold=True,
        )

        for choice_question in data.choice_questions:
            stem_para = _add_para(
                doc,
                f"{choice_question.number}. {choice_question.question}（    ）",
                east_asian="宋体",
                size=SIZE_S4,
            )
            _apply_para_format(stem_para, space_after=Pt(0))

            opt_para = doc.add_paragraph()
            _apply_para_format(opt_para, space_after=Pt(0))
            tab_stops = opt_para.paragraph_format.tab_stops
            for i in range(1, len(choice_question.options)):
                tab_stops.add_tab_stop(Cm(i * 4), alignment=WD_ALIGN_PARAGRAPH.LEFT)
            opt_text = "\t".join(
                [f"{opt.label}. {opt.text}" for opt in choice_question.options]
            )
            opt_run = opt_para.add_run(opt_text)
            _set_run_font(opt_run, "Times New Roman", "宋体", SIZE_S4)

        _add_para(doc, "", size=SIZE_S4, space_after=Pt(0))
        section_num += 1

    # ── 判断题 ──
    if data.true_false_questions:
        label = section_labels[section_num - 1]
        _add_para(
            doc,
            f"{label}、判断题：（{_section_score_text(tf_scores)}）",
            east_asian="宋体",
            size=SIZE_X3,
            bold=True,
        )

        for true_false_question in data.true_false_questions:
            _add_para(
                doc,
                f"{true_false_question.number}. {true_false_question.question}（    ）",
                east_asian="宋体",
                size=SIZE_S4,
            )

        _add_para(doc, "", size=SIZE_S4, space_after=Pt(0))
        section_num += 1

    # ── 填空题 ──
    if data.fill_blank_questions:
        label = section_labels[section_num - 1]
        _add_para(
            doc,
            f"{label}、填空题：（{_section_score_text(fill_scores)}）",
            east_asian="宋体",
            size=SIZE_X3,
            bold=True,
        )

        for fill_blank_question in data.fill_blank_questions:
            _add_para(
                doc,
                f"{fill_blank_question.number}. {fill_blank_question.question}",
                east_asian="宋体",
                size=SIZE_S4,
            )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _build_answer_paper(data: WrittenExamExportRequest) -> BytesIO:
    """生成答案卷：逐题列出题号与答案。"""
    doc = Document()

    # 分值分配与试卷卷保持一致
    total_questions = (
        len(data.choice_questions)
        + len(data.true_false_questions)
        + len(data.fill_blank_questions)
    )
    all_scores = _calc_question_scores(total_questions)
    choice_scores = all_scores[: len(data.choice_questions)]
    tf_scores = all_scores[
        len(data.choice_questions) : len(data.choice_questions)
        + len(data.true_false_questions)
    ]
    fill_scores = all_scores[
        len(data.choice_questions) + len(data.true_false_questions) :
    ]

    _add_para(
        doc,
        f"{data.title}（参考答案）",
        western="Times New Roman",
        east_asian="黑体",
        size=SIZE_4,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # 落款（出卷人/出卷时间/考核时间）按业务要求不打印在试卷上

    _add_para(doc, "", size=SIZE_S4, space_after=Pt(0))

    section_num = 1
    section_labels = ["一", "二", "三", "四", "五"]

    if data.choice_questions:
        label = section_labels[section_num - 1]
        _add_para(
            doc,
            f"{label}、选择题答案（{_section_score_text(choice_scores)}）",
            east_asian="宋体",
            size=SIZE_X3,
            bold=True,
        )
        for choice_question in data.choice_questions:
            _add_para(
                doc,
                f"{choice_question.number}. {choice_question.question}",
                east_asian="宋体",
                size=SIZE_S4,
            )
            _add_para(
                doc,
                f"    答案：{choice_question.answer or ''}",
                east_asian="宋体",
                size=SIZE_S4,
                bold=True,
            )
        _add_para(doc, "", size=SIZE_S4, space_after=Pt(0))
        section_num += 1

    if data.true_false_questions:
        label = section_labels[section_num - 1]
        _add_para(
            doc,
            f"{label}、判断题答案（{_section_score_text(tf_scores)}）",
            east_asian="宋体",
            size=SIZE_X3,
            bold=True,
        )
        for true_false_question in data.true_false_questions:
            _add_para(
                doc,
                f"{true_false_question.number}. {true_false_question.question}",
                east_asian="宋体",
                size=SIZE_S4,
            )
            _add_para(
                doc,
                f"    答案：{true_false_question.answer or ''}",
                east_asian="宋体",
                size=SIZE_S4,
                bold=True,
            )
        _add_para(doc, "", size=SIZE_S4, space_after=Pt(0))
        section_num += 1

    if data.fill_blank_questions:
        label = section_labels[section_num - 1]
        _add_para(
            doc,
            f"{label}、填空题答案（{_section_score_text(fill_scores)}）",
            east_asian="宋体",
            size=SIZE_X3,
            bold=True,
        )
        for fill_blank_question in data.fill_blank_questions:
            _add_para(
                doc,
                f"{fill_blank_question.number}. {fill_blank_question.question}",
                east_asian="宋体",
                size=SIZE_S4,
            )
            _add_para(
                doc,
                f"    答案：{fill_blank_question.answer or ''}",
                east_asian="宋体",
                size=SIZE_S4,
                bold=True,
            )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_written_exam_zip(data: WrittenExamExportRequest) -> BytesIO:
    """导出笔试资料 zip：试卷卷（无答案）+ 答案卷，两份 docx 分离。"""
    safe_title = _safe_filename(data.title)
    question_buf = _build_question_paper(data)
    answer_buf = _build_answer_paper(data)

    zip_buf = BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{safe_title}_试卷.docx", question_buf.getvalue())
        zf.writestr(f"{safe_title}_答案.docx", answer_buf.getvalue())
    zip_buf.seek(0)
    return zip_buf


# ─── 字号常量 ───
SIZE_S4 = Pt(12)  # 小四
SIZE_4 = Pt(14)  # 四号
SIZE_X3 = Pt(15)  # 小三号


def _set_run_font(
    run: Any, western: str, east_asian: str, size: Pt, bold: bool = False
) -> None:
    """设置 run 的字体：西文字体 + 中文字体 + 字号 + 加粗."""
    run.font.name = western
    run.font.size = size
    run.font.bold = bold
    r = run._element
    run_properties = r.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()
    run_fonts.set(qn("w:eastAsia"), east_asian)


def _apply_para_format(
    para: Any,
    space_before: Pt = Pt(0),
    space_after: Pt = Pt(0),
    line_spacing: Any = Pt(20),
    alignment: Any = None,
    first_line_indent: Pt = Pt(0),
) -> None:
    """统一设置段落格式."""
    pf = para.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    pf.first_line_indent = first_line_indent
    if alignment is not None:
        para.alignment = alignment


def _add_para(
    doc: Any,
    text: str,
    western: str = "Times New Roman",
    east_asian: str = "宋体",
    size: Pt = SIZE_S4,
    bold: bool = False,
    alignment: Any = None,
    line_spacing: Pt = Pt(20),
    **fmt_kwargs: Any,
) -> Any:
    """添加一个段落并统一设置字体和段落格式。

    默认字体：中文宋体，英文/数字 Times New Roman。
    默认行距：20 磅。
    """
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, western, east_asian, size, bold)
    _apply_para_format(
        para, alignment=alignment, line_spacing=line_spacing, **fmt_kwargs
    )
    return para


def generate_exam_docx(data: ExamExportRequest) -> BytesIO:
    """根据试卷数据生成 Word 文档."""
    doc = Document()

    # ── 试卷标题（黑体 四号 居中 加粗）──
    _add_para(
        doc,
        data.title,
        western="Times New Roman",
        east_asian="黑体",
        size=SIZE_4,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # ── 表头：姓名、部门、分数 ──
    _add_para(
        doc,
        "姓 名：                部 门：                    分 数：        ",
        east_asian="宋体",
        size=SIZE_S4,
    )

    # 落款（出卷人/出卷时间/考核时间）按业务要求不打印在试卷上

    # 表头结束后空一行
    _add_para(doc, "", size=SIZE_S4, space_after=Pt(0))

    # ── 一、选择题 ──
    _add_para(
        doc,
        "一、选择题：（共 50 分，每题 10 分）",
        east_asian="宋体",
        size=SIZE_X3,
        bold=True,
    )

    for choice_question in data.choice_questions:
        answer = choice_question.answer or ""

        # 题干段落
        stem_para = _add_para(
            doc,
            f"{choice_question.number}. {choice_question.question}（{answer}  ）",
            east_asian="宋体",
            size=SIZE_S4,
        )
        # 题干段落后无额外间距，与选项紧密衔接
        _apply_para_format(stem_para, space_after=Pt(0))

        # 选项段落（四个选项横向排列）
        opt_para = doc.add_paragraph()
        _apply_para_format(opt_para, space_after=Pt(0))

        # 设置制表位，使选项均匀分布（约每 4cm 一个）
        tab_stops = opt_para.paragraph_format.tab_stops
        for i in range(1, len(choice_question.options)):
            tab_stops.add_tab_stop(Cm(i * 4), alignment=WD_ALIGN_PARAGRAPH.LEFT)

        # 用制表符连接选项
        opt_text = "\t".join(
            [f"{opt.label}. {opt.text}" for opt in choice_question.options]
        )
        opt_run = opt_para.add_run(opt_text)
        _set_run_font(opt_run, "Times New Roman", "宋体", SIZE_S4)

    # 选择题结束后空一行
    _add_para(doc, "", size=SIZE_S4, space_after=Pt(0))

    # ── 二、判断题 ──
    _add_para(
        doc,
        "二、判断题：（共 50 分，每题 10 分）",
        east_asian="宋体",
        size=SIZE_X3,
        bold=True,
    )

    for true_false_question in data.true_false_questions:
        answer = true_false_question.answer or ""
        _add_para(
            doc,
            (
                f"{true_false_question.number}. {true_false_question.question}"
                f"（{answer}  ）"
            ),
            east_asian="宋体",
            size=SIZE_S4,
        )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
