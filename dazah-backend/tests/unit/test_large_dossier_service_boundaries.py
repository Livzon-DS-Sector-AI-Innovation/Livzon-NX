from __future__ import annotations

from datetime import datetime
from pathlib import Path
from types import SimpleNamespace as _SimpleNamespace
from typing import Any, cast
from unittest.mock import AsyncMock, MagicMock
from uuid import uuid4

import pytest
from docx import Document

from app.modules.dossier_writer import service as dossier_service

SimpleNamespace: Any = _SimpleNamespace


class Dump:
    def __init__(self: Any, **values: Any) -> None:
        self.values = values
        for key, value in values.items():
            setattr(self, key, value)

    def model_dump(self: Any, **_kwargs: Any) -> Any:
        return dict(self.values)


def _service(tmp_path: Any) -> Any:
    service = dossier_service.DossierService.__new__(dossier_service.DossierService)
    service.db = AsyncMock()
    service.repo = AsyncMock()
    service.storage_root = tmp_path
    return service


def _dossier(tmp_path: Any, **overrides: Any) -> Any:
    dossier_id = uuid4()
    values = {
        "id": dossier_id,
        "product_name": "新产品",
        "manufacturer": "新工厂",
        "sterile_type": "无菌",
        "template_original_product_name": "旧产品",
        "template_original_manufacturer": "旧工厂",
        "source_templates_path": str(tmp_path / "source"),
        "working_path": str(tmp_path / "working"),
        "assets_path": str(tmp_path / "assets"),
        "outputs_path": str(tmp_path / "outputs"),
    }
    values.update(overrides)
    return SimpleNamespace(**values)


def _chapter(**overrides: Any) -> Any:
    chapter_id = uuid4()
    values = {  # type: ignore[var-annotated]
        "id": chapter_id,
        "product_dossier_id": uuid4(),
        "parent_id": None,
        "chapter_code": "3.2.S.1",
        "chapter_title": "基本信息",
        "level": 1,
        "sort_order": 0,
        "has_content": True,
        "has_assets": False,
        "assets": [],
        "source_file": None,
        "working_file": None,
    }
    values.update(overrides)
    return SimpleNamespace(**values)


def test_dossier_sort_heading_and_filename_helpers(tmp_path: Any) -> Any:
    service = _service(tmp_path)
    assert dossier_service._chapter_sort_key("") == ()
    assert dossier_service._chapter_sort_key("3.2.S.1.X") == (3, 2, 100, 1, 999)
    assert service._detect_heading_level("Heading 1", "x") == 1
    assert service._detect_heading_level("heading2", "x") == 2
    assert service._detect_heading_level("Heading 3", "x") == 3
    assert service._detect_heading_level("Heading 4", "x") == 4
    assert service._detect_heading_level("Heading Custom", "x") == 2
    assert service._detect_heading_level("Normal", "3.2.1 title") == 3
    assert service._detect_heading_level("Normal", "plain") == 0
    assert service._extract_chapter_code("3.2.1 title") == "3.2.1"
    assert service._extract_chapter_code("S.1 title") == "S.1"
    assert service._extract_chapter_code("plain") is None
    assert service._clean_chapter_title("3.2.S.1 标题") == "1 标题"
    dossier = _dossier(tmp_path)
    assert (
        service._generate_working_filename("旧产品_模板.docx", dossier)
        == "新产品_模板_working.docx"
    )


def test_dossier_storage_and_docx_replacement(tmp_path: Any) -> Any:
    service = _service(tmp_path)
    paths = service._create_storage_dirs("d1")
    assert all(Path(path).is_dir() for path in paths.values())

    source = tmp_path / "replace.docx"
    document = Document()
    document.add_paragraph("旧产品 旧工厂 非无菌")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "旧产品"
    document.save(source)
    dossier = _dossier(tmp_path)
    service._replace_basic_info(source, dossier)
    updated = Document(source)
    assert "新产品 新工厂 无菌" in updated.paragraphs[0].text
    assert updated.tables[0].cell(0, 0).text == "新产品"

    untouched = tmp_path / "untouched.docx"
    Document().save(untouched)
    service._replace_basic_info(
        untouched,
        _dossier(
            tmp_path,
            template_original_product_name=None,
            template_original_manufacturer=None,
            sterile_type="其他",
        ),
    )
    service._delete_storage_dirs("d1")
    assert not (tmp_path / "products" / "d1").exists()
    service._delete_storage_dirs("missing")


@pytest.mark.asyncio
async def test_dossier_create_update_delete_transaction_boundaries(
    tmp_path: Any, monkeypatch: Any
) -> Any:
    service = _service(tmp_path)
    data = cast(Any, Dump)(
        product_name="P",
        manufacturer="M",
        sterile_type="无菌",
        template_original_product_name=None,
        template_original_manufacturer=None,
    )
    service.repo.check_duplicate.return_value = object()
    with pytest.raises(ValueError, match="已存在"):
        await service.create_product_dossier(data)
    service.db.commit.assert_not_awaited()

    created = _dossier(tmp_path)
    service.repo.check_duplicate.return_value = None
    service.repo.create_product_dossier.return_value = created
    monkeypatch.setattr(
        service,
        "_create_storage_dirs",
        lambda _id: {
            "source_templates": "s",
            "working": "w",
            "assets": "a",
            "outputs": "o",
        },
    )
    monkeypatch.setattr(service, "_create_m3_chapters", AsyncMock())
    monkeypatch.setattr(service, "init_chapter_ai_config", AsyncMock())
    execute_result: Any = MagicMock()
    execute_result.scalar_one.return_value = created
    service.db.execute.return_value = execute_result
    assert await service.create_product_dossier(data) is created
    service.db.commit.assert_awaited_once()

    service.repo.get_product_dossier.return_value = None
    assert await service.get_product_dossier(created.id) is None
    service.repo.get_product_dossier.return_value = created
    service.repo.count_chapters.return_value = 3
    assert (await service.get_product_dossier(created.id)).chapter_count == 3

    second = _dossier(tmp_path)
    service.repo.list_product_dossiers.return_value = ([created, second], 2)
    service.repo.count_chapters.side_effect = [2, 4]
    items, total = await service.list_product_dossiers(1, 5)
    assert total == 2
    assert [item.chapter_count for item in items] == [2, 4]

    service.db.commit.reset_mock()
    service.repo.update_product_dossier.return_value = created
    assert (
        await service.update_product_dossier(
            created.id, cast(Any, Dump)(status="active")
        )
        is created
    )
    service.db.commit.assert_awaited_once()
    service.db.commit.reset_mock()
    service.repo.get_product_dossier.return_value = created
    assert (
        await service.update_product_dossier(created.id, cast(Any, Dump)()) is created
    )
    service.db.commit.assert_not_awaited()

    service.repo.delete_product_dossier.return_value = False
    assert await service.delete_product_dossier(created.id) is False
    service.db.commit.assert_not_awaited()
    service.repo.delete_product_dossier.return_value = True
    monkeypatch.setattr(service, "_delete_storage_dirs", MagicMock())
    assert await service.delete_product_dossier(created.id) is True
    service.db.commit.assert_awaited_once()


@pytest.mark.asyncio
async def test_dossier_template_save_existing_and_new(tmp_path: Any) -> Any:
    service = _service(tmp_path)
    dossier = _dossier(tmp_path)
    service.repo.get_product_dossier.return_value = None
    with pytest.raises(ValueError, match="不存在"):
        await service.save_template_file(dossier.id, "a.docx", b"x")

    service.repo.get_product_dossier.return_value = dossier
    existing: Any = SimpleNamespace(file_path="", file_size=0)
    service.repo.get_template_by_filename.return_value = existing
    assert await service.save_template_file(dossier.id, "a.docx", b"abc") is existing
    assert existing.file_size == 3
    service.db.flush.assert_awaited_once()
    assert (Path(dossier.source_templates_path) / "a.docx").read_bytes() == b"abc"

    service.repo.get_template_by_filename.return_value = None
    service.repo.create_template.side_effect = lambda item: item
    created = await service.save_template_file(dossier.id, "b.pdf", b"data")
    assert created.original_filename == "b.pdf"
    assert created.file_size == 4


@pytest.mark.asyncio
async def test_dossier_parse_success_and_failure_transactions(
    tmp_path: Any, monkeypatch: Any
) -> Any:
    service = _service(tmp_path)
    dossier = _dossier(tmp_path)
    service.repo.get_product_dossier.return_value = None
    assert await service.parse_templates(dossier.id) == {
        "success": False,
        "message": "品种资料不存在",
        "error": "NOT_FOUND",
    }

    service.repo.get_product_dossier.return_value = dossier
    service.repo.list_templates.return_value = []
    failed = await service.parse_templates(dossier.id)
    assert failed["success"] is False
    assert "没有上传模板文件" in failed["error"]

    matching: Any = SimpleNamespace(original_filename="3.2.S.1_template.docx")
    unmatched: Any = SimpleNamespace(original_filename="unknown.docx")
    chapter = _chapter(chapter_code="3.2.S.1")
    service.repo.list_templates.return_value = [matching, unmatched]
    service.repo.get_chapter_tree.side_effect = [[], [chapter]]
    monkeypatch.setattr(service, "_create_m3_chapters", AsyncMock())
    monkeypatch.setattr(service, "_create_working_copy_for_chapter", AsyncMock())
    monkeypatch.setattr(
        dossier_service,
        "match_file_to_chapter",
        lambda filename: "3.2.S.1" if filename.startswith("3.") else None,
    )
    success = await service.parse_templates(dossier.id)
    assert success["success"] is True
    assert success["matched_count"] == 1
    service._create_m3_chapters.assert_awaited_once()
    service._create_working_copy_for_chapter.assert_awaited_once()


@pytest.mark.asyncio
async def test_dossier_docx_parse_tree_and_detail(
    tmp_path: Any, monkeypatch: Any
) -> Any:
    service = _service(tmp_path)
    dossier = _dossier(tmp_path)
    Path(dossier.working_path).mkdir(parents=True)
    source = tmp_path / "template.docx"
    doc = Document()
    doc.add_heading("3.2.1 一级", level=1)
    doc.add_paragraph("")
    doc.add_heading("3.2.1.1 二级", level=2)
    doc.save(source)
    template: Any = SimpleNamespace(
        file_path=str(source), original_filename=source.name
    )
    monkeypatch.setattr(service, "_replace_basic_info", MagicMock())
    chapters = await service._parse_docx_template(dossier, template)
    assert [chapter.level for chapter in chapters] == [1, 2]
    assert chapters[0].working_file.endswith("_working.docx")

    root = _chapter(chapter_code="3.2.S.2", sort_order=2)
    child = _chapter(
        parent_id=root.id,
        chapter_code="3.2.S.2.1",
        sort_order=1,
        assets=[object()],
    )
    earlier = _chapter(chapter_code="3.2.S.1", sort_order=1)
    tree = service._build_chapter_tree([root, child, earlier])
    assert [node.chapter_code for node in tree] == ["3.2.S.1", "3.2.S.2"]
    assert tree[1].children[0].asset_count == 1
    service.repo.get_chapter_tree.return_value = [root, child]
    assert len(await service.get_chapter_tree(dossier.id)) == 1

    service.repo.get_chapter.return_value = None
    assert await service.get_chapter_detail(root.id) is None
    asset: Any = SimpleNamespace(
        id=uuid4(),
        original_filename="a.pdf",
        file_type="pdf",
        file_size=3,
        uploaded_at=datetime.now(),
    )
    root.assets = [asset]
    root.product_dossier_id = dossier.id
    service.repo.get_chapter.return_value = root
    detail = await service.get_chapter_detail(root.id)
    assert detail is not None
    assert detail.assets[0].original_filename == "a.pdf"


@pytest.mark.asyncio
async def test_dossier_asset_upload_delete_and_category_boundaries(
    tmp_path: Any, monkeypatch: Any
) -> Any:
    service = _service(tmp_path)
    dossier = _dossier(tmp_path)
    chapter = _chapter(product_dossier_id=dossier.id)
    service.repo.get_chapter.return_value = None
    with pytest.raises(ValueError, match="章节不存在"):
        await service.upload_chapter_asset(chapter.id, "a.pdf", b"x")
    service.repo.get_chapter.return_value = chapter
    service.repo.get_product_dossier.return_value = None
    with pytest.raises(ValueError, match="品种资料不存在"):
        await service.upload_chapter_asset(chapter.id, "a.pdf", b"x")

    service.repo.get_product_dossier.return_value = dossier
    monkeypatch.setattr(service, "_suggest_category", AsyncMock(return_value=uuid4()))
    service.repo.create_asset.side_effect = lambda item: item
    asset = await service.upload_chapter_asset(chapter.id, "a.PDF", b"abc")
    assert asset.file_type == "pdf"
    assert asset.file_size == 3
    service.repo.update_chapter.assert_awaited_with(chapter.id, has_assets=True)

    service.repo.list_assets.return_value = [asset]
    assert await service.list_chapter_assets(chapter.id) == [asset]
    service.repo.get_asset.return_value = None
    assert await service.delete_asset(asset.id) is False
    service.repo.get_asset.return_value = asset
    service.repo.count_assets.return_value = 0
    assert await service.delete_asset(asset.id) is True
    service.repo.update_chapter.assert_awaited_with(chapter.id, has_assets=False)
    assert not Path(asset.file_path).exists()

    category_service = _service(tmp_path)
    result: Any = MagicMock()
    result.scalars.return_value.all.return_value = []
    category_service.db.execute.return_value = result
    assert await category_service._suggest_category("3.2.S.1", "test.pdf") is None
    by_name: Any = SimpleNamespace(id=uuid4(), category_name="稳定性", description=None)
    by_description: Any = SimpleNamespace(
        id=uuid4(), category_name="其他", description="长期 稳定 关键字"
    )
    result.scalars.return_value.all.return_value = [by_name, by_description]
    assert (
        await category_service._suggest_category("3.2.S.1", "稳定性数据.pdf")
        == by_name.id
    )
    assert (
        await category_service._suggest_category("3.2.S.1", "长期关键字报告.pdf")
        == by_description.id
    )


@pytest.mark.asyncio
async def test_dossier_export_preview_and_matching_paths(
    tmp_path: Any, monkeypatch: Any
) -> Any:
    service = _service(tmp_path)
    dossier = _dossier(tmp_path)
    chapter = _chapter(product_dossier_id=dossier.id, working_file="source.docx")
    working = Path(dossier.working_path)
    working.mkdir(parents=True)
    source = working / "source.docx"
    doc = Document()
    doc.add_paragraph("正文")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    doc.save(source)  # type: ignore[arg-type]

    service.repo.get_product_dossier.return_value = None
    assert (await service.export_dossier(dossier.id))["success"] is False
    service.repo.get_product_dossier.return_value = dossier
    service.repo.get_chapter_tree.return_value = [chapter]
    one = await service.export_dossier(dossier.id, [chapter.id])
    assert one["success"] is True
    assert "3.2.S.1" in one["filename"]

    second = _chapter(
        product_dossier_id=dossier.id,
        chapter_code="3.2.S.2",
        working_file="source.docx",
    )
    service.repo.get_chapter_tree.return_value = [chapter, second]
    many = await service.export_dossier(dossier.id, [chapter.id, second.id])
    assert many["success"] is True
    assert "3.2.S.1-3.2.S.2" in many["filename"]

    service.repo.get_chapter.return_value = None
    assert (await service.get_chapter_preview(chapter.id))["success"] is False
    service.repo.get_chapter.return_value = _chapter(working_file=None)
    assert "无工作副本" in (await service.get_chapter_preview(chapter.id))["message"]
    service.repo.get_chapter.return_value = chapter
    service.repo.get_product_dossier.return_value = None
    assert "品种不存在" in (await service.get_chapter_preview(chapter.id))["message"]
    service.repo.get_product_dossier.return_value = dossier
    preview = await service.get_chapter_preview(chapter.id)
    assert preview["success"] is True
    assert preview["paragraphs"][0]["text"] == "正文"
    assert preview["tables"] == [[["A", "B"]]]

    service.repo.get_product_dossier.return_value = None
    missing = await service.match_assets_to_chapters(dossier.id)
    assert missing["matched_count"] == 0
    service.repo.get_product_dossier.return_value = dossier
    service.repo.list_templates.return_value = []
    assert (
        "无模板文件" in (await service.match_assets_to_chapters(dossier.id))["message"]
    )
    service.repo.list_templates.return_value = [
        SimpleNamespace(original_filename="match.docx"),
        SimpleNamespace(original_filename="unknown.docx"),
    ]
    service.repo.get_chapter_tree.return_value = [chapter]
    monkeypatch.setattr(
        dossier_service,
        "match_file_to_chapter",
        lambda filename: chapter.chapter_code if filename == "match.docx" else None,
    )
    monkeypatch.setattr(service, "_create_working_copy_for_chapter", AsyncMock())
    matched = await service.match_assets_to_chapters(dossier.id)
    assert matched["matched_count"] == 1
    assert matched["unmatched_files"] == ["unknown.docx"]


@pytest.mark.asyncio
async def test_dossier_working_copy_and_m3_creation(
    tmp_path: Any, monkeypatch: Any
) -> Any:
    service = _service(tmp_path)
    dossier = _dossier(tmp_path)
    source = tmp_path / "source.docx"
    Document().save(source)
    template: Any = SimpleNamespace(
        file_path=str(source), original_filename="source.docx"
    )
    chapter = _chapter()
    monkeypatch.setattr(service, "_replace_basic_info", MagicMock())
    await service._create_working_copy_for_chapter(dossier, template, chapter)
    service.repo.update_chapter.assert_awaited_once()
    assert (Path(dossier.working_path) / "3_2_S_1_source.docx").exists()

    created_ids: dict[Any, Any] = {}

    async def create_chapter(item: Any) -> Any:
        item.id = uuid4()
        created_ids[item.chapter_code] = item.id
        return item

    service.repo.create_chapter.side_effect = create_chapter
    await service._create_m3_chapters(dossier.id)
    assert len(created_ids) == len(dossier_service.M3_CHAPTERS)  # type: ignore[attr-defined]
