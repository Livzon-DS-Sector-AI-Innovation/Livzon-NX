from pathlib import Path
from types import SimpleNamespace as _SimpleNamespace
from typing import Any
from unittest.mock import AsyncMock
from uuid import uuid4

import pytest
from docx import Document

from app.modules.dossier_writer.ai_fill_service import AIFillService

SimpleNamespace: Any = _SimpleNamespace


def make_service() -> AIFillService:
    return AIFillService(SimpleNamespace())


def test_asset_category_filtering_prefers_exact_matches_and_supports_legacy() -> None:
    service = make_service()
    quality: Any = SimpleNamespace(_category_name="质量标准")
    process: Any = SimpleNamespace(_category_name="工艺资料")
    uncategorized: Any = SimpleNamespace()

    assets = [quality, process, uncategorized]
    assert service._filter_assets_by_category(assets, "_default") == assets
    assert service._filter_assets_by_category(assets, "质量标准") == [quality]
    assert service._filter_assets_by_category(assets, "不存在") == []
    assert service._filter_assets_by_category(
        [SimpleNamespace(), SimpleNamespace()],
        "质量标准",
    )


def test_docx_resolution_and_table_extraction(tmp_path: Path) -> None:
    service = make_service()
    source = tmp_path / "standard.docx"
    document = Document()
    table = document.add_table(rows=5, cols=3)
    table.rows[0].cells[0].text = "标题"
    table.rows[1].cells[0].text = "检验项目"
    table.rows[1].cells[1].text = "企业内控标准"
    table.rows[2].cells[0].text = "性状"
    table.rows[2].cells[1].text = "白色粉末"
    table.rows[3].cells[0].text = ""
    table.rows[4].cells[0].text = "备注：内部使用"
    document.save(source)  # type: ignore[arg-type]

    assert service._resolve_docx_path(source) == source
    assert service._resolve_docx_path(tmp_path / "missing.docx") is None

    old_doc = tmp_path / "standard.doc"
    old_doc.write_bytes(b"legacy")
    assert service._resolve_docx_path(old_doc) == source
    assert service._resolve_docx_path(tmp_path / "unsupported.pdf") is None

    extracted = service._extract_table_from_asset(
        SimpleNamespace(file_path=str(source))
    )
    assert extracted == [["性状", "白色粉末", ""]]
    assert (
        service._extract_table_from_asset(
            SimpleNamespace(file_path=str(tmp_path / "missing.docx"))
        )
        is None
    )


def test_paragraph_and_table_fill_actions_cover_primary_and_fallback_paths() -> None:
    service = make_service()
    document = Document()
    document.add_paragraph("产品名称：旧名称")
    document.add_paragraph("生产企业:旧企业")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "规格"
    table.cell(0, 1).text = "旧规格"
    table.cell(1, 0).text = "包装"
    table.cell(1, 1).text = "旧包装"

    assert service._fill_paragraph_replace(
        document,
        {"paragraph_index": 0},
        "新产品",
    )
    assert document.paragraphs[0].text == "产品名称：新产品"
    assert service._fill_paragraph_replace(
        document,
        {"paragraph_index": 99, "keyword": "生产企业"},
        "新企业",
    )
    assert document.paragraphs[1].text == "生产企业:新企业"
    assert not service._fill_paragraph_replace(
        document,
        {"keyword": "不存在"},
        "值",
    )

    assert service._fill_table_cell(
        document,
        {"table_index": 0, "keyword": "规格"},
        "10 mg",
    )
    assert table.cell(0, 1).text == "10 mg"
    assert service._fill_table_cell(
        document,
        {"table_index": 99, "keyword": "包装"},
        "铝塑",
    )
    assert table.cell(1, 1).text == "铝塑"
    assert not service._fill_table_cell(
        document,
        {"table_index": 99, "keyword": "不存在"},
        "值",
    )

    assert service._execute_fill(
        document,
        {
            "fill_action": "replace_after_colon",
            "target": {"paragraph_index": 0},
        },
        "再次更新",
    )
    assert service._execute_fill(
        document,
        {
            "fill_action": "fill_table_cell",
            "target": {"table_index": 0, "keyword": "规格"},
        },
        "20 mg",
    )
    assert not service._execute_fill(
        document,
        {"fill_action": "unsupported", "target": {}},
        "值",
    )


def test_table_row_replacement_and_fallback_fill() -> None:
    service = make_service()
    document = Document()
    document.add_paragraph("批准人：旧姓名")
    document.add_table(rows=1, cols=1)
    data_table = document.add_table(rows=4, cols=3)
    data_table.rows[0].cells[0].text = "检验项目"
    data_table.rows[1].cells[0].text = "序号"
    data_table.rows[1].cells[1].text = "项目"
    data_table.rows[1].cells[2].text = "标准"
    data_table.rows[2].cells[0].text = "1"
    data_table.rows[2].cells[1].text = "旧项目"
    data_table.rows[2].cells[2].text = "旧标准"
    data_table.rows[3].cells[0].text = "备注"

    assert service._fill_table_rows(
        document,
        {"table_index": 1, "header_rows": 2},
        [["1", "性状", "白色粉末"], ["2", "含量", "98.0%～102.0%"]],
    )
    assert any("性状" in cell.text for row in data_table.rows for cell in row.cells)
    assert any("含量" in cell.text for row in data_table.rows for cell in row.cells)
    assert not service._fill_table_rows(
        document,
        {"table_index": 99},
        [["x"]],
    )
    assert not service._fill_table_rows(
        document,
        {"table_index": 1},
        [],
    )

    assert service._fallback_fill(
        document,
        "批准人",
        "张三",
        "text",
    )
    assert document.paragraphs[0].text == "批准人：张三"
    assert service._fallback_fill(
        document,
        "项目",
        "更新值",
        "text",
    )
    assert not service._fallback_fill(
        document,
        "完全不存在",
        "值",
        "text",
    )


@pytest.mark.anyio
async def test_preview_extraction_combines_fixed_table_ai_missing_and_image_fields(
    monkeypatch: Any,
) -> None:
    service = make_service()
    mappings = [
        SimpleNamespace(
            id=uuid4(),
            source_type="fixed",
            field_name="申报类型",
            field_type="text",
            fixed_value="化学药品",
            source_category=None,
            extraction_prompt=None,
            appendix_slot=None,
        ),
        SimpleNamespace(
            id=uuid4(),
            source_type="asset_extract",
            field_name="产品名称",
            field_type="text",
            fixed_value=None,
            source_category="质量标准",
            extraction_prompt="提取产品名称",
            appendix_slot=None,
        ),
        SimpleNamespace(
            id=uuid4(),
            source_type="asset_extract",
            field_name="检验项目",
            field_type="table",
            fixed_value=None,
            source_category="质量标准",
            extraction_prompt=None,
            appendix_slot=None,
        ),
        SimpleNamespace(
            id=uuid4(),
            source_type="asset_extract",
            field_name="缺失字段",
            field_type="text",
            fixed_value=None,
            source_category="不存在分类",
            extraction_prompt=None,
            appendix_slot=None,
        ),
        SimpleNamespace(
            id=uuid4(),
            source_type="asset_image",
            field_name="工艺流程图",
            field_type="image_appendix",
            fixed_value=None,
            source_category="工艺资料",
            extraction_prompt=None,
            appendix_slot="附录1",
        ),
    ]
    quality_asset: Any = SimpleNamespace(
        _category_name="质量标准",
        file_path="/tmp/standard.docx",
        original_filename="standard.docx",
    )
    service.get_field_mappings = AsyncMock(return_value=mappings)  # type: ignore[method-assign]
    service.get_chapter_assets = AsyncMock(return_value=[quality_asset])  # type: ignore[method-assign]
    service.extractor = SimpleNamespace(
        extract=lambda path: {"text": "产品名称：阿莫西林"}
    )
    service.llm = SimpleNamespace(
        chat_json=AsyncMock(
            return_value={
                "fields": [
                    {
                        "field_name": "产品名称",
                        "value": "阿莫西林",
                        "confidence": 0.98,
                        "source": "质量标准",
                    }
                ]
            }
        )
    )
    monkeypatch.setattr(
        service,
        "_extract_table_from_asset",
        lambda asset: [["性状", "白色粉末"]],
    )

    result = await service.preview_extraction(
        SimpleNamespace(product_name="阿莫西林"),
        SimpleNamespace(id=uuid4(), chapter_code="3.2.S.1"),
    )

    assert result["success"] is True
    by_name = {field["field_name"]: field for field in result["fields"]}
    assert by_name["申报类型"]["value"] == "化学药品"
    assert by_name["产品名称"]["value"] == "阿莫西林"
    assert by_name["检验项目"]["value"] == [["性状", "白色粉末"]]
    assert by_name["缺失字段"]["value"] is None
    assert by_name["工艺流程图"]["value"] == "附录1"

    service.get_field_mappings.return_value = []
    no_mapping = await service.preview_extraction(
        SimpleNamespace(product_name="阿莫西林"),
        SimpleNamespace(id=uuid4(), chapter_code="empty"),
    )
    assert no_mapping["success"] is False

    service.get_field_mappings.return_value = mappings
    service.get_chapter_assets.return_value = []
    no_assets = await service.preview_extraction(
        SimpleNamespace(product_name="阿莫西林"),
        SimpleNamespace(id=uuid4(), chapter_code="3.2.S.1"),
    )
    assert no_assets["message"] == "请先上传素材"


@pytest.mark.anyio
async def test_confirm_and_fill_updates_document_and_records_results(
    tmp_path: Path,
) -> None:
    working_file = "chapter.docx"
    document = Document()
    document.add_paragraph("产品名称：旧名称")
    document.save(tmp_path / working_file)  # type: ignore[arg-type]

    db: Any = SimpleNamespace(add=lambda value: None, commit=AsyncMock())
    service = AIFillService(db)
    service.llm = SimpleNamespace(
        chat_json=AsyncMock(
            return_value={
                "fills": [
                    {
                        "field_name": "产品名称",
                        "fill_action": "replace_after_colon",
                        "target": {"paragraph_index": 0},
                    }
                ]
            }
        )
    )
    service.get_chapter_assets = AsyncMock(return_value=[])  # type: ignore[method-assign]
    service.get_field_mappings = AsyncMock(return_value=[])  # type: ignore[method-assign]
    dossier: Any = SimpleNamespace(id=uuid4(), working_path=str(tmp_path))
    chapter: Any = SimpleNamespace(
        id=uuid4(),
        working_file=working_file,
        chapter_code="3.2.S.1",
    )

    result = await service.confirm_and_fill(
        dossier,
        chapter,
        [
            {"field_name": "产品名称", "field_type": "text", "value": "新产品"},
            {"field_name": "空字段", "field_type": "text", "value": None},
            {
                "field_name": "工艺流程图",
                "field_type": "image_appendix",
                "value": "已插入：第1页",
            },
        ],
    )

    assert result["success"] is True
    assert result["message"] == "填充完成: 2/3 个字段"
    saved = Document(tmp_path / working_file)  # type: ignore[arg-type]
    assert saved.paragraphs[0].text == "产品名称：新产品"
    db.commit.assert_awaited_once()

    missing = await service.confirm_and_fill(
        SimpleNamespace(id=uuid4(), working_path=str(tmp_path)),
        SimpleNamespace(
            id=uuid4(),
            working_file="missing.docx",
            chapter_code="3.2.S.1",
        ),
        [],
    )
    assert missing["success"] is False


@pytest.mark.anyio
async def test_preview_page_splits_persists_ai_page_classification() -> None:
    added: list[Any] = []
    db: Any = SimpleNamespace(add=added.append, commit=AsyncMock())
    service = AIFillService(db)
    service.extractor = SimpleNamespace(
        extract=lambda path: {
            "page_texts": [
                {"page": 1, "text": "生产许可证"},
                {"page": 2, "text": "检验报告"},
            ],
            "page_count": 2,
        }
    )
    service.llm = SimpleNamespace(
        chat_json=AsyncMock(
            return_value={
                "pages": [
                    {
                        "page_number": 1,
                        "page_type": "许可证",
                        "content_summary": "生产许可",
                        "appendix_slot": "附录1",
                    },
                    {
                        "page_number": 2,
                        "page_type": "检验报告",
                        "content_summary": "批次检验",
                        "appendix_slot": "附录2",
                    },
                ]
            }
        )
    )

    result = await service.preview_page_splits(
        SimpleNamespace(id=uuid4(), file_path="/tmp/material.pdf"),
        ["附录1", "附录2"],
    )

    assert result["success"] is True
    assert result["page_count"] == 2
    assert len(added) == 2
    assert added[0].ocr_text == "生产许可证"
    db.commit.assert_awaited_once()

    not_pdf = await service.preview_page_splits(
        SimpleNamespace(id=uuid4(), file_path="/tmp/material.docx"),
        [],
    )
    assert not_pdf["success"] is False


@pytest.mark.anyio
async def test_database_query_helpers_shape_categories_mappings_and_assets() -> None:
    category: Any = SimpleNamespace(
        id=uuid4(),
        category_name="质量标准",
        category_type="document",
        appendix_slot="附录1",
        description="质量标准资料",
        sort_order=1,
    )
    mapping: Any = SimpleNamespace(id=uuid4(), field_name="产品名称")
    asset: Any = SimpleNamespace(id=uuid4())
    results = [
        SimpleNamespace(
            scalars=lambda: SimpleNamespace(all=lambda: [category]),
        ),
        SimpleNamespace(
            scalars=lambda: SimpleNamespace(all=lambda: [mapping]),
        ),
        SimpleNamespace(all=lambda: [(asset, "质量标准")]),
    ]
    db: Any = SimpleNamespace(execute=AsyncMock(side_effect=results))
    service = AIFillService(db)

    categories = await service.get_asset_categories("3.2.S.1")
    mappings = await service.get_field_mappings("3.2.S.1")
    assets = await service.get_chapter_assets(uuid4())

    assert categories == [
        {
            "id": str(category.id),
            "category_name": "质量标准",
            "category_type": "document",
            "appendix_slot": "附录1",
            "description": "质量标准资料",
            "sort_order": 1,
        }
    ]
    assert mappings == [mapping]
    assert assets == [asset]
    assert asset._category_name == "质量标准"


@pytest.mark.anyio
async def test_confirm_page_splits_inserts_images_and_updates_split_records(
    tmp_path: Path,
    monkeypatch: Any,
) -> None:
    working_file = "chapter.docx"
    Document().save(tmp_path / working_file)  # type: ignore[arg-type]
    asset_id = uuid4()
    split_id = uuid4()
    asset: Any = SimpleNamespace(id=asset_id, file_path="/tmp/source.pdf")
    split_record: Any = SimpleNamespace(
        appendix_slot=None,
        image_path=None,
        status="pending",
    )

    async def fake_get(model: Any, object_id: Any) -> Any:
        if object_id == asset_id:
            return asset
        if object_id == split_id:
            return split_record
        return None

    db: Any = SimpleNamespace(get=fake_get, commit=AsyncMock())
    service = AIFillService(db)
    service.extractor = SimpleNamespace(
        pdf_page_to_image=lambda path, page: Path("/tmp/page-1.png")
    )
    monkeypatch.setattr(
        service,
        "_insert_image_at_appendix",
        lambda document, slot, image_path: slot == "附录1",
    )

    result = await service.confirm_page_splits_and_insert(
        SimpleNamespace(working_path=str(tmp_path)),
        SimpleNamespace(working_file=working_file),
        [
            {
                "split_id": str(split_id),
                "appendix_slot": "附录1",
                "asset_id": str(asset_id),
                "page_number": 1,
            },
            {"appendix_slot": None, "asset_id": str(asset_id), "page_number": 2},
            {
                "split_id": "not-a-uuid",
                "appendix_slot": "附录2",
                "asset_id": str(asset_id),
                "page_number": 3,
            },
        ],
    )

    assert result["inserted_count"] == 1
    assert split_record.appendix_slot == "附录1"
    assert split_record.status == "inserted"
    assert split_record.image_path == "/tmp/page-1.png"
    db.commit.assert_awaited_once()

    missing = await service.confirm_page_splits_and_insert(
        SimpleNamespace(working_path=str(tmp_path)),
        SimpleNamespace(working_file="missing.docx"),
        [],
    )
    assert missing["success"] is False


@pytest.mark.anyio
async def test_auto_image_insert_rejects_missing_or_unmatched_assets(
    tmp_path: Path,
) -> None:
    missing_asset: Any = SimpleNamespace(
        _category_name="工艺资料",
        original_filename="missing.pdf",
        file_path=str(tmp_path / "missing.pdf"),
    )
    mapping: Any = SimpleNamespace(source_category="工艺资料")
    db: Any = SimpleNamespace(
        get=AsyncMock(return_value=mapping),
        execute=AsyncMock(
            return_value=SimpleNamespace(
                scalars=lambda: SimpleNamespace(all=lambda: []),
            )
        ),
    )
    service = AIFillService(db)
    chapter: Any = SimpleNamespace(chapter_code="3.2.S.2")

    assert not await service._auto_insert_image(
        Document(),
        "工艺流程图片",
        {"field_mapping_id": str(uuid4()), "value": "附录1"},
        chapter,
        [missing_asset],
    )

    db.get.return_value = None
    assert not await service._auto_insert_image(
        Document(),
        "工艺流程图片",
        {"value": "附录1"},
        chapter,
        [],
    )
