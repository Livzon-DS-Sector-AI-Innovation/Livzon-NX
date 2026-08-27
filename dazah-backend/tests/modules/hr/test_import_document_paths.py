from __future__ import annotations

from io import BytesIO
from types import SimpleNamespace
from unittest.mock import AsyncMock
from uuid import uuid4

import pytest
from docx import Document
from fastapi import UploadFile
from openpyxl import Workbook

from app.core.exceptions import AppException
from app.modules.hr import position_training_api, trainer_api


def _position_training_doc() -> bytes:
    document = Document()
    document.add_paragraph("部门：质量部          岗位：质量员")
    table = document.add_table(rows=6, cols=5)
    table.cell(0, 0).text = "序号"
    table.cell(0, 1).text = "教材名称"
    table.cell(0, 2).text = "教材编号"
    table.cell(0, 3).text = "考核方式"
    table.cell(0, 4).text = "备注"
    table.cell(1, 0).text = "部门级"
    table.cell(2, 0).text = "1"
    table.cell(2, 1).text = "GMP"
    table.cell(2, 2).text = "SOP-1"
    table.cell(2, 3).text = "考试"
    table.cell(2, 4).text = "必修"
    table.cell(3, 0).text = "岗位级"
    table.cell(4, 0).text = "2"
    table.cell(4, 1).text = "偏差处理"
    table.cell(4, 2).text = "SOP-2"
    table.cell(4, 3).text = "实操"
    table.cell(4, 4).text = ""
    table.cell(5, 0).text = "……"
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _trainer_doc() -> bytes:
    document = Document()
    document.add_paragraph("说明")
    table = document.add_table(rows=3, cols=6)
    for index, value in enumerate(
        ["姓名", "部门", "岗位", "批准时间", "批准人", "备注"]
    ):
        table.cell(0, index).text = value
    values = ["张三", "质量部", "质量员", "2026.08.26", "李四", "内部"]
    for index, value in enumerate(values):
        table.cell(1, index).text = value
    table.cell(2, 0).text = ""
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_position_training_and_trainer_document_parsers() -> None:
    parsed = position_training_api._parse_import_docx(_position_training_doc())
    assert parsed["department"] == "质量部"
    assert parsed["position"] == "质量员"
    assert [item["level"] for item in parsed["items"]] == ["部门级", "岗位级"]

    assert trainer_api._parse_date_value("2026-08-26").isoformat() == "2026-08-26"
    assert trainer_api._parse_date_value("2026年8月26日").isoformat() == "2026-08-26"
    assert trainer_api._parse_date_value("") is None
    assert trainer_api._parse_date_value("bad") is None
    rows = trainer_api._parse_docx_trainers(_trainer_doc())
    assert rows[0]["name"] == "张三"
    assert rows[0]["approval_date"].isoformat() == "2026-08-26"

    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["培训师姓名", "部门", "职位", "批准日期", "批准人", "备注"])
    sheet.append(["王五", "生产部", "主管", "2026/8/26", "赵六", ""])
    sheet.append([None, None, None, None, None, None])
    stream = BytesIO()
    workbook.save(stream)
    excel_rows = trainer_api._parse_excel_trainers(stream.getvalue())
    assert excel_rows[0]["name"] == "王五"
    assert trainer_api._match_trainer_header(["姓名", "部门"])["name"] == 0
    empty = Document()
    empty_stream = BytesIO()
    empty.save(empty_stream)
    with pytest.raises(AppException):
        trainer_api._parse_docx_trainers(empty_stream.getvalue())


@pytest.mark.asyncio
async def test_position_training_import_creates_and_appends_lists(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    current_user = SimpleNamespace(id=uuid4(), name="测试用户")
    monkeypatch.setattr(
        position_training_api,
        "read_upload_secure",
        AsyncMock(return_value=("岗位培训清单.docx", _position_training_doc())),
    )
    monkeypatch.setattr(
        "app.modules.hr.training_dept_resolver.resolve_training_department",
        AsyncMock(side_effect=lambda _db, name: name),
    )
    created = SimpleNamespace(id=uuid4())
    service = SimpleNamespace(
        repo=SimpleNamespace(list_lists=AsyncMock(return_value=([], 0))),
        create=AsyncMock(return_value=created),
        batch_update_items=AsyncMock(),
    )
    monkeypatch.setattr(
        position_training_api,
        "PositionTrainingListService",
        lambda _db: service,
    )

    result = await position_training_api.import_position_training_list(
        UploadFile(filename="岗位培训清单.docx", file=BytesIO(b"ignored")),
        SimpleNamespace(),
        current_user,
    )
    import json

    first_result = json.loads(result.body)
    assert first_result["data"]["matched"] is False
    service.create.assert_awaited_once()

    existing = SimpleNamespace(
        id=uuid4(),
        position="质量员",
        items=[
            SimpleNamespace(
                level="部门级",
                textbook_name="GMP",
                textbook_code="old",
                assessment_method="考试",
                remarks="",
                is_deleted=False,
            )
        ],
    )
    service.repo.list_lists.return_value = ([existing], 1)
    result = await position_training_api.import_position_training_list(
        UploadFile(filename="岗位培训清单.docx", file=BytesIO(b"ignored")),
        SimpleNamespace(),
        current_user,
    )
    second_result = json.loads(result.body)
    assert second_result["data"]["matched"] is True
    assert second_result["data"]["imported"] == 1
    service.batch_update_items.assert_awaited_once()

    with pytest.raises(AppException, match="暂不支持"):
        monkeypatch.setattr(
            position_training_api,
            "read_upload_secure",
            AsyncMock(return_value=("岗位培训清单.doc", b"ignored")),
        )
        await position_training_api.import_position_training_list(
            UploadFile(filename="岗位培训清单.doc", file=BytesIO(b"ignored")),
            SimpleNamespace(),
            current_user,
        )
