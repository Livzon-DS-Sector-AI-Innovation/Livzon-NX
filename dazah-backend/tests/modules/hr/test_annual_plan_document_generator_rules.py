from __future__ import annotations

from io import BytesIO
from types import SimpleNamespace as _SimpleNamespace
from typing import Any

import pytest
from docx import Document
from docx.enum.text import WD_UNDERLINE

from app.modules.hr import annual_plan_document_generator as generator

SimpleNamespace: Any = _SimpleNamespace


def _item(index: int, *, deleted: bool = False) -> SimpleNamespace:
    return SimpleNamespace(
        training_type="内训" if index % 2 else "外训",
        training_month=f"{index}月",
        content_textbook=f"课程{index}",
        target_audience_new="全员",
        instructor=f"讲师{index}",
        assessment_method="考试",
        is_deleted=deleted,
    )


def _template(path: Any) -> None:
    doc = Document()
    title = doc.add_paragraph()
    placeholder = title.add_run("______")
    placeholder.underline = WD_UNDERLINE.SINGLE
    title.add_run("年度部门培训计划表")
    info = doc.add_paragraph()
    info.add_run("部门：")
    info.add_run("          ")
    info.add_run("版本：")
    table = doc.add_table(rows=1, cols=7)
    headers = [
        "序号",
        "培训类型",
        "培训时间",
        "培训内容",
        "培训对象",
        "授课人",
        "考核方式",
    ]
    for idx, value in enumerate(headers):
        table.cell(0, idx).text = value
    for marker in ("", "……"):
        row = table.add_row()
        row.cells[0].text = marker
        row.cells[1].text = "□内训 □外训"
    remark = table.add_row()
    remark.cells[0].text = "备注"
    approval = table.add_row()
    approval.cells[0].text = "制表人签名"
    doc.save(path)


def test_find_template_raises_for_missing_template(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Any
) -> None:
    monkeypatch.chdir(tmp_path)
    monkeypatch.setattr(generator, "__file__", str(tmp_path / "pkg" / "module.py"))
    with pytest.raises(FileNotFoundError):
        generator._find_template("部门级")


def test_table_helpers_map_fill_clone_and_clear_rows() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=7)
    headers = ["序号", "培训类型", "月度", "教材", "对象", "授课", "考核"]
    for idx, value in enumerate(headers):
        table.cell(0, idx).text = value
    data = table.add_row()
    remark = table.add_row()
    remark.cells[0].text = "备注"
    approval = table.add_row()
    approval.cells[0].text = "签名"
    header, rows, remark_idx, approvals = generator._get_table_info(table)
    assert (header, rows, remark_idx, approvals) == (0, [1], 2, [3])
    mapping = generator._get_column_map(table, header, "部门级")
    generator._fill_row(data, _item(1), 1, mapping)
    assert data.cells[0].text == "1"
    assert data.cells[2].text == "1月"
    assert "内训" in data.cells[1].text
    clone = generator._clone_row_template(table, 1)
    assert "课程1" not in "".join(clone.itertext())
    generator._clear_row(data)
    assert all(not cell.text for cell in data.cells)


def test_year_and_training_type_helpers_rebuild_expected_runs() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("年度公司培训计划表")
    generator._fill_year_in_paragraphs(doc, 2026)
    assert paragraph.text.startswith("2026年度")
    table = doc.add_table(rows=1, cols=1)
    generator._fill_training_type_cell(table.cell(0, 0), "内训+外训")
    xml = table.cell(0, 0)._tc.xml
    assert "0052" in xml
    assert "内训" in table.cell(0, 0).text
    assert "外训" in table.cell(0, 0).text


def test_generate_annual_plan_document_fills_and_expands_template(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Any
) -> None:
    template = tmp_path / "template.docx"
    _template(template)
    monkeypatch.setattr(generator, "_find_template", lambda level: template)
    plan: Any = SimpleNamespace(
        plan_level="部门级",
        year=2026,
        department="质量部",
        version="V2",
        remarks="重点培训",
    )
    output = generator.generate_annual_plan_doc(
        plan, [_item(1), _item(2), _item(3), _item(4, deleted=True)]
    )
    assert isinstance(output, BytesIO)
    generated = Document(output)
    full_text = "\n".join(p.text for p in generated.paragraphs)
    assert "2026年度部门培训计划表" in full_text
    assert "质量部" in full_text
    assert "V2" in full_text
    table_text = "\n".join(
        cell.text for row in generated.tables[0].rows for cell in row.cells
    )
    assert "课程1" in table_text
    assert "课程2" in table_text
    assert "课程3" in table_text
    assert "课程4" not in table_text
    assert "重点培训" in table_text


def test_generate_annual_plan_document_removes_unused_rows(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Any
) -> None:
    template = tmp_path / "template.docx"
    _template(template)
    monkeypatch.setattr(generator, "_find_template", lambda level: template)
    plan: Any = SimpleNamespace(
        plan_level="公司级",
        year=2027,
        department="",
        version=None,
        remarks=None,
    )
    generated = Document(generator.generate_annual_plan_doc(plan, [_item(1)]))
    table_text = "\n".join(
        cell.text for row in generated.tables[0].rows for cell in row.cells
    )
    assert "课程1" in table_text
    assert "……" not in table_text
