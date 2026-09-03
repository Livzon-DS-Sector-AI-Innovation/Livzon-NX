from __future__ import annotations

from datetime import date
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from typing import Any
from unittest.mock import AsyncMock

import openpyxl
import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fastapi import UploadFile
from openpyxl import Workbook

from app.modules.hr import (
    employee_training_list_document_generator as employee_list_generator,
)
from app.modules.hr import (
    exam_score_parser,
    notification_document_generator,
    oral_exam_document_generator,
    plan_tracking_document_generator,
    position_training_confirmation_generator,
    signin_document_generator,
    trainer_document_generator,
    training_evaluation_document_generator,
)
from app.modules.hr.schemas import (
    OralExamExportRequest,
    OralExamPersonItem,
    OralExamQuestionItem,
    TrainingNotificationInput,
    TrainingSignInSheetInput,
)


def _docx_table(path: Path, rows: int, cols: int) -> None:
    doc = Document()
    table = doc.add_table(rows=rows, cols=cols)
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].add_run("")
    doc.save(path)


def _oral_template(path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=0, cols=7)

    def add_row(first: str = "", values: dict[int, str] | None = None) -> None:
        cells = table.add_row().cells
        cells[0].text = first
        for index, value in (values or {}).items():
            cells[index].text = value

    add_row(values={1: "模板", 5: "模板"})
    add_row("1")
    add_row("……")
    add_row("序号")
    add_row("1", {4: "合格□不合格□"})
    add_row("……")
    add_row(values={0: "评估人"})
    # 模拟真实模板预置的可编辑例外标记（重复 ID 场景由克隆行放大）
    question_cell = table.rows[1].cells[1]
    start = OxmlElement("w:permStart")
    start.set(qn("w:id"), "0")
    start.set(qn("w:edGrp"), "everyone")
    question_cell.paragraphs[0]._p.append(start)
    end = OxmlElement("w:permEnd")
    end.set(qn("w:id"), "0")
    question_cell.paragraphs[0]._p.append(end)
    doc.save(path)


def _perm_starts(element: Any) -> list[Any]:
    return list(element.iter(qn("w:permStart")))


def _notification_template(path: Path) -> None:
    doc = Document()
    labels = [
        "一、培训内容",
        "",
        "二、培训对象",
        "",
        "三、培训时间",
        "",
        "培训地点",
        "",
        "备用",
        "备用",
        "备用",
        "备用",
        "备用",
        "备用",
        "六、培训考核：",
        "备用",
        "特此通知",
        "",
        "",
        "",
    ]
    for index, label in enumerate(labels):
        paragraph = doc.add_paragraph()
        paragraph.add_run(label)
        if index == 14:
            paragraph.add_run("")
    doc.save(path)


@pytest.mark.asyncio
async def test_oral_evaluation_and_notification_generators(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    oral_path = tmp_path / "oral.docx"
    _oral_template(oral_path)
    monkeypatch.setattr(
        oral_exam_document_generator, "_find_template", lambda: oral_path
    )
    oral_data = OralExamExportRequest(
        training_content="GMP",
        training_date="2026-08-20",
        questions=[
            OralExamQuestionItem(no="1", question="问题一", answer="答案一"),
            OralExamQuestionItem(no="2", question="问题二", answer="答案二"),
        ],
        persons=[
            OralExamPersonItem(
                name="张三", department="质量部", question_nos="1", result="合格"
            ),
            OralExamPersonItem(
                name="李四", department="质量部", question_nos="2", result="不合格"
            ),
        ],
        assessor="王五",
    )
    output = oral_exam_document_generator.generate_oral_exam_result(oral_data)
    assert isinstance(output, BytesIO)
    saved = Document(BytesIO(output.getvalue()))
    assert "问题一" in saved.tables[0].cell(1, 1).text
    assert "☑" in saved.tables[0].cell(5, 4).text
    assert (
        oral_exam_document_generator._find_row(saved.tables[0], lambda row: False) == -1
    )
    # 可编辑例外区：模板旧标记清空重建，ID 唯一且全部为"每个人"
    starts = _perm_starts(saved.element.body)
    marker_ids = [s.get(qn("w:id")) for s in starts]
    assert "0" not in marker_ids
    assert len(marker_ids) == len(set(marker_ids))
    assert all(s.get(qn("w:edGrp")) == "everyone" for s in starts)
    # 培训内容/日期、克隆出的问题行/人员行、评估人均可编辑
    assert len(_perm_starts(saved.tables[0].cell(1, 1)._tc)) == 1
    assert len(_perm_starts(saved.tables[0].cell(2, 1)._tc)) == 1
    assert len(_perm_starts(saved.tables[0].cell(5, 1)._tc)) == 1
    assert len(_perm_starts(saved.tables[0].cell(6, 5)._tc)) == 1
    # 锁定区保持锁定：标题行、题头行、序号表头无任何例外标记
    assert not _perm_starts(saved.tables[0].rows[0]._tr)
    assert not _perm_starts(saved.tables[0].rows[3]._tr)

    notification_path = tmp_path / "notification.docx"
    _notification_template(notification_path)
    monkeypatch.setattr(
        notification_document_generator, "_find_template", lambda: notification_path
    )
    notification = notification_document_generator.generate_training_notification(
        TrainingNotificationInput(
            department="质量部",
            training_date=date(2026, 8, 20),
            subject="偏差调查",
            training_time_start="09:00",
            training_time_end="10:00",
            location="一号会议室",
            trainee_names=["张三", "李四"],
            issuer_department="人力资源部",
            issue_date=date(2026, 8, 19),
            assessment_method="口试",
        )
    )
    notification_doc = Document(BytesIO(notification.getvalue()))
    text = "\n".join(p.text for p in notification_doc.paragraphs)
    assert "偏差调查" in text
    assert "一号会议室" in text
    assert "人力资源部" in text


@pytest.mark.asyncio
async def test_training_evaluation_and_signin_generators(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    evaluation_path = tmp_path / "evaluation.docx"
    _docx_table(evaluation_path, 20, 7)
    monkeypatch.setattr(
        training_evaluation_document_generator,
        "_find_template",
        lambda: evaluation_path,
    )
    values = {
        "subject": "年度培训",
        "training_date": date(2026, 8, 20),
        "duration_hours": 2,
        "training_method": "其他：线上",
        "instructor": "讲师",
        "target_dept_person": "质量部",
        "expected_count": 10,
        "actual_count": 9,
        "absent_count": 1,
        "textbook": "SOP",
        "absent_handling": "补训",
        "need_retraining": True,
        "retraining_info": "下周补训",
        "assessment_method": "笔试",
        "excellent_count": 2,
        "good_count": 3,
        "pass_count": 3,
        "fail_count": 1,
        "absent_exam_count": 1,
        "fail_handling": "补考",
        "makeup_count": 1,
        "makeup_pass_count": 1,
        "makeup_fail_count": 0,
        "makeup_fail_handling": "重新培训",
        "evaluation_result": "经考核，基本达到培训效果。",
        "evaluation_comment": "效果良好",
        "evaluator": "评估人",
        "evaluate_date": date(2026, 8, 21),
        "has_notification": True,
        "has_signin_sheet": False,
        "has_textbook": True,
        "has_exam_paper": False,
        "has_score_summary": True,
        "has_notification_qty": 1,
        "other_attachment": "照片",
    }
    evaluation = (
        training_evaluation_document_generator.generate_training_evaluation_doc(values)
    )
    assert len(evaluation.getvalue()) > 100
    orm = SimpleNamespace(**values, training_content="ORM培训")
    assert (
        len(
            training_evaluation_document_generator.generate_training_evaluation_doc_from_orm(
                orm
            ).getvalue()
        )
        > 100
    )
    assert training_evaluation_document_generator._d(None) == ""
    assert training_evaluation_document_generator._d(date(2026, 8, 20))
    assert training_evaluation_document_generator._int(None) == ""

    def signin_template(path: Path) -> None:
        doc = Document()
        for page in range(2):
            for index in range(4):
                doc.add_paragraph(f"page-{page}-{index}")
            table = doc.add_table(rows=23, cols=10)
            table.cell(0, 0).text = "培训日期"
            table.cell(1, 0).text = "面授 □ 实操 □ 其他 □"
            table.cell(2, 0).text = "应受训人数： 人"
            table.cell(4, 0).text = ""
            table.cell(4, 2).text = ""
            table.cell(4, 8).text = ""
            table.cell(5, 0).text = "序号"
            for row in range(6, 23):
                table.cell(row, 0).text = str(row - 5)
        doc.save(path)

    signin_path = tmp_path / "signin.docx"
    signin_template(signin_path)
    monkeypatch.setattr(
        signin_document_generator, "_find_template", lambda: signin_path
    )
    names = [f"员工{i}" for i in range(45)]
    signin = signin_document_generator.generate_training_sign_in_sheet(
        TrainingSignInSheetInput(
            training_date=date(2026, 8, 20),
            training_time_start="09:00",
            training_time_end="10:00",
            department="质量部",
            training_subject="GMP",
            topic="偏差",
            instructor="讲师",
            training_method="其他：线上",
            employee_names=names,
            remarks="备注",
        )
    )
    assert len(Document(BytesIO(signin.getvalue())).tables) == 2


def test_excel_generators_and_trainer_export(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    plan_path = tmp_path / "plan.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    for row in range(1, 15):
        for column in range(1, 11):
            sheet.cell(row=row, column=column).value = ""
    workbook.save(plan_path)
    monkeypatch.setattr(
        plan_tracking_document_generator, "_find_template", lambda: plan_path
    )
    records = [
        SimpleNamespace(
            training_content="内容",
            actual_time="2026-08-20",
            target_audience="质量部",
            training_type="部门级",
            tracking_assessment_method="笔试",
            is_completed=value,
            tracker="张三",
            track_date=date(2026, 8, 20),
            remarks="备注",
        )
        for value in (True, False, None)
    ]
    output = plan_tracking_document_generator.generate_plan_tracking_excel(
        records * 3, year=2026, month=8, plan_level="部门级"
    )
    assert openpyxl.load_workbook(BytesIO(output.getvalue())).active["A4"].value

    employee_path = tmp_path / "employee.xlsx"
    employee_book = Workbook()
    employee_sheet = employee_book.active
    employee_sheet.title = "Sheet1"
    for row in range(1, 57):
        for column in range(1, 6):
            employee_sheet.cell(row=row, column=column).value = ""
    employee_book.save(employee_path)
    monkeypatch.setattr(
        employee_list_generator, "_find_template", lambda: employee_path
    )
    employee_output = employee_list_generator.generate_employee_training_list(
        "质量部",
        "张三",
        [
            {
                "training_datetime": "09:00~10:00",
                "training_date": date(2026, 8, 20),
                "training_content": "GMP",
                "personal_score": "合格",
                "remarks": None,
            }
        ]
        * 50,
        year=2026,
    )
    assert (
        openpyxl.load_workbook(BytesIO(employee_output.getvalue())).active["A4"].value
    )

    trainer_path = tmp_path / "trainer.docx"
    trainer_doc = Document()
    trainer_table = trainer_doc.add_table(rows=2, cols=5)
    for row in trainer_table.rows:
        for cell in row.cells:
            cell.text = ""
    trainer_doc.save(trainer_path)
    monkeypatch.setattr(
        trainer_document_generator, "_find_template", lambda: trainer_path
    )
    trainer_output = trainer_document_generator.generate_trainer_list(
        [
            SimpleNamespace(
                name="张三",
                department="质量部",
                position="质量员",
                approval_date=date(2026, 8, 20),
                remarks="",
            ),
            SimpleNamespace(
                name="李四",
                department="生产部",
                position="操作员",
                approval_date=None,
                remarks="备注",
            ),
        ]
    )
    assert len(trainer_output.getvalue()) > 100


@pytest.mark.asyncio
async def test_position_confirmation_offer_and_exam_score_parsers(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    confirmation_path = tmp_path / "confirmation.xlsx"
    confirmation_book = Workbook()
    confirmation_sheet = confirmation_book.active
    for row in range(1, 18):
        for column in range(1, 15):
            confirmation_sheet.cell(row=row, column=column).value = ""
    confirmation_book.save(confirmation_path)
    monkeypatch.setattr(
        position_training_confirmation_generator,
        "_find_template",
        lambda: confirmation_path,
    )
    generate_confirmation = (
        position_training_confirmation_generator.generate_position_training_confirmation
    )
    confirmation = generate_confirmation(
        "张三",
        "E001",
        "质量部",
        "质量员",
        "2026-08-20",
        "入职",
        [
            {
                "textbook_name": "GMP",
                "textbook_code": "SOP-1",
                "training_date": "2026-08-20",
                "assessment_result": "合格",
            }
        ]
        * 5,
    )
    assert len(confirmation.getvalue()) > 100

    xlsx = Workbook()
    sheet = xlsx.active
    sheet.append(["姓名", "成绩"])
    sheet.append(["张三", 96])
    sheet.append(["李四", 88.5])
    xlsx_buffer = BytesIO()
    xlsx.save(xlsx_buffer)
    parsed_xlsx = exam_score_parser._parse_xlsx(xlsx_buffer.getvalue())
    assert [item.name for item in parsed_xlsx] == ["张三", "李四"]

    fallback = Workbook()
    fallback_sheet = fallback.active
    fallback_sheet.append(["列A", "列B", "列C"])
    fallback_sheet.append(["张三", 96, "x"])
    fallback_sheet.append(["李四", 88, "x"])
    fallback_buffer = BytesIO()
    fallback.save(fallback_buffer)
    assert len(exam_score_parser._parse_xlsx(fallback_buffer.getvalue())) == 2

    doc = Document()
    doc.add_paragraph("周亚学96 秦亚瑞 93")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "成绩"
    table.cell(1, 0).text = "王五"
    table.cell(1, 1).text = "91"
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    parsed_docx = exam_score_parser._parse_docx(doc_buffer.getvalue())
    assert any(item.name == "周亚学" for item in parsed_docx)
    assert any(item.name == "王五" for item in parsed_docx)
    assert exam_score_parser.format_score_summary(parsed_docx)
    assert (
        exam_score_parser.extract_personal_score("张三", "张三96、李四95", "笔试")
        == "96"
    )
    assert exam_score_parser.extract_personal_score("张三", "/", "口试") == "合格"
    assert exam_score_parser.extract_personal_score("赵六", "张三96", "笔试") is None
    assert exam_score_parser._is_chinese_name("张三")
    assert exam_score_parser._is_score("99.5")
    assert not exam_score_parser._is_score("101")

    monkeypatch.setattr(
        exam_score_parser,
        "read_upload_secure",
        AsyncMock(return_value=("scores.xlsx", xlsx_buffer.getvalue())),
    )
    parsed_upload = await exam_score_parser.parse_exam_scores(
        UploadFile(filename="scores.xlsx", file=BytesIO(xlsx_buffer.getvalue()))
    )
    assert len(parsed_upload) == 2


def test_offer_pdf_generation(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    fitz = pytest.importorskip("fitz")
    from app.modules.hr import offer_pdf_generator

    template = tmp_path / "offer.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "**")
    page.insert_text((50, 70), "*")
    page.insert_text((50, 90), offer_pdf_generator.DATE_MARKER)
    doc.save(template)
    doc.close()
    font = Path(r"C:\Windows\Fonts\arial.ttf")
    if not font.exists():
        pytest.skip("系统未提供可用于 PDF 覆盖测试的字体")
    monkeypatch.setattr(offer_pdf_generator, "_FONT_FILE", str(font))
    output = offer_pdf_generator.generate_offer_pdf(
        str(template), "张三", "质量工程师", str(tmp_path)
    )
    assert Path(output).exists()
