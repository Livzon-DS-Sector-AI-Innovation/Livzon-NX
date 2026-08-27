from __future__ import annotations

from datetime import UTC, date, datetime, time
from types import SimpleNamespace as _SimpleNamespace
from typing import Any
from unittest.mock import AsyncMock, Mock
from uuid import uuid4

import pytest
from docx import Document

from app.core.exceptions import AppException
from app.modules.hr import service, training_dept_resolver

SimpleNamespace: Any = _SimpleNamespace


def _training(repo: SimpleNamespace, session: SimpleNamespace | None = None) -> Any:
    instance = service.TrainingLedgerService.__new__(service.TrainingLedgerService)
    instance.repo = repo
    instance.session = session or SimpleNamespace()
    return instance


def test_feishu_value_and_date_helpers_cover_supported_shapes() -> None:
    assert service._extract_text([{"text": "张三"}]) == "张三"
    assert service._extract_text({"text": "李四"}) == "李四"
    assert service._extract_text({"value": [{"text": "王五"}]}) == "王五"
    assert service._extract_text(None) == ""
    assert service._extract_text(3) == "3"
    assert service._extract_number(3.8) == 3
    assert service._extract_number({"value": [7]}) == 7
    assert service._extract_number({}) is None
    epoch_ms = datetime(2026, 8, 20, tzinfo=UTC).timestamp() * 1000
    assert service._ms_to_date(epoch_ms) == date(2026, 8, 20)
    assert service._ms_to_date(0) is None
    assert service._parse_date("2026-08-20T12:00:00") == date(2026, 8, 20)
    assert service._parse_date("invalid") is None
    assert service._parse_date(-1) is None
    assert service._parse_contract_date("2026/08/20") == date(2026, 8, 20)
    assert service._parse_contract_date("20260820") == date(2026, 8, 20)
    assert service._parse_contract_date("") is None


def test_parse_feishu_employee_record_maps_fields_and_sync_date() -> None:
    milliseconds = datetime(2026, 8, 20, tzinfo=UTC).timestamp() * 1000
    record = {
        "record_id": "rec1",
        "updated_time": "2026-08-20T08:30:00Z",
        "fields": {
            "序号": {"value": [2]},
            "工号": [{"text": "E001"}],
            "姓名": "张三",
            "部门": "质量部",
            "职位": "QA",
            "技能证书": "GMP",
            "拟转正日期": "2026-09-20",
            "进厂时间": milliseconds,
            "首次签订合同日期": milliseconds,
            "合同截止日期（2）": "2027/08/20",
            "手机": "13800000000",
            "备注": "重点培养",
        },
    }
    parsed = service._parse_feishu_record(record)
    assert parsed["feishu_record_id"] == "rec1"
    assert parsed["seq_number"] == 2
    assert parsed["employee_number"] == "E001"
    assert parsed["department"] == "质量部"
    assert parsed["position"] == "QA"
    assert parsed["qualifications"] == ["GMP"]
    assert parsed["hire_date"] == date(2026, 8, 20)
    assert parsed["contract_end_2"] == date(2027, 8, 20)
    assert parsed["remarks"] == ["重点培养"]
    assert parsed["status"] == "在职"
    assert parsed["feishu_synced_at"] == date(2026, 8, 20)

    invalid = service._parse_feishu_record({"updated_time": "bad", "fields": {}})
    assert invalid["feishu_synced_at"] == date.today()


def test_department_mapping_serialization_and_training_name_helpers() -> None:
    mapping: Any = SimpleNamespace(
        id=uuid4(),
        source_name="QA",
        target_name="质量部",
        match_level="exact",
        mapping_type="alias",
        priority=10,
        enabled=True,
        remark=None,
        created_at=datetime(2026, 8, 20),
        updated_at=None,
    )
    result = service._dept_mapping_to_dict(mapping)
    assert result["source_name"] == "QA"
    assert result["created_at"].startswith("2026-08-20")
    assert result["updated_at"] is None
    assert service._split_trainees(None) == []
    assert service._split_trainees("张三、李四 张三；王五") == ["张三", "李四", "王五"]
    records = [
        SimpleNamespace(
            session_id="s1",
            training_date=date.today(),
            training_datetime="a",
            training_content="x",
        ),
        SimpleNamespace(
            session_id="s1",
            training_date=date.today(),
            training_datetime="b",
            training_content="y",
        ),
        SimpleNamespace(
            session_id=None,
            training_date=date.today(),
            training_datetime="a",
            training_content="x",
        ),
        SimpleNamespace(
            session_id=None,
            training_date=date.today(),
            training_datetime="a",
            training_content="x",
        ),
    ]
    assert len(service._dedupe_training_records(records)) == 2


def test_training_datetime_parser_and_free_slot_suggestions() -> None:
    assert service.TrainingLedgerService._parse_datetime_range(None) == (None, None)
    assert service.TrainingLedgerService._parse_datetime_range("全天") == (None, None)
    assert service.TrainingLedgerService._parse_datetime_range("19:00~21:00") == (
        time(19),
        time(21),
    )
    assert (
        service.TrainingLedgerService._suggest_free_slots(time(9), time(10), []) == []
    )
    suggestions = service.TrainingLedgerService._suggest_free_slots(
        time(9), time(10), [(time(8), time(9)), (time(14), time(15))]
    )
    assert suggestions[0] == {"start": "09:00", "end": "10:00"}
    assert len(suggestions) <= 3


@pytest.mark.anyio
async def test_training_department_and_mapping_operations(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    mapping: Any = SimpleNamespace(
        id=uuid4(),
        source_name="QA",
        target_name="质量部",
        match_level="exact",
        mapping_type="alias",
        priority=1,
        enabled=True,
        remark=None,
        created_at=None,
        updated_at=None,
        updated_by=None,
    )
    repo: Any = SimpleNamespace(
        list_all_training_departments=AsyncMock(return_value=["质量部", "生产部"]),
        list_custom_training_departments=AsyncMock(return_value=["生产部"]),
        add_custom_training_department=AsyncMock(
            return_value=SimpleNamespace(name="研发部", id=uuid4())
        ),
        delete_custom_training_department=AsyncMock(return_value=True),
        list_dept_mappings=AsyncMock(return_value=[mapping]),
        create_dept_mapping=AsyncMock(return_value=mapping),
        get_dept_mapping=AsyncMock(return_value=mapping),
        delete_dept_mapping=AsyncMock(),
    )
    session: Any = SimpleNamespace(
        flush=AsyncMock(), refresh=AsyncMock(), rollback=AsyncMock()
    )
    instance = _training(repo, session)
    monkeypatch.setattr(
        training_dept_resolver,
        "resolve_training_department",
        AsyncMock(
            side_effect=lambda session, name: "研发部" if name == "研发" else name
        ),
    )
    invalidate: Any = Mock()
    monkeypatch.setattr(
        training_dept_resolver, "invalidate_training_dept_mapping_cache", invalidate
    )

    added = await instance.add_custom_training_department("研发")
    assert added["name"] == "研发部"
    with pytest.raises(AppException):
        await instance.add_custom_training_department("质量部")
    with pytest.raises(AppException):
        await instance.delete_custom_training_department("质量部")
    assert await instance.delete_custom_training_department("生产部") is True
    assert (await instance.list_dept_mappings())[0]["source_name"] == "QA"

    payload: Any = SimpleNamespace(
        source_name="研发",
        target_name="研发部",
        match_level="exact",
        mapping_type="alias",
        priority=2,
        enabled=True,
        remark="映射",
    )
    created = await instance.create_dept_mapping(payload, user_id=uuid4())
    assert created["target_name"] == "质量部"
    payload.source_name = "QA"
    payload.target_name = "质量部"
    payload.priority = 1
    with pytest.raises(AppException):
        await instance.create_dept_mapping(payload)

    update: Any = SimpleNamespace(priority=5)
    updated = await instance.update_dept_mapping(mapping.id, update, user_id=uuid4())
    assert updated["priority"] == 5
    assert await instance.delete_dept_mapping(mapping.id) is True
    repo.get_dept_mapping.return_value = None
    assert await instance.delete_dept_mapping(mapping.id) is False
    with pytest.raises(AppException):
        await instance.update_dept_mapping(mapping.id, update)
    assert invalidate.call_count == 3


def test_training_type_and_plan_metadata_parsing() -> None:
    assert service._parse_training_type("内训 外训") is None
    assert service._parse_training_type("☑内训 □外训") == "内训"
    assert service._parse_training_type("□内训 √外训") == "外训"
    assert service._parse_training_type("☑内训 ☑外训") == "内训+外训"
    assert service._parse_training_type("内训") == "内训"
    assert service._parse_training_type("未知") is None
    doc = Document()
    doc.add_paragraph("2026年度部门培训计划表")
    doc.add_paragraph("部门：质量部 版本：V2")
    assert service._detect_plan_meta_from_doc(doc) == ("部门级", "质量部", "V2")
    company = Document()
    company.add_paragraph("2026年度公司培训计划表")
    assert service._detect_plan_meta_from_doc(company) == ("公司级", None, None)


def test_parse_plan_items_from_doc_skips_non_data_rows_and_extracts_remarks() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=7)
    headers = [
        "序号",
        "培训类型",
        "培训时间",
        "培训内容",
        "培训对象",
        "授课人",
        "考核方式",
    ]
    for idx, text_value in enumerate(headers):
        table.cell(0, idx).text = text_value
    data = table.add_row().cells
    values = ["1", "☑内训 □外训", "8月", "GMP", "全员", "张老师", "考试"]
    for idx, text_value in enumerate(values):
        data[idx].text = text_value
    empty = table.add_row().cells
    empty[1].text = "内训 外训"
    approval = table.add_row().cells
    approval[0].text = "制表人签名"
    remark = table.add_row().cells
    remark[0].text = "备注"
    remark[1].text = "重点课程"
    remark[2].text = "重点课程"

    items, plan_remark = service._parse_plan_items_from_doc(doc)
    assert items == [
        {
            "training_type": "内训",
            "training_month": "8月",
            "content_textbook": "GMP",
            "target_audience_new": "全员",
            "instructor": "张老师",
            "assessment_method": "考试",
            "sort_order": 0,
        }
    ]
    assert plan_remark == "重点课程"
