"""培训附件 Word 生成器单元测试（模板保真：附件： + 序号/文件名称/文件编号表格）."""

from io import BytesIO

import pytest
from docx import Document

from app.modules.hr import training_attachment_document_generator as generator
from app.modules.hr.schemas import TrainingAttachmentItem


@pytest.fixture
def training_template(tmp_path, monkeypatch):
    path = tmp_path / "培训附件.docx"
    document = Document()
    document.add_paragraph("附件：")
    table = document.add_table(rows=5, cols=3)
    for index, value in enumerate(("序号", "文件名称", "文件编号")):
        table.cell(0, index).text = value
    document.save(path)
    monkeypatch.setattr(generator, "_find_template", lambda: path)
    return path


def _rows_of(buffer: BytesIO) -> list[list[str]]:
    doc = Document(buffer)
    assert doc.tables, "模板中应存在表格"
    return [[c.text.strip() for c in row.cells] for row in doc.tables[0].rows]


def test_generate_with_more_items_than_preset_rows(training_template):
    """多于模板预置行（4 行）时按最后一行克隆扩展，序号/名称/编号正确填充."""
    items = [
        TrainingAttachmentItem(name=f"文件{i}", code=f"编号{i:03d}")
        for i in range(1, 8)
    ]
    rows = _rows_of(generator.generate_training_attachment(items))

    assert len(rows) == 8  # 表头 + 7 数据行
    assert rows[0] == ["序号", "文件名称", "文件编号"]  # 表头保留
    for i, expected in enumerate(range(1, 8), start=1):
        assert rows[i][0] == str(expected)
        assert rows[i][1] == f"文件{expected}"
        assert rows[i][2] == f"编号{expected:03d}"


def test_generate_with_fewer_items_than_preset_rows(training_template):
    """少于模板预置行时删除多余行，仅保留实际数据行."""
    items = [
        TrainingAttachmentItem(name="A", code="A1"),
        TrainingAttachmentItem(name="B", code="B1"),
    ]
    rows = _rows_of(generator.generate_training_attachment(items))

    assert len(rows) == 3  # 表头 + 2 数据行
    assert rows[1] == ["1", "A", "A1"]
    assert rows[2] == ["2", "B", "B1"]


def test_generate_with_empty_items(training_template):
    """空清单时仅保留表头行，不报错."""
    rows = _rows_of(generator.generate_training_attachment([]))
    assert len(rows) == 1
    assert rows[0] == ["序号", "文件名称", "文件编号"]


def test_generate_with_missing_code(training_template):
    """无编号条目仅填名称，编号列留空."""
    items = [TrainingAttachmentItem(name="无编号文件")]
    rows = _rows_of(generator.generate_training_attachment(items))
    assert rows[1] == ["1", "无编号文件", ""]
