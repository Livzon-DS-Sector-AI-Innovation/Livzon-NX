from __future__ import annotations

from types import SimpleNamespace as _SimpleNamespace
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.modules.hr import position_training_document_generator as generator

SimpleNamespace: Any = _SimpleNamespace


def _item(level: str, order: int = 0) -> SimpleNamespace:
    return SimpleNamespace(
        level=level,
        textbook_name="GMP基础",
        textbook_code="SMP-001",
        assessment_method="考试",
        remarks="必修",
        sort_order=order,
        is_deleted=False,
    )


def _template_table() -> Any:
    doc = Document()
    table = doc.add_table(rows=16, cols=5)
    table.cell(0, 0).text = "部门级"
    for row_index in (1, 9):
        headers = ["序号", "教材名称", "教材编码", "考核方式", "备注"]
        for idx, value in enumerate(headers):
            table.cell(row_index, idx).text = value
    for row_index in range(2, 7):
        table.cell(row_index, 0).text = str(row_index - 1)
    table.cell(7, 0).text = "……"
    table.cell(8, 0).text = "岗位级"
    for row_index in range(10, 15):
        table.cell(row_index, 0).text = str(row_index - 9)
    table.cell(15, 0).text = "……"
    return doc, table


def test_fill_cell_and_paragraph_helpers() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    generator._fill_cell(table.cell(0, 0), "内容", 10.5, bold=True, shade="D7D7D7")
    assert table.cell(0, 0).text == "内容"
    paragraph = doc.add_paragraph("部门：")
    for index, label in enumerate(("", " 岗位：")):
        if label:
            paragraph.add_run(label)
        start = OxmlElement("w:permStart")
        start.set(qn("w:id"), str(index))
        paragraph._p.append(start)
        paragraph.add_run("____")
        end = OxmlElement("w:permEnd")
        end.set(qn("w:id"), str(index))
        paragraph._p.append(end)
    generator._fill_dept_pos_paragraph(doc, "质量部", "QA")
    assert "质量部" in paragraph.text
    assert "QA" in paragraph.text
    generator._enable_update_fields(doc)
    assert "updateFields" in doc.settings.element.xml


def test_build_table_replaces_template_rows_and_keeps_empty_placeholder() -> None:
    _, table = _template_table()
    generator._build_table(table, [_item("部门级")], [])
    all_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    assert "GMP基础" in all_text
    assert "SMP-001" in all_text
    assert all_text.count("……") >= 1
    assert any(row.cells[0].text == "部门级" for row in table.rows)
    assert any(row.cells[0].text == "岗位级" for row in table.rows)


def test_build_table_handles_empty_department_and_position_data() -> None:
    _, table = _template_table()
    generator._build_table(table, [], [])
    assert len(table.rows) == 16
    assert sum(row.cells[0].text == "……" for row in table.rows) == 2


def test_clone_fill_and_permission_markers_cover_xml_variants() -> None:
    _, table = _template_table()
    header = table.rows[1]._tr
    clone = generator._clone_data_row(header)
    generator._fill_row_element(clone, 1, _item("部门级"))
    xml = clone.xml
    assert "GMP基础" in xml
    assert "permStart" in xml
    assert "permEnd" in xml
    empty = clone.__class__()
    generator._add_row_perm_markers(empty)


def test_placeholder_noops_when_items_or_header_missing() -> None:
    _, table = _template_table()
    before = len(table.rows)
    generator._ensure_placeholder(table, "部门级", [_item("部门级")])
    assert len(table.rows) == before
    generator._ensure_placeholder(table, "不存在", [])
    assert len(table.rows) == before
