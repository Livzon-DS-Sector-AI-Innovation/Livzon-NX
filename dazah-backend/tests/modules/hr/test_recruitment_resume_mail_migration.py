from __future__ import annotations

import json
from email.message import EmailMessage
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock, Mock

import pytest
from docx import Document
from pypdf import PdfWriter

from app.core.llm import LLMOutputError
from app.modules.hr import mail_fetcher, recruitment_service, resume_watcher


@pytest.mark.asyncio
async def test_resume_watcher_helpers_scan_and_match(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    hash_file = tmp_path / resume_watcher.HASH_RECORD_FILE
    hash_file.write_text(json.dumps({"hashes": ["from-json"]}), encoding="utf-8")
    (tmp_path / "简历_from-name_candidate.pdf").write_bytes(b"root")
    processed = tmp_path / "processed"
    processed.mkdir()
    (processed / "processed.docx").write_bytes(b"processed-content")
    processed_hash = resume_watcher.hashlib.sha256(b"processed-content").hexdigest()[
        :12
    ]
    assert {"from-json", "from-name", processed_hash}.issubset(
        resume_watcher._load_all_hashes(tmp_path)
    )

    hash_file.write_text("not-json", encoding="utf-8")
    resume_watcher._save_hashes(tmp_path, {"z", "a"})
    assert json.loads(hash_file.read_text(encoding="utf-8")) == ["a", "z"]

    session = SimpleNamespace(commit=AsyncMock())

    class _SessionContext:
        async def __aenter__(self) -> SimpleNamespace:
            return session

        async def __aexit__(self, *_args: object) -> None:
            return None

    import app.core.database as database
    import app.shared.config_reader as config_reader

    monkeypatch.setattr(
        resume_watcher,
        "get_module_setting",
        AsyncMock(return_value=resume_watcher._LEGACY_DESKTOP_WATCH_DIR),
    )
    monkeypatch.setattr(database, "async_session_factory", lambda: _SessionContext())
    set_setting = AsyncMock()
    monkeypatch.setattr(config_reader, "set_module_setting", set_setting)
    assert await resume_watcher._resolve_watch_dir() == Path(
        resume_watcher.DEFAULT_RESUME_WATCH_DIR
    )
    set_setting.assert_awaited_once()
    session.commit.assert_awaited_once()

    monkeypatch.setattr(
        resume_watcher,
        "get_module_setting",
        AsyncMock(return_value=str(tmp_path / "custom")),
    )
    assert await resume_watcher._resolve_watch_dir() == tmp_path / "custom"

    new_file = tmp_path / "new.pdf"
    failed_file = tmp_path / "retry.docx"
    duplicate = tmp_path / "duplicate.pdf"
    new_file.write_bytes(b"new")
    failed_file.write_bytes(b"failed")
    duplicate.write_bytes(b"duplicate")
    duplicate_hash = resume_watcher.hashlib.sha256(b"duplicate").hexdigest()[:12]
    hash_file.write_text(json.dumps([duplicate_hash]), encoding="utf-8")
    monkeypatch.setattr(
        resume_watcher,
        "_resolve_watch_dir",
        AsyncMock(return_value=tmp_path),
    )
    submit_job = AsyncMock(side_effect=[RuntimeError("queue unavailable"), None, None])
    import app.core.jobs as jobs

    monkeypatch.setattr(jobs, "submit_job", submit_job)
    result = await resume_watcher.scan_watched_folder()
    assert result["new_files"] == 3
    assert not duplicate.exists()
    assert submit_job.await_count == 3

    assert resume_watcher._normalize_education("") == ""
    assert resume_watcher._normalize_education("硕士研究生") == "硕士"
    assert resume_watcher._normalize_education("中专") == "其他"
    assert resume_watcher._normalize_fit_level("强烈推荐") == "非常满足"
    assert resume_watcher._normalize_fit_level("未知") == "低"
    assert await resume_watcher._match_job_position("", {"j": "工程师"}) is None
    assert (
        await resume_watcher._match_job_position("工程师", {"j": "工程师"}) == "工程师"
    )

    import app.core.llm as llm

    monkeypatch.setattr(
        llm,
        "llm_client",
        SimpleNamespace(chat_json=AsyncMock(return_value={"job_title": "研发工程师"})),
    )
    assert (
        await resume_watcher._match_job_position(
            "开发", {"j": "研发工程师", "k": "质量员"}
        )
        == "研发工程师"
    )
    llm.llm_client.chat_json = AsyncMock(return_value={"job_title": "不存在"})
    assert await resume_watcher._match_job_position("开发", {"j": "研发工程师"}) is None
    llm.llm_client.chat_json = AsyncMock(side_effect=LLMOutputError("bad"))
    assert await resume_watcher._match_job_position("开发", {"j": "研发工程师"}) is None


def test_resume_text_extraction_and_normalization(tmp_path: Path) -> None:
    docx_path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("候选人经历")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "技能"
    table.cell(0, 1).text = "GMP"
    doc.save(docx_path)
    assert "候选人经历" in resume_watcher._extract_resume_text(docx_path)
    assert "技能 | GMP" in resume_watcher._extract_resume_text(docx_path)

    pdf_path = tmp_path / "resume.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    with pdf_path.open("wb") as stream:
        writer.write(stream)
    assert resume_watcher._extract_resume_text(pdf_path) == ""
    with pytest.raises(ValueError, match="Unsupported"):
        resume_watcher._extract_resume_text(tmp_path / "resume.txt")


@pytest.mark.asyncio
async def test_process_resume_creates_candidate_with_attachment_failure(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    resume_path = tmp_path / "candidate.docx"
    doc = Document()
    doc.add_paragraph("张三 本科 GMP")
    doc.save(resume_path)

    bitable = SimpleNamespace(
        upload_attachment=AsyncMock(side_effect=RuntimeError("upload failed"))
    )

    class _Repo:
        def __init__(self) -> None:
            self.updated = AsyncMock()

        async def get_job_names(self) -> dict[str, str]:
            return {"job-1": "质量工程师"}

        async def create_candidate(self, fields: dict[str, object]) -> dict[str, str]:
            assert fields["job_id"] == "质量工程师"
            return {"id": "candidate-1"}

        async def _get_client(self) -> SimpleNamespace:
            return bitable

        async def update_candidate(self, *_args: object) -> None:
            await self.updated()

    monkeypatch.setattr(
        "app.modules.hr.recruitment_repository.RecruitmentBitableRepo", _Repo
    )
    import app.core.llm as llm

    monkeypatch.setattr(
        llm,
        "llm_client",
        SimpleNamespace(
            chat_json=AsyncMock(
                return_value={
                    "name": "张三",
                    "phone": "13800000000",
                    "email": "z@example.com",
                    "target_position": "质量工程师",
                    "education_level": "本科",
                    "work_years": 3,
                    "skills": ["GMP", "Excel"],
                    "match_rate": 80,
                    "resume_score": 75,
                    "fit_level": "推荐",
                    "reason": "经历匹配",
                }
            )
        ),
    )
    monkeypatch.setattr(
        resume_watcher,
        "_resolve_watch_dir",
        AsyncMock(return_value=tmp_path),
    )
    await resume_watcher.process_single_resume(str(resume_path), "hash-1")
    assert not resume_path.exists()
    assert "hash-1" in resume_watcher._load_all_hashes(tmp_path)

    await resume_watcher.process_single_resume(str(tmp_path / "missing.docx"), "hash-2")


@pytest.mark.asyncio
async def test_mail_fetcher_disabled_success_and_error_paths(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    settings = {
        "HR_MAIL_IMAP_HOST": "imap.example.com",
        "HR_MAIL_IMAP_PORT": "993",
        "HR_MAIL_IMAP_USER": "hr@example.com",
        "HR_MAIL_IMAP_PASS": "encrypted",
        "HR_MAIL_FETCH_ENABLED": "true",
    }

    async def get_setting(_module: str, key: str, default: str | None = None) -> str:
        return settings.get(key, default or "")

    monkeypatch.setattr(mail_fetcher, "get_module_setting", get_setting)
    save_status = AsyncMock()
    monkeypatch.setattr(mail_fetcher, "_save_fetch_status", save_status)
    settings["HR_MAIL_FETCH_ENABLED"] = "false"
    disabled = await mail_fetcher.fetch_resumes_from_mail()
    assert disabled == {"status": "not_configured_or_disabled", "fetched": 0}
    save_status.assert_awaited_once_with(0, "not_configured_or_disabled")

    settings["HR_MAIL_FETCH_ENABLED"] = "true"
    monkeypatch.setattr("app.core.llm.decrypt_api_key", lambda _value: "plain-password")
    monkeypatch.setattr(
        "app.modules.hr.resume_watcher._resolve_watch_dir",
        AsyncMock(return_value=tmp_path),
    )
    message = EmailMessage()
    message["Subject"] = "resume"
    message.set_content("see attachment")
    message.add_attachment(
        b"resume-by-mail",
        maintype="application",
        subtype="pdf",
        filename="candidate.pdf",
    )
    message_bytes = message.as_bytes()

    class _Mail:
        def login(self, *_args: object) -> tuple[str, list[bytes]]:
            return "OK", []

        def select(self, *_args: object) -> tuple[str, list[bytes]]:
            return "OK", []

        def search(self, *_args: object) -> tuple[str, list[bytes]]:
            return "OK", [b"1"]

        def fetch(self, *_args: object) -> tuple[str, list[tuple[bytes, bytes]]]:
            return "OK", [(b"meta", message_bytes)]

        def store(self, *_args: object) -> tuple[str, list[bytes]]:
            return "OK", []

        def logout(self) -> tuple[str, list[bytes]]:
            return "BYE", []

    monkeypatch.setattr(mail_fetcher.imaplib, "IMAP4_SSL", lambda *_args: _Mail())
    save_status.reset_mock()
    fetched = await mail_fetcher.fetch_resumes_from_mail(force_redownload=True)
    assert fetched == {"status": "ok", "fetched": 1, "scanned": 1}
    assert (tmp_path / "candidate.pdf").read_bytes() == b"resume-by-mail"
    assert save_status.await_args.args == (1, "ok")

    monkeypatch.setattr(
        mail_fetcher.imaplib,
        "IMAP4_SSL",
        Mock(side_effect=RuntimeError("imap down")),
    )
    errored = await mail_fetcher.fetch_resumes_from_mail()
    assert errored["status"] == "error"
    assert errored["scanned"] == 0


@pytest.mark.asyncio
async def test_recruitment_batch_analysis_covers_vision_skip_and_failure(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class _Client:
        async def request(self, _method: str, path: str) -> dict[str, object]:
            cid = path.rsplit("/", 1)[-1]
            if cid == "failed":
                raise RuntimeError("record unavailable")
            if cid == "vision":
                return {
                    "record": {
                        "record_id": cid,
                        "fields": {
                            "姓名": "视觉候选人",
                            "简历附件": [
                                {"file_token": "file-1", "name": "resume.pdf"}
                            ],
                        },
                    }
                }
            if cid == "skip":
                return {"record": {"record_id": cid, "fields": {}}}
            return {
                "record": {
                    "record_id": cid,
                    "name": "文本候选人",
                    "education": "本科",
                    "skills": "GMP",
                    "fields": {"姓名": "文本候选人", "学历": "本科", "技能标签": "GMP"},
                }
            }

        def _path(self, _table: str, suffix: str) -> str:
            return suffix

        async def download_file(self, _token: str) -> bytes:
            return b"pdf-content"

    client = _Client()
    bitable = SimpleNamespace(client=client, _path=client._path)
    repo = SimpleNamespace(
        _get_client=AsyncMock(return_value=bitable),
        update_candidate=AsyncMock(),
        get_job=AsyncMock(
            return_value={
                "职位名称": "质量工程师",
                "任职要求": "GMP",
                "要求技能": ["GMP"],
            }
        ),
        list_candidates=AsyncMock(
            return_value=(
                [
                    {"id": "text", "match_rate": None},
                    {"id": "already", "match_rate": 90},
                ],
                2,
            )
        ),
    )
    service = recruitment_service.RecruitmentService.__new__(
        recruitment_service.RecruitmentService
    )
    service.repo = repo
    service.analyze_resume = AsyncMock(
        return_value={
            "match_rate": 72,
            "resume_score": 70,
            "fit_level": "推荐",
            "reason": "ok",
        }
    )
    service._try_vision_analysis = AsyncMock(
        return_value={"match_rate": 91, "resume_score": 90, "reason": "vision"}
    )

    result = await service.batch_analyze(
        candidate_ids=["text", "vision", "skip", "failed"], job_id="job-1"
    )
    assert result == {"total": 4, "success": 2, "skipped": 1, "failed": 1}
    assert repo.update_candidate.await_count == 2
    assert repo.update_candidate.await_args_list[0].args[1]["fit_level"] == "高"

    empty = await service.batch_analyze(candidate_ids=[])
    assert empty["total"] == 1
