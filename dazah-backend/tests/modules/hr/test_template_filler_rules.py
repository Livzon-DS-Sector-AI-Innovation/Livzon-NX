from __future__ import annotations

from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.modules.hr import template_filler as filler


def _cell(text: str = "") -> Any:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = text
    return cell


def _append_sym(paragraph: Any, option: str, blank: bool = False) -> None:
    run = paragraph.add_run()
    sym = OxmlElement("w:sym")
    sym.set(qn("w:font"), "Wingdings 2")
    sym.set(qn("w:char"), filler.SYM_UNCHECKED)
    run._r.append(sym)
    paragraph.add_run(option)
    if blank:
        paragraph.add_run("   ")


def test_fill_after_label_uses_blank_run_and_append_fallback() -> None:
    cell = _cell()
    paragraph = cell.paragraphs[0]
    paragraph.add_run("培训日期：")
    paragraph.add_run("   ")
    paragraph.add_run("  ")
    filler.fill_after_label(cell, "2026-08-20")
    assert "2026-08-20" in cell.text
    assert paragraph.runs[-1].text == ""

    fallback = _cell("部门：")
    filler.fill_after_label(fallback, "质量部")
    assert fallback.text == "部门：质量部"
    empty = _cell()
    empty.paragraphs[0]._p.clear_content()
    filler.fill_after_label(empty, None)
    assert empty.text == ""


def test_fill_after_phrase_updates_only_matching_area() -> None:
    cell = _cell()
    paragraph = cell.paragraphs[0]
    paragraph.add_run("应到：")
    paragraph.add_run(" ")
    paragraph.add_run("人")
    filler.fill_after_phrase(cell, "应到", 10)
    assert "10" in cell.text
    before = cell.text
    filler.fill_after_phrase(cell, "不存在", 20)
    assert cell.text == before
    append = _cell("实到：")
    filler.fill_after_phrase(append, "实到", 9)
    assert append.text.endswith("9")


def test_whole_cell_rewrite_append_and_paragraph_value() -> None:
    cell = _cell()
    paragraph = cell.paragraphs[0]
    paragraph.add_run("旧")
    paragraph.add_run("值")
    filler.fill_whole_cell(cell, "新值")
    assert cell.text == "新值"
    filler.rewrite_cell_runs(cell, "重写")
    assert cell.text == "重写"
    filler.append_value(cell, None)
    filler.append_value(cell, "追加")
    assert cell.text.endswith("追加")
    filler.replace_text_in_cell(cell, "追加", "替换")
    assert cell.text.endswith("替换")

    doc = Document()
    p = doc.add_paragraph()
    filler.set_paragraph_value(p, "内容")
    assert p.text == "内容"
    p.add_run("残留")
    filler.set_paragraph_value(p, None)
    assert p.text == ""


def test_sym_group_checks_option_and_writes_trailing_value() -> None:
    cell = _cell()
    paragraph = cell.paragraphs[0]
    _append_sym(paragraph, "内训")
    _append_sym(paragraph, "其他：", blank=True)
    filler.set_sym_group(
        cell,
        "其他：外部课程",
        ["内训", "其他"],
        write_trailing={"其他": "外部课程"},
    )
    syms = [run._r.find(qn("w:sym")) for run in paragraph.runs]
    chars = [sym.get(qn("w:char")) for sym in syms if sym is not None]
    assert chars == [filler.SYM_UNCHECKED, filler.SYM_CHECKED]
    assert "外部课程" in cell.text
    filler.set_sym_group(cell, None, ["内训", "其他"])
    syms = [run._r.find(qn("w:sym")) for run in paragraph.runs]
    assert all(
        sym.get(qn("w:char")) == filler.SYM_UNCHECKED for sym in syms if sym is not None
    )


def test_fill_date_paragraph_uses_placeholders_and_fallback() -> None:
    doc = Document()
    p = doc.add_paragraph()
    for value in ("  ", "年", "  ", "月", "  ", "日"):
        p.add_run(value)
    filler.fill_date_paragraph(p, 2026, 8, 20)
    assert "2026年" in p.text
    assert "8月" in p.text
    assert "20日" in p.text
    fallback = doc.add_paragraph("日期")
    filler.fill_date_paragraph(fallback, 2027, 1, 2)
    assert fallback.text == "2027年1月2日"


def test_header_source_clone_and_delete_rows() -> None:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = ""
    table.cell(0, 1).text = "表头"
    table.cell(1, 0).text = "数据"
    assert filler.header_fmt_source(table).text == "表头"
    filler.clone_row_after(table, 1)
    assert len(table.rows) == 3
    filler.delete_row(table, 2)
    assert len(table.rows) == 2


def test_blank_run_and_clone_format_helpers() -> None:
    doc = Document()
    p = doc.add_paragraph()
    source = p.add_run("source")
    source.bold = True
    blank = p.add_run(" ")
    assert filler._is_blank_run(blank)
    target = p.add_run("target")
    filler._clone_rpr_to(source, target)
    assert target.bold is True
    filler._clone_rpr_to(None, target)
    filler._apply_std_font(target)
    assert target.bold is False
    filler._clear_paragraph_except(p, target)
    assert p.text == "target"
