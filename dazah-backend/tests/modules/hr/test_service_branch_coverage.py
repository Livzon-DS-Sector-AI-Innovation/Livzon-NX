from __future__ import annotations

from datetime import date
from io import BytesIO
from types import SimpleNamespace
from unittest.mock import AsyncMock
from uuid import uuid4

import pytest
from docx import Document

from app.modules.hr import service


class _Result:
    def __init__(self, values: list[object] | None = None) -> None:
        self.values = values or []

    def scalars(self) -> _Result:
        return self

    def all(self) -> list[object]:
        return self.values


def _offboarding_record(**overrides: object) -> SimpleNamespace:
    values: dict[str, object] = {
        "id": uuid4(),
        "seq_number": 1,
        "employee_number": "1001",
        "name": "张三",
        "domain_account": "zhangsan",
        "gender": "男",
        "ethnic_group": "汉",
        "native_place": "珠海",
        "political_status": "党员",
        "marital_status": "已婚",
        "health_status": "健康",
        "household_type": "城镇",
        "status_category": "正式",
        "birth_year": 1990,
        "birth_month": 2,
        "birth_day": 3,
        "age": 36,
        "id_card": "440000000000000000",
        "id_card_expiry": "2035-01-01",
        "current_address": "珠海",
        "phone": "13800000000",
        "email": "zhangsan@example.com",
        "emergency_contact_name": "李四",
        "emergency_contact_phone": "13900000000",
        "emergency_contact_relation": "配偶",
        "department": "质量部",
        "sub_department": "QA",
        "position": "质量员",
        "level": "主管",
        "employment_type": "正式",
        "probation_status": "已转正",
        "probation_effective_date": date(2020, 4, 1),
        "hire_date": date(2020, 1, 1),
        "work_start_date": date(2018, 1, 1),
        "factory_entry_date": date(2020, 1, 1),
        "work_years": "6",
        "offboarding_date": date(2026, 8, 27),
        "offboarding_type": "辞职",
        "reason": "个人原因",
        "notes": "已完成交接",
        "handover_status": "已完成",
        "archive_number": "档案-1001",
        "qualifications": ["质量管理体系内审员", "注册安全工程师"],
        "certificate_number": "CERT-1",
        "certificate_review_date": date(2027, 1, 1),
        "education": "本科",
        "degree": "学士",
        "school": "中山大学",
        "major": "药学",
        "graduation_date": date(2012, 6, 1),
        "qualification_type": "工程师",
        "contract_start_date": date(2020, 1, 1),
        "contract_end_date": date(2027, 1, 1),
        "contract_start_2": date(2024, 1, 1),
        "contract_end_2": "2028-01-01",
        "contract_start_3": "2028-01-02",
        "contract_end_3": "2031-01-01",
        "contract_start_4": "2031-01-02",
        "contract_end_4": "2034-01-01",
        "contract_start_5": "2034-01-02",
        "contract_end_5": "2037-01-01",
        "contract_start_6": "2037-01-02",
        "work_experience_1": "制药企业",
        "work_experience_2": "质量管理",
        "work_experience_3": None,
        "work_experience_4": None,
        "feishu_record_id": None,
        "feishu_synced_at": None,
    }
    values.update(overrides)
    return SimpleNamespace(**values)


def _modern_employee(**overrides: object) -> SimpleNamespace:
    values: dict[str, object] = {
        "archive_number": "A-1001",
        "employee_number": "1001",
        "domain_account": "zhangsan",
        "department": "质量部",
        "sub_department": "QA",
        "position": "质量员",
        "name": "张三",
        "level": "主管",
        "gender": "男",
        "ethnic_group": "汉",
        "native_place": "珠海",
        "id_card": "440000000000000000",
        "id_card_expiry": "2035-01-01",
        "marital_status": "已婚",
        "political_status": "党员",
        "current_address": "珠海",
        "household_type": "城镇",
        "education": "本科",
        "degree": "学士",
        "school": "中山大学",
        "major": "药学",
        "graduation_date": date(2012, 6, 1),
        "certificate_number": "CERT-1",
        "qualification_type": "工程师",
        "certificate_review_date": date(2027, 1, 1),
        "work_start_date": date(2018, 1, 1),
        "hire_date": date(2020, 1, 1),
        "factory_entry_date": date(2020, 1, 1),
        "work_years": "6",
        "status": "在职",
        "planned_probation_date": date(2020, 4, 1),
        "probation_status": "已转正",
        "probation_effective_date": date(2020, 4, 1),
        "contract_start_date": date(2020, 1, 1),
        "contract_end_date": date(2027, 1, 1),
        "contract_start_2": date(2024, 1, 1),
        "contract_end_2": "2028-01-01",
        "contract_start_3": date(2028, 1, 1),
        "contract_end_3": "2031-01-01",
        "contract_start_4": date(2031, 1, 1),
        "contract_end_4": "2034-01-01",
        "contract_start_5": date(2034, 1, 1),
        "contract_end_5": "2037-01-01",
        "contract_start_6": "2037-01-01",
        "contract_end_6": "2040-01-01",
        "employment_type": "正式",
        "phone": "13800000000",
        "email": "zhangsan@example.com",
        "emergency_contact_name": "李四",
        "emergency_contact_relation": "配偶",
        "emergency_contact_phone": "13900000000",
        "health_status": "健康",
        "last_working_day": date(2026, 8, 27),
        "work_experience_1": "制药企业",
        "work_experience_2": "质量管理",
        "work_experience_3": None,
        "work_experience_4": None,
        "status_category": "正式",
        "offboarding_type": None,
        "offboarding_reason": None,
        "qualifications": ["内审员"],
        "remarks": "保留备注",
    }
    values.update(overrides)
    return SimpleNamespace(**values)


@pytest.mark.asyncio
async def test_offboarding_feishu_sync_builds_all_fields_and_updates() -> None:
    record = _offboarding_record()
    employee = SimpleNamespace(
        status="离职",
        planned_probation_date=date(2020, 4, 1),
    )
    client = SimpleNamespace(
        create_record=AsyncMock(return_value={"record_id": "rec-new"}),
        update_record=AsyncMock(),
    )
    instance = service.OffboardingRecordService.__new__(
        service.OffboardingRecordService
    )
    instance.repo = SimpleNamespace(update=AsyncMock())
    instance._get_offboarding_bitable = AsyncMock(return_value=(client, "tbl-off"))

    await instance._sync_to_feishu(record, employee, is_create=True)
    assert record.feishu_record_id == "rec-new"
    client.create_record.assert_awaited_once()
    fields = client.create_record.await_args.args[1]
    assert fields["工号"] == 1001
    assert fields["出生年月"] == "1990-02-03"
    assert fields["技能证书"] == "质量管理体系内审员、注册安全工程师"
    instance.repo.update.assert_awaited_once_with(record)

    record.feishu_record_id = "rec-old"
    await instance._sync_to_feishu(record, employee, is_create=False)
    client.update_record.assert_awaited_once()
    assert client.update_record.await_args.args[:2] == ("tbl-off", "rec-old")


@pytest.mark.asyncio
async def test_offboarding_feishu_delete_handles_missing_and_errors() -> None:
    instance = service.OffboardingRecordService.__new__(
        service.OffboardingRecordService
    )
    record = _offboarding_record(feishu_record_id=None)
    instance._get_offboarding_bitable = AsyncMock()
    await instance._delete_from_feishu(record)
    instance._get_offboarding_bitable.assert_not_awaited()

    record.feishu_record_id = "rec-1"
    client = SimpleNamespace(
        delete_record=AsyncMock(side_effect=RuntimeError("failed"))
    )
    instance._get_offboarding_bitable.return_value = (client, "tbl-off")
    with pytest.raises(RuntimeError):
        await instance._delete_from_feishu(record)
    client.delete_record.assert_awaited_once_with("tbl-off", "rec-1")


@pytest.mark.asyncio
async def test_position_transfer_notifications_cover_send_dedupe_and_rejection(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    record = SimpleNamespace(
        id=uuid4(),
        employee_name="张三",
        department_before="质量部",
        apply_department="生产部",
        original_position="质量员",
        apply_position="主管",
        approval_flow={
            "current_step": 0,
            "steps": [
                {
                    "node": "manager",
                    "label": "部门经理",
                    "signer": "李四",
                    "signer_open_id": "ou-manager",
                }
            ],
        },
        feishu_approval_message_id=None,
    )
    cache_get = AsyncMock(return_value=None)
    cache_set = AsyncMock()
    monkeypatch.setattr("app.core.redis.cache_get", cache_get)
    monkeypatch.setattr("app.core.redis.cache_set", cache_set)
    sender = AsyncMock(return_value="msg-1")
    monkeypatch.setattr(
        "app.platform.integrations.feishu.notification.send_user_card_with_message_id",
        sender,
    )
    instance = service.PositionTransferRecordService.__new__(
        service.PositionTransferRecordService
    )
    await instance._notify_next_approver(record)
    assert record.feishu_approval_message_id == "msg-1"
    cache_set.assert_awaited_once_with(
        f"hr:position_transfer:notify:{record.id}:0", "1", ex=3600
    )
    sender.assert_awaited_once()

    cache_get.return_value = "1"
    await instance._notify_next_approver(record)
    assert sender.await_count == 1

    record.approval_flow["steps"][0]["signer_open_id"] = ""
    cache_get.return_value = None
    await instance._notify_next_approver(record)
    assert cache_set.await_count == 1

    rejected_sender = AsyncMock()
    monkeypatch.setattr(
        "app.platform.integrations.feishu.notification.send_user_card_with_message_id",
        rejected_sender,
    )
    instance._get_open_id_by_name = AsyncMock(return_value="ou-applicant")
    await instance._notify_applicant_rejected(record)
    rejected_sender.assert_awaited_once()

    instance._get_open_id_by_name.return_value = None
    await instance._notify_applicant_rejected(record)
    assert rejected_sender.await_count == 1


def test_employee_bitable_fields_cover_modern_and_legacy_shapes() -> None:
    modern = _modern_employee()
    instance = service.EmployeeService.__new__(service.EmployeeService)
    fields = instance._to_bitable_fields(modern)
    assert fields["工号"] == 1001
    assert fields["技能证书"] == ["内审员"]
    assert fields["备注"] == "保留备注"


@pytest.mark.asyncio
async def test_employee_sync_and_legacy_sync_paths(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    employee = _modern_employee(feishu_record_id=None)
    repo = SimpleNamespace(update=AsyncMock(), get_by_employee_number=AsyncMock())
    instance = service.EmployeeService.__new__(service.EmployeeService)
    instance.repo = repo
    instance.session = SimpleNamespace()
    bitable = SimpleNamespace(
        find_by_employee_number=AsyncMock(return_value=None),
        create=AsyncMock(return_value="rec-created"),
        update=AsyncMock(),
    )
    instance._get_bitable = AsyncMock(return_value=bitable)
    assert await instance._sync_single_to_feishu(employee) == "rec-created"
    assert employee.feishu_record_id == "rec-created"

    employee.feishu_record_id = "rec-existing"
    assert await instance._sync_single_to_feishu(employee) == "rec-existing"
    bitable.update.assert_awaited()

    bitable.find_by_employee_number.return_value = SimpleNamespace(
        record_id="rec-found"
    )
    employee.feishu_record_id = None
    assert await instance._sync_single_to_feishu(employee) == "rec-found"

    legacy = service.EmployeeService.__new__(service.EmployeeService)
    legacy.repo = SimpleNamespace(
        upsert_by_employee_number=AsyncMock(),
        get_by_employee_number=AsyncMock(return_value=SimpleNamespace(created_at=None)),
    )
    legacy.bitable = SimpleNamespace(
        table_id="tbl-employee",
        client=SimpleNamespace(
            search_records=AsyncMock(
                return_value=[
                    {"record_id": "r1", "fields": {"工号": "1001", "姓名": "张三"}},
                    {"record_id": "r2", "fields": {"姓名": "无工号"}},
                ]
            )
        ),
    )
    result = await legacy.sync_from_feishu()
    assert result["total"] == 2
    assert result["failed"] == 1


@pytest.mark.asyncio
async def test_training_members_import_and_training_plan_docx_import(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    rows = [
        SimpleNamespace(name="张三", employee_no="1001", department="质量部"),
        SimpleNamespace(name="公用账号", employee_no=None, department="质量部"),
        SimpleNamespace(name="李四", employee_no="1002", department="冻结部门"),
    ]
    session = SimpleNamespace(execute=AsyncMock(return_value=_Result(rows)))
    member_repo = SimpleNamespace(upsert_member=AsyncMock())
    member_instance = service.EmployeeTrainingListService.__new__(
        service.EmployeeTrainingListService
    )
    member_instance.session = session
    member_instance.member_repo = member_repo
    monkeypatch.setattr(
        "app.modules.hr.training_dept_resolver._load_mappings",
        AsyncMock(
            return_value=[
                {
                    "mapping_type": "candidate_source",
                    "source_name": "注册部",
                    "target_name": "质量部",
                }
            ]
        ),
    )
    monkeypatch.setattr(
        "app.modules.hr.training_dept_resolver.resolve_training_department",
        AsyncMock(return_value="质量部"),
    )
    result = await member_instance.import_feishu_members(department="注册部")
    assert result["total"] == 1
    member_repo.upsert_member.assert_awaited_once_with(
        department="注册部", name="张三", employee_number="1001", source="feishu"
    )

    doc = Document()
    doc.add_paragraph("2026年度部门培训计划表")
    doc.add_paragraph("部门：质量部 版本：V2")
    table = doc.add_table(rows=2, cols=7)
    headers = [
        "序号",
        "培训类型",
        "培训时间",
        "培训内容",
        "培训对象",
        "授课人",
        "考核方式",
    ]
    values = ["1", "☑内训", "1月", "GMP基础", "新员工", "王老师", "考试"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for cell, value in zip(table.rows[1].cells, values):
        cell.text = value

    plan = SimpleNamespace(id=uuid4())
    plan_repo = SimpleNamespace(
        get_by_year_and_department=AsyncMock(return_value=None),
        create=AsyncMock(return_value=plan),
    )
    plan_instance = service.AnnualTrainingPlanService.__new__(
        service.AnnualTrainingPlanService
    )
    plan_instance.session = SimpleNamespace()
    plan_instance.repo = plan_repo
    item_service = SimpleNamespace(batch_update_items=AsyncMock())
    content = BytesIO()
    doc.save(content)
    imported = await plan_instance.import_from_docx(
        content.getvalue(),
        year=2026,
        plan_level=None,
        department=None,
        item_service=item_service,
    )
    assert imported == {"plan_id": str(plan.id), "imported_count": 1}
    item_service.batch_update_items.assert_awaited_once()


@pytest.mark.asyncio
async def test_training_conflict_ignores_invalid_sessions_and_supports_exclusion() -> (
    None
):
    training_date = date(2026, 8, 27)
    current_session_id = uuid4()
    ledger = SimpleNamespace(
        training_datetime="2026.08.27 09:00~10:00",
        training_subject="GMP",
        instructor="王老师",
        teaching_dept="质量部",
        trainees="张三、李四",
    )
    sessions = [
        SimpleNamespace(
            id=current_session_id,
            time_start="invalid",
            time_end="10:00",
            topic="坏数据",
            department="质量部",
            instructor="王老师",
            employee_names=["张三"],
        ),
        SimpleNamespace(
            id=uuid4(),
            time_start="14:00",
            time_end="15:00",
            topic="下午培训",
            department="生产部",
            instructor="李老师",
            employee_names="非数组",
        ),
    ]
    result = _Result(sessions)
    instance = service.TrainingLedgerService.__new__(service.TrainingLedgerService)
    instance.repo = SimpleNamespace(list_by_date=AsyncMock(return_value=[ledger]))
    instance.session = SimpleNamespace(execute=AsyncMock(return_value=result))
    conflict = await instance.check_conflict(
        training_date,
        "11:00",
        "12:00",
        "王老师",
        ["张三"],
        exclude_session_id=str(current_session_id),
    )
    assert conflict["has_conflict"] is False
