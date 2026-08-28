from __future__ import annotations

from io import BytesIO

from docx import Document
from openpyxl import Workbook  # type: ignore[import-untyped]

from app.modules.hr import attachment_parser as parser


def _xlsx_bytes(*, summary_only: bool = False) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "全部附件" if summary_only else "附件一 培训签到"
    ws.append(["附件一"] if summary_only else ["培训签到表"])
    ws.append(["序号", "名称", "结果"])
    ws.append([1, "张三", "合格"])
    if summary_only:
        ws.append(["附件二"])
        ws.append(["序号", "名称"])
        ws.append([2, "李四"])
    else:
        ws2 = wb.create_sheet("其他")
        ws2.append(["附件2 培训记录"])
        ws2.append(["名称", "日期"])
        ws2.append(["课程", "2026-08-20"])
    stream = BytesIO()
    wb.save(stream)
    wb.close()
    return stream.getvalue()


def _docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("年度培训材料")
    doc.add_paragraph("附件一 培训签到")
    doc.add_paragraph("张三 已签到")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "结果"
    table.cell(1, 0).text = "张三"
    table.cell(1, 1).text = "合格"
    doc.add_paragraph("附件2 考核记录")
    doc.add_paragraph("李四 90分")
    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


def test_annex_normalization_and_reference_extraction() -> None:
    assert parser._cn_to_int("") is None
    assert parser._cn_to_int("12") == 12
    assert parser._cn_to_int("一") == 1
    assert parser._cn_to_int("十") == 10
    assert parser._cn_to_int("二十一") == 21
    assert parser.normalize_annex_no("参见附件１２") == "附件12"
    assert parser.normalize_annex_no("无引用") is None
    assert parser.extract_annex_refs("附件一、附件2、附件一") == ["附件1", "附件2"]
    assert parser.extract_annex_refs(None) == []
    assert parser.strip_punct("附件（1）：签到-表") == "附件签到表"
    assert parser._file_ext("FILE.XLSX") == "xlsx"
    assert parser._file_ext("README") == ""


def test_parse_xlsx_sections_from_sheet_names_and_first_cells() -> None:
    drafts = parser.parse_sections("plan.xlsx", _xlsx_bytes())
    assert [(d.annex_no, d.source_ref) for d in drafts] == [
        ("附件1", "附件一 培训签到"),
        ("附件2", "其他"),
    ]
    summary = parser.parse_sections("plan.xlsx", _xlsx_bytes(summary_only=True))
    assert [d.annex_no for d in summary] == ["附件1", "附件2"]
    assert parser.parse_sections("plan.pdf", b"pdf") == []


def test_parse_docx_sections_and_build_outlines() -> None:
    data = _docx_bytes()
    drafts = parser.parse_sections("plan.docx", data)
    assert [d.annex_no for d in drafts] == ["附件1", "附件2"]
    outline = parser.build_outline("plan.docx", data)
    assert outline["kind"] == "docx"
    assert any(line["text"].startswith("附件一") for line in outline["lines"])
    xlsx_outline = parser.build_outline("plan.xlsx", _xlsx_bytes())
    assert xlsx_outline["kind"] == "xlsx"
    assert len(xlsx_outline["sheets"]) == 2
    assert parser.build_outline("plan.pdf", b"pdf") == {"kind": "unknown"}


def test_xlsx_preview_supports_single_sheet_and_whole_workbook() -> None:
    data = _xlsx_bytes()
    single = parser.build_preview("plan.xlsx", data, "xlsx_sheet", "附件一 培训签到")
    assert single["kind"] == "table"
    assert single["header"] == ["序号", "名称", "结果"]
    assert single["rows"][0] == ["1", "张三", "合格"]
    whole = parser.build_preview("plan.xlsx", data, "whole_file", None)
    assert whole["kind"] == "tables"
    assert len(whole["tables"]) == 2


def test_docx_preview_supports_section_and_whole_document() -> None:
    data = _docx_bytes()
    drafts = parser.parse_sections("plan.docx", data)
    section = parser.build_preview(
        "plan.docx", data, "docx_section", drafts[0].source_ref
    )
    assert section["title"].startswith("附件一")
    assert any(block["type"] == "table" for block in section["blocks"])
    assert not any(
        block.get("text", "").startswith("附件2") for block in section["blocks"]
    )
    whole = parser.build_preview("plan.docx", data, "whole_file", None)
    assert whole["kind"] == "doc"
    assert any(block["type"] == "table" for block in whole["blocks"])


def test_unsupported_preview_returns_download_message() -> None:
    preview = parser.build_preview("scan.pdf", b"pdf", "whole_file", None)
    assert preview["kind"] == "doc"
    assert "下载" in preview["blocks"][0]["text"]
