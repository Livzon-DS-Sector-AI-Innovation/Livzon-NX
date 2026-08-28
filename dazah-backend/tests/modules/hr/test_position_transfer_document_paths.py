from __future__ import annotations

from datetime import date
from io import BytesIO
from types import SimpleNamespace
from typing import Any
from unittest.mock import AsyncMock
from uuid import uuid4

import pytest
from docx import Document

from app.core.exceptions import AppException
from app.modules.hr import service


def _template_bytes() -> bytes:
    document = Document()
    table = document.add_table(rows=12, cols=8)

    def put_runs(cell: Any, *values: str) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        for value in values:
            paragraph.add_run(value)

    for row in table.rows:
        for cell in row.cells:
            cell.text = "占位"
    put_runs(table.cell(2, 0), "签名：", "    ", "日期：")
    for row_index in (3, 4, 5, 6, 8, 10, 11):
        put_runs(table.cell(row_index, 0), "签名：", "    ", "日期：")
    put_runs(table.cell(7, 0), "签名：")
    put_runs(table.cell(7, 3), "日期：")
    put_runs(table.cell(9, 0), "签名：")
    put_runs(table.cell(9, 2), "日期：")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


@pytest.mark.asyncio
async def test_position_transfer_document_fills_all_approval_rows() -> None:
    template = SimpleNamespace(template_data=_template_bytes())
    instance = service.PositionTransferRecordService.__new__(
        service.PositionTransferRecordService
    )
    instance.session = SimpleNamespace(
        execute=AsyncMock(
            return_value=SimpleNamespace(scalar_one_or_none=lambda: template)
        )
    )
    steps = [
        {"node": node, "signer": f"{node}负责人", "date": "2026-08-26"}
        for node in (
            "origin_direct_leader",
            "origin_manager",
            "origin_director",
            "origin_vp",
            "target_direct_leader",
            "target_manager",
            "target_director",
            "target_vp",
            "hr",
            "executive_vp",
            "general_manager",
        )
    ]
    record = SimpleNamespace(
        id=uuid4(),
        employee_name="张三",
        department_before="生产部",
        original_position="操作员",
        effective_date=date(2026, 9, 1),
        apply_department="质量部",
        apply_position="质量员",
        contact_phone="13800000000",
        applicant_signature="张三",
        applicant_confirmation_date=date(2026, 8, 26),
        approval_flow={"applicant_date": "2026-08-25", "steps": steps},
    )

    rendered = await instance._fill_word_template(record)

    assert rendered[:2] == b"PK"
    assert len(rendered) > 1000

    instance.session.execute.return_value = SimpleNamespace(
        scalar_one_or_none=lambda: None
    )
    with pytest.raises(AppException, match="模板未配置"):
        await instance._fill_word_template(record)
