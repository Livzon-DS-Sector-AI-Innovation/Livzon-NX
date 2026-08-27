import pytest
from docx import Document

from app.modules.hr import offboarding_document_generator as generator


@pytest.fixture
def termination_template(tmp_path, monkeypatch):
    path = tmp_path / "解除劳动合同单.docx"
    document = Document()
    document.add_paragraph(
        "{姓名} {性别} {身份证号} {入职日期} {现家庭地址} YYYY年MM月DD日"
    )
    document.save(path)
    monkeypatch.setattr(generator, "TEMPLATE_DIR", tmp_path)
    return path


def _extract_all_text(doc) -> str:
    """提取文档中所有文本（段落 + 表格）"""
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)


def test_generate_termination_notice(termination_template):
    employee = {
        "name": "张三",
        "gender": "男",
        "id_card": "640102199001011234",
        "hire_date": "2020年01月01日",
        "current_address": "宁夏平罗县",
    }

    result = generator.generate_termination_notice(employee)

    assert result is not None
    assert result.readable()

    from docx import Document

    doc = Document(result)
    full_text = _extract_all_text(doc)

    assert "张三" in full_text
    assert "640102199001011234" in full_text
    assert "2020年01月01日" in full_text
    assert "宁夏平罗县" in full_text


def test_generate_termination_notice_empty_fields(termination_template):
    employee = {}
    result = generator.generate_termination_notice(employee)
    assert result is not None
    assert result.readable()
