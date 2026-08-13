import json
import uuid
from datetime import UTC, datetime, timedelta
from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document

from app.modules.agent.schemas import (
    AgentBackendV2Event,
    AgentChatRequest,
    AgentToolExecuteRequest,
    AgentTrustedSubject,
)
from app.modules.agent.service import (
    HUMAN_DECISION_REQUIRED_MESSAGE,
    AgentService,
    normalize_agent_basic_command,
)
from app.modules.procurement.contract_generator import generate_contract
from app.modules.procurement.schemas import ContractGenerateRequest


class AllowAllAccessScopeService:
    async def require_tool_access(self, *args, **kwargs):
        return None


class FakeDb:
    committed = False

    def __init__(self) -> None:
        self.added = []

    async def commit(self) -> None:
        self.committed = True

    async def flush(self) -> None:
        return None

    async def scalar(self, _query):
        return None

    async def get(self, model, item_id):
        return SimpleNamespace(
            id=item_id,
            role="admin",
            status="active",
            is_deleted=False,
            tenant_key="test",
        )

    def add(self, item) -> None:
        self.added.append(item)


def _subject() -> AgentTrustedSubject:
    return AgentTrustedSubject(
        tenant_id="test",
        user_id=uuid.uuid4(),
        source="internal",
    )


@pytest.mark.parametrize(
    "command",
    ["/new", "/restart", "/reset", "/新建会话"],
)
def test_session_reset_command_aliases(command: str) -> None:
    assert normalize_agent_basic_command(command) == "new"


def test_basic_help_and_status_commands() -> None:
    assert normalize_agent_basic_command(" /HELP ") == "help"
    assert normalize_agent_basic_command("/status") == "status"
    assert normalize_agent_basic_command("普通问题") is None
    assert normalize_agent_basic_command("/restrat") is None


@pytest.mark.anyio
async def test_restart_command_archives_old_web_session_and_creates_new_one() -> None:
    user_id = uuid.uuid4()
    old_session = SimpleNamespace(
        id=uuid.uuid4(),
        user_id=user_id,
        status="active",
        context={"channel": "web"},
        updated_by=user_id,
    )
    new_session = SimpleNamespace(
        id=uuid.uuid4(),
        user_id=user_id,
        status="active",
        context={"channel": "web"},
        updated_by=user_id,
    )
    archived: list[uuid.UUID] = []
    messages: list[SimpleNamespace] = []

    class ResetRepository:
        async def get_session(self, _db, session_id):
            return old_session if session_id == old_session.id else None

        async def archive_session(self, _db, *, session, user_id):
            session.status = "archived"
            archived.append(session.id)
            return session

        async def create_session(self, _db, *, user_id, context, title):
            new_session.user_id = user_id
            new_session.context = context
            new_session.title = title
            return new_session

        async def add_message(
            self, _db, *, session_id, role, content, metadata, user_id
        ):
            item = SimpleNamespace(
                id=uuid.uuid4(),
                session_id=session_id,
                role=role,
                content=content,
                message_metadata=metadata,
                created_at=datetime.now(UTC),
            )
            messages.append(item)
            return item

        async def list_messages(self, _db, *, session_id):
            return [item for item in messages if item.session_id == session_id]

    service = AgentService(SimpleNamespace())
    service.repo = ResetRepository()
    current_user = SimpleNamespace(id=user_id)

    session, history, _message = await service._prepare_chat_context(
        object(),
        request=AgentChatRequest(
            session_id=old_session.id,
            message="/restart",
            context={"channel": "web"},
        ),
        current_user=current_user,
    )

    assert archived == [old_session.id]
    assert old_session.status == "archived"
    assert session.id == new_session.id
    assert history[-1]["content"] == "/restart"


class FakeAgentRepository:
    def __init__(self) -> None:
        self.session = SimpleNamespace(
            id=uuid.uuid4(),
            user_id=None,
            context={},
            title=None,
        )
        self.messages = []
        self.tool_calls = []
        self.confirmations = {}

    async def create_session(self, db, *, user_id, context, title):
        self.session.user_id = user_id
        self.session.context = context
        self.session.title = title
        return self.session

    async def add_message(
        self,
        db,
        *,
        session_id,
        role,
        content,
        metadata=None,
        user_id=None,
    ):
        message = SimpleNamespace(
            id=uuid.uuid4(),
            session_id=session_id,
            role=role,
            content=content,
            message_metadata=metadata or {},
            created_at=datetime.now(UTC),
            user_id=user_id,
        )
        self.messages.append(message)
        return message

    async def list_messages(self, db, *, session_id, limit=20):
        return self.messages[-limit:]

    async def list_session_attachments(self, db, *, session_id, user_id, limit=50):
        return []

    async def create_tool_call(
        self,
        db,
        *,
        session_id,
        operation,
        request_payload,
    ):
        call = SimpleNamespace(
            session_id=session_id,
            operation=operation,
            request_payload=request_payload,
            status="started",
            response_payload=None,
            error_message=None,
        )
        self.tool_calls.append(call)
        return call

    async def finish_tool_call(
        self,
        db,
        call,
        *,
        status,
        response_payload=None,
        error_message=None,
    ):
        call.status = status
        call.response_payload = response_payload
        call.error_message = error_message
        return call

    async def create_confirmation(
        self,
        db,
        *,
        session_id,
        user_id,
        operation,
        summary,
        risk_level,
        request_payload,
        expires_at,
    ):
        confirmation = SimpleNamespace(
            id=uuid.uuid4(),
            session_id=session_id,
            user_id=user_id,
            operation=operation,
            summary=summary,
            risk_level=risk_level,
            status="pending",
            request_payload=request_payload,
            expires_at=expires_at,
        )
        self.confirmations[confirmation.id] = confirmation
        return confirmation

    async def get_confirmation(self, db, confirmation_id):
        return self.confirmations.get(confirmation_id)

    async def get_confirmation_for_update(self, db, confirmation_id):
        return self.confirmations.get(confirmation_id)

    async def mirror_external_confirmation(self, db, **values):
        confirmation_id = values.pop("confirmation_id")
        existing = self.confirmations.get(confirmation_id)
        if existing is not None:
            return existing
        confirmation = SimpleNamespace(
            id=confirmation_id,
            status="pending",
            **values,
        )
        self.confirmations[confirmation_id] = confirmation
        return confirmation

    async def finish_external_confirmation(
        self,
        db,
        confirmation,
        *,
        status,
        result_payload,
        user_id,
    ):
        confirmation.status = status
        confirmation.result_payload = result_payload
        confirmation.updated_by = user_id
        return confirmation

    async def expire_confirmation(self, db, confirmation, *, user_id):
        confirmation.status = "expired"
        confirmation.updated_by = user_id
        return confirmation

    async def list_pending_confirmations(
        self,
        db,
        *,
        session_id=None,
        user_id=None,
    ):
        now = datetime.now(UTC)
        return [
            confirmation
            for confirmation in self.confirmations.values()
            if confirmation.status == "pending"
            and confirmation.expires_at > now
            and (session_id is None or confirmation.session_id == session_id)
            and (user_id is None or confirmation.user_id == user_id)
        ]


class PolicyOnlyAgentService(AgentService):
    async def _call_hermes(self, **kwargs):
        raise AssertionError("policy-blocked messages must not reach Hermes")


@pytest.mark.anyio
async def test_expired_confirmation_is_returned_as_terminal_state() -> None:
    repo = FakeAgentRepository()
    service = AgentService(settings=SimpleNamespace(), repo=repo)
    user = SimpleNamespace(id=uuid.uuid4())
    confirmation = await repo.create_confirmation(
        FakeDb(),
        session_id=repo.session.id,
        user_id=user.id,
        operation="identity.deliver_feishu_message",
        summary="过期投递确认",
        risk_level="medium",
        request_payload={},
        expires_at=datetime.now(UTC) - timedelta(seconds=1),
    )

    resolved, result = await service.execute_confirmation(
        FakeDb(),
        confirmation_id=confirmation.id,
        current_user=user,
    )

    assert result is None
    assert resolved.status == "expired"
    assert resolved.updated_by == user.id


class StreamingAgentService(AgentService):
    async def _call_hermes_stream(self, **kwargs):
        trace_id = uuid.uuid4()
        run_id = uuid.uuid4()
        yield AgentBackendV2Event(
            trace_id=trace_id,
            run_id=run_id,
            sequence=1,
            type="accepted",
            data={},
        )
        yield AgentBackendV2Event(
            trace_id=trace_id,
            run_id=run_id,
            sequence=2,
            type="text_delta",
            data={"text": "你好，"},
        )
        yield AgentBackendV2Event(
            trace_id=trace_id,
            run_id=run_id,
            sequence=3,
            type="ping",
            data={"ts": 1},
        )
        yield AgentBackendV2Event(
            trace_id=trace_id,
            run_id=run_id,
            sequence=4,
            type="text_delta",
            data={"text": "我正在查询。"},
        )
        yield AgentBackendV2Event(
            trace_id=trace_id,
            run_id=run_id,
            sequence=5,
            type="finished",
            data={
                "message": "你好，我正在查询。",
                "pending_confirmations": [],
                "tool_trace": [{"tool": "search"}],
            },
        )


class ErrorStreamingAgentService(AgentService):
    async def _call_hermes_stream(self, **kwargs):
        trace_id = uuid.uuid4()
        run_id = uuid.uuid4()
        yield AgentBackendV2Event(
            trace_id=trace_id,
            run_id=run_id,
            sequence=1,
            type="accepted",
            data={},
        )
        yield AgentBackendV2Event(
            trace_id=trace_id,
            run_id=run_id,
            sequence=2,
            type="text_delta",
            data={"text": "处理中"},
        )
        yield AgentBackendV2Event(
            trace_id=trace_id,
            run_id=run_id,
            sequence=3,
            type="error",
            data={"message": "上游中断"},
        )


class CrashingStreamingAgentService(AgentService):
    async def _call_hermes_stream(self, **kwargs):
        raise ValueError("internal contract mismatch")
        yield  # pragma: no cover


def parse_sse_events(frames: list[str]) -> list[tuple[str, dict]]:
    events = []
    for frame in frames:
        event = "message"
        data_lines = []
        for line in frame.strip().splitlines():
            if line.startswith("event:"):
                event = line.removeprefix("event:").strip()
            elif line.startswith("data:"):
                data_lines.append(line.removeprefix("data:").strip())
        if data_lines:
            events.append((event, json.loads("\n".join(data_lines))))
    return events


def test_human_decision_policy_detects_delegated_approval() -> None:
    assert AgentService._is_human_decision_required_message("帮我审批通过这条采购申请")
    assert AgentService._is_human_decision_required_message("直接驳回这个申请")
    assert AgentService._is_human_decision_required_message("请重启仓储飞书 WebSocket")
    assert AgentService._is_human_decision_required_message("请批准这条采购申请")


def test_human_decision_policy_allows_read_only_approval_queries() -> None:
    assert not AgentService._is_human_decision_required_message("查看待审批采购申请")
    assert not AgentService._is_human_decision_required_message(
        "给我查询待审批采购申请明细"
    )


def test_human_decision_policy_allows_confirmed_feishu_message_send() -> None:
    assert not AgentService._is_human_decision_required_message(
        "请通过飞书向但昊发送交互卡片，等待我点击确认执行后再发送"
    )
    assert not AgentService._is_human_decision_required_message(
        "此操作需要先生成确认卡片，不要直接执行"
    )
    assert not AgentService._is_human_decision_required_message(
        "如果我同意，请生成待确认项"
    )


def test_assistant_metadata_exposes_structured_query_evidence() -> None:
    metadata = AgentService._assistant_metadata(
        {
            "tool_trace": [
                {
                    "operation": "quality.list_deviations",
                    "params": {"status": "open", "page_size": 5},
                    "ok": True,
                }
            ]
        },
        {"scope": ["quality"]},
    )

    evidence = metadata["evidence"]
    assert evidence["sources"] == [
        {
            "operation": "quality.list_deviations",
            "module": "quality",
            "params": {"status": "open", "page_size": 5},
            "ok": True,
        }
    ]
    assert evidence["scope"] == ["quality"]
    assert evidence["truncated"] is True


@pytest.mark.anyio
async def test_chat_returns_policy_refusal_for_delegated_approval() -> None:
    repo = FakeAgentRepository()
    service = PolicyOnlyAgentService(settings=SimpleNamespace(), repo=repo)
    user = SimpleNamespace(id=uuid.uuid4(), name="测试用户")
    db = FakeDb()

    response = await service.chat(
        db,
        request=AgentChatRequest(message="帮我审批通过这条采购申请"),
        current_user=user,
    )

    assert response.message.content == HUMAN_DECISION_REQUIRED_MESSAGE
    assert response.message.metadata["policy"] == "human_decision_required"
    assert response.pending_confirmations == []
    assert db.committed is True


@pytest.mark.anyio
async def test_stream_chat_emits_agent_backend_v2_events() -> None:
    repo = FakeAgentRepository()
    service = StreamingAgentService(settings=SimpleNamespace(), repo=repo)
    user = SimpleNamespace(id=uuid.uuid4(), name="测试用户")
    db = FakeDb()

    stream = service.stream_chat(
        db,
        request=AgentChatRequest(message="查一下库存"),
        current_user=user,
    )
    first_frame = await anext(stream)

    assert db.committed is True

    frames = [first_frame, *[frame async for frame in stream]]
    events = parse_sse_events(frames)

    assert [event for event, _ in events] == [
        "accepted",
        "text_delta",
        "ping",
        "text_delta",
        "finished",
    ]
    assert events[0][1]["sequence"] == 1
    assert events[0][1]["data"]["session_id"] == str(repo.session.id)
    assert events[1][1]["data"]["text"] == "你好，"
    assert events[3][1]["data"]["text"] == "我正在查询。"
    assert events[-1][1]["data"]["message"]["content"] == "你好，我正在查询。"
    assert events[-1][1]["data"]["tool_trace"] == [{"tool": "search"}]
    assert db.committed is True


@pytest.mark.anyio
async def test_stream_chat_persists_partial_response_when_client_stops() -> None:
    repo = FakeAgentRepository()
    service = StreamingAgentService(settings=SimpleNamespace(), repo=repo)
    user = SimpleNamespace(id=uuid.uuid4(), name="测试用户")
    db = FakeDb()

    stream = service.stream_chat(
        db,
        request=AgentChatRequest(message="生成长回复"),
        current_user=user,
    )
    await anext(stream)  # accepted
    await anext(stream)  # first text delta
    await stream.aclose()

    assert [message.role for message in repo.messages] == ["user", "assistant"]
    assert repo.messages[-1].content == "你好，"
    assert repo.messages[-1].message_metadata["generation_status"] == "stopped"
    assert db.committed is True


@pytest.mark.anyio
async def test_stream_chat_error_does_not_store_assistant_message() -> None:
    repo = FakeAgentRepository()
    service = ErrorStreamingAgentService(settings=SimpleNamespace(), repo=repo)
    user = SimpleNamespace(id=uuid.uuid4(), name="测试用户")
    db = FakeDb()

    frames = [
        frame
        async for frame in service.stream_chat(
            db,
            request=AgentChatRequest(message="查一下库存"),
            current_user=user,
        )
    ]
    events = parse_sse_events(frames)

    assert [event for event, _ in events] == [
        "accepted",
        "text_delta",
        "error",
    ]
    assert events[-1][1]["data"]["message"] == "上游中断"
    assert [message.role for message in repo.messages] == ["user"]
    assert db.committed is True


@pytest.mark.anyio
async def test_stream_chat_pipeline_exception_emits_stable_error_event() -> None:
    repo = FakeAgentRepository()
    service = CrashingStreamingAgentService(settings=SimpleNamespace(), repo=repo)
    user = SimpleNamespace(id=uuid.uuid4(), name="测试用户")

    frames = [
        frame
        async for frame in service.stream_chat(
            FakeDb(),
            request=AgentChatRequest(message="分析附件"),
            current_user=user,
        )
    ]
    events = parse_sse_events(frames)

    assert [event for event, _ in events] == ["error"]
    assert events[0][1]["data"]["code"] == "agent.internal_error"
    assert "contract mismatch" not in events[0][1]["data"]["message"]
    assert [message.role for message in repo.messages] == ["user"]


@pytest.mark.anyio
async def test_stream_chat_policy_refusal_returns_v2_finished_without_hermes() -> None:
    repo = FakeAgentRepository()
    service = PolicyOnlyAgentService(settings=SimpleNamespace(), repo=repo)
    user = SimpleNamespace(id=uuid.uuid4(), name="测试用户")
    db = FakeDb()

    frames = [
        frame
        async for frame in service.stream_chat(
            db,
            request=AgentChatRequest(message="帮我审批通过这条采购申请"),
            current_user=user,
        )
    ]
    events = parse_sse_events(frames)

    assert [event for event, _ in events] == ["accepted", "finished"]
    assert events[-1][1]["data"]["message"]["content"] == (
        HUMAN_DECISION_REQUIRED_MESSAGE
    )
    assert events[-1][1]["data"]["pending_confirmations"] == []
    assert db.committed is True


@pytest.mark.anyio
async def test_high_risk_tool_execution_returns_policy_refusal() -> None:
    repo = FakeAgentRepository()
    service = AgentService(
        settings=SimpleNamespace(),
        repo=repo,
        access_scope_service=AllowAllAccessScopeService(),
    )

    response = await service.execute_tool(
        FakeDb(),
        request=AgentToolExecuteRequest(
            operation="procurement.approve_purchase_request",
            params={"request_id": str(uuid.uuid4())},
            subject=_subject(),
        ),
    )

    assert response.ok is False
    assert response.requires_confirmation is False
    assert response.confirmation is None
    assert response.meta["policy"] == "human_decision_required"
    assert response.data["message"] == HUMAN_DECISION_REQUIRED_MESSAGE
    assert repo.tool_calls[0].status == "rejected_by_policy"


@pytest.mark.anyio
async def test_resolve_pending_confirmation_from_assistant_text_id() -> None:
    repo = FakeAgentRepository()
    service = AgentService(
        settings=SimpleNamespace(),
        repo=repo,
        access_scope_service=AllowAllAccessScopeService(),
    )
    confirmation = await repo.create_confirmation(
        FakeDb(),
        session_id=repo.session.id,
        user_id=uuid.uuid4(),
        operation="identity.deliver_feishu_message",
        summary="向但昊发送飞书卡片消息",
        risk_level="medium",
        request_payload={"user_ids": [str(uuid.uuid4())]},
        expires_at=datetime.now(UTC) + timedelta(minutes=5),
    )

    pending = await service._resolve_pending_confirmations(
        FakeDb(),
        {
            "message": (
                "已生成确认项，需要你确认后执行。\n"
                f"发送确认编号：{confirmation.id} 状态：待确认"
            ),
            "pending_confirmations": [],
            "tool_trace": [],
        },
    )

    assert pending == [confirmation]


@pytest.mark.anyio
async def test_resolve_multiple_pending_confirmations_from_same_result() -> None:
    repo = FakeAgentRepository()
    service = AgentService(settings=SimpleNamespace(), repo=repo)
    first_confirmation = await repo.create_confirmation(
        FakeDb(),
        session_id=repo.session.id,
        user_id=uuid.uuid4(),
        operation="identity.deliver_feishu_message",
        summary="向张三发送飞书文本消息",
        risk_level="medium",
        request_payload={"user_ids": [str(uuid.uuid4())]},
        expires_at=datetime.now(UTC) + timedelta(minutes=5),
    )
    second_confirmation = await repo.create_confirmation(
        FakeDb(),
        session_id=repo.session.id,
        user_id=uuid.uuid4(),
        operation="identity.deliver_feishu_message",
        summary="向李四发送飞书卡片消息",
        risk_level="medium",
        request_payload={"user_ids": [str(uuid.uuid4())]},
        expires_at=datetime.now(UTC) + timedelta(minutes=5),
    )

    pending = await service._resolve_pending_confirmations(
        FakeDb(),
        {
            "message": (
                "已生成两个确认项。\n"
                f"第一个确认编号：{first_confirmation.id}\n"
                f"第二个确认编号：{second_confirmation.id}"
            ),
            "pending_confirmations": [
                {"id": str(first_confirmation.id)},
                {"id": str(second_confirmation.id)},
            ],
            "tool_trace": [],
        },
    )

    assert pending == [first_confirmation, second_confirmation]


@pytest.mark.anyio
async def test_resolve_pending_confirmations_includes_session_pending() -> None:
    repo = FakeAgentRepository()
    service = AgentService(settings=SimpleNamespace(), repo=repo)
    first_confirmation = await repo.create_confirmation(
        FakeDb(),
        session_id=repo.session.id,
        user_id=uuid.uuid4(),
        operation="identity.deliver_feishu_message",
        summary="向张三发送飞书文本消息",
        risk_level="medium",
        request_payload={"user_ids": [str(uuid.uuid4())]},
        expires_at=datetime.now(UTC) + timedelta(minutes=5),
    )
    second_confirmation = await repo.create_confirmation(
        FakeDb(),
        session_id=repo.session.id,
        user_id=uuid.uuid4(),
        operation="identity.deliver_feishu_message",
        summary="向李四发送飞书卡片消息",
        risk_level="medium",
        request_payload={"user_ids": [str(uuid.uuid4())]},
        expires_at=datetime.now(UTC) + timedelta(minutes=5),
    )

    pending = await service._resolve_pending_confirmations(
        FakeDb(),
        {
            "message": f"已生成最新确认项：{second_confirmation.id}",
            "pending_confirmations": [{"id": str(second_confirmation.id)}],
            "tool_trace": [],
        },
        session_id=repo.session.id,
    )

    assert pending == [first_confirmation, second_confirmation]


@pytest.mark.anyio
async def test_feishu_native_confirmation_is_mirrored_without_cli_payload() -> None:
    repo = FakeAgentRepository()
    service = AgentService(settings=SimpleNamespace(), repo=repo)
    user_id = uuid.uuid4()
    confirmation_id = uuid.uuid4()

    pending = await service._resolve_pending_confirmations(
        FakeDb(),
        {
            "pending_confirmations": [
                {
                    "id": str(confirmation_id),
                    "operation": "sheets +cells-set",
                    "summary": "修改飞书资源",
                    "risk_level": "medium",
                    "status": "pending",
                    "expires_at": (
                        datetime.now(UTC) + timedelta(minutes=5)
                    ).isoformat(),
                    "resource_domain": "feishu_native",
                    "resource": "飞书资源 …ample",
                    "impact_count": 2,
                    "reason": "修改既有值",
                    "preview": "不得复制到平台的正文",
                }
            ]
        },
        session_id=repo.session.id,
        user_id=user_id,
    )

    assert [item.id for item in pending] == [confirmation_id]
    mirrored = pending[0].request_payload
    assert mirrored["resource_domain"] == "feishu_native"
    assert mirrored["remote_confirmation_id"] == str(confirmation_id)
    assert mirrored["impact_count"] == 2
    assert "preview" not in mirrored


@pytest.mark.anyio
async def test_execute_feishu_native_confirmation_uses_remote_result(
    monkeypatch,
) -> None:
    repo = FakeAgentRepository()
    service = AgentService(settings=SimpleNamespace(), repo=repo)
    user = SimpleNamespace(id=uuid.uuid4())
    confirmation = await repo.create_confirmation(
        FakeDb(),
        session_id=None,
        user_id=user.id,
        operation="base +record-delete",
        summary="删除飞书记录",
        risk_level="high",
        request_payload={
            "resource_domain": "feishu_native",
            "remote_confirmation_id": str(uuid.uuid4()),
        },
        expires_at=datetime.now(UTC) + timedelta(minutes=5),
    )

    async def resolve_remote(*args, **kwargs):
        return {"ok": False, "status": "verification_failed"}

    monkeypatch.setattr(service, "_resolve_feishu_native_confirmation", resolve_remote)
    resolved, result = await service.execute_confirmation(
        FakeDb(),
        confirmation_id=confirmation.id,
        current_user=user,
    )

    assert resolved.status == "failed"
    assert result is not None and result.ok is False
    assert result.data["status"] == "verification_failed"


def test_contract_generation_sample_request_is_normalized() -> None:
    service = AgentService(settings=SimpleNamespace())

    request = service._normalize_tool_request(
        AgentToolExecuteRequest(
            operation="procurement.generate_contract",
            params={"department": "行政部", "month": "2026年7月"},
            reason="生成办公用品耗材采购合同示例（行政部 2026年7月）",
            subject=_subject(),
        )
    )

    assert request.params == {}
    assert request.body["category"] == "consumables"
    assert request.body["title"] == "办公用品耗材采购合同"
    assert request.body["contract_number"].startswith("AI-CONSUMABLES-")
    assert request.body["items"][0]["name"] == "办公用品耗材"
    assert request.body["items"][0]["department"] == "行政部"


@pytest.mark.anyio
async def test_agent_lists_all_contract_templates_with_fields() -> None:
    repo = FakeAgentRepository()
    service = AgentService(
        settings=SimpleNamespace(),
        repo=repo,
        access_scope_service=AllowAllAccessScopeService(),
    )

    response = await service.execute_tool(
        FakeDb(),
        request=AgentToolExecuteRequest(
            operation="procurement.list_contract_templates",
            subject=_subject(),
        ),
    )

    templates = response.data["templates"]
    by_category = {template["category"]: template for template in templates}

    assert response.ok is True
    assert set(by_category) == {
        "fixed-assets",
        "consumables",
        "hardware",
        "raw-materials",
    }
    assert by_category["fixed-assets"]["template"]["file"] == "fixed-assets.docx"
    assert by_category["fixed-assets"]["template"]["exists"] is True
    assert by_category["fixed-assets"]["template"]["size"] > 0
    assert "contract_number" in by_category["consumables"]["required_fields"]
    assert "seller.name" in by_category["raw-materials"]["required_fields"]
    assert by_category["hardware"]["item_required_fields"] == [
        "name",
        "quantity",
        "unit_price",
    ]
    assert repo.tool_calls[0].status == "succeeded"


@pytest.mark.anyio
async def test_agent_get_contract_template_normalizes_chinese_category() -> None:
    repo = FakeAgentRepository()
    service = AgentService(
        settings=SimpleNamespace(),
        repo=repo,
        access_scope_service=AllowAllAccessScopeService(),
    )

    response = await service.execute_tool(
        FakeDb(),
        request=AgentToolExecuteRequest(
            operation="procurement.get_contract_template",
            params={"category": "办公用品耗材"},
            subject=_subject(),
        ),
    )

    field_names = {field["name"] for field in response.data["fields"]}

    assert response.ok is True
    assert response.data["category"] == "consumables"
    assert response.data["template"]["file"] == "consumables.docx"
    assert "buyer_invoice_recipient" in field_names


def test_contract_generation_sample_matches_template_and_exports_docx() -> None:
    service = AgentService(settings=SimpleNamespace())
    request = service._normalize_tool_request(
        AgentToolExecuteRequest(
            operation="procurement.generate_contract",
            params={"department": "行政部", "month": "2026年7月"},
            reason="生成办公用品耗材采购合同示例（行政部 2026年7月）",
            subject=_subject(),
        )
    )
    payload = ContractGenerateRequest.model_validate(request.body)

    buffer, filename, media_type = generate_contract(payload)
    document = Document(BytesIO(buffer.getvalue()))
    text = "\n".join(
        [
            *(paragraph.text for paragraph in document.paragraphs),
            *(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ),
        ]
    )

    assert filename.startswith("耗材合同_")
    assert filename.endswith(".docx")
    assert media_type.endswith(
        "vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(buffer.getvalue()) > 1000
    assert "办公用品耗材" in text
    assert "行政部" in text


@pytest.mark.anyio
async def test_contract_generation_without_items_returns_validation_error() -> None:
    repo = FakeAgentRepository()
    service = AgentService(
        settings=SimpleNamespace(AGENT_WRITE_CONFIRM_TTL_SECONDS=300),
        repo=repo,
    )

    response = await service.execute_tool(
        FakeDb(),
        request=AgentToolExecuteRequest(
            operation="procurement.generate_contract",
            params={"category": "耗材"},
            reason="生成耗材采购合同",
            subject=_subject(),
        ),
    )

    assert response.ok is False
    assert response.requires_confirmation is False
    assert response.meta["validation"] == "failed"
    assert "合同明细" in response.data["message"]
    assert repo.tool_calls[0].status == "invalid_request"
