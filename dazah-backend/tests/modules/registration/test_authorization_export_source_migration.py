from __future__ import annotations

import xml.etree.ElementTree as ET
from datetime import datetime, timedelta
from pathlib import Path
from unittest.mock import AsyncMock, Mock
from uuid import uuid4
from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.modules.registration.schemas import (
    AuthorizationFdaRecord,
    AuthorizationLedgerGroupedOverview,
    AuthorizationLedgerMainRead,
    AuthorizationLedgerUpdateRead,
)
from app.modules.registration.service import authorization_export
from app.modules.registration.service.authorization import AuthorizationLetterService

_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _set_vmerge(cell, value: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    vmerge = OxmlElement("w:vMerge")
    vmerge.set(qn("w:val"), value)
    tc_pr.append(vmerge)


def _build_fda_template(path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=2, cols=7)
    headers = ["序号", "公司", "地址", "编号", "LOA日期", "递交日期", "章节"]
    for index, value in enumerate(headers):
        table.cell(0, index).text = value
        table.cell(1, index).text = f"模板{index + 1}"
    doc.save(path)


def _build_market_template(path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=3, cols=9)
    headers = [
        "序号",
        "授权文件名称",
        "质量标准",
        "单位名称/国家",
        "客户信息编号",
        "用途",
        "授权日期",
        "经手人",
        "备注",
    ]
    for index, value in enumerate(headers):
        table.cell(0, index).text = value

    for index in range(6):
        table.cell(1, index).text = f"主行{index + 1}"
        table.cell(2, index).text = ""
        _set_vmerge(table.cell(1, index), "restart")
        _set_vmerge(table.cell(2, index), "continue")

    table.cell(1, 6).text = "2025.01.01"
    table.cell(1, 7).text = "王五"
    table.cell(1, 8).text = "主备注"
    table.cell(2, 6).text = "2025.02.02"
    table.cell(2, 7).text = "张三"
    table.cell(2, 8).text = "更新备注"
    doc.save(path)


def _build_plain_market_template(path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=4, cols=9)
    headers = [
        "序号",
        "授权文件名称",
        "质量标准",
        "单位名称/国家",
        "客户信息编号",
        "用途",
        "授权日期",
        "经手人",
        "备注",
    ]
    for index, value in enumerate(headers):
        table.cell(0, index).text = value

    row_values = [
        "1",
        "模板文件",
        "USP",
        "模板公司/国家",
        "KH-001",
        "模板用途",
        "2025.01.01",
        "王五",
        "模板备注",
    ]
    for index, value in enumerate(row_values):
        table.cell(1, index).text = value
    doc.save(path)


def _build_partial_merge_market_template(path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=3, cols=9)
    headers = [
        "序号",
        "授权文件名称",
        "质量标准",
        "单位名称/国家",
        "客户信息编号",
        "用途",
        "授权日期",
        "经手人",
        "备注",
    ]
    for index, value in enumerate(headers):
        table.cell(0, index).text = value

    base_values = [
        "1",
        "模板文件",
        "USP",
        "模板公司/国家",
        "KH-001",
        "模板用途",
        "2025.01.01",
        "王五",
        "模板备注",
    ]
    for index, value in enumerate(base_values):
        table.cell(1, index).text = value
    for index in range(4):
        _set_vmerge(table.cell(1, index), "restart")

    update_values = [
        "",
        "",
        "",
        "",
        "KH-001",
        "模板用途",
        "2025.02.02",
        "张三",
        "更新备注",
    ]
    for index, value in enumerate(update_values):
        table.cell(2, index).text = value
    for index in range(4):
        _set_vmerge(table.cell(2, index), "continue")
    doc.save(path)


def _read_table_rows(docx_path: Path) -> list[ET.Element]:
    with ZipFile(docx_path, "r") as zipped:
        document_xml = ET.fromstring(zipped.read("word/document.xml"))
    table = document_xml.find(".//w:tbl", _NS)
    assert table is not None
    return table.findall("w:tr", _NS)


def _row_cells(row: ET.Element) -> list[ET.Element]:
    return row.findall("w:tc", _NS)


def _cell_text(cell: ET.Element) -> str:
    return "".join(node.text or "" for node in cell.findall(".//w:t", _NS))


def _cell_vmerge(cell: ET.Element) -> str | None:
    vmerge = cell.find("w:tcPr/w:vMerge", _NS)
    if vmerge is None:
        return None
    return vmerge.get(f"{{{_NS['w']}}}val") or "continue"


def _build_grouped_ledger_main(
    *,
    product_name: str = "多拉菌素",
    market_name: str = "博茨瓦纳",
    source_sequence: str = "1",
    authorization_file_name: str = "授权文件A",
    quality_standard: str = "USP",
    company_name: str = "客户公司",
    country: str = "博茨瓦纳",
    customer_code: str = "KH-001",
    purpose: str = "注册",
    status: str = "已递交",
    updates: list[tuple[str | None, str | None, str | None]] | None = None,
) -> AuthorizationLedgerMainRead:
    now = datetime(2026, 1, 1, 12, 0, 0)
    main_id = uuid4()
    update_values = updates or [("2026.01.01", "王五", "首版备注")]
    return AuthorizationLedgerMainRead(
        id=main_id,
        product_name=product_name,
        market_name=market_name,
        source_sequence=source_sequence,
        authorization_file_name=authorization_file_name,
        quality_standard=quality_standard,
        company_name=company_name,
        country=country,
        customer_code=customer_code,
        purpose=purpose,
        status=status,
        created_at=now,
        updated_at=now,
        updates=[
            AuthorizationLedgerUpdateRead(
                id=uuid4(),
                ledger_main_id=main_id,
                sort_order=index,
                authorization_date=authorization_date,
                handler=handler,
                remarks=remarks,
                created_at=now + timedelta(minutes=index),
                updated_at=now + timedelta(minutes=index),
            )
            for index, (authorization_date, handler, remarks) in enumerate(
                update_values, start=1
            )
        ],
    )


def test_render_fda_export_fills_template_rows(tmp_path, monkeypatch) -> None:
    template_path = tmp_path / "fda-template.docx"
    _build_fda_template(template_path)

    monkeypatch.setattr(
        authorization_export, "resolve_fda_template", lambda product_name: template_path
    )
    monkeypatch.setattr(authorization_export, "_get_upload_dir", lambda: tmp_path)

    artifact = authorization_export.render_fda_export(
        product_name="多拉菌素",
        records=[
            AuthorizationFdaRecord(
                product_name="多拉菌素",
                sequence=1,
                company_name="客户A",
                address="地址A",
                reference_number="REF-001",
                loa_date="2026.01.01",
                submission_date="2026.01.15",
                referenced_sections="3.2.S",
            ),
            AuthorizationFdaRecord(
                product_name="多拉菌素",
                sequence=2,
                company_name="客户B",
                address="地址B",
                reference_number="REF-002",
                loa_date="2026.02.01",
                submission_date="2026.02.18",
                referenced_sections="3.2.P",
            ),
        ],
    )

    doc = Document(str(artifact.file_path))
    table = doc.tables[0]
    assert len(table.rows) == 3
    assert table.cell(1, 0).text == "1"
    assert table.cell(1, 1).text == "客户A"
    assert table.cell(1, 6).text == "3.2.S"
    assert table.cell(2, 0).text == "2"
    assert table.cell(2, 1).text == "客户B"
    assert table.cell(2, 6).text == "3.2.P"


def test_render_fda_export_supports_six_column_templates(tmp_path, monkeypatch) -> None:
    template_path = tmp_path / "fda-template-6col.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=6)
    headers = [
        "Company name",
        "Address",
        "Reference Number",
        "Date of LOA",
        "Date of Submission",
        "Sections referenced",
    ]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
        table.cell(1, index).text = f"模板{index + 1}"
    doc.save(str(template_path))

    monkeypatch.setattr(
        authorization_export, "resolve_fda_template", lambda product_name: template_path
    )
    monkeypatch.setattr(authorization_export, "_get_upload_dir", lambda: tmp_path)

    artifact = authorization_export.render_fda_export(
        product_name="洛伐他汀",
        records=[
            AuthorizationFdaRecord(
                product_name="洛伐他汀",
                sequence=1,
                company_name="Biocon Limited",
                address="India",
                reference_number="DMF # 018539",
                loa_date="December 18, 2018",
                submission_date="July 20, 2005",
                referenced_sections="Entire DMF",
            ),
        ],
    )

    doc = Document(str(artifact.file_path))
    table = doc.tables[0]
    assert len(table.rows) == 2
    assert table.cell(1, 0).text == "Biocon Limited"
    assert table.cell(1, 1).text == "India"
    assert table.cell(1, 5).text == "Entire DMF"


def test_render_market_export_preserves_update_row_merge_structure(
    tmp_path, monkeypatch
) -> None:
    template_path = tmp_path / "market-template.docx"
    _build_market_template(template_path)

    monkeypatch.setattr(
        authorization_export,
        "resolve_market_template",
        lambda product_name, market_name: template_path,
    )
    monkeypatch.setattr(authorization_export, "_get_upload_dir", lambda: tmp_path)

    artifact = authorization_export.render_market_export(
        product_name="多拉菌素",
        market_name="博茨瓦纳",
        records=[
            _build_grouped_ledger_main(
                updates=[
                    ("2026.01.01", "王五", "首版备注"),
                    ("2026.2.3", "张三", "第一次更新"),
                    ("2026.3.4", "李四", "第二次更新"),
                ]
            )
        ],
    )

    rows = _read_table_rows(artifact.file_path)
    assert len(rows) == 4

    base_cells = _row_cells(rows[1])
    update1_cells = _row_cells(rows[2])
    update2_cells = _row_cells(rows[3])

    assert _cell_text(base_cells[0]) == "1"
    assert _cell_text(base_cells[1]) == "授权文件A"
    assert _cell_text(base_cells[6]) == "2026.01.01"
    assert _cell_text(base_cells[8]) == "首版备注"

    for index in range(6):
        assert _cell_vmerge(base_cells[index]) == "restart"
        assert _cell_vmerge(update1_cells[index]) == "continue"
        assert _cell_vmerge(update2_cells[index]) == "continue"

    assert _cell_text(update1_cells[6]) == "2026.2.3"
    assert _cell_text(update1_cells[7]) == "张三"
    assert _cell_text(update1_cells[8]) == "第一次更新"
    assert _cell_text(update2_cells[6]) == "2026.3.4"
    assert _cell_text(update2_cells[7]) == "李四"
    assert _cell_text(update2_cells[8]) == "第二次更新"


def test_render_market_export_does_not_append_blank_update_row_for_base_only_record(
    tmp_path, monkeypatch
) -> None:
    template_path = tmp_path / "market-template.docx"
    _build_market_template(template_path)

    monkeypatch.setattr(
        authorization_export,
        "resolve_market_template",
        lambda product_name, market_name: template_path,
    )
    monkeypatch.setattr(authorization_export, "_get_upload_dir", lambda: tmp_path)

    artifact = authorization_export.render_market_export(
        product_name="多拉菌素",
        market_name="博茨瓦纳",
        records=[
            _build_grouped_ledger_main(updates=[("2026.01.01", "王五", "仅主备注")])
        ],
    )

    rows = _read_table_rows(artifact.file_path)
    assert len(rows) == 2

    base_cells = _row_cells(rows[1])
    assert _cell_text(base_cells[8]) == "仅主备注"
    for index in range(6):
        assert _cell_vmerge(base_cells[index]) == "restart"


def test_render_market_export_supports_plain_row_templates(
    tmp_path, monkeypatch
) -> None:
    template_path = tmp_path / "plain-market-template.docx"
    _build_plain_market_template(template_path)

    monkeypatch.setattr(
        authorization_export,
        "resolve_market_template",
        lambda product_name, market_name: template_path,
    )
    monkeypatch.setattr(authorization_export, "_get_upload_dir", lambda: tmp_path)

    artifact = authorization_export.render_market_export(
        product_name="多拉菌素",
        market_name="欧盟",
        records=[
            _build_grouped_ledger_main(
                market_name="欧盟",
                country="爱尔兰",
                updates=[
                    ("2026.01.01", "王五", "首版备注"),
                    ("2026.2.3", "张三", "第一次更新"),
                ],
            )
        ],
    )

    rows = _read_table_rows(artifact.file_path)
    assert len(rows) == 3

    base_cells = _row_cells(rows[1])
    update_cells = _row_cells(rows[2])
    assert _cell_text(base_cells[0]) == "1"
    assert _cell_text(base_cells[1]) == "授权文件A"
    for index in range(6):
        assert _cell_vmerge(base_cells[index]) == "restart"
        assert _cell_vmerge(update_cells[index]) == "continue"
        assert _cell_text(update_cells[index]) == ""
    assert _cell_text(update_cells[6]) == "2026.2.3"
    assert _cell_text(update_cells[7]) == "张三"
    assert _cell_text(update_cells[8]) == "第一次更新"


def test_render_market_export_supports_partial_merge_templates(
    tmp_path, monkeypatch
) -> None:
    template_path = tmp_path / "partial-merge-market-template.docx"
    _build_partial_merge_market_template(template_path)

    monkeypatch.setattr(
        authorization_export,
        "resolve_market_template",
        lambda product_name, market_name: template_path,
    )
    monkeypatch.setattr(authorization_export, "_get_upload_dir", lambda: tmp_path)

    artifact = authorization_export.render_market_export(
        product_name="多拉菌素",
        market_name="英国",
        records=[
            _build_grouped_ledger_main(
                market_name="英国",
                country="爱尔兰",
                purpose="英国注册",
                updates=[
                    ("2026.01.01", "王五", "首版备注"),
                    ("2026.2.3", "张三", "第一次更新"),
                ],
            )
        ],
    )

    rows = _read_table_rows(artifact.file_path)
    assert len(rows) == 3

    base_cells = _row_cells(rows[1])
    update_cells = _row_cells(rows[2])
    for index in range(4):
        assert _cell_vmerge(base_cells[index]) == "restart"
        assert _cell_vmerge(update_cells[index]) == "continue"
        assert _cell_text(update_cells[index]) == ""
    assert _cell_text(update_cells[4]) == "KH-001"
    assert _cell_text(update_cells[5]) == "英国注册"
    assert _cell_text(update_cells[6]) == "2026.2.3"
    assert _cell_text(update_cells[7]) == "张三"
    assert _cell_text(update_cells[8]) == "第一次更新"


def test_resolve_fda_template_matches_all_products_in_source_tree(
    tmp_path, monkeypatch
) -> None:
    expected_paths: dict[str, Path] = {}
    for product_name, relative_path in [
        (
            "阿魏酸钠",
            Path("阿魏酸钠")
            / "阿魏酸钠-list-of-authorized-parties-to-incorporate-by-reference.docx",
        ),
        (
            "艾普拉唑",
            Path("艾普拉唑")
            / "艾普拉唑-list-of-authorized-parties-to-incorporate-by-reference.docx",
        ),
        (
            "头孢他啶",
            Path("头孢他啶")
            / "头孢他啶-list-of-authorized-parties-to-incorporate-by-reference.docx",
        ),
    ]:
        template_path = tmp_path / relative_path
        template_path.parent.mkdir(parents=True, exist_ok=True)
        _build_fda_template(template_path)
        expected_paths[product_name] = template_path

    monkeypatch.setattr(
        authorization_export, "_get_authorization_source_dir", lambda: tmp_path
    )

    for product_name, expected_path in expected_paths.items():
        assert authorization_export.resolve_fda_template(product_name) == expected_path


@pytest.mark.asyncio
async def test_export_ledger_entries_uses_grouped_records(monkeypatch) -> None:
    grouped_records = [
        _build_grouped_ledger_main(
            product_name="多拉菌素",
            market_name="欧盟",
            country="爱尔兰",
            updates=[
                ("2026.01.01", "王五", "首版备注"),
                ("2026.02.02", "张三", "第一次更新"),
            ],
        )
    ]
    render_mock = Mock(return_value=object())
    service = AuthorizationLetterService(AsyncMock())
    monkeypatch.setattr(
        service,
        "list_grouped_ledger_mains",
        AsyncMock(
            return_value=(
                grouped_records,
                AuthorizationLedgerGroupedOverview(
                    total_main_records=1,
                    total_update_records=2,
                    total_products=1,
                    total_markets=1,
                    submitted_main_records=1,
                    pending_main_records=0,
                ),
            )
        ),
    )
    monkeypatch.setattr(
        "app.modules.registration.service.authorization.render_market_export",
        render_mock,
    )

    artifact = await service.export_ledger_entries(
        product_name="多拉菌素", market_name="欧盟"
    )

    assert artifact is render_mock.return_value
    service.list_grouped_ledger_mains.assert_awaited_once_with(
        product_name="多拉菌素",
        market_name="欧盟",
        status=None,
        keyword=None,
    )
    render_mock.assert_called_once_with(
        product_name="多拉菌素",
        market_name="欧盟",
        records=grouped_records,
    )
