from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path
from types import SimpleNamespace as _SimpleNamespace
from typing import Any
from uuid import uuid4

import pytest
from docx import Document

from app.core.exceptions import AppException, NotFoundException
from app.modules.registration.schemas import (
    AuthorizationFdaRecord,
    AuthorizationLedgerMainRead,
    AuthorizationLedgerUpdateRead,
)
from app.modules.registration.service import authorization_export as export

SimpleNamespace: Any = _SimpleNamespace


def _update(
    order: int, date: str, handler: str, remarks: str
) -> AuthorizationLedgerUpdateRead:
    now = datetime.now(UTC)
    return AuthorizationLedgerUpdateRead(
        id=uuid4(),
        ledger_main_id=uuid4(),
        sort_order=order,
        authorization_date=date,
        handler=handler,
        remarks=remarks,
        created_at=now,
        updated_at=now,
    )


def _main_record(
    *updates: AuthorizationLedgerUpdateRead,
) -> AuthorizationLedgerMainRead:
    now = datetime.now(UTC)
    return AuthorizationLedgerMainRead(
        id=uuid4(),
        product_name="产品A",
        market_name="美国",
        source_sequence="1",
        authorization_file_name="授权文件",
        quality_standard="USP",
        company_name="公司",
        country="美国",
        customer_code="C001",
        purpose="申报",
        status="submitted",
        created_at=now,
        updated_at=now,
        updates=list(updates),
    )


def _save_table(path: Path, columns: int, rows: int = 2) -> None:
    document = Document()
    table = document.add_table(rows=rows, cols=columns)
    for column in range(columns):
        table.cell(0, column).text = f"H{column}"
        if rows > 1:
            table.cell(1, column).text = "模板"
    document.save(path)  # type: ignore[arg-type]


def test_file_name_and_value_helpers_cover_all_conventions(tmp_path: Path) -> None:
    root = tmp_path / "sources"
    detail = root / "产品A授权客户明细"
    detail.mkdir(parents=True)
    fda = root / f"Product{export.FDA_SUFFIX}.docx"
    ledger = root / "授权书台帐-产品A-美国.doc"
    nested = detail / "授权-客户-欧盟.docx"
    assert export._extract_product_name(fda, root) == "Product"
    assert export._extract_product_name(ledger, root) == "产品A"
    assert export._extract_product_name(nested, root) == "产品A"
    assert export._extract_product_name(root / "plain.docx", root) == "plain"
    assert export._extract_market_name(fda) == "FDA"
    assert export._extract_market_name(ledger) == "美国"
    assert export._extract_market_name(nested) == "欧盟"
    assert export._extract_market_name(root / "plain.docx") is None
    assert export._should_skip_source_file(root / "~$temp.docx") is True
    assert export._normalize_match_text(" A  B ") == "ab"
    assert export._normalize_line_value(None) == "-"
    assert export._normalize_line_value(" x ") == "x"
    assert export._build_fda_download_name("产品A", fda).endswith(".docx")
    assert export._build_market_download_name("产品A", "美国", ledger).endswith(".docx")


def test_template_resolution_finds_matches_and_reports_missing(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    root = tmp_path / "sources"
    root.mkdir()
    fda = root / f"Product A{export.FDA_SUFFIX}.docx"
    fda.write_bytes(b"x")
    market = root / "授权书台帐-产品A-美国.docx"
    market.write_bytes(b"x")
    (root / "~$ignored.docx").write_bytes(b"x")
    monkeypatch.setattr(export, "_get_authorization_source_dir", lambda: root)
    assert export.resolve_fda_template("ProductA") == fda
    assert export.resolve_market_template("产品A", "美国") == market
    with pytest.raises(AppException, match="未配置 FDA"):
        export.resolve_fda_template("missing")
    with pytest.raises(AppException, match="未找到对应"):
        export.resolve_market_template("产品A", "欧盟")

    monkeypatch.setattr(
        export, "_get_authorization_source_dir", lambda: tmp_path / "missing"
    )
    with pytest.raises(NotFoundException):
        export.resolve_fda_template("x")
    with pytest.raises(NotFoundException):
        export.resolve_market_template("x", "y")


def test_market_document_lines_sort_updates_and_fill_defaults() -> None:
    later = _update(2, "2026.02.01", "李四", "更新")
    first = _update(1, "2026.01.01", "张三", "首次")
    lines = export._build_market_document_lines(_main_record(later, first))
    assert [line.row_type for line in lines] == ["base", "update"]
    assert lines[0].authorization_date == "2026.01.01"
    assert lines[0].company_country == "公司/美国"
    assert lines[1].handler == "李四"
    empty = _main_record()
    empty.company_name = None
    empty.country = None
    assert export._build_market_document_lines(empty)[0].company_country == "-"


def test_docx_xml_helpers_set_text_rows_and_vertical_merges() -> None:
    document = Document()
    table = document.add_table(rows=2, cols=9)
    row = table.rows[1]._tr
    export._set_row_values(row, [str(index) for index in range(9)])
    assert export._get_cell_text(export._get_row_cell_elements(row)[0]) == "0"
    export._set_docx_cell_text(export._get_row_cell_elements(row)[0], "a\nb")
    assert "a" in export._get_cell_text(export._get_row_cell_elements(row)[0])
    with pytest.raises(AppException, match="列数不足"):
        export._set_row_values(row, ["x"] * 10)

    cell = export._get_row_cell_elements(row)[0]
    assert export._get_vmerge_state(cell) is None
    export._set_vmerge_state(cell, "restart")
    assert export._get_vmerge_state(cell) == "restart"
    export._set_vmerge_state(cell, None)
    assert export._get_vmerge_state(cell) is None
    export._apply_vertical_merge(row, state="continue")
    assert export._has_continue_merge(row) is True
    assert export._has_any_vertical_merge(row) is True
    assert export._is_market_update_template_row(row) is True
    assert export._is_market_base_template_row(row) is False
    export._apply_vertical_merge(row, state="restart")
    assert export._is_market_base_template_row(row) is True
    assert export._is_blank_market_row(table.rows[0]._tr) is True


def test_market_layout_uses_merge_blank_and_base_fallbacks() -> None:
    document = Document()
    table = document.add_table(rows=3, cols=9)
    for index in range(9):
        table.cell(1, index).text = "base"
    export._apply_vertical_merge(table.rows[1]._tr, state="restart")
    export._apply_vertical_merge(table.rows[2]._tr, state="continue")
    rows = table._tbl.findall(export.qn("w:tr"))  # type: ignore[attr-defined]
    layout = export._resolve_market_template_layout(rows)
    assert layout.base_template is not None
    assert layout.update_template is not None

    short = Document().add_table(rows=1, cols=9)
    with pytest.raises(AppException, match="格式不正确"):
        export._resolve_market_template_layout(short._tbl.findall(export.qn("w:tr")))  # type: ignore[attr-defined]


def test_soffice_and_word_converters_cover_success_failure_and_timeout(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = tmp_path / "template.doc"
    output = tmp_path / "result.docx"
    source.write_bytes(b"doc")
    monkeypatch.setattr(export, "_find_soffice", lambda: None)
    assert export._convert_doc_with_soffice(source, output) is False

    monkeypatch.setattr(export, "_find_soffice", lambda: Path("office"))
    monkeypatch.setattr(
        export.subprocess,  # type: ignore[attr-defined]
        "run",
        lambda *_args, **_kwargs: SimpleNamespace(
            returncode=1, stderr="failed", stdout=""
        ),
    )
    assert export._convert_doc_with_soffice(source, output) is False
    monkeypatch.setattr(
        export.subprocess,  # type: ignore[attr-defined]
        "run",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(FileNotFoundError()),
    )
    assert export._convert_doc_with_soffice(source, output) is False
    assert export._convert_doc_with_word(source, output) is False

    monkeypatch.setattr(
        export.subprocess,  # type: ignore[attr-defined]
        "run",
        lambda *_args, **_kwargs: SimpleNamespace(returncode=0, stderr="", stdout=""),
    )
    output.write_bytes(b"docx")
    assert export._convert_doc_with_word(source, output) is True


def test_prepare_artifacts_copy_docx_and_reject_unconvertible_doc(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    upload = tmp_path / "uploads"
    upload.mkdir()
    monkeypatch.setattr(export, "_get_upload_dir", lambda: upload)
    template = tmp_path / "template.docx"
    template.write_bytes(b"docx")
    artifact = export._prepare_market_export_artifact(template, "产品A", "美国")
    assert artifact.file_path.read_bytes() == b"docx"
    copied = export._copy_to_temp(template, "copy.docx")
    assert copied.file_path.exists()

    legacy = tmp_path / "template.doc"
    legacy.write_bytes(b"doc")
    monkeypatch.setattr(export, "_convert_doc_with_soffice", lambda *_args: False)
    monkeypatch.setattr(export, "_convert_doc_with_word", lambda *_args: False)
    with pytest.raises(AppException, match="无法归一化"):
        export._prepare_market_export_artifact(legacy, "产品A", "美国")


@pytest.mark.parametrize("columns", [6, 7])
def test_render_fda_export_supports_six_and_seven_column_templates(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, columns: int
) -> None:
    template = tmp_path / f"fda-{columns}.docx"
    _save_table(template, columns)
    upload = tmp_path / "uploads"
    upload.mkdir()
    monkeypatch.setattr(export, "resolve_fda_template", lambda _product: template)
    monkeypatch.setattr(export, "_get_upload_dir", lambda: upload)
    record = AuthorizationFdaRecord(
        sequence=1,
        company_name="公司",
        address="地址",
        reference_number="REF",
        loa_date="2026.01.01",
        submission_date="2026.02.01",
        referenced_sections="3.2.S",
    )
    artifact = export.render_fda_export(product_name="产品A", records=[record])
    document = Document(artifact.file_path)  # type: ignore[arg-type]
    assert len(document.tables[0].rows) == 2
    assert "公司" in [cell.text for cell in document.tables[0].rows[1].cells]


def test_render_market_export_adds_update_rows_and_synthesized_merges(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    template = tmp_path / "market.docx"
    _save_table(template, 9, rows=3)
    upload = tmp_path / "uploads"
    upload.mkdir()
    monkeypatch.setattr(export, "resolve_market_template", lambda *_args: template)
    monkeypatch.setattr(export, "_get_upload_dir", lambda: upload)
    record = _main_record(
        _update(1, "2026.01.01", "张三", "首次"),
        _update(2, "2026.02.01", "李四", "更新"),
    )
    artifact = export.render_market_export(
        product_name="产品A", market_name="美国", records=[record]
    )
    document = Document(artifact.file_path)  # type: ignore[arg-type]
    assert len(document.tables[0].rows) == 3
    assert document.tables[0].cell(2, 7).text == "李四"
