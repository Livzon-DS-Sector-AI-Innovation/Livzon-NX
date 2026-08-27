from datetime import UTC, date, datetime
from io import BytesIO
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock
from uuid import uuid4

import pytest
from docx import Document
from openpyxl import Workbook

from app.core.exceptions import AppException
from app.core.llm.exceptions import LLMOutputError, LLMProviderError, LLMRateLimitError
from app.modules.quality.service import quality_ai as service


def _deviation(**overrides: object) -> SimpleNamespace:
    values = {
        "id": uuid4(),
        "deviation_code": "PC-1",
        "title": "设备异常",
        "department": "质量部",
        "discovery_date": date(2026, 8, 20),
        "discovery_time": "10:00",
        "discovery_location": "车间",
        "level": "major",
        "description": "发现异常",
        "immediate_actions": "隔离",
        "affected_items": "产品A",
        "batch_number": "B1",
        "root_cause_analysis": "待调查",
        "corrective_actions": "待制定",
        "is_deleted": False,
        "status": "pending_ai_analysis",
        "status_updated_at": None,
        "updated_by": None,
    }
    values.update(overrides)
    return SimpleNamespace(**values)


def _capa(**overrides: object) -> SimpleNamespace:
    values = {
        "id": uuid4(),
        "capa_code": "CAPA-PC-1",
        "title": "改进措施",
        "deviation_id": None,
        "source": "deviation",
        "source_code": "PC-1",
        "category": "设备",
        "department": "质量部",
        "affected_product": "产品A",
        "non_conformity_description": "不符合",
        "root_cause_analysis": "原因",
        "capa_content": "措施",
        "capa_items": [],
        "execution_status": "进行中",
        "evaluation_result": None,
        "closure_date": date(2026, 9, 1),
        "is_deleted": False,
    }
    values.update(overrides)
    return SimpleNamespace(**values)


def _change(**overrides: object) -> SimpleNamespace:
    values = {
        "id": uuid4(),
        "change_code": "CH-1",
        "serial_number": "001",
        "applicant_department": "质量部",
        "change_object": "设备",
        "change_content": "更换设备",
        "impact_assessment": "待评估",
        "change_level": "major",
        "application_date": date(2026, 8, 20),
        "planned_approval_date": date(2026, 8, 22),
        "execution_date": None,
        "closure_date": None,
        "is_deleted": False,
    }
    values.update(overrides)
    return SimpleNamespace(**values)


def test_quality_ai_normalization_and_prompts_cover_entity_variants() -> None:
    nested = {
        "at": datetime(2026, 8, 20, tzinfo=UTC),
        "id": uuid4(),
        "items": [date(2026, 8, 21)],
    }
    converted = service._to_iso(nested)
    assert converted["at"].startswith("2026-08-20")
    assert isinstance(converted["id"], str)
    assert converted["items"] == ["2026-08-21"]

    assert service._normalize_result(
        {"risks": "风险", "suggestions": 1, "missing_info": None}
    )["risks"] == ["风险"]
    assert service._normalize_result({})["disclaimer"]
    payload = {
        "structured_fields": {
            "preliminary_cause_analysis": "根因",
            "capa_suggestions": "措施",
        }
    }
    fields = service._build_applicable_fields(
        "deviation", "deviation_analysis", payload
    )
    assert {item.field_key for item in fields} == {
        "root_cause_analysis",
        "corrective_actions",
    }
    assert (
        service._build_applicable_fields("deviation", "capa_suggestion", payload)[
            0
        ].field_key
        == "corrective_actions"
    )
    assert service._build_applicable_fields(
        "capa",
        "capa_review",
        {
            "structured_fields": {
                "root_cause_analysis": "r",
                "capa_content": "c",
                "effectiveness_review": "e",
            }
        },
    )
    assert (
        service._build_applicable_fields(
            "change", "change_impact", {"structured_fields": {"impact_assessment": "i"}}
        )[0].field_key
        == "impact_assessment"
    )
    assert service._build_applicable_fields("unknown", "other", payload) == []

    deviation = _deviation()
    capa = _capa()
    change = _change()
    assert service._build_deviation_snapshot(deviation)["deviation_code"] == "PC-1"
    assert service._build_capa_snapshot(capa)["capa_code"] == "CAPA-PC-1"
    assert service._build_change_snapshot(change)["change_code"] == "CH-1"
    assert service._extract_deviation_code_from_capa_code("CAPA-PC-1") == "PC-1"
    assert service._extract_deviation_code_from_capa_code("bad") is None
    assert "capa_suggestion" in service._deviation_prompt({}, "capa_suggestion")
    assert "关联偏差" in service._capa_prompt(
        {
            "linked_deviation": {"code": "PC-1"},
            "analysis_context": {"match_rule": "source_code"},
        }
    )
    assert "未识别" in service._capa_prompt(
        {"linked_deviation": None, "analysis_context": {}}
    )
    assert "变更记录" in service._change_prompt({})
    assert "补充说明" in service._deviation_conversation_prompt({}, "补充", "附件")


@pytest.mark.asyncio
async def test_quality_ai_config_related_capa_and_log_helpers(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    async def config_ok(_kind: str) -> object:
        return SimpleNamespace(model_name="test")

    monkeypatch.setattr(service, "get_config", config_ok)
    assert (await service._require_quality_ai_config()).model_name == "test"

    async def config_missing(_kind: str) -> object:
        from app.core.llm.exceptions import LLMConfigError

        raise LLMConfigError("missing")

    monkeypatch.setattr(service, "get_config", config_missing)
    with pytest.raises(AppException) as exc:
        await service._require_quality_ai_config()
    assert exc.value.status_code == 503

    deviation = _deviation()
    db = SimpleNamespace(get=AsyncMock(return_value=deviation))
    capa = _capa(deviation_id=deviation.id, source=None, source_code=None)
    linked, rule = await service._resolve_related_deviation_for_capa(db, capa)
    assert linked is deviation and rule == "deviation_id"
    db.get.return_value = None
    result = SimpleNamespace(scalar_one_or_none=lambda: deviation)
    db.execute = AsyncMock(return_value=result)
    capa = _capa(deviation_id=None, source="deviation", source_code="PC-1")
    linked, rule = await service._resolve_related_deviation_for_capa(db, capa)
    assert linked is deviation and rule == "source_code"
    db.execute.return_value = SimpleNamespace(scalar_one_or_none=lambda: None)
    capa = _capa(
        deviation_id=None, source=None, source_code=None, capa_code="CAPA-PC-1"
    )
    db.execute.side_effect = [SimpleNamespace(scalar_one_or_none=lambda: deviation)]
    linked, rule = await service._resolve_related_deviation_for_capa(db, capa)
    assert linked is deviation and rule == "capa_code"
    db.execute.side_effect = None
    db.execute.return_value = SimpleNamespace(scalar_one_or_none=lambda: None)
    assert await service._resolve_related_deviation_for_capa(
        db, _capa(deviation_id=None, source=None, source_code=None, capa_code="NONE")
    ) == (None, None)
    snapshot = await service._build_capa_analysis_snapshot(
        db, _capa(deviation_id=None, source=None, source_code=None, capa_code="NONE")
    )
    assert snapshot["analysis_context"]["analysis_basis"] == "capa_only"

    db.add = lambda item: setattr(item, "id", uuid4())
    db.flush = AsyncMock()
    log = await service._create_log(
        db,
        entity_type="deviation",
        entity_id=deviation.id,
        analysis_type="deviation_analysis",
        input_snapshot={},
        model_name="m",
        status="completed",
    )
    assert log.entity_type == "deviation"


def _llm_result() -> dict[str, object]:
    return {
        "summary": "摘要",
        "risk_level": "低",
        "risks": [],
        "suggestions": [],
        "missing_info": [],
        "structured_fields": {"preliminary_cause_analysis": "根因"},
    }


@pytest.mark.asyncio
async def test_quality_ai_analysis_success_and_provider_mapping(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    config = SimpleNamespace(model_name="test-model")
    monkeypatch.setattr(
        service, "_require_quality_ai_config", AsyncMock(return_value=config)
    )
    monkeypatch.setattr(
        type(service.llm_client),
        "chat_json",
        AsyncMock(return_value=_llm_result()),
    )
    log = SimpleNamespace(
        id=uuid4(),
        entity_type="deviation",
        entity_id=uuid4(),
        analysis_type="deviation_analysis",
        input_snapshot={},
        output_payload=_llm_result(),
        model_name="test-model",
        status="completed",
        error_message=None,
        is_applied=False,
        is_deleted=False,
        created_at=datetime.now(UTC),
        created_by=None,
        applied_at=None,
        applied_by=None,
    )
    monkeypatch.setattr(service, "_create_log", AsyncMock(return_value=log))

    async def get_entity(model: object, entity_id: object) -> object:
        if model is service.Deviation:
            return _deviation(id=entity_id)
        if model is service.CAPA:
            return capa
        return change

    db = SimpleNamespace(
        get=AsyncMock(side_effect=get_entity),
        add=MagicMock(),
        flush=AsyncMock(),
        commit=AsyncMock(),
        execute=AsyncMock(return_value=SimpleNamespace(scalar_one=lambda: log)),
    )
    entity = await service.analyze_deviation_record(
        db, log.entity_id, "system", transition_status=True
    )
    assert entity.status == "completed"

    capa = _capa(id=uuid4())
    log.entity_id = capa.id
    log.entity_type = "capa"
    log.analysis_type = "capa_review"
    monkeypatch.setattr(
        service, "_build_capa_analysis_snapshot", AsyncMock(return_value={"capa": {}})
    )
    db.get.return_value = capa
    assert (
        await service.analyze_capa_record(db, capa.id, "system")
    ).entity_type == "capa"

    change = _change(id=uuid4())
    log.entity_id = change.id
    log.entity_type = "change"
    log.analysis_type = "change_impact"
    db.get.return_value = change
    assert (
        await service.analyze_change_record(db, change.id, "system")
    ).entity_type == "change"

    for error in (LLMOutputError("bad"), LLMProviderError("provider")):
        monkeypatch.setattr(
            type(service.llm_client), "chat_json", AsyncMock(side_effect=error)
        )
        db.get.return_value = change
        with pytest.raises(RuntimeError):
            await service.analyze_change_record(db, change.id, "system")
    monkeypatch.setattr(
        type(service.llm_client),
        "chat_json",
        AsyncMock(side_effect=LLMRateLimitError("rate")),
    )
    with pytest.raises(RuntimeError):
        await service.analyze_change_record(db, change.id, "system")


@pytest.mark.asyncio
async def test_quality_ai_apply_and_attachment_parsers(
    monkeypatch: pytest.MonkeyPatch, tmp_path
) -> None:
    assert service._truncate_text(" x ") == "x"
    assert service._truncate_text("x" * 5, 3) == "xxx..."
    first = SimpleNamespace(
        is_deleted=False,
        parse_status="completed",
        file_name="a.txt",
        parsed_summary=None,
        parsed_text="x" * 500,
    )
    second = SimpleNamespace(
        is_deleted=False,
        parse_status="completed",
        file_name="b.txt",
        parsed_summary="摘要",
        parsed_text=None,
    )
    skipped = SimpleNamespace(
        is_deleted=True,
        parse_status="completed",
        file_name="c",
        parsed_summary="x",
        parsed_text=None,
    )
    assert "a.txt" in service._build_attachment_summary_from_rows(
        [first, second, skipped]
    )

    document = Document()
    document.add_paragraph("正文")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "列1"
    table.cell(0, 1).text = "值1"
    stream = BytesIO()
    document.save(stream)
    assert "正文" in service._extract_docx_text(stream.getvalue())
    assert service._extract_legacy_doc_text(b"\x00" + "文本".encode()) == "文本"

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "库存"
    sheet.append(["物料", "数量"])
    sheet.append(["乙醇", 2])
    excel = BytesIO()
    workbook.save(excel)
    parsed, summary = await service._parse_attachment_content(
        "stock.xlsx",
        excel.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    assert "乙醇" in (parsed or "") and summary
    parsed, summary = await service._parse_attachment_content(
        "notes.doc", "测试".encode("gbk"), "application/msword"
    )
    assert parsed == "测试" and summary == "测试"
    with pytest.raises(AppException):
        await service._parse_attachment_content("bad.pdf", b"x", "application/pdf")

    entity_id = uuid4()
    log = SimpleNamespace(
        id=uuid4(),
        entity_type="deviation",
        entity_id=entity_id,
        analysis_type="deviation_analysis",
        status="completed",
        output_payload={
            "structured_fields": {
                "preliminary_cause_analysis": "新根因",
                "capa_suggestions": "新措施",
            }
        },
        input_snapshot={},
        model_name="m",
        error_message=None,
        is_applied=False,
        is_deleted=False,
        created_at=datetime.now(UTC),
        created_by=None,
        applied_at=None,
        applied_by=None,
    )
    deviation = _deviation(id=entity_id)
    db = SimpleNamespace(
        get=AsyncMock(side_effect=[log, deviation]),
        commit=AsyncMock(),
        execute=AsyncMock(return_value=SimpleNamespace(scalar_one=lambda: log)),
    )
    result = await service.apply_ai_log(
        db, log.id, ["root_cause_analysis", "corrective_actions"], "system"
    )
    assert result.is_applied is True
    assert deviation.root_cause_analysis == "新根因"
    db.get.side_effect = [log, deviation]
    with pytest.raises(AppException):
        await service.apply_ai_log(db, log.id, ["invalid"], "system")
