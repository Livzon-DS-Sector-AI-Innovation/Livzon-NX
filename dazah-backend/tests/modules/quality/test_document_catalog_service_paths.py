from __future__ import annotations

from datetime import date, datetime
from io import BytesIO
from unittest.mock import AsyncMock
from uuid import uuid4

import pytest
from docx import Document
from openpyxl import Workbook

from app.modules.quality.models.document_catalog import (
    DocumentDepartment,
    DocumentEntry,
)
from app.modules.quality.service import document_catalog


class _Result:
    def __init__(
        self,
        value: object | None = None,
        rows: list[object] | None = None,
    ) -> None:
        self.value = value
        self.rows = rows or []

    def scalar_one_or_none(self) -> object | None:
        return self.value

    def scalars(self) -> _Result:
        return self

    def all(self) -> list[object]:
        return self.rows


class _Db:
    def __init__(self, results: list[_Result]) -> None:
        self.execute = AsyncMock(side_effect=results)
        self.added: list[object] = []
        self.add = self.added.append
        self.flush = AsyncMock(side_effect=self._flush)

    async def _flush(self) -> None:
        for item in self.added:
            if getattr(item, "id", None) is None:
                item.id = uuid4()
            if getattr(item, "created_at", None) is None:
                item.created_at = datetime(2026, 8, 26)
            if getattr(item, "updated_at", None) is None:
                item.updated_at = datetime(2026, 8, 26)
            if getattr(item, "is_deleted", None) is None:
                item.is_deleted = False


def _workbook_bytes() -> bytes:
    workbook = Workbook()
    first = workbook.active
    assert first is not None
    first.title = "质量部"
    second = workbook.create_sheet("生产部")
    for sheet in (first, second):
        sheet.append(["序号", "文件名称", "文件编码", "生效日期"])
    first.append(["", "偏差处理程序", "QA-001", "2026-08-20"])
    first.append([2, "变更控制程序", "QA-002", datetime(2026, 8, 21)])
    first.append([None, None, None, None])
    second.append(["bad", "生产记录", "PR-001", "无效日期"])
    stream = BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def test_document_catalog_parses_workbook_and_docx_variants() -> None:
    parsed = document_catalog.parse_document_catalog_workbook(_workbook_bytes())
    assert list(parsed) == ["质量部", "生产部"]
    assert parsed["质量部"][0]["seq_no"] is None
    assert parsed["质量部"][0]["effective_date"] == date(2026, 8, 20)
    assert parsed["质量部"][1]["effective_date"] == date(2026, 8, 21)
    assert parsed["生产部"][0]["effective_date"] is None
    assert parsed["生产部"][0]["effective_date_text"] == "无效日期"
    assert document_catalog.parse_effective_date("2026年13月01日")[0] is None
    assert document_catalog.parse_effective_date(date(2026, 8, 26))[0] == date(
        2026, 8, 26
    )

    document = Document()
    document.add_paragraph("部门：质量部")
    table = document.add_table(rows=3, cols=4)
    for index, value in enumerate(("序号", "文件名称", "文件编码", "生效日期")):
        table.cell(0, index).text = value
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "偏差程序"
    table.cell(1, 2).text = "QA-003"
    table.cell(1, 3).text = "2026/08/22"
    table.cell(2, 0).text = "2"
    table.cell(2, 1).text = ""
    doc_stream = BytesIO()
    document.save(doc_stream)
    department, entries = document_catalog.parse_document_catalog_docx(
        doc_stream.getvalue(), fallback_name="质量部文件目录.docx"
    )
    assert department == "质量部"
    assert entries[0]["seq_no"] == 1
    assert entries[0]["effective_date"] == date(2026, 8, 22)

    empty_document = Document()
    empty_stream = BytesIO()
    empty_document.save(empty_stream)
    assert document_catalog.parse_document_catalog_docx(
        empty_stream.getvalue(), fallback_name="生产部文件目录.docx"
    ) == ("生产部", [])


@pytest.mark.asyncio
async def test_import_document_catalog_restores_departments_and_keeps_attachments() -> (
    None
):
    restored_department = DocumentDepartment(name="质量部", sort_order=9)
    restored_department.id = uuid4()
    restored_department.is_deleted = True
    old_entry = DocumentEntry(
        department_id=restored_department.id,
        name="旧偏差程序",
        code="QA-001",
        attachments=[{"storage_key": "quality/old.pdf"}],
    )
    old_entry.id = uuid4()
    old_entry.is_deleted = False
    db = _Db(
        [
            _Result(restored_department),
            _Result(rows=[old_entry]),
            _Result(),
            _Result(),
            _Result(rows=[]),
            _Result(),
        ]
    )

    result = await document_catalog.import_document_catalog(
        db,
        _workbook_bytes(),
        source_file="目录.xlsx",
        filename="目录.xlsx",
    )

    assert result["department_count"] == 2
    assert result["entry_count"] == 3
    assert restored_department.is_deleted is False
    imported_entries = [item for item in db.added if isinstance(item, DocumentEntry)]
    preserved = next(item for item in imported_entries if item.code == "QA-001")
    assert preserved.attachments == [{"storage_key": "quality/old.pdf"}]
    assert len(imported_entries) == 3
    assert db.flush.await_count == 4
