from __future__ import annotations

from datetime import UTC, date, datetime
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock
from uuid import uuid4

import pytest
from docx import Document

from app.core.exceptions import AppException, DuplicateException, NotFoundException
from app.modules.quality.schemas.contacts import (
    CreateDepartmentContactRequest,
    UpdateDepartmentContactRequest,
)
from app.modules.quality.service import (
    department_contacts,
    document_catalog_export,
    quality_statistics,
)
from app.modules.registration.repository.fee import RegistrationFeeRepository


class _Result:
    def __init__(
        self,
        *,
        scalar_value: object = None,
        one: object = None,
        scalars: list[object] | None = None,
        rows: list[object] | None = None,
    ) -> None:
        self.scalar_value = scalar_value
        self.one = one
        self.scalar_values = scalars or []
        self.rows = rows or []

    def scalar_one(self) -> object:
        return self.scalar_value if self.scalar_value is not None else self.one

    def scalar_one_or_none(self) -> object:
        return self.one

    def scalar(self) -> object:
        return self.scalar_value

    def scalars(self) -> _Result:
        return self

    def all(self) -> list[object]:
        return self.scalar_values if self.scalar_values else self.rows


@pytest.mark.asyncio
async def test_quality_statistics_cover_local_and_feishu_branches(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    now = datetime.now(UTC)
    deviations = [
        SimpleNamespace(
            status="closed",
            department="质量部",
            level="major",
            investigation_completed_at=now,
            created_at=now,
        ),
        SimpleNamespace(
            status="investigating",
            department=None,
            level="minor",
            investigation_completed_at=None,
            created_at=now,
        ),
    ]
    db = SimpleNamespace(execute=AsyncMock(return_value=_Result(scalars=deviations)))
    monkeypatch.setattr(
        quality_statistics,
        "_fetch_feishu_records",
        AsyncMock(return_value=[{"fields": {}}]),
    )
    deviation_stats = await quality_statistics.get_deviation_statistics(db)
    assert deviation_stats.total == 2
    assert deviation_stats.closed_count == 1
    assert deviation_stats.capa_total == 1
    assert {item["name"] for item in deviation_stats.department_distribution} == {
        "质量部",
        "未知",
    }

    import app.modules.quality.service.quality_feishu_sync as sync_module

    runtime = SimpleNamespace(
        is_enabled=lambda: True,
        get_entity_config=lambda *_args, **_kwargs: {},
    )
    monkeypatch.setattr(
        sync_module.feishu_sync,
        "_resolve_runtime",
        AsyncMock(return_value=runtime),
    )
    monkeypatch.setattr(
        sync_module,
        "_get_mapped_field_value",
        lambda _entity, fields, key: fields.get(key),
    )
    monkeypatch.setattr(
        sync_module,
        "_normalize_text",
        lambda value: str(value or "").strip(),
    )
    monkeypatch.setattr(
        quality_statistics,
        "_fetch_feishu_records",
        AsyncMock(
            return_value=[
                {
                    "fields": {
                        "状态": "已关闭",
                        "来源": "偏差",
                        "CAPA类型": "纠正",
                        "责任部门": "质量部",
                    }
                },
                {
                    "fields": {
                        "状态": "进行中",
                        "来源": "投诉",
                        "类型": "预防",
                        "部门": "生产部",
                    }
                },
            ]
        ),
    )
    capa_stats = await quality_statistics.get_capa_statistics(db)
    assert capa_stats.total == 2
    assert capa_stats.closed_count == 1
    assert {item["name"] for item in capa_stats.department_distribution} == {
        "质量部",
        "生产部",
    }

    change_db = SimpleNamespace(
        execute=AsyncMock(
            side_effect=[
                _Result(scalar_value=4),
                _Result(rows=[("major", 2), (None, 2)]),
                _Result(
                    rows=[
                        (date.today(), None, None),
                        (None, date.today(), None),
                        (None, None, date.today()),
                        (None, None, date(2020, 1, 1)),
                    ]
                ),
                _Result(rows=[("质量部", 3), (None, 1)]),
                _Result(rows=[("工艺", 4)]),
                _Result(scalar_value=6),
                _Result(scalar_value=2),
                _Result(scalar_value=3),
            ]
        )
    )
    change_stats = await quality_statistics.get_change_statistics(change_db)
    assert change_stats.total == 4
    assert change_stats.closed_count == 1
    assert change_stats.delay_count == 1
    assert change_stats.action_plan_total == 6
    assert change_stats.action_plan_overdue == 2
    assert change_stats.action_plan_confirmed == 3


def _contact(**overrides: object) -> SimpleNamespace:
    values: dict[str, object] = {
        "id": uuid4(),
        "name": "张三",
        "department": "质量部",
        "enterprise_email": "z@example.com",
        "open_id": "ou-z",
        "department_head_name": "李四",
        "department_head_enterprise_email": "l@example.com",
        "department_head_open_id": "ou-l",
        "feishu_record_id": "rec-1",
        "created_at": datetime(2026, 8, 20, tzinfo=UTC),
        "updated_at": datetime(2026, 8, 21, tzinfo=UTC),
        "is_deleted": False,
    }
    values.update(overrides)
    return SimpleNamespace(**values)


@pytest.mark.asyncio
async def test_department_contacts_helpers_crud_and_feishu(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    record = _contact()
    db = SimpleNamespace(
        execute=AsyncMock(
            side_effect=[
                _Result(scalar_value=1),
                _Result(scalars=[record]),
            ]
        ),
        add=MagicMock(),
        commit=AsyncMock(),
        rollback=AsyncMock(),
    )
    listing = await department_contacts.get_department_contact_list(db)
    assert listing["total"] == 1
    assert listing["items"][0]["department"] == "质量部"

    assert (
        department_contacts._normalize_feishu_contact_value(
            [{"text": "张三"}, {"email": "z@example.com"}]
        )
        == "张三 / z@example.com"
    )
    assert department_contacts._normalize_feishu_contact_value(0) == "0"
    assert (
        department_contacts._normalize_feishu_contact_value({"name": "质量部"})
        == "质量部"
    )
    assert department_contacts._normalize_feishu_contact_value({}) is None
    assert (
        department_contacts._normalize_feishu_contact_person_id(
            {"value": [{"id": "ou-z"}]}
        )
        == "ou-z"
    )
    assert department_contacts._normalize_feishu_contact_person_id([]) is None
    assert department_contacts._format_feishu_contact_datetime(1_750_000_000_000)
    assert department_contacts._format_feishu_contact_datetime(None) == ""
    serialized = department_contacts._serialize_feishu_department_contact(
        {
            "record_id": "rec-1",
            "created_time": 1_750_000_000_000,
            "last_modified_time": 1_750_000_000_001,
            "fields": {
                "姓名 (人员 )": [{"id": "ou-z", "name": "张三"}],
                "部门": [{"text": "质量部"}],
                "Open ID": "ou-z",
                "上级负责人姓名 (人员 )": {"value": [{"id": "ou-l"}]},
            },
        }
    )
    assert serialized.bitable_user_id == "ou-z"
    assert serialized.department_head_bitable_user_id == "ou-l"

    unique_db = SimpleNamespace(
        execute=AsyncMock(return_value=_Result(one=None)),
        add=MagicMock(),
        commit=AsyncMock(),
        rollback=AsyncMock(),
    )
    created = await department_contacts.upsert_department_contact(
        unique_db,
        CreateDepartmentContactRequest(
            name="张三", department="质量部", open_id="ou-z"
        ),
        "质量部",
        "user-1",
    )
    assert created == {"success": True}
    unique_db.commit.assert_awaited_once()
    with pytest.raises(AppException, match="required"):
        await department_contacts.upsert_department_contact(
            unique_db,
            UpdateDepartmentContactRequest(name="改名"),
            None,
            "user-1",
        )

    updated_record = _contact()
    update_db = SimpleNamespace(
        execute=AsyncMock(side_effect=[_Result(one=updated_record), _Result(one=None)]),
        commit=AsyncMock(),
        rollback=AsyncMock(),
    )
    updated = await department_contacts.update_department_contact(
        update_db,
        updated_record.id,
        UpdateDepartmentContactRequest(name="新姓名", open_id="ou-new"),
    )
    assert updated == {"success": True}
    assert updated_record.name == "新姓名"
    assert updated_record.open_id == "ou-new"

    duplicate_db = SimpleNamespace(execute=AsyncMock(return_value=_Result(one=record)))
    with pytest.raises(DuplicateException):
        await department_contacts._ensure_department_contact_open_id_unique(
            duplicate_db, "ou-z"
        )
    missing_db = SimpleNamespace(
        execute=AsyncMock(return_value=_Result(one=None)),
        rollback=AsyncMock(),
        commit=AsyncMock(),
    )
    with pytest.raises(NotFoundException):
        await department_contacts.update_department_contact(
            missing_db, uuid4(), UpdateDepartmentContactRequest(name="x")
        )
    with pytest.raises(NotFoundException):
        await department_contacts.delete_department_contact(missing_db, uuid4())

    delete_record = _contact()
    delete_db = SimpleNamespace(
        execute=AsyncMock(return_value=_Result(one=delete_record)),
        commit=AsyncMock(),
        rollback=AsyncMock(),
    )
    assert await department_contacts.delete_department_contact(
        delete_db, delete_record.id
    ) == {"success": True}
    assert delete_record.is_deleted is True

    import app.modules.quality.service.quality_feishu_sync as sync_module

    runtime = SimpleNamespace(
        app_id="app",
        app_secret="secret",
        is_enabled=lambda: True,
        get_entity_config=lambda *_args, **_kwargs: SimpleNamespace(
            app_token="token", table_id="table"
        ),
    )
    monkeypatch.setattr(
        sync_module.feishu_sync,
        "_resolve_runtime",
        AsyncMock(return_value=runtime),
    )
    client = SimpleNamespace(
        list_all_records=AsyncMock(
            return_value=[
                {
                    "record_id": "r2",
                    "fields": {"部门": "生产部", "姓名 (人员 )": "李四"},
                },
                {
                    "record_id": "r1",
                    "fields": {"部门": "质量部", "姓名 (人员 )": "张三"},
                },
            ]
        )
    )
    monkeypatch.setattr(
        "app.platform.integrations.feishu.bitable.BitableClient",
        lambda **_kwargs: client,
    )
    feishu_listing = await department_contacts.get_department_contact_list_from_feishu(
        SimpleNamespace(), page=1, page_size=1
    )
    assert feishu_listing["total"] == 2
    assert feishu_listing["items"][0]["department"] == "生产部"

    runtime.is_enabled = lambda: False
    with pytest.raises(AppException):
        await department_contacts.get_department_contact_list_from_feishu(
            SimpleNamespace()
        )


def _entry(name: str = "SOP") -> SimpleNamespace:
    return SimpleNamespace(
        name=name,
        code="SOP-1",
        effective_date=date(2026, 8, 20),
        effective_date_text=None,
    )


def test_document_catalog_export_template_and_fallback(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    source = tmp_path / "source.docx"
    source_doc = Document()
    source_doc.add_paragraph("部门：旧部门")
    table = source_doc.add_table(rows=2, cols=4)
    table.cell(0, 0).text = "序号"
    table.cell(0, 1).text = "文件名称"
    table.cell(0, 2).text = "文件编码"
    table.cell(0, 3).text = "生效日期"
    table.cell(1, 0).text = "0"
    table.cell(1, 1).text = "旧文件"
    table.cell(1, 2).text = "旧编码"
    table.cell(1, 3).text = "2020.01.01"
    source_doc.save(source)
    seeded_dir = tmp_path / "templates"
    seeded_dir.mkdir()
    monkeypatch.setattr(
        document_catalog_export, "get_templates_dir", lambda: seeded_dir
    )
    seeded = document_catalog_export.seed_template_from_source(str(source), "质量部")
    assert seeded.exists()
    assert document_catalog_export.find_template("质量部") == seeded
    monkeypatch.setattr(document_catalog_export, "find_template", lambda _name: seeded)
    exported = document_catalog_export.export_document_catalog_docx(
        [_entry()], "质量部"
    )
    exported_doc = Document(BytesIO(exported))
    assert "部门：质量部" in exported_doc.paragraphs[0].text
    assert "SOP" in exported_doc.tables[0].cell(1, 1).text

    monkeypatch.setattr(document_catalog_export, "find_template", lambda _name: None)
    fallback = document_catalog_export.export_document_catalog_docx([], "注册部")
    assert len(Document(BytesIO(fallback)).tables) == 1
    assert document_catalog_export._find_target_table(Document()) is None


@pytest.mark.asyncio
async def test_registration_fee_repository_filters_summaries_and_contacts() -> None:
    entry = SimpleNamespace(
        id=uuid4(),
        fee_type="官方费",
        amount=Decimal("10.00"),
        payment_status="待支付",
        agency_name="代理",
        expense_content="注册",
        handler="张三",
        remarks="备注",
        is_deleted=False,
    )
    contact = SimpleNamespace(id=uuid4(), test_item="稳定性", agency_name="机构")
    rows = [
        SimpleNamespace(fee_type="官方费", total_amount=Decimal("10"), record_count=1),
        SimpleNamespace(
            payment_status="待支付", total_amount=Decimal("10"), record_count=1
        ),
        SimpleNamespace(year=2026, total_amount=Decimal("10"), record_count=1),
        SimpleNamespace(
            year=2026, fee_type="官方费", total_amount=Decimal("10"), record_count=1
        ),
        SimpleNamespace(agency_name="代理", total_amount=Decimal("10"), record_count=1),
    ]
    db = SimpleNamespace(
        execute=AsyncMock(
            side_effect=[
                _Result(scalar_value=1),
                _Result(one=entry),
                _Result(scalars=[entry]),
                _Result(one=entry),
                _Result(rows=[rows[0]]),
                _Result(rows=[rows[1]]),
                _Result(rows=[rows[2]]),
                _Result(rows=[rows[3]]),
                _Result(rows=[rows[4]]),
                _Result(scalar_value=Decimal("10")),
                _Result(scalar_value=Decimal("10")),
                _Result(scalar_value=Decimal("0")),
                _Result(scalar_value=1),
                _Result(one=contact),
                _Result(scalars=[contact]),
                _Result(one=contact),
            ]
        ),
        add=MagicMock(),
        add_all=MagicMock(),
        flush=AsyncMock(),
    )
    repo = RegistrationFeeRepository(db)
    assert await repo.count_entries(year_from=2025) == 1
    assert await repo.get_by_id(entry.id) is entry
    assert await repo.list_entries(
        fee_type="官方费",
        payment_status="待支付",
        project_name="项目",
        product_name="产品",
        country="中国",
        year=2026,
        year_from=2025,
        keyword="代理",
    ) == [entry]
    assert await repo.create_entry(entry) is entry
    assert (
        await repo.update_entry(entry, {"remarks": "新备注", "unknown": "x"}) is entry
    )
    await repo.soft_delete(entry)
    assert entry.is_deleted is True
    assert (await repo.get_fee_type_summaries(year_from=2025))[0][
        "total_amount"
    ] == Decimal("10")
    assert (await repo.get_payment_status_summaries())[0]["payment_status"] == "待支付"
    assert (await repo.get_year_summaries())[0]["year"] == 2026
    assert (await repo.get_year_fee_type_summaries())[0]["fee_type"] == "官方费"
    assert (await repo.get_agency_summaries())[0]["agency_name"] == "代理"
    assert await repo.get_total_amount() == Decimal("10")
    assert await repo.get_pending_amount() == Decimal("10")
    assert await repo.get_paid_amount() == Decimal("0")
    assert await repo.count_inspection_contacts() == 1
    assert await repo.get_inspection_contact_by_id(contact.id) is contact
    assert await repo.list_inspection_contacts(keyword="稳定") == [contact]
    assert await repo.create_inspection_contact(contact) is contact
    assert await repo.create_inspection_contacts([contact]) == [contact]
    assert (
        await repo.update_inspection_contact(contact, {"agency_name": "新机构"})
        is contact
    )
    await repo.soft_delete_inspection_contact(contact)
    assert contact.is_deleted is True
