from __future__ import annotations

from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace as _SimpleNamespace
from typing import Any
from unittest.mock import AsyncMock

import pytest
from docx import Document

from app.modules.quality.service import quality_import_export as service

SimpleNamespace: Any = _SimpleNamespace


def _docx_table(headers: list[str], rows: list[list[str]]) -> bytes:
    document = Document()
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            row.cells[index].text = value
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _empty_docx() -> bytes:
    stream = BytesIO()
    Document().save(stream)
    return stream.getvalue()


def _capa_docx(rows: list[list[str]]) -> bytes:
    return _docx_table(
        [
            "CAPA编号",
            "事件部门",
            "涉及产品",
            "来源编号",
            "CAPA简述",
            "CAPA效果评估",
            "关闭日期",
            "QA质量员/日期",
        ],
        rows,
    )


def _deviation_docx(rows: list[list[str]]) -> bytes:
    return _docx_table(
        [
            "偏差编号",
            "产品名称/批号",
            "偏差简要描述",
            "偏差是否曾发生",
            "根本原因",
            "偏差等级",
            "调查完成时间",
            "纠正预防措施",
            "产品/物料处理结果",
            "事件部门",
        ],
        rows,
    )


def test_import_parsers_cover_supported_and_invalid_values() -> None:
    now = datetime.now(UTC)
    assert service._parse_date(None) is None
    assert service._parse_date(now) is now
    assert service._parse_date("2026.08.20").date().isoformat() == "2026-08-20"  # type: ignore[union-attr]
    assert service._parse_date("20260820").date().isoformat() == "2026-08-20"  # type: ignore[union-attr]
    assert service._parse_date("2026-08-20 12:00:00").hour == 12  # type: ignore[union-attr]
    assert service._parse_date("bad") is None
    assert service._parse_bool(None) is None
    assert service._parse_bool("yes") is True
    assert service._parse_bool("否") is False
    assert service._parse_bool("unknown") is None
    assert service._parse_occurred_text(None) == (None, None)
    assert service._parse_occurred_text("是 编号：PC-1\nPC-2\n□否") == (
        True,
        "PC-1\nPC-2",
    )
    assert service._parse_occurred_text("是 编号：自定义\n□否") == (True, "自定义")
    assert service._parse_occurred_text("□是 编号：\n否") == (False, None)
    assert service._parse_occurred_text("未知") == (None, None)
    assert service._clean_text(None) == ""
    assert service._clean_text(" x ") == "x"
    assert service._parse_date_value("bad") is None
    assert service._is_empty_row({"a": "", "b": " "}) is True
    assert service._change_row_to_data({"变更控制号": "BG-1"})["change_code"] == "BG-1"


@pytest.mark.anyio
async def test_capa_preview_handles_no_table_valid_missing_and_duplicate(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    db: Any = SimpleNamespace()
    assert (await service.preview_capa_import(db, _empty_docx()))["total_rows"] == 0

    async def exists(_db: Any, code: str) -> bool:
        return code == "C-DUP"

    monkeypatch.setattr(service.repo, "exists_by_capa_code", exists)  # type: ignore[attr-defined]
    result = await service.preview_capa_import(
        db,
        _capa_docx(
            [
                ["C-NEW", "质量部", "A", "S", "简述", "有效", "", "张三2026.08.20"],
                ["C-DUP", "质量部", "A", "S", "重复", "", "", ""],
                ["", "质量部", "A", "S", "缺编号", "", "", ""],
            ]
        ),
    )
    assert result["valid_rows"] == 1
    assert len(result["error_rows"]) == 2


@pytest.mark.anyio
async def test_confirm_capa_import_covers_create_update_skip_duplicate_and_errors(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    db: Any = SimpleNamespace(commit=AsyncMock())
    existing = {
        "C-UP": SimpleNamespace(),
        "C-SKIP": SimpleNamespace(),
        "C-DUP": SimpleNamespace(),
    }

    async def get_by_code(_db: Any, code: str) -> Any:
        return existing.get(code)

    async def create(_db: Any, data: dict[str, Any]) -> Any:
        if data["capa_code"] == "C-FAIL":
            raise RuntimeError("create failed")

    update: Any = AsyncMock()
    monkeypatch.setattr(service.repo, "get_capa_by_code", get_by_code)  # type: ignore[attr-defined]
    monkeypatch.setattr(service.repo, "create_capa", create)  # type: ignore[attr-defined]
    monkeypatch.setattr(service.repo, "update_capa", update)  # type: ignore[attr-defined]
    rows = [
        ["C-NEW", "质量部", "A", "S", "新增", "有效", "2026-08-20", "张三2026.08.20"],
        ["C-FAIL", "质量部", "A", "S", "失败", "", "", "无日期人员"],
        ["", "质量部", "A", "S", "缺编号", "", "", ""],
    ]
    created = await service.confirm_capa_import(db, _capa_docx(rows))
    assert created["success_count"] == 1
    assert created["error_count"] == 2

    updated = await service.confirm_capa_import(
        db,
        _capa_docx([["C-UP", "质量部", "A", "S", "更新", "有效", "", "李四"]]),
        update_existing=True,
    )
    assert updated["update_count"] == 1

    skipped = await service.confirm_capa_import(
        db,
        _capa_docx([["C-SKIP", "", "", "", "", "", "", ""]]),
    )
    assert skipped["skip_count"] == 1
    duplicate = await service.confirm_capa_import(
        db,
        _capa_docx([["C-DUP", "", "", "", "", "", "", ""]]),
        skip_duplicates=False,
    )
    assert duplicate["error_count"] == 1
    assert (await service.confirm_capa_import(db, _empty_docx()))["error_count"] == 1


@pytest.mark.anyio
async def test_export_capa_template_and_data_with_and_without_template(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    db: Any = SimpleNamespace()
    capa: Any = SimpleNamespace(
        capa_code="C-1",
        created_at=datetime(2026, 1, 2),
        department="质量部",
        affected_product="产品A",
        source_code="SRC",
        title="CAPA描述",
        evaluation_result="有效",
        closure_date=datetime(2026, 3, 4),
        qa_confirmer="张三",
        qa_confirm_date=datetime(2026, 3, 5),
    )
    monkeypatch.setattr(service.repo, "get_capas", AsyncMock(return_value=([capa], 1)))  # type: ignore[attr-defined]
    template = service.export_capas_template()
    assert Document(BytesIO(template)).tables
    generated = await service.export_capas(
        db,
        closure_date_from="2026-01-01",
        closure_date_to="2026-12-31",
    )
    assert Document(BytesIO(generated)).tables[0].cell(1, 0).text == "C-1"
    templated = await service.export_capas(db, template_content=template)
    assert Document(BytesIO(templated)).tables[0].cell(1, 8).text == "张三2026.03.05"


@pytest.mark.anyio
async def test_deviation_preview_handles_no_table_valid_missing_and_duplicate(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    db: Any = SimpleNamespace()
    assert (await service.preview_deviation_import(db, _empty_docx()))[
        "total_rows"
    ] == 0

    async def exists(_db: Any, code: str) -> bool:
        return code == "PC-DUP"

    monkeypatch.setattr(service.repo, "exists_by_deviation_code", exists)  # type: ignore[attr-defined]
    result = await service.preview_deviation_import(
        db,
        _deviation_docx(
            [
                [
                    "PC-NEW",
                    "A\nB1",
                    "描述",
                    "□是\n否",
                    "原因",
                    "次要偏差",
                    "",
                    "措施",
                    "放行",
                    "质量部",
                ],
                ["PC-DUP", "", "重复", "", "", "", "", "", "", ""],
                ["", "", "缺编号", "", "", "", "", "", "", ""],
            ]
        ),
    )
    assert result["valid_rows"] == 1
    assert len(result["error_rows"]) == 2


@pytest.mark.anyio
async def test_confirm_deviation_import_covers_all_duplicate_paths(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    db: Any = SimpleNamespace(commit=AsyncMock())
    existing = {
        "PC-DELETED": SimpleNamespace(is_deleted=True, deleted_by="x", deleted_at="x"),
        "PC-UP": SimpleNamespace(is_deleted=False),
        "PC-SKIP": SimpleNamespace(is_deleted=False),
        "PC-DUP": SimpleNamespace(is_deleted=False),
    }

    async def get_by_code(_db: Any, code: str) -> Any:
        return existing.get(code)

    async def create(_db: Any, data: dict[str, Any]) -> Any:
        if data["deviation_code"] == "PC-FAIL":
            raise RuntimeError("create failed")

    monkeypatch.setattr(
        service.repo,  # type: ignore[attr-defined]
        "get_deviation_by_code_include_deleted",
        get_by_code,
    )
    monkeypatch.setattr(service.repo, "create_deviation", create)  # type: ignore[attr-defined]
    update: Any = AsyncMock()
    monkeypatch.setattr(service.repo, "update_deviation", update)  # type: ignore[attr-defined]

    base = [
        "产品A\nB-1",
        "描述",
        "是 编号：PC-1\n□否",
        "原因",
        "严重偏差",
        "2026.08.20",
        "措施",
        "报废",
        "质量部",
    ]
    rows = [["PC-NEW", *base], ["PC-DELETED", *base], ["PC-FAIL", *base], ["", *base]]
    result = await service.confirm_deviation_import(db, _deviation_docx(rows))
    assert result["success_count"] == 1
    assert result["update_count"] == 1
    assert result["error_count"] == 2
    assert existing["PC-DELETED"].is_deleted is False

    updated = await service.confirm_deviation_import(
        db, _deviation_docx([["PC-UP", *base]]), update_existing=True
    )
    assert updated["update_count"] == 1
    skipped = await service.confirm_deviation_import(
        db, _deviation_docx([["PC-SKIP", *base]])
    )
    assert skipped["skip_count"] == 1
    duplicate = await service.confirm_deviation_import(
        db,
        _deviation_docx([["PC-DUP", *base]]),
        skip_duplicates=False,
    )
    assert duplicate["error_count"] == 1
    assert (await service.confirm_deviation_import(db, _empty_docx()))[
        "error_count"
    ] == 1


def test_deviation_row_values_cover_occurrence_level_dates_and_defaults() -> None:
    base = dict(
        deviation_code="PC-1",
        affected_items="产品A",
        batch_number="B-1",
        description="描述",
        title="标题",
        root_cause_analysis="原因",
        level="major",
        investigation_completed_at=datetime(2026, 8, 20),
        corrective_actions="措施",
        material_disposition="报废",
        status="closed",
        previous_occurrence_code="PC-OLD",
    )
    occurred = service._deviation_row_values(
        1, SimpleNamespace(**base, has_occurred_before=True), {"major": "严重偏差"}
    )
    assert occurred[2] == "产品A\nB-1"
    assert "PC-OLD" in occurred[4]
    assert occurred[-1] == "是"

    base.update(
        affected_items="—",
        batch_number="—",
        investigation_completed_at=None,
        material_disposition=None,
        status="draft",
    )
    not_occurred = service._deviation_row_values(
        2, SimpleNamespace(**base, has_occurred_before=False), {}
    )
    unknown = service._deviation_row_values(
        3, SimpleNamespace(**base, has_occurred_before=None), {}
    )
    assert not_occurred[2] == "—"
    assert not_occurred[4] == unknown[4]
    assert not_occurred[-1] == "否"


def test_xml_cell_text_helpers_cover_paragraph_growth_and_shrink() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "line1\nline2\nline3"
    cell = table.cell(0, 0)._tc
    service._set_cell_text(cell, "a\nb")
    assert [line for line in table.cell(0, 0).text.splitlines() if line] == ["a", "b"]
    service._set_cell_text(cell, "single")
    assert [line for line in table.cell(0, 0).text.splitlines() if line] == ["single"]

    empty_cell = table.cell(0, 1)._tc
    for paragraph in list(
        empty_cell.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
        )
    ):
        empty_cell.remove(paragraph)
    service._set_cell_text(empty_cell, "created")
    assert table.cell(0, 1).text == "created"


@pytest.mark.anyio
async def test_export_deviations_uses_fallback_and_template(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    db: Any = SimpleNamespace()
    deviation: Any = SimpleNamespace(
        deviation_code="PC-1",
        affected_items="产品A",
        batch_number="B-1",
        description="描述",
        title="标题",
        has_occurred_before=False,
        previous_occurrence_code=None,
        root_cause_analysis="原因",
        level="minor",
        investigation_completed_at=None,
        corrective_actions="措施",
        material_disposition="放行",
        status="draft",
    )
    monkeypatch.setattr(
        service.repo,  # type: ignore[attr-defined]
        "get_deviations",
        AsyncMock(return_value=([deviation], 1)),
    )
    monkeypatch.setattr(service, "TEMPLATE_PATH", tmp_path / "missing.docx")
    fallback = await service.export_deviations(db)
    assert Document(BytesIO(fallback)).tables[0].cell(1, 1).text == "PC-1"

    template = _docx_table([f"H{i}" for i in range(11)], [["template"] * 11])
    templated = await service.export_deviations(db, template_content=template)
    assert Document(BytesIO(templated)).tables[0].cell(1, 1).text == "PC-1"
