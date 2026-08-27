from __future__ import annotations

import io
from typing import Any
from unittest.mock import AsyncMock

import pytest
from docx import Document
from httpx import AsyncClient

from app.modules.quality.api import quality_management as quality_api
from app.modules.quality.service import quality_feishu_pages as service


@pytest.mark.anyio
async def test_list_report_records_adds_record_id_alias(
    db_session: Any,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        service.quality_management_service,  # type: ignore[attr-defined]
        "get_deviation_report_record_list",
        AsyncMock(
            return_value={
                "items": [
                    {
                        "id": "rec_report_001",
                        "deviation_code": "PC-2607001",
                        "feishu_base_record_id": "rec_report_001",
                    }
                ],
                "total": 1,
                "page": 1,
                "page_size": 20,
            }
        ),
    )

    result = await service.list_report_records(db_session, page=1, page_size=20)

    assert result["items"][0]["id"] == "rec_report_001"
    assert result["items"][0]["record_id"] == "rec_report_001"


@pytest.mark.anyio
async def test_create_investigation_push_record_posts_feishu_fields_only(
    db_session: Any,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        service.tracking_service,  # type: ignore[attr-defined]
        "_resolve_selected_submitter_contact",
        AsyncMock(
            return_value={
                "name": "张起智",
                "open_id": "ou_submitter_001",
                "department": "质量部",
                "department_head_name": "部门负责人甲",
            }
        ),
    )
    monkeypatch.setattr(
        service.feishu_sync_service,  # type: ignore[attr-defined]
        "_resolve_contact_bitable_user_value",
        AsyncMock(
            side_effect=[
                [{"id": "ou_submitter_001"}],
                [{"id": "ou_head_001"}],
                None,
                None,
            ]
        ),
    )
    create_mock: Any = AsyncMock(
        return_value={"record_id": "rec_push_001", "table_id": "tbl_push"}
    )
    monkeypatch.setattr(service, "_create_entity_record", create_mock)
    monkeypatch.setattr(
        service,
        "_get_investigation_push_record",
        AsyncMock(
            return_value={
                "id": "rec_push_001",
                "record_id": "rec_push_001",
                "deviation_code": "PC-2607001",
                "push_round": "第1次",
            }
        ),
    )

    result = await service.create_investigation_push_record(
        db_session,
        {
            "deviation_code": "PC-2607001",
            "push_round": "第1次",
            "investigation_report_url": "https://example.com/report.pdf",
            "submitter_open_id": "ou_submitter_001",
        },
    )

    assert result["record_id"] == "rec_push_001"
    create_mock.assert_awaited_once()
    call = create_mock.await_args
    assert call.args[0] is db_session
    assert call.args[1] == "deviation_investigation_push_record"
    assert call.kwargs["search_conditions"] == [
        ("偏差编号", "PC-2607001"),
        ("第N次推送", "第1次"),
    ]
    assert call.args[2]["偏差编号"] == "PC-2607001"
    assert call.args[2]["第N次推送"] == "第1次"
    assert call.args[2]["偏差调查报告"] == {
        "link": "https://example.com/report.pdf",
        "text": "https://example.com/report.pdf",
        "type": "url",
    }
    assert call.args[2]["提交人"] == [{"id": "ou_submitter_001"}]


@pytest.mark.anyio
async def test_create_deviation_ledger_record_generates_code_and_writes_feishu(
    db_session: Any,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        service.quality_management_service,  # type: ignore[attr-defined]
        "_generate_monthly_deviation_code",
        AsyncMock(return_value="PC-2607009"),
    )
    create_mock: Any = AsyncMock(
        return_value={"record_id": "rec_dev_001", "table_id": "tbl_dev"}
    )
    monkeypatch.setattr(service, "_create_entity_record", create_mock)
    monkeypatch.setattr(
        service,
        "get_deviation_ledger_record",
        AsyncMock(
            return_value={
                "id": "rec_dev_001",
                "record_id": "rec_dev_001",
                "deviation_code": "PC-2607009",
            }
        ),
    )

    result = await service.create_deviation_ledger_record(
        db_session,
        {
            "description": "洁净区压差异常",
            "affected_items": "原料A/批号001",
            "level": "major",
            "is_closed": True,
            "close_time": "2026-07-04T10:00:00+00:00",
        },
    )

    assert result["record_id"] == "rec_dev_001"
    create_mock.assert_awaited_once()
    call = create_mock.await_args
    assert call.args[1] == "deviation_ledger"
    assert call.kwargs["search_conditions"] == [("偏差编号", "PC-2607009")]
    assert call.args[2]["偏差编号"] == "PC-2607009"
    assert call.args[2]["偏差简要描述"] == "洁净区压差异常"
    assert call.args[2]["产品名称/批号"] == "原料A/批号001"
    assert call.args[2]["偏差等级"] == "major"
    assert call.args[2]["是否关闭"] == "是"


@pytest.mark.anyio
async def test_deviation_ledger_round_trip_reads_back_latest_feishu_fields(
    db_session: Any,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        service.quality_management_service,  # type: ignore[attr-defined]
        "_generate_monthly_deviation_code",
        AsyncMock(return_value="PC-2607010"),
    )
    create_mock: Any = AsyncMock(
        return_value={"record_id": "rec_dev_001", "table_id": "tbl_dev"}
    )
    get_mock: Any = AsyncMock(
        side_effect=[
            {
                "id": "rec_dev_001",
                "record_id": "rec_dev_001",
                "deviation_code": "PC-2607010",
                "description": "初始偏差描述",
            },
            {
                "id": "rec_dev_001",
                "record_id": "rec_dev_001",
                "deviation_code": "PC-2607010",
                "description": "更新后的偏差描述",
            },
        ]
    )
    monkeypatch.setattr(service, "_create_entity_record", create_mock)
    monkeypatch.setattr(service, "get_deviation_ledger_record", get_mock)

    created = await service.create_deviation_ledger_record(
        db_session,
        {
            "description": "初始偏差描述",
        },
    )
    loaded = await service.get_deviation_ledger_record(db_session, created["record_id"])

    assert created["description"] == "初始偏差描述"
    assert loaded["description"] == "更新后的偏差描述"
    assert loaded["record_id"] == "rec_dev_001"
    assert get_mock.await_count == 2


@pytest.mark.anyio
async def test_list_deviation_ledger_records_filters_by_record_ids(
    db_session: Any,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        service,
        "_resolve_runtime_entity",
        AsyncMock(return_value=(None, object())),
    )
    monkeypatch.setattr(
        service,
        "_search_entity_records",
        AsyncMock(
            return_value=[
                {"record_id": "rec_dev_001"},
                {"record_id": "rec_dev_002"},
            ]
        ),
    )
    monkeypatch.setattr(
        service,
        "_map_deviation_ledger_base_item",
        lambda record, _entity: {
            "id": record["record_id"],
            "record_id": record["record_id"],
            "deviation_code": record["record_id"].replace("rec_dev_", "PC-2607"),
            "title": "",
            "description": "",
            "status": "draft",
            "created_at": None,
            "feishu_source_updated_at": None,
        },
    )

    result = await service.list_deviation_ledger_records(
        db_session,
        record_ids=["rec_dev_002"],
        page=1,
        page_size=20,
    )

    assert [item["record_id"] for item in result["items"]] == ["rec_dev_002"]
    assert result["total"] == 1


@pytest.mark.anyio
async def test_report_records_api_returns_feishu_page_service_payload(
    client: AsyncClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    list_mock: Any = AsyncMock(
        return_value={
            "items": [
                {
                    "id": "rec_report_001",
                    "record_id": "rec_report_001",
                    "deviation_code": "PC-2607001",
                }
            ],
            "total": 1,
            "page": 1,
            "page_size": 20,
        }
    )
    monkeypatch.setattr(
        quality_api.quality_feishu_pages,  # type: ignore[attr-defined]
        "list_report_records",
        list_mock,
    )

    response = await client.get(
        "/api/v1/quality/deviation-report-records?page=1&page_size=20"
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["data"][0]["record_id"] == "rec_report_001"
    assert payload["meta"]["total"] == 1
    list_mock.assert_awaited_once()


@pytest.mark.anyio
async def test_investigation_push_create_api_posts_to_feishu_service(
    client: AsyncClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    create_mock: Any = AsyncMock(
        return_value={"record_id": "rec_push_001", "id": "rec_push_001"}
    )
    monkeypatch.setattr(
        quality_api.quality_feishu_pages,  # type: ignore[attr-defined]
        "create_investigation_push_record",
        create_mock,
    )

    response = await client.post(
        "/api/v1/quality/deviation-investigation-push-records",
        json={"deviation_code": "PC-2607001", "push_round": "第1次"},
    )

    assert response.status_code == 200
    assert response.json()["data"]["record_id"] == "rec_push_001"
    create_mock.assert_awaited_once()


@pytest.mark.anyio
async def test_deviation_ledger_api_roundtrip_delegates_to_feishu_page_service(
    client: AsyncClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    list_mock: Any = AsyncMock(
        return_value={
            "items": [
                {
                    "id": "rec_dev_001",
                    "record_id": "rec_dev_001",
                    "deviation_code": "PC-2607001",
                }
            ],
            "total": 1,
            "page": 1,
            "page_size": 20,
        }
    )
    get_mock: Any = AsyncMock(
        return_value={
            "id": "rec_dev_001",
            "record_id": "rec_dev_001",
            "deviation_code": "PC-2607001",
        }
    )
    create_mock: Any = AsyncMock(
        return_value={
            "id": "rec_dev_001",
            "record_id": "rec_dev_001",
            "deviation_code": "PC-2607001",
        }
    )
    update_mock: Any = AsyncMock(
        return_value={
            "id": "rec_dev_001",
            "record_id": "rec_dev_001",
            "deviation_code": "PC-2607001",
        }
    )
    delete_mock: Any = AsyncMock(return_value=None)
    monkeypatch.setattr(
        quality_api.quality_feishu_pages,  # type: ignore[attr-defined]
        "list_deviation_ledger_records",
        list_mock,
    )
    monkeypatch.setattr(
        quality_api.quality_feishu_pages,  # type: ignore[attr-defined]
        "get_deviation_ledger_record",
        get_mock,
    )
    monkeypatch.setattr(
        quality_api.quality_feishu_pages,  # type: ignore[attr-defined]
        "create_deviation_ledger_record",
        create_mock,
    )
    monkeypatch.setattr(
        quality_api.quality_feishu_pages,  # type: ignore[attr-defined]
        "update_deviation_ledger_record",
        update_mock,
    )
    monkeypatch.setattr(
        quality_api.quality_feishu_pages,  # type: ignore[attr-defined]
        "delete_deviation_ledger_record",
        delete_mock,
    )

    list_response = await client.get(
        "/api/v1/quality/deviation-ledger-records?page=1&page_size=20"
    )
    assert list_response.status_code == 200
    assert list_response.json()["data"][0]["record_id"] == "rec_dev_001"

    detail_response = await client.get(
        "/api/v1/quality/deviation-ledger-records/rec_dev_001"
    )
    assert detail_response.status_code == 200
    assert detail_response.json()["data"]["record_id"] == "rec_dev_001"

    create_response = await client.post(
        "/api/v1/quality/deviation-ledger-records",
        json={"description": "洁净区压差异常"},
    )
    assert create_response.status_code == 200
    assert create_response.json()["data"]["record_id"] == "rec_dev_001"

    update_response = await client.put(
        "/api/v1/quality/deviation-ledger-records/rec_dev_001",
        json={"description": "已更新"},
    )
    assert update_response.status_code == 200
    assert update_response.json()["data"]["record_id"] == "rec_dev_001"

    delete_response = await client.delete(
        "/api/v1/quality/deviation-ledger-records/rec_dev_001"
    )
    assert delete_response.status_code == 200
    assert delete_response.json()["data"]["success"] is True
    delete_mock.assert_awaited_once()


@pytest.mark.anyio
async def test_deviation_ledger_export_api_streams_docx(
    client: AsyncClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    list_mock: Any = AsyncMock(
        return_value={
            "items": [
                {
                    "record_id": "rec_dev_001",
                    "deviation_code": "PC-2607001",
                    "affected_items": "原料A",
                    "batch_number": "BATCH-001",
                    "description": "洁净区压差异常",
                    "has_occurred_before": True,
                    "root_cause_analysis": "空调机组波动",
                    "level": "major",
                    "investigation_completed_at": "2026-07-04T10:00:00+00:00",
                    "corrective_actions": "复核空调系统参数",
                    "material_disposition": "隔离待评估",
                    "status": "closed",
                }
            ],
            "total": 1,
            "page": 1,
            "page_size": 10000,
        }
    )
    monkeypatch.setattr(
        quality_api.quality_feishu_pages,  # type: ignore[attr-defined]
        "list_deviation_ledger_records",
        list_mock,
    )

    response = await client.get(
        "/api/v1/quality/deviation-ledger-records/export?deviation_code=PC-2607001"
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "deviation-ledger.docx" in response.headers["content-disposition"]

    doc = Document(io.BytesIO(response.content))
    table = doc.tables[0]
    assert table.cell(1, 1).text == "PC-2607001"
    assert table.cell(1, 3).text == "洁净区压差异常"
    assert table.cell(1, 10).text == "是"

    list_mock.assert_awaited_once()


@pytest.mark.anyio
async def test_deviation_ledger_export_api_supports_selected_record_ids(
    client: AsyncClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    list_mock: Any = AsyncMock(
        return_value={
            "items": [
                {
                    "record_id": "rec_dev_002",
                    "deviation_code": "PC-2607002",
                    "description": "称量记录缺页",
                    "status": "draft",
                }
            ],
            "total": 1,
            "page": 1,
            "page_size": 10000,
        }
    )
    monkeypatch.setattr(
        quality_api.quality_feishu_pages,  # type: ignore[attr-defined]
        "list_deviation_ledger_records",
        list_mock,
    )

    response = await client.get(
        "/api/v1/quality/deviation-ledger-records/export?record_ids=rec_dev_002"
    )

    assert response.status_code == 200
    list_mock.assert_awaited_once()
    call = list_mock.await_args
    assert call.kwargs["record_ids"] == ["rec_dev_002"]


@pytest.mark.anyio
async def test_deviation_ledger_single_export_api_streams_single_docx(
    client: AsyncClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    get_mock: Any = AsyncMock(
        return_value={
            "record_id": "rec_dev_001",
            "deviation_code": "PC-2607001",
            "affected_items": "原料A",
            "batch_number": "BATCH-001",
            "description": "洁净区压差异常",
            "has_occurred_before": True,
            "root_cause_analysis": "空调机组波动",
            "level": "major",
            "investigation_completed_at": "2026-07-04T10:00:00+00:00",
            "corrective_actions": "复核空调系统参数",
            "material_disposition": "隔离待评估",
            "status": "closed",
        }
    )
    monkeypatch.setattr(
        quality_api.quality_feishu_pages,  # type: ignore[attr-defined]
        "get_deviation_ledger_record",
        get_mock,
    )

    response = await client.get(
        "/api/v1/quality/deviation-ledger-records/rec_dev_001/export"
    )

    assert response.status_code == 200
    assert "PC-2607001.docx" in response.headers["content-disposition"]
    doc = Document(io.BytesIO(response.content))
    table = doc.tables[0]
    assert len(table.rows) == 2
    assert table.cell(1, 1).text == "PC-2607001"
    get_mock.assert_awaited_once()
