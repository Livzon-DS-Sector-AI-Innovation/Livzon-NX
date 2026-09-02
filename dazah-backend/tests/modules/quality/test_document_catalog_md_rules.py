from __future__ import annotations

import base64
import builtins
from io import BytesIO
from typing import Any

import pytest
from docx import Document

from app.core.exceptions import AppException
from app.modules.quality.service import document_catalog_docx_md as docx_md
from app.modules.quality.service import document_catalog_md as md

# 1x1 透明 PNG，用于内存构造含图片的 docx
PNG_1PX: bytes = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


def _is_section_number(text: str) -> bool:
    return docx_md._is_section_number(text)


def _build_template_docx() -> bytes:
    """构造公司 GMP 模板 docx：首页审批格 + 正文段落 + 主内容表（含数据表、
    嵌套表格与内嵌图片）。"""
    document = Document()

    # 首页审批格（应整体跳过，仅生效日期进入 md 头部）
    header_table = document.add_table(rows=2, cols=2)
    header_table.cell(0, 0).text = "生效日期"
    header_table.cell(0, 1).text = "2026-08-01"
    header_table.cell(1, 0).text = "分发-1"
    header_table.cell(1, 1).text = "QA部门"

    # 正文段落（有序号，转标题）
    document.add_paragraph("1. 目的：建立本文件")

    # 主内容表：序号列 + 内容列 + 两个空列；第 5/6 行构成修订数据表，
    # 末行附嵌套表格与内嵌图片
    main_table = document.add_table(rows=7, cols=4)
    rows_content = [
        ("1", "目的", "", ""),
        ("2", "职责", "", ""),
        ("2.1", "车间主任负责执行", "", ""),
        ("3", "程序", "", ""),
        ("", "版本", "修订内容", "修订内容"),
        ("", "01", "首次颁发", "首次颁发"),
        ("4", "附录", "", ""),
    ]
    for r, cells in enumerate(rows_content):
        for c, text in enumerate(cells):
            main_table.cell(r, c).text = text

    nested = main_table.cell(6, 1).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "嵌套内容"
    picture_run = main_table.cell(6, 2).paragraphs[0].add_run()
    picture_run.add_picture(BytesIO(PNG_1PX))

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_heading_level_and_section_number_helpers() -> None:
    assert _is_section_number("1") is True
    assert _is_section_number("1.2.3") is True
    assert _is_section_number("1.2.") is True
    assert _is_section_number("1、") is True
    assert _is_section_number("1a") is False
    assert _is_section_number("") is False
    assert docx_md._heading_level("1") == 2
    assert docx_md._heading_level("1.1") == 3
    assert docx_md._heading_level("2.3.4") == 4
    assert docx_md._heading_level("2.3.4.5") == 0


def test_dedupe_table_columns_keeps_first_unique_column() -> None:
    rows = [["a", "a", "b"], ["1", "1", "2"]]
    assert docx_md._dedupe_table_columns(rows) == [["a", "b"], ["1", "2"]]
    assert docx_md._dedupe_table_columns([]) == []


def test_convert_docx_content_to_md_full_template() -> None:
    content = _build_template_docx()
    markdown, images = docx_md.convert_docx_content_to_md(
        content, "SMP-QA-001-02偏差处理程序.docx"
    )

    # md 头部：文件名提取编号/标题，首页表格提取生效日期
    assert markdown.startswith("# 偏差处理程序")
    assert "**文件编号**: SMP-QA-001-02" in markdown
    assert "**生效日期**: 2026-08-01" in markdown

    # 首页审批格整体跳过
    assert "分发-1" not in markdown
    assert "QA部门" not in markdown

    # 正文段落与主内容表序号转标题层级
    assert "## 1 目的" in markdown
    assert "## 1 目的：建立本文件" in markdown
    assert "### 2.1 车间主任负责执行" in markdown
    assert "## 3 程序" in markdown
    assert "## 4 附录" in markdown

    # 数据表（修订简历）保留为 md 表格，重复列去重
    assert "版本 | 修订内容 |" in markdown
    assert "01 | 首次颁发 |" in markdown

    # 嵌套表格保留为 md 表格
    assert "| 嵌套内容 |" in markdown

    # 图片原位引用且被提取
    assert "![image](img_000.png)" in markdown
    assert len(images) == 1
    assert images[0].name == "img_000.png"
    assert images[0].data == PNG_1PX
    assert images[0].content_type == "image/png"


def test_convert_docx_content_to_md_rejects_corrupt_file() -> None:
    with pytest.raises(AppException, match="解析失败"):
        docx_md.convert_docx_content_to_md(b"not-a-docx", "bad.docx")


def test_find_soffice_prefers_config_then_path(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(md, "DOC_CONVERTER_BIN", "configured")
    monkeypatch.setattr(md.os.path, "exists", lambda path: path == "configured")  # type: ignore[attr-defined]
    assert md._find_soffice() == "configured"

    monkeypatch.setattr(md, "DOC_CONVERTER_BIN", "")
    monkeypatch.setattr(
        md.shutil,  # type: ignore[attr-defined]
        "which",
        lambda name: "office-bin" if name == "libreoffice" else None,
    )
    assert md._find_soffice() == "office-bin"
    monkeypatch.setattr(md.shutil, "which", lambda _name: None)  # type: ignore[attr-defined]
    assert md._find_soffice() is None


def test_soffice_conversion_handles_missing_binary_and_process_error(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(md, "_find_soffice", lambda: None)
    assert md._convert_doc_via_soffice(b"doc", "test.doc") == b""

    monkeypatch.setattr(md, "_find_soffice", lambda: "office-bin")
    monkeypatch.setattr(
        md.subprocess,  # type: ignore[attr-defined]
        "run",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(OSError("failed")),
    )
    assert md._convert_doc_via_soffice(b"doc", "test.doc") == b""


def test_convert_legacy_to_docx_uses_soffice_or_reports_unsupported(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(md, "_convert_doc_via_soffice", lambda *_args: b"docx")
    assert md.convert_legacy_to_docx(b"doc", "test.doc") == b"docx"
    assert md.convert_legacy_to_docx(b"wps", "test.wps") == b"docx"

    monkeypatch.setattr(md, "_convert_doc_via_soffice", lambda *_args: b"")
    original_import = builtins.__import__

    def import_without_win32(name: str, *args: Any, **kwargs: Any) -> Any:
        if name == "win32com.client":
            raise ImportError(name)
        return original_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", import_without_win32)
    with pytest.raises(AppException, match="不支持该格式转换"):
        md.convert_legacy_to_docx(b"doc", "test.doc")
    with pytest.raises(AppException, match="不支持该格式转换"):
        md.convert_legacy_to_docx(b"wps", "test.wps")


def test_convert_word_attachment_runs_legacy_conversion_for_doc_and_wps(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    legacy_calls: list[str] = []

    def fake_legacy(content: bytes, file_name: str) -> bytes:
        legacy_calls.append(file_name)
        return content

    monkeypatch.setattr(md, "convert_legacy_to_docx", fake_legacy)
    monkeypatch.setattr(
        md,
        "convert_docx_content_to_md",
        lambda content, name: (f"# {name}", []),
    )

    docx_result = md.convert_word_attachment("标准.docx", b"docx")
    assert docx_result == ("# 标准.docx", [])
    assert legacy_calls == []

    assert md.convert_word_attachment("标准.doc", b"doc")[0] == "# 标准.doc"
    assert md.convert_word_attachment("标准.WPS", b"wps")[0] == "# 标准.WPS"
    assert legacy_calls == ["标准.doc", "标准.WPS"]
