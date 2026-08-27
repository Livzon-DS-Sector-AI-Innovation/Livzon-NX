from __future__ import annotations

import io
from datetime import UTC, datetime

from docx import Document

from app.modules.quality.service import deviation_ledger_export


def test_generate_deviation_ledger_export_docx_uses_template_layout() -> None:
    docx_bytes = deviation_ledger_export.generate_deviation_ledger_export_docx(
        [
            {
                "deviation_code": "PC-2607001",
                "affected_items": "原料A",
                "batch_number": "BATCH-001",
                "description": "洁净区压差异常",
                "has_occurred_before": True,
                "root_cause_analysis": "空调机组波动",
                "level": "major",
                "investigation_completed_at": datetime(2026, 7, 4, 10, 30, tzinfo=UTC),
                "corrective_actions": "复核空调系统参数",
                "material_disposition": "隔离待评估",
                "status": "closed",
            },
            {
                "deviation_code": "PC-2607002",
                "product_batch": "原料B\nBATCH-002",
                "description": "称量记录缺页",
                "has_occurred_before": False,
                "root_cause_analysis": "记录回收不完整",
                "level": "微小",
                "investigation_completed_at": "2026-07-05T08:00:00+00:00",
                "corrective_actions": "补充培训",
                "material_disposition": "已补录",
                "is_closed": False,
            },
        ]
    )

    doc = Document(io.BytesIO(docx_bytes))
    assert len(doc.tables) == 1

    table = doc.tables[0]
    assert len(table.rows) == 3
    assert table.cell(1, 0).text == "1"
    assert table.cell(1, 1).text == "PC-2607001"
    assert table.cell(1, 2).text == "原料A\nBATCH-001"
    assert table.cell(1, 6).text == "重大"
    assert table.cell(1, 7).text == "2026.07.04"
    assert table.cell(1, 10).text == "是"

    assert table.cell(2, 0).text == "2"
    assert table.cell(2, 1).text == "PC-2607002"
    assert table.cell(2, 2).text == "原料B\nBATCH-002"
    assert table.cell(2, 4).text == "否"
    assert table.cell(2, 6).text == "微小"
    assert table.cell(2, 7).text == "2026.07.05"
    assert table.cell(2, 10).text == "否"
