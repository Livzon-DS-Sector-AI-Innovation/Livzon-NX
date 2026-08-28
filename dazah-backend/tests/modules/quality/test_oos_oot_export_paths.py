from __future__ import annotations

import io
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock

import pytest
from docx import Document

from app.modules.quality.service import oos_oot_export


def _write_template(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=11)
    for index, (header, _field) in enumerate(oos_oot_export.LEDGER_COLUMNS):
        table.cell(0, index).text = header
        table.cell(1, index).text = "模板"
        table.cell(1, index).paragraphs[0].runs[0].font.name = "Arial"
    document.save(path)


def test_fill_template_maps_dates_formats_fields_and_limits_rows(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    _write_template(tmp_path / oos_oot_export.OOS_TEMPLATE)
    monkeypatch.setattr(oos_oot_export, "_TEMPLATE_DIR", tmp_path)

    output = oos_oot_export._fill_template(
        oos_oot_export.OOS_TEMPLATE,
        [
            {
                "date": "2026-08-26T12:30:00Z",
                "material_name": "物料A",
                "batch_number": "B-001",
                "investigation_code": "OOS-001",
                "problem_description": "描述",
                "root_cause": "原因",
                "corrective_actions": "措施",
                "final_disposition": "放行",
                "registrant": "张三",
                "remark": "备注",
            },
            {"date": "not-a-date", "material_name": "物料B"},
        ],
        max_rows=1,
    )

    assert isinstance(output, io.BytesIO)
    document = Document(output)
    table = document.tables[0]
    assert len(table.rows) == 2
    assert [cell.text for cell in table.rows[1].cells] == [
        "1",
        "2026/08/26",
        "物料A",
        "B-001",
        "OOS-001",
        "描述",
        "原因",
        "措施",
        "放行",
        "张三",
        "备注",
    ]


def test_fill_template_keeps_one_blank_row_and_fallback_date(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    _write_template(tmp_path / oos_oot_export.OOT_TEMPLATE)
    monkeypatch.setattr(oos_oot_export, "_TEMPLATE_DIR", tmp_path)

    output = oos_oot_export._fill_template(oos_oot_export.OOT_TEMPLATE, [])
    row = Document(output).tables[0].rows[1]
    assert row.cells[0].text == "1"
    assert row.cells[1].text == ""
    assert oos_oot_export._format_date("2026/08/26") == "2026/08/26"
    assert oos_oot_export._format_date(None) == ""


@pytest.mark.asyncio
async def test_export_wrappers_use_records_and_fallback_to_empty_template(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    _write_template(tmp_path / oos_oot_export.OOS_TEMPLATE)
    _write_template(tmp_path / oos_oot_export.OOT_TEMPLATE)
    monkeypatch.setattr(oos_oot_export, "_TEMPLATE_DIR", tmp_path)
    monkeypatch.setattr(
        oos_oot_export,
        "list_oos_ledger_records",
        AsyncMock(return_value={"items": [{"material_name": "OOS物料"}]}),
    )
    monkeypatch.setattr(
        oos_oot_export,
        "list_oot_ledger_records",
        AsyncMock(side_effect=RuntimeError("source unavailable")),
    )

    oos_output = await oos_oot_export.export_oos_ledger(SimpleNamespace())
    oot_output = await oos_oot_export.export_oot_ledger(SimpleNamespace())

    assert Document(oos_output).tables[0].rows[1].cells[2].text == "OOS物料"
    assert Document(oot_output).tables[0].rows[1].cells[0].text == "1"
