"""偏差工作台 API 接口测试（AsyncClient 真调路由 + mock llm_client）。"""

from __future__ import annotations

import io
from collections.abc import AsyncIterator
from types import SimpleNamespace as _SimpleNamespace
from typing import Any

import pytest
from docx import Document
from httpx import AsyncClient
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.llm import LLMConfigError, LLMOutputError

SimpleNamespace: Any = _SimpleNamespace


@pytest.fixture(autouse=True)
def _disable_minio(monkeypatch: pytest.MonkeyPatch) -> None:
    """本地测试环境无 MinIO 容器，降级为本地存储路径（外部依赖按规范 mock）。"""
    monkeypatch.setattr(
        "app.modules.quality.service.quality_attachment.minio_enabled",
        lambda: False,
    )


@pytest.fixture(autouse=True)
async def _clean_workbench_tables(db_session: AsyncSession) -> AsyncIterator[Any]:
    await db_session.execute(text("CREATE SCHEMA IF NOT EXISTS quality"))
    await db_session.execute(
        text(
            """
            CREATE TABLE IF NOT EXISTS quality.deviation_workbench_settings (
                id UUID PRIMARY KEY,
                report_system_prompt TEXT NOT NULL DEFAULT '',
                created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
                updated_at TIMESTAMPTZ NOT NULL DEFAULT now(),
                created_by UUID NULL,
                updated_by UUID NULL,
                is_deleted BOOLEAN NOT NULL DEFAULT FALSE
            )
            """
        )
    )
    await db_session.execute(
        text(
            """
            CREATE TABLE IF NOT EXISTS quality.deviation_workbench_reports (
                id UUID PRIMARY KEY,
                code VARCHAR(255) NOT NULL,
                source_type VARCHAR(50) NOT NULL DEFAULT 'manual',
                source_record_id VARCHAR(100) NULL,
                deviation_summary TEXT NULL,
                manual_text TEXT NULL,
                attachments JSON NULL,
                context_snapshot JSON NULL,
                report_payload JSON NULL,
                report_md TEXT NULL,
                model_name VARCHAR(255) NULL,
                status VARCHAR(50) NOT NULL DEFAULT 'processing',
                error_message TEXT NULL,
                deleted_by UUID NULL,
                deleted_at TIMESTAMPTZ NULL,
                created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
                updated_at TIMESTAMPTZ NOT NULL DEFAULT now(),
                created_by UUID NULL,
                updated_by UUID NULL,
                is_deleted BOOLEAN NOT NULL DEFAULT FALSE
            )
            """
        )
    )
    await db_session.execute(
        text("DELETE FROM quality.deviation_workbench_reports")
    )
    await db_session.execute(
        text("DELETE FROM quality.deviation_workbench_settings")
    )
    await db_session.commit()
    yield
    await db_session.execute(
        text("DELETE FROM quality.deviation_workbench_reports")
    )
    await db_session.execute(
        text("DELETE FROM quality.deviation_workbench_settings")
    )
    await db_session.commit()


def _async_config_stub() -> SimpleNamespace:
    return SimpleNamespace(model_name="test-model")


def _valid_report_payload() -> dict[str, Any]:
    return {
        "deviation_summary": "测试偏差：灌装压塞压力超上限",
        "analysis": {
            "人": "操作人员未复核压塞压力",
            "机": "压塞机压力传感器漂移",
            "料": "胶塞来料批次差异",
            "法": "SOP 未规定压塞压力复核频次",
            "环": "洁净区压差波动",
            "测": "在线压力监测未设趋势报警",
        },
        "direct_cause": "压塞压力传感器漂移未及时校准",
        "root_cause": "预防性维护未覆盖传感器校准且缺少趋势报警",
        "conclusion": "偏差属实，直接原因与根本原因已定位",
        "recommendations": ["建立传感器校准台账", "增加压塞压力趋势报警"],
        "referenced_sources": ["模型通用知识"],
    }


@pytest.fixture
async def _fake_llm_success(monkeypatch: pytest.MonkeyPatch) -> None:
    async def _fake_chat_json(
        self: Any,  # noqa: ANN001
        messages: list[dict[Any, Any]],
        expected_keys: Any = None,  # noqa: ANN001
        temperature: Any = None,  # noqa: ANN001
        config_type: Any = "text",  # noqa: ANN001
    ) -> dict[str, Any]:
        return _valid_report_payload()

    async def _async_config(*_args: Any, **_kwargs: Any) -> SimpleNamespace:
        return _async_config_stub()

    async def _async_list_entries(*_args: Any, **_kwargs: Any) -> tuple[Any, int]:
        return ([], 0)

    monkeypatch.setattr(
        "app.core.llm.client.LLMClient.chat_json", _fake_chat_json
    )
    monkeypatch.setattr(
        "app.modules.quality.service.deviation_workbench.get_config",
        _async_config,
    )
    # 隔离文件管理检索，避免依赖 document_catalog 表结构
    monkeypatch.setattr(
        "app.modules.quality.service.deviation_workbench.list_document_entries",
        _async_list_entries,
    )


@pytest.mark.anyio
async def test_get_and_update_settings(client: AsyncClient) -> None:
    get_resp = await client.get("/api/v1/quality/deviation-workbench/settings")
    assert get_resp.status_code == 200
    assert "report_system_prompt" in get_resp.json()["data"]

    update_resp = await client.put(
        "/api/v1/quality/deviation-workbench/settings",
        json={"report_system_prompt": "你是质量调查专家，从 5M1E 分析。"},
    )
    assert update_resp.status_code == 200
    assert (
        update_resp.json()["data"]["report_system_prompt"]
        == "你是质量调查专家，从 5M1E 分析。"
    )


@pytest.mark.anyio
async def test_analyze_manual_success(
    client: AsyncClient, db_session: AsyncSession, _fake_llm_success: None
) -> None:
    response = await client.post(
        "/api/v1/quality/deviation-workbench/analyze",
        json={
            "source_type": "manual",
            "manual_text": "灌装线压塞压力超上限，怀疑传感器漂移。",
            "attachments": [],
        },
    )
    assert response.status_code == 200
    data = response.json()["data"]
    assert data["code"].startswith("WB-")
    assert data["status"] == "completed"
    assert data["report_payload"]["root_cause"]
    assert "## 二、人机料法环测" in data["report_md"]
    assert data["context_snapshot"] is not None


@pytest.mark.anyio
async def test_analyze_retrieves_training_ledgers(
    client: AsyncClient, db_session: AsyncSession, monkeypatch: pytest.MonkeyPatch
) -> None:
    """偏差工作台分析时应检索培训台账并注入上下文快照。"""
    captured_prompt: dict[str, str] = {}

    async def _fake_chat_json(
        self: Any,  # noqa: ANN001
        messages: list[dict[Any, Any]],
        expected_keys: Any = None,  # noqa: ANN001
        temperature: Any = None,  # noqa: ANN001
        config_type: Any = "text",  # noqa: ANN001
    ) -> dict[str, Any]:
        captured_prompt["prompt"] = messages[-1]["content"] if messages else ""
        return _valid_report_payload()

    async def _async_config(*_args: Any, **_kwargs: Any) -> SimpleNamespace:
        return _async_config_stub()

    async def _async_list_entries(*_args: Any, **_kwargs: Any) -> tuple[Any, int]:
        return ([], 0)

    async def _fake_training(
        _db: Any, keywords: list[str], *, limit: int = 5
    ) -> list[dict[str, Any]]:
        captured_prompt["training_keywords"] = ",".join(keywords)
        return [
            {
                "training_date": "2026-08-01",
                "training_subject": "压塞机操作与压力复核培训",
                "training_content": "压塞压力标准 0.40MPa，超限须停机上报",
                "training_method": "现场",
                "duration_hours": 2,
                "trainer": "质量部",
                "instructor": "张三",
                "teaching_dept": "质量部",
                "trainees": "灌装线操作工",
                "training_type": "质量类",
                "assessment_result": "合格",
            }
        ]

    monkeypatch.setattr("app.core.llm.client.LLMClient.chat_json", _fake_chat_json)
    monkeypatch.setattr(
        "app.modules.quality.service.deviation_workbench.get_config", _async_config
    )
    monkeypatch.setattr(
        "app.modules.quality.service.deviation_workbench.list_document_entries",
        _async_list_entries,
    )
    monkeypatch.setattr(
        "app.modules.hr.public_api.query_training_ledgers", _fake_training
    )

    response = await client.post(
        "/api/v1/quality/deviation-workbench/analyze",
        json={
            "source_type": "manual",
            "manual_text": "灌装线压塞压力超上限，怀疑传感器漂移。",
            "attachments": [],
        },
    )
    assert response.status_code == 200
    data = response.json()["data"]
    assert data["status"] == "completed"
    training = (data["context_snapshot"] or {}).get("training_ledgers") or []
    assert len(training) == 1
    assert training[0]["training_subject"] == "压塞机操作与压力复核培训"
    assert "压塞机操作与压力复核培训" in captured_prompt["prompt"]


@pytest.mark.anyio
async def test_analyze_no_input_rejected(client: AsyncClient) -> None:
    response = await client.post(
        "/api/v1/quality/deviation-workbench/analyze",
        json={"source_type": "manual", "manual_text": "  ", "attachments": []},
    )
    assert response.status_code == 400


@pytest.mark.anyio
async def test_analyze_llm_not_configured(
    client: AsyncClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    async def _raise_config(*_args: Any, **_kwargs: Any) -> Any:
        raise LLMConfigError("未配置")

    async def _async_list_entries(*_args: Any, **_kwargs: Any) -> tuple[Any, int]:
        return ([], 0)

    monkeypatch.setattr(
        "app.modules.quality.service.deviation_workbench.get_config",
        _raise_config,
    )
    monkeypatch.setattr(
        "app.modules.quality.service.deviation_workbench.list_document_entries",
        _async_list_entries,
    )

    response = await client.post(
        "/api/v1/quality/deviation-workbench/analyze",
        json={"source_type": "manual", "manual_text": "灌装压塞压力超上限。"},
    )
    assert response.status_code == 200
    data = response.json()["data"]
    assert data["status"] == "failed"
    assert "AI 服务尚未配置" in data["error_message"]


@pytest.mark.anyio
async def test_analyze_llm_output_error(
    client: AsyncClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    async def _raise_output(
        self: Any,  # noqa: ANN001
        *args: Any,  # noqa: ANN002
        **kwargs: Any,  # noqa: ANN003
    ) -> dict[str, Any]:
        raise LLMOutputError("输出格式错误")

    monkeypatch.setattr(
        "app.core.llm.client.LLMClient.chat_json", _raise_output
    )

    async def _async_config(*_args: Any, **_kwargs: Any) -> SimpleNamespace:
        return _async_config_stub()

    async def _async_list_entries(*_args: Any, **_kwargs: Any) -> tuple[Any, int]:
        return ([], 0)

    monkeypatch.setattr(
        "app.modules.quality.service.deviation_workbench.get_config",
        _async_config,
    )
    monkeypatch.setattr(
        "app.modules.quality.service.deviation_workbench.list_document_entries",
        _async_list_entries,
    )

    response = await client.post(
        "/api/v1/quality/deviation-workbench/analyze",
        json={"source_type": "manual", "manual_text": "灌装压塞压力超上限。"},
    )
    assert response.status_code == 200
    assert response.json()["data"]["status"] == "failed"


@pytest.mark.anyio
async def test_upload_workbench_attachment(client: AsyncClient) -> None:
    document = Document()
    document.add_paragraph("附件说明：压塞机已完成常规点检。")
    buffer = io.BytesIO()
    document.save(buffer)

    response = await client.post(
        "/api/v1/quality/deviation-workbench/attachments",
        files={
            "file": (
                "note.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200
    descriptor = response.json()["data"]
    assert descriptor["file_name"] == "note.docx"
    assert descriptor["converted"] is True
    assert descriptor["storage_key"]
    assert descriptor["converted_md_key"]


@pytest.mark.anyio
async def test_delete_workbench_attachment(client: AsyncClient) -> None:
    document = Document()
    document.add_paragraph("附件说明：压塞机已完成常规点检。")
    buffer = io.BytesIO()
    document.save(buffer)

    uploaded = await client.post(
        "/api/v1/quality/deviation-workbench/attachments",
        files={
            "file": (
                "cleanup.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert uploaded.status_code == 200
    descriptor = uploaded.json()["data"]
    keys = [
        descriptor["storage_key"],
        descriptor.get("converted_md_key"),
        *(descriptor.get("asset_keys") or []),
    ]

    response = await client.delete(
        "/api/v1/quality/deviation-workbench/attachments",
        params=[("keys", key) for key in keys],
    )
    assert response.status_code == 200
    assert response.json()["data"]["deleted"] == len(keys)


@pytest.mark.anyio
async def test_list_and_delete_reports(
    client: AsyncClient, db_session: AsyncSession, _fake_llm_success: None
) -> None:
    created = await client.post(
        "/api/v1/quality/deviation-workbench/analyze",
        json={"source_type": "manual", "manual_text": "灌装压塞压力超上限。"},
    )
    assert created.status_code == 200
    report_id = created.json()["data"]["id"]

    listed = await client.get(
        "/api/v1/quality/deviation-workbench/reports",
        params={"status": "completed", "keyword": "灌装"},
    )
    assert listed.status_code == 200
    assert listed.json()["meta"]["total"] >= 1

    detail = await client.get(
        f"/api/v1/quality/deviation-workbench/reports/{report_id}"
    )
    assert detail.status_code == 200
    assert detail.json()["data"]["id"] == report_id

    deleted = await client.delete(
        f"/api/v1/quality/deviation-workbench/reports/{report_id}"
    )
    assert deleted.status_code == 200
    gone = await client.get(
        f"/api/v1/quality/deviation-workbench/reports/{report_id}"
    )
    assert gone.status_code == 404
