from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document

from app.modules.quality.service import change_ledger_export


def test_generate_change_ledger_export_docx_uses_template_layout() -> None:
    docx_bytes = change_ledger_export.generate_change_ledger_export_docx(
        [
            {
                "change_code": "BG-2607001",
                "applicant_department": "质量部",
                "change_object": "洁净区空调机组",
                "change_content": "新增压差趋势监测\n补充偏差预警联动",
                "change_level": "中度变更",
                "application_date": datetime(2026, 7, 4, 10, 30, tzinfo=timezone.utc),
                "planned_approval_date": "2026-07-06",
                "execution_date": "2026/07/08",
                "closure_date": "2026.07.10",
            },
            {
                "serial_number": "8",
                "change_code": "BG-2607002",
                "applicant_department": "采购部",
                "change_object": "甘油供应商",
                "change_content": "增加甘油供应商",
                "change_level": "一般变更",
                "application_date": "2026-07-05T08:00:00+00:00",
                "planned_approval_date": None,
                "execution_date": None,
                "status": "正在进行",
            },
        ]
    )

    doc = Document(io.BytesIO(docx_bytes))
    assert len(doc.tables) == 1

    table = doc.tables[0]
    assert len(table.rows) == 3

    header_texts = [cell.text for cell in table.rows[0].cells]
    assert header_texts == [
        "序号",
        "变更控制号",
        "变更申请部门",
        "变更对象",
        "变更内容",
        "变更等级",
        "变更申请日期",
        "变更计划批准日期",
        "变更正式执行日期",
        "变更关闭日期",
    ]

    assert table.cell(1, 0).text == "1"
    assert table.cell(1, 1).text == "BG-2607001"
    assert table.cell(1, 2).text == "质量部"
    assert table.cell(1, 4).text == "新增压差趋势监测\n补充偏差预警联动"
    assert table.cell(1, 6).text == "2026.07.04"
    assert table.cell(1, 7).text == "2026.07.06"
    assert table.cell(1, 8).text == "2026.07.08"
    assert table.cell(1, 9).text == "2026.07.10"

    assert table.cell(2, 0).text == "8"
    assert table.cell(2, 1).text == "BG-2607002"
    assert table.cell(2, 2).text == "采购部"
    assert table.cell(2, 5).text == "一般变更"
    assert table.cell(2, 6).text == "2026.07.05"
    assert table.cell(2, 7).text == ""
    assert table.cell(2, 8).text == ""
    assert table.cell(2, 9).text == "正在进行"
