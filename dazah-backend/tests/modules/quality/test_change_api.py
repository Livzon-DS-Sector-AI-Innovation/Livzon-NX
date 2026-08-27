from __future__ import annotations

import io
import uuid
from datetime import date
from typing import Any

import pytest
from docx import Document
from httpx import AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession

from app.modules.quality.models.change_control import ChangeControl


@pytest.fixture(autouse=True)
async def _clean_change_controls(db_session: AsyncSession) -> Any:
    await db_session.execute(
        ChangeControl.__table__.delete()  # type: ignore[attr-defined]
    )
    await db_session.commit()
    yield
    await db_session.execute(
        ChangeControl.__table__.delete()  # type: ignore[attr-defined]
    )
    await db_session.commit()


@pytest.mark.anyio
async def test_list_changes_api_returns_filtered_rows(
    client: AsyncClient,
    db_session: AsyncSession,
) -> None:
    db_session.add(
        ChangeControl(
            id=uuid.uuid4(),
            serial_number="1",
            change_code="BG-2026-001",
            applicant_department="质量部",
            change_object="反应釜",
            change_content="更换搅拌电机",
            change_level="二级",
            application_date=date(2026, 6, 30),
        )
    )
    await db_session.commit()

    response = await client.get(
        "/api/v1/quality/changes", params={"change_code": "BG-2026-001"}
    )

    assert response.status_code == 200
    body = response.json()
    assert body["meta"]["total"] == 1
    assert body["data"][0]["change_code"] == "BG-2026-001"


@pytest.mark.anyio
async def test_change_crud_api_roundtrip(client: AsyncClient) -> None:
    create_response = await client.post(
        "/api/v1/quality/changes",
        json={
            "serial_number": "2",
            "change_code": "BG-2026-002",
            "applicant_department": "工程部",
            "change_object": "纯化水系统",
            "change_content": "新增监测点位",
            "change_level": "一级",
            "application_date": "2026-07-02",
        },
    )
    assert create_response.status_code == 200
    change_id = create_response.json()["data"]["id"]

    detail_response = await client.get(f"/api/v1/quality/changes/{change_id}")
    assert detail_response.status_code == 200
    assert detail_response.json()["data"]["change_code"] == "BG-2026-002"

    update_response = await client.put(
        f"/api/v1/quality/changes/{change_id}",
        json={"change_content": "新增在线监测点位", "closure_date": "2026-07-15"},
    )
    assert update_response.status_code == 200
    assert update_response.json()["data"]["success"] is True

    delete_response = await client.delete(f"/api/v1/quality/changes/{change_id}")
    assert delete_response.status_code == 200
    assert delete_response.json()["data"]["success"] is True


@pytest.mark.anyio
async def test_batch_delete_changes_api(
    client: AsyncClient,
    db_session: AsyncSession,
) -> None:
    first = ChangeControl(
        id=uuid.uuid4(),
        serial_number="3",
        change_code="BG-2026-003",
        applicant_department="生产部",
    )
    second = ChangeControl(
        id=uuid.uuid4(),
        serial_number="4",
        change_code="BG-2026-004",
        applicant_department="设备部",
    )
    db_session.add_all([first, second])
    await db_session.commit()

    response = await client.post(
        "/api/v1/quality/changes/batch-delete",
        json={"ids": [str(first.id), str(second.id)]},
    )

    assert response.status_code == 200
    assert response.json()["data"]["deleted"] == 2


@pytest.mark.anyio
async def test_export_changes_api_defaults_to_filtered_scope(
    client: AsyncClient,
    db_session: AsyncSession,
) -> None:
    db_session.add_all(
        [
            ChangeControl(
                id=uuid.uuid4(),
                serial_number="1",
                change_code="BG-2026-101",
                applicant_department="质量部",
                change_object="洁净区",
                change_content="更新压差监测规则",
                change_level="一级",
                application_date=date(2026, 7, 1),
            ),
            ChangeControl(
                id=uuid.uuid4(),
                serial_number="2",
                change_code="BG-2026-102",
                applicant_department="生产部",
                change_object="配液罐",
                change_content="更换喷淋球",
                change_level="二级",
                application_date=date(2026, 7, 2),
            ),
        ]
    )
    await db_session.commit()

    response = await client.get(
        "/api/v1/quality/changes/export",
        params={"change_code": "BG-2026-101"},
    )

    assert response.status_code == 200
    assert "change-controls.docx" in response.headers["content-disposition"]
    doc = Document(io.BytesIO(response.content))
    table = doc.tables[0]
    assert len(table.rows) == 2
    assert table.cell(1, 1).text == "BG-2026-101"
    assert table.cell(1, 2).text == "质量部"


@pytest.mark.anyio
async def test_export_changes_api_supports_single_scope(
    client: AsyncClient,
    db_session: AsyncSession,
) -> None:
    change = ChangeControl(
        id=uuid.uuid4(),
        serial_number="9",
        change_code="BG-2026-201",
        applicant_department="设备部",
        change_object="制水系统",
        change_content="新增在线电导监测",
        change_level="一级",
        application_date=date(2026, 7, 3),
    )
    db_session.add(change)
    await db_session.commit()

    response = await client.get(
        "/api/v1/quality/changes/export",
        params={"scope": "single", "change_id": str(change.id)},
    )

    assert response.status_code == 200
    assert "BG-2026-201.docx" in response.headers["content-disposition"]
    doc = Document(io.BytesIO(response.content))
    table = doc.tables[0]
    assert len(table.rows) == 2
    assert table.cell(1, 0).text == "9"
    assert table.cell(1, 1).text == "BG-2026-201"


@pytest.mark.anyio
async def test_export_changes_api_supports_selected_scope(
    client: AsyncClient,
    db_session: AsyncSession,
) -> None:
    first = ChangeControl(
        id=uuid.uuid4(),
        serial_number="3",
        change_code="BG-2026-301",
        applicant_department="工程部",
        change_object="空调箱",
        change_content="更换过滤器",
        change_level="三级",
        application_date=date(2026, 7, 4),
    )
    second = ChangeControl(
        id=uuid.uuid4(),
        serial_number="4",
        change_code="BG-2026-302",
        applicant_department="仓储部",
        change_object="冷库",
        change_content="增加温度巡检频次",
        change_level="二级",
        application_date=date(2026, 7, 5),
    )
    third = ChangeControl(
        id=uuid.uuid4(),
        serial_number="5",
        change_code="BG-2026-303",
        applicant_department="质量部",
        change_object="称量间",
        change_content="调整清场频次",
        change_level="一级",
        application_date=date(2026, 7, 6),
    )
    db_session.add_all([first, second, third])
    await db_session.commit()

    response = await client.get(
        "/api/v1/quality/changes/export",
        params=[
            ("scope", "selected"),
            ("change_ids", str(second.id)),
            ("change_ids", str(first.id)),
        ],
    )

    assert response.status_code == 200
    assert "change-controls.docx" in response.headers["content-disposition"]
    doc = Document(io.BytesIO(response.content))
    table = doc.tables[0]
    assert len(table.rows) == 3
    assert table.cell(1, 1).text == "BG-2026-302"
    assert table.cell(2, 1).text == "BG-2026-301"


@pytest.mark.anyio
async def test_export_changes_api_validates_single_scope_params(
    client: AsyncClient,
) -> None:
    response = await client.get(
        "/api/v1/quality/changes/export",
        params={"scope": "single"},
    )

    assert response.status_code == 400
    assert response.json()["message"] == "scope=single 时必须提供 change_id"


@pytest.mark.anyio
async def test_export_changes_api_validates_selected_scope_params(
    client: AsyncClient,
) -> None:
    response = await client.get(
        "/api/v1/quality/changes/export",
        params={"scope": "selected"},
    )

    assert response.status_code == 400
    assert response.json()["message"] == "scope=selected 时必须提供 change_ids"


@pytest.mark.anyio
async def test_export_changes_api_validates_scope_enum(
    client: AsyncClient,
) -> None:
    response = await client.get(
        "/api/v1/quality/changes/export",
        params={"scope": "all"},
    )

    assert response.status_code == 422
    body = response.json()
    assert body["message"] == "请求参数校验失败"
    assert "scope" in body["detail"]
